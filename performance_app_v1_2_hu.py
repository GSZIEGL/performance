# performance_app_v1_2_hu.py
# AI-alapú Performance Ajánlórendszer - MVP V1
# Streamlit app: upload -> standardization -> KPI engine -> insight/recommendation engine

from __future__ import annotations

import io
import re
from datetime import datetime
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st

try:
    from docx import Document
    from docx.shared import Inches, Pt
except Exception:
    Document = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
except Exception:
    SimpleDocTemplate = None


# -----------------------------
# Page config
# -----------------------------
st.set_page_config(
    page_title="Performance Ajánlórendszer MVP",
    page_icon="⚽",
    layout="wide",
)


# -----------------------------
# Constants / metric mapping
# -----------------------------
STANDARD_COLUMNS = {
    "player_name": ["Játékos neve", "Player", "Player Name", "Name", "Név"],
    "session_type": ["Típus", "Type", "Session Type", "Edzés/Meccs"],
    "session_name": ["Szakasz neve", "Session", "Session Name"],
    "start_time": ["Kezdési idő", "Start Time", "Start", "Dátum", "Date"],
    "end_time": ["Befejezési idő", "End Time", "End"],
    "duration": ["Időtartam", "Duration", "Time"],
    "total_distance": ["Tel\xadjes táv [m]", "Teljes táv [m]", "Total Distance", "Distance", "Össztáv"],
    "distance_per_min": ["Táv/perc [m/min]", "Distance/min", "Distance Per Min", "m/min"],
    "max_speed": ["Maximális sebesség [km/h]", "Max Speed", "Maximum Speed"],
    "avg_speed": ["Átlagsebesség [km/h]", "Average Speed", "Avg Speed"],
    "sprints": ["Sprintek", "Sprints", "Sprint Count"],
    "speed_zone_3": ["Táv a sebesség célzónában 3 [m] (14.40 - 19.79 km/h)"],
    "speed_zone_4": ["Táv a sebesség célzónában 4 [m] (19.80 - 24.99 km/h)"],
    "speed_zone_5": ["Táv a sebesség célzónában 5 [m] (25.00- km/h)"],
    "training_load": ["Edzési terhelési pontérték", "Terhelési pont", "Player Load", "Load"],
    "cardio_load": ["Kardióterhelés", "Cardio Load"],
    "recovery_hours": ["Regenerálódási idő [h]", "Recovery Time", "Recovery"],
    "muscle_load": ["Izomterhelés", "Muscle Load"],
    "hr_avg": ["Átlagos pulzus [bpm]", "Average HR", "Avg HR"],
    "hr_max": ["Maximális pulzus [bpm]", "Max HR", "Maximum HR"],
    "hrv": ["HRV (RMSSD)", "HRV", "RMSSD"],
    "acc_low": ["Gyorsulások száma (2.00 - 2.49 m/s²)"],
    "acc_mid": ["Gyorsulások száma (2.50 - 2.99 m/s²)"],
    "acc_high": ["Gyorsulások száma (3.00 - 50.00 m/s²)"],
    "dec_low": ["Gyorsulások száma (-2.49 - -2.00 m/s²)"],
    "dec_mid": ["Gyorsulások száma (-2.99 - -2.50 m/s²)"],
    "dec_high": ["Gyorsulások száma (-50.00 - -3.00 m/s²)"],
}

CORE_REQUIRED = ["player_name", "session_type", "start_time"]


@dataclass
class Insight:
    title: str
    severity: str
    observation: str
    impact: str
    recommendation: str
    scope: str = "Csapat"

    def as_dict(self) -> Dict[str, str]:
        return {
            "Súlyosság": self.severity,
            "Szint": self.scope,
            "Megállapítás": self.title,
            "Mit látunk?": self.observation,
            "Miért fontos?": self.impact,
            "Javaslat": self.recommendation,
        }


# -----------------------------
# Utility functions
# -----------------------------
def clean_col_name(col: object) -> str:
    if col is None:
        return ""
    return str(col).replace("\u00ad", "").strip()


def find_oszlop(df: pd.DataFrame, aliases: List[str]) -> Optional[str]:
    cleaned = {clean_col_name(c).lower(): c for c in df.oszlops}
    for alias in aliases:
        key = clean_col_name(alias).lower()
        if key in cleaned:
            return cleaned[key]
    # fuzzy contains fallback
    for alias in aliases:
        a = clean_col_name(alias).lower()
        for c in df.oszlops:
            cc = clean_col_name(c).lower()
            if a and (a in cc or cc in a):
                return c
    return None


def to_numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def duration_to_minutes(x) -> float:
    if pd.isna(x):
        return np.nan
    if isinstance(x, pd.Timedelta):
        return x.total_seconds() / 60
    if hasattr(x, "hour") and hasattr(x, "minute") and hasattr(x, "second"):
        return x.hour * 60 + x.minute + x.second / 60
    if isinstance(x, str):
        # Try formats like HH:MM:SS
        try:
            td = pd.to_timedelta(x)
            return td.total_seconds() / 60
        except Exception:
            return np.nan
    if isinstance(x, (int, float)):
        # Excel time fraction fallback: values under 2 are probably day fractions
        if x < 2:
            return x * 24 * 60
        return float(x)
    return np.nan


def normalize_session_type(x: object) -> str:
    text = str(x).strip().lower()
    if "meccs" in text or "match" in text or "game" in text:
        return "Meccs"
    if "edzés" in text or "training" in text:
        return "Edzés"
    return str(x).strip() if str(x).strip() else "Ismeretlen"


@st.cache_data(show_spinner=False)
def read_excel_all(file) -> Dict[str, pd.DataFrame]:
    return pd.read_excel(file, sheet_name=None)


def standardize_dataframe(raw: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, Optional[str]], List[str]]:
    df = raw.copy()
    df.oszlops = [clean_col_name(c) for c in df.oszlops]

    mapping: Dict[str, Optional[str]] = {}
    out = pd.DataFrame()

    for std_col, aliases in STANDARD_COLUMNS.items():
        source = find_oszlop(df, aliases)
        mapping[std_col] = source
        if source is not None:
            out[std_col] = df[source]

    missing_core = [c for c in CORE_REQUIRED if c not in out.oszlops]
    if missing_core:
        return out, mapping, missing_core

    out["player_name"] = out["player_name"].astype(str).str.strip()
    out["session_type"] = out["session_type"].apply(normalize_session_type)
    out["start_time"] = pd.to_datetime(out["start_time"], errors="coerce")
    out["session_date"] = out["start_time"].dt.date
    out["week"] = out["start_time"].dt.to_period("W").astype(str)

    if "duration" in out.oszlops:
        out["duration_min"] = out["duration"].apply(duration_to_minutes)
    else:
        out["duration_min"] = np.nan

    numeric_cols = [
        "total_distance", "distance_per_min", "max_speed", "avg_speed", "sprints",
        "speed_zone_3", "speed_zone_4", "speed_zone_5", "training_load", "cardio_load",
        "recovery_hours", "muscle_load", "hr_avg", "hr_max", "hrv", "acc_low", "acc_mid",
        "acc_high", "dec_low", "dec_mid", "dec_high",
    ]
    for col in numeric_cols:
        if col in out.oszlops:
            out[col] = to_numeric(out[col])

    # Derived fields
    for col in ["speed_zone_3", "speed_zone_4", "speed_zone_5", "acc_low", "acc_mid", "acc_high", "dec_low", "dec_mid", "dec_high"]:
        if col not in out.oszlops:
            out[col] = 0

    out["hsr_distance"] = out[["speed_zone_4", "speed_zone_5"]].sum(axis=1, min_count=1)
    out["sprint_distance"] = out["speed_zone_5"]
    out["acc_count"] = out[["acc_low", "acc_mid", "acc_high"]].sum(axis=1, min_count=1)
    out["dec_count"] = out[["dec_low", "dec_mid", "dec_high"]].sum(axis=1, min_count=1)
    out["high_efforts"] = out[["acc_mid", "acc_high", "dec_mid", "dec_high"]].sum(axis=1, min_count=1)

    if "distance_per_min" not in out.oszlops or out["distance_per_min"].isna().all():
        if "total_distance" in out.oszlops:
            out["distance_per_min"] = out["total_distance"] / out["duration_min"]

    # Remove empty/bad rows
    out = out.dropna(subset=["start_time"])
    out = out[out["player_name"].str.len() > 0]
    out = out[~out["player_name"].str.lower().str.contains("benchmark|átlag|összesen", na=False)]

    return out, mapping, []


def aggregate_weekly(df: pd.DataFrame) -> pd.DataFrame:
    agg_map = {
        "total_distance": "sum",
        "training_load": "sum",
        "muscle_load": "sum",
        "hsr_distance": "sum",
        "sprint_distance": "sum",
        "sprints": "sum",
        "acc_count": "sum",
        "dec_count": "sum",
        "high_efforts": "sum",
        "duration_min": "sum",
        "distance_per_min": "mean",
        "max_speed": "max",
        "hr_avg": "mean",
        "hrv": "mean",
    }
    usable = {k: v for k, v in agg_map.items() if k in df.oszlops}
    return df.groupby(["week", "session_type"], as_index=False).agg(usable)


def player_weekly(df: pd.DataFrame) -> pd.DataFrame:
    agg_map = {
        "total_distance": "sum",
        "training_load": "sum",
        "muscle_load": "sum",
        "hsr_distance": "sum",
        "sprint_distance": "sum",
        "sprints": "sum",
        "acc_count": "sum",
        "dec_count": "sum",
        "high_efforts": "sum",
        "duration_min": "sum",
        "distance_per_min": "mean",
        "max_speed": "max",
        "hr_avg": "mean",
        "hrv": "mean",
    }
    usable = {k: v for k, v in agg_map.items() if k in df.oszlops}
    return df.groupby(["player_name", "week", "session_type"], as_index=False).agg(usable)


def pct_change(current: float, previous: float) -> Optional[float]:
    if previous is None or pd.isna(previous) or previous == 0 or pd.isna(current):
        return None
    return (current - previous) / previous


def team_insights(df: pd.DataFrame, selected_week: str) -> List[Insight]:
    insights: List[Insight] = []
    weeks = sorted(df["week"].dropna().unique().tolist())
    week_idx = weeks.index(selected_week) if selected_week in weeks else len(weeks) - 1
    previous_week = weeks[week_idx - 1] if week_idx > 0 else None

    current = df[df["week"] == selected_week]
    prev = df[df["week"] == previous_week] if previous_week else pd.DataFrame()
    cur_train = current[current["session_type"] == "Edzés"]
    cur_match = current[current["session_type"] == "Meccs"]
    prev_train = prev[prev["session_type"] == "Edzés"] if not prev.empty else pd.DataFrame()

    # 1. Sprint underload: training sprint per session/player vs match sprint per session/player
    if not cur_train.empty and not cur_match.empty:
        train_sprint = cur_train["sprint_distance"].mean()
        match_sprint = cur_match["sprint_distance"].mean()
        if match_sprint and match_sprint > 0:
            ratio = train_sprint / match_sprint
            if ratio < 0.65:
                insights.append(Insight(
                    "Kevés sprintterhelés", "KRITIKUS",
                    f"Az edzés sprintterhelése a meccsterhelés kb. {ratio:.0%}-a.",
                    "A nagy intenzitású terhelés jelentősen elmaradhat a mérkőzésigénytől.",
                    "Érdemes lehet célzott sprint- vagy nagysebességű blokkokat beépíteni, ha ez illeszkedik a heti periodizációhoz.",
                ))
            elif ratio < 0.80:
                insights.append(Insight(
                    "Kevés sprintterhelés", "FIGYELMEZTETÉS",
                    f"Az edzés sprintterhelése a meccsterhelés kb. {ratio:.0%}-a.",
                    "A meccsigényhez képest mérsékelt intenzitási hiány látható.",
                    "Érdemes ellenőrizni, hogy tudatos visszaterhelésről vagy nem kívánt intenzitáshiányról van-e szó.",
                ))

    # 2. Weekly load change
    if not cur_train.empty and not prev_train.empty and "training_load" in df.oszlops:
        cur_load = cur_train["training_load"].sum()
        prev_load = prev_train["training_load"].sum()
        chg = pct_change(cur_load, prev_load)
        if chg is not None:
            if chg > 0.25:
                insights.append(Insight(
                    "Heti terhelési kiugrás", "FIGYELMEZTETÉS",
                    f"Az edzés terhelési pontértéke {chg:.0%}-kal nőtt az előző héthez képest.",
                    "A hirtelen terhelésemelkedés ronthatja a heti frissességet.",
                    "Érdemes figyelni a következő edzések intenzitását és a játékosok egyéni válaszait.",
                ))
            elif chg < -0.25:
                insights.append(Insight(
                    "Heti terheléscsökkenés", "INFORMÁCIÓ",
                    f"Az edzés terhelési pontértéke {abs(chg):.0%}-kal csökkent az előző héthez képest.",
                    "Ez lehet tudatos tapering, de lehet nem kívánt terhelésvesztés is.",
                    "Érdemes kontextusba helyezni: meccs előtti frissítés, hiányzók vagy edzéstartalom-váltás okozta-e.",
                ))

    # 3. Match intensity gap
    if not cur_train.empty and not cur_match.empty:
        train_int = cur_train["distance_per_min"].mean()
        match_int = cur_match["distance_per_min"].mean()
        if match_int and match_int > 0:
            ratio = train_int / match_int
            if ratio < 0.85:
                insights.append(Insight(
                    "Meccsintenzitástól való elmaradás", "FIGYELMEZTETÉS",
                    f"Az edzés átlagos táv/perc értéke a meccs kb. {ratio:.0%}-a.",
                    "A csapat edzésintenzitása elmaradhat a mérkőzés tempójától.",
                    "Érdemes lehet rövidebb, intenzívebb játékszituációkat vagy tempóváltásokat használni.",
                ))

    # 4. High deceleration load
    if not cur_train.empty and not prev_train.empty:
        cur_dec = cur_train["dec_count"].mean()
        prev_dec = prev_train["dec_count"].mean()
        chg = pct_change(cur_dec, prev_dec)
        if chg is not None and chg > 0.35:
            insights.append(Insight(
                "Magas lassítási terhelés", "FIGYELMEZTETÉS",
                f"Az átlagos lassításszám {chg:.0%}-kal nőtt az előző héthez képest.",
                "A lassítások nagy neuromuszkuláris terhelést jelenthetnek.",
                "Érdemes figyelni a regenerációra és a következő edzés excentrikus terhelésére.",
            ))

    # 5. Max speed suppression
    if not cur_train.empty and not prev_train.empty:
        cur_speed = cur_train["max_speed"].max()
        prev_speed = prev_train["max_speed"].max()
        chg = pct_change(cur_speed, prev_speed)
        if chg is not None and chg < -0.05:
            insights.append(Insight(
                "Maximális sebesség visszaesése", "INFORMÁCIÓ",
                f"A heti maximális sebesség {abs(chg):.0%}-kal alacsonyabb az előző hétnél.",
                "Ez jelezhet alacsonyabb neuromuszkuláris frissességet, de lehet edzéstartalom-függő is.",
                "Érdemes megnézni, volt-e valódi max sebességű inger a héten.",
            ))

    # 6. Player outlier alerts
    if not cur_train.empty and "training_load" in cur_train.oszlops:
        player_load = cur_train.groupby("player_name")["training_load"].sum().dropna()
        if len(player_load) >= 5 and player_load.mean() > 0:
            mean = player_load.mean()
            high = player_load[player_load > mean * 1.25].sort_values(ascending=False)
            low = player_load[player_load < mean * 0.75].sort_values(ascending=True)
            if len(high) > 0:
                names = ", ".join(high.head(3).index.tolist())
                insights.append(Insight(
                    "Kiugró játékosterhelések", "FIGYELMEZTETÉS",
                    f"Néhány játékos jelentősen a csapatátlag felett terhelődött: {names}.",
                    "Egyéni terheléskülönbség alakult ki a héten.",
                    "Érdemes egyéni szinten ránézni a következő edzés terhelésére és a játékpercekre.",
                    scope="Játékos",
                ))
            if len(low) > 0:
                names = ", ".join(low.head(3).index.tolist())
                insights.append(Insight(
                    "Alacsony játékosterhelések", "INFORMÁCIÓ",
                    f"Néhány játékos jelentősen a csapatátlag alatt terhelődött: {names}.",
                    "Terheléslemaradás alakulhat ki, főleg ha ez több héten át fennáll.",
                    "Érdemes ellenőrizni a hiányzásokat, játékperceket és egyéni kiegészítő munkát.",
                    scope="Játékos",
                ))

    # If no insights, return positive summary
    if not insights:
        insights.append(Insight(
            "Stabil hét", "INFORMÁCIÓ",
            "Nem látható kiemelt negatív eltérés az aktuális hét fő mutatóiban.",
            "A csapat terhelési profilja stabilnak tűnik az elérhető adatok alapján.",
            "Érdemes tovább figyelni a sprint- és intenzitási trendeket, különösen meccs előtti héten.",
        ))

    severity_rank = {"KRITIKUS": 0, "FIGYELMEZTETÉS": 1, "INFORMÁCIÓ": 2}
    return sorted(insights, key=lambda x: severity_rank.get(x.severity, 9))[:8]


def style_severity(sev: str) -> str:
    if sev == "KRITIKUS":
        return "🔴 KRITIKUS"
    if sev == "FIGYELMEZTETÉS":
        return "🟠 FIGYELMEZTETÉS"
    return "🔵 INFORMÁCIÓ"


def metric_card(label: str, value: object, help_text: str = ""):
    st.metric(label, value if value is not None else "—", help=help_text)


def _safe_filename_week(week: str) -> str:
    return re.sub(r"[^0-9A-Za-z_-]+", "_", str(week)).strip("_") or "week"


def build_insight_export_df(insights: List[Insight]) -> pd.DataFrame:
    """Create a user-friendly, export-ready insight table."""
    df = pd.DataFrame([i.as_dict() for i in insights])
    if df.empty:
        return df
    order = {"KRITIKUS": 1, "FIGYELMEZTETÉS": 2, "INFORMÁCIÓ": 3}
    df.insert(0, "Prioritás", df["Súlyosság"].map(order).fillna(9).astype(int))
    df["Súlyosság"] = df["Súlyosság"].map({
        "KRITIKUS": "🔴 KRITIKUS",
        "FIGYELMEZTETÉS": "🟠 FIGYELMEZTETÉS",
        "INFORMÁCIÓ": "🔵 INFORMÁCIÓ",
    }).fillna(df["Súlyosság"])
    return df


def insights_to_excel_bytes(insights_df: pd.DataFrame, selected_week: str) -> bytes:
    """Export insights to a formatted Excel file with wrapped text and readable oszlops."""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        insights_df.to_excel(writer, index=False, sheet_name="Insightok")
        wb = writer.book
        ws = writer.sheets["Insightok"]

        from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
        from openpyxl.utils import get_oszlop_letter

        # Title rows
        ws.insert_rows(1, 3)
        ws["A1"] = "Performance megállapítások és javaslatok"
        ws["A2"] = f"Hét: {selected_week}"
        ws["A3"] = f"Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ws.merge_cells(start_row=1, start_oszlop=1, end_row=1, end_oszlop=max(1, len(insights_df.oszlops)))
        ws.merge_cells(start_row=2, start_oszlop=1, end_row=2, end_oszlop=max(1, len(insights_df.oszlops)))
        ws.merge_cells(start_row=3, start_oszlop=1, end_row=3, end_oszlop=max(1, len(insights_df.oszlops)))
        ws["A1"].font = Font(bold=True, size=16)
        ws["A2"].font = Font(bold=True, size=11)
        ws["A3"].font = Font(italic=True, size=10)

        header_row = 4
        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(color="FFFFFF", bold=True)
        thin = Side(style="thin", color="D9E2F3")

        for cell in ws[header_row]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

        # Readable widths; wrap long text.
        widths = {
            "Prioritás": 10,
            "Súlyosság": 16,
            "Szint": 14,
            "Megállapítás": 24,
            "Mit látunk?": 45,
            "Miért fontos?": 45,
            "Javaslat": 55,
        }
        for idx, col_name in enumerate(insights_df.oszlops, start=1):
            ws.oszlop_dimensions[get_oszlop_letter(idx)].width = widths.get(str(col_name), 22)

        for row in ws.iter_rows(min_row=5, max_row=ws.max_row, max_col=ws.max_oszlop):
            max_lines = 1
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
                text = str(cell.value or "")
                max_lines = max(max_lines, min(6, (len(text) // 45) + 1))
            ws.row_dimensions[row[0].row].height = max(24, max_lines * 18)

        ws.freeze_panes = "A5"
        ws.auto_filter.ref = ws.dimensions
    return output.getvalue()


def insights_to_word_bytes(insights_df: pd.DataFrame, selected_week: str) -> Optional[bytes]:
    """Export insights to a Word report. Requires python-docx."""
    if Document is None:
        return None
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    doc.add_heading("Performance megállapítások és javaslatok", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Hét: {selected_week}\n").bold = True
    p.add_run(f"Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for _, row in insights_df.iterrows():
        doc.add_heading(f"{row.get('Severity', '')} · {row.get('Title', '')}", level=2)
        for label in ["Mit látunk?", "Miért fontos?", "Javaslat"]:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(row.get(label, "")))

    doc.add_page_break()
    doc.add_heading("Export table", level=2)
    table = doc.add_table(rows=1, cols=len(insights_df.oszlops))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(insights_df.oszlops):
        hdr_cells[i].text = str(col)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for _, row in insights_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(insights_df.oszlops):
            cells[i].text = str(row.get(col, ""))
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def insights_to_pdf_bytes(insights_df: pd.DataFrame, selected_week: str) -> Optional[bytes]:
    """Export insights to a landscape PDF with wrapped text. Requires reportlab."""
    if SimpleDocTemplate is None:
        return None
    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=0.8 * cm,
        leftMargin=0.8 * cm,
        topMargin=0.8 * cm,
        bottomMargin=0.8 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    normal = ParagraphStyle("NormalWrapped", parent=styles["Normal"], fontSize=7, leading=9)
    header = ParagraphStyle("HeaderWrapped", parent=styles["Normal"], fontSize=7, leading=9, textColor=colors.white)

    story = [
        Paragraph("Performance megállapítások és javaslatok", title_style),
        Paragraph(f"Hét: {selected_week} · Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]),
        Spacer(1, 0.25 * cm),
    ]

    cols = ["Súlyosság", "Szint", "Megállapítás", "Mit látunk?", "Miért fontos?", "Javaslat"]
    cols = [c for c in cols if c in insights_df.oszlops]
    table_data = [[Paragraph(str(c), header) for c in cols]]
    for _, row in insights_df.iterrows():
        table_data.append([Paragraph(str(row.get(c, "")), normal) for c in cols])

    col_widths_map = {
        "Súlyosság": 2.4 * cm,
        "Szint": 2.0 * cm,
        "Megállapítás": 3.2 * cm,
        "Mit látunk?": 6.2 * cm,
        "Miért fontos?": 6.2 * cm,
        "Javaslat": 7.0 * cm,
    }
    table = Table(table_data, colWidths=[col_widths_map.get(c, 3 * cm) for c in cols], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BFBFBF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FC")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    doc.build(story)
    return output.getvalue()


# -----------------------------
# UI
# -----------------------------
st.title("⚽ Performance Ajánlórendszer – MVP")
st.caption("Adatfeltöltés → KPI-k → szakmai megállapítások → edzői javaslatok")

with st.sidebar:
    st.header("1) Adatfeltöltés")
    uploaded = st.file_uploader("Tölts fel Excel fájlt", type=["xlsx", "xls"])
    st.divider()
    st.markdown("**MVP fókusz:** edzői döntéstámogatás, nem sérülésdiagnosztika.")

if uploaded is None:
    st.info("Tölts fel egy GPS/terhelési Excel fájlt a kezdéshez.")
    st.stop()

sheets = read_excel_all(uploaded)
sheet_names = list(sheets.keys())

with st.sidebar:
    selected_sheet = st.selectbox("Melyik munkalapot használjuk?", sheet_names, index=0)

raw_df = sheets[selected_sheet]
df, mapping, missing_core = standardize_dataframe(raw_df)

if missing_core:
    st.error(f"Hiányzó alapmezők: {', '.join(missing_core)}")
    st.write("Oszlopmapping:", mapping)
    st.stop()

# Sidebar filters
weeks = sorted(df["week"].dropna().unique().tolist())
players = sorted(df["player_name"].dropna().unique().tolist())
session_types = sorted(df["session_type"].dropna().unique().tolist())

with st.sidebar:
    st.header("2) Szűrők")
    selected_week = st.selectbox("Hét", weeks, index=len(weeks) - 1 if weeks else 0)
    selected_types = st.multiselect("Típus", session_types, default=session_types)
    selected_players = st.multiselect("Játékosok", players, default=players)

filtered = df[(df["week"] == selected_week) & (df["session_type"].isin(selected_types)) & (df["player_name"].isin(selected_players))]

# Tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "Csapat áttekintő", "Megállapítások és javaslatok", "Játékosmonitoring", "Adatminőség", "Nyers adatok"
])

with tab1:
    st.subheader("Csapat áttekintő")
    st.caption("A főoldal célja: 30 másodperc alatt látni, mi fontos a héten.")

    col1, col2, col3, col4 = st.oszlops(4)
    with col1:
        metric_card("Játékosok", filtered["player_name"].nunique())
    with col2:
        metric_card("Sessionök száma", len(filtered))
    with col3:
        td = filtered["total_distance"].sum() if "total_distance" in filtered.oszlops else np.nan
        metric_card("Össztáv", f"{td:,.0f} m" if pd.notna(td) else "—")
    with col4:
        load = filtered["training_load"].sum() if "training_load" in filtered.oszlops else np.nan
        metric_card("Terhelési pont", f"{load:,.0f}" if pd.notna(load) else "—")

    weekly = aggregate_weekly(df[df["session_type"].isin(selected_types)])
    st.markdown("### Heti trendek")
    chart_metric = st.selectbox(
        "Trendmutató",
        [m for m in ["training_load", "total_distance", "sprint_distance", "hsr_distance", "distance_per_min", "max_speed", "dec_count"] if m in weekly.oszlops],
        index=0,
    )
    fig = px.line(
        weekly,
        x="week",
        y=chart_metric,
        color="session_type",
        markers=True,
        title=f"Heti trend: {chart_metric}",
    )
    st.plotly_chart(fig, use_container_width=True)

    st.markdown("### Edzés vs meccs profil az aktuális héten")
    profile_cols = ["total_distance", "distance_per_min", "sprint_distance", "hsr_distance", "max_speed", "dec_count"]
    usable_profile = [c for c in profile_cols if c in filtered.oszlops]
    if usable_profile:
        prof = filtered.groupby("session_type", as_index=False)[usable_profile].mean(numeric_only=True)
        st.dataframe(prof, use_container_width=True)

with tab2:
    st.subheader("Megállapítások és javaslatok")
    st.caption("Szabályalapú performance motor: AI nélkül is ad szakmai következtetést és javaslatot.")
    insights = team_insights(df[df["player_name"].isin(selected_players)], selected_week)
    for ins in insights:
        with st.container(border=True):
            st.markdown(f"### {style_severity(ins.severity)} · {ins.title}")
            c1, c2, c3 = st.oszlops(3)
            with c1:
                st.markdown("**Observation**")
                st.write(ins.observation)
            with c2:
                st.markdown("**Impact**")
                st.write(ins.impact)
            with c3:
                st.markdown("**Recommendation**")
                st.write(ins.recommendation)

    st.markdown("### Exportálható megállapítási tábla")
    insight_export_df = build_insight_export_df(insights)

    st.dataframe(
        insight_export_df,
        use_container_width=True,
        hide_index=True,
        oszlop_config={
            "Mit látunk?": st.column_config.TextColumn("Mit látunk?", width="large"),
            "Miért fontos?": st.column_config.TextColumn("Miért fontos?", width="large"),
            "Javaslat": st.column_config.TextColumn("Javaslat", width="large"),
            "Megállapítás": st.column_config.TextColumn("Megállapítás", width="medium"),
        },
    )

    st.markdown("#### Riport exportálása")
    safe_week = _safe_filename_week(selected_week)
    c1, c2, c3, c4 = st.oszlops(4)
    with c1:
        st.download_button(
            "⬇️ Excel riport",
            data=insights_to_excel_bytes(insight_export_df, selected_week),
            file_name=f"performance_riport_{safe_week}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    with c2:
        csv_bytes = insight_export_df.to_csv(index=False).encode("utf-8-sig")
        st.download_button(
            "⬇️ CSV",
            data=csv_bytes,
            file_name=f"performance_riport_{safe_week}.csv",
            mime="text/csv",
            use_container_width=True,
        )
    with c3:
        word_bytes = insights_to_word_bytes(insight_export_df, selected_week)
        if word_bytes is not None:
            st.download_button(
                "⬇️ Word riport",
                data=word_bytes,
                file_name=f"performance_riport_{safe_week}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.info("Word exporthoz add hozzá a requirements.txt fájlhoz: python-docx")
    with c4:
        pdf_bytes = insights_to_pdf_bytes(insight_export_df, selected_week)
        if pdf_bytes is not None:
            st.download_button(
                "⬇️ PDF riport",
                data=pdf_bytes,
                file_name=f"performance_riport_{safe_week}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        else:
            st.info("PDF exporthoz add hozzá a requirements.txt fájlhoz: reportlab")

with tab3:
    st.subheader("Játékosmonitoring")
    pw = player_weekly(df)
    pw_current = pw[(pw["week"] == selected_week) & (pw["player_name"].isin(selected_players))]
    if not pw_current.empty:
        metric = st.selectbox(
            "Játékos rangsor mutató",
            [m for m in ["training_load", "total_distance", "sprint_distance", "hsr_distance", "max_speed", "dec_count", "high_efforts"] if m in pw_current.oszlops],
            index=0,
        )
        rank = pw_current.groupby("player_name", as_index=False)[metric].sum().sort_values(metric, ascending=False)
        fig = px.bar(rank.head(25), x="player_name", y=metric, title=f"Játékosrangsor: {metric}")
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)

        st.markdown("### Játékos heti összesítő tábla")
        st.dataframe(pw_current, use_container_width=True)
    else:
        st.info("Nincs adat az aktuális szűrésre.")

with tab4:
    st.subheader("Adatminőség")
    st.caption("Ez azért fontos, mert a sportadat a valóságban mindig kicsit koszos.")

    c1, c2, c3, c4 = st.oszlops(4)
    with c1:
        metric_card("Nyers sor", len(raw_df))
    with c2:
        metric_card("Standardizált sor", len(df))
    with c3:
        metric_card("Hetek", df["week"].nunique())
    with c4:
        metric_card("Típusok", ", ".join(session_types))

    st.markdown("### Automatikus oszlopmapping")
    map_df = pd.DataFrame([
        {"standard mező": k, "forrás oszlop": v if v is not None else "NINCS"}
        for k, v in mapping.items()
    ])
    st.dataframe(map_df, use_container_width=True)

    st.markdown("### Hiányzó értékek a standard mezőkben")
    na = df.isna().mean().sort_values(ascending=False).reset_index()
    na.oszlops = ["oszlop", "hiányzó_arány"]
    st.dataframe(na.head(30), use_container_width=True)

with tab5:
    st.subheader("Nyers / standardizált adatok")
    st.markdown("### Standardizált adat")
    st.dataframe(df.head(500), use_container_width=True)
    st.markdown("### Nyers adat")
    st.dataframe(raw_df.head(200), use_container_width=True)


st.divider()
st.caption("MVP V1.2 HU – szabályalapú motor + vizualizált Excel/Word/PDF export. Következő lépés: finomított szabályok, benchmarkok, majd AI-os heti összefoglaló.")
