# performance_app_v2_0_hu.py
# AI-assisted Performance Recommendation System - magyar Streamlit MVP
# Upload -> standardizálás -> KPI-k -> szabályalapú insightok -> coach-friendly javaslatok -> Excel/Word/PDF export

from __future__ import annotations

import html
import io
import hashlib
import os
import json
import re
import unicodedata
import zipfile
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st
try:
    import pdfplumber
except Exception:
    pdfplumber = None

PYMUPDF_AVAILABLE = True
try:
    import fitz  # PyMuPDF
except Exception:
    PYMUPDF_AVAILABLE = False
    fitz = None

PYPDF_AVAILABLE = True
try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PYPDF_AVAILABLE = False
        PdfReader = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
except Exception:
    Document = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except Exception:
    SimpleDocTemplate = None
    pdfmetrics = None
    TTFont = None


try:
    from supabase import create_client
except Exception:
    create_client = None

FPI_IMPORT_ENGINE_VERSION = "FPI_TACTICAL_MERGE_V147_SUMMARY_LAYOUT_FITNESS_DEPTH_2026_07_11"

# -----------------------------------------------------------------------------
# Oldalbeállítás
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="Football Performance Intelligence",
    page_icon="⚽",
    layout="wide",
)

st.markdown(
    """
    <style>
    .insight-card {
        border-radius: 16px;
        padding: 18px 20px;
        margin-bottom: 16px;
        background: rgba(31, 41, 55, 0.72);
        border: 1px solid rgba(255,255,255,0.10);
        box-shadow: 0 6px 18px rgba(0,0,0,0.12);
    }
    .insight-title {
        font-size: 1.15rem;
        font-weight: 800;
        margin-bottom: 6px;
    }
    .pill {
        display: inline-block;
        padding: 4px 10px;
        border-radius: 999px;
        font-size: 0.82rem;
        font-weight: 700;
        margin-right: 8px;
    }
    .pill-critical { background:#7f1d1d; color:#fecaca; }
    .pill-warning { background:#78350f; color:#fde68a; }
    .pill-info { background:#1e3a8a; color:#bfdbfe; }
    .insight-label { font-weight: 800; margin-top: 10px; }
    .insight-text { line-height: 1.45; margin-bottom: 6px; }
    .wrap-table table { width: 100%; border-collapse: collapse; table-layout: fixed; }
    .wrap-table th, .wrap-table td {
        border: 1px solid rgba(255,255,255,0.12);
        padding: 10px;
        vertical-align: top;
        white-space: normal !important;
        overflow-wrap: break-word;
        word-break: normal;
        font-size: 0.92rem;
    }
    .wrap-table th { background: rgba(30, 64, 175, 0.60); font-weight: 800; }
    .wrap-table tr:nth-child(even) { background: rgba(255,255,255,0.03); }
    .priority-card {
        border-radius: 16px;
        padding: 16px 18px;
        margin-bottom: 12px;
        background: rgba(17, 24, 39, 0.78);
        border-left: 7px solid #22c55e;
        box-shadow: 0 6px 18px rgba(0,0,0,0.10);
    }
    .micro-pill {
        display: inline-block;
        padding: 6px 10px;
        border-radius: 10px;
        background: rgba(59,130,246,.18);
        border: 1px solid rgba(147,197,253,.22);
        margin: 3px;
        font-weight: 700;
    }

    .score-card {
        border-radius: 18px;
        padding: 18px 20px;
        margin-bottom: 14px;
        background: linear-gradient(135deg, rgba(15,23,42,.92), rgba(30,41,59,.82));
        border: 1px solid rgba(148,163,184,.22);
        box-shadow: 0 8px 24px rgba(0,0,0,.14);
    }
    .score-number {
        font-size: 2.4rem;
        font-weight: 900;
        line-height: 1;
    }
    .score-label {
        color: #e0f2fe;
        font-weight: 800;
        margin-top: 6px;
    }
    .mini-muted {
        color: #dbeafe;
        font-size: .92rem;
        line-height: 1.35;
    }

    .hero-box {border-radius:24px;padding:24px 28px;margin-bottom:20px;background:radial-gradient(circle at top left,rgba(34,197,94,.22),transparent 34%),radial-gradient(circle at bottom right,rgba(59,130,246,.24),transparent 30%),linear-gradient(135deg,rgba(2,6,23,.96),rgba(15,23,42,.88));border:1px solid rgba(148,163,184,.22);box-shadow:0 18px 45px rgba(0,0,0,.28)}
    .hero-title {font-size:2.1rem;font-weight:950;letter-spacing:-.04em;margin-bottom:4px}.hero-sub{color:#dbeafe;font-size:1.02rem;line-height:1.45}
    .premium-kpi{border-radius:20px;padding:18px;background:linear-gradient(145deg,rgba(15,23,42,.94),rgba(30,41,59,.78));border:1px solid rgba(148,163,184,.20);box-shadow:0 10px 28px rgba(0,0,0,.18);min-height:120px}
    .premium-kpi-label{color:#dbeafe;font-size:.86rem;font-weight:800;text-transform:uppercase;letter-spacing:.06em}.premium-kpi-value{font-size:2rem;font-weight:950;margin-top:8px;line-height:1}.premium-kpi-note{color:#dbeafe;font-size:.86rem;margin-top:9px}
    .risk-high{border-left:8px solid #ef4444!important}.risk-medium{border-left:8px solid #f59e0b!important}.risk-low{border-left:8px solid #22c55e!important}
    .section-chip{display:inline-block;padding:5px 11px;border-radius:999px;background:rgba(34,197,94,.15);border:1px solid rgba(34,197,94,.25);color:#bbf7d0;font-weight:850;margin:2px 4px 8px 0}

    .intro-card {
        border-radius: 22px;
        padding: 22px 24px;
        margin-bottom: 16px;
        background: linear-gradient(135deg, rgba(15,23,42,.95), rgba(30,41,59,.82));
        border: 1px solid rgba(148,163,184,.24);
        box-shadow: 0 12px 34px rgba(0,0,0,.20);
    }
    .intro-card h2, .intro-card h3 {
        margin-top: 0;
        margin-bottom: 8px;
    }
    .intro-grid {
        display: grid;
        grid-template-columns: repeat(2, minmax(0, 1fr));
        gap: 14px;
        margin-top: 14px;
    }
    .feature-box {
        border-radius: 18px;
        padding: 16px 18px;
        background: rgba(15,23,42,.72);
        border: 1px solid rgba(148,163,184,.18);
    }
    .feature-title {
        font-weight: 900;
        font-size: 1.02rem;
        margin-bottom: 6px;
        color: #bfdbfe;
    }
    .feature-text {
        color: #e0f2fe;
        line-height: 1.45;
        font-size: .94rem;
    }
    .export-panel {
        border-radius: 22px;
        padding: 20px;
        background: radial-gradient(circle at top right, rgba(34,197,94,.16), transparent 35%),
                    linear-gradient(135deg, rgba(2,6,23,.94), rgba(15,23,42,.86));
        border: 1px solid rgba(148,163,184,.24);
        margin-bottom: 18px;
    }
   
 .fpi-clean-card {
   border-radius: 20px;
   padding: 18px 20px;
   background: #ffffff;
   color: #0f172a;
   border: 1px solid #e2e8f0;
   box-shadow: 0 10px 28px rgba(15,23,42,.08);
   margin-bottom: 14px;
 }
 .fpi-dark-card {
   border-radius: 20px;
   padding: 18px 20px;
   background: linear-gradient(135deg,#0f172a,#1e293b);
   color: #f8fafc;
   border: 1px solid rgba(148,163,184,.28);
   box-shadow: 0 14px 36px rgba(15,23,42,.22);
   margin-bottom: 14px;
 }
 .fpi-muted { color:#64748b; }
 .fpi-dark-card .fpi-muted { color:#cbd5e1; }
 .export-panel p, .export-panel div, .export-panel h3 { color: #f8fafc; }
 .feature-box, .intro-card, .score-card, .priority-card, .insight-card { color: #f8fafc; }
 .feature-text, .mini-muted { color: #dbeafe !important; }


 /* ===== FPI WOW UI refresh ===== */
 .stApp {
   background:
     radial-gradient(circle at 10% 5%, rgba(34,197,94,.18), transparent 28%),
     radial-gradient(circle at 90% 0%, rgba(59,130,246,.20), transparent 30%),
     linear-gradient(135deg, #020617 0%, #0f172a 52%, #111827 100%);
 }
 .block-container { padding-top: 1.1rem; padding-bottom: 2rem; }
 .fpi-hero-wow {
   border-radius: 30px; padding: 28px 30px; margin: 8px 0 22px 0;
   background: radial-gradient(circle at top left, rgba(34,197,94,.30), transparent 34%),
               radial-gradient(circle at bottom right, rgba(59,130,246,.34), transparent 30%),
               linear-gradient(135deg, rgba(15,23,42,.98), rgba(30,41,59,.86));
   border: 1px solid rgba(226,232,240,.18);
   box-shadow: 0 24px 70px rgba(0,0,0,.35);
   color: #f8fafc;
 }
 .fpi-hero-wow h1 { margin: 0; font-size: 2.55rem; line-height: 1; letter-spacing: -0.055em; font-weight: 950; }
 .fpi-hero-wow p { margin: 10px 0 0 0; color: #cbd5e1; font-size: 1.02rem; line-height: 1.45; }
 .fpi-chip-row { margin-top: 16px; }
 .fpi-chip-wow {
   display: inline-block; padding: 7px 12px; margin: 4px 6px 0 0; border-radius: 999px;
   background: rgba(15,23,42,.55); border: 1px solid rgba(148,163,184,.28);
   color: #dbeafe; font-weight: 800; font-size: .86rem;
 }
 .fpi-summary-card {
   border-radius: 24px; padding: 20px 22px; margin: 10px 0 18px 0;
   background: rgba(248,250,252,.97); color: #0f172a;
   border: 1px solid rgba(226,232,240,.95); box-shadow: 0 18px 46px rgba(15,23,42,.18);
 }
 .fpi-summary-card h3 { margin: 0 0 10px 0; font-size: 1.18rem; font-weight: 950; color: #0f172a; }
 .fpi-summary-grid { display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 10px 14px; margin-top: 12px; }
 .fpi-summary-item { border-radius: 16px; padding: 12px 13px; background: #f8fafc; border: 1px solid #e2e8f0; }
 .fpi-summary-label { font-size: .75rem; text-transform: uppercase; letter-spacing: .06em; color: #64748b; font-weight: 900; margin-bottom: 4px; }
 .fpi-summary-value { color: #0f172a; font-weight: 800; line-height: 1.35; }
 .fpi-action-card {
   border-radius: 20px; padding: 16px 18px; margin: 12px 0;
   background: linear-gradient(135deg, #ecfdf5, #eff6ff);
   color: #0f172a; border: 1px solid #bfdbfe; box-shadow: 0 10px 28px rgba(37,99,235,.12);
 }
 .fpi-action-card b { color:#0f172a; }
 .fpi-kpi-panel {
   border-radius: 22px; padding: 17px 18px;
   background: linear-gradient(145deg, rgba(255,255,255,.98), rgba(239,246,255,.96));
   color: #0f172a; border: 1px solid rgba(219,234,254,.95);
   box-shadow: 0 14px 34px rgba(15,23,42,.14); min-height: 118px;
 }
 .fpi-kpi-panel .label { color: #64748b; font-size: .78rem; font-weight: 900; text-transform: uppercase; letter-spacing: .07em; }
 .fpi-kpi-panel .value { color: #0f172a; font-size: 2.05rem; font-weight: 950; letter-spacing: -.04em; margin-top: 6px; }
 .fpi-kpi-panel .note { color: #475569; font-size: .86rem; margin-top: 5px; line-height: 1.35; }
 .fpi-section-title { color: #f8fafc; font-size: 1.35rem; font-weight: 950; margin: 22px 0 10px 0; letter-spacing: -.03em; }
 .insight-card, .priority-card, .score-card, .intro-card, .feature-box, .export-panel { color: #f8fafc !important; }
 .feature-text, .mini-muted, .premium-kpi-note, .hero-sub { color: #dbeafe !important; }


 /* ===== FPI contrast/readability fix ===== */
 html, body, .stApp, [data-testid="stAppViewContainer"] {
   color: #f8fafc !important;
 }
 [data-testid="stSidebar"] {
   background: linear-gradient(180deg, #020617, #0f172a) !important;
 }
 [data-testid="stSidebar"] * {
   color: #f8fafc !important;
 }
 .stMarkdown, .stText, p, span, label, div {
   text-rendering: optimizeLegibility;
 }
 .fpi-hero-wow,
 .fpi-hero-wow *,
 .fpi-dark-card,
 .fpi-dark-card *,
 .fpi-glass,
 .fpi-glass *,
 .intro-card,
 .intro-card *,
 .feature-box,
 .feature-box *,
 .score-card,
 .score-card *,
 .priority-card,
 .priority-card *,
 .insight-card,
 .insight-card *,
 .export-panel,
 .export-panel * {
   color: #f8fafc !important;
 }
 .fpi-hero-wow p,
 .hero-sub,
 .feature-text,
 .mini-muted,
 .premium-kpi-note,
 .score-label {
   color: #dbeafe !important;
 }
 .fpi-summary-card,
 .fpi-summary-card *,
 .fpi-summary-item,
 .fpi-summary-item *,
 .fpi-action-card,
 .fpi-action-card *,
 .fpi-kpi-panel,
 .fpi-kpi-panel * {
   color: #0f172a !important;
 }
 .fpi-summary-label,
 .fpi-kpi-panel .label,
 .fpi-kpi-panel .note,
 .fpi-muted {
   color: #475569 !important;
 }
 .fpi-chip-wow,
 .section-chip,
 .micro-pill {
   color: #f8fafc !important;
   background: rgba(15,23,42,.78) !important;
   border-color: rgba(226,232,240,.32) !important;
 }
 .pill-critical { background:#991b1b !important; color:#ffffff !important; }
 .pill-warning { background:#92400e !important; color:#ffffff !important; }
 .pill-info { background:#1d4ed8 !important; color:#ffffff !important; }
 .wrap-table th {
   background: #1e3a8a !important;
   color: #ffffff !important;
 }
 .wrap-table td {
   color: #f8fafc !important;
   background: rgba(15,23,42,.72) !important;
 }
 .wrap-table tr:nth-child(even) td {
   background: rgba(30,41,59,.82) !important;
 }
 .stDataFrame, .stDataFrame * {
   color: inherit;
 }
 div[data-testid="stMetricValue"] {
   color: #f8fafc !important;
 }
 div[data-testid="stMetricLabel"] {
   color: #dbeafe !important;
 }
 .stAlert, .stAlert * {
   color: #0f172a !important;
 }
 button, button * {
   font-weight: 800 !important;
 }


 /* ===== FPI tab + CTA readability fix ===== */
 button[data-baseweb="tab"] {
   color: #cbd5e1 !important;
   font-weight: 850 !important;
   opacity: 1 !important;
   background: transparent !important;
 }
 button[data-baseweb="tab"] * {
   color: #cbd5e1 !important;
   font-weight: 850 !important;
   opacity: 1 !important;
 }
 button[data-baseweb="tab"][aria-selected="true"],
 button[data-baseweb="tab"][aria-selected="true"] * {
   color: #ffffff !important;
   font-weight: 950 !important;
 }
 button[data-baseweb="tab"]:hover,
 button[data-baseweb="tab"]:hover * {
   color: #ffffff !important;
 }
 div[data-baseweb="tab-list"] {
   gap: 8px;
   background: rgba(15,23,42,.42);
   border-radius: 18px;
   padding: 8px 10px;
   border: 1px solid rgba(148,163,184,.18);
 }
 div[data-baseweb="tab-border"] {
   background-color: #ef4444 !important;
   height: 3px !important;
 }
 .stDownloadButton > button {
   background: linear-gradient(135deg,#22c55e,#16a34a) !important;
   color: #ffffff !important;
   border: 0 !important;
   border-radius: 16px !important;
   font-weight: 950 !important;
   min-height: 48px !important;
   box-shadow: 0 14px 34px rgba(34,197,94,.24) !important;
 }
 .stDownloadButton > button:hover {
   background: linear-gradient(135deg,#16a34a,#15803d) !important;
   color: #ffffff !important;
   transform: translateY(-1px);
   box-shadow: 0 18px 42px rgba(34,197,94,.32) !important;
 }
 .stDownloadButton > button * {
   color: #ffffff !important;
   font-weight: 950 !important;
 }
 .stButton > button {
   border-radius: 14px !important;
   font-weight: 900 !important;
 }
 input, textarea, [data-baseweb="input"] input {
   color: #0f172a !important;
   background: #ffffff !important;
 }
 [data-testid="stSidebar"] .stAlert,
 [data-testid="stSidebar"] .stAlert * {
   color: #0f172a !important;
 }
 [data-testid="stSidebar"] .stAlert {
   background: #dbeafe !important;
   border-radius: 14px !important;
 }


 /* ===== FPI HARD READABILITY PATCH ===== */
 :root {
   --fpi-dark: #020617;
   --fpi-panel: #0f172a;
   --fpi-panel-2: #111827;
   --fpi-white: #f8fafc;
   --fpi-muted-light: #dbeafe;
   --fpi-text-dark: #0f172a;
   --fpi-green: #22c55e;
   --fpi-blue: #2563eb;
 }

 /* Main text on dark background */
 .stApp, .stApp p, .stApp span, .stApp label, .stApp div {
   text-shadow: none !important;
 }

 /* Sidebar: everything readable */
 [data-testid="stSidebar"] {
   background: linear-gradient(180deg, #020617 0%, #0b1220 100%) !important;
 }
 [data-testid="stSidebar"] h1,
 [data-testid="stSidebar"] h2,
 [data-testid="stSidebar"] h3,
 [data-testid="stSidebar"] p,
 [data-testid="stSidebar"] label,
 [data-testid="stSidebar"] span,
 [data-testid="stSidebar"] div {
   color: #f8fafc !important;
 }

 /* Sidebar input fields */
 [data-testid="stSidebar"] input,
 [data-testid="stSidebar"] textarea,
 [data-testid="stSidebar"] [data-baseweb="input"],
 [data-testid="stSidebar"] [data-baseweb="input"] > div {
   background: #ffffff !important;
   color: #0f172a !important;
   border-color: #cbd5e1 !important;
   border-radius: 12px !important;
 }
 [data-testid="stSidebar"] input::placeholder {
   color: #64748b !important;
   opacity: 1 !important;
 }

 /* Sidebar buttons: no white-on-white */
 [data-testid="stSidebar"] .stButton > button,
 [data-testid="stSidebar"] button {
   background: linear-gradient(135deg,#2563eb,#1d4ed8) !important;
   color: #ffffff !important;
   border: 1px solid rgba(255,255,255,.25) !important;
   border-radius: 14px !important;
   font-weight: 900 !important;
   min-height: 44px !important;
 }
 [data-testid="stSidebar"] .stButton > button *,
 [data-testid="stSidebar"] button * {
   color: #ffffff !important;
   font-weight: 900 !important;
 }
 [data-testid="stSidebar"] .stButton > button:hover,
 [data-testid="stSidebar"] button:hover {
   background: linear-gradient(135deg,#1d4ed8,#1e40af) !important;
 }

 /* Alerts in sidebar */
 [data-testid="stSidebar"] .stAlert {
   background: #dbeafe !important;
   border: 1px solid #93c5fd !important;
   border-radius: 16px !important;
 }
 [data-testid="stSidebar"] .stAlert *,
 [data-testid="stSidebar"] .stAlert p,
 [data-testid="stSidebar"] .stAlert div {
   color: #0f172a !important;
 }

 /* Toggle label */
 [data-testid="stSidebar"] [data-testid="stWidgetLabel"] * {
   color: #f8fafc !important;
 }

 /* Top KPI labels: force light */
 div[data-testid="stMetric"],
 div[data-testid="stMetric"] * {
   color: #f8fafc !important;
 }
 div[data-testid="stMetricLabel"],
 div[data-testid="stMetricLabel"] *,
 div[data-testid="stMetricDelta"],
 div[data-testid="stMetricDelta"] * {
   color: #dbeafe !important;
   opacity: 1 !important;
 }
 div[data-testid="stMetricValue"],
 div[data-testid="stMetricValue"] * {
   color: #ffffff !important;
   opacity: 1 !important;
 }

 /* Weekly summary / action boxes: force light cards with dark text */
 .fpi-action-card,
 .fpi-action-card *,
 .fpi-summary-card,
 .fpi-summary-card *,
 .fpi-summary-item,
 .fpi-summary-item * {
   background-color: #f8fafc !important;
   color: #0f172a !important;
   opacity: 1 !important;
   text-shadow: none !important;
 }
 .fpi-action-card {
   background: linear-gradient(135deg,#f8fafc,#e0f2fe) !important;
   border: 1px solid #93c5fd !important;
   box-shadow: 0 18px 42px rgba(15,23,42,.24) !important;
 }
 .fpi-summary-label {
   color: #334155 !important;
 }
 .fpi-summary-value {
   color: #0f172a !important;
 }

 /* Any markdown card that has blue background but dark text */
 .stMarkdown div[style*="background"],
 .stMarkdown div[style*="background"] p,
 .stMarkdown div[style*="background"] span {
   opacity: 1 !important;
 }

 /* Tabs readable */
 button[data-baseweb="tab"],
 button[data-baseweb="tab"] *,
 [role="tab"],
 [role="tab"] * {
   color: #e2e8f0 !important;
   opacity: 1 !important;
   font-weight: 850 !important;
 }
 button[data-baseweb="tab"][aria-selected="true"],
 button[data-baseweb="tab"][aria-selected="true"] *,
 [role="tab"][aria-selected="true"],
 [role="tab"][aria-selected="true"] * {
   color: #ffffff !important;
   font-weight: 950 !important;
 }
 button[data-baseweb="tab"]:hover,
 button[data-baseweb="tab"]:hover *,
 [role="tab"]:hover,
 [role="tab"]:hover * {
   color: #ffffff !important;
 }

 /* Download buttons: strong green CTA */
 .stDownloadButton > button,
 .stDownloadButton > button * {
   background: linear-gradient(135deg,#22c55e,#16a34a) !important;
   color: #ffffff !important;
   border: none !important;
   font-weight: 950 !important;
   opacity: 1 !important;
 }
 .stDownloadButton > button {
   border-radius: 16px !important;
   min-height: 48px !important;
   box-shadow: 0 16px 36px rgba(34,197,94,.28) !important;
 }

 /* Data tables: keep default white readable */
 [data-testid="stDataFrame"] {
   background: #ffffff !important;
   color: #0f172a !important;
 }
 [data-testid="stDataFrame"] * {
   color: #0f172a !important;
 }

 /* Selectboxes / dropdowns in main */
 [data-baseweb="select"] div,
 [data-baseweb="select"] span {
   color: #0f172a !important;
 }


 /* Final direct summary readability override */
 .fpi-readable-summary, .fpi-readable-summary * {
   color: #0f172a !important;
   opacity: 1 !important;
   text-shadow: none !important;
 }
 div[style*="Heti vezetői összefoglaló"],
 div[style*="Heti vezetői összefoglaló"] * {
   color: #0f172a !important;
   opacity: 1 !important;
   text-shadow: none !important;
 }
 /* Expander / Smart mapper readable */
 div[data-testid="stExpander"] {
   background: #ffffff !important;
   border: 1px solid #cbd5e1 !important;
   border-radius: 18px !important;
 }
 div[data-testid="stExpander"],
 div[data-testid="stExpander"] *,
 div[data-testid="stExpander"] p,
 div[data-testid="stExpander"] label,
 div[data-testid="stExpander"] span {
   color: #0f172a !important;
   opacity: 1 !important;
 }
 div[data-testid="stExpander"] input,
 div[data-testid="stExpander"] textarea,
 div[data-testid="stExpander"] [data-baseweb="select"] *,
 div[data-testid="stExpander"] [data-baseweb="input"] * {
   color: #0f172a !important;
   background: #ffffff !important;
 }


 /* ===== Weekly summary final readability override ===== */
 pre, pre * {
   color: #000000 !important;
 }
 div[style*="Heti vezetői összefoglaló"],
 div[style*="Heti vezetői összefoglaló"] *,
 div[style*="border-left:10px solid #2563eb"],
 div[style*="border-left:10px solid #2563eb"] * {
   color: #000000 !important;
   opacity: 1 !important;
   text-shadow: none !important;
 }

 
 /* ===== V7.3 Tactical upload readability patch ===== */
 [data-testid="stFileUploader"] {
     background: #ffffff !important;
     color: #0f172a !important;
     border: 1px solid #cbd5e1 !important;
     border-radius: 16px !important;
     padding: 10px !important;
 }
 [data-testid="stFileUploader"] *,
 [data-testid="stFileUploader"] label,
 [data-testid="stFileUploader"] span,
 [data-testid="stFileUploader"] small,
 [data-testid="stFileUploader"] p,
 [data-testid="stFileUploader"] div {
     color: #0f172a !important;
     opacity: 1 !important;
 }
 [data-testid="stFileUploader"] button,
 [data-testid="stFileUploader"] button * {
     color: #ffffff !important;
     background: linear-gradient(135deg,#2563eb,#1d4ed8) !important;
     border-radius: 12px !important;
     font-weight: 900 !important;
 }
 [data-testid="stFileUploaderDeleteBtn"],
 [data-testid="stFileUploaderDeleteBtn"] *,
 button[title="Delete"],
 button[title="Remove"],
 button[aria-label*="Delete"],
 button[aria-label*="Remove"] {
     color: #ffffff !important;
     background: #dc2626 !important;
     border-radius: 999px !important;
     opacity: 1 !important;
 }

 
 /* ===== V7.4 HARD readability patch for Tactical Pro+ ===== */
 [data-testid="stFileUploader"],
 [data-testid="stFileUploader"] section,
 [data-testid="stFileUploader"] div {
     background: #ffffff !important;
     color: #0f172a !important;
     border-color: #cbd5e1 !important;
 }
 [data-testid="stFileUploader"] label,
 [data-testid="stFileUploader"] span,
 [data-testid="stFileUploader"] small,
 [data-testid="stFileUploader"] p,
 [data-testid="stFileUploader"] svg {
     color: #0f172a !important;
     fill: #0f172a !important;
     opacity: 1 !important;
 }
 [data-testid="stFileUploader"] button {
     background: #2563eb !important;
     color: #ffffff !important;
     border-radius: 12px !important;
     border: 1px solid #1d4ed8 !important;
 }
 [data-testid="stFileUploader"] button *,
 [data-testid="stFileUploader"] button svg {
     color: #ffffff !important;
     fill: #ffffff !important;
 }
 [data-testid="stFileUploaderDeleteBtn"],
 [data-testid="stFileUploaderDeleteBtn"] *,
 button[title*="Delete"],
 button[aria-label*="Delete"],
 button[title*="Remove"],
 button[aria-label*="Remove"] {
     background: #dc2626 !important;
     color: #ffffff !important;
     fill: #ffffff !important;
     opacity: 1 !important;
 }
 div[data-testid="stExpander"],
 div[data-testid="stExpander"] *,
 div[data-testid="stExpander"] label,
 div[data-testid="stExpander"] p,
 div[data-testid="stExpander"] span {
     color: #0f172a !important;
 }
 div[data-testid="stExpander"] {
     background: #ffffff !important;
 }
 [data-baseweb="select"] *,
 [data-baseweb="popover"] *,
 [data-baseweb="menu"] * {
     color: #0f172a !important;
     background: #ffffff !important;
 }
 .tactical-readable-box {
     background:#ffffff;
     color:#0f172a;
     border:1px solid #cbd5e1;
     border-radius:18px;
     padding:16px 18px;
     margin:10px 0;
     box-shadow:0 8px 24px rgba(15,23,42,.10);
 }
 .tactical-readable-box * { color:#0f172a !important; }

 </style>
    """,
    unsafe_allow_html=True,
)




# -----------------------------------------------------------------------------
# V115 - Light premium UI + readability + no duplicate imports
# -----------------------------------------------------------------------------
def _fpi_apply_v115_light_ui_patch() -> None:
    """Világos, prémium Football Performance Intelligence (FPI) arculat + erős kontrasztjavítás Streamlit elemekre."""
    st.markdown(
        """
        <style>
        :root {
            --fpi-bg: #eef5f3;
            --fpi-bg-2: #f8fbff;
            --fpi-panel: #ffffff;
            --fpi-ink: #0f172a;
            --fpi-muted: #475569;
            --fpi-line: #d8e3ea;
            --fpi-green: #0f766e;
            --fpi-blue: #2563eb;
            --fpi-cyan: #0891b2;
            --fpi-orange: #f97316;
            --fpi-purple: #7c3aed;
        }
        .stApp {
            background:
                radial-gradient(circle at 8% 4%, rgba(20,184,166,.18), transparent 30%),
                radial-gradient(circle at 88% 2%, rgba(37,99,235,.14), transparent 30%),
                linear-gradient(135deg, #edf7f4 0%, #f8fbff 46%, #eef2ff 100%) !important;
            color: var(--fpi-ink) !important;
        }
        .block-container { padding-top: 1.2rem; }
        .stApp h1, .stApp h2, .stApp h3, .stApp h4,
        .stApp p, .stApp label, .stApp span, .stApp div {
            color: var(--fpi-ink) !important;
            text-shadow: none !important;
        }
        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, #ffffff 0%, #f1f8f6 100%) !important;
            border-right: 1px solid var(--fpi-line) !important;
        }
        [data-testid="stSidebar"] * { color: var(--fpi-ink) !important; }
        [data-testid="stSidebar"] .stButton > button,
        [data-testid="stSidebar"] button {
            background: linear-gradient(135deg, #0f766e, #2563eb) !important;
            color:#ffffff !important;
            border:0 !important;
        }
        [data-testid="stSidebar"] .stButton > button *,
        [data-testid="stSidebar"] button * { color:#ffffff !important; }

        /* Cards and hero: komoly, de nem fekete */
        .hero-box, .fpi-hero-wow, .fpi-dark-card, .intro-card, .feature-box,
        .score-card, .priority-card, .insight-card, .export-panel,
        .premium-kpi, .fpi-mode-card, .fpi-flow-step, .fpi-decision-card {
            background: linear-gradient(135deg, rgba(255,255,255,.98), rgba(239,247,255,.96)) !important;
            border: 1px solid var(--fpi-line) !important;
            color: var(--fpi-ink) !important;
            box-shadow: 0 14px 38px rgba(15,23,42,.10) !important;
        }
        .hero-title, .fpi-hero-wow h1, .fpi-landing-title,
        .insight-title, .feature-title, .score-number, .score-label,
        .premium-kpi-label, .premium-kpi-value, .premium-kpi-note,
        .hero-sub, .feature-text, .mini-muted, .fpi-hero-wow p,
        .fpi-mode-card h3, .fpi-mode-card p, .fpi-flow-step b, .fpi-flow-step div,
        .fpi-decision-card b, .fpi-decision-card span {
            color: var(--fpi-ink) !important;
        }
        .fpi-landing-hero {
            background: linear-gradient(135deg, #ffffff 0%, #e0f2fe 55%, #ecfdf5 100%) !important;
            border: 1px solid var(--fpi-line) !important;
            box-shadow: 0 20px 54px rgba(15,23,42,.13) !important;
        }
        .fpi-landing-kicker { color: var(--fpi-green) !important; }
        .fpi-landing-sub { color: var(--fpi-muted) !important; }
        .section-chip, .fpi-chip-wow, .micro-pill {
            background: #e0f2fe !important;
            border: 1px solid #bae6fd !important;
            color: #075985 !important;
        }
        .pill-critical { background:#fee2e2 !important; color:#991b1b !important; }
        .pill-warning { background:#ffedd5 !important; color:#9a3412 !important; }
        .pill-info { background:#dbeafe !important; color:#1e40af !important; }

        /* Dropdown/readability hard fix */
        [data-baseweb="select"], [data-baseweb="select"] *,
        [data-baseweb="popover"], [data-baseweb="popover"] *,
        [data-baseweb="menu"], [data-baseweb="menu"] *,
        [role="listbox"], [role="listbox"] *,
        [role="option"], [role="option"] * {
            background-color: #ffffff !important;
            color: #0f172a !important;
            opacity: 1 !important;
        }
        [role="option"]:hover, [role="option"][aria-selected="true"] {
            background-color: #dbeafe !important;
            color: #0f172a !important;
        }
        input, textarea, [data-baseweb="input"], [data-baseweb="input"] *,
        [data-baseweb="textarea"], [data-baseweb="textarea"] * {
            background: #ffffff !important;
            color: #0f172a !important;
            border-color: #cbd5e1 !important;
        }

        /* Dashboard / metodología / intelligence olvashatóság */
        .stInfo, .stInfo *, .stAlert, .stAlert * {
            color: #0f172a !important;
        }
        div[data-testid="stMetric"], div[data-testid="stMetric"] * {
            color: #0f172a !important;
        }
        .wrap-table th { background: #dbeafe !important; color: #0f172a !important; }
        .wrap-table td, .wrap-table tr:nth-child(even) td {
            background: #ffffff !important;
            color: #0f172a !important;
            border-color: #dbeafe !important;
        }
        pre, pre *, code, code * { color:#0f172a !important; background:#f8fafc !important; }
        div[style*="Heti vezetői összefoglaló"],
        div[style*="Heti vezetői összefoglaló"] *,
        div[style*="border-left:10px solid #2563eb"],
        div[style*="border-left:10px solid #2563eb"] *,
        div[style*="background:#0f172a"],
        div[style*="background: #0f172a"],
        div[style*="background:#111827"],
        div[style*="background: #111827"],
        div[style*="background:#020617"],
        div[style*="background: #020617"] {
            background: #ffffff !important;
            color: #0f172a !important;
            border-color: #dbeafe !important;
            opacity: 1 !important;
        }
        div[data-testid="stExpander"] {
            background:#ffffff !important;
            border:1px solid var(--fpi-line) !important;
            border-radius:18px !important;
        }
        div[data-testid="stExpander"] * { color:#0f172a !important; }
        [data-testid="stFileUploader"], [data-testid="stFileUploader"] section,
        [data-testid="stFileUploader"] div {
            background:#ffffff !important;
            color:#0f172a !important;
            border-color:#cbd5e1 !important;
        }
        [data-testid="stFileUploader"] * { color:#0f172a !important; fill:#0f172a !important; }
        [data-testid="stFileUploader"] button, [data-testid="stFileUploader"] button * {
            background:#2563eb !important;
            color:#ffffff !important;
            fill:#ffffff !important;
        }
        .stDownloadButton > button, .stButton > button {
            border-radius: 15px !important;
            font-weight: 900 !important;
        }
        .stDownloadButton > button {
            background: linear-gradient(135deg,#0f766e,#14b8a6) !important;
            color:#ffffff !important;
            border:0 !important;
        }
        .stDownloadButton > button * { color:#ffffff !important; }
        button[data-baseweb="tab"], button[data-baseweb="tab"] * {
            color:#334155 !important;
            background: transparent !important;
        }
        button[data-baseweb="tab"][aria-selected="true"], button[data-baseweb="tab"][aria-selected="true"] * {
            color:#0f766e !important;
            font-weight:950 !important;
        }
        div[data-baseweb="tab-list"] {
            background: rgba(255,255,255,.88) !important;
            border: 1px solid var(--fpi-line) !important;
        }
        div[data-baseweb="tab-border"] { background-color:#0f766e !important; }

        /* V116: BaseWeb legördülők + léptetők végső olvashatósági javítása */
        div[data-baseweb="popover"], div[data-baseweb="popover"] *,
        div[data-baseweb="menu"], div[data-baseweb="menu"] *,
        ul[role="listbox"], ul[role="listbox"] *,
        li[role="option"], li[role="option"] *,
        div[role="option"], div[role="option"] * {
            background: #ffffff !important;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            opacity: 1 !important;
        }
        li[role="option"]:hover, div[role="option"]:hover,
        li[role="option"][aria-selected="true"], div[role="option"][aria-selected="true"] {
            background: #e0f2fe !important;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
        }
        div[data-baseweb="select"] > div,
        div[data-baseweb="select"] input,
        div[data-baseweb="select"] span,
        div[data-baseweb="select"] svg,
        div[data-baseweb="input"] > div,
        div[data-baseweb="input"] input {
            background: #ffffff !important;
            color: #0f172a !important;
            fill: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            opacity: 1 !important;
        }
        div[data-baseweb="input"] button,
        div[data-baseweb="input"] button *,
        button[aria-label="Increment"], button[aria-label="Decrement"],
        button[aria-label="Növelés"], button[aria-label="Csökkentés"] {
            background: #e2e8f0 !important;
            color: #0f172a !important;
            fill: #0f172a !important;
            border-color: #cbd5e1 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

_fpi_apply_v115_light_ui_patch()

# =========================================================
# V117 - Final readable UI controls patch
# =========================================================
def _fpi_apply_v117_control_readability_patch() -> None:
    """Végső kontrasztjavítás: minden gomb, lenyíló, léptető, popover olvasható.
    A korábbi sötét gomb/CSS szabályokat ez felülírja.
    """
    st.markdown(
        """
        <style>
        :root {
            --fpi-control-bg: #ffffff;
            --fpi-control-text: #0f172a;
            --fpi-control-border: #cbd5e1;
            --fpi-primary: #0f766e;
            --fpi-primary-2: #2563eb;
            --fpi-soft-blue: #e0f2fe;
            --fpi-soft-green: #ecfdf5;
        }

        /* A Streamlit/BaseWeb gyakran inline style-t ad a select és number_input elemeknek.
           Ezért több szelektorral, végső sorrendben írjuk felül. */
        div[data-baseweb="select"],
        div[data-baseweb="select"] > div,
        div[data-baseweb="select"] div,
        div[data-baseweb="select"] span,
        div[data-baseweb="select"] input,
        div[data-baseweb="select"] svg,
        div[data-baseweb="input"],
        div[data-baseweb="input"] > div,
        div[data-baseweb="input"] div,
        div[data-baseweb="input"] input,
        [data-testid="stNumberInput"] div,
        [data-testid="stNumberInput"] input,
        [data-testid="stTextInput"] div,
        [data-testid="stTextInput"] input,
        [data-testid="stDateInput"] div,
        [data-testid="stDateInput"] input {
            background-color: var(--fpi-control-bg) !important;
            color: var(--fpi-control-text) !important;
            -webkit-text-fill-color: var(--fpi-control-text) !important;
            fill: var(--fpi-control-text) !important;
            opacity: 1 !important;
            border-color: var(--fpi-control-border) !important;
        }

        /* Lenyitott dropdown menü */
        div[data-baseweb="popover"],
        div[data-baseweb="popover"] div,
        div[data-baseweb="popover"] span,
        div[data-baseweb="popover"] ul,
        div[data-baseweb="popover"] li,
        div[data-baseweb="menu"],
        div[data-baseweb="menu"] div,
        div[data-baseweb="menu"] span,
        ul[role="listbox"],
        ul[role="listbox"] *,
        li[role="option"],
        li[role="option"] *,
        div[role="option"],
        div[role="option"] * {
            background-color: #ffffff !important;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            fill: #0f172a !important;
            opacity: 1 !important;
        }
        li[role="option"]:hover,
        div[role="option"]:hover,
        li[role="option"][aria-selected="true"],
        div[role="option"][aria-selected="true"] {
            background-color: #dbeafe !important;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
        }

        /* Number input léptető gombok */
        [data-testid="stNumberInput"] button,
        [data-testid="stNumberInput"] button *,
        div[data-baseweb="input"] button,
        div[data-baseweb="input"] button *,
        button[aria-label="Increment"], button[aria-label="Increment"] *,
        button[aria-label="Decrement"], button[aria-label="Decrement"] *,
        button[aria-label="Növelés"], button[aria-label="Növelés"] *,
        button[aria-label="Csökkentés"], button[aria-label="Csökkentés"] * {
            background: #e2e8f0 !important;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            fill: #0f172a !important;
            border-color: #cbd5e1 !important;
            opacity: 1 !important;
        }

        /* Általános gombok: ne maradjanak fekete, szöveg nélküli téglalapok. */
        .stButton > button,
        .stButton > button *,
        button[kind="secondary"],
        button[kind="secondary"] *,
        button[data-testid="baseButton-secondary"],
        button[data-testid="baseButton-secondary"] * {
            background: linear-gradient(135deg, var(--fpi-primary), var(--fpi-primary-2)) !important;
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
            fill: #ffffff !important;
            border: 0 !important;
            border-radius: 16px !important;
            font-weight: 900 !important;
            opacity: 1 !important;
        }
        .stButton > button:hover,
        button[data-testid="baseButton-secondary"]:hover {
            filter: brightness(1.05) !important;
            transform: translateY(-1px);
        }

        /* File uploader gomb marad olvasható. */
        [data-testid="stFileUploader"] button,
        [data-testid="stFileUploader"] button * {
            background: linear-gradient(135deg, #2563eb, #0f766e) !important;
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
            fill: #ffffff !important;
        }

        /* Sötét maradék panelek kényszerített világosítása. */
        div[style*="background:#0f172a"], div[style*="background: #0f172a"],
        div[style*="background:#111827"], div[style*="background: #111827"],
        div[style*="background:#020617"], div[style*="background: #020617"],
        div[style*="background:rgba(15,23,42"], div[style*="background: rgba(15,23,42"],
        div[style*="background:rgba(17,24,39"], div[style*="background: rgba(17,24,39"] {
            background: linear-gradient(135deg, #ffffff, #f0f9ff) !important;
            color: #0f172a !important;
            border-color: #dbeafe !important;
        }
        div[style*="background:#0f172a"] *, div[style*="background: #0f172a"] *,
        div[style*="background:#111827"] *, div[style*="background: #111827"] *,
        div[style*="background:#020617"] *, div[style*="background: #020617"] *,
        div[style*="background:rgba(15,23,42"] *, div[style*="background: rgba(15,23,42"] *,
        div[style*="background:rgba(17,24,39"] *, div[style*="background: rgba(17,24,39"] * {
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
            opacity: 1 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

_fpi_apply_v117_control_readability_patch()


# =========================================================
# V131 - Methodology static table readability fix
# =========================================================
def _fpi_apply_v131_methodology_table_css() -> None:
    st.markdown(
        """
        <style>
        .fpi-method-table-wrap {
            width: 100%;
            overflow-x: auto;
            margin: 10px 0 18px 0;
            border-radius: 18px;
            border: 1px solid #cbd5e1;
            background: #ffffff;
            box-shadow: 0 10px 28px rgba(15,23,42,.08);
        }
        .fpi-method-table {
            width: 100%;
            border-collapse: collapse;
            table-layout: auto;
            background: #ffffff;
            color: #0f172a;
        }
        .fpi-method-table th {
            background: #e0f2fe;
            color: #0f172a;
            text-align: left;
            font-weight: 900;
            padding: 12px 14px;
            border-bottom: 1px solid #cbd5e1;
            white-space: nowrap;
        }
        .fpi-method-table td {
            color: #0f172a;
            padding: 11px 14px;
            border-bottom: 1px solid #e2e8f0;
            vertical-align: top;
            line-height: 1.4;
            background: #ffffff;
        }
        .fpi-method-table tr:nth-child(even) td { background: #f8fafc; }
        .fpi-method-table tr:last-child td { border-bottom: none; }
        .fpi-method-table-wrap *, .fpi-method-table * {
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

_fpi_apply_v131_methodology_table_css()

# =========================================================
# V119 - Nuclear light UI override: no black controls/tables
# =========================================================
def _fpi_apply_v119_all_light_readable_patch() -> None:
    """Minden maradék fekete/sötét Streamlit, BaseWeb és inline panel világosítása.
    Szándékosan a korábbi UI patchek után fut.
    """
    st.markdown(
        """
        <style>
        :root {
            --fpi-bg-main:#f6f9fc;
            --fpi-surface:#ffffff;
            --fpi-surface-2:#f1f5f9;
            --fpi-surface-3:#eaf3ff;
            --fpi-text:#0f172a;
            --fpi-text-muted:#334155;
            --fpi-border:#cbd5e1;
            --fpi-accent:#0f766e;
            --fpi-accent-2:#2563eb;
        }

        html, body, .stApp, [data-testid="stAppViewContainer"], .main {
            background: linear-gradient(135deg, #f8fafc 0%, #edf6ff 48%, #eefbf5 100%) !important;
            color: var(--fpi-text) !important;
        }

        /* Alapszöveg: mindenhol sötét, világos felületen */
        .stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp h5, .stApp h6,
        .stApp p, .stApp label, .stApp span, .stApp div,
        [data-testid="stWidgetLabel"], [data-testid="stWidgetLabel"] * {
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            text-shadow: none !important;
            opacity: 1 !important;
        }

        /* Minden régi sötét card/panel világos legyen */
        .hero-box, .fpi-hero-wow, .fpi-dark-card, .intro-card, .feature-box,
        .score-card, .priority-card, .insight-card, .export-panel, .premium-kpi,
        .fpi-mode-card, .fpi-flow-step, .fpi-decision-card, .fpi-glass,
        .fpi-clean-card, .fpi-summary-card, .fpi-action-card,
        div[class*="dark"], div[class*="card"], div[class*="panel"] {
            background: linear-gradient(135deg, #ffffff 0%, #f1f7ff 100%) !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            border-color: var(--fpi-border) !important;
            box-shadow: 0 10px 28px rgba(15,23,42,.08) !important;
        }
        .hero-box *, .fpi-hero-wow *, .fpi-dark-card *, .intro-card *, .feature-box *,
        .score-card *, .priority-card *, .insight-card *, .export-panel *, .premium-kpi *,
        .fpi-mode-card *, .fpi-flow-step *, .fpi-decision-card *, .fpi-glass *,
        .fpi-clean-card *, .fpi-summary-card *, .fpi-action-card * {
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            fill: var(--fpi-text) !important;
        }

        /* Inline style-ból érkező fekete/sötét hátterek felülírása */
        div[style*="background:#0"], div[style*="background: #0"],
        div[style*="background:#1"], div[style*="background: #1"],
        div[style*="background:#2"], div[style*="background: #2"],
        div[style*="background-color:#0"], div[style*="background-color: #0"],
        div[style*="background-color:#1"], div[style*="background-color: #1"],
        div[style*="background-color:#2"], div[style*="background-color: #2"],
        div[style*="rgba(15,23,42"], div[style*="rgba(15, 23, 42"],
        div[style*="rgba(17,24,39"], div[style*="rgba(17, 24, 39"],
        div[style*="rgba(2,6,23"], div[style*="rgba(2, 6, 23"] {
            background: linear-gradient(135deg, #ffffff 0%, #eef6ff 100%) !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            border-color: var(--fpi-border) !important;
        }
        div[style*="background:#0"] *, div[style*="background: #0"] *,
        div[style*="background:#1"] *, div[style*="background: #1"] *,
        div[style*="background:#2"] *, div[style*="background: #2"] *,
        div[style*="background-color:#0"] *, div[style*="background-color: #0"] *,
        div[style*="background-color:#1"] *, div[style*="background-color: #1"] *,
        div[style*="background-color:#2"] *, div[style*="background-color: #2"] *,
        div[style*="rgba(15,23,42"] *, div[style*="rgba(15, 23, 42"] *,
        div[style*="rgba(17,24,39"] *, div[style*="rgba(17, 24, 39"] *,
        div[style*="rgba(2,6,23"] *, div[style*="rgba(2, 6, 23"] * {
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            fill: var(--fpi-text) !important;
        }

        /* Selectbox, multiselect, dropdown, nyitott lista */
        [data-baseweb="select"], [data-baseweb="select"] *,
        [data-baseweb="popover"], [data-baseweb="popover"] *,
        [data-baseweb="menu"], [data-baseweb="menu"] *,
        [data-baseweb="option"], [data-baseweb="option"] *,
        [role="listbox"], [role="listbox"] *,
        [role="option"], [role="option"] *,
        [role="combobox"], [role="combobox"] * {
            background-color: #ffffff !important;
            background-image: none !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            fill: var(--fpi-text) !important;
            border-color: var(--fpi-border) !important;
            opacity: 1 !important;
        }
        [data-baseweb="popover"], [data-baseweb="menu"], [role="listbox"] {
            border: 1px solid var(--fpi-border) !important;
            box-shadow: 0 18px 40px rgba(15,23,42,.16) !important;
            border-radius: 14px !important;
        }
        [role="option"]:hover, [role="option"][aria-selected="true"],
        [data-baseweb="option"]:hover {
            background-color: #dbeafe !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
        }

        /* Inputok, dátum, szám, léptető */
        input, textarea,
        [data-baseweb="input"], [data-baseweb="input"] *,
        [data-baseweb="textarea"], [data-baseweb="textarea"] *,
        [data-testid="stTextInput"], [data-testid="stTextInput"] *,
        [data-testid="stDateInput"], [data-testid="stDateInput"] *,
        [data-testid="stNumberInput"], [data-testid="stNumberInput"] *,
        [data-testid="stTimeInput"], [data-testid="stTimeInput"] * {
            background-color: #ffffff !important;
            background-image: none !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            fill: var(--fpi-text) !important;
            border-color: var(--fpi-border) !important;
            opacity: 1 !important;
        }
        [data-testid="stNumberInput"] button, [data-testid="stNumberInput"] button *,
        button[aria-label*="Increment"], button[aria-label*="Increment"] *,
        button[aria-label*="Decrement"], button[aria-label*="Decrement"] *,
        button[aria-label*="Növel"], button[aria-label*="Növel"] *,
        button[aria-label*="Csökk"], button[aria-label*="Csökk"] * {
            background: #e2e8f0 !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            fill: var(--fpi-text) !important;
            border: 1px solid var(--fpi-border) !important;
        }

        /* Gombok: sötét helyett szürke/kékes, olvasható betűvel */
        .stButton > button, .stButton > button *,
        button[data-testid^="baseButton"], button[data-testid^="baseButton"] *,
        button[kind], button[kind] * {
            background: linear-gradient(135deg, #e2e8f0, #dbeafe) !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            fill: var(--fpi-text) !important;
            border: 1px solid #b6c8dc !important;
            border-radius: 16px !important;
            font-weight: 900 !important;
            opacity: 1 !important;
            box-shadow: 0 8px 18px rgba(15,23,42,.08) !important;
        }
        .stButton > button:hover, button[data-testid^="baseButton"]:hover, button[kind]:hover {
            background: linear-gradient(135deg, #cbd5e1, #bfdbfe) !important;
            color: var(--fpi-text) !important;
        }
        .stDownloadButton > button, .stDownloadButton > button * {
            background: linear-gradient(135deg, #0f766e, #2563eb) !important;
            color: #ffffff !important;
            -webkit-text-fill-color: #ffffff !important;
            fill: #ffffff !important;
            border: 0 !important;
        }

        /* File uploader, Smart Mapper, Expander */
        [data-testid="stFileUploader"], [data-testid="stFileUploader"] *,
        [data-testid="stExpander"], [data-testid="stExpander"] *,
        details[data-testid="stExpander"], details[data-testid="stExpander"] * {
            background: #ffffff !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            fill: var(--fpi-text) !important;
            border-color: var(--fpi-border) !important;
            opacity: 1 !important;
        }
        [data-testid="stFileUploader"], [data-testid="stExpander"], details[data-testid="stExpander"] {
            border: 1px solid var(--fpi-border) !important;
            border-radius: 18px !important;
            box-shadow: 0 8px 22px rgba(15,23,42,.07) !important;
        }

        /* Táblázatok: ne legyen fekete háttér */
        [data-testid="stDataFrame"], [data-testid="stDataFrame"] *,
        .stDataFrame, .stDataFrame *,
        table, thead, tbody, tr, th, td,
        .wrap-table table, .wrap-table th, .wrap-table td {
            background-color: #ffffff !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            border-color: #dbe3ea !important;
            opacity: 1 !important;
        }
        thead, th, .wrap-table th {
            background-color: #e0f2fe !important;
            color: #0f172a !important;
            font-weight: 900 !important;
        }
        tbody tr:nth-child(even) td, .wrap-table tr:nth-child(even) td {
            background-color: #f8fafc !important;
        }

        /* Alert, info, markdown code */
        [data-testid="stAlert"], [data-testid="stAlert"] *,
        pre, pre *, code, code * {
            background: #f8fafc !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
            border-color: var(--fpi-border) !important;
        }

        /* Sidebar is legyen világos */
        [data-testid="stSidebar"], [data-testid="stSidebar"] * {
            background-color: #f8fafc !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
        }
        [data-testid="stSidebar"] .stButton > button,
        [data-testid="stSidebar"] .stButton > button * {
            background: linear-gradient(135deg, #e2e8f0, #dbeafe) !important;
            color: var(--fpi-text) !important;
            -webkit-text-fill-color: var(--fpi-text) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

_fpi_apply_v119_all_light_readable_patch()


# -----------------------------------------------------------------------------
# Oszlopmapping
# -----------------------------------------------------------------------------
STANDARD_COLUMNS = {
    "player_name": ["Játékos neve", "Player", "Player Name", "Name", "name", "Név", "Nev", "Játékos", "Jatekos", "Athlete", "Athlete Name", "athlete_name", "Player full name", "Full Name"],
    "session_type": ["Típus", "Type", "Session Type", "Edzés/Meccs", "SessionType", "Activity Type", "Drill Type", "Event Type", "Training/Match", "Tags", "day_name"],
    "session_name": ["Szakasz neve", "Session", "Session Name", "Activity", "Drill", "Exercise", "Event", "Session title", "Session Title", "activity_name", "Split Name"],
    "position": ["Poszt", "Position", "Player Position", "Role", "Playing Position", "Post", "Pos", "position_name"],
    "start_time": ["Kezdési idő", "Start Time", "Start", "Dátum", "Date", "Session Date", "Day", "Datum", "Kezdés", "Start date", "StartTime", "Split", "date", "start_time", "Split Start Time"],
    "end_time": ["Befejezési idő", "End Time", "End", "Finish", "Befejezés", "EndTime"],
    "duration": ["Időtartam", "Duration", "Time", "Minutes", "Idő", "Időtartam [perc]", "Duration [min]", "Duration min", "Duration (mm:ss)", "total_duration"],
    "match_minutes": ["Játékperc", "Játékpercek", "Meccsperc", "Meccspercek", "Minutes played", "Minutes Played", "Playing Time", "Match Minutes", "Match minutes", "Player minutes", "On pitch minutes"],
    "total_distance": ["Teljes táv [m]", "Tel\xadjes táv [m]", "Total Distance", "Distance", "Össztáv", "Total distance (m)", "Total Dist", "Dist Total", "Distance [m]", "TD", "Total Distance m", "total_distance", "Distance (km)"],
    "distance_per_min": ["Táv/perc [m/min]", "Distance/min", "Distance Per Min", "m/min", "Distance per minute", "m per min", "m/minute", "Rel Distance", "Distance Per Min (m/min)", "meterage_per_minute"],
    "max_speed": ["Maximális sebesség [km/h]", "Max Speed", "Maximum Speed", "Top Speed", "Peak Speed", "Max Velocity", "Vmax", "Top Speed (km/h)", "max_vel"],
    "avg_speed": ["Átlagsebesség [km/h]", "Average Speed", "Avg Speed", "Mean Speed"],
    "sprints": ["Sprintek", "Sprints", "Sprint Count", "Number of Sprints", "Sprint #", "Sprint efforts", "Sprints count  ()", "Sprints count"],
    "speed_zone_3": ["Táv a sebesség célzónában 3 [m] (14.40 - 19.79 km/h)"],
    "speed_zone_4": ["Táv a sebesség célzónában 4 [m] (19.80 - 24.99 km/h)", "Distance(4+5)  (m)", "Distance(4+5)", "Distance 4+5", "HSR Distance", "velocity2_band3_total_distance"],
    "speed_zone_5": ["Táv a sebesség célzónában 5 [m] (25.00- km/h)", "Total sprints distance  (m)", "Total sprints distance", "Sprint distance", "Sprint Distance", "Sprint Distance (m)", "velocity2_band4_total_distance"],
    "training_load": ["Edzési terhelési pontérték", "Terhelési pont", "Player Load", "Load", "Training Load", "Total Load", "Workload", "Load Score", "total_player_load", "HMLD"],
    "cardio_load": ["Kardióterhelés", "Cardio Load"],
    "recovery_hours": ["Regenerálódási idő [h]", "Recovery Time", "Recovery"],
    "muscle_load": ["Izomterhelés", "Muscle Load", "Muscular Load", "Mechanical Load"],
    "hr_avg": ["Átlagos pulzus [bpm]", "Average HR", "Avg HR", "Mean HR", "HR avg", "Avg Heart Rate"],
    "hr_max": ["Maximális pulzus [bpm]", "Max HR", "Maximum HR", "Peak HR", "Max Heart Rate"],
    "hrv": ["HRV (RMSSD)", "HRV", "RMSSD", "HRV RMSSD"],
    "acc_low": ["Gyorsulások száma (2.00 - 2.49 m/s²)"],
    "acc_mid": ["Gyorsulások száma (2.50 - 2.99 m/s²)"],
    "acc_high": ["Gyorsulások száma (3.00 - 50.00 m/s²)", "Total Accelerations  ()", "Total Accelerations", "Accelerations (2+3)  ()", "Accelerations (2+3)", "Accelerations", "gen2_acceleration_band3plus_total_effort_count"],
    "dec_low": ["Gyorsulások száma (-2.49 - -2.00 m/s²)"],
    "dec_mid": ["Gyorsulások száma (-2.99 - -2.50 m/s²)"],
    "dec_high": ["Gyorsulások száma (-50.00 - -3.00 m/s²)", "Total Decelerations  ()", "Total Decelerations", "Decelerations (2+3)  ()", "Decelerations (2+3)", "Decelerations"],
    "high_efforts": ["High Efforts", "High Effort", "High efforts", "Nagy intenzitású erőfeszítések", "Nagy intenzitású akciók", "Explosive Efforts", "Explosive efforts", "High Intensity Efforts", "HIE", "Efforts High", "HI Efforts"],
}

CORE_REQUIRED = ["player_name", "session_type", "start_time"]
SEVERITY_RANK = {"KRITIKUS": 0, "FIGYELMEZTETÉS": 1, "INFORMÁCIÓ": 2}

METRIC_LABELS = {
    "training_load": "Terhelési pont",
    "total_distance": "Össztáv",
    "sprint_distance": "Sprinttáv",
    "hsr_distance": "Nagy sebességű táv",
    "distance_per_min": "Táv/perc",
    "max_speed": "Maximális sebesség",
    "dec_count": "Lassítások",
    "acc_count": "Gyorsulások",
    "high_efforts": "Nagy intenzitású erőfeszítések",
}

PLAYSTYLE_OPTIONS = {
    "Kiegyensúlyozott": "Általános, kiegyensúlyozott performance profil.",
    "Dominancia": "Labdabirtoklásra és kontrollra építő modell: stabil volumen, kontrollált sebesség- és sprintinger.",
    "Magas presszing": "Magas intenzitás, sok gyorsulás/lassítás, sok High Effort és erős munkasűrűség.",
    "Átmeneti játék": "Gyors átmenetek, magas sprint- és maximális sebességű inger igény.",
    "Direkt játék": "Mélységi futások, sprint- és HSR-expozíció, gyors előrehaladás.",
    "Pressing": "Régi kompatibilitási név: magas presszing profil.",
    "Transition": "Régi kompatibilitási név: átmeneti játék profil.",
    "Possession": "Régi kompatibilitási név: dominancia/possession profil.",
    "Low Block": "Rövidebb, robbanékony intenzív blokkok, kontrollált összterhelés.",
}



# -----------------------------------------------------------------------------
# Smart Mapper v2 - magyarázat és összevont GPS mezők
# -----------------------------------------------------------------------------
MAPPER_FIELD_INFO = {
    "player_name": ("Játékos neve", "A játékos azonosítója / neve.", "szöveg", "minden játékos szintű elemzés, risk motor", True),
    "session_type": ("Típus", "Az esemény típusa: Edzés vagy Meccs.", "szöveg", "edzés-meccs összevetés, readiness", True),
    "start_time": ("Kezdési idő / dátum – dátum/idő", "Az edzés vagy mérkőzés dátuma.", "dátum/idő", "heti bontás, mikrociklus, trendek", True),
    "duration": ("Időtartam", "Az esemény hossza.", "perc vagy hh:mm:ss", "táv/perc, intenzitás", False),
    "match_minutes": ("Játékperc / szereplési idő", "Meccsen ténylegesen pályán töltött perc. Ha nincs külön oszlop, az app meccsnél az időtartamot használja.", "perc vagy hh:mm:ss", "per90 normalizálás, meccs-edzés összevetés", False),
    "total_distance": ("Teljes táv", "Összes megtett távolság.", "méter", "terhelési volumen, heti load, risk", False),
    "distance_per_min": ("Táv/perc", "Relatív futóteljesítmény.", "m/perc", "meccsintenzitás, játékmodell", False),
    "max_speed": ("Maximális sebesség", "Legnagyobb elért sebesség.", "km/h", "max speed trend, frissesség", False),
    "sprints": ("Sprintek száma", "Sprint akciók darabszáma, nem méter.", "darab", "sprint expozíció", False),
    "speed_zone_4": ("Nagy sebességű futás / Zone 4 vagy 4+5", "Nagy sebességű távolság. Ha csak Distance(4+5) van, ide válaszd.", "méter", "HSR, játékmodell, load profil", False),
    "speed_zone_5": ("Sprint táv / Zone 5", "Sprint zónában megtett távolság. Ez méter, nem darabszám.", "méter", "sprintterhelés, sprint fit", False),
    "training_load": ("Edzési terhelési pont", "GPS/rendszer által számolt load pont.", "pont", "heti terhelés, risk", False),
    "muscle_load": ("Izomterhelés", "Mechanikus/izomterhelési mutató.", "pont", "neuromuszkuláris kockázat", False),
    "hr_avg": ("Átlagpulzus", "Átlagos pulzus.", "bpm", "belső terhelés", False),
    "hr_max": ("Max pulzus", "Maximális pulzus.", "bpm", "belső terhelés", False),
    "acc_high": ("Gyorsulások", "Nagy intenzitású vagy összesített gyorsulások száma.", "darab", "neuromuszkuláris terhelés", False),
    "dec_high": ("Lassítások", "Nagy intenzitású vagy összesített lassítások száma.", "darab", "excentrikus terhelés, risk", False),
    "high_efforts": ("High Efforts", "Nagy intenzitású akciók összesített mutatója.", "darab / pont", "pressing/transition profil", False),
}


def mapper_label(std_col: str) -> str:
    return MAPPER_FIELD_INFO.get(std_col, (std_col, "", "", "", False))[0]


def mapper_unit(std_col: str) -> str:
    return MAPPER_FIELD_INFO.get(std_col, ("", "", "", "", False))[2]


def mapper_desc(std_col: str) -> str:
    label, meaning, unit, used_for, required = MAPPER_FIELD_INFO.get(std_col, (std_col, "", "", "", False))
    extra = ""
    if std_col == "speed_zone_4":
        extra = " Összevont 4+5 zóna esetén ezt válaszd ide; az app HSR-ként kezeli."
    elif std_col == "speed_zone_5":
        extra = " Ide ne darabszámot válassz, hanem méter alapú sprinttávot."
    elif std_col == "sprints":
        extra = " Ide darabszám kell, nem sprinttáv méterben."
    elif std_col in ["acc_high", "dec_high"]:
        extra = " Ha csak összesített oszlop van, az is használható."
    return f"{meaning} Várt egység: {unit}. Ebből számolódik: {used_for}.{extra}"


def mapper_warning(std_col: str, source_col: object) -> str:
    s = _norm_mapping_text(source_col)
    if not s:
        return ""
    if std_col == "sprints" and any(x in s for x in ["distance", "tav", "meter", "metre"]):
        return "⚠ Sprintekhez darabszám kell, ez inkább sprinttávnak tűnik."
    if std_col == "speed_zone_5" and any(x in s for x in ["count", "darab"]):
        return "⚠ Sprint távhoz méter kell, ez inkább sprint darabszámnak tűnik."
    if std_col in ["speed_zone_4", "speed_zone_5", "total_distance"] and any(x in s for x in ["count", "darab"]):
        return "⚠ Ehhez távolság kellene, nem darabszám."
    if std_col in ["acc_high", "dec_high", "sprints"] and any(x in s for x in ["distance", "meter", "metre"]):
        return "⚠ Ehhez darabszám jellegű oszlop kellene, nem méter."
    return ""


def enhanced_mapping_quality_df(raw_df: pd.DataFrame, mapping: Dict[str, Optional[str]]) -> pd.DataFrame:
    rows = []
    for std_col, aliases in STANDARD_COLUMNS.items():
        src = mapping.get(std_col)
        score = smart_column_score(src, std_col, aliases) if src else 0
        required = MAPPER_FIELD_INFO.get(std_col, ("", "", "", "", std_col in CORE_REQUIRED))[4]
        rows.append({
            "App mező": mapper_label(std_col),
            "Technikai mező": std_col,
            "Mit jelent?": mapper_desc(std_col),
            "Várt egység": mapper_unit(std_col),
            "Kiválasztott Excel oszlop": src or "",
            "Bizonyosság": score,
            "Kötelező": "igen" if required else "nem",
            "Figyelmeztetés": mapper_warning(std_col, src),
        })
    return pd.DataFrame(rows)


def mapping_compatibility_score(mapping: Dict[str, Optional[str]]) -> Tuple[int, List[str]]:
    weights = {
        "player_name": 20, "session_type": 15, "start_time": 15,
        "total_distance": 8, "duration": 5, "distance_per_min": 7,
        "max_speed": 8, "sprints": 5, "speed_zone_4": 5,
        "speed_zone_5": 7, "training_load": 5, "match_minutes": 4, "acc_high": 3,
        "dec_high": 3, "high_efforts": 4,
    }
    total = sum(weights.values())
    got = 0
    missing = []
    for k, w in weights.items():
        if mapping.get(k):
            got += w
        else:
            missing.append(mapper_label(k))
    return int(round(got / total * 100)), missing


def render_mapping_score(mapping: Dict[str, Optional[str]]) -> None:
    score, missing = mapping_compatibility_score(mapping)
    if score >= 85:
        st.success(f"GPS fájl kompatibilitás: {score}% – nagyon jó")
    elif score >= 65:
        st.warning(f"GPS fájl kompatibilitás: {score}% – használható, de néhány mező hiányzik")
    else:
        st.error(f"GPS fájl kompatibilitás: {score}% – javítsd a mappinget")
    if missing:
        st.caption("Hiányzó fontos mezők: " + ", ".join(missing[:8]) + ("..." if len(missing) > 8 else ""))


def normalize_combined_fields(out: pd.DataFrame, mapping: Dict[str, Optional[str]]) -> pd.DataFrame:
    """Összevont GPS mezők kezelése, pl. Distance(4+5).
    Ha csak 4+5 oszlop van, azt HSR-ként használjuk. A sprint distance ilyenkor becslés,
    de az eredeti összevont értéket nem duplázzuk a HSR számításnál.
    """
    out = out.copy()
    src_z4 = str(mapping.get("speed_zone_4") or "").lower()
    src_z5 = str(mapping.get("speed_zone_5") or "").lower()
    combined_45 = any(x in src_z4 for x in ["4+5", "hsr", "high speed"])
    out["combined_zone_4_5_used"] = bool(combined_45)

    if "speed_zone_4" not in out.columns:
        out["speed_zone_4"] = 0
    if "speed_zone_5" not in out.columns:
        out["speed_zone_5"] = 0

    z5_missing = out["speed_zone_5"].fillna(0).sum() == 0
    if z5_missing and combined_45 and out["speed_zone_4"].fillna(0).sum() > 0:
        out["speed_zone_5"] = out["speed_zone_4"] * 0.30
        out["estimated_sprint_distance"] = True
    else:
        out["estimated_sprint_distance"] = False

    # HSR: külön z4+z5 esetén összeg, összevont 4+5 esetén maga az összevont oszlop.
    if combined_45:
        out["hsr_distance"] = out["speed_zone_4"]
    else:
        out["hsr_distance"] = out[["speed_zone_4", "speed_zone_5"]].sum(axis=1, min_count=1)
    out["sprint_distance"] = out["speed_zone_5"]
    return out

# -----------------------------------------------------------------------------
# V4.5 - Smart Mapper fallback helpers
# -----------------------------------------------------------------------------
def _norm_mapping_text(text: object) -> str:
    import unicodedata
    text = str(text or "").lower().replace("\u00ad", " ")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()

def smart_column_score(source_col: object, std_col: str, aliases: List[str]) -> int:
    src = _norm_mapping_text(source_col)
    if not src:
        return 0
    alias_norms = [_norm_mapping_text(a) for a in aliases] + [_norm_mapping_text(std_col)]
    best = 0
    for a in alias_norms:
        if not a:
            continue
        if src == a:
            best = max(best, 100)
        elif src.replace(" ", "") == a.replace(" ", ""):
            best = max(best, 96)
        elif a in src or src in a:
            best = max(best, 86 if len(a) >= 4 else 70)
        else:
            src_tokens = set(src.split())
            a_tokens = set(a.split())
            if src_tokens and a_tokens:
                overlap = len(src_tokens & a_tokens) / max(1, len(a_tokens))
                best = max(best, int(overlap * 78))
    hints = {
        "player_name": ["player", "jatekos", "nev", "name", "athlete"],
        "session_type": ["type", "tipus", "match", "game", "training", "edzes", "meccs"],
        "start_time": ["date", "datum", "start", "day", "ido"],
        "duration": ["duration", "minutes", "perc", "time"],
        "match_minutes": ["minutes played", "match minutes", "jatekp", "jatekperc", "meccsperc", "playing time"],
        "total_distance": ["distance", "tav", "dist", "total"],
        "distance_per_min": ["min", "minute", "rel", "per"],
        "max_speed": ["max", "speed", "velocity", "vmax"],
        "sprints": ["sprint"],
        "training_load": ["load", "terheles", "workload"],
        "speed_zone_4": ["zone 4", "z4", "19", "24", "hsr"],
        "speed_zone_5": ["zone 5", "z5", "25", "sprint distance"],
        "acc_high": ["acc", "acceleration", "gyorsulas", "3"],
        "dec_high": ["dec", "deceleration", "lassitas", "-3"],
        "high_efforts": ["high effort", "effort"],
    }
    if std_col in hints:
        hit = sum(1 for h in hints[std_col] if h in src)
        if hit:
            best = max(best, min(94, 58 + hit * 14))
    return int(best)

def suggest_mapping(raw_df: pd.DataFrame) -> Dict[str, Optional[str]]:
    suggestions: Dict[str, Optional[str]] = {}
    used = set()
    for std_col, aliases in STANDARD_COLUMNS.items():
        scored = []
        for c in raw_df.columns:
            if c in used:
                continue
            scored.append((smart_column_score(c, std_col, aliases), c))
        scored.sort(reverse=True, key=lambda x: x[0])
        if scored and scored[0][0] >= 58:
            suggestions[std_col] = scored[0][1]
            used.add(scored[0][1])
        else:
            suggestions[std_col] = None
    return suggestions

def mapping_quality_df(raw_df: pd.DataFrame, mapping: Dict[str, Optional[str]]) -> pd.DataFrame:
    rows = []
    for std_col, aliases in STANDARD_COLUMNS.items():
        src = mapping.get(std_col)
        score = smart_column_score(src, std_col, aliases) if src else 0
        rows.append({
            "Standard mező": std_col,
            "Felismert forrásoszlop": src or "",
            "Bizonyosság": score,
            "Kötelező": "igen" if std_col in CORE_REQUIRED else "nem",
            "Magyar név": METRIC_LABELS.get(std_col, std_col),
        })
    return pd.DataFrame(rows)

def export_mapping_profile(mapping: Dict[str, Optional[str]], profile_name: str = "GPS mapping profil") -> bytes:
    payload = {
        "profile_name": profile_name,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "mapping": mapping,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")

def load_mapping_profile(uploaded_file) -> Dict[str, Optional[str]]:
    try:
        payload = json.loads(uploaded_file.read().decode("utf-8"))
        return payload.get("mapping", {})
    except Exception:
        return {}

@dataclass
class Insight:
    title: str
    severity: str
    observation: str
    impact: str
    recommendation: str
    scope: str = "Csapat"

    def as_dict(self) -> Dict[str, str]:
        return {
            "Súlyosság": self.severity,
            "Terület": self.scope,
            "Megállapítás": self.title,
            "Mit látunk?": self.observation,
            "Miért fontos?": self.impact,
            "Javaslat": self.recommendation,
        }


# -----------------------------------------------------------------------------
# Segédfüggvények
# -----------------------------------------------------------------------------
def clean_col_name(col: object) -> str:
    if col is None:
        return ""
    return str(col).replace("\u00ad", "").strip()


def find_column(df: pd.DataFrame, aliases: List[str]) -> Optional[str]:
    cleaned = {clean_col_name(c).lower(): c for c in df.columns}
    for alias in aliases:
        key = clean_col_name(alias).lower()
        if key in cleaned:
            return cleaned[key]

    # óvatos fuzzy fallback
    for alias in aliases:
        a = clean_col_name(alias).lower()
        for c in df.columns:
            cc = clean_col_name(c).lower()
            if a and (a in cc or cc in a):
                return c
    return None


def to_numeric(series: pd.Series) -> pd.Series:
    if series.dtype == object:
        return pd.to_numeric(series.astype(str).str.replace(",", ".", regex=False), errors="coerce")
    return pd.to_numeric(series, errors="coerce")


def duration_to_minutes(x) -> float:
    if pd.isna(x):
        return np.nan
    if isinstance(x, pd.Timedelta):
        return x.total_seconds() / 60
    if hasattr(x, "hour") and hasattr(x, "minute") and hasattr(x, "second"):
        return x.hour * 60 + x.minute + x.second / 60
    if isinstance(x, str):
        try:
            td = pd.to_timedelta(x)
            return td.total_seconds() / 60
        except Exception:
            return np.nan
    if isinstance(x, (int, float, np.number)):
        # Excel időtört: 0.5 = 12 óra
        if float(x) < 2:
            return float(x) * 24 * 60
        return float(x)
    return np.nan


def normalize_session_type(x: object) -> str:
    text = str(x).strip().lower()
    if "meccs" in text or "match" in text or "game" in text or "mérk" in text:
        return "Meccs"
    if "edzés" in text or "training" in text or "train" in text:
        return "Edzés"
    return str(x).strip() if str(x).strip() else "Ismeretlen"



def _plausible_datetime(dt: object) -> Optional[pd.Timestamp]:
    """Csak életszerű sportadat-dátumot enged át. Mezszám/ID ne lehessen év."""
    try:
        ts = pd.to_datetime(dt, errors="coerce", dayfirst=True)
        if pd.isna(ts):
            return None
        if 2015 <= int(ts.year) <= 2035:
            return ts
    except Exception:
        return None
    return None


def extract_date_from_text(text: object) -> Optional[pd.Timestamp]:
    """Dátum kinyerése szövegből, de csak valódi dátummintából.

    Fontos: a puszta számokat (pl. mezszám: 49, 9946) nem értelmezzük dátumként.
    Ez akadályozza meg a 9946-W01 típusú hibás hetek képződését.
    """
    s = str(text or "").strip()
    if not s or s.lower() in ["nan", "none", "nat"]:
        return None

    # Puszta szám nem dátum, kivéve Excel serial tartomány.
    if re.fullmatch(r"\d+(\.\d+)?", s):
        try:
            val = float(s)
            if 25000 < val < 90000:
                return _plausible_datetime(pd.to_datetime(val, unit="D", origin="1899-12-30", errors="coerce"))
        except Exception:
            pass
        return None

    patterns = [
        (r"(20\d{2})[-._/ ](\d{1,2})[-._/ ](\d{1,2})", "ymd"),
        (r"(\d{1,2})[-._/ ](\d{1,2})[-._/ ](20\d{2})", "dmy"),
    ]
    for pat, mode in patterns:
        m = re.search(pat, s)
        if m:
            try:
                if mode == "ymd":
                    return _plausible_datetime(f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}")
                return _plausible_datetime(f"{m.group(3)}-{int(m.group(2)):02d}-{int(m.group(1)):02d}")
            except Exception:
                pass

    # 02.01. jellegű lapnév / split. Csak akkor, ha tényleg dátumszerű pontozás van benne.
    m = re.search(r"(?<!\d)(\d{1,2})\.(\d{1,2})\.(?!\d)", s)
    if m:
        year = datetime.now().year
        return _plausible_datetime(f"{year}-{int(m.group(2)):02d}-{int(m.group(1)):02d}")

    # Általános parser csak akkor, ha van benne dátumelválasztó vagy hónap/év jelleg.
    if not re.search(r"[-./:]|20\d{2}", s):
        return None
    return _plausible_datetime(pd.to_datetime(s, errors="coerce", dayfirst=True))


def parse_datetime_series(series: pd.Series, fallback_source: Optional[pd.Series] = None, sheet_name: str = "") -> pd.Series:
    """Robusztus, de védett dátumfelismerés.

    Sorrend:
    1) Ha az oszlop eleve dátum/datetime, azt használja.
    2) Szöveges dátumokat parse-ol.
    3) Csak sikertelen esetben néz Split/lapnév fallbacket.
    4) 2015-2035 közötti dátumokat fogad el.
    """
    n = len(series) if series is not None else (len(fallback_source) if fallback_source is not None else 0)
    if series is None:
        base = pd.Series([pd.NaT] * n)
    else:
        # Excel/pandas datetime oszlopnál ez a legfontosabb út: ne fusson rá text-rescue.
        base = pd.to_datetime(series, errors="coerce", dayfirst=True)
        base = base.apply(lambda x: _plausible_datetime(x) if pd.notna(x) else pd.NaT)

        success_rate = float(base.notna().mean()) if len(base) else 0.0
        if success_rate < 0.60:
            text_parsed = series.apply(extract_date_from_text)
            text_parsed = pd.to_datetime(text_parsed, errors="coerce")
            if len(text_parsed) and text_parsed.notna().mean() > success_rate:
                base = base.fillna(text_parsed)

    if fallback_source is not None and (len(base) == 0 or base.notna().mean() < 0.60):
        fallback = fallback_source.apply(extract_date_from_text)
        fallback = pd.to_datetime(fallback, errors="coerce")
        base = base.fillna(fallback)

    if len(base) and base.notna().mean() < 0.60:
        sheet_dt = extract_date_from_text(sheet_name)
        if sheet_dt is not None and pd.notna(sheet_dt):
            base = base.fillna(sheet_dt)

    return pd.to_datetime(base, errors="coerce")


def make_iso_week_series(dt_series: pd.Series) -> pd.Series:
    """Egységes ISO hét címke: 2025-W29.

    V6.2 WEEK RESCUE:
    - csak a dátumrészt használja, az időpontot figyelmen kívül hagyja;
    - ha egy rövid, összefüggő dátumtartományból irreálisan sok hét keletkezne,
      akkor a teljes blokkot a kezdődátum ISO hetére menti.
    Ez védi az appot attól, hogy egymás utáni napokat külön heteknek vegyen,
    ha a dátumcellában időpont vagy extra szöveg is szerepel.
    """
    dts = pd.to_datetime(dt_series, errors="coerce")
    date_only = dts.dt.normalize() if hasattr(dts, "dt") else pd.to_datetime(dts, errors="coerce")
    weeks = date_only.apply(lambda x: f"{int(x.isocalendar().year)}-W{int(x.isocalendar().week):02d}" if pd.notna(x) else np.nan)

    valid = date_only.dropna()
    if len(valid) >= 2:
        span_days = int((valid.max() - valid.min()).days)
        unique_weeks = weeks.dropna().nunique()
        if span_days <= 10 and unique_weeks > 3:
            anchor = valid.min()
            rescue_week = f"{int(anchor.isocalendar().year)}-W{int(anchor.isocalendar().week):02d}"
            weeks = weeks.where(date_only.isna(), rescue_week)
            st.session_state["week_rescue_applied"] = {
                "reason": "short_date_span_many_weeks",
                "span_days": span_days,
                "original_unique_weeks": int(unique_weeks),
                "rescued_week": rescue_week,
                "date_min": str(valid.min()),
                "date_max": str(valid.max()),
            }
        else:
            st.session_state["week_rescue_applied"] = {
                "reason": "not_needed",
                "span_days": span_days,
                "unique_weeks": int(unique_weeks),
                "date_min": str(valid.min()),
                "date_max": str(valid.max()),
            }
    return weeks


def sheet_is_likely_helper(sheet_name: str, raw_df: Optional[pd.DataFrame] = None) -> bool:
    """Segédlapok kizárása az összesített adatlapból."""
    name = _norm_mapping_text(sheet_name)
    helper_tokens = [
        "dashboard", "riport", "report", "summary", "osszefoglalo", "összefoglaló",
        "benchmark", "benchmarks", "mapping", "mapper", "settings", "beallitas",
        "útmutató", "utmutato", "guide", "help", "readme", "sablon", "template",
        "pivot", "chart", "diagram", "grafikon", "calc", "szamitas", "reference",
        "lista", "players", "jatekoslista", "metadata"
    ]
    if any(tok in name for tok in helper_tokens):
        return True
    if raw_df is None or raw_df.empty:
        return True
    sample = " ".join([str(x).lower() for x in raw_df.head(8).fillna("").astype(str).values.ravel()[:200]])
    data_hints = ["name", "player", "játékos", "jatekos", "total distance", "teljes táv", "duration", "split", "date", "dátum", "datum"]
    return sum(1 for h in data_hints if h in sample) < 2

def detect_header_row(raw_df: pd.DataFrame) -> Optional[int]:
    """Általános fejlécfelismerés GPS/Excel exportokhoz.

    Nem fix sorra épít. Az első 50 sort pontozza magyar és angol GPS kulcsszavak alapján.
    Így működik akkor is, ha a fejléc az első sorban van (MegyeI.xlsx Data lap), és akkor is,
    ha néhány üres/logó/meta sor előzi meg.
    """
    if raw_df is None or raw_df.empty:
        return None

    header_keywords = {
        "player": ["jatekos", "játékos", "player", "athlete", "name", "nev", "név"],
        "date": ["kezdesi ido", "kezdési idő", "start time", "session date", "date", "datum", "dátum", "split"],
        "type": ["tipus", "típus", "session type", "type", "training", "match", "edzes", "edzés", "meccs"],
        "duration": ["idotartam", "időtartam", "duration", "minutes", "time"],
        "distance": ["teljes tav", "teljes táv", "total distance", "distance", "dist", "tav perc", "táv perc"],
        "speed": ["maximalis sebesseg", "maximális sebesség", "max speed", "top speed", "velocity"],
        "load": ["edzesi terheles", "edzési terhelés", "training load", "player load", "workload", "load"],
        "intensity": ["sprintek", "sprints", "gyorsulas", "gyorsulás", "lassitas", "lassítás", "high efforts"],
    }

    max_scan = min(50, len(raw_df))
    best_i: Optional[int] = None
    best_score = -999

    for i in range(max_scan):
        row = raw_df.iloc[i].tolist()
        cells = [str(v).strip() for v in row if str(v).strip().lower() not in ["nan", "none", ""]]
        if not cells:
            continue
        joined_raw = " | ".join(cells)
        joined_norm = _norm_mapping_text(joined_raw)

        score = 0
        matched_groups = 0
        for _, kws in header_keywords.items():
            if any(_norm_mapping_text(k) in joined_norm for k in kws):
                score += 10
                matched_groups += 1

        # A fejléc többnyire sok szöveges cellát tartalmaz, kevés tiszta számot.
        numeric_like = sum(1 for c in cells if re.fullmatch(r"\d+(\.\d+)?", c))
        text_like = len(cells) - numeric_like
        if text_like >= 4:
            score += 8
        if numeric_like > max(3, text_like):
            score -= 18

        # Következő sorban legyen valamennyi adat, különben lehet címblokk.
        if i + 1 < len(raw_df):
            next_vals = [str(v).strip() for v in raw_df.iloc[i + 1].tolist() if str(v).strip().lower() not in ["nan", "none", ""]]
            if len(next_vals) >= 3:
                score += 4

        if matched_groups >= 2 and score > best_score:
            best_score = score
            best_i = i

    return best_i if best_score >= 24 else None

def normalize_uploaded_sheet(raw_df: pd.DataFrame, sheet_name: str = "") -> pd.DataFrame:
    """Egy munkalap megtisztítása az app számára.
    Kezeli:
    - üres első sor
    - 2. sorban lévő fejléc
    - több meccslap
    - hiányzó Típus / Kezdési idő / Szakasz neve mező
    """
    if raw_df is None or raw_df.empty:
        return pd.DataFrame()

    df = raw_df.copy()

    # 1) Ha header=None-ból jön, keressük meg a valódi fejlécsort.
    header_row = detect_header_row(df)
    if header_row is not None:
        new_cols = []
        for j, x in enumerate(df.iloc[header_row].tolist()):
            name = clean_col_name(x)
            new_cols.append(name if name and name.lower() not in ["nan", "none"] else f"col_{j+1}")
        df = df.iloc[header_row + 1:].copy()
        df.columns = new_cols
    else:
        # 2) Ha mégis header=0 jelleggel jönne.
        df.columns = [clean_col_name(c) if clean_col_name(c) else f"col_{i+1}" for i, c in enumerate(df.columns)]

    # Üres és Unnamed oszlopok törlése.
    df = df.dropna(how="all")
    keep_cols = []
    for c in df.columns:
        cs = str(c).strip()
        if not cs:
            continue
        if cs.startswith("Unnamed"):
            continue
        if cs.lower() in ["nan", "none"]:
            continue
        keep_cols.append(c)
    df = df[keep_cols] if keep_cols else df

    # Cellák whitespace tisztítása szöveges oszlopokban.
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].apply(lambda x: str(x).strip() if pd.notna(x) else x)

    # Tipikus GPS-export: Name nevű oszlop legyen biztosan felismerhető.
    # Nem nevezzük át player_name-re itt, csak adunk magyar alias oszlopot, ha kell.
    if "Name" in df.columns and "Játékos neve" not in df.columns:
        df["Játékos neve"] = df["Name"]

    # Ha nincs típus, ebből a workbookból ez meccsadat.
    if not any(c in df.columns for c in ["Típus", "Type", "Session Type", "Edzés/Meccs"]):
        df["Típus"] = "Meccs"

    if not any(c in df.columns for c in ["Szakasz neve", "Session Name", "Session"]):
        df["Szakasz neve"] = sheet_name or "GPS mérkőzés"

    # Kezdési idő kinyerése Splitből vagy lapnévből, ha nincs explicit dátum.
    has_start = any(c in df.columns for c in ["Kezdési idő", "Start Time", "Start", "Date", "Dátum", "Session Date"])
    if not has_start:
        dates = []
        split_col = "Split" if "Split" in df.columns else None
        for _, row in df.iterrows():
            dt = extract_date_from_text(row.get(split_col, "")) if split_col else None
            if dt is None or pd.isna(dt):
                dt = extract_date_from_text(sheet_name)
            if dt is None or pd.isna(dt):
                dt = pd.to_datetime("2026-01-01")
            dates.append(dt.strftime("%Y-%m-%d 17:00"))
        df["Kezdési idő"] = dates

    return df


def prepare_uploaded_sheets(sheets: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
    """Munkalapok normalizálása.
    - Data/adat lap elsődleges és első helyen jelenik meg.
    - Segédlapok nem kerülnek az összesített adatlapba.
    - Ha van Data lap, az lesz a default; ha nincs, csak a valószínű adatlapok összesülnek.
    """
    prepared: Dict[str, pd.DataFrame] = {}
    relevant_frames: List[pd.DataFrame] = []
    data_like_names = []

    def is_data_sheet_name(n: str) -> bool:
        nn = _norm_mapping_text(n)
        return nn in ["data", "adat", "adatok", "gps data", "gps", "raw data", "nyers adat", "nyers adatok"]

    ordered_items = sorted(sheets.items(), key=lambda kv: (0 if is_data_sheet_name(kv[0]) else 1, kv[0]))
    for name, raw in ordered_items:
        clean = normalize_uploaded_sheet(raw, name)
        if clean is None or clean.empty:
            continue
        prepared[name] = clean
        if not sheet_is_likely_helper(name, raw):
            relevant_frames.append(clean)
            data_like_names.append(name)

    final: Dict[str, pd.DataFrame] = {}
    data_names = [n for n in prepared if is_data_sheet_name(n)]
    if data_names:
        # Data lap legyen az elsődleges: a felhasználó ezt látja elsőként.
        for n in data_names:
            final[n] = prepared[n]
        # Az összesített lap csak releváns adatlapokból épül, segédlapok nélkül.
        frames = [prepared[n] for n in data_like_names if n in prepared]
        if frames:
            final["Összes releváns adatlap"] = pd.concat(frames, ignore_index=True)
    elif relevant_frames:
        final["Összes releváns adatlap"] = pd.concat(relevant_frames, ignore_index=True)

    for name, df_sheet in prepared.items():
        if name not in final:
            final[name] = df_sheet
    return final if final else sheets

@st.cache_data(show_spinner=False)
def read_excel_all(file) -> Dict[str, pd.DataFrame]:
    # header=None kell, mert sok GPS exportnál az első sor üres,
    # a valódi fejléc a 2. sorban van.
    return pd.read_excel(file, sheet_name=None, header=None)


# =========================================================
# V136 - ZIP / multi-session GPS import
# =========================================================
def _fpi_read_csv_bytes_v136(data: bytes, name: str = "") -> pd.DataFrame:
    """CSV beolvasás több tipikus kódolással és automatikus elválasztóval."""
    last_err = None
    for enc in ["utf-8-sig", "utf-8", "latin2", "cp1250", "iso-8859-2"]:
        try:
            return pd.read_csv(io.BytesIO(data), sep=None, engine="python", encoding=enc, header=None)
        except Exception as e:
            last_err = e
    raise ValueError(f"CSV beolvasási hiba ({name}): {last_err}")


def _fpi_read_gps_upload_to_sheets_v136(uploaded_file) -> Tuple[Dict[str, pd.DataFrame], pd.DataFrame, str, bool]:
    """Egy Excel/CSV vagy ZIP-ben lévő több GPS fájl beolvasása egységes sheet-dict formába.

    A ZIP mód lényege:
    - minden belső Excel/CSV fájlt beolvasunk,
    - a sheet neve tartalmazza a fájlnevet is,
    - a meglévő normalize_uploaded_sheet + Smart Mapper ugyanúgy fut rajtuk,
    - végül a prepare_uploaded_sheets összefűzi az összes releváns adatlapot.
    """
    if uploaded_file is None:
        return {}, pd.DataFrame(), "", False

    try:
        data = uploaded_file.getvalue()
    except Exception:
        data = uploaded_file.read()

    upload_name = getattr(uploaded_file, "name", "gps_upload")
    signature = hashlib.md5(data or b"").hexdigest()
    ext = Path(str(upload_name)).suffix.lower()

    sheets: Dict[str, pd.DataFrame] = {}
    report_rows: List[Dict[str, object]] = []

    def _add_report(name: str, status: str, rows: int = 0, sheets_n: int = 0, note: str = "") -> None:
        report_rows.append({
            "Fájl": name,
            "Státusz": status,
            "Sor": int(rows or 0),
            "Lap": int(sheets_n or 0),
            "Megjegyzés": note,
        })

    def _read_single_file(data_bytes: bytes, name: str) -> None:
        low = str(name).lower()
        try:
            if low.endswith((".xlsx", ".xls")):
                xls_sheets = pd.read_excel(io.BytesIO(data_bytes), sheet_name=None, header=None)
                local_count = 0
                local_rows = 0
                for sheet_name, df_sheet in (xls_sheets or {}).items():
                    if df_sheet is None or df_sheet.empty:
                        continue
                    key = f"{Path(name).stem} / {sheet_name}"
                    # Ütközés esetén ne írjuk felül.
                    if key in sheets:
                        key = f"{Path(name).stem} / {sheet_name} / {len(sheets)+1}"
                    sheets[key] = df_sheet
                    local_count += 1
                    local_rows += len(df_sheet)
                _add_report(name, "OK", local_rows, local_count, "Excel beolvasva")
            elif low.endswith(".csv"):
                df_csv = _fpi_read_csv_bytes_v136(data_bytes, name)
                key = Path(name).stem
                if key in sheets:
                    key = f"{Path(name).stem} / {len(sheets)+1}"
                sheets[key] = df_csv
                _add_report(name, "OK", len(df_csv), 1, "CSV beolvasva")
            else:
                _add_report(name, "Kihagyva", 0, 0, "Nem támogatott fájltípus")
        except Exception as e:
            _add_report(name, "Hiba", 0, 0, str(e)[:220])

    is_zip = ext == ".zip"
    if is_zip:
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = [n for n in zf.namelist() if not n.endswith("/") and not Path(n).name.startswith(("~$", "."))]
                supported = [n for n in names if n.lower().endswith((".xlsx", ".xls", ".csv"))]
                if not supported:
                    _add_report(upload_name, "Hiba", 0, 0, "A ZIP nem tartalmaz támogatott Excel/CSV fájlt")
                for n in supported:
                    _read_single_file(zf.read(n), n)
        except Exception as e:
            _add_report(upload_name, "Hiba", 0, 0, f"ZIP beolvasási hiba: {e}")
    else:
        _read_single_file(data, upload_name)

    report_df = pd.DataFrame(report_rows)
    return sheets, report_df, signature, is_zip


def _fpi_render_gps_import_report_v136(report_df: pd.DataFrame, is_zip: bool) -> None:
    """Import ellenőrző blokk ZIP vagy egyfájlos GPS importhoz."""
    if report_df is None or report_df.empty:
        return
    ok = int((report_df["Státusz"] == "OK").sum()) if "Státusz" in report_df.columns else 0
    bad = int((report_df["Státusz"] == "Hiba").sum()) if "Státusz" in report_df.columns else 0
    skipped = int((report_df["Státusz"] == "Kihagyva").sum()) if "Státusz" in report_df.columns else 0
    total_rows = int(pd.to_numeric(report_df.get("Sor", 0), errors="coerce").fillna(0).sum()) if "Sor" in report_df.columns else 0

    with st.expander("📦 Feltöltött GPS fájlok ellenőrzése", expanded=bool(is_zip)):
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("OK fájl", ok)
        c2.metric("Hibás fájl", bad)
        c3.metric("Kihagyott", skipped)
        c4.metric("Nyers sor", total_rows)
        st.caption("ZIP mód esetén minden belső Excel/CSV fájlra ugyanaz a normalizálás és Smart Mapper logika fut. Hibás fájl nem állítja meg az egész importot.")
        st.dataframe(report_df, use_container_width=True, hide_index=True)


def standardize_dataframe(raw: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, Optional[str]], List[str]]:
    df = raw.copy()
    df.columns = [clean_col_name(c) for c in df.columns]

    mapping: Dict[str, Optional[str]] = {}
    out = pd.DataFrame()

    for std_col, aliases in STANDARD_COLUMNS.items():
        source = find_column(df, aliases)
        mapping[std_col] = source
        if source is not None:
            out[std_col] = df[source]

    # GPS export fallback: ha van Name oszlop, az legyen játékosnév.
    if "player_name" not in out.columns and "Name" in df.columns:
        out["player_name"] = df["Name"]
        mapping["player_name"] = "Name"
    if "session_type" not in out.columns and "Típus" in df.columns:
        out["session_type"] = df["Típus"]
        mapping["session_type"] = "Típus"
    if "start_time" not in out.columns and "Kezdési idő" in df.columns:
        out["start_time"] = df["Kezdési idő"]
        mapping["start_time"] = "Kezdési idő"

    missing_core = [c for c in CORE_REQUIRED if c not in out.columns]
    if missing_core:
        return out, mapping, missing_core

    out["player_name"] = out["player_name"].astype(str).str.strip()
    out["session_type"] = out["session_type"].apply(normalize_session_type)
    fallback_source = df["Split"] if "Split" in df.columns else None
    out["start_time"] = parse_datetime_series(out["start_time"], fallback_source=fallback_source)
    out["session_date"] = out["start_time"].dt.date
    out["week"] = make_iso_week_series(out["start_time"])

    if "duration" in out.columns:
        out["duration_min"] = out["duration"].apply(duration_to_minutes)
    else:
        out["duration_min"] = np.nan

    numeric_cols = [
        "total_distance", "distance_per_min", "max_speed", "avg_speed", "sprints", "match_minutes",
        "speed_zone_3", "speed_zone_4", "speed_zone_5", "training_load", "cardio_load",
        "recovery_hours", "muscle_load", "hr_avg", "hr_max", "hrv", "high_efforts", "acc_low", "acc_mid",
        "acc_high", "dec_low", "dec_mid", "dec_high",
    ]
    for col in numeric_cols:
        if col in out.columns:
            out[col] = to_numeric(out[col])

    for col in ["speed_zone_3", "speed_zone_4", "speed_zone_5", "acc_low", "acc_mid", "acc_high", "dec_low", "dec_mid", "dec_high"]:
        if col not in out.columns:
            out[col] = 0

    out["hsr_distance"] = out[["speed_zone_4", "speed_zone_5"]].sum(axis=1, min_count=1)
    out["sprint_distance"] = out["speed_zone_5"]
    out["acc_count"] = out[["acc_low", "acc_mid", "acc_high"]].sum(axis=1, min_count=1)
    out["dec_count"] = out[["dec_low", "dec_mid", "dec_high"]].sum(axis=1, min_count=1)
    if "high_efforts" not in out.columns or out["high_efforts"].isna().all():
        out["high_efforts"] = out[["acc_mid", "acc_high", "dec_mid", "dec_high"]].sum(axis=1, min_count=1)
    else:
        out["high_efforts"] = to_numeric(out["high_efforts"])

    out = normalize_combined_fields(out, mapping)

    if "distance_per_min" not in out.columns or out["distance_per_min"].isna().all():
        if "total_distance" in out.columns:
            out["distance_per_min"] = out["total_distance"] / out["duration_min"]

    out = out.dropna(subset=["start_time"])
    out = out[out["player_name"].str.len() > 0]
    out = out[~out["player_name"].str.lower().str.contains("benchmark|átlag|atlag|összesen|osszesen", na=False)]
    out = finalize_exposure_columns(add_position_group(out))

    return out, mapping, []

def apply_mapping_to_raw(raw: pd.DataFrame, mapping: Dict[str, Optional[str]]) -> Tuple[pd.DataFrame, Dict[str, Optional[str]], List[str]]:
    df = raw.copy()
    df.columns = [clean_col_name(c) for c in df.columns]
    out = pd.DataFrame()
    fixed_mapping: Dict[str, Optional[str]] = {}
    for std_col in STANDARD_COLUMNS:
        source = mapping.get(std_col)
        source = clean_col_name(source) if source else None
        fixed_mapping[std_col] = source
        if source and source in df.columns:
            out[std_col] = df[source]
    missing_core = [c for c in CORE_REQUIRED if c not in out.columns]
    if missing_core:
        return out, fixed_mapping, missing_core

    out["player_name"] = out["player_name"].astype(str).str.strip()
    out["session_type"] = out["session_type"].apply(normalize_session_type)
    fallback_source = df["Split"] if "Split" in df.columns else None
    out["start_time"] = parse_datetime_series(out["start_time"], fallback_source=fallback_source)
    out["session_date"] = out["start_time"].dt.date
    out["week"] = make_iso_week_series(out["start_time"])
    if "duration" in out.columns:
        out["duration_min"] = out["duration"].apply(duration_to_minutes)
    else:
        out["duration_min"] = np.nan

    numeric_cols = [
        "total_distance", "distance_per_min", "max_speed", "avg_speed", "sprints", "match_minutes",
        "speed_zone_3", "speed_zone_4", "speed_zone_5", "training_load", "cardio_load",
        "recovery_hours", "muscle_load", "hr_avg", "hr_max", "hrv", "high_efforts", "acc_low", "acc_mid",
        "acc_high", "dec_low", "dec_mid", "dec_high",
    ]
    for col in numeric_cols:
        if col in out.columns:
            out[col] = to_numeric(out[col])
    for col in ["speed_zone_3", "speed_zone_4", "speed_zone_5", "acc_low", "acc_mid", "acc_high", "dec_low", "dec_mid", "dec_high"]:
        if col not in out.columns:
            out[col] = 0

    out["hsr_distance"] = out[["speed_zone_4", "speed_zone_5"]].sum(axis=1, min_count=1)
    out["sprint_distance"] = out["speed_zone_5"]
    out["acc_count"] = out[["acc_low", "acc_mid", "acc_high"]].sum(axis=1, min_count=1)
    out["dec_count"] = out[["dec_low", "dec_mid", "dec_high"]].sum(axis=1, min_count=1)
    if "high_efforts" not in out.columns or out["high_efforts"].isna().all():
        out["high_efforts"] = out[["acc_mid", "acc_high", "dec_mid", "dec_high"]].sum(axis=1, min_count=1)
    else:
        out["high_efforts"] = to_numeric(out["high_efforts"])
    out = normalize_combined_fields(out, mapping)

    if "distance_per_min" not in out.columns or out["distance_per_min"].isna().all():
        if "total_distance" in out.columns:
            out["distance_per_min"] = out["total_distance"] / out["duration_min"]
    out = out.dropna(subset=["start_time"])
    out = out[out["player_name"].str.len() > 0]
    out = out[~out["player_name"].str.lower().str.contains("benchmark|átlag|atlag|összesen|osszesen", na=False)]
    out = finalize_exposure_columns(add_position_group(out))
    return out, fixed_mapping, []




def render_emergency_mapper(raw_df: pd.DataFrame, current_mapping: Dict[str, Optional[str]], missing_core: List[str]) -> None:
    """Mapper blokk akkor is, ha az app még nem tud elemezni.
    Cél: ne fusson hibára / st.stop előtt legyen javítási lehetőség.
    """
    st.markdown("### 🧭 Smart Excel Mapper – kötelező mezők javítása")
    st.info("A fájl szerkezete eltér a várt sablontól. Állítsd be a kötelező mezőket, majd kattints az alkalmazásra.")
    render_mapping_score(current_mapping)
    st.markdown("#### App mezők magyarázata")
    st.dataframe(enhanced_mapping_quality_df(raw_df, current_mapping), use_container_width=True)

    if raw_df is None or raw_df.empty:
        st.error("Nincs beolvasható nyers adat.")
        return

    cols = [""] + sorted([str(c) for c in raw_df.columns], key=lambda x: x.lower())
    manual = dict(current_mapping or {})

    col_a, col_b, col_c = st.columns(3)
    with col_a:
        default = manual.get("player_name") or ("Name" if "Name" in raw_df.columns else ("Játékos neve" if "Játékos neve" in raw_df.columns else ""))
        manual["player_name"] = st.selectbox(
            "Játékos neve / Name – szöveg",
            cols,
            index=cols.index(default) if default in cols else 0,
            key="emergency_map_player_name",
        ) or None

    with col_b:
        default = manual.get("session_type") or ("Típus" if "Típus" in raw_df.columns else "")
        manual["session_type"] = st.selectbox(
            "Típus – Edzés vagy Meccs",
            cols,
            index=cols.index(default) if default in cols else 0,
            key="emergency_map_session_type",
        ) or None

    with col_c:
        default = manual.get("start_time") or ("Kezdési idő" if "Kezdési idő" in raw_df.columns else ("Split" if "Split" in raw_df.columns else ""))
        manual["start_time"] = st.selectbox(
            "Kezdési idő / dátum – dátum/idő",
            cols,
            index=cols.index(default) if default in cols else 0,
            key="emergency_map_start_time",
        ) or None

    with st.expander("Haladó oszlopok mappingje", expanded=False):
        optional_fields = [
            "duration", "match_minutes", "total_distance", "distance_per_min", "max_speed", "sprints",
            "speed_zone_4", "speed_zone_5", "training_load", "muscle_load",
            "hr_avg", "hr_max", "acc_high", "dec_high", "high_efforts"
        ]
        for std_col in optional_fields:
            default = manual.get(std_col) or ""
            manual[std_col] = st.selectbox(
                f"{mapper_label(std_col)} | {mapper_unit(std_col)}",
                cols,
                index=cols.index(default) if default in cols else 0,
                key=f"emergency_map_{std_col}",
                help=mapper_desc(std_col),
            ) or None
            warn = mapper_warning(std_col, manual.get(std_col))
            if warn:
                st.warning(warn)

    c1, c2 = st.columns(2)
    with c1:
        if st.button("✅ Mapping alkalmazása", use_container_width=True, key="apply_emergency_mapping"):
            mapped_df, fixed_mapping, missing = apply_mapping_to_raw(raw_df, manual)
            if missing:
                st.error(f"Még hiányzik: {', '.join(missing)}")
            else:
                st.session_state["mapped_df_override"] = mapped_df
                st.session_state["manual_mapping"] = fixed_mapping
                st.success("Mapping alkalmazva. Újratöltöm az elemzést.")
                st.rerun()

    with c2:
        if st.button("♻️ Mapping törlése", use_container_width=True, key="clear_emergency_mapping"):
            st.session_state.pop("mapped_df_override", None)
            st.session_state.pop("manual_mapping", None)
            st.rerun()

    st.caption("Tipp: Distance(4+5) = összevont nagysebességű táv, ezt a Zone 4 / HSR mezőhöz érdemes tenni. Sprints count = darab, Total sprints distance = méter.")
    st.markdown("#### Beolvasott oszlopok")
    st.write(list(raw_df.columns))
    st.markdown("#### Adat előnézet")
    st.dataframe(raw_df.head(8), use_container_width=True)



# -----------------------------------------------------------------------------
# V6.0 - Kapus + játékperc + exposure normalizálás
# -----------------------------------------------------------------------------
KEEPER_POSITION_TOKENS = ["gk", "goalkeeper", "kapus", "kapusposzt", "kap", "keeper"]
FIELD_METRICS_FOR_EXPOSURE = [
    "total_distance", "hsr_distance", "sprint_distance", "sprints", "acc_count", "dec_count",
    "high_efforts", "training_load", "muscle_load", "distance_per_min", "max_speed",
]
KEEPER_BENCHMARK_WEIGHTS = {
    "total_distance": 0.40,
    "distance_per_min": 0.55,
    "hsr_distance": 0.15,
    "sprint_distance": 0.10,
    "sprints": 0.15,
    "acc_count": 0.55,
    "dec_count": 0.55,
    "high_efforts": 0.60,
    "training_load": 0.70,
    "muscle_load": 0.70,
    "max_speed": 0.45,
}

def is_keeper_position_value(value: object) -> bool:
    txt = _norm_mapping_text(value)
    if not txt:
        return False
    return any(tok in txt for tok in KEEPER_POSITION_TOKENS)

def add_keeper_flag_from_position(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    if "position" in out.columns:
        out["is_goalkeeper"] = out["position"].apply(is_keeper_position_value).astype(bool)
    elif "position_group" in out.columns:
        out["is_goalkeeper"] = out["position_group"].astype(str).str.lower().str.contains("kapus|gk|goalkeeper", na=False)
    else:
        out["is_goalkeeper"] = False
    return out

def render_keeper_controls_and_apply(df: pd.DataFrame) -> pd.DataFrame:
    """Kapusok kezelése.
    - Ha van poszt, automatikusan felismeri a kapusokat.
    - Ha nincs poszt / nincs felismerés, a választók között rákérdez és a kiválasztott játékosokat kapusként kezeli.
    """
    if df is None or df.empty or "player_name" not in df.columns:
        return df
    out = add_keeper_flag_from_position(df)
    players_all = sorted(out["player_name"].dropna().astype(str).unique().tolist())
    auto_keepers = sorted(out.loc[out.get("is_goalkeeper", False), "player_name"].dropna().astype(str).unique().tolist()) if "is_goalkeeper" in out.columns else []
    has_position_col = (
        "position" in out.columns
        and out["position"].notna().any()
        and out["position"].astype(str).str.strip().replace({"": np.nan, "nan": np.nan, "None": np.nan}).notna().any()
    )
    if not has_position_col:
        st.warning("Nincs felismerhető Poszt oszlop. Kérlek add meg, vannak-e kapusok az adatokban.")
    with st.expander("Kapusok és játékpercek", expanded=(not has_position_col)):
        if has_position_col:
            st.caption("Poszt oszlop alapján automatikus kapusfelismerés aktív. Ezt felülírhatod, ha szükséges.")
            st.write("Felismert kapusok: " + (", ".join(auto_keepers) if auto_keepers else "nincs"))
            manual_override = st.checkbox("Kapuslista kézi felülírása", value=False, key="keeper_manual_override")
            if manual_override:
                selected = st.multiselect("Kapusok kiválasztása", players_all, default=auto_keepers, key="manual_keeper_players")
                out["is_goalkeeper"] = out["player_name"].astype(str).isin(selected)
        else:
            st.markdown("**Kapusok kézi megadása szükséges**")
            has_keepers = st.radio(
                "Szerepelnek kapusok az adatok között?",
                ["Igen", "Nem"],
                horizontal=True,
                key="has_goalkeepers_without_position",
                help="Mivel nincs Poszt oszlop, az app csak így tudja külön kezelni a kapusokat."
            )
            if has_keepers == "Igen":
                selected = st.multiselect(
                    "Kapusok kiválasztása",
                    players_all,
                    default=st.session_state.get("manual_keeper_players_no_pos", []),
                    key="manual_keeper_players_no_pos",
                    help="A kiválasztott játékosokra kapus-specifikus súlyozás kerül."
                )
                out["is_goalkeeper"] = out["player_name"].astype(str).isin(selected)
                if not selected:
                    st.warning("Válaszd ki a kapus(oka)t, különben minden játékos mezőnyjátékosként kerül értékelésre.")
            else:
                out["is_goalkeeper"] = False
        st.caption("A csapatszintű meccs-edzés összevetésben a mezőnyjátékos-profil az alap; kapusokra külön súlyozott benchmark kerül.")
    return finalize_exposure_columns(out)

def finalize_exposure_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Játékpercek és terhelési exposure normalizálása.
    Fontos: egy meccsen 14 játékos szereplése nem 14x90 perc. A számítás a tényleges
    játékospercet használja, ha van; különben meccsnél a row szintű duration_min lesz az alap.
    Edzésnél a résztvevői percre és az egy főre jutó terhelésre normalizál.
    """
    if df is None or df.empty:
        return df
    out = df.copy()
    if "is_goalkeeper" not in out.columns:
        out = add_keeper_flag_from_position(out)
    if "duration_min" not in out.columns:
        out["duration_min"] = np.nan
    if "match_minutes" in out.columns:
        out["match_minutes"] = out["match_minutes"].apply(duration_to_minutes)
    else:
        out["match_minutes"] = np.nan
    # player_minutes: meccsnél tényleges játékperc, ha van; ha nincs, duration_min. Edzésnél duration_min.
    out["player_minutes"] = np.where(
        out.get("session_type", "").astype(str).eq("Meccs"),
        out["match_minutes"].where(out["match_minutes"].notna() & (out["match_minutes"] > 0), out["duration_min"]),
        out["duration_min"],
    )
    out["player_minutes"] = pd.to_numeric(out["player_minutes"], errors="coerce")
    out.loc[out["player_minutes"] <= 0, "player_minutes"] = np.nan
    group_cols = ["week", "session_date", "session_type"]
    if "session_name" in out.columns:
        group_cols.append("session_name")
    safe_group_cols = [c for c in group_cols if c in out.columns]
    if safe_group_cols:
        out["session_player_count"] = out.groupby(safe_group_cols)["player_name"].transform("nunique")
        out["field_player_count"] = out[~out["is_goalkeeper"]].groupby(safe_group_cols)["player_name"].transform("nunique")
        out["team_player_minutes"] = out.groupby(safe_group_cols)["player_minutes"].transform("sum")
        field_minutes = out.loc[~out["is_goalkeeper"]].groupby(safe_group_cols)["player_minutes"].sum().rename("field_team_minutes")
        out = out.merge(field_minutes.reset_index(), on=safe_group_cols, how="left")
    else:
        out["session_player_count"] = out["player_name"].nunique() if "player_name" in out.columns else np.nan
        out["field_player_count"] = out.loc[~out["is_goalkeeper"], "player_name"].nunique() if "player_name" in out.columns else np.nan
        out["team_player_minutes"] = out["player_minutes"].sum()
        out["field_team_minutes"] = out.loc[~out["is_goalkeeper"], "player_minutes"].sum()
    out["keeper_weight"] = 1.0
    for metric, weight in KEEPER_BENCHMARK_WEIGHTS.items():
        if metric in out.columns:
            out[f"{metric}_keeper_adjusted"] = np.where(out["is_goalkeeper"], out[metric] / max(weight, 0.01), out[metric])
    for metric in FIELD_METRICS_FOR_EXPOSURE:
        if metric in out.columns:
            out[f"{metric}_per90"] = np.where(out["player_minutes"] > 0, out[metric] / out["player_minutes"] * 90, np.nan)
            out[f"{metric}_per_player"] = np.where(out["session_player_count"] > 0, out[metric] / out["session_player_count"], np.nan)
            out[f"{metric}_per_field_player"] = np.where(out["field_player_count"] > 0, out[metric] / out["field_player_count"], np.nan)
    return out

def field_players_only(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    if "is_goalkeeper" in df.columns:
        return df[~df["is_goalkeeper"].fillna(False)].copy()
    return df.copy()

def aggregate_weekly(df: pd.DataFrame) -> pd.DataFrame:
    # Csapatszintű összevetésnél a mezőnyjátékos-profil az alap, hogy a kapusok ne húzzák le
    # a sprint/HSR/intenzitás mutatókat. A kapusok külön szerepként maradnak a játékosszintű elemzésben.
    work = field_players_only(finalize_exposure_columns(df)) if df is not None and not df.empty else pd.DataFrame()
    if work.empty:
        return pd.DataFrame()
    agg_map = {
        "total_distance": "sum",
        "training_load": "sum",
        "muscle_load": "sum",
        "hsr_distance": "sum",
        "sprint_distance": "sum",
        "sprints": "sum",
        "acc_count": "sum",
        "dec_count": "sum",
        "high_efforts": "sum",
        "duration_min": "sum",
        "player_minutes": "sum",
        "team_player_minutes": "max",
        "field_team_minutes": "max",
        "session_player_count": "max",
        "field_player_count": "max",
        "distance_per_min": "mean",
        "max_speed": "max",
        "hr_avg": "mean",
        "hrv": "mean",
    }
    usable = {k: v for k, v in agg_map.items() if k in work.columns}
    res = work.groupby(["week", "session_type"], as_index=False).agg(usable)
    # Per90 / per participant normalizált csapatmutatók: különösen meccs-edzés arányhoz.
    for metric in ["total_distance", "hsr_distance", "sprint_distance", "sprints", "high_efforts", "training_load"]:
        if metric in res.columns and "field_team_minutes" in res.columns:
            res[f"{metric}_per90_team"] = np.where(res["field_team_minutes"] > 0, res[metric] / res["field_team_minutes"] * 90, np.nan)
        if metric in res.columns and "field_player_count" in res.columns:
            res[f"{metric}_per_field_player"] = np.where(res["field_player_count"] > 0, res[metric] / res["field_player_count"], np.nan)
    return res


def player_weekly(df: pd.DataFrame) -> pd.DataFrame:
    work = finalize_exposure_columns(df) if df is not None and not df.empty else pd.DataFrame()
    if work.empty:
        return pd.DataFrame()
    agg_map = {
        "total_distance": "sum",
        "training_load": "sum",
        "muscle_load": "sum",
        "hsr_distance": "sum",
        "sprint_distance": "sum",
        "sprints": "sum",
        "acc_count": "sum",
        "dec_count": "sum",
        "high_efforts": "sum",
        "duration_min": "sum",
        "player_minutes": "sum",
        "distance_per_min": "mean",
        "max_speed": "max",
        "hr_avg": "mean",
        "hrv": "mean",
        "is_goalkeeper": "max",
    }
    usable = {k: v for k, v in agg_map.items() if k in work.columns}
    res = work.groupby(["player_name", "week", "session_type"], as_index=False).agg(usable)
    for metric in ["total_distance", "hsr_distance", "sprint_distance", "sprints", "high_efforts", "training_load", "acc_count", "dec_count"]:
        if metric in res.columns and "player_minutes" in res.columns:
            res[f"{metric}_per90"] = np.where(res["player_minutes"] > 0, res[metric] / res["player_minutes"] * 90, np.nan)
    return res


def pct_change(current: float, previous: float) -> Optional[float]:
    if previous is None or pd.isna(previous) or previous == 0 or pd.isna(current):
        return None
    return (current - previous) / previous


def available_metric_options(df: pd.DataFrame, desired: List[str]) -> List[str]:
    return [m for m in desired if m in df.columns and not df[m].isna().all()]


def metric_name(metric: str) -> str:
    return METRIC_LABELS.get(metric, metric)


def format_week_label(week_value: object) -> str:
    """Rövid, dashboard-barát hétfelirat.
    A pandas periodből érkező '2026-06-01/2026-06-07' helyett pl. '2026-W23'.
    """
    text = str(week_value or "")
    try:
        if "/" in text:
            start = pd.to_datetime(text.split("/")[0], errors="coerce")
            if pd.notna(start):
                iso = start.isocalendar()
                return f"{int(iso.year)}-W{int(iso.week):02d}"
        dt = pd.to_datetime(text, errors="coerce")
        if pd.notna(dt):
            iso = dt.isocalendar()
            return f"{int(iso.year)}-W{int(iso.week):02d}"
    except Exception:
        pass
    return text


def team_insights(df: pd.DataFrame, selected_week: str) -> List[Insight]:
    insights: List[Insight] = []
    weeks = sorted(df["week"].dropna().unique().tolist())
    week_idx = weeks.index(selected_week) if selected_week in weeks else len(weeks) - 1
    previous_week = weeks[week_idx - 1] if week_idx > 0 else None

    current = df[df["week"] == selected_week]
    prev = df[df["week"] == previous_week] if previous_week else pd.DataFrame()
    cur_train = current[current["session_type"] == "Edzés"]
    cur_match = current[current["session_type"] == "Meccs"]
    prev_train = prev[prev["session_type"] == "Edzés"] if not prev.empty else pd.DataFrame()

    # 1. Sprint underload
    if not cur_train.empty and not cur_match.empty and "sprint_distance" in df.columns:
        train_sprint = cur_train["sprint_distance"].mean()
        match_sprint = cur_match["sprint_distance"].mean()
        if pd.notna(match_sprint) and match_sprint > 0:
            ratio = train_sprint / match_sprint
            if ratio < 0.65:
                insights.append(Insight(
                    "Alacsony sprintterhelés", "KRITIKUS",
                    f"Az edzések átlagos sprintterhelése a meccsterhelés kb. {ratio:.0%}-a.",
                    "A nagy intenzitású terhelés jelentősen elmaradhat attól, amit a mérkőzés megkövetel.",
                    "Érdemes célzott sprint- vagy nagysebességű blokkokat beépíteni, ha ez illeszkedik a hét terhelési tervéhez.",
                ))
            elif ratio < 0.80:
                insights.append(Insight(
                    "Alacsony sprintterhelés", "FIGYELMEZTETÉS",
                    f"Az edzések átlagos sprintterhelése a meccsterhelés kb. {ratio:.0%}-a.",
                    "A meccsigényhez képest mérsékelt intenzitási hiány látható.",
                    "Érdemes ellenőrizni, hogy tudatos visszaterhelésről vagy nem kívánt intenzitáshiányról van-e szó.",
                ))

    # 2. Weekly load change
    if not cur_train.empty and not prev_train.empty and "training_load" in df.columns:
        cur_load = cur_train["training_load"].sum()
        prev_load = prev_train["training_load"].sum()
        chg = pct_change(cur_load, prev_load)
        if chg is not None:
            if chg > 0.25:
                insights.append(Insight(
                    "Heti terhelési kiugrás", "FIGYELMEZTETÉS",
                    f"Az edzés terhelési pontértéke {chg:.0%}-kal nőtt az előző héthez képest.",
                    "A hirtelen terhelésemelkedés ronthatja a frissességet és növelheti a következő napok terhelési kockázatát.",
                    "Érdemes figyelni a következő edzések intenzitását, illetve a játékosok egyéni reakcióit.",
                ))
            elif chg < -0.25:
                insights.append(Insight(
                    "Heti terheléscsökkenés", "INFORMÁCIÓ",
                    f"Az edzés terhelési pontértéke {abs(chg):.0%}-kal csökkent az előző héthez képest.",
                    "Ez lehet tudatos frissítés, de lehet nem kívánt terhelésvesztés is.",
                    "Érdemes kontextusba helyezni: meccs előtti könnyítés, hiányzók vagy edzéstartalom-váltás okozta-e.",
                ))

    # 3. Match intensity gap
    if not cur_train.empty and not cur_match.empty and "distance_per_min" in df.columns:
        train_int = cur_train["distance_per_min"].mean()
        match_int = cur_match["distance_per_min"].mean()
        if pd.notna(match_int) and match_int > 0:
            ratio = train_int / match_int
            if ratio < 0.85:
                insights.append(Insight(
                    "Meccsintenzitási eltérés", "FIGYELMEZTETÉS",
                    f"Az edzések átlagos táv/perc értéke a meccs kb. {ratio:.0%}-a.",
                    "A csapat edzésintenzitása elmaradhat a mérkőzés tempójától.",
                    "Érdemes lehet rövidebb, intenzívebb játékszituációkat vagy tudatos tempóváltásokat használni.",
                ))

    # 4. High deceleration load
    if not cur_train.empty and not prev_train.empty and "dec_count" in df.columns:
        cur_dec = cur_train["dec_count"].mean()
        prev_dec = prev_train["dec_count"].mean()
        chg = pct_change(cur_dec, prev_dec)
        if chg is not None and chg > 0.35:
            insights.append(Insight(
                "Magas lassítási terhelés", "FIGYELMEZTETÉS",
                f"Az átlagos lassításszám {chg:.0%}-kal nőtt az előző héthez képest.",
                "A lassítások jelentős neuromuszkuláris terhelést jelenthetnek.",
                "Érdemes figyelni a regenerációra és a következő edzés excentrikus terhelésére.",
            ))

    # 5. Max speed suppression
    if not cur_train.empty and not prev_train.empty and "max_speed" in df.columns:
        cur_speed = cur_train["max_speed"].max()
        prev_speed = prev_train["max_speed"].max()
        chg = pct_change(cur_speed, prev_speed)
        if chg is not None and chg < -0.05:
            insights.append(Insight(
                "Maximális sebesség visszaesése", "INFORMÁCIÓ",
                f"A heti legmagasabb maximális sebesség {abs(chg):.0%}-kal alacsonyabb az előző hétnél.",
                "Ez jelezhet alacsonyabb neuromuszkuláris frissességet, de lehet edzéstartalom-függő is.",
                "Érdemes megnézni, volt-e valódi maximális sebességű inger a héten.",
            ))

    # 6. Player outlier alerts
    if not cur_train.empty and "training_load" in cur_train.columns:
        player_load = cur_train.groupby("player_name")["training_load"].sum().dropna()
        if len(player_load) >= 5 and player_load.mean() > 0:
            mean = player_load.mean()
            high = player_load[player_load > mean * 1.25].sort_values(ascending=False)
            low = player_load[player_load < mean * 0.75].sort_values(ascending=True)
            if len(high) > 0:
                names = ", ".join(high.head(3).index.tolist())
                insights.append(Insight(
                    "Kiugró játékosterhelés", "FIGYELMEZTETÉS",
                    f"Néhány játékos jelentősen a csapatátlag felett terhelődött: {names}.",
                    "Egyéni terheléskülönbség alakult ki a héten.",
                    "Érdemes egyéni szinten ránézni a következő edzés terhelésére és a játékpercekre.",
                    scope="Játékos",
                ))
            if len(low) > 0:
                names = ", ".join(low.head(3).index.tolist())
                insights.append(Insight(
                    "Alacsony játékosterhelés", "INFORMÁCIÓ",
                    f"Néhány játékos jelentősen a csapatátlag alatt terhelődött: {names}.",
                    "Terheléslemaradás alakulhat ki, főleg ha ez több héten át fennáll.",
                    "Érdemes ellenőrizni a hiányzásokat, játékperceket és az egyéni kiegészítő munkát.",
                    scope="Játékos",
                ))

    if not insights:
        insights.append(Insight(
            "Stabil hét", "INFORMÁCIÓ",
            "Nem látható kiemelt negatív eltérés az aktuális hét fő mutatóiban.",
            "A csapat terhelési profilja stabilnak tűnik az elérhető adatok alapján.",
            "Érdemes tovább figyelni a sprint- és intenzitási trendeket, különösen meccs előtti héten.",
        ))

    return sorted(insights, key=lambda x: SEVERITY_RANK.get(x.severity, 9))[:8]



# -----------------------------------------------------------------------------
# V2 - Football Intelligence Layer
# -----------------------------------------------------------------------------
def pdf_safe_text(text: object) -> str:
    """PDF exporthoz stabil magyar szöveg.
    Cloud környezetben a hosszú ő/ű néha hibásan jelenik meg, ezért PDF-ben
    rövid párra normalizáljuk. UI, Word és Excel exportban marad az eredeti.
    """
    if text is None or (isinstance(text, float) and pd.isna(text)):
        return ""
    text = str(text)
    replacements = {"ő": "ö", "Ő": "Ö", "ű": "ü", "Ű": "Ü"}
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def day_label_from_delta(delta_days: int) -> str:
    if delta_days == 0:
        return "MD"
    if delta_days < 0:
        return f"MD{delta_days}"
    return f"MD+{delta_days}"


def detect_match_day(week_df: pd.DataFrame) -> Optional[pd.Timestamp]:
    if week_df.empty:
        return None
    tmp = week_df.copy()
    tmp["session_date_dt"] = pd.to_datetime(tmp["session_date"], errors="coerce")
    matches = tmp[tmp["session_type"] == "Meccs"]
    if not matches.empty:
        # Ha több meccs van, a hét fő meccsének az utolsót vesszük.
        return pd.to_datetime(matches["session_date_dt"].max())
    return None


def build_microcycle_table(df: pd.DataFrame, selected_week: str) -> pd.DataFrame:
    week_df = df[df["week"] == selected_week].copy()
    if week_df.empty:
        return pd.DataFrame()
    match_day = detect_match_day(week_df)
    week_df["session_date_dt"] = pd.to_datetime(week_df["session_date"], errors="coerce")

    agg_spec = {
        "training_load": "sum",
        "total_distance": "sum",
        "sprint_distance": "sum",
        "hsr_distance": "sum",
        "distance_per_min": "mean",
        "max_speed": "max",
        "acc_count": "sum",
        "dec_count": "sum",
        "high_efforts": "sum",
        "player_name": "nunique",
    }
    usable = {k: v for k, v in agg_spec.items() if k in week_df.columns}
    daily = week_df.groupby(["session_date_dt", "session_type"], as_index=False).agg(usable)
    daily = daily.rename(columns={"player_name": "játékosok"})
    if match_day is not None:
        daily["md_delta"] = (daily["session_date_dt"] - match_day).dt.days
        daily["md_label"] = daily["md_delta"].apply(day_label_from_delta)
    else:
        daily["md_delta"] = np.nan
        daily["md_label"] = "Nincs meccs"

    # Napi load fallback: ahol nincs training_load, ott total_distance alapján értékelünk.
    if "training_load" in daily.columns and not daily["training_load"].isna().all():
        daily["load_index"] = daily["training_load"]
        daily["load_index_label"] = "Terhelési pont"
    elif "total_distance" in daily.columns:
        daily["load_index"] = daily["total_distance"]
        daily["load_index_label"] = "Össztáv"
    else:
        daily["load_index"] = np.nan
        daily["load_index_label"] = "Nincs load mutató"
    return daily.sort_values("session_date_dt")


def microcycle_insights(df: pd.DataFrame, selected_week: str) -> List[Insight]:
    insights: List[Insight] = []
    daily = build_microcycle_table(df, selected_week)
    if daily.empty:
        return insights
    if "MD" not in daily["md_label"].values:
        insights.append(Insight(
            "Mikrociklus kontextus hiányzik", "INFORMÁCIÓ",
            "Az aktuális héten nem található meccs típusú session, ezért az MD-napok nem értelmezhetők automatikusan.",
            "Meccsnap nélkül a heti struktúrát csak általános terhelési trendként lehet értékelni.",
            "Ha van mérkőzés, érdemes a session típusát 'Meccs'-ként jelölni, hogy a rendszer MD-1 / MD-2 / MD-3 logikával is tudjon gondolkodni.",
            scope="Mikrociklus",
        ))
        return insights

    train = daily[daily["session_type"] == "Edzés"].copy()
    match = daily[daily["session_type"] == "Meccs"].copy()

    # Maximális sebességű inger: történt-e értelmezhető sprintinger a meccs előtti napokban?
    if not train.empty and "sprint_distance" in daily.columns and not match.empty:
        match_sprint = match["sprint_distance"].mean()
        max_training_sprint = train["sprint_distance"].max()
        if pd.notna(match_sprint) and match_sprint > 0:
            ratio = max_training_sprint / match_sprint if pd.notna(max_training_sprint) else 0
            if ratio < 0.15:
                insights.append(Insight(
                    "Hiányzó maximális sebességű inger", "KRITIKUS",
                    f"A héten nem látszik érdemi sprintinger: a legmagasabb edzésnapi sprintterhelés a meccs kb. {ratio:.0%}-a.",
                    "Meccsigényhez képest alacsony lehetett a maximális sebességű inger, ami a felkészítés minőségét ronthatja.",
                    "Érdemes lehet egy rövid, kontrollált maximális sebességű inger blokkot betervezni a megfelelő napon, ha a heti cél és a játékosállapot engedi.",
                    scope="Mikrociklus",
                ))
            elif ratio < 0.30:
                insights.append(Insight(
                    "Alacsony maximális sebességű inger", "FIGYELMEZTETÉS",
                    f"A legmagasabb edzésnapi sprintterhelés a meccs kb. {ratio:.0%}-a.",
                    "Volt sprintinger, de a meccsigényhez képest visszafogott lehetett.",
                    "Érdemes ellenőrizni, hogy ez tudatos frissítés vagy nem kívánt intenzitáshiány volt-e.",
                    scope="Mikrociklus",
                ))

    # MD-2 túl magas load
    md2 = daily[daily["md_label"] == "MD-2"]
    md3_4 = daily[daily["md_label"].isin(["MD-3", "MD-4"])]
    if not md2.empty and not md3_4.empty:
        md2_load = md2["load_index"].sum()
        peak_early = md3_4["load_index"].max()
        if pd.notna(md2_load) and pd.notna(peak_early) and peak_early > 0 and md2_load > peak_early * 0.80:
            insights.append(Insight(
                "Magas MD-2 terhelés", "FIGYELMEZTETÉS",
                "Az MD-2 nap terhelése közel volt a hét korábbi fő terhelési napjához.",
                "A mérkőzés előtti 48 órában a túl magas load ronthatja a frissességet.",
                "Érdemes lehet az MD-2 napot jobban kontrollálni, és a fő terhelési ingert inkább MD-3 / MD-4 környékére helyezni.",
                scope="Mikrociklus",
            ))

    # MD-1 activation kontroll
    md1 = daily[daily["md_label"] == "MD-1"]
    if not md1.empty and not md2.empty:
        md1_load = md1["load_index"].sum()
        md2_load = md2["load_index"].sum()
        if pd.notna(md1_load) and pd.notna(md2_load) and md2_load > 0 and md1_load > md2_load * 0.85:
            insights.append(Insight(
                "MD-1 activation túl erős lehet", "FIGYELMEZTETÉS",
                "Az MD-1 load nem csökkent érdemben az MD-2 naphoz képest.",
                "A meccs előtti aktiváció célja általában a frissítés, nem egy újabb nagy terhelési inger.",
                "Érdemes lehet az MD-1 napot rövidebb, frissebb, idegrendszeri aktivációs jelleggel tartani.",
                scope="Mikrociklus",
            ))

    # Pozitív taper insight, ha van adat és nincs túl magas MD-1/MD-2
    if not md1.empty and not md3_4.empty:
        md1_load = md1["load_index"].sum()
        peak_early = md3_4["load_index"].max()
        if pd.notna(md1_load) and pd.notna(peak_early) and peak_early > 0 and md1_load < peak_early * 0.45:
            insights.append(Insight(
                "Megfelelő tapering jel", "INFORMÁCIÓ",
                "A meccs előtti nap terhelése jelentősen alacsonyabb volt a hét fő terhelési napjához képest.",
                "Ez támogathatja a mérkőzésnapi frissességet.",
                "Érdemes megtartani ezt a struktúrát, ha a mérkőzésnapi teljesítmény is visszaigazolja.",
                scope="Mikrociklus",
            ))

    return insights


def playstyle_insights(df: pd.DataFrame, selected_week: str, playstyle: str) -> List[Insight]:
    insights: List[Insight] = []
    if playstyle == "Kiegyensúlyozott":
        return insights
    current = df[df["week"] == selected_week]
    train = current[current["session_type"] == "Edzés"]
    match = current[current["session_type"] == "Meccs"]
    if train.empty:
        return insights

    def mean_ratio(metric: str) -> Optional[float]:
        if metric not in current.columns or match.empty:
            return None
        m = match[metric].mean()
        t = train[metric].mean()
        if pd.isna(m) or m == 0 or pd.isna(t):
            return None
        return t / m

    intensity_ratio = mean_ratio("distance_per_min")
    sprint_ratio = mean_ratio("sprint_distance")
    effort_ratio = mean_ratio("high_efforts")

    if playstyle == "Pressing":
        if intensity_ratio is not None and intensity_ratio < 0.90:
            insights.append(Insight(
                "Pressing profil: intenzitáshiány", "FIGYELMEZTETÉS",
                f"A heti edzésintenzitás a meccsintenzitás kb. {intensity_ratio:.0%}-a.",
                "Pressing játékmodellnél fontos, hogy a csapat rendszeresen találkozzon magas munkasűrűségű helyzetekkel.",
                "Érdemes lehet rövidebb, nagyobb nyomású játékokat vagy pressing-specifikus blokkokat használni.",
                scope="Játékmodell",
            ))
        if effort_ratio is not None and effort_ratio < 0.75:
            insights.append(Insight(
                "Pressing profil: kevés nagy intenzitású erőfeszítés", "FIGYELMEZTETÉS",
                f"A high effort profil a meccsigény kb. {effort_ratio:.0%}-a.",
                "A sok gyorsulás, lassítás és ismételt intenzív akció kulcs a pressing identitáshoz.",
                "Érdemes lehet növelni az ismételt intenzív akciókat tartalmazó gyakorlatok arányát.",
                scope="Játékmodell",
            ))

    elif playstyle == "Transition":
        if sprint_ratio is not None and sprint_ratio < 0.80:
            insights.append(Insight(
                "Transition profil: alacsony sprintinger", "FIGYELMEZTETÉS",
                f"A heti sprintprofil a meccsigény kb. {sprint_ratio:.0%}-a.",
                "Átmenetekre építő játékmodellben fontos a rendszeres nagysebességű inger.",
                "Érdemes célzott átmeneti játékokat vagy nagy területű sprinthelyzeteket beépíteni.",
                scope="Játékmodell",
            ))

    elif playstyle == "Possession":
        if intensity_ratio is not None and intensity_ratio > 1.15:
            insights.append(Insight(
                "Possession profil: magas edzésintenzitás", "INFORMÁCIÓ",
                f"Az edzésintenzitás meghaladta a meccsintenzitást: kb. {intensity_ratio:.0%}.",
                "Ez lehet tudatos túlterhelés, de possession modellnél a kontrollált terhelés is fontos.",
                "Érdemes ellenőrizni, hogy az intenzitás a játékelvek tanulását vagy inkább csak a fizikai terhelést szolgálta-e.",
                scope="Játékmodell",
            ))

    elif playstyle == "Low Block":
        if sprint_ratio is not None and sprint_ratio < 0.60:
            insights.append(Insight(
                "Low block profil: kevés átmeneti sprintinger", "FIGYELMEZTETÉS",
                f"A sprintprofil a meccsigény kb. {sprint_ratio:.0%}-a.",
                "Mélyebb védekezésnél is fontosak lehetnek a rövid, robbanékony átmenetek.",
                "Érdemes lehet kontrollált kontra- és visszarendeződési szituációkat alkalmazni.",
                scope="Játékmodell",
            ))
    return insights



def html_linebreaks(text: object) -> str:
    """Streamlit markdown/HTML megjelenítéshez: a literal \\n karaktereket valódi sortöréssé alakítja."""
    safe = str(text or "").replace("\\n", "\n").replace("\\r", "")
    return html.escape(safe).replace("\n", "<br>")


def week_label_short(week_value: object) -> str:
    """Hosszú pandas week labelből rövid, dashboard-barát hét címke."""
    txt = str(week_value or "")
    if "/" in txt:
        try:
            start = pd.to_datetime(txt.split("/")[0])
            return f"{start.strftime('%Y')}-W{int(start.isocalendar().week):02d}"
        except Exception:
            return txt.split("/")[0]
    return txt


def render_fpi_hero() -> None:
    st.markdown(
        """
        <style>
        .fpi-hero-wow{
            background:linear-gradient(135deg,#f8fafc 0%,#e0f2fe 55%,#ecfdf5 100%) !important;
            color:#0f172a !important;
            border:1px solid #bfdbfe !important;
            box-shadow:0 22px 60px rgba(15,23,42,.14) !important;
        }
        .fpi-hero-wow h1{color:#0f172a !important;}
        .fpi-hero-wow p{color:#334155 !important;}
        .fpi-hero-wow *{color:#0f172a !important;}
        .fpi-chip-wow{background:#ffffff !important;border:1px solid #bfdbfe !important;color:#0f172a !important;}
        .fpi-top-nav-card{background:#ffffff;border:1px solid #dbeafe;border-radius:22px;padding:14px 16px;margin:8px 0 14px 0;box-shadow:0 10px 28px rgba(15,23,42,.08);}
        .fpi-top-nav-card b{color:#0f172a;}
        .fpi-top-nav-card span{color:#475569;}
        </style>
        <div class="fpi-hero-wow">
          <h1>⚽ Football Performance Intelligence</h1>
          <p>GPS adatokból és opcionális taktikai PDF/Excel inputból vezetői döntéstámogatás, edzői prioritások és exportálható performance riportok.</p>
          <div class="fpi-chip-row">
            <span class="fpi-chip-wow">⚡ Readiness</span>
            <span class="fpi-chip-wow">🎯 Coaching priorities</span>
            <span class="fpi-chip-wow">🧠 Smart mapper</span>
            <span class="fpi-chip-wow">📄 Executive export</span>
            <span class="fpi-chip-wow">🧩 Tactical Pro+ opcionális</span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )



def clean_literal_newlines_for_display(text: object) -> str:
    raw = str(text or "")
    raw = raw.replace("\\r", "")
    raw = raw.replace("\\n", "\n")
    return raw


def render_leader_summary_direct(text: object) -> None:
    """Fixen olvasható vezetői összefoglaló fehér kártyán, sötét betűkkel."""
    raw = clean_literal_newlines_for_display(text)
    lines = [ln.strip().lstrip("-• ").strip() for ln in raw.splitlines() if ln.strip()]

    # Biztonsági javítás: ha a Játékmodell sorban maradtak benne literal \n részek.
    normalized = []
    for ln in lines:
        if "\\n" in ln:
            normalized.extend([p.strip().lstrip("-• ").strip() for p in ln.replace("\\n", "\n").splitlines() if p.strip()])
        else:
            normalized.append(ln)
    lines = normalized

    html_rows = []
    for ln in lines:
        esc = html.escape(ln)
        if esc.startswith("Hét:"):
            label, value = "📅 Hét", esc.replace("Hét:", "").strip()
        elif esc.startswith("Legfontosabb üzenet:"):
            label, value = "⚡ Legfontosabb üzenet", esc.replace("Legfontosabb üzenet:", "").strip()
        elif esc.startswith("Mit látunk?"):
            label, value = "🔎 Mit látunk?", esc.replace("Mit látunk?", "").strip()
        elif esc.startswith("Javaslat:"):
            label, value = "🎯 Javaslat", esc.replace("Javaslat:", "").strip()
        elif esc.startswith("Második fontos téma:"):
            label, value = "➕ Második fontos téma", esc.replace("Második fontos téma:", "").strip()
        elif esc.startswith("Játékmodell:"):
            label, value = "♟️ Játékmodell", esc.replace("Játékmodell:", "").strip()
        elif esc.startswith("Meccskészültség:"):
            label, value = "🚀 Meccskészültség", esc.replace("Meccskészültség:", "").strip()
        elif esc.startswith("Periodizációs besorolás:"):
            label, value = "📈 Periodizáció", esc.replace("Periodizációs besorolás:", "").strip()
        else:
            label, value = "", esc

        if label:
            html_rows.append(
                f"""
                <div style="display:grid;grid-template-columns:230px 1fr;gap:14px;padding:10px 0;border-bottom:1px solid #e2e8f0;">
                    <div style="color:#1e3a8a!important;font-weight:950!important;">{label}</div>
                    <div style="color:#0f172a!important;font-weight:800!important;line-height:1.45;">{value}</div>
                </div>
                """
            )
        else:
            html_rows.append(
                f"""
                <div style="padding:8px 0;color:#0f172a!important;font-weight:800!important;line-height:1.45;">{value}</div>
                """
            )

    st.markdown(
        f"""
        <div style="
            background:#ffffff!important;
            color:#0f172a!important;
            border:1px solid #bfdbfe;
            border-radius:22px;
            padding:20px 22px;
            margin:12px 0 22px 0;
            box-shadow:0 18px 44px rgba(15,23,42,.22);
        ">
            <div style="font-size:1.25rem;font-weight:950;color:#0f172a!important;margin-bottom:10px;">
                Heti vezetői összefoglaló
            </div>
            {''.join(html_rows)}
        </div>
        """,
        unsafe_allow_html=True,
    )






def create_sample_input_template_bytes() -> Optional[bytes]:
    """Minta Excel sablon: Adatok + Használati útmutató.
    Csak openpyxl-t használ, hogy Streamlit Cloudon ne kelljen xlsxwriter.
    """
    try:
        import io
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
        from openpyxl.worksheet.table import Table, TableStyleInfo
    except Exception:
        return None

    output = io.BytesIO()

    columns = [
        "Játékos neve",
        "Típus",
        "Szakasz neve",
        "Poszt",
        "Kezdési idő",
        "Befejezési idő",
        "Időtartam",
        "Teljes táv [m]",
        "Táv/perc [m/min]",
        "Maximális sebesség [km/h]",
        "Sprintek",
        "Táv a sebesség célzónában 4 [m] (19.80 - 24.99 km/h)",
        "Táv a sebesség célzónában 5 [m] (25.00- km/h)",
        "Edzési terhelési pontérték",
        "Izomterhelés",
        "Átlagos pulzus [bpm]",
        "Maximális pulzus [bpm]",
        "HRV (RMSSD)",
        "Gyorsulások száma (2.50 - 2.99 m/s²)",
        "Gyorsulások száma (3.00 - 50.00 m/s²)",
        "Gyorsulások száma (-2.99 - -2.50 m/s²)",
        "Gyorsulások száma (-50.00 - -3.00 m/s²)",
    ]

    sample_rows = [
        [
            "Minta Játékos 1", "Edzés", "MD-4 edzés", "CM",
            "2026-06-01 10:00", "2026-06-01 11:20", "01:20:00",
            7200, 90, 29.4, 12, 620, 180, 410, 58, 148, 188, 62, 22, 8, 20, 7
        ],
        [
            "Minta Játékos 1", "Meccs", "Bajnoki mérkőzés", "CM",
            "2026-06-07 17:00", "2026-06-07 18:45", "01:45:00",
            10600, 101, 31.2, 24, 1050, 690, 620, 84, 162, 194, 55, 34, 14, 31, 12
        ],
        [
            "Minta Játékos 2", "Edzés", "MD-4 edzés", "W",
            "2026-06-01 10:00", "2026-06-01 11:20", "01:20:00",
            7600, 95, 30.7, 15, 760, 260, 445, 63, 151, 190, 59, 25, 10, 22, 9
        ],
        [
            "Minta Játékos 2", "Meccs", "Bajnoki mérkőzés", "W",
            "2026-06-07 17:00", "2026-06-07 18:45", "01:45:00",
            11100, 106, 32.0, 29, 1200, 760, 650, 88, 165, 197, 53, 38, 16, 33, 13
        ],
    ]

    guide_rows = [
        ["Téma", "Leírás"],
        ["Cél", "Ez a sablon mutatja, milyen szerkezetű Excel tölthető fel a Performance Intelligence appba."],
        ["Fontos", "Az edzés- és meccsadatokat EGYMÁS ALÁ kell halmozni ugyanazon az Adatok munkalapon."],
        ["Egy sor jelentése", "Egy játékos egy edzésen vagy mérkőzésen mért adata."],
        ["Típus oszlop", "Csak ilyen értékeket használj: Edzés vagy Meccs."],
        ["Dátum/idő", "A Kezdési idő lehet például: 2026-06-01 10:00."],
        ["Időtartam", "Javasolt formátum: 01:20:00 vagy percben megadott szám."],
        ["Kötelező mezők", "Játékos neve, Típus, Kezdési idő."],
        ["Ajánlott fő mezők", "Teljes táv, Táv/perc, Maximális sebesség, Sprintek, sebességzónák, terhelési pont, gyorsulás/lassítás."],
        ["Ha más a GPS-export fejléce", "A Smart Excel Mapper segít felismerni és kézzel összepárosítani az oszlopokat."],
        ["Tipp", "Ne készíts külön munkalapot edzésenként. Minden esemény menjen az Adatok munkalapra egymás alá."],
    ]

    wb = Workbook()
    ws = wb.active
    ws.title = "Adatok"
    guide = wb.create_sheet("Használati útmutató")

    ws.append(columns)
    for row in sample_rows:
        ws.append(row)

    guide.append(guide_rows[0])
    for row in guide_rows[1:]:
        guide.append(row)

    # Styles
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="1E3A8A")
    guide_fill = PatternFill("solid", fgColor="166534")
    white_font = Font(bold=True, color="FFFFFF")
    dark_font = Font(color="0F172A")
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(wrap_text=True, vertical="center", horizontal="center")

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = white_font
        cell.alignment = center
        cell.border = border

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = dark_font
            cell.alignment = wrap
            cell.border = border

    for cell in guide[1]:
        cell.fill = guide_fill
        cell.font = white_font
        cell.alignment = center
        cell.border = border

    for row in guide.iter_rows(min_row=2):
        for cell in row:
            cell.font = dark_font
            cell.alignment = wrap
            cell.border = border

    # Widths
    ws.column_dimensions["A"].width = 20
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 20
    ws.column_dimensions["D"].width = 10
    ws.column_dimensions["E"].width = 20
    ws.column_dimensions["F"].width = 20
    ws.column_dimensions["G"].width = 14
    for col in range(8, len(columns) + 1):
        ws.column_dimensions[chr(64 + col) if col <= 26 else "Z"].width = 18

    guide.column_dimensions["A"].width = 26
    guide.column_dimensions["B"].width = 95

    ws.freeze_panes = "A2"
    guide.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    # Table if possible
    try:
        tab = Table(displayName="PerformanceInputTable", ref=ws.dimensions)
        style = TableStyleInfo(name="TableStyleMedium2", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
        tab.tableStyleInfo = style
        ws.add_table(tab)
    except Exception:
        pass

    try:
        tab2 = Table(displayName="PerformanceGuideTable", ref=guide.dimensions)
        style2 = TableStyleInfo(name="TableStyleMedium4", showFirstColumn=False, showLastColumn=False, showRowStripes=True, showColumnStripes=False)
        tab2.tableStyleInfo = style2
        guide.add_table(tab2)
    except Exception:
        pass

    wb.save(output)
    output.seek(0)
    return output.getvalue()


def render_weekly_summary_card(summary_text: object) -> None:
    """Fixen olvasható heti vezetői összefoglaló.
    Fehér háttér + fekete szöveg + valódi sortörések.
    """
    raw = str(summary_text or "")
    raw = raw.replace("\\r", "").replace("\\n", "\n")
    raw = raw.replace("\r", "")
    safe = html.escape(raw)

    st.markdown(
        f"""
        <div style="
            background:#ffffff !important;
            color:#000000 !important;
            border-left:10px solid #2563eb;
            border-radius:18px;
            padding:24px 26px;
            margin:12px 0 22px 0;
            box-shadow:0 10px 28px rgba(0,0,0,.20);
        ">
            <div style="
                color:#000000 !important;
                font-size:1.25rem;
                font-weight:950;
                margin-bottom:14px;
            ">
                Heti vezetői összefoglaló
            </div>
            <pre style="
                color:#000000 !important;
                background:#ffffff !important;
                font-family:Segoe UI, Arial, sans-serif;
                font-size:16px;
                font-weight:700;
                line-height:1.7;
                white-space:pre-wrap;
                margin:0;
            ">{safe}</pre>
        </div>
        """,
        unsafe_allow_html=True,
    )



def build_weekly_summary(insights: List[Insight], selected_week: str, playstyle: str) -> str:
    if not insights:
        return (
            f"Hét: {week_label_short(selected_week)}\n\n"
            "Legfontosabb üzenet: Stabil hét\n"
            "- Mit látunk? Nem látható kiemelt negatív eltérés az aktuális hét fő mutatóiban.\n"
            "- Javaslat: Érdemes tovább figyelni a sprint- és intenzitási trendeket.\n"
            f"- Játékmodell: {playstyle}"
        )

    critical = [i for i in insights if i.severity == "KRITIKUS"]
    warning = [i for i in insights if i.severity == "FIGYELMEZTETÉS"]
    info = [i for i in insights if i.severity == "INFORMÁCIÓ"]
    main = critical[0] if critical else (warning[0] if warning else info[0])

    second = None
    for i in insights:
        if i.title != main.title:
            second = i
            break

    lines = [
        f"Hét: {week_label_short(selected_week)}",
        "",
        f"Legfontosabb üzenet: {main.title}",
        f"- Mit látunk? {main.observation}",
        f"- Javaslat: {main.recommendation}",
    ]
    if second is not None:
        lines.append(f"- Második fontos téma: {second.title}")
    lines.append(f"- Játékmodell: {playstyle}")
    return "\n".join(lines)

def top_coaching_priorities(insights: List[Insight], limit: int = 3) -> List[Dict[str, str]]:
    ordered = sorted(insights, key=lambda x: SEVERITY_RANK.get(x.severity, 9))
    selected = []
    seen = set()
    for ins in ordered:
        if ins.title in seen:
            continue
        seen.add(ins.title)
        selected.append({
            "Cím": ins.title,
            "Súlyosság": ins.severity,
            "Teendő": ins.recommendation,
            "Miért": ins.impact,
        })
        if len(selected) >= limit:
            break
    return selected


def render_coaching_priorities(priorities: List[Dict[str, str]]) -> None:
    if not priorities:
        st.info("Nincs kiemelt edzői teendő az aktuális szűrés alapján.")
        return
    for idx, item in enumerate(priorities, start=1):
        st.markdown(
            f"""
            <div class="priority-card">
                <div style="font-size:1.05rem;font-weight:850;">{idx}. {html.escape(item.get('Cím', ''))}</div>
                <div style="margin-top:6px;"><b>Teendő:</b><br>{html.escape(item.get('Teendő', ''))}</div>
                <div style="margin-top:6px;"><b>Miért:</b><br>{html.escape(item.get('Miért', ''))}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )



# -----------------------------------------------------------------------------
# V2.5 - Performance Memory + Meccskészültség + Adaptive Intelligence
# -----------------------------------------------------------------------------
MEMORY_FILE = Path("performance_memory_v25.csv")


def score_to_label(score: float) -> str:
    if pd.isna(score):
        return "Nincs elég adat"
    if score >= 80:
        return "Jó meccskészültség"
    if score >= 65:
        return "Elfogadható, figyelendő"
    if score >= 50:
        return "Közepes kockázat"
    return "Alacsony readiness"


def score_to_color(score: float) -> str:
    if pd.isna(score):
        return "#94a3b8"
    if score >= 80:
        return "#22c55e"
    if score >= 65:
        return "#84cc16"
    if score >= 50:
        return "#f59e0b"
    return "#ef4444"


def trend_label(values: List[float], tolerance: float = 0.03) -> str:
    clean = [v for v in values if pd.notna(v)]
    if len(clean) < 3:
        return "nincs elég adat"
    first, last = clean[0], clean[-1]
    if first == 0:
        return "nincs elég adat"
    change = (last - first) / abs(first)
    if change > tolerance:
        return "emelkedő"
    if change < -tolerance:
        return "csökkenő"
    return "stabil"


def classify_week_context(weekly_row: pd.Series, history: pd.DataFrame) -> str:
    if history.empty or len(history) < 3:
        return "Tanuló hét"
    load = weekly_row.get("training_load", np.nan)
    sprint = weekly_row.get("sprint_distance", np.nan)
    intensity = weekly_row.get("distance_per_min", np.nan)

    hist_load = history["training_load"].dropna() if "training_load" in history.columns else pd.Series(dtype=float)
    hist_sprint = history["sprint_distance"].dropna() if "sprint_distance" in history.columns else pd.Series(dtype=float)

    if len(hist_load) >= 3 and pd.notna(load):
        load_mean = hist_load.mean()
        if load_mean > 0:
            load_ratio = load / load_mean
            if load_ratio > 1.18:
                return "Terhelésépítő / overload hét"
            if load_ratio < 0.82:
                return "Recovery / alulterhelt hét"

    if len(hist_sprint) >= 3 and pd.notna(sprint):
        sprint_mean = hist_sprint.mean()
        if sprint_mean > 0 and sprint / sprint_mean < 0.75:
            return "Intenzitáshiányos hét"

    if pd.notna(intensity) and "distance_per_min" in history.columns:
        int_hist = history["distance_per_min"].dropna()
        if len(int_hist) >= 3 and int_hist.mean() > 0 and intensity / int_hist.mean() > 1.12:
            return "Intenzitásfókuszú hét"

    return "Normál hét"


def build_weekly_fingerprints(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty or "week" not in df.columns:
        return pd.DataFrame()

    agg_spec = {
        "training_load": "sum",
        "total_distance": "sum",
        "sprint_distance": "sum",
        "hsr_distance": "sum",
        "distance_per_min": "mean",
        "max_speed": "max",
        "acc_count": "sum",
        "dec_count": "sum",
        "high_efforts": "sum",
        "duration_min": "sum",
        "hrv": "mean",
        "player_name": "nunique",
    }
    usable = {k: v for k, v in agg_spec.items() if k in df.columns}
    weekly = df.groupby("week", as_index=False).agg(usable)
    weekly = weekly.rename(columns={"player_name": "players"})

    # Csak edzésekből külön load profil, ha van elég adat.
    train = df[df["session_type"] == "Edzés"].copy()
    if not train.empty:
        train_agg = train.groupby("week", as_index=False).agg({
            k: v for k, v in usable.items() if k != "player_name"
        })
        train_agg = train_agg.add_prefix("train_")
        train_agg = train_agg.rename(columns={"train_week": "week"})
        weekly = weekly.merge(train_agg, on="week", how="left")

    # Meccsprofil külön.
    match = df[df["session_type"] == "Meccs"].copy()
    if not match.empty:
        match_agg = match.groupby("week", as_index=False).agg({
            k: v for k, v in usable.items() if k != "player_name"
        })
        match_agg = match_agg.add_prefix("match_")
        match_agg = match_agg.rename(columns={"match_week": "week"})
        weekly = weekly.merge(match_agg, on="week", how="left")

    weekly = weekly.sort_values("week").reset_index(drop=True)

    # Rolling baseline-ok.
    for metric in ["training_load", "sprint_distance", "distance_per_min", "max_speed", "dec_count", "high_efforts"]:
        if metric in weekly.columns:
            weekly[f"{metric}_rolling4"] = weekly[metric].rolling(4, min_periods=2).mean()
            weekly[f"{metric}_change"] = weekly[metric].pct_change()

    contexts = []
    for idx, row in weekly.iterrows():
        contexts.append(classify_week_context(row, weekly.iloc[:idx]))
    weekly["periodizacios_tipus"] = contexts
    return weekly


def calculate_readiness_score(df: pd.DataFrame, selected_week: str, playstyle: str) -> Tuple[int, Dict[str, float], List[str]]:
    daily = build_microcycle_table(df, selected_week)
    weekly = build_weekly_fingerprints(df)
    current_week = weekly[weekly["week"] == selected_week]
    if current_week.empty:
        return 50, {}, ["Nincs elég adat a readiness számításhoz."]

    score = 75.0
    components = {}
    reasons = []

    row = current_week.iloc[0]

    # 1. Load trend komponens
    load_change = row.get("training_load_change", np.nan)
    if pd.notna(load_change):
        if load_change > 0.25:
            score -= 12
            reasons.append("A heti terhelés jelentősen nőtt az előző héthez képest.")
        elif load_change < -0.30:
            score -= 6
            reasons.append("A heti terhelés jelentősen visszaesett, ami alulterheltséget is jelezhet.")
        else:
            score += 4
        components["load_trend"] = max(0, min(100, 80 - abs(load_change) * 60))
    else:
        components["load_trend"] = 55

    # 2. Maximális sebességű inger
    match = daily[daily["session_type"] == "Meccs"] if not daily.empty else pd.DataFrame()
    train = daily[daily["session_type"] == "Edzés"] if not daily.empty else pd.DataFrame()
    speed_component = 60
    if not match.empty and not train.empty and "sprint_distance" in daily.columns:
        match_sprint = match["sprint_distance"].mean()
        max_train_sprint = train["sprint_distance"].max()
        if pd.notna(match_sprint) and match_sprint > 0 and pd.notna(max_train_sprint):
            ratio = max_train_sprint / match_sprint
            speed_component = min(100, ratio / 0.35 * 100)
            if ratio < 0.15:
                score -= 15
                reasons.append("A héten alig látszik maximális sebességű inger a meccsigényhez képest.")
            elif ratio < 0.30:
                score -= 7
                reasons.append("A maximális sebességű inger visszafogott volt.")
            else:
                score += 5
    components["speed_exposure"] = max(0, min(100, speed_component))

    # 3. Taper / MD-1 / MD-2 kontroll
    taper_component = 65
    if not daily.empty and "MD" in daily["md_label"].values:
        md1 = daily[daily["md_label"] == "MD-1"]
        md2 = daily[daily["md_label"] == "MD-2"]
        md34 = daily[daily["md_label"].isin(["MD-3", "MD-4"])]
        if not md1.empty and not md34.empty:
            peak_early = md34["load_index"].max()
            md1_load = md1["load_index"].sum()
            if pd.notna(peak_early) and peak_early > 0 and pd.notna(md1_load):
                ratio = md1_load / peak_early
                taper_component = max(0, min(100, (1.0 - ratio) * 130))
                if ratio > 0.65:
                    score -= 10
                    reasons.append("Az MD-1 terhelés magas lehetett a frissességhez képest.")
                elif ratio < 0.45:
                    score += 6
        if not md2.empty and not md34.empty:
            peak_early = md34["load_index"].max()
            md2_load = md2["load_index"].sum()
            if pd.notna(peak_early) and peak_early > 0 and pd.notna(md2_load) and md2_load / peak_early > 0.80:
                score -= 8
                reasons.append("Az MD-2 terhelés közel volt a hét fő terhelési napjához.")
    components["tapering"] = max(0, min(100, taper_component))

    # 4. Játékmodell illeszkedés
    playstyle_component = 70
    current = df[df["week"] == selected_week]
    cur_train = current[current["session_type"] == "Edzés"]
    cur_match = current[current["session_type"] == "Meccs"]
    if not cur_train.empty and not cur_match.empty:
        def ratio(metric):
            if metric not in current.columns:
                return np.nan
            m = cur_match[metric].mean()
            t = cur_train[metric].mean()
            return t / m if pd.notna(m) and m > 0 and pd.notna(t) else np.nan

        int_ratio = ratio("distance_per_min")
        spr_ratio = ratio("sprint_distance")
        eff_ratio = ratio("high_efforts")

        if playstyle == "Pressing":
            vals = [v for v in [int_ratio / 0.90 if pd.notna(int_ratio) else np.nan,
                                eff_ratio / 0.75 if pd.notna(eff_ratio) else np.nan] if pd.notna(v)]
        elif playstyle == "Transition":
            vals = [spr_ratio / 0.80] if pd.notna(spr_ratio) else []
        elif playstyle == "Possession":
            vals = [1.0 - max(0, (int_ratio - 1.15)) if pd.notna(int_ratio) else np.nan]
            vals = [v for v in vals if pd.notna(v)]
        elif playstyle == "Low Block":
            vals = [spr_ratio / 0.60] if pd.notna(spr_ratio) else []
        else:
            vals = [1.0]

        if vals:
            playstyle_component = max(0, min(100, np.mean([min(1.15, v) for v in vals]) * 85))
            if playstyle_component < 60:
                score -= 8
                reasons.append(f"A heti profil nem illeszkedik elég jól a(z) {playstyle} játékmodellhez.")
            elif playstyle_component > 80:
                score += 3

    components["playstyle_fit"] = max(0, min(100, playstyle_component))

    # 5. Meccskészültség végső skála
    score = int(max(0, min(100, round(score))))
    if not reasons:
        reasons.append("A readiness fő komponensei stabilnak tűnnek.")
    return score, components, reasons


def build_pattern_insights(df: pd.DataFrame, selected_week: str) -> List[Insight]:
    insights: List[Insight] = []
    weekly = build_weekly_fingerprints(df)
    if weekly.empty or selected_week not in weekly["week"].values:
        return insights

    idx = weekly.index[weekly["week"] == selected_week][0]
    hist = weekly.iloc[max(0, idx-5):idx+1].copy()

    if len(hist) < 3:
        insights.append(Insight(
            "Kevés történeti adat", "INFORMÁCIÓ",
            "A multi-week mintázatokhoz legalább 3 hét adat szükséges.",
            "Kevesebb adat mellett a rendszer inkább óvatos heti következtetéseket tud adni.",
            "Érdemes több heti adatot egy fájlban feltölteni vagy használni a performance memória funkciót.",
            scope="Memory",
        ))
        return insights

    # Sprint trend
    if "sprint_distance" in hist.columns:
        label = trend_label(hist["sprint_distance"].tolist())
        if label == "csökkenő":
            insights.append(Insight(
                "Többhetes sprintcsökkenés", "FIGYELMEZTETÉS",
                "Az elmúlt hetek sprintprofilja csökkenő mintázatot mutat.",
                "A tartós maximális sebességű inger hiány hosszabb távon ronthatja a maximális sebesség és intenzitás fenntartását.",
                "Érdemes célzottan megvizsgálni, hogy tudatos terheléscsökkentésről vagy nem kívánt intenzitáshiányról van-e szó.",
                scope="Memory",
            ))

    # Load trend
    if "training_load" in hist.columns:
        label = trend_label(hist["training_load"].tolist())
        if label == "emelkedő":
            insights.append(Insight(
                "Többhetes terhelésnövekedés", "FIGYELMEZTETÉS",
                "Az elmúlt hetek terhelési profilja emelkedő mintázatot mutat.",
                "A folyamatos load növekedés hasznos lehet építő periódusban, de frissességi problémát is okozhat.",
                "Érdemes ellenőrizni, hogy a növekvő terhelés illeszkedik-e a periodizációs célhoz.",
                scope="Memory",
            ))

    # Max speed trend
    if "max_speed" in hist.columns:
        label = trend_label(hist["max_speed"].tolist(), tolerance=0.02)
        if label == "csökkenő":
            insights.append(Insight(
                "Max sebesség csökkenő trend", "INFORMÁCIÓ",
                "Az elmúlt hetek maximális sebességprofilja enyhén csökkenő mintázatot mutat.",
                "Ez lehet edzéstartalom-függő, de jelezhet neuromuszkuláris frissességi problémát is.",
                "Érdemes ellenőrizni, hogy volt-e megfelelő maximális sebességű inger a hetekben.",
                scope="Memory",
            ))

    return insights


def memory_file_exists() -> bool:
    return MEMORY_FILE.exists()


def load_memory_df() -> pd.DataFrame:
    if not MEMORY_FILE.exists():
        return pd.DataFrame()
    try:
        return pd.read_csv(MEMORY_FILE, parse_dates=["start_time"])
    except Exception:
        return pd.DataFrame()


def save_to_memory(df: pd.DataFrame) -> Tuple[bool, str]:
    try:
        existing = load_memory_df()
        combined = pd.concat([existing, df], ignore_index=True) if not existing.empty else df.copy()
        # Duplikáció minimalizálása.
        key_cols = [c for c in ["player_name", "start_time", "session_type", "session_name", "total_distance", "training_load"] if c in combined.columns]
        if key_cols:
            combined = combined.drop_duplicates(subset=key_cols, keep="last")
        combined.to_csv(MEMORY_FILE, index=False, encoding="utf-8-sig")
        return True, f"Memory mentve: {len(combined)} sor."
    except Exception as exc:
        return False, f"Memory mentés sikertelen: {exc}"


def merge_with_memory(current_df: pd.DataFrame, use_memory: bool) -> pd.DataFrame:
    if not use_memory:
        return current_df
    mem = load_memory_df()
    if mem.empty:
        return current_df
    combined = pd.concat([mem, current_df], ignore_index=True)
    key_cols = [c for c in ["player_name", "start_time", "session_type", "session_name", "total_distance", "training_load"] if c in combined.columns]
    if key_cols:
        combined = combined.drop_duplicates(subset=key_cols, keep="last")
    # Dátummezők visszaállítása.
    if "start_time" in combined.columns:
        combined["start_time"] = pd.to_datetime(combined["start_time"], errors="coerce")
        combined["session_date"] = combined["start_time"].dt.date
        combined["week"] = make_iso_week_series(combined["start_time"])
    return combined


def build_adaptive_recommendations(
    insights: List[Insight],
    readiness_score: int,
    periodization_type: str,
    pattern_insights: List[Insight],
    playstyle: str
) -> List[Dict[str, str]]:
    recs = []

    if readiness_score < 55:
        recs.append({
            "Cím": "Meccskészültség elsődleges figyelem",
            "Súlyosság": "KRITIKUS",
            "Teendő": "A következő edzésen érdemes óvatosan bánni a terheléssel, és külön figyelni, mennyire frissek a játékosok.",
            "Miért": "A meccskészültségi pontszám alacsony, vagyis több jel is arra utal, hogy a csapat nem teljesen friss.",
        })
    elif readiness_score < 70:
        recs.append({
            "Cím": "Meccskészültség kontroll",
            "Súlyosság": "FIGYELMEZTETÉS",
            "Teendő": "A következő 1–2 edzésen ne a plusz terhelés legyen a fő cél, hanem a frissesség stabil megtartása.",
            "Miért": "A csapat nincs rossz állapotban, de nem is tűnik teljesen frissnek. Ilyenkor a túl nagy változtatás visszaüthet.",
        })

    if "overload" in periodization_type.lower() or "terhelésépítő" in periodization_type.lower():
        recs.append({
            "Cím": "Overload hét kontrollja",
            "Súlyosság": "FIGYELMEZTETÉS",
            "Teendő": "Ellenőrizd, hogy a magasabb terhelés tervezett-e, és legyen meg a következő frissítő blokk.",
            "Miért": "Terhelésépítő héten a következő napok/frissítés minősége kulcsfontosságú.",
        })
    elif "recovery" in periodization_type.lower() or "alulterhelt" in periodization_type.lower():
        recs.append({
            "Cím": "Könnyebb hét értelmezése",
            "Súlyosság": "INFORMÁCIÓ",
            "Teendő": "Tisztázd, hogy ez szándékosan könnyebb hét volt-e. Ha igen, rendben lehet; ha nem, akkor hiányozhatott a megfelelő edzésinger.",
            "Miért": "Az alacsonyabb terhelés rövid távon frissíthet, de ha több héten át így marad, a csapat intenzitása visszaeshet.",
        })

    for ins in pattern_insights:
        if ins.severity in ["KRITIKUS", "FIGYELMEZTETÉS"]:
            recs.append({
                "Cím": ins.title,
                "Súlyosság": ins.severity,
                "Teendő": ins.recommendation,
                "Miért": ins.impact,
            })

    # Alap insightokból pótoljuk 3 elemre.
    for item in top_coaching_priorities(insights, limit=5):
        if len(recs) >= 3:
            break
        if item["Cím"] not in [r["Cím"] for r in recs]:
            recs.append(item)

    return recs[:3]


def render_score_card(title: str, score: int, subtitle: str, reasons: List[str]) -> None:
    color = score_to_color(score)
    reason_html = "<br>".join([f"• {html.escape(str(r))}" for r in reasons[:4]])
    st.markdown(
        f"""
        <div class="score-card">
            <div class="score-number" style="color:{color};">{score}/100</div>
            <div class="score-label">{html.escape(title)}</div>
            <div class="mini-muted">{html.escape(subtitle)}</div>
            <div class="mini-muted" style="margin-top:10px;">{reason_html}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def severity_icon(sev: str) -> str:
    if sev == "KRITIKUS":
        return "🔴"
    if sev == "FIGYELMEZTETÉS":
        return "🟠"
    return "🔵"


def severity_class(sev: str) -> str:
    if sev == "KRITIKUS":
        return "pill-critical"
    if sev == "FIGYELMEZTETÉS":
        return "pill-warning"
    return "pill-info"


def metric_card(label: str, value: object, help_text: str = ""):
    st.metric(label, value if value is not None else "—", help=help_text)


def _safe_filename_week(week: str) -> str:
    return re.sub(r"[^0-9A-Za-z_-]+", "_", str(week)).strip("_") or "het"


def build_insight_export_df(insights: List[Insight]) -> pd.DataFrame:
    df = pd.DataFrame([i.as_dict() for i in insights])
    if df.empty:
        return df
    order = {"KRITIKUS": 1, "FIGYELMEZTETÉS": 2, "INFORMÁCIÓ": 3}
    df.insert(0, "Prioritás", df["Súlyosság"].map(order).fillna(9).astype(int))
    df["Súlyosság"] = df["Súlyosság"].map({
        "KRITIKUS": "🔴 KRITIKUS",
        "FIGYELMEZTETÉS": "🟠 FIGYELMEZTETÉS",
        "INFORMÁCIÓ": "🔵 INFORMÁCIÓ",
    }).fillna(df["Súlyosság"])
    return df


def render_insight_cards(insights: List[Insight]) -> None:
    for ins in insights:
        pill = severity_class(ins.severity)
        st.markdown(
            f"""
            <div class="insight-card">
                <div>
                    <span class="pill {pill}">{severity_icon(ins.severity)} {html.escape(ins.severity)}</span>
                    <span class="pill pill-info">{html.escape(ins.scope)}</span>
                </div>
                <div class="insight-title">{html.escape(ins.title)}</div>
                <div class="insight-label">Mit látunk?</div>
                <div class="insight-text">{html.escape(ins.observation)}</div>
                <div class="insight-label">Miért fontos?</div>
                <div class="insight-text">{html.escape(ins.impact)}</div>
                <div class="insight-label">Javaslat</div>
                <div class="insight-text">{html.escape(ins.recommendation)}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def render_wrapped_table(df: pd.DataFrame) -> None:
    if df.empty:
        st.info("Nincs megjeleníthető adat.")
        return
    table_html = ['<div class="wrap-table"><table>']
    table_html.append("<thead><tr>" + "".join(f"<th>{html.escape(str(c))}</th>" for c in df.columns) + "</tr></thead>")
    table_html.append("<tbody>")
    for _, row in df.iterrows():
        table_html.append("<tr>" + "".join(f"<td>{html.escape(str(row.get(c, '')))}</td>" for c in df.columns) + "</tr>")
    table_html.append("</tbody></table></div>")
    st.markdown("".join(table_html), unsafe_allow_html=True)


def insights_to_excel_bytes(insights_df: pd.DataFrame, selected_week: str) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        insights_df.to_excel(writer, index=False, sheet_name="Megállapítások", startrow=3)
        wb = writer.book
        ws = writer.sheets["Megállapítások"]

        from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        max_col = max(1, len(insights_df.columns))
        ws["A1"] = "Performance megállapítások és javaslatok"
        ws["A2"] = f"Hét: {selected_week}"
        ws["A3"] = f"Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max_col)
        ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=max_col)
        ws["A1"].font = Font(bold=True, size=16)
        ws["A2"].font = Font(bold=True, size=11)
        ws["A3"].font = Font(italic=True, size=10)

        header_row = 4
        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(color="FFFFFF", bold=True)
        thin = Side(style="thin", color="D9E2F3")

        for cell in ws[header_row]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

        widths = {
            "Prioritás": 10,
            "Súlyosság": 18,
            "Terület": 14,
            "Megállapítás": 28,
            "Mit látunk?": 58,
            "Miért fontos?": 58,
            "Javaslat": 68,
        }
        for idx, col_name in enumerate(insights_df.columns, start=1):
            ws.column_dimensions[get_column_letter(idx)].width = widths.get(str(col_name), 22)

        for row in ws.iter_rows(min_row=5, max_row=ws.max_row, max_col=ws.max_column):
            max_lines = 1
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
                text = str(cell.value or "")
                max_lines = max(max_lines, min(8, (len(text) // 50) + 1))
            ws.row_dimensions[row[0].row].height = max(26, max_lines * 17)

        ws.freeze_panes = "A5"
        ws.auto_filter.ref = ws.dimensions
    return output.getvalue()


def insights_to_word_bytes(insights_df: pd.DataFrame, selected_week: str) -> Optional[bytes]:
    if Document is None:
        return None
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    doc.add_heading("Performance megállapítások és javaslatok", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Hét: {selected_week}\n").bold = True
    p.add_run(f"Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for _, row in insights_df.iterrows():
        doc.add_heading(f"{row.get('Súlyosság', '')} · {row.get('Megállapítás', '')}", level=2)
        meta = doc.add_paragraph()
        meta.add_run("Terület: ").bold = True
        meta.add_run(str(row.get("Terület", "")))
        for label in ["Mit látunk?", "Miért fontos?", "Javaslat"]:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(row.get(label, "")))

    doc.add_page_break()
    doc.add_heading("Összesítő tábla", level=2)
    table = doc.add_table(rows=1, cols=len(insights_df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(insights_df.columns):
        hdr_cells[i].text = str(col)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for _, row in insights_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(insights_df.columns):
            cells[i].text = str(row.get(col, ""))
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _register_pdf_font() -> Tuple[str, str]:
    """ReportLab alapfontjai nem kezelik jól a magyar ékezeteket.
    Ezért megpróbálunk DejaVuSans betűtípust regisztrálni, ami Streamlit Cloudon
    és Linuxon általában elérhető. Ha nem található, visszaesünk Helvetica-ra.
    """
    if pdfmetrics is None or TTFont is None:
        return "Helvetica", "Helvetica-Bold"

    font_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        "/usr/local/share/fonts/DejaVuSans.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    bold_candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed-Bold.ttf",
        "/usr/local/share/fonts/DejaVuSans-Bold.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
    ]

    normal_path = next((fp for fp in font_candidates if Path(fp).exists()), None)
    bold_path = next((fp for fp in bold_candidates if Path(fp).exists()), None)

    if not normal_path:
        return "Helvetica", "Helvetica-Bold"

    try:
        pdfmetrics.registerFont(TTFont("DejaVuSans", normal_path))
        if bold_path:
            pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", bold_path))
            return "DejaVuSans", "DejaVuSans-Bold"
        return "DejaVuSans", "DejaVuSans"
    except Exception:
        return "Helvetica", "Helvetica-Bold"


def insights_to_pdf_bytes(insights_df: pd.DataFrame, selected_week: str) -> Optional[bytes]:
    if SimpleDocTemplate is None:
        return None

    font_name, font_bold = _register_pdf_font()

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=0.8 * cm,
        leftMargin=0.8 * cm,
        topMargin=0.8 * cm,
        bottomMargin=0.8 * cm,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "MagyarCim",
        parent=styles["Title"],
        fontName=font_bold,
        fontSize=17,
        leading=21,
        spaceAfter=10,
    )
    meta_style = ParagraphStyle(
        "MagyarMeta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        leading=11,
    )
    normal = ParagraphStyle(
        "MagyarNormalWrapped",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=7,
        leading=9,
        wordWrap="CJK",
    )
    header = ParagraphStyle(
        "MagyarHeaderWrapped",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=7,
        leading=9,
        textColor=colors.white,
        wordWrap="CJK",
    )

    story = [
        Paragraph(pdf_safe_text("Performance megállapítások és javaslatok"), title_style),
        Paragraph(pdf_safe_text(f"Hét: {selected_week} · Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}"), meta_style),
        Spacer(1, 0.25 * cm),
    ]

    cols = ["Súlyosság", "Terület", "Megállapítás", "Mit látunk?", "Miért fontos?", "Javaslat"]
    cols = [c for c in cols if c in insights_df.columns]

    def safe_paragraph(value, style):
        text = "" if pd.isna(value) else str(value)
        text = pdf_safe_text(text)
        return Paragraph(html.escape(text).replace("\n", "<br/>",), style)

    table_data = [[safe_paragraph(c, header) for c in cols]]
    for _, row in insights_df.iterrows():
        table_data.append([safe_paragraph(row.get(c, ""), normal) for c in cols])

    col_widths_map = {
        "Súlyosság": 2.6 * cm,
        "Terület": 2.0 * cm,
        "Megállapítás": 3.8 * cm,
        "Mit látunk?": 6.0 * cm,
        "Miért fontos?": 6.0 * cm,
        "Javaslat": 7.2 * cm,
    }
    table = Table(table_data, colWidths=[col_widths_map.get(c, 3 * cm) for c in cols], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), font_bold),
        ("FONTNAME", (0, 1), (-1, -1), font_name),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BFBFBF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FC")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)

    # V5 PDF pages - past/current/next split planning
    try:
        _past_review = globals().get("past_week_review_df", pd.DataFrame())
        _past_text = globals().get("past_week_review_text", "")
        _current_plan = globals().get("current_remaining_plan_df", pd.DataFrame())
        _current_text = globals().get("current_remaining_text", "")
        _next_week_plan = globals().get("next_week_plan_df", pd.DataFrame())
        _next_week_text = globals().get("next_week_plan_text", "")
        _player_actions = globals().get("player_next_actions_df", pd.DataFrame())

        try:
            from reportlab.platypus import PageBreak
        except Exception:
            PageBreak = None

        if PageBreak:
            story.append(PageBreak())
        story.append(P("Múlt hét / aktuális hét / jövő hét", title))
        story.append(P("Teljes múlt heti értékelés, aktuális hét hátralévő napjai, és jövő heti mikrociklus terv.", subtitle))

        story.append(section_bar("1) Múlt hét -> javaslat erre a hétre"))
        story.append(data_table([[P("Összefoglaló", header)], [P(_past_text, body)]], [27.7 * cm], header_bg="#1E3A8A", row_bgs=[colors.HexColor("#EFF6FF")]))
        if _past_review is not None and not _past_review.empty:
            cols = [c for c in ["Prioritás", "Múlt heti megállapítás", "Súlyosság", "Javaslat erre a hétre"] if c in _past_review.columns]
            pdata = [[P(c, header) for c in cols]]
            for _, rr in _past_review.head(6).iterrows():
                pdata.append([P(rr.get(c, ""), tiny) for c in cols])
            story.append(data_table(pdata, [1.7 * cm, 7.0 * cm, 3.0 * cm, 16.0 * cm][:len(cols)], header_bg="#0F172A", row_bgs=[colors.white, colors.HexColor("#F8FAFC")]))

        story += [Spacer(1, 0.25 * cm)]
        story.append(section_bar("2) Aktuális hét -> hátralévő napok"))
        story.append(data_table([[P("Összefoglaló", header)], [P(_current_text, body)]], [27.7 * cm], header_bg="#166534", row_bgs=[colors.HexColor("#ECFDF5")]))
        if _current_plan is not None and not _current_plan.empty:
            cols = [c for c in ["Hátralévő pont", "Fókusz", "Ajánlott terhelés", "Javaslat"] if c in _current_plan.columns]
            cdata = [[P(c, header) for c in cols]]
            for _, rr in _current_plan.head(6).iterrows():
                cdata.append([P(rr.get(c, ""), tiny) for c in cols])
            story.append(data_table(cdata, [3.2 * cm, 4.8 * cm, 3.5 * cm, 16.2 * cm][:len(cols)], header_bg="#166534", row_bgs=[colors.white, colors.HexColor("#ECFDF5")]))

        if PageBreak:
            story.append(PageBreak())
        story.append(P("Jövő heti mikrociklus terv", title))
        story.append(P(_next_week_text, subtitle))
        if _next_week_plan is not None and not _next_week_plan.empty:
            cols = [c for c in ["Nap", "Szerep", "Fő cél", "Ajánlott terhelés", "Javaslat", "Tervezési alap"] if c in _next_week_plan.columns]
            ndata = [[P(c, header) for c in cols]]
            for _, rr in _next_week_plan.head(8).iterrows():
                ndata.append([P(rr.get(c, ""), tiny) for c in cols])
            story.append(data_table(ndata, [1.7 * cm, 3.2 * cm, 4.0 * cm, 3.0 * cm, 8.8 * cm, 7.0 * cm][:len(cols)], header_bg="#312E81", row_bgs=[colors.white, colors.HexColor("#F5F3FF")]))

        if _player_actions is not None and not _player_actions.empty:
            story += [Spacer(1, 0.25 * cm)]
            story.append(section_bar("Játékosszintű következő teendők"))
            cols = [c for c in ["Játékos", "Prioritás", "Holnap / következő edzés", "Következő hét", "Indok"] if c in _player_actions.columns]
            adata = [[P(c, header) for c in cols]]
            for _, rr in _player_actions.head(8).iterrows():
                adata.append([P(rr.get(c, ""), tiny) for c in cols])
            story.append(data_table(adata, [3.3 * cm, 2.5 * cm, 7.0 * cm, 7.0 * cm, 7.9 * cm][:len(cols)], header_bg="#7F1D1D", row_bgs=[colors.white, colors.HexColor("#FFF7ED")]))
    except Exception:
        pass


    doc.build(story)
    return output.getvalue()



# -----------------------------------------------------------------------------
# V3 - Premium visualization + Player Risk + Positional layer + improved PDF
# -----------------------------------------------------------------------------
POSITION_GROUPS = {
    "Kapus": ["gk", "kapus", "goalkeeper"],
    "Belső védő": ["cb", "belső védő", "center back", "centre back"],
    "Szélső védő": ["fb", "wb", "szélső védő", "fullback", "wingback"],
    "Középpályás": ["cm", "dm", "am", "középpályás", "midfielder"],
    "Szélső": ["winger", "szélső", "wide"],
    "Csatár": ["st", "cf", "csatár", "striker", "forward"],
}

def infer_position_group(value: object) -> str:
    if value is None or pd.isna(value): return "Ismeretlen"
    text = str(value).lower().strip()
    for group, keys in POSITION_GROUPS.items():
        if any(k in text for k in keys): return group
    return str(value).strip() if str(value).strip() else "Ismeretlen"

def add_position_group(df: pd.DataFrame) -> pd.DataFrame:
    out=df.copy(); out["position_group"] = out["position"].apply(infer_position_group) if "position" in out.columns else "Ismeretlen"; return out

def calculate_player_risk(df: pd.DataFrame, selected_week: str) -> pd.DataFrame:
    pw=player_weekly(df)
    if pw.empty or selected_week not in pw["week"].values: return pd.DataFrame()
    weeks=sorted(pw["week"].dropna().unique().tolist()); idx=weeks.index(selected_week)
    cur=pw[pw["week"]==selected_week].copy(); hist=pw[pw["week"].isin(weeks[max(0,idx-4):idx])].copy(); rows=[]
    for _, row in cur.iterrows():
        player=row["player_name"]; hp=hist[hist["player_name"]==player]; score=20; reasons=[]
        is_gk = bool(row.get("is_goalkeeper", False))
        # Kapusnál nem büntetjük ugyanúgy az alacsony sprint/HSR profilt.
        # A risk fő fókusza nála a teljes load, high efforts, lassítás/gyorsulás és játékperc.
        metric_pack = [
            ("training_load", .35 if is_gk else .30, -0.40 if is_gk else -0.35, "Terhelési pont"),
            ("high_efforts", .45 if is_gk else .35, -0.45 if is_gk else -0.40, "High Efforts"),
            ("dec_count", .45 if is_gk else .35, -0.45 if is_gk else -0.40, "Lassítás"),
        ]
        if not is_gk:
            metric_pack.append(("sprint_distance", .45, -0.45, "Sprinttáv"))
        for metric, hi, lo, lab in metric_pack:
            if metric in row.index and metric in hp.columns:
                v=row.get(metric,np.nan); base=hp[metric].mean()
                if pd.notna(v) and pd.notna(base) and base!=0:
                    d=(v-base)/base
                    if d>hi: score+=18; reasons.append(f"{lab}: +{d:.0%} a saját átlaghoz képest")
                    elif d<lo: score+=8; reasons.append(f"{lab}: {d:.0%} a saját átlaghoz képest")
        if "player_minutes" in row.index and "player_minutes" in hp.columns:
            v=row.get("player_minutes",np.nan); base=hp["player_minutes"].mean()
            if pd.notna(v) and pd.notna(base) and base>0 and abs((v-base)/base)>.35:
                d=(v-base)/base; score+=8; reasons.append(f"Játékperc/exposure: {d:+.0%} a saját átlaghoz képest")
        if (not is_gk) and "max_speed" in row.index and "max_speed" in hp.columns:
            v=row.get("max_speed",np.nan); base=hp["max_speed"].max()
            if pd.notna(v) and pd.notna(base) and base>0 and (v-base)/base < -.06:
                score+=14; reasons.append(f"Max sebesség: {(v-base)/base:.0%} a saját csúcshoz képest")
        score=int(max(0,min(100,score))); level="Magas" if score>=70 else ("Közepes" if score>=45 else "Alacsony")
        role = "Kapus" if is_gk else "Mezőny"
        rows.append({"Játékos":player,"Szerep":role,"Típus":row.get("session_type",""),"Játékperc":round(row.get("player_minutes",0) or 0,1),"Kockázati pontszám":score,"Kockázati szint":level,"Fő okok":"; ".join(reasons[:3]) if reasons else "Nincs jelentős eltérés a saját előzményhez képest."})
    res=pd.DataFrame(rows); return res.sort_values("Kockázati pontszám",ascending=False) if not res.empty else res

def render_premium_kpi(label: str, value: str, note: str="", color: str="#22c55e") -> None:
    st.markdown(f"""<div class='premium-kpi'><div class='premium-kpi-label'>{html.escape(label)}</div><div class='premium-kpi-value' style='color:{color};'>{html.escape(str(value))}</div><div class='premium-kpi-note'>{html.escape(note)}</div></div>""", unsafe_allow_html=True)

def render_hero(selected_week: str, selected_playstyle: str, readiness_score: int, periodization_type: str) -> None:
    color=score_to_color(readiness_score) if "score_to_color" in globals() else "#22c55e"; label=score_to_label(readiness_score) if "score_to_label" in globals() else ""
    st.markdown(f"""<div class='hero-box'><div class='hero-title'>Football Performance Intelligence</div><div class='hero-sub'><span class='section-chip'>Hét: {html.escape(str(selected_week))}</span><span class='section-chip'>Játékmodell: {html.escape(str(selected_playstyle))}</span><span class='section-chip'>Meccskészültség: <b style='color:{color};'>{readiness_score}/100</b> - {html.escape(label)}</span><span class='section-chip'>Periodizáció: {html.escape(str(periodization_type))}</span></div></div>""", unsafe_allow_html=True)

def render_risk_cards(risk_df: pd.DataFrame, limit: int=5) -> None:
    if risk_df.empty: st.info("Nincs elég adat játékosszintű risk engine számításhoz."); return
    for _, row in risk_df.head(limit).iterrows():
        level=row.get("Kockázati szint","Alacsony"); css="risk-high" if level=="Magas" else ("risk-medium" if level=="Közepes" else "risk-low")
        st.markdown(f"""<div class='insight-card {css}'><div class='insight-title'>{html.escape(str(row.get('Játékos','')))} · {html.escape(str(level))} kockázat · {row.get('Kockázati pontszám',0)}/100</div><div class='insight-label'>Fő okok</div><div class='insight-text'>{html.escape(str(row.get('Fő okok','')))}</div></div>""", unsafe_allow_html=True)





# -----------------------------------------------------------------------------
# V5 - Past / current / next microcycle planning
# -----------------------------------------------------------------------------
def week_completeness_summary(df: pd.DataFrame, selected_week: str) -> Dict[str, object]:
    """Részleges hét értelmezése.
    Nem kell teljes hét: 1-2-3 edzés alapján is adunk folyamat közbeni visszajelzést.
    """
    if df is None or df.empty or "week" not in df.columns:
        return {
            "status": "Nincs adat",
            "sessions": 0,
            "train_days": 0,
            "match_days": 0,
            "days": 0,
            "last_day": None,
            "message": "Nincs értelmezhető heti adat.",
        }

    week_df = df[df["week"] == selected_week].copy()
    if week_df.empty:
        return {
            "status": "Nincs adat",
            "sessions": 0,
            "train_days": 0,
            "match_days": 0,
            "days": 0,
            "last_day": None,
            "message": "Az aktuális hétre nincs adat.",
        }

    week_df["session_date_dt"] = pd.to_datetime(week_df["session_date"], errors="coerce")
    sessions = len(week_df[["session_date_dt", "session_type"]].drop_duplicates()) if "session_type" in week_df.columns else week_df["session_date_dt"].nunique()
    train_days = week_df.loc[week_df["session_type"] == "Edzés", "session_date_dt"].dt.date.nunique() if "session_type" in week_df.columns else 0
    match_days = week_df.loc[week_df["session_type"] == "Meccs", "session_date_dt"].dt.date.nunique() if "session_type" in week_df.columns else 0
    days = week_df["session_date_dt"].dt.date.nunique()
    last_day = week_df["session_date_dt"].max()

    if match_days > 0:
        status = "Teljes / meccsel együtt értelmezhető hét"
    elif train_days <= 2:
        status = "Aktuális / folyamatban lévő hét - korai jelzés"
    elif train_days <= 4:
        status = "Aktuális / folyamatban lévő hét - tervezési pont"
    else:
        status = "Majdnem teljes edzéshet"

    message = (
        f"{status}. Eddig {train_days} edzésnap és {match_days} meccsnap látható. "
        "A javaslatokat a rendelkezésre álló adatokhoz igazítjuk, nem feltételezzük, hogy a hét teljes."
    )

    return {
        "status": status,
        "sessions": sessions,
        "train_days": train_days,
        "match_days": match_days,
        "days": days,
        "last_day": last_day,
        "message": message,
    }


def get_surrounding_week_context(df: pd.DataFrame, selected_week: str) -> Dict[str, Optional[str]]:
    weeks_sorted = sorted(df["week"].dropna().unique().tolist()) if df is not None and not df.empty and "week" in df.columns else []
    if selected_week not in weeks_sorted:
        return {"previous_week": None, "current_week": selected_week, "next_data_week": None}
    i = weeks_sorted.index(selected_week)
    return {
        "previous_week": weeks_sorted[i - 1] if i > 0 else None,
        "current_week": selected_week,
        "next_data_week": weeks_sorted[i + 1] if i < len(weeks_sorted) - 1 else None,
    }


def build_past_week_review(
    df: pd.DataFrame,
    selected_week: str,
    playstyle: str,
) -> Tuple[pd.DataFrame, str]:
    """Múlt hét teljes elemzése + javaslat az aktuális hétre."""
    ctx = get_surrounding_week_context(df, selected_week)
    prev_week = ctx.get("previous_week")
    if not prev_week:
        return pd.DataFrame(), "Nincs előző hét adat. Az aktuális javaslat csak a kiválasztott hét alapján készül."

    prev_insights = (
        team_insights(df, prev_week)
        + microcycle_insights(df, prev_week)
        + playstyle_insights(df, prev_week, playstyle)
        + build_pattern_insights(df, prev_week)
    )
    prev_insights = sorted(prev_insights, key=lambda x: SEVERITY_RANK.get(x.severity, 9))[:10]
    prev_readiness, _, _ = calculate_readiness_score(df, prev_week, playstyle)

    prev_fp = build_weekly_fingerprints(df)
    prev_period = "Nincs elég adat"
    if prev_fp is not None and not prev_fp.empty and "week" in prev_fp.columns:
        row = prev_fp[prev_fp["week"] == prev_week]
        if not row.empty and "periodizacios_tipus" in row.columns:
            prev_period = row["periodizacios_tipus"].iloc[0]

    rows = []
    for i, ins in enumerate(prev_insights, 1):
        rows.append({
            "Prioritás": i,
            "Múlt heti megállapítás": ins.title,
            "Súlyosság": ins.severity,
            "Mit láttunk?": ins.observation,
            "Javaslat erre a hétre": ins.recommendation,
        })

    if not rows:
        rows.append({
            "Prioritás": 1,
            "Múlt heti megállapítás": "Stabil hét",
            "Súlyosság": "INFORMÁCIÓ",
            "Mit láttunk?": "Nem látható kiemelt negatív eltérés.",
            "Javaslat erre a hétre": "A struktúra megtartható, de sprint/load/risk monitoring maradjon aktív.",
        })

    txt = (
        f"Múlt hét: {format_week_label(prev_week)}\n"
        f"Readiness: {prev_readiness}/100 ({score_to_label(prev_readiness)})\n"
        f"Periodizáció: {prev_period}\n\n"
        "Ebből erre a hétre a fő cél: a kritikus/figyelmeztető pontok kezelése, "
        "miközben a frissességet és az egyéni risket kontroll alatt tartjuk."
    )
    return pd.DataFrame(rows), txt


def build_current_remaining_days_plan(
    df: pd.DataFrame,
    selected_week: str,
    playstyle: str,
    readiness_score: int,
    periodization_type: str,
    player_risk_df: pd.DataFrame,
) -> Tuple[pd.DataFrame, str]:
    """Aktuális hét: eddigi feltöltött napok alapján javaslat a hátralévő napokra."""
    week_df = df[df["week"] == selected_week].copy() if df is not None and "week" in df.columns else pd.DataFrame()
    if week_df.empty:
        return pd.DataFrame(), "Nincs aktuális heti adat."

    week_df["session_date_dt"] = pd.to_datetime(week_df["session_date"], errors="coerce")
    last_day = week_df["session_date_dt"].max()
    match_day = detect_match_day(week_df)
    has_match = match_day is not None
    ws = week_completeness_summary(df, selected_week)

    if has_match and pd.notna(last_day) and match_day is not None and last_day.date() >= match_day.date():
        md_slots = ["MD+1", "Következő edzés", "Jövő hét előkészítés"]
    elif has_match:
        md_slots = ["Hátralévő edzés", "MD-2", "MD-1", "MD"]
    else:
        if ws.get("train_days", 0) <= 1:
            md_slots = ["Holnap", "Következő 2. edzés", "Következő 3. edzés", "Hétvégi meccs / referencia"]
        elif ws.get("train_days", 0) <= 3:
            md_slots = ["Holnap", "Hátralévő fő edzés", "Meccs előtti aktiváció", "Hétvégi meccs / referencia"]
        else:
            md_slots = ["Következő edzés", "Aktiváció", "Meccs / referencia"]

    high_risk = []
    if player_risk_df is not None and not player_risk_df.empty and "Kockázati szint" in player_risk_df.columns:
        high_risk = player_risk_df.loc[player_risk_df["Kockázati szint"] == "Magas", "Játékos"].head(5).tolist() if "Játékos" in player_risk_df.columns else []

    low_sprint = False
    intensity_gap = False
    train = week_df[week_df["session_type"] == "Edzés"] if "session_type" in week_df.columns else pd.DataFrame()
    match = week_df[week_df["session_type"] == "Meccs"] if "session_type" in week_df.columns else pd.DataFrame()
    if not train.empty:
        if "sprint_distance" in train.columns and train["sprint_distance"].mean() < 250:
            low_sprint = True
        if not match.empty and "distance_per_min" in week_df.columns and match["distance_per_min"].mean() > 0:
            intensity_gap = (train["distance_per_min"].mean() / match["distance_per_min"].mean()) < 0.88

    rows = []
    for slot in md_slots:
        if "MD+1" in slot:
            focus = "Regeneráció / pótló terhelés"
            rec = "Sokat játszóknak regeneráció, keveset játszóknak kontrollált kiegészítő blokk."
            load = "egyéni"
        elif "MD-1" in slot or "aktiváció" in slot.lower():
            focus = "Frissesség"
            rec = "Rövid aktiváció, alacsony volumen. Ne legyen új fárasztó inger."
            load = "alacsony"
        elif "MD-2" in slot:
            focus = "Taktikai kontroll"
            rec = "Taktikai fókusz, kevés lassítás/excentrikus terhelés. Magas risk játékosoknál limit."
            load = "alacsony-közepes"
        elif low_sprint or intensity_gap:
            focus = "Hiányzó intenzitás pótlása"
            rec = "Rövid, kontrollált sprint/max sebesség vagy magas tempójú játékblokk, kis volumennel."
            load = "közepes"
        else:
            focus = "Struktúra megtartása"
            rec = "A heti terv tartható, de egyéni risk és readiness kontroll javasolt."
            load = "közepes"

        if high_risk and load != "alacsony":
            rec += " Magas risk játékosoknál egyéni csökkentés: " + ", ".join(high_risk[:3]) + "."

        rows.append({
            "Hátralévő pont": slot,
            "Fókusz": focus,
            "Ajánlott terhelés": load,
            "Javaslat": rec,
            "Miért?": f"Eddig feltöltve: {ws.get('train_days', 0)} edzésnap, {ws.get('match_days', 0)} meccsnap. Readiness: {readiness_score}/100.",
        })

    txt = (
        f"Aktuális hét állapota: {ws.get('status')}\n"
        f"{ws.get('message')}\n\n"
        "A hátralévő napokra a javaslat nem teljes heti minősítés, hanem folyamat közbeni döntéstámogatás."
    )
    return pd.DataFrame(rows), txt


def build_next_microcycle_plan(
    df: pd.DataFrame,
    selected_week: str,
    playstyle: str,
    readiness_score: int,
    periodization_type: str,
    player_risk_df: pd.DataFrame,
) -> pd.DataFrame:
    """Jövő hét / következő mikrociklus javaslat MD-bontásban az eddigiek alapján."""
    week_df = df[df["week"] == selected_week].copy() if df is not None and "week" in df.columns else pd.DataFrame()

    low_sprint = False
    intensity_gap = False
    high_risk_players = []
    if player_risk_df is not None and not player_risk_df.empty:
        high_risk_players = player_risk_df.loc[player_risk_df["Kockázati szint"] == "Magas", "Játékos"].head(5).tolist() if "Kockázati szint" in player_risk_df.columns and "Játékos" in player_risk_df.columns else []

    if not week_df.empty:
        train = week_df[week_df["session_type"] == "Edzés"] if "session_type" in week_df.columns else pd.DataFrame()
        match = week_df[week_df["session_type"] == "Meccs"] if "session_type" in week_df.columns else pd.DataFrame()
        if not train.empty and not match.empty:
            if "sprint_distance" in week_df.columns and match["sprint_distance"].mean() > 0:
                low_sprint = (train["sprint_distance"].mean() / match["sprint_distance"].mean()) < 0.75
            if "distance_per_min" in week_df.columns and match["distance_per_min"].mean() > 0:
                intensity_gap = (train["distance_per_min"].mean() / match["distance_per_min"].mean()) < 0.88
        elif not train.empty:
            if "sprint_distance" in train.columns and train["sprint_distance"].mean() < 250:
                low_sprint = True

    if readiness_score < 60:
        global_tone = "frissítés és terheléskontroll"
        volume = "alacsony-közepes"
    elif "alulterhelt" in str(periodization_type).lower() or low_sprint:
        global_tone = "kontrollált inger pótlása"
        volume = "közepes"
    else:
        global_tone = "struktúra megtartása és finomhangolás"
        volume = "közepes"

    md_structure = [
        ("MD-4", "Fő terhelési nap", "volumen + játékmodell-specifikus intenzitás"),
        ("MD-3", "Sebesség / intenzitás nap", "rövid maximális sebesség vagy high effort blokk"),
        ("MD-2", "Kontrollált taktikai nap", "alacsonyabb neuromuszkuláris teher, kevesebb lassítás"),
        ("MD-1", "Aktiváció", "rövid, frissítő, nem fárasztó inger"),
        ("MD", "Mérkőzés", "referencianap"),
        ("MD+1", "Regeneráció / pótlás", "játszók regeneráció, kevesebbet játszók kiegészítő munka"),
    ]

    rows = []
    for md, role, base_goal in md_structure:
        if md == "MD-4":
            recommendation = "Legyen ez a hét fő terhelési napja. Játékmodellhez illeszkedő játékok, nagyobb volumen."
            if readiness_score < 60:
                recommendation = "Csak óvatos fő nap: ne legyen nagy load spike, inkább kontrollált volumen."
        elif md == "MD-3":
            if low_sprint:
                recommendation = "Tegyél be rövid, kontrollált sprint/max sebesség blokkot. Kevés ismétlés, jó minőség."
            elif intensity_gap:
                recommendation = "Rövidebb, magasabb tempójú játékokkal közelítsd a meccsintenzitást."
            else:
                recommendation = "Tartsd meg a minőségi sebesség/intenzitás ingert, de ne volumenből oldd meg."
        elif md == "MD-2":
            recommendation = "Taktikai fókusz, kontrollált terhelés. Kerüld a túl sok lassítást és excentrikus terhet."
            if high_risk_players:
                recommendation += " Magas risk játékosoknál egyéni limit javasolt."
        elif md == "MD-1":
            recommendation = "Rövid aktiváció, frissesség megtartása. Ne legyen új terhelési inger."
        elif md == "MD":
            recommendation = "Mérkőzés referencia. A következő heti edzésterhelést ehhez viszonyítsd."
        else:
            recommendation = "Regeneráció a sokat játszóknak, kiegészítő kontrollált terhelés a kevesebbet játszóknak."

        rows.append({
            "Nap": md,
            "Szerep": role,
            "Fő cél": base_goal,
            "Ajánlott terhelés": volume if md in ["MD-4", "MD-3"] else "alacsony-közepes" if md in ["MD-2", "MD+1"] else "alacsony" if md == "MD-1" else "meccs",
            "Javaslat": recommendation,
            "Tervezési alap": f"A jelenlegi hét értelmezése: {global_tone}. Readiness: {readiness_score}/100, periodizáció: {periodization_type}.",
        })

    return pd.DataFrame(rows)


def build_player_next_actions(player_risk_df: pd.DataFrame, df: pd.DataFrame, selected_week: str) -> pd.DataFrame:
    """Játékosszintű holnapi/következő heti teendők."""
    if player_risk_df is None or player_risk_df.empty:
        return pd.DataFrame(columns=["Játékos", "Prioritás", "Holnap / következő edzés", "Következő hét", "Indok"])

    rows = []
    for _, r in player_risk_df.head(12).iterrows():
        name = r.get("Játékos", r.get("player_name", ""))
        level = str(r.get("Kockázati szint", ""))
        score = r.get("Kockázati pontszám", r.get("Risk score", ""))
        reason = r.get("Fő okok", r.get("Fő ok", ""))

        if level == "Magas":
            tomorrow = "Terheléskontroll, extra sprint/lassítás kerülése, frissességi check-in."
            next_week = "Egyéni limit MD-4/MD-3 napon; minőségi, de alacsony volumenű sebességinger."
            priority = "Magas"
        elif level == "Közepes":
            tomorrow = "Normál edzés, de sprint/lassítás mennyiség figyelése."
            next_week = "Fokozatos terhelés, ne legyen hirtelen load spike."
            priority = "Közepes"
        else:
            tomorrow = "Normál terhelés folytatható."
            next_week = "Csapatprogram szerint, egyéni monitoringgal."
            priority = "Alacsony"

        rows.append({
            "Játékos": name,
            "Prioritás": priority,
            "Holnap / következő edzés": tomorrow,
            "Következő hét": next_week,
            "Indok": f"Risk: {score}. {reason}",
        })

    return pd.DataFrame(rows)


def build_next_week_plan_v5(
    df: pd.DataFrame,
    selected_week: str,
    playstyle: str,
    readiness_score: int,
    periodization_type: str,
    player_risk_df: pd.DataFrame,
    past_review_df: pd.DataFrame,
    current_remaining_df: pd.DataFrame,
) -> Tuple[pd.DataFrame, str]:
    """Jövő hét: múlt hét + aktuális részadat + risk alapján."""
    base_plan = build_next_microcycle_plan(df, selected_week, playstyle, readiness_score, periodization_type, player_risk_df)
    if base_plan is None or base_plan.empty:
        return base_plan, "Nincs elég adat következő hét tervhez."

    high_risk_count = 0
    if player_risk_df is not None and not player_risk_df.empty and "Kockázati szint" in player_risk_df.columns:
        high_risk_count = int((player_risk_df["Kockázati szint"] == "Magas").sum())

    past_warn = 0
    if past_review_df is not None and not past_review_df.empty and "Súlyosság" in past_review_df.columns:
        past_warn = int(past_review_df["Súlyosság"].astype(str).isin(["KRITIKUS", "FIGYELMEZTETÉS"]).sum())

    base_plan = base_plan.copy()
    base_plan["Tervezési alap"] = (
        f"Múlt heti figyelmeztetések: {past_warn}; aktuális readiness: {readiness_score}/100; "
        f"magas risk játékosok: {high_risk_count}; periodizáció: {periodization_type}."
    )

    txt = (
        "A jövő heti mikrociklus a múlt teljesebb értékelésére, az aktuális hét eddig feltöltött napjaira "
        "és a játékosszintű risk jelzésekre épül."
    )
    return base_plan, txt


def build_premium_pdf_bytes(
    insights_df: pd.DataFrame,
    selected_week: str,
    readiness_score: int,
    periodization_type: str,
    weekly_summary_text: str,
    coaching_priorities: List[Dict[str, str]],
    risk_df: pd.DataFrame,
    playstyle: str
) -> Optional[bytes]:
    """Éles, látványos Pro PDF riport valós feltöltött adatokból.
    A minta PDF vizuális és tartalmi logikáját követi.
    """
    if SimpleDocTemplate is None:
        return None

    try:
        from reportlab.platypus import PageBreak
    except Exception:
        PageBreak = None

    font_name, font_bold = _register_pdf_font()
    output = io.BytesIO()

    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=0.7 * cm,
        leftMargin=0.7 * cm,
        topMargin=0.65 * cm,
        bottomMargin=0.65 * cm,
    )

    styles = getSampleStyleSheet()

    title = ParagraphStyle(
        "LiveReportTitle",
        parent=styles["Title"],
        fontName=font_bold,
        fontSize=23,
        leading=27,
        alignment=1,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=5,
    )
    subtitle = ParagraphStyle(
        "LiveReportSubtitle",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#475569"),
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "LiveReportH2",
        parent=styles["Heading2"],
        fontName=font_bold,
        fontSize=15,
        leading=18,
        textColor=colors.HexColor("#1F4E78"),
        spaceBefore=7,
        spaceAfter=5,
    )
    body = ParagraphStyle(
        "LiveReportBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8.2,
        leading=10.5,
        textColor=colors.HexColor("#111827"),
    )
    small = ParagraphStyle(
        "LiveReportSmall",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=7.1,
        leading=9,
        textColor=colors.HexColor("#111827"),
    )
    tiny = ParagraphStyle(
        "LiveReportTiny",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=6.5,
        leading=8.2,
        textColor=colors.HexColor("#111827"),
    )
    header = ParagraphStyle(
        "LiveReportHeader",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=7.2,
        leading=9,
        alignment=1,
        textColor=colors.white,
    )
    kpi_label = ParagraphStyle(
        "KpiLabel",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=7.2,
        leading=9,
        alignment=1,
        textColor=colors.white,
    )
    kpi_value = ParagraphStyle(
        "KpiValue",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=17,
        leading=20,
        alignment=1,
        textColor=colors.white,
    )
    kpi_note = ParagraphStyle(
        "KpiNote",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=7,
        leading=8.5,
        alignment=1,
        textColor=colors.white,
    )

    def clean_text(v: object) -> str:
        txt = pdf_safe_text(v)
        txt = str(txt or "")
        txt = txt.replace("\\r", "").replace("\\n", "\n")
        return txt

    def P(v, style=body):
        return Paragraph(html.escape(clean_text(v)).replace("\n", "<br/>"), style)

    def section_bar(text: str):
        t = Table([[P(text, h2)]], colWidths=[27.7 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#DFF2FF")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#93C5FD")),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    def data_table(data, col_widths, header_bg="#0F172A", row_bgs=None):
        if row_bgs is None:
            row_bgs = [colors.white, colors.HexColor("#F8FAFC")]
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header_bg)),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), row_bgs),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    def severity_count(label: str) -> int:
        if insights_df is None or insights_df.empty or "Súlyosság" not in insights_df.columns:
            return 0
        return int((insights_df["Súlyosság"].astype(str).str.upper() == label).sum())

    high_risk_count, medium_risk_count = _fpi_count_risk_levels_v126(risk_df)

    critical_count = severity_count("KRITIKUS")
    warning_count = severity_count("FIGYELMEZTETÉS")
    insight_count = len(insights_df) if insights_df is not None else 0

    sprint_fit = "—"
    load_change = "—"
    try:
        full_text = " ".join(insights_df.astype(str).values.flatten().tolist()) if insights_df is not None and not insights_df.empty else ""
        m = re.search(r"kb\.\s*(\d+)%", full_text)
        if m:
            sprint_fit = f"{m.group(1)}%"
        m2 = re.search(r"(\d+)%[-–]kal nőtt", full_text)
        if m2:
            load_change = f"+{m2.group(1)}%"
    except Exception:
        pass

    readiness_label = score_to_label(readiness_score)
    top_risk_label = f"{high_risk_count} fő" if high_risk_count else "OK"

    story = []
    story.append(P("Football Performance Intelligence", title))
    story.append(P(
        f"Éles vezetői riport | Hét: {format_week_label(selected_week)} | Játékmodell: {playstyle} | Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        subtitle,
    ))

    # KPI strip
    kpis = [
        ("TEAM READINESS", f"{readiness_score}/100", readiness_label, "#166534" if readiness_score >= 75 else "#1E3A8A" if readiness_score >= 60 else "#991B1B"),
        ("LOAD CHANGE", load_change, "heti terhelési jel", "#8B2C13"),
        ("SPRINT FIT", sprint_fit, "edzés vs meccsigény", "#1E3A8A"),
        ("TOP RISK", top_risk_label, "magas egyéni kockázat", "#7F1D1D" if high_risk_count else "#166534"),
        ("INSIGHT", str(insight_count), "automatikus megállapítás", "#0F172A"),
    ]

    kpi_tables = []
    for label, value, note, color in kpis:
        kt = Table(
            [[P(label, kpi_label)], [P(value, kpi_value)], [P(note, kpi_note)]],
            colWidths=[5.1 * cm],
            rowHeights=[0.55 * cm, 0.9 * cm, 0.55 * cm],
        )
        kt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(color)),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.white),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        kpi_tables.append(kt)

    strip = Table([kpi_tables], colWidths=[5.25 * cm] * 5)
    strip.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story += [strip, Spacer(1, 0.28 * cm)]

    # Executive summary
    story.append(section_bar("Vezetői összefoglaló – mit kell tudni 30 másodperc alatt?"))

    summary_clean = clean_text(weekly_summary_text)
    summary_lines = [x.strip().lstrip("-• ").strip() for x in summary_clean.splitlines() if x.strip()]
    main_msg = ""
    observation = ""
    recommendation = ""
    second_topic = ""
    for line in summary_lines:
        if line.startswith("Legfontosabb üzenet:"):
            main_msg = line.replace("Legfontosabb üzenet:", "").strip()
        elif line.startswith("Mit látunk?"):
            observation = line.replace("Mit látunk?", "").strip()
        elif line.startswith("Javaslat:"):
            recommendation = line.replace("Javaslat:", "").strip()
        elif line.startswith("Második fontos téma:"):
            second_topic = line.replace("Második fontos téma:", "").strip()

    if not main_msg and summary_lines:
        main_msg = summary_lines[0]
    if not observation:
        observation = "A heti adatok alapján több automatikus performance jelzés készült."
    if not recommendation:
        recommendation = "A következő edzés tervezésénél a readiness, a sprintinger és az egyéni risk jelzések együtt értelmezendők."
    if not second_topic:
        second_topic = "Egyéni és mikrociklus kontroll"

    exec_data = [
        [P("Fő üzenet", header), P("Értelmezés", header), P("Azonnali döntés", header)],
        [
            P(main_msg, body),
            P(f"{observation}\nMeccskészültség: {readiness_score}/100 ({readiness_label}). Periodizáció: {periodization_type}.", body),
            P(recommendation, body),
        ],
        [
            P(second_topic, body),
            P(f"Magas risk: {high_risk_count} fő, közepes risk: {medium_risk_count} fő. Insightok száma: {insight_count}.", body),
            P("A top edzői teendőket és játékos risk sort érdemes stábértekezleten külön átvenni.", body),
        ],
    ]
    story.append(data_table(exec_data, [5.9 * cm, 9.9 * cm, 11.9 * cm], header_bg="#1E3A8A", row_bgs=[colors.HexColor("#F8FAFC"), colors.HexColor("#EFF6FF")]))
    story += [Spacer(1, 0.25 * cm)]

    # Top insights
    story.append(section_bar("Top automatikus megállapítások"))
    insight_cols = ["Megállapítás", "Súlyosság", "Mit látunk?", "Javaslat"]
    insight_cols = [c for c in insight_cols if insights_df is not None and c in insights_df.columns]
    top_data = [[P("#", header)] + [P(c, header) for c in insight_cols]]
    if insights_df is not None and not insights_df.empty and insight_cols:
        for i, (_, r) in enumerate(insights_df.head(6).iterrows(), 1):
            top_data.append([P(str(i), small)] + [P(r.get(c, ""), small) for c in insight_cols])
    else:
        top_data.append([P("1", small), P("Nincs kiemelt megállapítás", small), P("INFORMÁCIÓ", small), P("A feltöltött adatok alapján nincs kritikus jelzés.", small), P("Normál monitoring folytatható.", small)])

    widths = [0.8 * cm, 5.3 * cm, 3.0 * cm, 10.0 * cm, 8.6 * cm][:len(top_data[0])]
    story.append(data_table(top_data, widths, header_bg="#0F172A"))

    if PageBreak:
        story.append(PageBreak())
    else:
        story.append(Spacer(1, 1 * cm))

    # PAGE 2
    story.append(P("Edzői döntéstámogatás", title))
    story.append(P("Prioritások, játékos risk és gyakorlati javaslatok valós feltöltött adatok alapján.", subtitle))
    story.append(section_bar("Top edzői teendők"))

    pr_data = [[P("Prioritás", header), P("Teendő", header), P("Miért fontos?", header), P("Mikor?", header)]]
    if coaching_priorities:
        for i, item in enumerate(coaching_priorities[:6], 1):
            pr_data.append([
                P(str(i), small),
                P(item.get("Teendő", item.get("Cím", "")), small),
                P(item.get("Miért", ""), small),
                P(item.get("Mikor", "Következő edzés / heti review"), small),
            ])
    else:
        pr_data.append([P("1", small), P("Normál monitoring", small), P("Nincs kiemelt edzői beavatkozás.", small), P("Heti review", small)])

    story.append(data_table(pr_data, [1.8 * cm, 10.0 * cm, 10.0 * cm, 5.9 * cm], header_bg="#166534", row_bgs=[colors.HexColor("#ECFDF5"), colors.HexColor("#F8FAFC")]))
    story += [Spacer(1, 0.25 * cm)]

    story.append(section_bar("Játékos risk tábla – vezetői gyorsnézet"))
    risk_cols = ["Játékos", "Kockázati szint", "Kockázati pontszám", "Fő okok"]
    risk_cols = [c for c in risk_cols if risk_df is not None and c in risk_df.columns]
    risk_data = [[P(c, header) for c in risk_cols]]
    if risk_df is not None and not risk_df.empty and risk_cols:
        for _, r in risk_df.head(10).iterrows():
            risk_data.append([P(r.get(c, ""), small) for c in risk_cols])
    else:
        risk_cols = ["Játékos", "Kockázati szint", "Kockázati pontszám", "Fő okok"]
        risk_data = [[P(c, header) for c in risk_cols]]
        risk_data.append([P("Nincs magas risk", small), P("Alacsony", small), P("—", small), P("A feltöltött hét alapján nincs kiemelt egyéni kockázat.", small)])

    risk_widths = {
        "Játékos": 4.7 * cm,
        "Kockázati szint": 3.0 * cm,
        "Kockázati pontszám": 3.0 * cm,
        "Fő okok": 17.0 * cm,
    }
    story.append(data_table(risk_data, [risk_widths.get(c, 4 * cm) for c in risk_cols], header_bg="#7F1D1D", row_bgs=[colors.white, colors.HexColor("#FFF7ED")]))

    if PageBreak:
        story.append(PageBreak())
    else:
        story.append(Spacer(1, 1 * cm))

    # PAGE 3
    story.append(P("Mikrociklus és játékmodell illeszkedés", title))
    story.append(P("A riport nem csak adatot mutat, hanem edzői döntésre fordítja le a GPS-profilt.", subtitle))
    story.append(section_bar("Mikrociklus szerkezet – automatikus értelmezés"))

    micro_rows = [
        ["Komponens", "Aktuális jel", "Értékelés"],
        ["Readiness", f"{readiness_score}/100", readiness_label],
        ["Periodizáció", periodization_type, "A heti load és frissességi jelzések alapján."],
        ["Sprint fit", sprint_fit, "A meccsigényhez viszonyított sprintinger becsült jele."],
        ["Magas risk", f"{high_risk_count} fő", "Egyéni kontroll szükséges, ha 1 főnél több magas risk van."],
        ["Insight súlyosság", f"{critical_count} kritikus / {warning_count} figyelmeztetés", "A heti fő fókusz a kritikus és figyelmeztető jelzésekből jön."],
    ]
    story.append(data_table([[P(c, header) for c in micro_rows[0]]] + [[P(c, small) for c in row] for row in micro_rows[1:]], [5.8 * cm, 6.8 * cm, 15.1 * cm], header_bg="#1E3A8A", row_bgs=[colors.HexColor("#EFF6FF"), colors.white]))
    story += [Spacer(1, 0.25 * cm)]

    story.append(section_bar(f"Játékmodell illeszkedés – {playstyle} profil"))

    if str(playstyle).lower() == "pressing":
        components = [
            ["Táv/perc", "Meccsprofil 90%+", "Ellenőrizendő", "Pressing modellhez magas munkasűrűség kell."],
            ["High effort", "Meccsprofil 75%+", "Ellenőrizendő", "Ismételt intenzív akciók kulcsfontosságúak."],
            ["Sprintprofil", "Meccsprofil 80%+", sprint_fit, "Célzott sprintinger javasolt, ha alacsony."],
            ["Lassítási terhelés", "Kontrollált", "Risk alapján", "Excentrikus/regenerációs kockázatot jelezhet."],
        ]
    else:
        components = [
            ["Táv/perc", "Játékmodellhez illeszkedő", "Ellenőrizendő", "A heti tempó illeszkedjen a meccsigényhez."],
            ["Sprintprofil", "Meccsprofilhoz közeli", sprint_fit, "A gyors munka ne maradjon tartósan alacsony."],
            ["Load kontroll", "Ne legyen nagy kiugrás", load_change, "A hirtelen terhelésemelkedés readiness kockázatot adhat."],
            ["Játékos risk", "Egyéni eltérések kontrollja", top_risk_label, "A csapatátlag elfedhet egyéni problémákat."],
        ]

    jd = [[P("Komponens", header), P("Cél", header), P("Aktuális hét", header), P("Értékelés", header)]]
    for row in components:
        jd.append([P(c, small) for c in row])

    story.append(data_table(jd, [5.4 * cm, 6.2 * cm, 4.6 * cm, 11.5 * cm], header_bg="#312E81", row_bgs=[colors.white, colors.HexColor("#F5F3FF")]))
    story += [Spacer(1, 0.25 * cm)]

    story.append(section_bar("Vizuális gyorsjelentés – vezetői dashboard nézet"))
    dash_kpis = [
        ("READINESS", f"{readiness_score}%", "meccskészültség", "#166534" if readiness_score >= 75 else "#1E3A8A"),
        ("SPRINT FIT", sprint_fit, "sebességkitettség", "#1E3A8A"),
        ("LOAD", load_change, "heti terhelésjel", "#8B2C13"),
        ("PLAYER RISK", top_risk_label, "egyéni kontroll", "#7F1D1D" if high_risk_count else "#166534"),
        ("COACHING", f"{len(coaching_priorities[:6])} task", "automatikus teendő", "#0F172A"),
    ]

    mini_tables = []
    for label, val, note, color in dash_kpis:
        mt = Table(
            [[P(label, kpi_label)], [P(val, kpi_value)], [P(note, kpi_note)]],
            colWidths=[5.1 * cm],
            rowHeights=[0.5 * cm, 0.8 * cm, 0.5 * cm],
        )
        mt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(color)),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.white),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        mini_tables.append(mt)
    story.append(Table([mini_tables], colWidths=[5.25 * cm] * 5))

    if PageBreak:
        story.append(PageBreak())
    else:
        story.append(Spacer(1, 1 * cm))

    # PAGE 4
    story.append(P("Teljes Pro riport – részletes insight tábla", title))
    story.append(P("A Pro export célja, hogy a vezetői összkép mellett a szakmai stáb részletes megállapításokat is megkapjon.", subtitle))
    story.append(section_bar("Mit kap ebből a stáb?"))

    pro_rows = [
        ["Modul", "Mit ad?", "Kinek hasznos?"],
        ["Vezetői riport", "30 másodperces összefoglaló, readiness, top risk, top teendő.", "sportigazgató, vezetőedző"],
        ["Mikrociklus", "MD-logika, tapering, sprintinger, frissességi jelzések.", "vezetőedző, erőnléti edző"],
        ["Játékos risk motor", "Egyéni terhelési eltérések, sprint, lassítás, load kiugrás.", "erőnléti edző, rehabilitáció"],
        ["Játékmodell illeszkedés", "Fizikai profil összevetése a választott játékmodellel.", "szakmai stáb"],
        ["Export központ", "PDF, Word, Excel riportok vezetői és szakmai felhasználásra.", "klubvezetés, stáb"],
    ]

    story.append(data_table([[P(c, header) for c in pro_rows[0]]] + [[P(c, small) for c in row] for row in pro_rows[1:]], [5.0 * cm, 14.5 * cm, 8.2 * cm], header_bg="#0F172A", row_bgs=[colors.white, colors.HexColor("#F8FAFC")]))
    story += [Spacer(1, 0.25 * cm)]

    story.append(section_bar("Részletes insight tábla"))
    if insights_df is not None and not insights_df.empty:
        detail_cols = ["Súlyosság", "Terület", "Megállapítás", "Mit látunk?", "Javaslat"]
        detail_cols = [c for c in detail_cols if c in insights_df.columns]
        detail_data = [[P(c, header) for c in detail_cols]]
        for _, r in insights_df.head(12).iterrows():
            detail_data.append([P(r.get(c, ""), tiny) for c in detail_cols])
        detail_widths = {
            "Súlyosság": 3.2 * cm,
            "Terület": 2.4 * cm,
            "Megállapítás": 4.8 * cm,
            "Mit látunk?": 8.6 * cm,
            "Javaslat": 8.7 * cm,
        }
        story.append(data_table(detail_data, [detail_widths.get(c, 4 * cm) for c in detail_cols], header_bg="#1F4E78", row_bgs=[colors.white, colors.HexColor("#F7F9FC")]))
    else:
        story.append(P("Nincs részletes insight adat.", body))

    doc.build(story)
    output.seek(0)
    return output.getvalue()


# -----------------------------------------------------------------------------
# V3.1 - Magyar demo oldal + Executive Export Center
# -----------------------------------------------------------------------------
def render_system_intro_page() -> None:
    st.markdown(
        """
        <div class="intro-card">
            <h2>⚽ Football Performance Intelligence Platform</h2>
            <div class="hero-sub">
                Automatikus heti teljesítményintelligencia futballcsapatok számára.
                Ez nem egyszerű GPS dashboard, hanem egy döntéstámogató performance engine,
                amely football kontextusban értelmezi a terhelési adatokat.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Mit csinálunk?")
    st.markdown(
        """
        A rendszer a feltöltött GPS / terhelési adatokat automatikusan feldolgozza, majd
        edzői nyelvre fordítja: kiemeli a legfontosabb heti megállapításokat,
        figyelmeztet a problémás mintázatokra, és javaslatokat ad a következő döntésekhez.
        """
    )

    st.markdown("### Mitől extra?")
    st.markdown(
        """
        Nem csak számokat és grafikonokat mutat. A rendszer figyelembe veszi a mikrociklust,
        a meccsnaphoz viszonyított terhelést, a játékmodellt, a többhetes mintázatokat,
        a játékosok egyéni eltéréseit és a match readiness állapotot.
        """
    )

    st.markdown(
        """
        <div class="intro-grid">
            <div class="feature-box">
                <div class="feature-title">1. Mikrociklus intelligencia</div>
                <div class="feature-text">MD-4 / MD-3 / MD-2 / MD-1 logika, tapering, maximális sebességű inger és heti struktúra értékelése.</div>
            </div>
            <div class="feature-box">
                <div class="feature-title">2. Meccskészültség score</div>
                <div class="feature-text">0–100-as meccskészültségi pontszám load trend, frissesség, maximális sebességű inger és játékmodell alapján.</div>
            </div>
            <div class="feature-box">
                <div class="feature-title">3. Játékos kockázati motor</div>
                <div class="feature-text">Automatikusan jelzi, ha egy játékos terhelése, sprintprofilja vagy max sebessége eltér a saját múltjától.</div>
            </div>
            <div class="feature-box">
                <div class="feature-title">4. Performance memória</div>
                <div class="feature-text">Többhetes történetet épít, trendeket keres, és nem csak az aktuális hetet nézi elszigetelten.</div>
            </div>
            <div class="feature-box">
                <div class="feature-title">5. Coaching priorities</div>
                <div class="feature-text">Nem 40 figyelmeztetést ad, hanem kiemeli a hét 3 legfontosabb edzői teendőjét.</div>
            </div>
            <div class="feature-box">
                <div class="feature-title">6. Executive export</div>
                <div class="feature-text">Egy kattintással vezetőedzői PDF / Word / Excel riport készül magyar nyelven.</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Kinek készült?")
    st.markdown(
        """
        Elsősorban NB2 / NB3 kluboknak, akadémiáknak, utánpótlás elitműhelyeknek
        és olyan stáboknak, ahol van GPS adat, de nincs külön analyst vagy sport scientist csapat.
        """
    )

    st.markdown("### Egy mondatban")
    st.success("A rendszer analyst gondolkodást ad a stábnak, de nem kell hozzá külön analyst csapat.")

    st.markdown("### Egyszerű nyelven a fő fogalmak")
    render_wrapped_table(build_plain_language_explanation())


def build_executive_summary_df(
    selected_week: str,
    selected_playstyle: str,
    readiness_score: int,
    periodization_type: str,
    weekly_summary_text: str,
    high_risk_count: int,
    medium_risk_count: int,
) -> pd.DataFrame:
    return pd.DataFrame([
        {"Elem": "Hét", "Érték": selected_week},
        {"Elem": "Játékmodell", "Érték": selected_playstyle},
        {"Elem": "Meccskészültség", "Érték": f"{readiness_score}/100 – {score_to_label(readiness_score)}"},
        {"Elem": "Periodizáció", "Érték": periodization_type},
        {"Elem": "Magas kockázatos játékosok", "Érték": str(high_risk_count)},
        {"Elem": "Közepes kockázatos játékosok", "Érték": str(medium_risk_count)},
        {"Elem": "Heti vezetői összefoglaló", "Érték": weekly_summary_text},
    ])


def build_executive_excel_bytes(
    executive_df: pd.DataFrame,
    insights_df: pd.DataFrame,
    priorities_df: pd.DataFrame,
    risk_df: pd.DataFrame,
    weekly_fingerprints: pd.DataFrame,
) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        executive_df.to_excel(writer, index=False, sheet_name="Vezetői összefoglaló")
        insights_df.to_excel(writer, index=False, sheet_name="Insightok")
        priorities_df.to_excel(writer, index=False, sheet_name="Edzői teendők")
        if risk_df is not None and not risk_df.empty:
            risk_df.to_excel(writer, index=False, sheet_name="Játékos risk")
        if weekly_fingerprints is not None and not weekly_fingerprints.empty:
            weekly_fingerprints.to_excel(writer, index=False, sheet_name="Memory trendek")

        from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"
            header_fill = PatternFill("solid", fgColor="0F172A")
            header_font = Font(color="FFFFFF", bold=True)
            thin = Side(style="thin", color="CBD5E1")
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

            for row in ws.iter_rows(min_row=2):
                max_lines = 1
                for cell in row:
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
                    text = str(cell.value or "")
                    max_lines = max(max_lines, min(8, len(text) // 45 + 1))
                ws.row_dimensions[row[0].row].height = max(24, max_lines * 16)

            for idx, col_cells in enumerate(ws.columns, start=1):
                values = [str(c.value or "") for c in col_cells[:80]]
                width = min(68, max(14, max(len(v) for v in values[:80]) + 2))
                if ws.title in ["Insightok", "Edzői teendők"]:
                    width = min(60, max(width, 28))
                ws.column_dimensions[get_column_letter(idx)].width = width

    return output.getvalue()


def build_executive_word_bytes(
    executive_df: pd.DataFrame,
    priorities_df: pd.DataFrame,
    insights_df: pd.DataFrame,
    risk_df: pd.DataFrame,
    selected_week: str,
) -> Optional[bytes]:
    if Document is None:
        return None
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    doc.add_heading("Football Performance Intelligence – vezetői riport", level=1)
    doc.add_paragraph(f"Hét: {selected_week}")
    doc.add_paragraph(f"Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Vezetői összefoglaló", level=2)
    for _, row in executive_df.iterrows():
        p = doc.add_paragraph()
        p.add_run(f"{row.get('Elem', '')}: ").bold = True
        p.add_run(str(row.get("Érték", "")))

    doc.add_heading("Top edzői teendők", level=2)
    for idx, row in priorities_df.iterrows():
        doc.add_heading(f"{idx + 1}. {row.get('Cím', '')}", level=3)
        p = doc.add_paragraph()
        p.add_run("Teendő: ").bold = True
        p.add_run(str(row.get("Teendő", "")))
        p = doc.add_paragraph()
        p.add_run("Miért fontos: ").bold = True
        p.add_run(str(row.get("Miért", "")))

    if risk_df is not None and not risk_df.empty:
        doc.add_heading("Top játékos kockázat", level=2)
        for _, row in risk_df.head(8).iterrows():
            p = doc.add_paragraph()
            p.add_run(f"{row.get('Játékos', '')} – {row.get('Kockázati szint', '')} ({row.get('Kockázati pontszám', '')}/100): ").bold = True
            p.add_run(str(row.get("Fő okok", "")))

    doc.add_heading("Insightok", level=2)
    for _, row in insights_df.iterrows():
        doc.add_heading(str(row.get("Megállapítás", "")), level=3)
        for label in ["Súlyosság", "Terület", "Mit látunk?", "Miért fontos?", "Javaslat"]:
            if label in insights_df.columns:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(row.get(label, "")))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()



# -----------------------------------------------------------------------------
# V3.2 - Magyar, edzői nyelvű kommunikációs réteg
# -----------------------------------------------------------------------------
def coach_friendly_phrase(text: object) -> str:
    """Analyst-jellegű megfogalmazásokat fordít edzőbarát magyar nyelvre."""
    if text is None or (isinstance(text, float) and pd.isna(text)):
        return ""
    s = str(text)

    replacements = {
        "A következő 1–2 edzésen ne a plusz terhelés legyen a fő cél, hanem a frissesség stabil megtartása.": 
            "A következő 1–2 edzésen ne a plusz terhelés legyen a fő cél, hanem a frissesség stabil megtartása. Ha hiányzik valamilyen inger, azt röviden és célzottan érdemes pótolni.",
        "A csapat nincs rossz állapotban, de nem is tűnik teljesen frissnek. Ilyenkor a túl nagy változtatás visszaüthet.": 
            "A csapat nincs rossz állapotban, de nem is tűnik teljesen frissnek. Ilyenkor a túl nagy változtatás könnyen visszaüthet.",
        "Tisztázd, hogy ez szándékosan könnyebb hét volt-e. Ha igen, rendben lehet; ha nem, akkor hiányozhatott a megfelelő edzésinger.": 
            "Érdemes tisztázni, hogy ez a hét szándékosan volt-e könnyebb. Ha igen, rendben lehet; ha nem, akkor hiányozhatott a megfelelő edzésinger.",
        "Az alacsonyabb terhelés rövid távon frissíthet, de ha több héten át így marad, a csapat intenzitása visszaeshet.": 
            "Az alacsonyabb terhelés rövid távon frissíthet, de ha több héten át így marad, a csapat intenzitása visszaeshet.",
        "Érdemes célzottan megvizsgálni, hogy tudatos terheléscsökkentésről vagy nem kívánt intenzitáshiányról van-e szó.": 
            "Nézd meg, hogy szándékosan csökkent-e a gyors munka. Ha nem, akkor érdemes visszatenni egy rövid, kontrollált maximális sebességű blokkot.",
        "A tartós speed exposure hiány hosszabb távon ronthatja a maximális sebesség és intenzitás fenntartását.": 
            "Ha hetekig kevés a maximális sebességű inger, a játékosok nehezebben tudják fenntartani a meccshez szükséges gyorsaságot és intenzitást.",
        "Meccskészültség kontroll": "Meccskészültség kontroll",
        "Könnyebb hét értelmezése": "Könnyebb hét értelmezése",
        "Többhetes sprintcsökkenés": "Több hete csökken a gyors munka",
        "speed exposure": "maximális sebességű inger",
        "Speed exposure": "Maximális sebességű inger",
        "load": "terhelés",
        "Load": "Terhelés",
        "readiness": "meccskészültség",
        "Readiness": "Meccskészültség",
        "risk": "kockázat",
        "Risk": "Kockázat",
        "player": "játékos",
        "Player": "Játékos",
        "overload": "túlterhelés",
        "Overload": "Túlterhelés",
        "recovery": "regeneráció",
        "Recovery": "Regeneráció",
    }

    for old, new in replacements.items():
        s = s.replace(old, new)

    return s


def humanize_insight(insight: Insight) -> Insight:
    """Insight objektum edzőbarát magyarítása."""
    return Insight(
        title=coach_friendly_phrase(insight.title),
        severity=insight.severity,
        observation=coach_friendly_phrase(insight.observation),
        impact=coach_friendly_phrase(insight.impact),
        recommendation=coach_friendly_phrase(insight.recommendation),
        scope=coach_friendly_phrase(insight.scope),
    )


def humanize_insights(insights: List[Insight]) -> List[Insight]:
    return [humanize_insight(i) for i in insights]


def humanize_priority_item(item: Dict[str, str]) -> Dict[str, str]:
    return {k: coach_friendly_phrase(v) for k, v in item.items()}


def humanize_priority_list(items: List[Dict[str, str]]) -> List[Dict[str, str]]:
    return [humanize_priority_item(i) for i in items]


def build_plain_language_explanation() -> pd.DataFrame:
    """Rövid magyarázó tábla a vezetőedzőnek: mit jelent egy-egy fogalom."""
    return pd.DataFrame([
        {
            "Fogalom": "Meccskészültség",
            "Egyszerű jelentés": "Mennyire tűnik késznek a csapat a következő meccsterhelésre.",
            "Mit nézünk?": "Terhelési trend, frissesség, MD-1/MD-2 terhelés, sprintinger, játékmodellhez való illeszkedés.",
        },
        {
            "Fogalom": "Könnyebb hét",
            "Egyszerű jelentés": "A hét terhelése alacsonyabb volt a megszokottnál.",
            "Mit nézünk?": "Ez lehet tudatos frissítés, de lehet nem kívánt alulterhelés is.",
        },
        {
            "Fogalom": "Maximális sebességű inger",
            "Egyszerű jelentés": "Kapott-e a csapat/játékos elég nagy sebességű futást a héten.",
            "Mit nézünk?": "Sprinttáv, max sebesség, meccsigényhez viszonyított arány.",
        },
        {
            "Fogalom": "Többhetes sprintcsökkenés",
            "Egyszerű jelentés": "Hetek óta csökken a gyors futások mennyisége.",
            "Mit nézünk?": "Sprinttáv trend, nagy sebességű táv, maximális sebesség.",
        },
        {
            "Fogalom": "Játékos kockázat",
            "Egyszerű jelentés": "Egy játékos mennyire tér el a saját megszokott terhelési profiljától.",
            "Mit nézünk?": "Terhelési ugrás, lassítások, sprinttáv, max sebesség-visszaesés.",
        },
    ])





# -----------------------------------------------------------------------------
# Marketing / demo minta PDF riport - V4.2
# -----------------------------------------------------------------------------
def build_marketing_sample_pdf_bytes() -> Optional[bytes]:
    """Látványos, többoldalas, kamu adatokkal készült teljes vezetői riport.
    Nem függ feltöltött adattól, ezért a kezdőoldalon is letölthető.
    """
    if SimpleDocTemplate is None:
        return None

    from reportlab.platypus import PageBreak, KeepTogether
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    font_name, font_bold = _register_pdf_font()
    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=landscape(A4),
        rightMargin=.85*cm,
        leftMargin=.85*cm,
        topMargin=.75*cm,
        bottomMargin=.75*cm,
    )

    styles = getSampleStyleSheet()
    title = ParagraphStyle("FPITitle", parent=styles["Title"], fontName=font_bold, fontSize=23, leading=27, textColor=colors.HexColor("#0f172a"), spaceAfter=8)
    subtitle = ParagraphStyle("FPISub", parent=styles["BodyText"], fontName=font_name, fontSize=9, leading=11, textColor=colors.HexColor("#475569"))
    h2 = ParagraphStyle("FPIH2", parent=styles["Heading2"], fontName=font_bold, fontSize=14, leading=17, textColor=colors.HexColor("#1e3a8a"), spaceBefore=6, spaceAfter=6)
    h3 = ParagraphStyle("FPIH3", parent=styles["Heading3"], fontName=font_bold, fontSize=11, leading=13, textColor=colors.HexColor("#0f172a"), spaceBefore=4, spaceAfter=4)
    body = ParagraphStyle("FPIBody", parent=styles["BodyText"], fontName=font_name, fontSize=8.3, leading=10.2, textColor=colors.HexColor("#111827"))
    body_white = ParagraphStyle("FPIBodyWhite", parent=styles["BodyText"], fontName=font_name, fontSize=8.2, leading=10.2, textColor=colors.white)
    small = ParagraphStyle("FPISmall", parent=styles["BodyText"], fontName=font_name, fontSize=7.3, leading=9, textColor=colors.HexColor("#64748b"))
    header = ParagraphStyle("FPIHeader", parent=styles["BodyText"], fontName=font_bold, fontSize=7.7, leading=9.2, textColor=colors.white, alignment=TA_CENTER)
    kpi_label = ParagraphStyle("KPILabel", parent=styles["BodyText"], fontName=font_bold, fontSize=7, leading=8, textColor=colors.HexColor("#bfdbfe"), alignment=TA_CENTER)
    kpi_value = ParagraphStyle("KPIValue", parent=styles["BodyText"], fontName=font_bold, fontSize=18, leading=21, textColor=colors.white, alignment=TA_CENTER)

    def P(text, style=body):
        return Paragraph(html.escape(pdf_safe_text(text)).replace("\n", "<br/>"), style)

    def section(title_text):
        return Table([[P(title_text, h2)]], colWidths=[27.2*cm], style=[
            ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#e0f2fe")),
            ("BOX", (0,0), (-1,-1), .5, colors.HexColor("#93c5fd")),
            ("LEFTPADDING", (0,0), (-1,-1), 8),
            ("RIGHTPADDING", (0,0), (-1,-1), 8),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ])

    def kpi_card(label, value, note, color="#1e3a8a"):
        data = [[P(label, kpi_label)], [P(value, kpi_value)], [P(note, body_white)]]
        t = Table(data, colWidths=[5.15*cm], rowHeights=[.55*cm, 1.0*cm, .8*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), colors.HexColor(color)),
            ("BOX", (0,0), (-1,-1), .6, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        return t

    def styled_table(data, widths, header_color="#0f172a", body_bg="#ffffff", alt_bg="#f8fafc"):
        t = Table(data, colWidths=widths, repeatRows=1)
        style = [
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor(header_color)),
            ("GRID", (0,0), (-1,-1), .35, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 5),
            ("RIGHTPADDING", (0,0), (-1,-1), 5),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ("BACKGROUND", (0,1), (-1,-1), colors.HexColor(body_bg)),
        ]
        for r in range(2, len(data), 2):
            style.append(("BACKGROUND", (0,r), (-1,r), colors.HexColor(alt_bg)))
        t.setStyle(TableStyle(style))
        return t

    story = []

    # PAGE 1 - Executive cover/dashboard
    story.append(P("Football Performance Intelligence", title))
    story.append(P("Teljes minta vezetői riport | Demo FC U19 | 2026-W22 | Játékmodell: Pressing | Minta GPS/terhelési adatok", subtitle))
    story.append(Spacer(1, 7))

    story.append(Table([[
        kpi_card("TEAM READINESS", "78/100", "elfogadható, figyelendő", "#14532d"),
        kpi_card("LOAD CHANGE", "+27%", "heti terhelési növekedés", "#7c2d12"),
        kpi_card("SPRINT FIT", "72%", "edzés vs meccsigény", "#1e3a8a"),
        kpi_card("TOP RISK", "2 fő", "magas egyéni kockázat", "#7f1d1d"),
        kpi_card("MD-1 TAPER", "OK", "frissességi jel pozitív", "#064e3b"),
    ]], colWidths=[5.35*cm]*5))
    story.append(Spacer(1, 8))

    story.append(KeepTogether([
        section("Vezetői összefoglaló - mit kell tudni 30 másodperc alatt?"),
        Spacer(1, 4),
        styled_table([
            [P("Fő üzenet", header), P("Értelmezés", header), P("Azonnali döntés", header)],
            [P("A csapat alapvetően mérkőzésképes, de a hét nem teljesen tiszta."), P("A readiness 78/100, miközben két játékosnál magasabb risk észlelhető. A heti load +27%, a sprintprofil pedig csak 72%-a a meccsigénynek."), P("A következő edzés fő célja ne a plusz terhelés, hanem a frissesség megtartása és az egyéni terheléskontroll legyen.")],
            [P("A mikrociklus struktúrája részben jó."), P("Az MD-1 taper pozitív jel, de az MD-2 terhelése közel került a hét fő terhelési napjához."), P("MD-2 napon kisebb neuromuszkuláris terhelés, MD-1-en rövid aktiváció javasolt.")],
        ], [5.8*cm, 13.0*cm, 8.4*cm], header_color="#1e3a8a", body_bg="#eff6ff")
    ]))
    story.append(Spacer(1, 7))

    story.append(KeepTogether([
        section("Top 5 automatikus megállapítás"),
        Spacer(1, 4),
        styled_table([
            [P("#", header), P("Megállapítás", header), P("Súlyosság", header), P("Mit látunk?", header), P("Javaslat", header)],
            [P("1"), P("Heti terhelési kiugrás"), P("FIGYELMEZTETÉS"), P("A heti edzés terhelési pontértéke 27%-kal nőtt az előző héthez képest."), P("Következő edzésen terheléskontroll és egyéni reakciók figyelése.")],
            [P("2"), P("Alacsony sprintinger"), P("FIGYELMEZTETÉS"), P("A heti sprintprofil a meccsigény kb. 72%-a."), P("Rövid, kontrollált maximális sebességű blokk, ha a heti cél engedi.")],
            [P("3"), P("MD-2 terhelés magas lehet"), P("FIGYELMEZTETÉS"), P("Az MD-2 load közel volt a hét fő terhelési napjához."), P("Meccs előtt 48 órával alacsonyabb neuromuszkuláris terhelés.")],
            [P("4"), P("Két játékos magas risk zónában"), P("KRITIKUS"), P("Nagy D. és Varga L. egyéni terhelése eltér a csapatprofiltól."), P("Egyéni terheléskorrekció és regenerációs visszajelzés ellenőrzése.")],
            [P("5"), P("Pozitív tapering jel"), P("INFORMÁCIÓ"), P("Az MD-1 terhelés alacsonyabb volt a hét fő napjához képest."), P("A struktúra megtartható, ha a meccsteljesítmény visszaigazolja.")],
        ], [1.0*cm, 5.0*cm, 3.0*cm, 9.0*cm, 9.2*cm], header_color="#0f172a")
    ]))

    story.append(PageBreak())

    # PAGE 2 - Coaching priorities and player risk
    story.append(P("Edzői döntéstámogatás", title))
    story.append(P("Prioritások, játékos risk és gyakorlati javaslatok", subtitle))
    story.append(Spacer(1, 6))

    story.append(KeepTogether([
        section("Top 6 edzői teendő"),
        Spacer(1, 4),
        styled_table([
            [P("Prioritás", header), P("Teendő", header), P("Miért fontos?", header), P("Mikor?", header)],
            [P("1. Sprintterhelés kontroll"), P("Nagy D. és Varga L. kapjon kontrollált nagysebességű ingert, de ne újabb volumennövelő blokkot."), P("Mindkét játékosnál magasabb sprinttáv és terhelési löket látható."), P("Következő edzés")],
            [P("2. MD-2 frissesség"), P("A mérkőzés előtti 48 órában csökkenteni kell a teljes loadot és a lassítási terhelést."), P("A túl magas MD-2 terhelés ronthatja a meccsnapi frissességet."), P("Meccs előtt 2 nap")],
            [P("3. Alacsony terhelésű játékosok"), P("Kiss R. és Farkas Z. részére kiegészítő egyéni munka vagy fokozatos visszaterhelés."), P("Két hete csapatátlag alatti load, lemaradó inger alakulhat ki."), P("Következő 2 edzés")],
            [P("4. High effort monitoring"), P("A nagy intenzitású erőfeszítéseket játékosonként kell figyelni, nem csak csapatszinten."), P("A csapatátlag elfedhet egyéni neuromuszkuláris kockázatot."), P("Minden edzés után")],
            [P("5. Posztcsoport összevetés"), P("Szélsők és középpályások sprint/HSR arányát külön értékelni."), P("A játékmodell más fizikai profilt vár el posztonként."), P("Heti review")],
            [P("6. Frissességi kérdőív"), P("A két magas risk játékosnál rövid szubjektív frissességi check-in."), P("A GPS nem mutatja önmagában a belső terhelésérzetet."), P("Edzés előtt")],
        ], [4.7*cm, 10.0*cm, 9.2*cm, 3.3*cm], header_color="#14532d", body_bg="#f0fdf4")
    ]))
    story.append(Spacer(1, 7))

    story.append(KeepTogether([
        section("Játékos risk tábla - vezetői gyorsnézet"),
        Spacer(1, 4),
        styled_table([
            [P("Játékos", header), P("Poszt", header), P("Risk", header), P("Fő eltérés", header), P("Ajánlott lépés", header)],
            [P("Nagy D."), P("CM"), P("Magas - 82"), P("Sprinttáv +55%, összterhelés +18%"), P("Load kontroll, regenerációs monitor, következő edzésen limitált extra munka.")],
            [P("Varga L."), P("W"), P("Magas - 79"), P("High effort és lassítás kiugrás"), P("Excentrikus terhelés csökkentése, frissesség ellenőrzés.")],
            [P("Farkas Z."), P("DM"), P("Közepes - 63"), P("Max sebesség trend csökkenő"), P("Rövid, kontrollált max sebességű inger.")],
            [P("Kiss R."), P("GK"), P("Közepes - 61"), P("Alacsony heti load"), P("Pozícióspecifikus kiegészítő blokk.")],
            [P("Tóth B."), P("CB"), P("Alacsony - 38"), P("Stabil terhelés, nincs kiugrás"), P("Normál terhelés folytatható.")],
            [P("Mészáros P."), P("ST"), P("Alacsony - 34"), P("Jó sprint kitettség, kontrollált load"), P("Jelenlegi struktúra megtartható.")],
        ], [3.4*cm, 2.0*cm, 2.8*cm, 8.2*cm, 10.8*cm], header_color="#7f1d1d", body_bg="#fff7ed")
    ]))

    story.append(PageBreak())

    # PAGE 3 - Microcycle and model fit
    story.append(P("Mikrociklus és játékmodell illeszkedés", title))
    story.append(P("A riport nem csak adatot mutat, hanem edzői döntésre fordítja le a GPS-profilt.", subtitle))
    story.append(Spacer(1, 6))

    story.append(KeepTogether([
        section("Mikrociklus szerkezet - MD logika"),
        Spacer(1, 4),
        styled_table([
            [P("Nap", header), P("Load index", header), P("Sprint", header), P("HSR", header), P("Értelmezés", header)],
            [P("MD-4"), P("Magas"), P("Közepes"), P("Magas"), P("Fő terhelési nap, megfelelő helyen a mikrociklusban.")],
            [P("MD-3"), P("Közepes"), P("Magas"), P("Közepes"), P("Jó nap maximális sebességű ingerre.")],
            [P("MD-2"), P("Közepesen magas"), P("Alacsony"), P("Közepes"), P("Kicsit magas a meccshez közel, frissességi kockázat.")],
            [P("MD-1"), P("Alacsony"), P("Minimális"), P("Alacsony"), P("Pozitív tapering jel, aktivációs napként megfelelő.")],
            [P("MD"), P("Meccs"), P("Referencia"), P("Referencia"), P("A hét értékelése ehhez viszonyítva történik.")],
        ], [3.0*cm, 3.2*cm, 3.2*cm, 3.2*cm, 14.6*cm], header_color="#1e3a8a", body_bg="#eff6ff")
    ]))
    story.append(Spacer(1, 7))

    story.append(KeepTogether([
        section("Játékmodell illeszkedés - Pressing profil"),
        Spacer(1, 4),
        styled_table([
            [P("Komponens", header), P("Cél", header), P("Aktuális hét", header), P("Értékelés", header)],
            [P("Táv/perc"), P("Meccsprofil 90%+"), P("88%"), P("Közel jó, de még figyelendő.")],
            [P("High effort"), P("Meccsprofil 75%+"), P("69%"), P("Pressing modellhez kissé kevés ismételt intenzív akció.")],
            [P("Sprintprofil"), P("Meccsprofil 80%+"), P("72%"), P("Transition/pressing elemekhez célzott sprintinger kell.")],
            [P("Lassítási terhelés"), P("Kontrollált, ne kiugró"), P("+31%"), P("Regenerációs és excentrikus terhelési kockázat.")],
        ], [5.5*cm, 6.2*cm, 5.0*cm, 10.5*cm], header_color="#312e81", body_bg="#f5f3ff")
    ]))
    story.append(Spacer(1, 7))

    story.append(KeepTogether([
        section("Vizuális gyorsjelentés - mit látna a vezető az appban?"),
        Spacer(1, 4),
        Table([
            [kpi_card("PRESSING FIT", "71%", "játékmodell illeszkedés", "#312e81"),
             kpi_card("MAX SPEED EXP.", "64%", "sebességkitettség", "#1e3a8a"),
             kpi_card("TAPERING", "84%", "meccs előtti frissesség", "#14532d"),
             kpi_card("PLAYER RISK", "2 high", "egyéni kontroll szükséges", "#7f1d1d"),
             kpi_card("COACHING", "6 task", "automatikus teendő", "#0f172a")]
        ], colWidths=[5.35*cm]*5)
    ]))

    story.append(PageBreak())

    # PAGE 4 - What Pro gives
    story.append(P("Mit kap a klub a Pro verzióban?", title))
    story.append(P("A minta riport célja, hogy megmutassa: a rendszer nem táblázatot ad, hanem döntéstámogató vezetői anyagot.", subtitle))
    story.append(Spacer(1, 7))

    story.append(KeepTogether([
        section("Pro riport tartalom"),
        Spacer(1, 4),
        styled_table([
            [P("Modul", header), P("Mit ad?", header), P("Kinek hasznos?", header)],
            [P("Vezetői riport"), P("30 másodperces vezetői összefoglaló, readiness, top risk, top teendő."), P("sportigazgató, vezetőedző")],
            [P("📅 Mikrociklus"), P("MD-4/MD-3/MD-2/MD-1 logika, tapering, sprintinger, frissességi jelzések."), P("vezetőedző, erőnléti edző")],
            [P("Játékos risk motor"), P("Egyéni terhelési eltérések, high effort, sprint, lassítás, load kiugrás."), P("erőnléti edző, rehabilitáció")],
            [P("Játékmodell illeszkedés"), P("Pressing, transition, possession vagy low block fizikai profil összevetése."), P("szakmai stáb")],
            [P("Export központ"), P("PDF, Word, Excel riportok vezetői és szakmai felhasználásra."), P("klubvezetés, stáb")],
            [P("Performance memory"), P("Több hét/szezon trendjei, periodizációs mintázatok, hosszú távú terhelésprofil."), P("klubszintű monitoring")],
        ], [5.1*cm, 14.5*cm, 7.6*cm], header_color="#0f172a", body_bg="#f8fafc")
    ]))
    story.append(Spacer(1, 7))

    story.append(KeepTogether([
        section("Demo vs Pro"),
        Spacer(1, 4),
        styled_table([
            [P("Funkció", header), P("Demo", header), P("Pro", header)],
            [P("Saját adat feltöltése"), P("Igen, korlátozott: max 8 játékos / 3 hét / 5000 sor"), P("Korlátlan")],
            [P("Vezetői riport"), P("Igen, preview"), P("Teljes")],
            [P("PDF / Word / Excel export"), P("Minta PDF és vízjeles előnézet"), P("Teljes export")],
            [P("Performance memory"), P("Nem"), P("Igen")],
            [P("Több szezon"), P("Nem"), P("Igen")],
            [P("Saját benchmark"), P("Nem"), P("Igen")],
        ], [9.0*cm, 9.0*cm, 9.0*cm], header_color="#14532d", body_bg="#f0fdf4")
    ]))

    doc.build(story)
    output.seek(0)
    return output.read()



# -----------------------------------------------------------------------------
# V4.3 - Productized Demo / Pro layer - unique Streamlit export keys
# -----------------------------------------------------------------------------
DEMO_PLAYER_LIMIT = 8
DEMO_WEEK_LIMIT = 3
DEMO_ROW_LIMIT = 5000

# -----------------------------------------------------------------------------
# V4.5 - Emailhez kötött aktiváló kód koncepció
# -----------------------------------------------------------------------------
# Éles verzióban ezt nem a kódban tároljuk, hanem Supabase táblában:
# licenses(email, activation_code_hash, plan, is_active, expires_at, max_users, club_name)
# A felhasználó megadja: email + aktiváló kód.
# Az app lekéri / ellenőrzi a hash-t, és ha aktív, Pro módot ad.
# Mostani MVP-ben marad az egyszerű tesztkód: .


# -----------------------------------------------------------------------------
# Supabase license layer
# -----------------------------------------------------------------------------
def get_secret_value(name: str, default: str = "") -> str:
    try:
        if name in st.secrets:
            return str(st.secrets[name])
    except Exception:
        pass
    return os.getenv(name, default)


def get_supabase_client():
    url = get_secret_value("SUPABASE_URL")
    key = get_secret_value("SUPABASE_ANON_KEY")
    if not url or not key or create_client is None:
        return None
    try:
        return create_client(url, key)
    except Exception:
        return None


def hash_license_key(raw_key: str) -> str:
    salt = get_secret_value("LICENSE_SALT", "performance-intelligence")
    return hashlib.sha256((salt + "::" + str(raw_key).strip()).encode("utf-8")).hexdigest()


def validate_license_supabase(email: str, license_key: str) -> Dict[str, object]:
    email = str(email or "").strip().lower()
    license_key = str(license_key or "").strip()
    if not email or not license_key:
        return {"ok": False, "message": "Add meg az e-mail címet és az aktiváló kódot."}

    fallback = get_secret_value("FALLBACK_PRO_CODE")
    if fallback and license_key == fallback:
        return {
            "ok": True,
            "email": email,
            "plan": "pro",
            "club_name": "Pro klub",
            "team_name": "Pro csapat",
            "license_id": "fallback",
            "message": "Pro hozzáférés aktív.",
        }

    client = get_supabase_client()
    if client is None:
        return {"ok": False, "message": "Supabase kapcsolat nincs beállítva. Demo módban folytatható."}

    try:
        key_hash = hash_license_key(license_key)
        resp = client.rpc("validate_license", {
            "p_email": email,
            "p_license_hash": key_hash,
        }).execute()
        data = resp.data
        if isinstance(data, list):
            data = data[0] if data else None

        if not data:
            return {"ok": False, "message": "Nem található aktív licenc ehhez az e-mailhez és kódhoz."}

        if data.get("ok"):
            return {
                "ok": True,
                "email": data.get("email", email),
                "plan": data.get("plan", "pro"),
                "club_name": data.get("club_name", ""),
                "team_name": data.get("team_name", ""),
                "license_id": data.get("license_id", ""),
                "message": "Pro hozzáférés aktív.",
            }

        return {"ok": False, "message": data.get("message", "A licenc nem aktív vagy lejárt.")}
    except Exception as exc:
        return {"ok": False, "message": f"Licencellenőrzési hiba: {exc}"}


def is_pro_mode() -> bool:
    lic = st.session_state.get("license_status", {})
    return bool(lic.get("ok"))


def is_demo_mode() -> bool:
    return not is_pro_mode()


def render_mode_badge() -> None:
    if is_pro_mode():
        lic = st.session_state.get("license_status", {})
        st.sidebar.success("Pro hozzáférés aktív")
        club = lic.get("club_name") or lic.get("email")
        if club:
            st.sidebar.caption(str(club))
        if lic.get("team_name"):
            st.sidebar.caption(f"Csapat: {lic.get('team_name')}")
    else:
        st.sidebar.success("Demo mód")
        st.sidebar.caption(f"Demo limit: max {DEMO_PLAYER_LIMIT} játékos · max {DEMO_WEEK_LIMIT} hét · max {DEMO_ROW_LIMIT} sor")


def render_license_panel() -> None:
    st.sidebar.markdown("### Belépés")
    if is_pro_mode():
        render_mode_badge()
        if st.sidebar.button("Kijelentkezés", use_container_width=True, key="logout_license"):
            st.session_state.pop("license_status", None)
            st.rerun()
        return

    email = st.sidebar.text_input("E-mail", value=st.session_state.get("user_email", ""), placeholder="nev@klub.hu", key="license_email")
    if email:
        st.session_state["user_email"] = email

    license_key = st.sidebar.text_input("Aktiváló kód", type="password", help="A klubhoz kapott aktiváló kód.", key="license_key")
    if st.sidebar.button("Pro aktiválása", use_container_width=True, key="activate_license"):
        result = validate_license_supabase(email, license_key)
        if result.get("ok"):
            st.session_state["license_status"] = result
            st.sidebar.success("Pro hozzáférés aktiválva.")
            st.rerun()
        else:
            st.sidebar.warning(result.get("message", "Sikertelen aktiválás."))

    render_mode_badge()


def pro_locked_box(feature: str) -> None:
    st.markdown(
        f"""
        <div class="export-panel">
            <h3 style="margin-top:0;">🔒 {html.escape(feature)}</h3>
            <p style="color:#e0f2fe;">
                Ez a funkció a Pro verzió része. A demo célja, hogy saját adaton is lásd az értéket,
                de a teljes riport/export és hosszabb trendek Pro módban érhetők el.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def apply_demo_limits(df: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, object]]:
    """Demo verzióban saját feltöltés limitálása: 8 játékos, 3 hét, 5000 sor."""
    info = {"limited": False, "original_rows": len(df), "original_players": 0, "original_weeks": 0}
    if df.empty or is_pro_mode():
        return df, info

    out = df.copy()
    players = sorted(out["player_name"].dropna().astype(str).unique().tolist()) if "player_name" in out.columns else []
    weeks = sorted(out["week"].dropna().astype(str).unique().tolist()) if "week" in out.columns else []
    info["original_players"] = len(players)
    info["original_weeks"] = len(weeks)

    keep_players = players[:DEMO_PLAYER_LIMIT]
    keep_weeks = weeks[-DEMO_WEEK_LIMIT:]
    if keep_players and "player_name" in out.columns:
        out = out[out["player_name"].isin(keep_players)]
    if keep_weeks and "week" in out.columns:
        out = out[out["week"].isin(keep_weeks)]
    if len(out) > DEMO_ROW_LIMIT:
        out = out.head(DEMO_ROW_LIMIT)

    info.update({
        "limited": len(out) != len(df) or len(players) > DEMO_PLAYER_LIMIT or len(weeks) > DEMO_WEEK_LIMIT,
        "visible_rows": len(out),
        "visible_players": out["player_name"].nunique() if "player_name" in out.columns else 0,
        "visible_weeks": out["week"].nunique() if "week" in out.columns else 0,
    })
    return out, info


def render_demo_limit_notice(info: Dict[str, object]) -> None:
    if not info or not info.get("limited"):
        return
    st.warning(
        "Demo limit aktív: "
        f"játékosok {info.get('original_players', 0)} → {info.get('visible_players', 0)}, "
        f"hetek {info.get('original_weeks', 0)} → {info.get('visible_weeks', 0)}, "
        f"sorok {info.get('original_rows', 0)} → {info.get('visible_rows', 0)}. "
        "A teljes keret Pro módban érhető el."
    )


@st.cache_data(show_spinner=False)
def build_demo_performance_data() -> pd.DataFrame:
    """Beépített mintaadat: 8 játékos, 3 hét, edzések + meccsek.
    Direkt olyan mintákkal, amelyekből működő readiness, mikrociklus és risk riport készül.
    """
    rng = np.random.default_rng(42)
    players = ["Kovács M.", "Nagy D.", "Szabó B.", "Tóth Á.", "Varga L.", "Farkas Z.", "Balogh P.", "Kiss R."]
    positions = ["CB", "FB", "CM", "AM", "W", "F", "DM", "GK"]
    start = pd.Timestamp("2026-05-18")
    rows = []
    for w in range(3):
        week_start = start + pd.Timedelta(days=7*w)
        # MD-4, MD-3, MD-2, MD-1, MD sessions
        sessions = [
            (week_start + pd.Timedelta(days=1), "Edzés", "MD-4 nagyobb terhelés", 1.08),
            (week_start + pd.Timedelta(days=2), "Edzés", "MD-3 intenzív játék", 1.00),
            (week_start + pd.Timedelta(days=3), "Edzés", "MD-2 taktikai", 0.78 + 0.10*w),
            (week_start + pd.Timedelta(days=4), "Edzés", "MD-1 aktiváció", 0.46 + 0.08*w),
            (week_start + pd.Timedelta(days=5), "Meccs", "Bajnoki mérkőzés", 1.28),
        ]
        for date, typ, sess, factor in sessions:
            for idx, (player, pos) in enumerate(zip(players, positions)):
                role_factor = 1.0 + (idx % 4) * 0.04
                noise = rng.normal(1.0, 0.08)
                base_dist = 5600 if typ == "Edzés" else 9200
                total_distance = max(1200, base_dist * factor * role_factor * noise)
                duration = 75 if typ == "Edzés" else 95
                dpm = total_distance / duration
                sprint_dist = max(20, (180 if typ == "Edzés" else 520) * factor * role_factor * rng.normal(1, .18))
                # direkt kiugrók: 2 játékos magas risk a 3. héten
                if w == 2 and player in ["Nagy D.", "Varga L."]:
                    sprint_dist *= 1.55
                    total_distance *= 1.18
                acc_high = max(0, rng.normal(14, 4) * factor)
                dec_high = max(0, rng.normal(13, 4) * factor)
                rows.append({
                    "Játékos neve": player,
                    "Típus": typ,
                    "Szakasz neve": sess,
                    "Poszt": pos,
                    "Kezdési idő": date + pd.Timedelta(hours=10),
                    "Időtartam": duration,
                    "Teljes táv [m]": round(total_distance, 0),
                    "Táv/perc [m/min]": round(dpm, 1),
                    "Maximális sebesség [km/h]": round(rng.normal(28.5 if typ == "Meccs" else 27.0, 1.2), 1),
                    "Sprintek": int(max(1, sprint_dist / 45)),
                    "Táv a sebesség célzónában 4 [m] (19.80 - 24.99 km/h)": round(sprint_dist * 1.7, 0),
                    "Táv a sebesség célzónában 5 [m] (25.00- km/h)": round(sprint_dist, 0),
                    "Edzési terhelési pontérték": round(total_distance / 80 + sprint_dist / 5 + acc_high * 2 + dec_high * 2, 0),
                    "Kardióterhelés": round(total_distance / 100, 0),
                    "Regenerálódási idő [h]": round(18 + factor * 8 + rng.normal(0, 2), 1),
                    "Izomterhelés": round(45 + factor * 15 + rng.normal(0, 4), 1),
                    "Átlagos pulzus [bpm]": round(136 + factor * 12 + rng.normal(0, 5), 0),
                    "Maximális pulzus [bpm]": round(178 + factor * 6 + rng.normal(0, 4), 0),
                    "HRV (RMSSD)": round(55 - factor * 8 + rng.normal(0, 5), 1),
                    "Gyorsulások száma (2.50 - 2.99 m/s²)": round(acc_high * 1.6, 0),
                    "Gyorsulások száma (3.00 - 50.00 m/s²)": round(acc_high, 0),
                    "Gyorsulások száma (-2.99 - -2.50 m/s²)": round(dec_high * 1.5, 0),
                    "Gyorsulások száma (-50.00 - -3.00 m/s²)": round(dec_high, 0),
                })
    return pd.DataFrame(rows)



# -----------------------------------------------------------------------------
# V5.8 - Product report pack: vezetői / erőnléti / mikrociklus PDF + minta PDF-ek
# -----------------------------------------------------------------------------

def _fpi_to_standard_if_needed(data: pd.DataFrame) -> pd.DataFrame:
    """Bármely támogatott Excelből vagy már standardizált DF-ből riportképes DF-et készít."""
    if data is None or data.empty:
        return pd.DataFrame()
    if {"player_name", "week", "start_time"}.issubset(set(data.columns)):
        out = data.copy()
    else:
        out, _, missing = standardize_dataframe(data.copy())
        if missing:
            return pd.DataFrame()
    try:
        out = add_position_group(out)
    except Exception:
        pass
    return out


def _fpi_latest_week(df: pd.DataFrame, selected_week: Optional[str] = None) -> Optional[str]:
    if df is None or df.empty or "week" not in df.columns:
        return selected_week
    weeks = sorted(df["week"].dropna().astype(str).unique().tolist())
    if selected_week in weeks:
        return selected_week
    return weeks[-1] if weeks else selected_week


def _fpi_report_context(data: pd.DataFrame, selected_week: Optional[str] = None, playstyle: str = "Kiegyensúlyozott") -> Dict[str, object]:
    """Riportok közös számítási magja. Ugyanezt használja a minta PDF és az éles PDF."""
    df = _fpi_to_standard_if_needed(data)
    week = _fpi_latest_week(df, selected_week)
    if df.empty or not week:
        return {"df": df, "selected_week": week, "error": "Nincs riportképes adat."}
    try:
        base = team_insights(df, week)
        micro = microcycle_insights(df, week)
        style = playstyle_insights(df, week, playstyle)
        pattern = build_pattern_insights(df, week)
        readiness_score, readiness_components, readiness_reasons = calculate_readiness_score(df, week, playstyle)
        fingerprints = build_weekly_fingerprints(df)
        current_fp = fingerprints[fingerprints["week"].astype(str) == str(week)] if not fingerprints.empty and "week" in fingerprints.columns else pd.DataFrame()
        periodization_type = current_fp["periodizacios_tipus"].iloc[0] if not current_fp.empty and "periodizacios_tipus" in current_fp.columns else "Nincs elég adat"
        insights = sorted(base + micro + style + pattern, key=lambda x: SEVERITY_RANK.get(x.severity, 9))[:18]
        summary = build_weekly_summary(insights, week, playstyle)
        insights = humanize_insights(insights)
        priorities = humanize_priority_list(build_adaptive_recommendations(insights, readiness_score, periodization_type, pattern, playstyle))
        summary = coach_friendly_phrase(summary)
        risk = calculate_player_risk(df, week)
        past_df, past_text = build_past_week_review(df, week, playstyle)
        current_df, current_text = build_current_remaining_days_plan(df, week, playstyle, readiness_score, periodization_type, risk)
        next_df, next_text = build_next_week_plan_v5(df, week, playstyle, readiness_score, periodization_type, risk, past_df, current_df)
        player_actions = build_player_next_actions(risk, df, week)
        weekly = aggregate_weekly(df) if "week" in df.columns else pd.DataFrame()
        player_week = player_weekly(df) if "week" in df.columns else pd.DataFrame()
        return {
            "df": df,
            "selected_week": week,
            "playstyle": playstyle,
            "readiness_score": readiness_score,
            "readiness_components": readiness_components,
            "readiness_reasons": readiness_reasons,
            "periodization_type": periodization_type,
            "insights": insights,
            "insights_df": build_insight_export_df(insights),
            "priorities": priorities,
            "summary": summary + f"\n\nMeccskészültség: {readiness_score}/100 ({score_to_label(readiness_score)})\nPeriodizációs besorolás: {periodization_type}",
            "risk_df": risk,
            "past_df": past_df,
            "past_text": past_text,
            "current_df": current_df,
            "current_text": current_text,
            "next_df": next_df,
            "next_text": next_text,
            "player_actions_df": player_actions,
            "weekly": weekly,
            "player_week": player_week,
            "fingerprints": fingerprints,
        }
    except Exception as exc:
        return {"df": df, "selected_week": week, "error": str(exc)}



def _fpi_pdf_tactical_source_status_v80(tactical_context: Optional[Dict[str, object]]) -> str:
    if not tactical_context:
        return "Nincs taktikai input – GPS-only riport."
    own_pdf = bool(tactical_context.get("has_own_pdf"))
    opp_pdf = bool(tactical_context.get("has_opp_pdf"))
    own_topics = tactical_context.get("own_topics", []) or []
    opp_topics = tactical_context.get("opp_topics", []) or []
    excel_ok = bool(tactical_context.get("has_own_team_excel") or tactical_context.get("has_opp_team_excel") or tactical_context.get("has_own_player_excel") or tactical_context.get("has_opp_player_excel"))
    if (own_topics or opp_topics) and excel_ok:
        return "PDF + Excel együtt – a következtetésekben mindkettő szerepel."
    if own_topics or opp_topics:
        return "PDF-alapú taktikai jelzések – Excel nélkül vagy korlátozott Excel-adattal."
    if (own_pdf or opp_pdf) and excel_ok:
        return "PDF feltöltve, de nem adott erős tématalálatot; a taktikai következtetések főleg az Excelből készültek."
    if excel_ok:
        return "Taktikai Excel-alapú következtetések."
    if own_pdf or opp_pdf:
        return "PDF feltöltve, de nincs kinyert / felismert taktikai téma."
    return "Nincs használható taktikai input."

def _fpi_pdf_short_tactical_topics_v80(tactical_context: Optional[Dict[str, object]], limit: int = 3) -> str:
    if not tactical_context:
        return "n.a."
    rows = []
    for side, key in [("Saját", "own_topics"), ("Ellenfél", "opp_topics")]:
        for r in (tactical_context.get(key, []) or [])[:limit]:
            tema = str(r.get("Téma", "")).strip()
            if tema:
                rows.append(f"{side}: {tema}")
    return "; ".join(rows[:limit]) if rows else "Nincs erős PDF-tématalálat."

def _fpi_pdf_primary_tactical_finding_v80(tactical_context: Optional[Dict[str, object]]) -> str:
    if not tactical_context:
        return "Nincs taktikai input."
    findings = tactical_context.get("tactical_findings", []) or []
    if findings:
        f = findings[0]
        return f"{f.get('Téma', '')}: {f.get('Edzői következtetés', '')}"
    risks = tactical_context.get("risks", []) or []
    if risks:
        return str(risks[0])
    return _fpi_pdf_tactical_source_status_v80(tactical_context)

def _fpi_pdf_next_integrated_decision_v80(tactical_context: Optional[Dict[str, object]], fallback: str) -> str:
    if tactical_context and tactical_context.get("md_plan"):
        try:
            md = tactical_context.get("md_plan")[0]
            return f"Következő mikrociklus-lépés: {md[0]} – {md[1]}. Indoklás: {md[2]}"
        except Exception:
            pass
    if tactical_context and tactical_context.get("plan_a"):
        return f"Meccsterv irány: {tactical_context.get('plan_a')}. Következő edzésen ehhez igazított erőnléti + taktikai cél."
    return fallback



def _fpi_compact_text_v81(x: object, max_len: int = 260) -> str:
    s = re.sub(r"\s+", " ", str(x or "")).strip()
    return s if len(s) <= max_len else s[:max_len-1].rstrip() + "…"

def _fpi_pdf_scope_line_v81(tactical_context: Optional[Dict[str, object]]) -> str:
    if not tactical_context:
        return "PDF: nincs feltöltött taktikai PDF."
    own_chars = int(tactical_context.get("own_pdf_chars", 0) or 0)
    opp_chars = int(tactical_context.get("opp_pdf_chars", 0) or 0)
    own_pages = int(tactical_context.get("own_pdf_pages", 0) or 0)
    opp_pages = int(tactical_context.get("opp_pdf_pages", 0) or 0)
    topics = len(tactical_context.get("own_topics", []) or []) + len(tactical_context.get("opp_topics", []) or [])
    if own_chars + opp_chars > 0:
        return f"PDF: {own_pages}+{opp_pages} oldalból {own_chars + opp_chars} karakter feldolgozva; felismert témák: {topics}. A riport a legerősebb PDF-témákat használja, a többi következtetés Excelből jön."
    if tactical_context.get("has_own_pdf") or tactical_context.get("has_opp_pdf"):
        return f"PDF: feltöltve ({own_pages}+{opp_pages} oldal), de nincs kinyert szöveg vagy erős tématalálat; taktikai döntés főleg Excelből."
    return "PDF: nincs használható PDF-jelzés; taktikai döntés Excelből."



def _fpi_clean_sentence_v82(x: object, max_len: int = 160) -> str:
    s = re.sub(r"\s+", " ", str(x or "")).strip()
    s = s.replace("->", "→")
    return s if len(s) <= max_len else s[:max_len-1].rstrip() + "…"

def _fpi_readiness_short_v82(score: int) -> str:
    try:
        score = int(score)
    except Exception:
        score = 70
    if score >= 80:
        return "jó állapot, vállalható intenzitás"
    if score >= 65:
        return "vállalható, de frissességre figyelni kell"
    if score >= 50:
        return "figyelendő, terheléskontroll szükséges"
    return "magas kockázat, óvatos mikrociklus"



def _fpi_norm_risk_level_v126(x: object) -> str:
    s = unicodedata.normalize("NFKD", str(x or "").strip().lower())
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    if "magas" in s or "high" in s:
        return "Magas"
    if "kozep" in s or "medium" in s:
        return "Közepes"
    if "alacsony" in s or "low" in s:
        return "Alacsony"
    return str(x or "").strip() or "Alacsony"


def _fpi_count_risk_levels_v126(risk_df: pd.DataFrame) -> Tuple[int, int]:
    """Játékosszintű risk darabszám.

    V12.7 javítás: nem sorokat számolunk, hanem egyedi játékosokat.
    Ha egy játékos több sorban szerepel, a legmagasabb kockázati szintje számít.
    Így nem fordulhat elő, hogy a PDF 5 közepes risket ír, miközben valójában
    csak 1 játékos közepes kockázatú.
    """
    if risk_df is None or risk_df.empty or "Kockázati szint" not in risk_df.columns:
        return 0, 0

    tmp = risk_df.copy()
    tmp["_risk_norm"] = tmp["Kockázati szint"].apply(_fpi_norm_risk_level_v126)
    tmp["_risk_rank"] = tmp["_risk_norm"].map({"Magas": 3, "Közepes": 2, "Alacsony": 1}).fillna(1).astype(int)

    # Lehetséges játékosnév-oszlopok több exportverzióból.
    player_col = None
    for c in ["Játékos", "Játékos neve", "player_name", "Player", "Név", "Nev", "Name"]:
        if c in tmp.columns:
            player_col = c
            break

    if player_col:
        tmp["_player_key"] = tmp[player_col].astype(str).str.strip().replace({"": np.nan, "nan": np.nan, "None": np.nan})
        known = tmp[tmp["_player_key"].notna()].copy()
        unknown = tmp[tmp["_player_key"].isna()].copy()
        if not known.empty:
            idx = known.groupby("_player_key")["_risk_rank"].idxmax()
            tmp = pd.concat([known.loc[idx], unknown], ignore_index=True)

    high = int((tmp["_risk_norm"] == "Magas").sum())
    med = int((tmp["_risk_norm"] == "Közepes").sum())
    return high, med


def _fpi_pdf_week_type_label_v126(ctx: Dict[str, object], demo_label: str = "") -> str:
    raw = str(ctx.get("periodization_type", "") or "").strip()
    raw = _fpi_normalize_coach_week_label_v121(raw) if "_fpi_normalize_coach_week_label_v121" in globals() else raw
    bad = {"", "n.a.", "na", "none", "nincs elég adat", "nincs eleg adat", "nan"}
    if raw.lower() in bad:
        return "Mérkőzésre felkészítő mikrociklus" if demo_label else "Stabilizáló mikrociklus"
    return raw


def _fpi_week_type_interpretation_v126(label: str) -> str:
    l = str(label or "").lower()
    if "regener" in l:
        return "A hét elsődleges célja a frissesség visszaépítése és a terhelési kockázat csökkentése."
    if "stabil" in l:
        return "A hét célja a terhelési szint megtartása, kontrollált sebességi és intenzitási ingerekkel."
    if "terhelésfokoz" in l or "terhelesfokoz" in l:
        return "A hét célja a fokozatos terhelésépítés, külön kontrollal a sprint és HSR ingerekre."
    if "fejleszt" in l:
        return "A hét célja célzott fizikai fejlesztés, de a readiness és játékoskockázat folyamatos kontrollja mellett."
    if "formaid" in l:
        return "A hét célja a meccs előtti frissítés: alacsonyabb volumen, rövid és minőségi intenzív ingerek."
    if "mérkőzés" in l or "merkozes" in l:
        return "A hét célja a mérkőzésre optimalizált mikrociklus: terhelésépítés, sebességi inger, majd frissítés."
    return "A hét típusa a terhelési trend, readiness és játékosszintű kockázatok alapján értelmezendő."


def _fpi_short_counts_sentence_v126(high_risk: int, med_risk: int) -> str:
    if high_risk or med_risk:
        return f"Játékoskockázat: {high_risk} magas, {med_risk} közepes."
    return "Játékoskockázat: nincs kiemelt magas vagy közepes jelzés."

def _fpi_top_tactical_messages_v82(tactical_context: Optional[Dict[str, object]], limit: int = 3, gps_context: Optional[Dict[str, object]] = None, readiness: Optional[int] = None, priorities: Optional[List[dict]] = None, week: Optional[str] = None) -> List[str]:
    if not _fpi_has_tactical_signal_v95(tactical_context):
        if gps_context is not None:
            return _fpi_gps_only_conclusions_v95(gps_context, priorities or [], int(readiness or 70), str(week or gps_context.get("selected_week", "")), limit=max(limit, 5))
        return ["GPS-only mód: nincs taktikai input, ezért a javaslat az erőnléti/GPS adatokra épül."]
    out = []
    findings = tactical_context.get("tactical_findings") or []
    pdf_first = [f for f in findings if str(f.get("Téma", "")).lower().startswith("pdf") or "PDF" in str(f.get("Forrás", ""))]
    other = [f for f in findings if f not in pdf_first]
    for f in (pdf_first + other)[:limit]:
        theme = str(f.get("Téma", "")).strip()
        decision = str(f.get("Edzői következtetés", "")).strip()
        if theme and decision:
            out.append(f"{theme}: {decision}")
    if not out:
        risks = tactical_context.get("risks", []) or []
        out = [str(x) for x in risks[:limit]]
    if not out:
        topics = []
        for key in ["opp_topics", "own_topics"]:
            for row in tactical_context.get(key, []) or []:
                if row.get("Téma"):
                    topics.append(str(row.get("Téma")))
        if topics:
            out.append("PDF témák: " + ", ".join(topics[:3]))
    return [_fpi_clean_sentence_v82(x, 180) for x in out[:limit]] or ["Nincs erős taktikai jelzés."]

def _fpi_top_fitness_messages_v82(ctx: Dict[str, object], priorities: List[dict], readiness: int, limit: int = 3) -> List[str]:
    out = []
    summary = str(ctx.get("summary", "") or "")
    # prefer short existing messages from summary
    for line in summary.splitlines():
        line = line.strip("- ").strip()
        if any(k in line.lower() for k in ["legfontosabb", "md-1", "stabil", "readiness", "terhelés"]):
            out.append(line)
    for p in priorities or []:
        t = p.get("Teendő", p.get("Cím", ""))
        if t:
            out.append(t)
    if not out:
        out.append(f"Readiness {readiness}/100 – {_fpi_readiness_short_v82(readiness)}.")
    return [_fpi_clean_sentence_v82(x, 170) for x in list(dict.fromkeys(out))[:limit]]

def _fpi_compact_player_risk_rows_v82(risk_df: pd.DataFrame, max_rows: int = 12) -> List[List[str]]:
    rows = [["Játékos", "Szint", "Miért fontos?"]]
    if risk_df is None or risk_df.empty:
        rows.append(["Nincs kiemelt", "Alacsony", "Nincs azonnali beavatkozási jelzés."])
        return rows
    df = risk_df.copy()
    level_col = "Kockázati szint" if "Kockázati szint" in df.columns else None
    reason_col = "Fő okok" if "Fő okok" in df.columns else "Fő ok" if "Fő ok" in df.columns else None
    player_col = "Játékos" if "Játékos" in df.columns else "player_name" if "player_name" in df.columns else df.columns[0]
    if level_col:
        df["_risk_norm_v126"] = df[level_col].apply(_fpi_norm_risk_level_v126)
        focus = df[df["_risk_norm_v126"].isin(["Magas", "Közepes"])]
        if not focus.empty:
            order = {"Magas": 0, "Közepes": 1, "Alacsony": 2}
            focus = focus.assign(_risk_order=focus["_risk_norm_v126"].map(order).fillna(9))
            df = focus.sort_values(["_risk_order"])
    for _, r in df.head(max_rows).iterrows():
        level = _fpi_norm_risk_level_v126(r.get(level_col, "Alacsony") if level_col else "Figyelendő")
        rows.append([
            _fpi_clean_sentence_v82(r.get(player_col, ""), 42),
            _fpi_clean_sentence_v82(level, 30),
            _fpi_clean_sentence_v82(r.get(reason_col, "Monitoring.") if reason_col else "Monitoring.", 105),
        ])
    return rows


def _fpi_md_plan_rows_v82(tactical_context: Optional[Dict[str, object]], gps_context: Optional[Dict[str, object]] = None, readiness: Optional[int] = None, priorities: Optional[List[dict]] = None, week: Optional[str] = None) -> List[Tuple[str, str, str]]:
    # output: nap, erőnléti cél, taktikai/GPS cél
    if not _fpi_has_tactical_signal_v95(tactical_context) and gps_context is not None:
        return _fpi_gps_only_md_plan_v95(gps_context, int(readiness or 70), priorities or [], str(week or gps_context.get("selected_week", "")))
    if tactical_context:
        md = tactical_context.get("md_plan", []) or []
        out = []
        for item in md:
            try:
                day, tactical_focus, why = item
            except Exception:
                continue
            fgoal = "Terheléskontroll"
            tgoal = str(tactical_focus)
            low = (str(tactical_focus) + " " + str(why)).lower()
            if "regener" in low:
                fgoal = "Regeneráció, frissítés"
            elif "hsr" in low or "sprint" in low or "átmenet" in low:
                fgoal = "HSR/sprint inger kontrolláltan"
            elif "aktiv" in low:
                fgoal = "Aktiváció, frissesség"
            elif "volumen" in low:
                fgoal = "Fő terhelési nap"
            elif "terheléskontroll" in low:
                fgoal = "Terhelés csökkentése"
            out.append((str(day), fgoal, _fpi_clean_sentence_v82(tgoal, 95)))
        # vezetői oldalon MD-4-től MD-1-ig a lényeg
        filtered = [x for x in out if any(k in x[0] for k in ["MD-4", "MD-3", "MD-2", "MD-1"])]
        return filtered[:4] if filtered else out[:4]
    return [
        ("MD-4", "Fő terhelési nap", "Saját játékmodell"),
        ("MD-3", "HSR/sprint inger", "Átmenetek"),
        ("MD-2", "Terhelés csökkentése", "Meccsterv"),
        ("MD-1", "Aktiváció", "Pontrúgás + frissítés"),
    ]

def _fpi_plan_why_v82(tactical_context: Optional[Dict[str, object]], readiness: int, gps_context: Optional[Dict[str, object]] = None, priorities: Optional[List[dict]] = None, week: Optional[str] = None) -> str:
    if not _fpi_has_tactical_signal_v95(tactical_context):
        if gps_context is not None:
            msgs = _fpi_gps_only_conclusions_v95(gps_context, priorities or [], readiness, str(week or gps_context.get("selected_week", "")), limit=2)
            return _fpi_clean_sentence_v82("; ".join(msgs), 260)
        return f"GPS alapján: readiness {readiness}/100, taktikai input nélkül."
    msgs = _fpi_top_tactical_messages_v82(tactical_context, 2)
    return _fpi_clean_sentence_v82("; ".join(msgs), 260)




def _fpi_uploaded_files_to_bytes_v88(files: List[object]) -> List[dict]:
    """Stabil UploadFile -> bytes mentés.
    Streamlit download/re-run közben a file_uploader objektum nem mindig olvasható újra megbízhatóan,
    ezért azonnal byte-listát mentünk session_state-be.
    """
    out = []
    for f in files or []:
        try:
            b = f.getvalue()
        except Exception:
            try:
                b = f.read()
            except Exception:
                b = b""
        if b:
            out.append({
                "name": getattr(f, "name", "uploaded.pdf"),
                "size": len(b),
                "md5": hashlib.md5(b).hexdigest(),
                "bytes": b,
            })
    return out

def _fpi_restore_uploaded_file_wrappers_v88(items: List[dict]) -> List[object]:
    class _FPIBytesPDF:
        def __init__(self, name, data):
            self.name = name
            self._data = data
        def getvalue(self):
            return self._data
        def read(self):
            return self._data
    return [_FPIBytesPDF(x.get("name", "uploaded.pdf"), x.get("bytes", b"")) for x in (items or []) if x.get("bytes")]

def _fpi_tactical_app_extract_pdf_pages_v88(file_bytes: bytes, max_pages: int = 80) -> List[dict]:
    """Régi Tactical app út + PyMuPDF fallback.
    Elsődleges: pdfplumber.open(io.BytesIO(bytes)).
    Fallback: PyMuPDF, ha pdfplumber nincs vagy hibázik.
    """
    out = []
    if not file_bytes:
        return out

    # 1) pdfplumber / pdfminer út
    try:
        if pdfplumber is not None:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                total_pages = min(len(pdf.pages), max_pages)
                for p in range(total_pages):
                    txt = pdf.pages[p].extract_text(x_tolerance=1, y_tolerance=3) or ""
                    if txt.strip():
                        out.append({
                            "page_index": p,
                            "page_number": p + 1,
                            "reader": "legacy_tactical_pdfplumber_v90",
                            "chars": len(txt),
                            "has_text": True,
                            "text": txt,
                        })
        if out:
            return out
    except Exception as e:
        out.append({
            "page_index": None,
            "page_number": None,
            "reader": "legacy_tactical_pdfplumber_v90",
            "chars": 0,
            "has_text": False,
            "error": str(e),
            "text": "",
        })

    # 2) PyMuPDF fallback
    try:
        if PYMUPDF_AVAILABLE and fitz is not None:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for p in range(min(len(doc), max_pages)):
                txt = doc[p].get_text("text") or ""
                if txt.strip():
                    out.append({
                        "page_index": p,
                        "page_number": p + 1,
                        "reader": "pymupdf_fallback_v90",
                        "chars": len(txt),
                        "has_text": True,
                        "text": txt,
                    })
            doc.close()
            if out:
                return out
    except Exception as e:
        out.append({
            "page_index": None,
            "page_number": None,
            "reader": "pymupdf_fallback_v90",
            "chars": 0,
            "has_text": False,
            "error": str(e),
            "text": "",
        })

    return out


def _fpi_tactical_app_combine_pdf_texts_v88(items: List[dict]) -> Tuple[str, List[dict]]:
    page_blocks = []
    texts = []
    for item in items or []:
        b = item.get("bytes", b"") if isinstance(item, dict) else b""
        name = item.get("name", "uploaded.pdf") if isinstance(item, dict) else "uploaded.pdf"
        pages = _fpi_tactical_app_extract_pdf_pages_v88(b)
        for p in pages:
            p["file"] = name
        page_blocks.extend(pages)
        texts.extend([x["text"] for x in pages if x.get("text", "").strip()])
    return "\n\n".join(texts), page_blocks




def _fpi_tactical_app_combine_uploadfiles_v89(files: List[object]) -> Tuple[str, List[dict], List[dict]]:
    """Aktuális file_uploader objektumból azonnal olvas. Ez a régi Tactical app útja."""
    items = _fpi_uploaded_files_to_bytes_v88(files or [])
    text, pages = _fpi_tactical_app_combine_pdf_texts_v88(items)
    return text, pages, items

def _fpi_store_current_pdf_text_v89(side: str, text: str, pages: List[dict], items: List[dict]) -> None:
    st.session_state[f"tactical_pro_{side}_pdf_text_store_v89"] = {
        "text": text or "",
        "pages": pages or [],
        "items": items or [],
        "chars": len(text or ""),
        "page_count": len([p for p in (pages or []) if p.get("has_text") or p.get("text")]),
        "files": [x.get("name", "") for x in (items or [])],
        "signature": "|".join([f"{x.get('name')}:{x.get('size')}:{x.get('md5')}" for x in (items or [])]),
        "reader": "direct_upload_legacy_tactical_reader_v89",
    }

def _fpi_get_pdf_text_store_v89(side: str) -> Dict[str, object]:
    return st.session_state.get(f"tactical_pro_{side}_pdf_text_store_v89") or {}

def _fpi_clear_pdf_side_state_v91(side: str) -> None:
    """Ha a file_uploader üres, ne maradjon bent régi PDF-context.
    Ez javítja azt, hogy a riport 'PDF feltöltve'-t írt, miközben a widgetben len=0.
    """
    for key in [
        f"tactical_pro_{side}_pdf_text_store_v89",
        f"tactical_pro_{side}_pdf_bytes_v88",
    ]:
        st.session_state.pop(key, None)

def _fpi_has_current_pdf_upload_v91(files: List[object]) -> bool:
    return bool(files) and len(files or []) > 0

def _fpi_pdf_upload_state_key_v92(side: str) -> str:
    return f"tactical_pro_{side}_pdf_upload_manager_v92"

def _fpi_set_pdf_upload_state_v92(side: str, files: List[object]) -> Dict[str, object]:
    """Stabil PDF upload manager.
    A file_uploader visszatérési értékét AZONNAL bytes + text formában tárolja.
    Innentől a Tactical Pro+ és az export ugyanebből dolgozik.
    """
    state = {
        "has_files": False,
        "files": [],
        "items": [],
        "text": "",
        "pages": [],
        "chars": 0,
        "page_count": 0,
        "reader": "not_run",
        "error": "",
    }
    if not files:
        st.session_state[_fpi_pdf_upload_state_key_v92(side)] = state
        _fpi_clear_pdf_side_state_v91(side)
        return state

    items = _fpi_uploaded_files_to_bytes_v88(files or [])
    state["has_files"] = bool(items)
    state["items"] = items
    state["files"] = [x.get("name", "") for x in items]

    if items:
        try:
            txt, pages = _fpi_tactical_app_combine_pdf_texts_v88(items)
            state["text"] = txt or ""
            state["pages"] = pages or []
            state["chars"] = len(txt or "")
            state["page_count"] = len([p for p in (pages or []) if p.get("has_text") or p.get("text")])
            state["reader"] = "legacy_tactical_reader_v92"
        except Exception as e:
            state["error"] = f"legacy reader error: {e}"

        if not state["text"]:
            try:
                wrappers = _fpi_restore_uploaded_file_wrappers_v88(items)
                txt2, pages2 = _fpi_tactical_extract_pdf_text(wrappers)
                state["text"] = txt2 or ""
                state["pages"] = pages2 or []
                state["chars"] = len(txt2 or "")
                state["page_count"] = len([p for p in (pages2 or []) if p.get("has_text") or p.get("text")])
                state["reader"] = "multi_reader_fallback_v92"
            except Exception as e:
                state["error"] = (state.get("error", "") + f" | fallback error: {e}").strip(" |")

    st.session_state[_fpi_pdf_upload_state_key_v92(side)] = state

    # Kompatibilitás a korábbi v88/v89 kulcsokkal
    st.session_state[f"tactical_pro_{side}_pdf_bytes_v88"] = items
    _fpi_store_current_pdf_text_v89(side, state["text"], state["pages"], items)
    return state

def _fpi_get_pdf_upload_state_v92(side: str) -> Dict[str, object]:
    return st.session_state.get(_fpi_pdf_upload_state_key_v92(side)) or {
        "has_files": False, "files": [], "items": [], "text": "", "pages": [],
        "chars": 0, "page_count": 0, "reader": "missing", "error": ""
    }

def _fpi_pdf_uploader_v92(label: str, side: str, key: str):
    """Izolált uploader: a widget key és a feldolgozó state külön van."""
    files = st.file_uploader(label, type=["pdf"], accept_multiple_files=True, key=key)
    state = _fpi_set_pdf_upload_state_v92(side, files or [])
    if state.get("has_files"):
        if state.get("chars", 0) > 0:
            st.success(f"PDF OK: {len(state.get('files', []))} fájl, {state.get('page_count', 0)} oldal, {state.get('chars', 0)} karakter.")
        else:
            st.error(f"PDF feltöltve, de nincs kinyert szöveg. Reader: {state.get('reader')}. Hiba: {state.get('error') or 'n.a.'}")
    else:
        st.info("Nincs PDF feltöltve ebben a mezőben.")
    return files, state






def _fpi_uploaded_file_signature_v87(files: List[object]) -> str:
    parts = []
    for f in files or []:
        try:
            b = f.getvalue()
            parts.append(f"{getattr(f, 'name', 'pdf')}:{len(b)}:{hashlib.md5(b[:200000]).hexdigest()}")
        except Exception:
            parts.append(str(getattr(f, "name", "unknown")))
    return "|".join(parts)

def _fpi_build_pdf_only_context_from_session_v87(gps_context: Dict[str, object]) -> Optional[Dict[str, object]]:
    """V8.9: export előtti PDF context rebuild a már kinyert PDF-SZÖVEGBŐL.
    Először text-store, utána byte-store, utána UploadFile fallback.
    """
    try:
        own_state_v92 = _fpi_get_pdf_upload_state_v92("own")
        opp_state_v92 = _fpi_get_pdf_upload_state_v92("opp")
        own_store = _fpi_get_pdf_text_store_v89("own")
        opp_store = _fpi_get_pdf_text_store_v89("opp")

        own_pdf_text = own_state_v92.get("text", "") or own_store.get("text", "") or ""
        opp_pdf_text = opp_state_v92.get("text", "") or opp_store.get("text", "") or ""
        own_pdf_pages = own_state_v92.get("pages", []) or own_store.get("pages", []) or []
        opp_pdf_pages = opp_state_v92.get("pages", []) or opp_store.get("pages", []) or []
        own_items = own_state_v92.get("items", []) or own_store.get("items", []) or []
        opp_items = opp_state_v92.get("items", []) or opp_store.get("items", []) or []

        if not own_pdf_text:
            own_items = st.session_state.get("tactical_pro_own_pdf_bytes_v88") or _fpi_uploaded_files_to_bytes_v88(st.session_state.get("tactical_pro_own_pdfs") or [])
            own_pdf_text, own_pdf_pages = _fpi_tactical_app_combine_pdf_texts_v88(own_items)
        if not opp_pdf_text:
            opp_items = st.session_state.get("tactical_pro_opp_pdf_bytes_v88") or _fpi_uploaded_files_to_bytes_v88(st.session_state.get("tactical_pro_opp_pdfs") or [])
            opp_pdf_text, opp_pdf_pages = _fpi_tactical_app_combine_pdf_texts_v88(opp_items)

        if not own_pdf_text and own_items:
            own_pdf_text, own_pdf_pages = _fpi_tactical_extract_pdf_text(_fpi_restore_uploaded_file_wrappers_v88(own_items))
        if not opp_pdf_text and opp_items:
            opp_pdf_text, opp_pdf_pages = _fpi_tactical_extract_pdf_text(_fpi_restore_uploaded_file_wrappers_v88(opp_items))
    except Exception:
        return None

    has_any_pdf = bool(own_items or opp_items or own_pdf_text or opp_pdf_text)
    if not has_any_pdf:
        return None

    own_pdf_insights = _fpi_tactical_pdf_insights(own_pdf_text) if own_pdf_text else {"blocks": {}, "topics": [], "raw_text_len": 0, "sportsbase_findings": [], "sportsbase_lines": []}
    opp_pdf_insights = _fpi_tactical_pdf_insights(opp_pdf_text) if opp_pdf_text else {"blocks": {}, "topics": [], "raw_text_len": 0, "sportsbase_findings": [], "sportsbase_lines": []}

    own_real_pages = len({(p.get("file"), p.get("page_number", p.get("page"))) for p in own_pdf_pages if p.get("has_text") or p.get("text")})
    opp_real_pages = len({(p.get("file"), p.get("page_number", p.get("page"))) for p in opp_pdf_pages if p.get("has_text") or p.get("text")})

    own_pdf_insights["pdf_uploaded"] = bool(own_items or own_pdf_text)
    own_pdf_insights["pdf_pages"] = own_real_pages
    own_pdf_insights["raw_text_len"] = len(own_pdf_text or "")
    own_pdf_insights["upload_signature"] = own_store.get("signature", "") or "|".join([f"{x.get('name')}:{x.get('size')}:{x.get('md5')}" for x in own_items])
    own_pdf_insights["reader_version"] = "V8.9 direct UploadFile text-store + legacy Tactical reader"
    own_pdf_insights["page_debug"] = own_pdf_pages[:10]

    opp_pdf_insights["pdf_uploaded"] = bool(opp_items or opp_pdf_text)
    opp_pdf_insights["pdf_pages"] = opp_real_pages
    opp_pdf_insights["raw_text_len"] = len(opp_pdf_text or "")
    opp_pdf_insights["upload_signature"] = opp_store.get("signature", "") or "|".join([f"{x.get('name')}:{x.get('size')}:{x.get('md5')}" for x in opp_items])
    opp_pdf_insights["reader_version"] = "V8.9 direct UploadFile text-store + legacy Tactical reader"
    opp_pdf_insights["page_debug"] = opp_pdf_pages[:10]

    merged_pdf_insights = _fpi_safe_merge_tactical_pdf_insights_v104(own_pdf_insights, opp_pdf_insights)

    previous = st.session_state.get("tactical_pro_context") or {}
    own_team_metrics = previous.get("own_team_metrics", {}) or {}
    opp_team_metrics = previous.get("opp_team_metrics", {}) or {}
    own_player_tables = previous.get("own_player_tables", {}) or {}
    opp_player_tables = previous.get("opp_player_tables", {}) or {}

    has_excel = bool(own_team_metrics or opp_team_metrics or previous.get("has_own_player_excel") or previous.get("has_opp_player_excel"))
    level_label = "Full Intelligence – GPS + frissített taktikai PDF + meglévő Excel context" if has_excel else "GPS + frissített taktikai PDF"

    tactical_ctx_for_plan = {
        "analysis_level_label": level_label,
        "pdf_insights": merged_pdf_insights,
        "team_metrics": opp_team_metrics,
        "player_tables": opp_player_tables,
        "own": {"pdf_insights": own_pdf_insights, "team_metrics": own_team_metrics, "player_tables": own_player_tables},
        "opponent": {"pdf_insights": opp_pdf_insights, "team_metrics": opp_team_metrics, "player_tables": opp_player_tables},
    }
    plan = _fpi_safe_build_adaptive_plan_v104(gps_context or {}, tactical_ctx_for_plan)
    ctx = _build_tactical_executive_context(gps_context or {}, tactical_ctx_for_plan, plan)
    ctx["forced_pdf_rebuild"] = True
    ctx["forced_pdf_signature"] = (own_pdf_insights.get("upload_signature", "") + "||" + opp_pdf_insights.get("upload_signature", ""))
    ctx["own_pdf_chars"] = len(own_pdf_text or "")
    ctx["opp_pdf_chars"] = len(opp_pdf_text or "")
    ctx["own_pdf_pages"] = own_real_pages
    ctx["opp_pdf_pages"] = opp_real_pages
    ctx["pdf_reader_debug"] = {
        "own_files": own_store.get("files", []) or [x.get("name") for x in own_items],
        "opp_files": opp_store.get("files", []) or [x.get("name") for x in opp_items],
        "own_chars": len(own_pdf_text or ""),
        "opp_chars": len(opp_pdf_text or ""),
        "own_pages": own_real_pages,
        "opp_pages": opp_real_pages,
        "source": "v89_text_store_first",
    }
    return ctx


def _fpi_context_for_export_v87(gps_context: Dict[str, object]) -> Optional[Dict[str, object]]:
    """Riportexporthoz mindig próbálunk friss PDF-contextet építeni.
    Ha sikerül és talált PDF-sorokat/következtetéseket, azt használjuk.
    """
    fresh = _fpi_build_pdf_only_context_from_session_v87(gps_context or {})
    old = st.session_state.get("tactical_pro_context") if "st" in globals() else None

    # V9.1: ha nincs aktuális PDF a widgetben és nincs text-store, ne használjunk régi PDF-es contextet.
    try:
        no_pdf_now = (
            not st.session_state.get("tactical_pro_own_pdfs")
            and not st.session_state.get("tactical_pro_opp_pdfs")
            and not (_fpi_get_pdf_text_store_v89("own").get("text"))
            and not (_fpi_get_pdf_text_store_v89("opp").get("text"))
        )
        if no_pdf_now and old:
            old = dict(old)
            old["has_own_pdf"] = False
            old["has_opp_pdf"] = False
            old["own_pdf_pages"] = 0
            old["opp_pdf_pages"] = 0
            old["own_pdf_chars"] = 0
            old["opp_pdf_chars"] = 0
            old["pdf_provider_lines"] = []
            old["pdf_provider_findings"] = []
    except Exception:
        pass

    if fresh:
        fresh_score = int(fresh.get("pdf_direct_findings_count", 0) or 0) + int(fresh.get("pdf_direct_lines_count", 0) or 0)
        old_score = int((old or {}).get("pdf_direct_findings_count", 0) or 0) + int((old or {}).get("pdf_direct_lines_count", 0) or 0)
        fresh_pages = int(fresh.get("own_pdf_pages", 0) or 0) + int(fresh.get("opp_pdf_pages", 0) or 0)
        fresh_chars = int(fresh.get("own_pdf_chars", 0) or 0) + int(fresh.get("opp_pdf_chars", 0) or 0)
        old_pages = int((old or {}).get("own_pdf_pages", 0) or 0) + int((old or {}).get("opp_pdf_pages", 0) or 0)
        old_chars = int((old or {}).get("own_pdf_chars", 0) or 0) + int((old or {}).get("opp_pdf_chars", 0) or 0)
        if fresh_pages > 0 or fresh_chars > 0 or fresh_score > 0 or (old_pages == 0 and old_chars == 0):
            st.session_state["tactical_pro_context"] = fresh
            return fresh
    return old



# =========================================================
# V9.3 - Edzés / meccs százalékos referencia motor
# =========================================================

FPI_NB2_ADULT_REFERENCE_RANGES_V93 = {
    "total_distance": {
        "label": "Volumen / össztáv",
        "unit": "m",
        "weekly_ref": "280–420%",
        "avg_ref": "60–100%",
        "low": 280,
        "high": 420,
        "avg_low": 60,
        "avg_high": 100,
        "explain": "Heti összmunka. A heti edzésösszeg jellemzően több meccsterhelésnyi, az egy edzés átlaga viszont általában 1 meccs alatt marad.",
    },
    "hsr_distance": {
        "label": "HSR",
        "unit": "m",
        "weekly_ref": "150–250%",
        "avg_ref": "35–70%",
        "low": 150,
        "high": 250,
        "avg_low": 35,
        "avg_high": 70,
        "explain": "High Speed Running: nagy sebességű futás, tipikusan kb. 19,8–20 km/h felett.",
    },
    "sprint_distance": {
        "label": "Sprint táv",
        "unit": "m",
        "weekly_ref": "100–200%",
        "avg_ref": "25–55%",
        "low": 100,
        "high": 200,
        "avg_low": 25,
        "avg_high": 55,
        "explain": "Sprintzónában megtett méter, jellemzően kb. 25 km/h felett vagy rendszer-specifikus sprintküszöb felett.",
    },
    "sprints": {
        "label": "Sprint count",
        "unit": "db",
        "weekly_ref": "100–220%",
        "avg_ref": "25–60%",
        "low": 100,
        "high": 220,
        "avg_low": 25,
        "avg_high": 60,
        "explain": "Sprintakciók darabszáma. Nem ugyanaz, mint a sprint táv.",
    },
    "high_efforts": {
        "label": "High Efforts",
        "unit": "db/pont",
        "weekly_ref": "150–280%",
        "avg_ref": "35–75%",
        "low": 150,
        "high": 280,
        "avg_low": 35,
        "avg_high": 75,
        "explain": "Nagy intenzitású akciók összesített mutatója. Rendszertől függően sprint, gyorsulás, lassítás vagy robbanékony effort is lehet benne.",
    },
    "training_load": {
        "label": "Load",
        "unit": "pont",
        "weekly_ref": "260–410%",
        "avg_ref": "55–95%",
        "low": 260,
        "high": 410,
        "avg_low": 55,
        "avg_high": 95,
        "explain": "GPS/rendszer által számolt terhelési pont. A heti load és az edzésátlag is külön értékelődik.",
    },
}

def _fpi_session_kind_v93(x: object) -> str:
    s = _norm_mapping_text(x)
    if any(k in s for k in ["edzes", "training", "train"]):
        return "training"
    if any(k in s for k in ["meccs", "merkozes", "match", "game"]):
        return "match"
    return "other"

def _fpi_safe_pct_v93(num: float, den: float) -> Optional[float]:
    try:
        num = float(num or 0)
        den = float(den or 0)
        if den <= 0:
            return None
        return num / den * 100
    except Exception:
        return None

def _fpi_ratio_status_v93(value: Optional[float], low: float, high: float) -> str:
    if value is None:
        return "n.a."
    if value < low:
        return "alacsony"
    if value > high:
        return "magas"
    return "célzónában"

def _fpi_ratio_note_v93(metric: str, weekly_pct: Optional[float], avg_pct: Optional[float]) -> str:
    ref = FPI_NB2_ADULT_REFERENCE_RANGES_V93.get(metric, {})
    status_w = _fpi_ratio_status_v93(weekly_pct, ref.get("low", 0), ref.get("high", 9999))
    status_a = _fpi_ratio_status_v93(avg_pct, ref.get("avg_low", 0), ref.get("avg_high", 9999))
    if weekly_pct is None:
        return "Nincs meccs referencia vagy nincs értelmezhető adat."
    if status_w == "alacsony" and metric in ["hsr_distance", "sprint_distance", "sprints"]:
        return "Heti inger alacsony lehet; érdemes kontrollált sebesség/sprint expozíciót tervezni."
    if status_w == "magas":
        return "Heti összterhelés magas; nézd meg, egy napra koncentrálódik-e."
    if avg_pct is not None and status_a == "magas":
        return "Az edzésátlag magas; egy-egy edzés közel meccsterhelésű lehet."
    if status_w == "célzónában" and status_a == "célzónában":
        return "A heti összeg és az edzésátlag is referenciazónában van."
    return f"Heti: {status_w}, edzésátlag: {status_a}."

def _fpi_event_count_v93(df: pd.DataFrame) -> int:
    if df is None or df.empty:
        return 0
    candidates = [c for c in ["session_id", "session_name", "start_time", "date", "session_date"] if c in df.columns]
    if candidates:
        tmp = df[candidates].astype(str).agg(" | ".join, axis=1)
        return int(tmp.nunique())
    return 1

def _fpi_match_ratio_reference_df_v93(df: pd.DataFrame, week: str) -> pd.DataFrame:
    """Edzés összes / meccs összes és edzésátlag / meccs arány.
    választott referenciazónával.
    """
    if df is None or df.empty:
        return pd.DataFrame()

    data = df.copy()
    if "week" in data.columns and week is not None:
        data = data[data["week"].astype(str) == str(week)].copy()
    if data.empty or "session_type" not in data.columns:
        return pd.DataFrame()

    data["_kind_v93"] = data["session_type"].apply(_fpi_session_kind_v93)
    train = data[data["_kind_v93"] == "training"].copy()
    match = data[data["_kind_v93"] == "match"].copy()

    rows = []
    training_events = max(_fpi_event_count_v93(train), 1) if not train.empty else 0

    for metric, ref in FPI_NB2_ADULT_REFERENCE_RANGES_V93.items():
        if metric not in data.columns:
            continue
        train_total = float(pd.to_numeric(train.get(metric, pd.Series(dtype=float)), errors="coerce").fillna(0).sum()) if not train.empty else 0.0
        match_total = float(pd.to_numeric(match.get(metric, pd.Series(dtype=float)), errors="coerce").fillna(0).sum()) if not match.empty else 0.0
        train_avg = train_total / training_events if training_events else 0.0
        weekly_pct = _fpi_safe_pct_v93(train_total, match_total)
        avg_pct = _fpi_safe_pct_v93(train_avg, match_total)

        rows.append({
            "Mutató": ref["label"],
            "Mit mér?": ref["explain"],
            "Edzés heti összes": train_total,
            "Meccs összes": match_total,
            "Edzés/heti meccs %": weekly_pct,
            "Edzésátlag/meccs %": avg_pct,
            "NB2 felnőtt heti ref.": ref["weekly_ref"],
            "NB2 felnőtt edzésátlag ref.": ref["avg_ref"],
            "Értékelés": _fpi_ratio_note_v93(metric, weekly_pct, avg_pct),
        })

    return pd.DataFrame(rows)

def _fpi_fmt_pct_v93(x: Optional[float]) -> str:
    if x is None or pd.isna(x):
        return "n.a."
    return f"{float(x):.0f}%"

def _fpi_fmt_num_v93(x: object) -> str:
    try:
        return f"{float(x):.0f}"
    except Exception:
        return "0"



# =========================================================
# V9.4 - Meccsnap / ellenfél / hétkeveredés guard
# =========================================================

def _fpi_iso_week_from_date_v94(d: object) -> Optional[str]:
    try:
        if d is None or pd.isna(d):
            return None
        ts = pd.to_datetime(d, errors="coerce")
        if pd.isna(ts):
            return None
        return f"{int(ts.isocalendar().year)}-W{int(ts.isocalendar().week):02d}"
    except Exception:
        return None

def _fpi_selected_match_context_v94() -> Dict[str, object]:
    return st.session_state.get("fpi_match_context_v94", {}) or {}

def _fpi_session_kind_simple_v94(x: object) -> str:
    return _fpi_session_kind_v93(x)

def _fpi_week_context_df_v94(df: pd.DataFrame, match_date: Optional[object] = None) -> pd.DataFrame:
    """Hetek és feltöltött események áttekintése.
    Cél: ne keveredjenek a különböző hetek / előző hét / aktuális hét / meccshét fájlok.
    """
    if df is None or df.empty or "week" not in df.columns:
        return pd.DataFrame()
    data = df.copy()
    if "session_date" in data.columns:
        data["_date_v94"] = pd.to_datetime(data["session_date"], errors="coerce")
    elif "start_time" in data.columns:
        data["_date_v94"] = pd.to_datetime(data["start_time"], errors="coerce")
    else:
        data["_date_v94"] = pd.NaT
    data["_kind_v94"] = data["session_type"].apply(_fpi_session_kind_simple_v94) if "session_type" in data.columns else "other"

    match_week = _fpi_iso_week_from_date_v94(match_date)
    today = pd.Timestamp.today().normalize()
    today_week = _fpi_iso_week_from_date_v94(today)

    rows = []
    for week, g in data.groupby("week", dropna=True):
        dates = g["_date_v94"].dropna()
        train_events = _fpi_event_count_v93(g[g["_kind_v94"] == "training"])
        match_events = _fpi_event_count_v93(g[g["_kind_v94"] == "match"])
        label = []
        if str(week) == str(today_week):
            label.append("aktuális hét")
        if match_week and str(week) == str(match_week):
            label.append("meccshét")
        if match_week and str(week) < str(match_week):
            label.append("előző / felvezető hét")
        if match_week and str(week) > str(match_week):
            label.append("meccs utáni / jövő hét")
        if not label:
            label.append("feltöltött hét")
        rows.append({
            "Hét": str(week),
            "Státusz": ", ".join(label),
            "Dátum min": dates.min().strftime("%Y-%m-%d") if len(dates) else "n.a.",
            "Dátum max": dates.max().strftime("%Y-%m-%d") if len(dates) else "n.a.",
            "Edzés esemény": int(train_events),
            "Meccs esemény": int(match_events),
            "Sor": int(len(g)),
            "Játékos": int(g["player_name"].nunique()) if "player_name" in g.columns else 0,
        })
    return pd.DataFrame(rows).sort_values("Hét")

def _fpi_match_week_warning_v94(df: pd.DataFrame, selected_week: str, match_date: Optional[object]) -> List[str]:
    warnings = []
    if df is None or df.empty or "week" not in df.columns:
        return ["Nincs értelmezhető hétadat."]
    weeks = sorted([str(x) for x in df["week"].dropna().unique()])
    match_week = _fpi_iso_week_from_date_v94(match_date)
    today_week = _fpi_iso_week_from_date_v94(pd.Timestamp.today().normalize())

    if match_date is None:
        warnings.append("Nincs megadva meccsnap. Add meg, hogy melyik mérkőzésre készül a riport.")
    elif match_week and str(selected_week) != str(match_week):
        warnings.append(f"A kiválasztott hét ({selected_week}) nem egyezik a meccsnap hetével ({match_week}). Ellenőrizd, hogy nem előző heti / másik heti fájlokat nézel-e.")

    if len(weeks) > 1:
        warnings.append(f"Több hét van a feltöltött adatban: {', '.join(weeks[:6])}{'…' if len(weeks)>6 else ''}. A riport a kiválasztott hétre készül, de a feltöltés több hetet tartalmaz.")
    if today_week and str(selected_week) != str(today_week):
        warnings.append(f"Mai nap alapján az aktuális hét: {today_week}. Most a kiválasztott hét: {selected_week}. Ez lehet szándékos, de érdemes ellenőrizni.")
    return warnings

def _fpi_match_context_label_v94() -> str:
    ctx = _fpi_selected_match_context_v94()
    opponent = ctx.get("opponent") or "n.a."
    md = ctx.get("match_date")
    mw = ctx.get("match_week") or "n.a."
    md_txt = str(md) if md else "n.a."
    return f"Ellenfél: {opponent} | Meccsnap: {md_txt} | Meccshét: {mw}"


def _fpi_demo_match_context_label_v122() -> str:
    """Minta PDF-ekhez: konkrét, eladható meccskontextus N/A és adatminőségi figyelmeztetés nélkül."""
    return "Mérkőzés: KTE U19 – Soroksár SC | Meccsnap: 2026.06.06. szombat | Meccshét: 2026-W23"


def _fpi_demo_week_check_text_v122() -> str:
    return "A riport egy mérkőzésre felkészítő mikrociklusra készült. A kiválasztott hét és a meccsnap összhangban van."


def _fpi_pdf_match_context_line_v122(demo_label: str = "") -> str:
    return _fpi_demo_match_context_label_v122() if demo_label else _fpi_match_context_label_v94()



# =========================================================
# V9.5 - KTE/Kecskemét saját csapat + GPS-only mikrociklus
# =========================================================

FPI_OWN_TEAM_ALIASES_V95 = [
    "kecskemet", "kecskemeti", "kecskeméti", "kecskemeti te", "kecskeméti te",
    "kecskemeti te hufbau", "kecskeméti te hufbau", "kecskemeti lc", "kecskemeti lc kte",
    "kte", "kte hufbau", "kecskemet te", "kecskemét", "kecskemét te",
]

def _fpi_norm_team_v95(x: object) -> str:
    s = unicodedata.normalize("NFKD", str(x or "").lower())
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()

def _fpi_is_own_team_v95(name: object) -> bool:
    n = _fpi_norm_team_v95(name)
    if not n:
        return False
    return any(alias in n for alias in [_fpi_norm_team_v95(a) for a in FPI_OWN_TEAM_ALIASES_V95])

def _fpi_pretty_team_name_v95(name: object) -> str:
    return "KTE / Kecskemét" if _fpi_is_own_team_v95(name) else str(name or "").strip()

def _fpi_mark_own_opponent_teams_v95(teams: Dict[str, str]) -> Dict[str, str]:
    out = dict(teams or {})
    a = out.get("team_a", "")
    b = out.get("team_b", "")
    if _fpi_is_own_team_v95(a):
        out["own_team"] = _fpi_pretty_team_name_v95(a)
        out["opponent_team"] = str(b or "Ellenfél")
        out["own_side"] = "a"
    elif _fpi_is_own_team_v95(b):
        out["own_team"] = _fpi_pretty_team_name_v95(b)
        out["opponent_team"] = str(a or "Ellenfél")
        out["own_side"] = "b"
    else:
        out["own_team"] = "KTE / Kecskemét"
        out["opponent_team"] = str(b or a or "Ellenfél")
        out["own_side"] = "unknown"
    return out

def _fpi_has_tactical_signal_v95(tactical_context: Optional[Dict[str, object]]) -> bool:
    if not tactical_context:
        return False
    keys = ["tactical_findings", "pdf_provider_lines", "pdf_provider_findings", "own_topics", "opp_topics", "team_comparison"]
    for k in keys:
        v = tactical_context.get(k)
        if isinstance(v, list) and len(v) > 0:
            return True
        if isinstance(v, dict) and len(v) > 0:
            return True
    if tactical_context.get("has_own_team_excel") or tactical_context.get("has_opp_team_excel") or tactical_context.get("has_own_pdf") or tactical_context.get("has_opp_pdf"):
        return True
    return False

def _fpi_gps_week_metrics_v95(ctx: Dict[str, object], week: str) -> Dict[str, object]:
    df = ctx.get("df")
    if not isinstance(df, pd.DataFrame) or df.empty:
        return {}
    d = df.copy()
    if "week" in d.columns:
        d = d[d["week"].astype(str) == str(week)].copy()
    if d.empty:
        return {}
    ratio = _fpi_match_ratio_reference_df_v97(d, week)
    ratio_map = {}
    if isinstance(ratio, pd.DataFrame) and not ratio.empty:
        for _, r in ratio.iterrows():
            ratio_map[str(r.get("Mutató", ""))] = {
                "weekly_pct": r.get("Edzés/heti meccs %"),
                "avg_pct": r.get("Edzésátlag/meccs %"),
                "eval": r.get("Értékelés", ""),
            }
    return {
        "ratio": ratio,
        "ratio_map": ratio_map,
        "periodization": ctx.get("periodization_type", "Nincs elég adat"),
        "summary": ctx.get("summary", ""),
    }

def _fpi_gps_only_conclusions_v95(ctx: Dict[str, object], priorities: List[dict], readiness: int, week: str, limit: int = 6) -> List[str]:
    out = []
    gps = _fpi_gps_week_metrics_v95(ctx, week)
    ratio = gps.get("ratio")
    if isinstance(ratio, pd.DataFrame) and not ratio.empty:
        for _, r in ratio.iterrows():
            mut = str(r.get("Mutató", ""))
            ev = str(r.get("Értékelés", ""))
            wp = _fpi_fmt_pct_v93(r.get("Edzés/heti meccs %"))
            ap = _fpi_fmt_pct_v93(r.get("Edzésátlag/meccs %"))
            if any(k in ev.lower() for k in ["alacsony", "magas", "kontroll", "referencia", "inger"]):
                out.append(f"{mut}: heti {wp}, edzésátlag {ap}. {ev}")
    for p in priorities or []:
        t = p.get("Teendő", p.get("Cím", "")) if isinstance(p, dict) else str(p)
        why = p.get("Miért", "") if isinstance(p, dict) else ""
        if t:
            out.append(f"{t}. {why}".strip())
    if readiness < 55:
        out.insert(0, f"Readiness {readiness}/100: óvatos, csökkentett volumenű hét javasolt.")
    elif readiness < 70:
        out.insert(0, f"Readiness {readiness}/100: vállalható hét, de a HSR/sprint inger adagolását kontrollálni kell.")
    else:
        out.insert(0, f"Readiness {readiness}/100: jó alapállapot, tervezhető specifikus HSR/sprint expozíció.")
    out.append(f"Periodizációs jelleg: {gps.get('periodization', 'Nincs elég adat')}.")
    # unique, compact
    uniq = []
    seen = set()
    for x in out:
        y = _fpi_clean_sentence_v82(x, 190)
        if y and y not in seen:
            uniq.append(y); seen.add(y)
    return uniq[:limit]

def _fpi_gps_only_md_plan_v95(ctx: Dict[str, object], readiness: int, priorities: List[dict], week: str) -> List[Tuple[str, str, str]]:
    """Taktikai input nélkül is GPS-alapú, változó mikrociklus.
    Nem fix sablon: a readiness, HSR/sprint/High Effort arányok és priority-k alapján állít fókuszt.
    """
    ratio = _fpi_match_ratio_reference_df_v97(ctx.get("df", pd.DataFrame()), week)
    low_hsr = low_sprint = high_vol = high_eff = False
    if isinstance(ratio, pd.DataFrame) and not ratio.empty:
        for _, r in ratio.iterrows():
            m = str(r.get("Mutató", "")).lower()
            eval_txt = str(r.get("Értékelés", "")).lower()
            weekly = r.get("Edzés/heti meccs %")
            if "hsr" in m and ("alacsony" in eval_txt or (pd.notna(weekly) and float(weekly) < 150)):
                low_hsr = True
            if "sprint" in m and ("alacsony" in eval_txt or (pd.notna(weekly) and float(weekly) < 100)):
                low_sprint = True
            if "volumen" in m and (pd.notna(weekly) and float(weekly) > 420):
                high_vol = True
            if "high efforts" in m and (pd.notna(weekly) and float(weekly) > 280):
                high_eff = True

    ptxt = " ".join([str(p.get("Teendő", p.get("Cím", ""))) for p in priorities or [] if isinstance(p, dict)]).lower()
    if "sprint" in ptxt:
        low_sprint = True
    if "hsr" in ptxt or "nagy sebess" in ptxt:
        low_hsr = True
    if "terhel" in ptxt and ("magas" in ptxt or "csökk" in ptxt):
        high_vol = True

    if readiness < 55:
        return [
            ("MD-4", "Regeneráció + alacsony/közepes volumen", "Terhelési visszarendezés, readiness javítása."),
            ("MD-3", "Rövid HSR/sprint expozíció", "Csak idegrendszeri inger, alacsony ismétlésszám."),
            ("MD-2", "Frissítés + technikai intenzitás", "Rövid blokkok, terhelés kontroll."),
            ("MD-1", "Aktiváció", "Rövid gyors döntések, friss láb prioritás."),
        ]
    if high_vol or high_eff:
        return [
            ("MD-4", "Volumen kontroll + aerob visszarendezés", "A heti összterhelés/High Effort magas, ne halmozzunk új terhelést."),
            ("MD-3", "Minőségi, rövid HSR/sprint inger", "Sebességexpozíció volumen nélkül."),
            ("MD-2", "Alacsony volumen + reakciógyorsaság", "Readiness megtartása, frissítés."),
            ("MD-1", "Aktiváció + mobilitás", "Idegrendszeri frissítés, minimális fárasztás."),
        ]
    if low_hsr or low_sprint:
        return [
            ("MD-4", "Stabil csapatvolumen", "Általános heti alap terhelés felépítése."),
            ("MD-3", "HSR/sprint expozíció nap", "Hiányzó nagy sebességű/sprint inger pótlása kontrolláltan."),
            ("MD-2", "High Effort + rövid intenzív blokkok", "Pressing/transition jellegű fizikai inger, de rövid volumen."),
            ("MD-1", "Aktiváció", "Frissítés, 3-5 rövid gyors akció."),
        ]
    return [
        ("MD-4", "Fő terhelési nap", "Volumen és HSR referenciazóna fenntartása."),
        ("MD-3", "Specifikus intenzitási nap", "HSR/sprint/High Effort kontrollált meccsinger."),
        ("MD-2", "Terhelés csökkentése", "Frissesség megtartása, egyéni risk kontroll."),
        ("MD-1", "Aktiváció", "Rövid gyorsasági és döntési inger."),
    ]


def build_fpi_product_pdf_bytes(
    data: pd.DataFrame,
    selected_week: Optional[str] = None,
    playstyle: str = "Kiegyensúlyozott",
    report_type: str = "full",
    demo_label: str = "",
    tactical_context: Optional[Dict[str, object]] = None,
) -> Optional[bytes]:
    """Egységes PDF motor.

    report_type:
    - executive: 1-2 oldalas vezetőedzői / sportigazgatói riport
    - fitness: részletes erőnléti riport
    - microcycle: múlt hét / aktuális hét / jövő hét terv
    - full: teljes döntéstámogató csomag
    """
    if SimpleDocTemplate is None:
        return None
    from reportlab.platypus import PageBreak

    ctx = _fpi_report_context(data, selected_week, playstyle)
    if ctx.get("error"):
        return None
    tactical_context = tactical_context if tactical_context is not None else st.session_state.get("tactical_pro_context", None)

    df = ctx["df"]
    week = ctx["selected_week"]
    readiness = int(ctx["readiness_score"])
    risk_df = ctx["risk_df"] if isinstance(ctx.get("risk_df"), pd.DataFrame) else pd.DataFrame()
    priorities = ctx.get("priorities", []) or []
    insights_df = ctx["insights_df"] if isinstance(ctx.get("insights_df"), pd.DataFrame) else pd.DataFrame()
    weekly = ctx["weekly"] if isinstance(ctx.get("weekly"), pd.DataFrame) else pd.DataFrame()
    player_week = ctx["player_week"] if isinstance(ctx.get("player_week"), pd.DataFrame) else pd.DataFrame()
    player_actions = ctx["player_actions_df"] if isinstance(ctx.get("player_actions_df"), pd.DataFrame) else pd.DataFrame()
    next_df = ctx["next_df"] if isinstance(ctx.get("next_df"), pd.DataFrame) else pd.DataFrame()
    current_df = ctx["current_df"] if isinstance(ctx.get("current_df"), pd.DataFrame) else pd.DataFrame()
    high_risk, med_risk = _fpi_count_risk_levels_v126(risk_df)
    week_type_label_v126 = _fpi_pdf_week_type_label_v126(ctx, demo_label)

    font_name, font_bold = _register_pdf_font()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=0.9*cm, leftMargin=0.9*cm, topMargin=0.7*cm, bottomMargin=0.7*cm)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("FPI58Title", parent=styles["Title"], fontName=font_bold, fontSize=20, leading=23, textColor=colors.HexColor("#0F172A"))
    sub = ParagraphStyle("FPI58Sub", parent=styles["Normal"], fontName=font_name, fontSize=8.8, leading=11, textColor=colors.HexColor("#334155"))
    h2 = ParagraphStyle("FPI58H2", parent=styles["Heading2"], fontName=font_bold, fontSize=11.2, leading=14, textColor=colors.HexColor("#0F172A"))
    body = ParagraphStyle("FPI58Body", parent=styles["Normal"], fontName=font_name, fontSize=8.0, leading=10.2, textColor=colors.HexColor("#111827"))
    small = ParagraphStyle("FPI58Small", parent=styles["Normal"], fontName=font_name, fontSize=6.6, leading=7.9, textColor=colors.HexColor("#111827"))
    head = ParagraphStyle("FPI58Head", parent=styles["Normal"], fontName=font_bold, fontSize=7.2, leading=8.8, alignment=1, textColor=colors.white)
    white_big = ParagraphStyle("FPI58WhiteBig", parent=styles["Normal"], fontName=font_bold, fontSize=16, leading=18, alignment=1, textColor=colors.white)
    white_small = ParagraphStyle("FPI58WhiteSmall", parent=styles["Normal"], fontName=font_name, fontSize=7.0, leading=8.5, alignment=1, textColor=colors.white)

    def clean(v: object) -> str:
        text = pdf_safe_text(v).replace("\r", "")
        text = re.sub(r"(?i)<br\s*/?>", "\n", text)
        text = re.sub(r"(?i)</?b>", "", text)
        text = re.sub(r"(?is)\{\s*['\"]Téma['\"].*?\}", "", text)
        text = re.sub(r"(?is)\{\s*['\"]Tema['\"].*?\}", "", text)
        text = re.sub(r"(?is)\{\s*['\"]Bizonyíték['\"].*?\}", "", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def P(v, style=body):
        return Paragraph(html.escape(clean(v)).replace("\n", "<br/>"), style)

    def section(text: str, color: str = "#DBEAFE"):
        t = Table([[P(text, h2)]], colWidths=[27.7*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), colors.HexColor(color)),
            ("BOX", (0,0), (-1,-1), 0.4, colors.HexColor("#93C5FD")),
            ("LEFTPADDING", (0,0), (-1,-1), 7), ("RIGHTPADDING", (0,0), (-1,-1), 7),
            ("TOPPADDING", (0,0), (-1,-1), 4), ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        return t

    def table(rows, widths, header_bg="#0F172A", row_bgs=None):
        if row_bgs is None:
            row_bgs = [colors.white, colors.HexColor("#F8FAFC")]
        t = Table(rows, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor(header_bg)),
            ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), row_bgs),
            ("LEFTPADDING", (0,0), (-1,-1), 4), ("RIGHTPADDING", (0,0), (-1,-1), 4),
            ("TOPPADDING", (0,0), (-1,-1), 3), ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ]))
        return t

    def kpi(label, value, note, color):
        t = Table([[P(label, white_small)], [P(value, white_big)], [P(note, white_small)]], colWidths=[5.25*cm], rowHeights=[0.5*cm, 0.85*cm, 0.5*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), colors.HexColor(color)),
            ("BOX", (0,0), (-1,-1), 0.4, colors.white),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING", (0,0), (-1,-1), 3), ("RIGHTPADDING", (0,0), (-1,-1), 3),
        ]))
        return t

    story = []
    label_prefix = f"{demo_label} | " if demo_label else ""
    report_names = {
        "executive": "Vezetői riport",
        "fitness": "Erőnléti szakmai riport",
        "microcycle": "Mikrociklus döntéstámogató riport",
        "full": "Teljes stáb riportcsomag",
    }

    def add_cover():
        story.append(P("Football Performance Intelligence", title))
        story.append(P(f"{label_prefix}{report_names.get(report_type, 'Riport')} | Hét: {format_week_label(str(week))} | Heti típus: {week_type_label_v126} | Játékmodell: {playstyle} | Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}", sub))
        story.append(P(_fpi_pdf_match_context_line_v122(demo_label), sub))
        story.append(Spacer(1, 0.25*cm))
        kpis = [
            kpi("READINESS", f"{readiness}/100", score_to_label(readiness), "#166534" if readiness >= 75 else "#1E3A8A" if readiness >= 60 else "#991B1B"),
            kpi("HIGH RISK", f"{high_risk} fő", "egyéni kontroll", "#7F1D1D" if high_risk else "#166534"),
            kpi("MEDIUM RISK", f"{med_risk} fő", "figyelendő játékos", "#92400E" if med_risk else "#166534"),
            kpi("INSIGHT", str(len(insights_df)), "automatikus megállapítás", "#0F172A"),
            kpi("HETEK", str(df['week'].nunique()) if 'week' in df.columns else "—", "elemzett adatbázis", "#1E3A8A"),
        ]
        story.append(Table([kpis], colWidths=[5.45*cm]*5))
        story.append(Spacer(1, 0.25*cm))

    def add_executive_page():
        story.append(section("1. Vezetői oldal – heti meccs- és edzésterv", "#DBEAFE"))

        has_tactical_signal = _fpi_has_tactical_signal_v95(tactical_context)
        tactical_plan = tactical_context.get("plan_a", "KIE – kiegyensúlyozott") if has_tactical_signal else "GPS-only – erőnléti fókuszú mikrociklus"
        if has_tactical_signal:
            fitness_insights_v146 = _fpi_contextual_gps_only_insights_v146(
                ctx, priorities, readiness, week, 5
            )
            coach_blocks_v146 = _fpi_coach_blocks_v146(
                tactical_context, ctx, readiness, priorities, week
            )
            tactical_insights_v146 = coach_blocks_v146["tactical"]
            match_plan_insights_v146 = coach_blocks_v146["plan"]
            team_tactical_insights_v146 = coach_blocks_v146["team"]
        else:
            fitness_insights_v146 = _fpi_contextual_gps_only_insights_v146(
                ctx, priorities, readiness, week, 7
            )
            tactical_insights_v146 = []
            match_plan_insights_v146 = []
            team_tactical_insights_v146 = []

        fitness_msgs = [
            _fpi_render_insight_text_v146(x) for x in fitness_insights_v146[:5 if not has_tactical_signal else 3]
        ]
        tactical_msgs = [
            _fpi_render_insight_text_v146(x) for x in tactical_insights_v146
        ]
        team_tactical_msgs = [
            _fpi_render_insight_text_v146(x) for x in team_tactical_insights_v146
        ]
        md_rows_simple = _fpi_contextual_md_plan_rows_v146(
            tactical_context,
            gps_context=ctx,
            readiness=readiness,
            priorities=priorities,
            week=week,
        )
        risk_rows_simple = _fpi_compact_player_risk_rows_v82(risk_df, 12)

        # 1) felső döntési sáv: több taktikai megállapítás és konkrét meccsterv
        opp_eval_exec = (tactical_context or {}).get("opponent_player_evaluation", []) if isinstance(tactical_context, dict) else []

        tactical_findings_text = "\n\n".join(
            [f"• {_fpi_render_insight_text_v146(x)}" for x in tactical_insights_v146[:5]]
        ) or "• Nincs elegendő taktikai input; a riport GPS-only módban készült."

        opponent_focus_lines = []
        for row in opp_eval_exec[:3]:
            row_insight = _fpi_normalize_any_insight_v146(row)
            name = str(row.get("Játékos", "Ismeretlen játékos")) if isinstance(row, dict) else "Ismeretlen játékos"
            role = _fpi_extract_coach_text_v145(row.get("Szerep", ""), 70) if isinstance(row, dict) else ""
            eval_text = _fpi_extract_coach_text_v145(row.get("Értelmezés", ""), 150) if isinstance(row, dict) else ""
            clean_line = _fpi_strip_raw_repr_v146(f"{name} – {role}: {eval_text}")
            if clean_line:
                opponent_focus_lines.append(f"• {clean_line}")
        opponent_focus_text = "\n\n".join(opponent_focus_lines) or "• Nincs ellenfél-játékos Excel vagy azonosítható játékosprofil."

        match_plan_text = "\n\n".join(
            [f"• {_fpi_render_insight_text_v146(x)}" for x in match_plan_insights_v146[:6]]
        ) or f"• {_fpi_hu_plain_text_v144(tactical_plan)}"

        fast_rows = [[P("Taktikai megállapítások", head), P("Ellenfél játékosok", head), P("Meccsterv", head)]]
        fast_rows.append([
            P(tactical_findings_text, body),
            P(opponent_focus_text, body),
            P(match_plan_text, body),
        ])
        story.append(table(fast_rows, [9.2*cm, 9.2*cm, 9.3*cm], header_bg="#1E3A8A", row_bgs=[colors.HexColor("#EFF6FF")]))
        story.append(Spacer(1, 0.18*cm))

        # 1. pont lezárása: a heti ciklusterv ugyanazon az oldalon, tömören
        story.append(section("Heti ciklusterv – erőnléti + taktikai cél", "#EDE9FE"))
        detailed_md_rows_v147 = _fpi_weekly_fitness_rows_v147(
            md_rows_simple, fitness_insights_v146, readiness
        )
        md_table = [[
            P("Nap", head),
            P("Erőnléti fókusz", head),
            P("Taktikai fókusz", head),
            P("Edzői megjegyzés", head),
        ]]
        for d, fgoal, tgoal, coach_note in detailed_md_rows_v147:
            md_table.append([
                P(d, small),
                P(fgoal, small),
                P(tgoal, small),
                P(coach_note, small),
            ])
        story.append(table(
            md_table,
            [2.4*cm, 8.2*cm, 8.6*cm, 8.5*cm],
            header_bg="#312E81",
            row_bgs=[colors.HexColor("#F5F3FF"), colors.white],
        ))

        # 2. pont – külön oldal
        story.append(PageBreak())
        story.append(section("2. Fő edzői üzenetek", "#DCFCE7"))
        story.append(P(
            "A taktikai és erőnléti üzenetek az adott ellenfélhez, heti terhelési állapothoz és mikrociklushoz igazodnak.",
            small,
        ))
        story.append(Spacer(1, 0.12*cm))
        msg_rows = [[P("Erőnléti üzenet", head), P("Csapatszintű taktikai üzenet", head)]]
        max_len = max(len(fitness_msgs), len(team_tactical_msgs), 4)
        for i in range(max_len):
            fm = fitness_msgs[i] if i < len(fitness_msgs) else ""
            tm = team_tactical_msgs[i] if i < len(team_tactical_msgs) else ""
            msg_rows.append([
                P(_fpi_strip_raw_repr_v146(_fpi_extract_coach_text_v145(fm, 420)), small),
                P(_fpi_strip_raw_repr_v146(_fpi_extract_coach_text_v145(tm, 420)), small),
            ])
        story.append(table(
            msg_rows,
            [13.8*cm, 13.9*cm],
            header_bg="#166534",
            row_bgs=[colors.HexColor("#ECFDF5"), colors.white],
        ))

        # 3. pont – külön oldal, minden játékosszintű információ egy helyen
        story.append(PageBreak())
        story.append(section("3. Játékosszintű fókusz", "#FEE2E2"))

        if opp_eval_exec:
            story.append(P("Ellenfél kulcsemberei és támadható pontjai", head))
            op_rows = [[
                P("Játékos", head),
                P("Szerep", head),
                P("Mutató + referencia", head),
                P("Értékelés", head),
                P("Meccstervi teendő", head),
            ]]
            for r in opp_eval_exec[:5]:
                op_rows.append([
                    P(_fpi_extract_coach_text_v145(r.get("Játékos", ""), 60), small),
                    P(_fpi_extract_coach_text_v145(r.get("Szerep", ""), 90), small),
                    P(_fpi_strip_raw_repr_v146(_fpi_extract_coach_text_v145(r.get("Bizonyíték", ""), 180)), small),
                    P(_fpi_strip_raw_repr_v146(_fpi_extract_coach_text_v145(r.get("Értelmezés", ""), 210)), small),
                    P(_fpi_strip_raw_repr_v146(_fpi_extract_coach_text_v145(r.get("Javaslat", ""), 230)), small),
                ])
            story.append(table(
                op_rows,
                [3.4*cm, 4.4*cm, 7.0*cm, 6.2*cm, 6.7*cm],
                header_bg="#991B1B",
                row_bgs=[colors.HexColor("#FEF2F2"), colors.white],
            ))
            story.append(Spacer(1, 0.20*cm))

        story.append(P("Saját játékosok terhelési kockázata", head))
        rr = [[P(c, head) for c in risk_rows_simple[0]]]
        for row in risk_rows_simple[1:]:
            rr.append([P(x, small) for x in row])
        story.append(table(
            rr,
            [6.4*cm, 4.2*cm, 17.1*cm],
            header_bg="#7F1D1D",
            row_bgs=[colors.white, colors.HexColor("#FEF2F2")],
        ))

        # Legalul: a GPS-only tudásmotor teljesebb erőnléti helyzetképe
        story.append(PageBreak())
        story.append(section("Kiegészítő erőnléti helyzetkép – GPS Intelligence", "#E0F2FE"))
        story.append(P(
            "Ez a szekció ugyanazt a kombinációs és ismétléscsökkentő logikát használja, mint a GPS-only riport.",
            small,
        ))
        story.append(Spacer(1, 0.12*cm))
        fitness_snapshot_v147 = _fpi_fitness_snapshot_rows_v147(
            _fpi_contextual_gps_only_insights_v146(ctx, priorities, readiness, week, 8),
            8,
        )
        fs_rows = [[P(c, head) for c in fitness_snapshot_v147[0]]]
        for row in fitness_snapshot_v147[1:]:
            fs_rows.append([P(x, small) for x in row])
        story.append(table(
            fs_rows,
            [5.2*cm, 10.8*cm, 11.7*cm],
            header_bg="#075985",
            row_bgs=[colors.HexColor("#F0F9FF"), colors.white],
        ))


    def add_fitness_page():
        story.append(section("2. Erőnléti szakmai oldal – GPS terhelés, trend és játékosszint", "#E0F2FE"))
        # heti csapatösszegzés
        wk = weekly.copy()
        if not wk.empty:
            wk = wk[wk["week"].astype(str) == str(week)].copy()
        cols = [c for c in ["session_type", "total_distance", "duration_min", "distance_per_min", "hsr_distance", "sprint_distance", "sprints", "high_efforts", "training_load", "max_speed"] if c in wk.columns]
        rows = [[P("Típus", head), P("Össztáv", head), P("Perc", head), P("m/perc", head), P("HSR", head), P("Sprint táv", head), P("Sprint", head), P("High Eff.", head), P("Load", head), P("Max seb.", head)]]
        if not wk.empty:
            for _, r in wk.head(8).iterrows():
                rows.append([
                    P(r.get("session_type", ""), small),
                    P(_fpi_fmt_thousands_v97(r.get("total_distance", 0)), small),
                    P(_fpi_fmt_thousands_v97(r.get("duration_min", 0)), small),
                    P(f"{r.get('distance_per_min', 0):.1f}", small),
                    P(_fpi_fmt_thousands_v97(r.get("hsr_distance", 0)), small),
                    P(_fpi_fmt_thousands_v97(r.get("sprint_distance", 0)), small),
                    P(_fpi_fmt_thousands_v97(r.get("sprints", 0)), small),
                    P(_fpi_fmt_thousands_v97(r.get("high_efforts", 0)), small),
                    P(_fpi_fmt_thousands_v97(r.get("training_load", 0)), small),
                    P(f"{r.get('max_speed', 0):.1f}", small),
                ])
        else:
            rows.append([P("Nincs adat", small)] + [P("—", small)]*9)
        story.append(table(rows, [2.8*cm, 3.0*cm, 2.5*cm, 2.5*cm, 3.0*cm, 3.0*cm, 2.4*cm, 2.5*cm, 2.5*cm, 3.0*cm], header_bg="#0369A1"))
        story.append(Spacer(1, 0.25*cm))

        ratio_df = _fpi_match_ratio_reference_df_v97(df, week)
        story.append(section("Edzés–meccs arányok – százalékos referencia NB2 felnőtt szinthez", "#FEF3C7"))
        if ratio_df.empty:
            story.append(Paragraph(pdf_safe_text("Nincs elég adat az edzés/meccs százalékos referencia kiszámításához. Legalább egy edzés és egy meccs típusú esemény szükséges."), body))
        else:
            rr = [[
                P("Mutató", head),
                P("Edzés heti összes / meccs", head),
                P("Edzésátlag / meccs", head),
                P("választott referencia", head),
                P("Értékelés", head),
            ]]
            for _, r in ratio_df.iterrows():
                rr.append([
                    P(str(r.get("Mutató", "")), small),
                    P(_fpi_fmt_pct_v93(r.get("Edzés/heti meccs %")), small),
                    P(_fpi_fmt_pct_v93(r.get("Edzésátlag/meccs %")), small),
                    P(f"Heti: {r.get('NB2 felnőtt heti ref.', '')}<br/>Edzésátlag: {r.get('NB2 felnőtt edzésátlag ref.', '')}", small),
                    P(str(r.get("Értékelés", "")), small),
                ])
            story.append(table(rr, [4.0*cm, 4.2*cm, 4.2*cm, 6.0*cm, 9.3*cm], header_bg="#92400E", row_bgs=[colors.HexColor("#FFFBEB"), colors.white]))
            story.append(Paragraph(pdf_safe_text("Értelmezés: a heti összes megmutatja, hogy a teljes edzésheted hány meccsnyi ingert adott. Az edzésátlag azt mutatja, hogy egy átlagos edzés hány %-a egy meccsnek. A kettőt együtt kell nézni."), small))
        story.append(Spacer(1, 0.25*cm))

        story.append(section("Fogalmak röviden – HSR, sprint és sprint expozíció", "#DBEAFE"))
        expl = [[P("Fogalom", head), P("Egyszerű jelentés", head)]]
        expl.append([P("HSR", small), P("High Speed Running: nagy sebességű futás, általában kb. 19,8–20 km/h felett. Ez még nem feltétlen maximális sprint.", small)])
        expl.append([P("Sprint", small), P("Sprintzóna vagy sprintakció. A sprint táv a sprintzónában megtett méter, a sprint count pedig a sprintakciók darabszáma.", small)])
        expl.append([P("Sprint expozíció", small), P("Nem csak mennyiség: azt jelenti, hogy a játékos elér-e nagy sebességű / maximális sebességhez közeli ingert a héten. Sérülésmegelőzés miatt fontos.", small)])
        expl.append([P("Load / terhelési pont", small), P("A GPS-rendszer összesített terhelési mutatója. Nem csak futótáv: rendszerfüggően tartalmazhat mozgásintenzitást, gyorsításokat, lassításokat és mechanikai terhelést is. A heti volumen és a játékoskockázat értelmezéséhez használjuk.", small)])
        story.append(table(expl, [4.0*cm, 23.7*cm], header_bg="#1D4ED8", row_bgs=[colors.HexColor("#EFF6FF"), colors.white]))
        story.append(Spacer(1, 0.25*cm))

        story.append(section("Top játékosok – heti terhelési profil", "#F0FDFA"))
        pw = player_week.copy()
        if not pw.empty and "week" in pw.columns:
            pw = pw[pw["week"].astype(str) == str(week)].copy()
            sort_col = "training_load" if "training_load" in pw.columns else "total_distance"
            pw = pw.sort_values(sort_col, ascending=False)
        prows = [[P("Játékos", head), P("Össztáv", head), P("HSR", head), P("Sprint táv", head), P("High Efforts", head), P("Load", head), P("Max seb.", head), P("Értelmezés", head)]]
        if not pw.empty:
            for _, r in pw.head(12).iterrows():
                interp = "Magas heti load – regeneráció kontroll" if float(r.get("training_load", 0) or 0) >= float(pw.get("training_load", pd.Series([0])).quantile(.75) or 0) else "Normál monitoring"
                prows.append([P(r.get("player_name", ""), small), P(_fpi_fmt_thousands_v97(r.get("total_distance",0)), small), P(_fpi_fmt_thousands_v97(r.get("hsr_distance",0)), small), P(_fpi_fmt_thousands_v97(r.get("sprint_distance",0)), small), P(_fpi_fmt_thousands_v97(r.get("high_efforts",0)), small), P(_fpi_fmt_thousands_v97(r.get("training_load",0)), small), P(f"{r.get('max_speed',0):.1f}", small), P(interp, small)])
        else:
            prows.append([P("Nincs játékosszintű adat", small)] + [P("—", small)]*7)
        story.append(table(prows, [4.5*cm, 3.0*cm, 3.0*cm, 3.0*cm, 3.0*cm, 3.0*cm, 2.8*cm, 5.4*cm], header_bg="#0F766E", row_bgs=[colors.white, colors.HexColor("#F0FDFA")]))
        story.append(Spacer(1, 0.2*cm))
        story.append(section("Erőnléti edzői értelmezés", "#FEF3C7"))
        notes = [
            "A riport nem váltja ki az erőnléti edző döntését: előkészíti a heti review-t, kiemeli az eltéréseket és egységes PDF-et ad a stábnak.",
            "A High Efforts külön mezőként szerepel; ha nincs külön oszlop, az app gyorsulás/lassítás alapon becsli.",
            "A 4-es és 5-ös sebességzóna külön is kezelhető, de összevont 4+5 export esetén is használható HSR-ként.",
        ]
        story.append(table([[P("Megjegyzés", head), P("Használat", head)]] + [[P(n, small), P("Heti review / stábmegbeszélés", small)] for n in notes], [20*cm, 7.7*cm], header_bg="#92400E", row_bgs=[colors.HexColor("#FFFBEB"), colors.white]))

    def add_micro_page():
        story.append(section("3. Mikrociklus oldal – múlt hét / aktuális hét / jövő hét", "#EDE9FE"))
        blocks = [
            ("Múlt hét teljes értékelése", ctx.get("past_text", ""), "#312E81"),
            ("Aktuális hét – hátralévő napok javaslata", ctx.get("current_text", ""), "#1E3A8A"),
            ("Jövő heti MD-bontású terv", ctx.get("next_text", ""), "#166534"),
        ]
        bdata = [[P("Szekció", head), P("Automatikus értékelés / javaslat", head)]]
        for name, txt, _ in blocks:
            bdata.append([P(name, body), P(txt or "Nincs elegendő adat a részletes szöveghez.", small)])
        story.append(table(bdata, [6.5*cm, 21.2*cm], header_bg="#312E81", row_bgs=[colors.HexColor("#F5F3FF"), colors.white]))
        story.append(Spacer(1, 0.25*cm))
        if not next_df.empty:
            story.append(section("Jövő hét – strukturált napi terv", "#DCFCE7"))
            cols = [c for c in next_df.columns if c in ["Nap", "MD", "Fókusz", "Terhelés", "Javaslat", "Megjegyzés", "Cél"]]
            if not cols:
                cols = list(next_df.columns[:5])
            nd = [[P(c, head) for c in cols]]
            for _, r in next_df.head(8).iterrows():
                nd.append([P(r.get(c, ""), small) for c in cols])
            story.append(table(nd, [27.7*cm/len(cols)]*len(cols), header_bg="#166534", row_bgs=[colors.HexColor("#ECFDF5"), colors.white]))
        if not player_actions.empty:
            story.append(Spacer(1, 0.25*cm))
            story.append(section("Játékosszintű teendők", "#FEE2E2"))
            cols = list(player_actions.columns[:5])
            ad = [[P(c, head) for c in cols]]
            for _, r in player_actions.head(12).iterrows():
                ad.append([P(r.get(c, ""), small) for c in cols])
            story.append(table(ad, [27.7*cm/len(cols)]*len(cols), header_bg="#7F1D1D", row_bgs=[colors.HexColor("#FEF2F2"), colors.white]))



    def _strategy_palette_pdf_rows():
        """Coach-friendly magyar magyarázat a taktikai stratégiai palettához."""
        return [
            ("KON", "Kontra mély blokkból", "Mélyebb védekezésből gyors, direkt támadásindítás."),
            ("GAT", "Gyors átmenet", "Labdaszerzés után gyors előrejáték, kevés passzból veszély."),
            ("BAT", "Középső blokk + átmenet", "Középső zónás védekezés, majd gyors átmeneti támadás."),
            ("KIE", "Kiegyensúlyozott", "Stabil, kockázatkerülőbb alapjáték, kontrollált intenzitással."),
            ("PRS", "Presszing + átmenet", "Aktív letámadás, labdaszerzés után gyors támadásvezetés."),
            ("MLT", "Magas letámadás", "Magas blokkból agresszív nyomás és korai labdaszerzés."),
            ("DOM", "Dominancia", "Labdabirtoklásra és területi fölényre épülő meccskontroll."),
            ("POZ", "Pozíciós támadás", "Türelmes építkezés, félterületek és szélesség használata."),
            ("LAB", "Labdatartás mélyebben", "Biztonságosabb labdatartás mélyebb szerkezetből."),
        ]

    def _tactical_plain_hungarian_explanation(tactical_context: Dict[str, object]) -> List[str]:
        risks = tactical_context.get("risks", []) or []
        plan_a = tactical_context.get("plan_a", "KIE – kiegyensúlyozott")
        out = [
            f"Javasolt alapirány: {plan_a}.",
            "Az értelmezés lényege: a GPS a saját csapat fizikai állapotát mutatja, a Tactical Pro+ pedig azt, hogy ehhez milyen ellenfélprofil és meccsterv illeszkedik.",
        ]
        if tactical_context.get("has_opp_pdf") or tactical_context.get("has_opp_team_excel"):
            out.append("Az ellenfélanyag alapján a rendszer kiemeli, hol várható veszély: átmenetek, szélső játék, pontrúgások, presszing vagy mély blokk.")
        if tactical_context.get("has_own_pdf") or tactical_context.get("has_own_team_excel"):
            out.append("A saját csapat taktikai anyaga segít eldönteni, hogy a javasolt terv illeszkedik-e a saját játékmodellhez.")
        if risks:
            out.append("Fő fókuszok: " + "; ".join(str(x) for x in risks[:3]) + ".")
        out.append("A javaslat nem kész meccsterv helyett van, hanem a stáb döntését készíti elő.")
        return out


    def _integrated_conclusion_rows(tactical_context: Dict[str, object]) -> List[Tuple[str, str]]:
        risks = tactical_context.get("risks", []) or []
        plan_a = tactical_context.get("plan_a", "KIE – kiegyensúlyozott")
        rows = [
            ("Vezetői konklúzió", f"A javasolt alapirány: {plan_a}. A döntés alapja: saját fizikai állapot + taktikai inputok + ellenfélprofil."),
            ("Fő közös üzenet", "A GPS nem külön, hanem a meccsterv megvalósíthatóságát mutatja: mennyire bírja el a csapat a javasolt taktikai intenzitást."),
        ]
        if risks:
            rows.append(("Top kockázat", "; ".join(str(x) for x in risks[:3])))
        else:
            rows.append(("Top kockázat", "Nincs erős taktikai input; ilyenkor GPS/readiness-alapú terhelési döntéstámogatás működik."))
        rows.append(("Edzői döntés", "A mikrociklusban minden napnak legyen erőnléti és taktikai célja, ne két külön riportként kezeljük."))
        return rows

    def _combined_md_rows(tactical_context: Dict[str, object]) -> List[Tuple[str, str, str, str]]:
        md = tactical_context.get("md_plan", []) or []
        out = []
        for item in md[:6]:
            try:
                day, focus, why = item
            except Exception:
                continue
            fitness_goal = "Terheléskontroll / readiness fenntartása"
            tactical_goal = str(focus)
            why_txt = str(why)
            f_low = tactical_goal.lower()
            if "sprint" in f_low or "hsr" in f_low:
                fitness_goal = "HSR / sprint inger kontrollált adagolása"
            elif "regener" in f_low:
                fitness_goal = "Regeneráció, frissítés, neuromuszkuláris visszarendezés"
            elif "aktiv" in f_low:
                fitness_goal = "Aktiváció, frissesség, döntési gyorsaság"
            elif "volumen" in f_low:
                fitness_goal = "Csapatvolumen és munkasűrűség felépítése"
            out.append((str(day), fitness_goal, tactical_goal, why_txt))
        if not out:
            out = [
                ("MD-4", "Csapatvolumen felépítése", "Saját játékmodell ismétlése", "GPS-only vagy hiányos taktikai input esetén alap mikrociklus."),
                ("MD-3", "HSR / sprint kontrollált inger", "Átmeneti játék vagy meccsintenzitás", "A meccsterhelés előkészítése."),
                ("MD-2", "Terheléscsökkentés", "Meccsterv finomítás", "Frissesség és taktikai tisztaság."),
                ("MD-1", "Aktiváció", "Pontrúgások, döntési gyorsaság", "Rövid, frissítő nap."),
            ]
        return out

    def add_tactical_executive_page():
        def _pdf_tactical_key_numbers_summary(metrics: Dict[str, float]) -> str:
            if not metrics:
                return "Nincs értelmezhető taktikai csapat KPI."
            label_map = {
                "possession_pct": "Labdabirtoklás", "shots": "Lövések", "xg": "xG", "entries_box": "Box entries",
                "pressing_success_pct": "Pressing %", "counterattacks": "Kontrák", "crosses": "Beadások", "corners": "Szögletek",
            }
            parts = []
            for k, lab in label_map.items():
                v = metrics.get(k)
                if v not in [None, 0, 0.0, ""]:
                    try:
                        fv = float(v)
                        if k in ["possession_pct", "pressing_success_pct"] and fv <= 1:
                            fv *= 100
                        parts.append(f"{lab}: {fv:.1f}{'%' if k in ['possession_pct','pressing_success_pct'] else ''}")
                    except Exception:
                        parts.append(f"{lab}: {v}")
            return " | ".join(parts[:6]) if parts else "Nincs kiemelkedő taktikai KPI."

        story.append(section("Integrált taktikai összegzés – GPS + Tactical", "#E0F2FE"))
        if not tactical_context:
            story.append(Paragraph(pdf_safe_text("Ehhez a riporthoz nem volt taktikai PDF/Excel feltöltve, ezért a vezetői értékelés GPS-only módban készült."), body))
            return

        plan_a = str(tactical_context.get("plan_a", "KIE – kiegyensúlyozott"))
        risks = tactical_context.get("risks", []) or []
        status = f"Saját PDF: {'igen' if tactical_context.get('has_own_pdf') else 'nem'} | Ellenfél PDF: {'igen' if tactical_context.get('has_opp_pdf') else 'nem'} | Taktikai Excel: {'igen' if (tactical_context.get('has_own_team_excel') or tactical_context.get('has_opp_team_excel')) else 'nem'}"
        rows = [[P("Terület", head), P("Vezetői üzenet", head)]]
        rows.append([P("Input státusz", small), P(status, small)])
        rows.append([P("Meccsterv irány", small), P(plan_a, small)])
        rows.append([P("Fő kockázat", small), P("; ".join(str(x) for x in risks[:3]) if risks else "Nincs erős taktikai kockázati jelzés.", small)])
        rows.append([P("Saját KPI", small), P(_pdf_tactical_key_numbers_summary(tactical_context.get("own_team_metrics", {}) or {}), small)])
        rows.append([P("Ellenfél KPI", small), P(_pdf_tactical_key_numbers_summary(tactical_context.get("opp_team_metrics", {}) or {}), small)])
        story.append(table(rows, [5.2*cm, 22.5*cm], header_bg="#1E3A8A", row_bgs=[colors.HexColor("#EFF6FF"), colors.white]))
        story.append(Spacer(1, 0.16*cm))

        if tactical_context.get("tactical_findings"):
            story.append(section("Top taktikai következtetések", "#FEF3C7"))
            f_rows = [[P("Téma", head), P("Bizonyíték", head), P("Edzői következtetés", head), P("Prioritás", head)]]
            for f in (tactical_context.get("tactical_findings") or [])[:6]:
                f_rows.append([P(str(f.get("Téma", "")), small), P(str(f.get("Bizonyíték", "")), small), P(str(f.get("Edzői következtetés", "")), small), P(str(f.get("Prioritás", "")), small)])
            story.append(table(f_rows, [5.2*cm, 8.2*cm, 10.2*cm, 4.1*cm], header_bg="#92400E", row_bgs=[colors.HexColor("#FFFBEB"), colors.white]))
            story.append(Spacer(1, 0.16*cm))

        opp_eval_pdf = tactical_context.get("opponent_player_evaluation", []) or []
        if opp_eval_pdf:
            story.append(section("Ellenfél játékosszintű értékelés", "#FEE2E2"))
            op_rows = [[P("Játékos", head), P("Szerep", head), P("Konkrét mutatók", head), P("Értékelés", head), P("Meccstervi teendő", head)]]
            for r in opp_eval_pdf[:8]:
                op_rows.append([P(r.get("Játékos", ""), small), P(r.get("Szerep", ""), small), P(r.get("Bizonyíték", ""), small), P(r.get("Értelmezés", ""), small), P(r.get("Javaslat", ""), small)])
            story.append(table(op_rows, [3.8*cm, 4.5*cm, 5.8*cm, 6.5*cm, 7.1*cm], header_bg="#991B1B", row_bgs=[colors.HexColor("#FEF2F2"), colors.white]))
            story.append(Spacer(1, 0.16*cm))

        if tactical_context.get("md_plan"):
            story.append(section("Integrált mikrociklus – rövid MD-terv", "#EDE9FE"))
            md_rows = [[P("Nap", head), P("Erőnléti cél", head), P("Taktikai cél", head), P("Indoklás", head)]]
            for a, b, c, d in _combined_md_rows(tactical_context)[:4]:
                md_rows.append([P(a, small), P(b, small), P(c, small), P(d, small)])
            story.append(table(md_rows, [3.0*cm, 7.5*cm, 8.0*cm, 9.2*cm], header_bg="#312E81", row_bgs=[colors.HexColor("#F5F3FF"), colors.white]))

    def add_methodology_page():
        story.append(section("Módszertani összefoglaló – hogyan számol a Football Performance Intelligence?", "#DBEAFE"))
        intro = (
            "A Football Performance Intelligence (FPI) döntéstámogató rendszer: a GPS-exportokból egységesíti a heti terhelési képet, "
            "kiemeli a kockázatokat és edzői/erőnléti javaslatokat ad. Nem helyettesíti a szakmai stábot; "
            "a végső döntés mindig edzői, erőnléti és orvosi kontroll mellett történik."
        )
        story.append(Paragraph(pdf_safe_text(intro), body))
        story.append(Spacer(1, 0.20*cm))

        meth_rows = [
            [P("Terület", head), P("Football Performance Intelligence metodika", head)],
            [P("Adatimport", small), P("A Data/Adat lap elsődleges. A segédlapokat az app igyekszik kizárni. A Smart Mapper magyar és angol GPS oszlopneveket is kezel.", small)],
            [P("Dátum és hét", small), P("A Week Rescue Engine a dátumot időponttal vagy extra szöveggel együtt is értelmezi, majd ISO hétre csoportosít. Rövid dátumtartományból képződő irreálisan sok hét esetén védelmi újraértelmezést alkalmaz.", small)],
            [P("Kapusok", small), P("Ha van Poszt/Position oszlop, a kapusok automatikusan felismerhetők. Ha nincs, az app kézi kapusválasztást kér. A kapusok sprint/HSR értelmezése csökkentett súlyú.", small)],
            [P("Játékpercek", small), P("A meccsterhelésnél az app figyelembe veszi, hogy nem minden játékos játszik 90 percet. Ahol elérhető, per90 és csapatperc alapú normalizálást alkalmaz.", small)],
            [P("Edzés-meccs normalizálás", small), P("Az összevetés nem csak nyers csapatösszeg alapján történik, mert edzésen és meccsen eltérhet a játékosszám és a játékidő. A résztvevők száma és az időtartam is számít.", small)],
            [P("Sebességzónák", small), P("A 4-es és 5-ös zóna külön vagy összevont 4+5 exportként is kezelhető. Összevont oszlopnál az app HSR-ként használja az értéket.", small)],
            [P("High Efforts", small), P("Ha külön High Efforts oszlop van, azt használja. Ha nincs, gyorsulás/lassítás jellegű mutatókból becsült nagy intenzitású akciót képez.", small)],
            [P("Benchmark", small), P("A jelenlegi benchmark általános referencia. Későbbi verzióban korosztály, szint, poszt és játékmodell szerint finomíthető.", small)],
            [P("Mikrociklus", small), P("A múlt hét, aktuális hét és következő hét javaslatai a volumen, HSR, sprint, High Efforts, readiness és játékosszintű risk jelzések alapján készülnek.", small)],
            [P("Korlátok", small), P("Az eredmények adatminőségtől, GPS-exporttól és mappingtől függenek. Hibás vagy hiányos input esetén szakmai ellenőrzés szükséges.", small)],
        ]
        story.append(table(meth_rows, [5.2*cm, 22.5*cm], header_bg="#1E3A8A", row_bgs=[colors.HexColor("#EFF6FF"), colors.white]))
        story.append(Spacer(1, 0.20*cm))
        story.append(Paragraph(pdf_safe_text(f"Technikai státusz: {FPI_IMPORT_ENGINE_VERSION} | Smart Mapper: aktív | Week Rescue: aktív | Keeper Logic: aktív | Minutes Normalization: aktív | Microcycle Engine: aktív"), small))


    add_cover()
    # V8.3 riportlogika:
    # - executive: csak a tiszta vezetői döntési oldal
    # - full: ugyanazzal kezd, majd részletes GPS, mikrociklus, taktikai és metodikai appendix
    if report_type == "executive":
        add_executive_page()
    elif report_type == "full":
        add_executive_page()
        story.append(PageBreak()); add_fitness_page()
        story.append(PageBreak()); add_micro_page()
        # V12.7: a Full Reportból kivesszük az integrált taktikai összegzés oldalt.
        # A vezetői oldal már tartalmazza a legfontosabb taktikai/GPS döntési információkat,
        # így a riport rövidebb, kevésbé ismétlődő és könnyebben átadható.
    elif report_type == "fitness":
        add_fitness_page()
    elif report_type == "microcycle":
        add_micro_page()
    else:
        add_executive_page()

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def _build_demo_tactical_context() -> Dict[str, object]:
    return {
        "version": "DEMO_TACTICAL_CONTEXT_V1",
        "analysis_level": "Full Intelligence DEMO",
        "has_own_pdf": True,
        "has_opp_pdf": True,
        "has_own_team_excel": True,
        "has_opp_team_excel": True,
        "has_own_player_excel": True,
        "has_opp_player_excel": True,
        "own_topics": [
            {"Téma": "Labdakihozatal / támadásépítés", "Bizonyosság": 88},
            {"Téma": "Presszing / letámadás", "Bizonyosság": 82},
        ],
        "opp_topics": [
            {"Téma": "Támadó átmenet / kontrák", "Bizonyosság": 91},
            {"Téma": "Szélső játék / oldali dominancia", "Bizonyosság": 78},
            {"Téma": "Pontrúgások", "Bizonyosság": 72},
        ],
        "own_team_metrics": {"possession_pct": 39, "shots": 18, "xg": 0.75, "entries_box": 9, "final_third_entries": 62, "key_passes": 3, "pressing_success_pct": 71, "passes_accurate_pct": 74, "recoveries": 72, "lost_balls": 48, "crosses": 24, "corners": 2},
        "opp_team_metrics": {"counterattacks": 8, "crosses": 21, "corners": 6, "shots": 10, "xg": 1.20, "key_passes": 7, "ppda": 10.5},
        "plan_a": "BAT – középső blokk + gyors átmenet, jobb oldali biztosítással",
        "risks": [
            "Ellenfél-kontrák / gyors átmenetek kezelése",
            "Szélső játék és beadások elleni védekezés",
            "Pontrúgás-védekezés és második labdák",
        ],
        "md_plan": [
            ("MD-4", "Volumen + saját játékmodell", "Saját build-up és presszing trigger ismétlése."),
            ("MD-3", "HSR / sprint exponálás + átmenetek", "Ellenfél kontraveszély miatt átmeneti futások kontrollált terheléssel."),
            ("MD-2", "Kontrák elleni biztosítás + rest defense", "Ellenfél gyors átmeneti profilja miatt."),
            ("MD-1", "Aktiváció + pontrúgás", "Pontrúgás-veszély és frissesség kezelése."),
        ],
        "player_focus": ["Jobb oldali védő: beadások elleni 1v1 fókusz", "6-os: rest defense pozíció", "9-es: second ball célpont"],
        "tactical_findings": [
            {
                "Téma": "Ellenfél kontraveszély",
                "Bizonyíték": "PDF: támadó átmenet / kontrák; Excel: ellenfél kontra és beadásprofil.",
                "Edzői következtetés": "MD-3 átmeneti játék + HSR, MD-2 rest defense biztosítás.",
                "Prioritás": "Magas",
            },
            {
                "Téma": "Szélső játék és beadások",
                "Bizonyíték": "PDF: szélső játék / oldali dominancia; Excel: magas beadásszám.",
                "Edzői következtetés": "Oldali 1v1 védekezés, beadásblokkolás és hosszú oldali zárás.",
                "Prioritás": "Magas",
            },
            {
                "Téma": "Saját presszing-opció",
                "Bizonyíték": "Saját csapat PDF: presszing; GPS readiness közepes-jó.",
                "Edzői következtetés": "Csak kontrollált szakaszokban érdemes magasabb nyomást alkalmazni.",
                "Prioritás": "Közepes",
            },
        ],
        "team_comparison": [
            {"Mutató": "Labdabirtoklás", "Saját": 54.0, "Ellenfél": 46.0, "Értelmezés": "Saját kontroll"},
            {"Mutató": "Kontrák", "Saját": 4.0, "Ellenfél": 8.0, "Értelmezés": "Ellenfél átmeneti veszély"},
            {"Mutató": "Beadások", "Saját": 14.0, "Ellenfél": 21.0, "Értelmezés": "Ellenfél szélső veszély"},
        ],
        "opponent_player_evaluation": [
            {"Játékos":"András Simon","Poszt":"CF","Szerep":"Befejező / kapura veszélyes játékos","Bizonyíték":"lövés: 4.0; xG: 0.6","Értelmezés":"aktív befejező, kapura veszélyes","Javaslat":"Boxon belüli felvétel korlátozása, második labdák kontrollja."},
            {"Játékos":"Bence Szabó","Poszt":"RM","Szerep":"Szélső / beadó veszély","Bizonyíték":"beadás: 8.0; kulcspassz: 3.0","Értelmezés":"oldali veszély / beadások jelentős száma","Javaslat":"Oldali 1v1 kontroll, beadásblokkolás, hosszú oldali zárás."},
            {"Játékos":"Zsolt János Magyar","Poszt":"AM","Szerep":"Kreatív kulcsjátékos","Bizonyíték":"kulcspassz: 4.0; progresszív passz: 7.0","Értelmezés":"jó kreatív kapcsolódási pont, több kulcspasszal","Javaslat":"Passzsáv zárása, testhelyzet kontroll, belső irány lezárása."},
        ],
    }




# =========================================================
# V9.7 - Edzői kontextus: korosztály/szint referencia + edzésnapok
# =========================================================

FPI_REFERENCE_AGE_OPTIONS_V112 = ["Felnőtt", "U21", "U19", "U17", "U16", "U15", "U14", "U13"]
FPI_REFERENCE_LEVEL_OPTIONS_V112 = ["NB I", "NB II", "NB III", "Akadémia", "Regionális", "Megye I", "Egyéb"]
FPI_REFERENCE_POSITION_OPTIONS_V112 = ["Kapus", "Középhátvéd", "Szélső hátvéd", "Védekező középpályás", "Középpályás", "Támadó középpályás", "Szélső", "Csatár"]
FPI_COACH_WEEK_OPTIONS_V112 = [
    "Regenerációs mikrociklus",
    "Stabilizáló mikrociklus",
    "Terhelésfokozó mikrociklus",
    "Fejlesztő mikrociklus",
    "Formaidőzítő mikrociklus",
    "Mérkőzésre felkészítő mikrociklus",
]
FPI_COACH_WEEK_LABEL_MAP_V121 = {
    "Regeneráló hét": "Regenerációs mikrociklus",
    "Fenntartó hét": "Stabilizáló mikrociklus",
    "Terhelő hét": "Terhelésfokozó mikrociklus",
    "Fejlesztő hét": "Fejlesztő mikrociklus",
    "Élező hét": "Formaidőzítő mikrociklus",
    "Meccsre felkészítő hét": "Mérkőzésre felkészítő mikrociklus",
    "Meccsre frissítő hét": "Formaidőzítő mikrociklus",
}

def _fpi_normalize_coach_week_label_v121(label: object) -> str:
    s = str(label or "").strip()
    return FPI_COACH_WEEK_LABEL_MAP_V121.get(s, s or "Edző által nem megadva")
FPI_PLAYMODEL_OPTIONS_V112 = ["Dominancia", "Magas presszing", "Átmeneti játék", "Direkt játék", "Kiegyensúlyozott"]

# V11.2 - Referencia motor V2
# A százalékos tartományok azt mutatják, hogy a heti edzésösszeg / edzésátlag hány %-a a meccsreferenciának.
FPI_REFERENCE_BASE_RANGES_V112 = {
    "total_distance": ("280–420%", "60–100%", 280, 420, 60, 100),
    "training_load": ("260–410%", "55–95%", 260, 410, 55, 95),
    "hsr_distance": ("150–250%", "35–70%", 150, 250, 35, 70),
    "sprint_distance": ("100–200%", "25–55%", 100, 200, 25, 55),
    "sprints": ("100–220%", "25–60%", 100, 220, 25, 60),
    "high_efforts": ("150–280%", "35–75%", 150, 280, 35, 75),
}

FPI_AGE_FACTOR_V112 = {
    "Felnőtt": 1.00, "U21": 0.97, "U19": 0.93, "U17": 0.86,
    "U16": 0.80, "U15": 0.73, "U14": 0.66, "U13": 0.58,
}
FPI_LEVEL_FACTOR_V112 = {
    "NB I": 1.08, "NB II": 1.00, "NB III": 0.92, "Akadémia": 0.94,
    "Regionális": 0.84, "Megye I": 0.78, "Egyéb": 0.82,
}
FPI_POSITION_FACTORS_V112 = {
    "Kapus": {"total_distance": .62, "training_load": .65, "hsr_distance": .35, "sprint_distance": .25, "sprints": .35, "high_efforts": .55},
    "Középhátvéd": {"total_distance": .92, "training_load": .92, "hsr_distance": .78, "sprint_distance": .70, "sprints": .75, "high_efforts": .82},
    "Szélső hátvéd": {"total_distance": 1.05, "training_load": 1.04, "hsr_distance": 1.16, "sprint_distance": 1.12, "sprints": 1.10, "high_efforts": 1.08},
    "Védekező középpályás": {"total_distance": 1.08, "training_load": 1.06, "hsr_distance": .96, "sprint_distance": .82, "sprints": .88, "high_efforts": .98},
    "Középpályás": {"total_distance": 1.10, "training_load": 1.08, "hsr_distance": 1.02, "sprint_distance": .92, "sprints": .96, "high_efforts": 1.02},
    "Támadó középpályás": {"total_distance": 1.00, "training_load": 1.00, "hsr_distance": 1.04, "sprint_distance": 1.02, "sprints": 1.00, "high_efforts": 1.04},
    "Szélső": {"total_distance": 1.02, "training_load": 1.02, "hsr_distance": 1.22, "sprint_distance": 1.25, "sprints": 1.22, "high_efforts": 1.12},
    "Csatár": {"total_distance": .96, "training_load": .98, "hsr_distance": 1.08, "sprint_distance": 1.18, "sprints": 1.15, "high_efforts": 1.05},
}
FPI_PLAYMODEL_FACTORS_V112 = {
    "Dominancia": {"total_distance": 1.03, "training_load": .98, "hsr_distance": .92, "sprint_distance": .88, "sprints": .90, "high_efforts": .94},
    "Magas presszing": {"total_distance": 1.04, "training_load": 1.08, "hsr_distance": 1.10, "sprint_distance": 1.08, "sprints": 1.12, "high_efforts": 1.22},
    "Átmeneti játék": {"total_distance": 1.00, "training_load": 1.03, "hsr_distance": 1.18, "sprint_distance": 1.22, "sprints": 1.18, "high_efforts": 1.10},
    "Direkt játék": {"total_distance": .98, "training_load": 1.02, "hsr_distance": 1.16, "sprint_distance": 1.18, "sprints": 1.12, "high_efforts": 1.02},
    "Kiegyensúlyozott": {"total_distance": 1.00, "training_load": 1.00, "hsr_distance": 1.00, "sprint_distance": 1.00, "sprints": 1.00, "high_efforts": 1.00},
    "Pressing": {"total_distance": 1.04, "training_load": 1.08, "hsr_distance": 1.10, "sprint_distance": 1.08, "sprints": 1.12, "high_efforts": 1.22},
    "Transition": {"total_distance": 1.00, "training_load": 1.03, "hsr_distance": 1.18, "sprint_distance": 1.22, "sprints": 1.18, "high_efforts": 1.10},
    "Possession": {"total_distance": 1.03, "training_load": .98, "hsr_distance": .92, "sprint_distance": .88, "sprints": .90, "high_efforts": .94},
    "Low Block": {"total_distance": .92, "training_load": .94, "hsr_distance": .88, "sprint_distance": .92, "sprints": .92, "high_efforts": .96},
}


def _fpi_scale_range_v112(rng: Tuple[str, str, float, float, float, float], factor: float) -> Tuple[str, str, float, float, float, float]:
    _, _, low, high, avg_low, avg_high = rng
    low2, high2 = int(round(low * factor)), int(round(high * factor))
    avg_low2, avg_high2 = int(round(avg_low * factor)), int(round(avg_high * factor))
    return (f"{low2}–{high2}%", f"{avg_low2}–{avg_high2}%", low2, high2, avg_low2, avg_high2)


def _fpi_build_reference_profile_v112(age: str, level: str, position: str, playmodel: str) -> Dict[str, object]:
    age = age if age in FPI_AGE_FACTOR_V112 else "Felnőtt"
    level = level if level in FPI_LEVEL_FACTOR_V112 else "NB II"
    position = position if position in FPI_POSITION_FACTORS_V112 else "Középpályás"
    playmodel = playmodel if playmodel in FPI_PLAYMODEL_FACTORS_V112 else "Kiegyensúlyozott"
    ranges = {}
    for metric, base in FPI_REFERENCE_BASE_RANGES_V112.items():
        factor = FPI_AGE_FACTOR_V112[age] * FPI_LEVEL_FACTOR_V112[level]
        factor *= FPI_POSITION_FACTORS_V112[position].get(metric, 1.0)
        factor *= FPI_PLAYMODEL_FACTORS_V112[playmodel].get(metric, 1.0)
        factor = max(0.25, min(1.45, factor))
        ranges[metric] = _fpi_scale_range_v112(base, factor)
    label = f"{age} / {level} / {position} / {playmodel}"
    return {"label": label, "age": age, "level": level, "position": position, "playmodel": playmodel, "ranges": ranges}


def _fpi_position_to_reference_position_v116(value: object, is_goalkeeper: bool = False) -> str:
    """Nyers poszt / position_group -> FPI referencia poszt.
    Ha van kapus jelölés, mindig Kapus. Ha nincs poszt, mezőnyátlagként középpályás fallbacket használunk,
    de a riportban jelezzük, hogy játékosposzt-alapú súlyozás csak posztoszlop esetén pontos.
    """
    if is_goalkeeper:
        return "Kapus"
    txt = _norm_mapping_text(value)
    if not txt:
        return "Középpályás"
    if any(x in txt for x in ["kapus", "goalkeeper", "keeper", "gk"]):
        return "Kapus"
    if any(x in txt for x in ["szelso hatved", "fullback", "wing back", "wingback", "fb", "rb", "lb"]):
        return "Szélső hátvéd"
    if any(x in txt for x in ["kozep hatved", "kozephatved", "centre back", "center back", "central defender", "cb"]):
        return "Középhátvéd"
    if any(x in txt for x in ["vedekezo kozeppalyas", "defensive mid", "dm", "cdm", "six", "6"]):
        return "Védekező középpályás"
    if any(x in txt for x in ["tamado kozeppalyas", "attacking mid", "am", "cam", "10"]):
        return "Támadó középpályás"
    if any(x in txt for x in ["szelso", "winger", "wide", "lw", "rw"]):
        return "Szélső"
    if any(x in txt for x in ["csatar", "striker", "forward", "fw", "st", "9"]):
        return "Csatár"
    if any(x in txt for x in ["kozeppalyas", "midfield", "cm", "8"]):
        return "Középpályás"
    # position_group régi címkék
    if "ved" in txt and "kozep" in txt:
        return "Középhátvéd"
    if "kozeppalya" in txt:
        return "Középpályás"
    return "Középpályás"


def _fpi_reference_profile_for_player_v116(row: pd.Series, age: str, level: str, playmodel: str) -> Dict[str, object]:
    pos_source = row.get("position", row.get("position_group", ""))
    ref_pos = _fpi_position_to_reference_position_v116(pos_source, bool(row.get("is_goalkeeper", False)))
    return _fpi_build_reference_profile_v112(age, level, ref_pos, playmodel)


def _fpi_composition_reference_ranges_v116(df: pd.DataFrame, week: str, metric: str) -> Tuple[str, str, float, float, float, float, str]:
    """Csapat referencia játékosösszetétel alapján.
    A korosztály + bajnoki szint + játékmodell globális, a poszt játékosonként jön a Poszt oszlopból.
    Kapusok kisebb súlyt kapnak a csapatszintű sebesség/HSR/sprint benchmarkban.
    """
    ctx = _fpi_get_coach_context_v97()
    age = str(ctx.get("reference_age") or "Felnőtt")
    level = str(ctx.get("reference_level") or "NB II")
    playmodel = str(ctx.get("playmodel_profile") or ctx.get("selected_playstyle") or "Kiegyensúlyozott")
    data = df.copy() if isinstance(df, pd.DataFrame) else pd.DataFrame()
    if not data.empty and "week" in data.columns and week:
        data = data[data["week"].astype(str) == str(week)]
    if data.empty:
        prof = _fpi_build_reference_profile_v112(age, level, "Középpályás", playmodel)
        a,b,c,d,e,f = prof["ranges"].get(metric, FPI_REFERENCE_BASE_RANGES_V112.get(metric, ("n.a.", "n.a.", 0, 9999, 0, 9999)))
        return a,b,c,d,e,f, f"{age} / {level} / játékosposzt nincs adat / {playmodel}"
    if "player_name" in data.columns:
        # egy játékos egyszer számítson, a leggyakoribb vagy utolsó posztjával
        sort_cols = [c for c in ["player_name", "position", "position_group", "is_goalkeeper"] if c in data.columns]
        players_df = data[sort_cols].drop_duplicates(subset=["player_name"], keep="last") if sort_cols else data.head(0)
    else:
        players_df = data.copy()
    if players_df.empty:
        prof = _fpi_build_reference_profile_v112(age, level, "Középpályás", playmodel)
        a,b,c,d,e,f = prof["ranges"].get(metric, FPI_REFERENCE_BASE_RANGES_V112.get(metric, ("n.a.", "n.a.", 0, 9999, 0, 9999)))
        return a,b,c,d,e,f, f"{age} / {level} / mezőnyátlag / {playmodel}"
    vals = []
    labels = []
    for _, row in players_df.iterrows():
        pos = _fpi_position_to_reference_position_v116(row.get("position", row.get("position_group", "")), bool(row.get("is_goalkeeper", False)))
        prof = _fpi_build_reference_profile_v112(age, level, pos, playmodel)
        rng = prof["ranges"].get(metric, FPI_REFERENCE_BASE_RANGES_V112.get(metric, ("n.a.", "n.a.", 0, 9999, 0, 9999)))
        # Kapus kisebb súllyal a csapatszintű benchmarkban, de nem tűnik el.
        w = 0.35 if pos == "Kapus" and metric in ["hsr_distance", "sprint_distance", "sprints", "total_distance"] else 1.0
        vals.append((rng[2], rng[3], rng[4], rng[5], w))
        labels.append(pos)
    sw = sum(v[4] for v in vals) or 1.0
    low = int(round(sum(v[0]*v[4] for v in vals)/sw)); high = int(round(sum(v[1]*v[4] for v in vals)/sw))
    avg_low = int(round(sum(v[2]*v[4] for v in vals)/sw)); avg_high = int(round(sum(v[3]*v[4] for v in vals)/sw))
    from collections import Counter
    comp = ", ".join(f"{k}:{v}" for k,v in Counter(labels).most_common(4))
    return f"{low}–{high}%", f"{avg_low}–{avg_high}%", low, high, avg_low, avg_high, f"{age} / {level} / játékosposzt-súlyozott ({comp}) / {playmodel}"

# V9.7 régi profilok megtartva kompatibilitás miatt, de a V2 motor már a fenti négy komponensből épít profilt.
FPI_REFERENCE_PROFILES_V97 = {
    "Felnőtt NB2": _fpi_build_reference_profile_v112("Felnőtt", "NB II", "Középpályás", "Kiegyensúlyozott"),
    "Felnőtt NB3 / megyei felnőtt": _fpi_build_reference_profile_v112("Felnőtt", "NB III", "Középpályás", "Kiegyensúlyozott"),
    "Akadémia U19": _fpi_build_reference_profile_v112("U19", "Akadémia", "Középpályás", "Kiegyensúlyozott"),
    "Akadémia U17": _fpi_build_reference_profile_v112("U17", "Akadémia", "Középpályás", "Kiegyensúlyozott"),
    "Megyei / grassroots U13-U15": _fpi_build_reference_profile_v112("U14", "Megye I", "Középpályás", "Kiegyensúlyozott"),
}

def _fpi_fmt_thousands_v97(x: object) -> str:
    try:
        return f"{float(x):,.0f}".replace(",", " ")
    except Exception:
        return "0"

def _fpi_get_coach_context_v97() -> Dict[str, object]:
    return st.session_state.get("fpi_coach_context_v97", {}) or {}

def _fpi_reference_profile_v97() -> Dict[str, object]:
    ctx = _fpi_get_coach_context_v97()
    if any(k in ctx for k in ["reference_age", "reference_level", "reference_position", "playmodel_profile"]):
        return _fpi_build_reference_profile_v112(
            str(ctx.get("reference_age") or "Felnőtt"),
            str(ctx.get("reference_level") or "NB II"),
            str(ctx.get("reference_position") or "Középpályás"),
            str(ctx.get("playmodel_profile") or ctx.get("selected_playstyle") or "Kiegyensúlyozott"),
        )
    key = ctx.get("reference_profile", "Felnőtt NB2")
    return FPI_REFERENCE_PROFILES_V97.get(key, FPI_REFERENCE_PROFILES_V97["Felnőtt NB2"])

def _fpi_session_plan_v97() -> List[Dict[str, object]]:
    return _fpi_get_coach_context_v97().get("session_plan", []) or []

def _fpi_periodization_label_v97(ctx: Dict[str, object]) -> str:
    return _fpi_normalize_coach_week_label_v121(_fpi_get_coach_context_v97().get("coach_week_type") or "Edző által nem megadva")

def _fpi_reference_ranges_for_metric_v97(metric: str) -> Tuple[str, str, float, float, float, float]:
    prof = _fpi_reference_profile_v97()
    return prof.get("ranges", {}).get(metric, FPI_REFERENCE_BASE_RANGES_V112.get(metric, ("n.a.", "n.a.", 0, 9999, 0, 9999)))

def _fpi_ratio_status_v97(value: Optional[float], low: float, high: float) -> str:
    return _fpi_ratio_status_v93(value, low, high)

def _fpi_match_ratio_reference_df_v97(df: pd.DataFrame, week: str) -> pd.DataFrame:
    # V11.6: a korosztály/szint/játékmodell globális, a referencia poszt játékosonként jön a Poszt oszlopból.
    base = _fpi_match_ratio_reference_df_v93(df, week)
    if base is None or base.empty:
        return base
    metric_by_label = {v["label"]: k for k, v in FPI_NB2_ADULT_REFERENCE_RANGES_V93.items()}
    metric_by_label.update({"Load": "training_load", "Terhelési pont": "training_load"})
    rows = []
    for _, r in base.iterrows():
        rr = r.to_dict()
        metric = metric_by_label.get(str(r.get("Mutató", "")))
        if metric:
            weekly_ref, avg_ref, low, high, avg_low, avg_high, ref_label = _fpi_composition_reference_ranges_v116(df, week, metric)
            weekly_pct = r.get("Edzés/heti meccs %")
            avg_pct = r.get("Edzésátlag/meccs %")
            rr["Profil heti ref."] = weekly_ref
            rr["Profil edzésátlag ref."] = avg_ref
            rr["NB2 felnőtt heti ref."] = weekly_ref
            rr["NB2 felnőtt edzésátlag ref."] = avg_ref
            rr["Referencia profil"] = ref_label
            status_w = _fpi_ratio_status_v97(weekly_pct, low, high)
            status_a = _fpi_ratio_status_v97(avg_pct, avg_low, avg_high)
            if weekly_pct is None or pd.isna(weekly_pct):
                rr["Értékelés"] = "Nincs meccs referencia vagy nincs értelmezhető adat."
            elif status_w == "alacsony" and metric in ["hsr_distance", "sprint_distance", "sprints"]:
                rr["Értékelés"] = "A heti sebességi/sprint inger a referenciazóna alatt van. Javasolt kontrollált sebességi expozíció."
            elif status_w == "magas":
                rr["Értékelés"] = "A heti összterhelés a referenciazóna felett van. Érdemes ellenőrizni a napokra bontott terhelést és az egyéni kockázatot."
            elif status_w == "célzónában" and status_a == "célzónában":
                rr["Értékelés"] = "A heti összeg és az edzésátlag is a referenciazónában van. A terhelési szerkezet fenntartható."
            else:
                rr["Értékelés"] = f"Heti érték: {status_w}. Edzésátlag: {status_a}."
        rows.append(rr)
    return pd.DataFrame(rows)

def _fpi_gps_only_md_plan_v97(ctx: Dict[str, object], readiness: int, priorities: List[dict], week: str) -> List[Tuple[str, str, str]]:
    coach_plan = _fpi_session_plan_v97()
    if coach_plan:
        base_focus = _fpi_gps_only_conclusions_v95(ctx, priorities, readiness, week, limit=4)
        rows = []
        for item in coach_plan:
            md = item.get("md", "MD?")
            kind = item.get("type", "Edzés")
            note = item.get("note", "")
            if str(kind).lower().startswith("pihen"):
                rows.append((md, "Pihenő / regeneráció", note or "Tervezett pihenőnap, readiness és frissesség megtartása."))
            elif str(kind).lower().startswith("meccs"):
                rows.append((md, "Meccsnap", note or "Mérkőzés: a heti terhelés célja erre a napra friss állapotot biztosítani."))
            else:
                focus = base_focus[min(len(rows), max(len(base_focus)-1, 0))] if base_focus else "GPS-alapú célzott edzés."
                rows.append((md, str(kind), note or focus))
        return rows[:7]
    return _fpi_gps_only_md_plan_v95(ctx, readiness, priorities, week)



# =========================================================
# V9.9 - GPS-only Mikrociklus AI Planner v2 + KPI doboz szövegfix
# =========================================================

def _fpi_short_week_type_v99(label: object) -> str:
    s = _fpi_normalize_coach_week_label_v121(label)
    mapping = {
        "Edző által nem megadva": "Nincs megadva",
        "Regenerációs mikrociklus": "Regenerációs",
        "Stabilizáló mikrociklus": "Stabilizáló",
        "Terhelésfokozó mikrociklus": "Terhelésfokozó",
        "Fejlesztő mikrociklus": "Fejlesztő",
        "Formaidőzítő mikrociklus": "Formaidőzítő",
        "Mérkőzésre felkészítő mikrociklus": "Meccsre készítő",
        "Tanuló hét": "Tanuló",
        "Vegyes hét": "Vegyes",
    }
    return mapping.get(s, s[:18] if len(s) > 18 else s)

def _fpi_week_sort_key_v99(w: object) -> Tuple[int, int, str]:
    s = str(w)
    m = re.search(r"(\d{4})-W(\d{1,2})", s)
    if m:
        return (int(m.group(1)), int(m.group(2)), s)
    return (0, 0, s)

def _fpi_previous_weeks_v99(df: pd.DataFrame, week: str, n: int = 4) -> List[str]:
    if df is None or df.empty or "week" not in df.columns:
        return []
    weeks = sorted([str(x) for x in df["week"].dropna().unique()], key=_fpi_week_sort_key_v99)
    if str(week) not in weeks:
        return weeks[-n:]
    idx = weeks.index(str(week))
    return weeks[max(0, idx-n):idx]

def _fpi_week_team_totals_v99(df: pd.DataFrame, weeks: List[str]) -> pd.DataFrame:
    if df is None or df.empty or "week" not in df.columns:
        return pd.DataFrame()
    d = df[df["week"].astype(str).isin([str(w) for w in weeks])].copy()
    if d.empty:
        return pd.DataFrame()
    metrics = [c for c in ["total_distance", "hsr_distance", "sprint_distance", "sprints", "high_efforts", "training_load"] if c in d.columns]
    if not metrics:
        return pd.DataFrame()
    for c in metrics:
        d[c] = pd.to_numeric(d[c], errors="coerce").fillna(0)
    out = d.groupby("week", as_index=False)[metrics].sum()
    out["_sort"] = out["week"].apply(_fpi_week_sort_key_v99)
    out = out.sort_values("_sort").drop(columns=["_sort"])
    return out

def _fpi_trend_label_v99(values: List[float]) -> str:
    vals = [float(v or 0) for v in values if v is not None]
    if len(vals) < 2:
        return "nincs trend"
    first = vals[0]
    last = vals[-1]
    if first <= 0 and last <= 0:
        return "nincs adat"
    pct = (last - first) / max(abs(first), 1) * 100
    if pct > 20:
        return "emelkedő"
    if pct < -20:
        return "csökkenő"
    return "stabil"

def _fpi_gps_trend_summary_v99(ctx: Dict[str, object], week: str) -> Dict[str, object]:
    df = ctx.get("df")
    if not isinstance(df, pd.DataFrame) or df.empty:
        return {"weeks": [], "signals": [], "totals": pd.DataFrame()}
    prev = _fpi_previous_weeks_v99(df, week, n=4)
    all_weeks = prev + [str(week)]
    totals = _fpi_week_team_totals_v99(df, all_weeks)
    signals = []
    if not totals.empty:
        for metric, label in [
            ("total_distance", "volumen"),
            ("hsr_distance", "HSR"),
            ("sprint_distance", "sprint táv"),
            ("sprints", "sprint count"),
            ("high_efforts", "High Efforts"),
            ("training_load", "Load"),
        ]:
            if metric in totals.columns:
                tr = _fpi_trend_label_v99(totals[metric].tolist())
                if tr in ["emelkedő", "csökkenő"]:
                    signals.append(f"{label}: {tr} trend az utolsó {len(totals)} hétben")
    return {"weeks": all_weeks, "signals": signals[:6], "totals": totals}


def _fpi_clean_profile_noise_v121(text: object) -> str:
    """A PDF mikrociklus szövegeiből kiszedi a gépi referencia-profil zajt."""
    s = str(text or "").strip()
    s = re.sub(r"\s*Profil:\s*.*?(?=$|\.)", "", s, flags=re.IGNORECASE)
    s = re.sub(r"\s+játékosposzt-súlyozott profilhoz képest", "", s, flags=re.IGNORECASE)
    s = re.sub(r"\s+játékosposzt-súlyozott referenciazónában", " referenciazónában", s, flags=re.IGNORECASE)
    s = re.sub(r"\s+", " ", s).strip(" .")
    return s + "." if s else ""


def _fpi_focus_from_kind_v121(md: object, kind: object, train_i: int, coach_week: str) -> str:
    k = _norm_mapping_text(kind)
    md_s = str(md or "")
    if "pihen" in k:
        return "Regeneráció"
    if "meccs" in k:
        return "Mérkőzés"
    if "aktiv" in k or md_s == "MD-1":
        return "Aktiválás"
    if "regener" in k:
        return "Regenerációs kontroll"
    if "hsr" in k or "sebesseg" in k:
        return "Sebességi terhelés"
    if "sprint" in k:
        return "Sprintinger"
    if "terhel" in _norm_mapping_text(coach_week) and train_i == 0:
        return "Terhelésépítés"
    if "formaidozito" in _norm_mapping_text(coach_week) or "merkozesre" in _norm_mapping_text(coach_week):
        return "Mérkőzésre hangolás" if train_i == 0 else "Minőségi intenzitás"
    if "stabilizalo" in _norm_mapping_text(coach_week):
        return "Terhelés stabilizálása" if train_i == 0 else "Kontrollált intenzitás"
    if train_i == 0:
        return "Terhelésépítés"
    if train_i == 1:
        return "Sebességi terhelés"
    if train_i == 2:
        return "Mérkőzésspecifikus terhelés"
    return "Kontrollált edzésterhelés"


def _fpi_status_and_recommendation_v121(md: object, focus: object, why: object, readiness: int, trend_txt: str) -> Tuple[str, str]:
    raw = _fpi_clean_profile_noise_v121(why)
    f = _norm_mapping_text(focus)
    status = raw
    rec = "A terhelést a játékosok egyéni állapotához és a heti célhoz igazítva érdemes adagolni."
    if not status:
        status = f"Readiness: {readiness}/100. Trend: {trend_txt}."
    if "piheno" in f or "regener" in f:
        status = "Tervezett regenerációs nap. Cél a frissesség visszaépítése és a neuromuszkuláris terhelés csökkentése."
        rec = "Alacsony volumen, mobilitás, regenerációs munka és egyéni kezelés javasolt."
    elif "aktiv" in f or str(md) == "MD-1":
        status = f"Readiness: {readiness}/100. A cél a frissesség megtartása, nem új terhelés felépítése."
        rec = "Rövid aktiváló edzés: reakció, gyors lábmunka, pontrúgás és alacsony összvolumen."
    elif "sprint" in f:
        rec = "Rövid, kontrollált sprint-expozíció javasolt hosszú pihenőkkel; kerülni kell a felesleges ismétlésszámot."
    elif "sebesseg" in f or "hsr" in f:
        rec = "Nagyobb területű játék vagy célzott sebességi blokk javasolt, kontrollált mennyiséggel."
    elif "terhelesepites" in f or "volumen" in _norm_mapping_text(raw):
        rec = "A volumen fokozatos emelése javasolt, de a magas kockázatú játékosoknál egyéni korrekcióval."
    elif "merkoz" in f:
        rec = "A gyakorlatok intenzitása legyen mérkőzésszerű, de az összterhelés maradjon kontrollált."
    elif "minosegi" in f:
        rec = "Kevés, jó minőségű intenzív inger javasolt; a fárasztó mennyiségi munka kerülendő."
    if "alacsony" in _norm_mapping_text(raw) and ("hsr" in _norm_mapping_text(raw) or "sprint" in _norm_mapping_text(raw)):
        rec = "Sebességi/sprint inger beépítése javasolt, de csak kontrollált adagban és megfelelő pihenőkkel."
    if "magas" in _norm_mapping_text(raw) or "kockazat" in _norm_mapping_text(raw):
        rec = "Csökkentett volumen vagy egyéni módosítás javasolt a veszélyeztetett játékosoknál."
    return (_fpi_clean_sentence_v82(status, 170), _fpi_clean_sentence_v82(rec, 170))


def _fpi_structured_md_rows_v121(rows: List[Tuple], readiness: int, trend_txt: str) -> List[Tuple[str, str, str, str]]:
    out = []
    for row in rows:
        if len(row) >= 4:
            md, focus, status, rec = row[:4]
        else:
            md, focus, why = row[:3]
            status, rec = _fpi_status_and_recommendation_v121(md, focus, why, readiness, trend_txt)
        out.append((str(md), str(focus), str(status), str(rec)))
    return out

def _fpi_gps_only_md_plan_v99(ctx: Dict[str, object], readiness: int, priorities: List[dict], week: str) -> List[Tuple[str, str, str, str]]:
    """GPS-only Mikrociklus AI Planner v2.
    Figyelembe veszi:
    - edzői hét típust
    - edző által megadott edzésszámot / MD napokat
    - aktuális heti referenciaarányokat
    - előző 4 hét trendjét
    - readiness és játékosrisk jelzéseket
    """
    coach_week = _fpi_normalize_coach_week_label_v121(_fpi_get_coach_context_v97().get("coach_week_type") or "Edző által nem megadva")
    coach_plan = _fpi_session_plan_v97()
    trend = _fpi_gps_trend_summary_v99(ctx, week)
    trend_txt = "; ".join(trend.get("signals", [])[:2]) if trend.get("signals") else "nincs erős 4 hetes trend"
    base = _fpi_gps_only_md_plan_v95(ctx, readiness, priorities, week)

    # hét típus szerinti globális irány
    week_low = coach_week.lower()
    if "regener" in week_low:
        default_focus = ("Regenerációs kontroll", f"Alacsonyabb volumen, kontrollált intenzitás. Trend: {trend_txt}.")
    elif "fenntart" in week_low:
        default_focus = ("Terhelés stabilizálása", f"A fő fizikai képességek fenntartása túlterhelés nélkül. Trend: {trend_txt}.")
    elif "terhel" in week_low:
        default_focus = ("Terhelésépítés", f"Volumen és load fokozatos építése a referenciazónákhoz igazítva. Trend: {trend_txt}.")
    elif "fejleszt" in week_low:
        default_focus = ("Célzott fejlesztő inger", f"Célzott HSR/sprint/High Effort inger játékoskockázat-kontrollal. Trend: {trend_txt}.")
    elif "elez" in week_low or "élez" in week_low:
        default_focus = ("Formaidőzítés", f"Rövid, minőségi intenzitás, alacsony fárasztás. Trend: {trend_txt}.")
    elif "felkeszit" in week_low or "felkész" in week_low or "friss" in week_low:
        default_focus = ("Mérkőzésre hangolás", f"Frissesség megtartása és célzott sebességinger. Trend: {trend_txt}.")
    else:
        default_focus = ("GPS-alapú fókusz", f"Readiness, referenciaarányok és trend alapján. Trend: {trend_txt}.")

    if coach_plan:
        # Adott MD-struktúrát megtartjuk, csak a tartalmat szabjuk GPS-adatokra.
        base_focus = _fpi_gps_only_conclusions_v95(ctx, priorities, readiness, week, limit=6)
        rows = []
        train_i = 0
        for item in coach_plan:
            md = item.get("md", "MD?")
            kind = item.get("type", "Edzés")
            note = item.get("note", "")
            k = str(kind).lower()
            if k.startswith("pihen"):
                rows.append((md, "Pihenő / regeneráció", note or f"Tervezett pihenőnap. {trend_txt}."))
            elif "regener" in k:
                rows.append((md, "Regeneráció", note or f"Visszarendezés, alacsony neuromuszkuláris terhelés. {trend_txt}."))
            elif "aktiv" in k:
                rows.append((md, "Aktiváció", note or "Rövid gyorsasági/reakció inger, fárasztás nélkül."))
            elif "meccs" in k:
                rows.append((md, "Meccsnap", note or "A hét célja erre a napra friss állapotot biztosítani."))
            else:
                # edzésnap: base plan + trend + week type
                if base_focus:
                    focus = base_focus[min(train_i, len(base_focus)-1)]
                elif train_i < len(base):
                    focus = base[train_i][2]
                else:
                    focus = default_focus[1]
                if note:
                    focus = note
                rows.append((md, default_focus[0] if train_i == 0 else str(kind), focus))
                train_i += 1
        return _fpi_structured_md_rows_v121(rows[:8], readiness, trend_txt)

    # Ha nincs megadott edzésstruktúra, akkor v95 terv, de hét típus / trend szerint átfogalmazva
    if "regener" in week_low and readiness < 70:
        return _fpi_structured_md_rows_v121([
            ("MD-4", "Regenerációs kontroll", f"Alacsonyabb volumen és kontrollált intenzitás. Trend: {trend_txt}."),
            ("MD-3", "Minőségi sebességinger", "Csak expozíció, nem mennyiségi sprintmunka."),
            ("MD-2", "Frissítés + egyéni kontroll", "Játékosmonitoring, HSR/sprint halmozás nélkül."),
            ("MD-1", "Aktiváció", "Rövid döntési és mozgásgyorsasági inger."),
        ], readiness, trend_txt)
    if "tanul" in week_low:
        return _fpi_structured_md_rows_v121([
            ("MD-4", "Közepes volumen + tanulási blokk", f"Technikai/taktikai tanulás fizikai túlterhelés nélkül. {trend_txt}."),
            ("MD-3", "Kontrollált HSR/sprint expozíció", "A sebességinger megmarad, de nem ez a fő terhelési cél."),
            ("MD-2", "Alacsony-közepes intenzitás", "Tanulási ismétlések, frissesség megtartása."),
            ("MD-1", "Aktiváció", "Rövid, tiszta, frissítő inger."),
        ], readiness, trend_txt)
    return _fpi_structured_md_rows_v121(base, readiness, trend_txt)

def _fpi_gps_only_conclusions_v99(ctx: Dict[str, object], priorities: List[dict], readiness: int, week: str, limit: int = 6) -> List[str]:
    out = _fpi_gps_only_conclusions_v95(ctx, priorities, readiness, week, limit=limit)
    trend = _fpi_gps_trend_summary_v99(ctx, week)
    if trend.get("signals"):
        out.append("4 hetes trend: " + "; ".join(trend.get("signals", [])[:3]) + ".")
    coach_week = _fpi_get_coach_context_v97().get("coach_week_type")
    if coach_week and coach_week != "Edző által nem megadva":
        out.append(f"Mikrociklus típusa: {_fpi_normalize_coach_week_label_v121(coach_week)}. A napi javaslatok ezt a szakmai keretet veszik alapul.")
    uniq, seen = [], set()
    for x in out:
        y = _fpi_clean_sentence_v82(x, 190)
        if y and y not in seen:
            uniq.append(y); seen.add(y)
    return uniq[:limit]


def build_fpi_gps_only_pdf_bytes(
    data: pd.DataFrame,
    selected_week: Optional[str] = None,
    playstyle: str = "Kiegyensúlyozott",
    demo_label: str = "",
) -> Optional[bytes]:
    """GPS-only PDF riport.
    Nincs taktikai blokk, nincs Tactical Pro+ kontextus. Csak GPS, readiness, terhelésarány,
    játékoskockázat, mikrociklus és használható erőnléti következtetések.
    """
    if SimpleDocTemplate is None:
        return None
    from reportlab.platypus import PageBreak

    ctx = _fpi_report_context(data, selected_week, playstyle)
    if ctx.get("error"):
        return None

    df = ctx["df"]
    week = ctx["selected_week"]
    readiness = int(ctx.get("readiness_score", 70) or 70)
    priorities = ctx.get("priorities", []) or []
    risk_df = ctx.get("risk_df") if isinstance(ctx.get("risk_df"), pd.DataFrame) else pd.DataFrame()
    weekly = ctx.get("weekly") if isinstance(ctx.get("weekly"), pd.DataFrame) else pd.DataFrame()
    player_week = ctx.get("player_week") if isinstance(ctx.get("player_week"), pd.DataFrame) else pd.DataFrame()
    periodization_type = _fpi_pdf_week_type_label_v126(ctx, demo_label)
    conclusions = _fpi_gps_only_conclusions_v99(ctx, priorities, readiness, str(week), limit=6)
    md_plan = _fpi_gps_only_md_plan_v99(ctx, readiness, priorities, str(week))
    ratio_df = _fpi_match_ratio_reference_df_v97(df, str(week))

    font_name, font_bold = _register_pdf_font()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=0.9*cm, leftMargin=0.9*cm, topMargin=0.7*cm, bottomMargin=0.7*cm)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("GPSOnlyTitle", parent=styles["Title"], fontName=font_bold, fontSize=20, leading=23, textColor=colors.HexColor("#0F172A"))
    sub = ParagraphStyle("GPSOnlySub", parent=styles["Normal"], fontName=font_name, fontSize=8.8, leading=11, textColor=colors.HexColor("#334155"))
    body = ParagraphStyle("GPSOnlyBody", parent=styles["Normal"], fontName=font_name, fontSize=8.0, leading=10.2, textColor=colors.HexColor("#111827"))
    small = ParagraphStyle("GPSOnlySmall", parent=styles["Normal"], fontName=font_name, fontSize=7.0, leading=8.6, textColor=colors.HexColor("#111827"))
    head = ParagraphStyle("GPSOnlyHead", parent=styles["Normal"], fontName=font_bold, fontSize=7.4, leading=8.8, textColor=colors.white)
    story = []

    def P(txt, style=body):
        return Paragraph(pdf_safe_text(txt), style)

    def section(txt, color="#DBEAFE"):
        return Table([[P(txt, ParagraphStyle("GPSSection"+str(len(story)), parent=body, fontName=font_bold, fontSize=11, leading=13, textColor=colors.HexColor("#0F172A")))]],
                     colWidths=[27.7*cm],
                     style=TableStyle([
                         ("BACKGROUND", (0,0), (-1,-1), colors.HexColor(color)),
                         ("BOX", (0,0), (-1,-1), 0.4, colors.HexColor("#CBD5E1")),
                         ("LEFTPADDING", (0,0), (-1,-1), 6),
                         ("RIGHTPADDING", (0,0), (-1,-1), 6),
                         ("TOPPADDING", (0,0), (-1,-1), 5),
                         ("BOTTOMPADDING", (0,0), (-1,-1), 5),
                     ]))

    def tbl(rows, widths, header_bg="#1E3A8A", row_bgs=None):
        t = Table(rows, colWidths=widths, repeatRows=1)
        style = [
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor(header_bg)),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), font_bold),
            ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING", (0,0), (-1,-1), 4),
            ("RIGHTPADDING", (0,0), (-1,-1), 4),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]
        if row_bgs:
            for i in range(1, len(rows)):
                style.append(("BACKGROUND", (0,i), (-1,i), row_bgs[(i-1) % len(row_bgs)]))
        t.setStyle(TableStyle(style))
        return t

    def kpi(label, value, note, color="#1E3A8A"):
        ps1 = ParagraphStyle("KPI1"+label, parent=body, fontName=font_bold, fontSize=8, leading=9, textColor=colors.white)
        ps2 = ParagraphStyle("KPI2"+label, parent=body, fontName=font_bold, fontSize=12.5, leading=14.5, textColor=colors.white)
        ps3 = ParagraphStyle("KPI3"+label, parent=body, fontName=font_name, fontSize=6.4, leading=7.6, textColor=colors.white)
        t = Table([[Paragraph(pdf_safe_text(label), ps1)], [Paragraph(pdf_safe_text(value), ps2)], [Paragraph(pdf_safe_text(note), ps3)]], colWidths=[5.45*cm], rowHeights=[0.50*cm, 0.72*cm, 0.48*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), colors.HexColor(color)),
            ("BOX", (0,0), (-1,-1), 0.4, colors.white),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
            ("LEFTPADDING", (0,0), (-1,-1), 5),
            ("RIGHTPADDING", (0,0), (-1,-1), 5),
            ("TOPPADDING", (0,0), (-1,-1), 4),
            ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ]))
        return t

    # Cover / executive GPS-only
    label_prefix = f"{demo_label} | " if demo_label else ""
    story.append(P("Football Performance Intelligence – GPS-only riport", title))
    story.append(P(f"{label_prefix}Hét: {format_week_label(str(week))} | Játékmodell: {playstyle} | Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}", sub))
    story.append(P(_fpi_pdf_match_context_line_v122(demo_label), sub))
    story.append(P(f"Referencia profil: {_fpi_reference_profile_v97().get('label', 'Felnőtt NB2')} | Heti típus: {periodization_type}", sub))
    story.append(P("Heti cél: " + _fpi_week_type_interpretation_v126(periodization_type), sub))
    story.append(Spacer(1, 0.25*cm))
    high_risk, med_risk = _fpi_count_risk_levels_v126(risk_df)
    story.append(Table([[
        kpi("READINESS", f"{readiness}/100", score_to_label(readiness), "#166534" if readiness >= 75 else "#1E3A8A" if readiness >= 60 else "#991B1B"),
        kpi("HETI TÍPUS", _fpi_short_week_type_v99(periodization_type), "mikrociklus profil", "#0F172A"),
        kpi("HIGH RISK", f"{high_risk} fő", "egyéni kontroll", "#7F1D1D" if high_risk else "#166534"),
        kpi("MEDIUM RISK", f"{med_risk} fő", "figyelendő", "#92400E" if med_risk else "#166534"),
        kpi("FORRÁS", "GPS only", "nincs taktikai input", "#0369A1"),
    ]], colWidths=[5.45*cm]*5))
    story.append(Spacer(1, 0.25*cm))

    story.append(section("1. GPS-only vezetői konklúziók", "#DBEAFE"))
    c_rows = [[P("#", head), P("Konklúzió", head)]]
    c_rows.append([P("1", small), P(f"Heti típus: {periodization_type}. {_fpi_week_type_interpretation_v126(periodization_type)}", small)])
    for i, c in enumerate(conclusions, 2):
        c_rows.append([P(str(i), small), P(c, small)])
    story.append(tbl(c_rows, [1.0*cm, 26.7*cm], header_bg="#1E3A8A", row_bgs=[colors.HexColor("#EFF6FF"), colors.white]))
    story.append(Spacer(1, 0.22*cm))

    story.append(section("2. Edzés–meccs arányok választott korosztály/szint referenciával", "#FEF3C7"))
    if ratio_df is None or ratio_df.empty:
        story.append(P("Nincs elég adat az edzés/meccs arányokhoz. Kell legalább egy edzés és egy meccs típusú esemény.", body))
    else:
        rr = [[P("Mutató", head), P("Heti edzés / meccs", head), P("Edzésátlag / meccs", head), P("Referencia", head), P("Értelmezés", head)]]
        for _, r in ratio_df.iterrows():
            rr.append([
                P(str(r.get("Mutató","")), small),
                P(_fpi_fmt_pct_v93(r.get("Edzés/heti meccs %")), small),
                P(_fpi_fmt_pct_v93(r.get("Edzésátlag/meccs %")), small),
                P(f"Heti: {r.get('NB2 felnőtt heti ref.','')}<br/>Edzésátlag: {r.get('NB2 felnőtt edzésátlag ref.','')}", small),
                P(str(r.get("Értékelés","")), small),
            ])
        story.append(tbl(rr, [4.0*cm, 4.0*cm, 4.0*cm, 5.6*cm, 10.1*cm], header_bg="#92400E", row_bgs=[colors.HexColor("#FFFBEB"), colors.white]))

    story.append(PageBreak())
    story.append(section("3. GPS-alapú mikrociklus terv", "#EDE9FE"))
    md_rows = [[P("Nap", head), P("Fókusz", head), P("Állapot", head), P("Edzői javaslat", head)]]
    for row in md_plan:
        if len(row) >= 4:
            d, fgoal, status, recommendation = row[:4]
        else:
            d, fgoal, why = row[:3]
            status, recommendation = _fpi_status_and_recommendation_v121(d, fgoal, why, readiness, "")
        md_rows.append([P(d, small), P(fgoal, small), P(status, small), P(recommendation, small)])
    story.append(tbl(md_rows, [2.5*cm, 6.0*cm, 9.2*cm, 10.0*cm], header_bg="#312E81", row_bgs=[colors.HexColor("#F5F3FF"), colors.white]))
    story.append(Spacer(1, 0.22*cm))

    trend_v99 = _fpi_gps_trend_summary_v99(ctx, str(week))
    if isinstance(trend_v99.get("totals"), pd.DataFrame) and not trend_v99.get("totals").empty:
        story.append(section("4. Előző hetek trendje – mikrociklus alapja", "#ECFDF5"))
        tr = [[P("Hét", head), P("Volumen", head), P("HSR", head), P("Sprint táv", head), P("Sprint db", head), P("High Eff.", head), P("Load", head)]]
        for _, r in trend_v99.get("totals").tail(5).iterrows():
            tr.append([
                P(str(r.get("week","")), small),
                P(_fpi_fmt_thousands_v97(r.get("total_distance",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("hsr_distance",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("sprint_distance",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("sprints",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("high_efforts",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("training_load",0)), small),
            ])
        story.append(tbl(tr, [3.2*cm, 4.0*cm, 3.7*cm, 4.0*cm, 3.3*cm, 4.0*cm, 3.5*cm], header_bg="#047857", row_bgs=[colors.HexColor("#ECFDF5"), colors.white]))
        if trend_v99.get("signals"):
            story.append(P("Trendjelzések: " + "; ".join(trend_v99.get("signals", [])[:4]) + ".", small))
        story.append(Spacer(1, 0.22*cm))

    story.append(section("5. Heti csapat GPS profil", "#E0F2FE"))
    wk = weekly.copy()
    if not wk.empty and "week" in wk.columns:
        wk = wk[wk["week"].astype(str) == str(week)].copy()
    gps_rows = [[P("Típus", head), P("Össztáv", head), P("Perc", head), P("m/perc", head), P("HSR", head), P("Sprint táv", head), P("Sprint db", head), P("High Eff.", head), P("Load", head)]]
    if not wk.empty:
        for _, r in wk.head(10).iterrows():
            gps_rows.append([
                P(r.get("session_type",""), small),
                P(_fpi_fmt_thousands_v97(r.get("total_distance",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("duration_min",0)), small),
                P(f"{r.get('distance_per_min',0):.1f}", small),
                P(_fpi_fmt_thousands_v97(r.get("hsr_distance",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("sprint_distance",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("sprints",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("high_efforts",0)), small),
                P(_fpi_fmt_thousands_v97(r.get("training_load",0)), small),
            ])
    else:
        gps_rows.append([P("Nincs adat", small)] + [P("—", small)]*8)
    story.append(tbl(gps_rows, [3.0*cm, 3.2*cm, 2.5*cm, 2.6*cm, 3.0*cm, 3.0*cm, 2.5*cm, 2.8*cm, 2.8*cm], header_bg="#0369A1"))

    story.append(Spacer(1, 0.22*cm))
    story.append(section("6. Játékosszintű monitoring", "#FEE2E2"))
    risk_rows = [[P("Játékos", head), P("Szint", head), P("Miért fontos?", head)]]
    if risk_df is not None and not risk_df.empty:
        player_col = "Játékos" if "Játékos" in risk_df.columns else "player_name" if "player_name" in risk_df.columns else risk_df.columns[0]
        level_col = "Kockázati szint" if "Kockázati szint" in risk_df.columns else None
        reason_col = "Fő okok" if "Fő okok" in risk_df.columns else "Fő ok" if "Fő ok" in risk_df.columns else None
        for _, r in risk_df.head(8).iterrows():
            risk_rows.append([
                P(str(r.get(player_col,"")), small),
                P(str(r.get(level_col,"Figyelendő")) if level_col else "Figyelendő", small),
                P(str(r.get(reason_col,"Monitoring")) if reason_col else "Monitoring", small),
            ])
    else:
        risk_rows.append([P("Nincs kiemelt", small), P("Alacsony", small), P("Nincs azonnali beavatkozási jelzés.", small)])
    story.append(tbl(risk_rows, [6.5*cm, 4.0*cm, 17.2*cm], header_bg="#7F1D1D", row_bgs=[colors.white, colors.HexColor("#FEF2F2")]))

    story.append(PageBreak())
    story.append(section("7. Fogalmak röviden", "#DCFCE7"))
    expl = [[P("Fogalom", head), P("Jelentés", head)]]
    expl.append([P("Volumen", small), P("Összmunka: például teljes táv, edzésidő, load vagy heti összterhelés.", small)])
    expl.append([P("HSR", small), P("High Speed Running: nagy sebességű futás, általában kb. 19,8–20 km/h felett. Nem feltétlen maximális sprint.", small)])
    expl.append([P("Sprint", small), P("Sprintzónában megtett méter vagy sprintakciók darabszáma. Általában magasabb küszöb, mint a HSR.", small)])
    expl.append([P("Sprint expozíció", small), P("A játékos kap-e maximális sebességhez közeli ingert a héten. Nem csak mennyiség, hanem sérülésmegelőzési inger.", small)])
    expl.append([P("Load / terhelési pont", small), P("A GPS-rendszer összesített terhelési mutatója. Rendszertől függően tartalmazhat futómennyiséget, intenzitást, gyorsítást, lassítást vagy mechanikai terhelést.", small)])
    expl.append([P("High Efforts", small), P("Nagy intenzitású akciók gyűjtőmutatója. GPS-rendszertől függően sprint, gyorsítás, lassítás vagy robbanékony effort is lehet benne.", small)])
    story.append(tbl(expl, [5.0*cm, 22.7*cm], header_bg="#166534", row_bgs=[colors.HexColor("#ECFDF5"), colors.white]))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def build_fpi_gps_only_sample_pdf_bytes() -> Optional[bytes]:
    demo_raw = build_demo_performance_data()
    demo_df, _, missing = standardize_dataframe(demo_raw)
    if missing:
        return None
    demo_df = add_position_group(demo_df)
    latest = _fpi_latest_week(demo_df)
    return build_fpi_gps_only_pdf_bytes(demo_df, latest, "Magas presszing", demo_label="MINTA RIPORT / KTE U19 – GPS only")


def build_fpi_sample_pdf_bytes(report_type: str = "full", include_tactical: bool = True) -> Optional[bytes]:
    demo_raw = build_demo_performance_data()
    demo_df, _, missing = standardize_dataframe(demo_raw)
    if missing:
        return None
    demo_df = add_position_group(demo_df)
    latest = _fpi_latest_week(demo_df)
    tactical_ctx = _build_demo_tactical_context() if include_tactical else None
    if include_tactical and isinstance(tactical_ctx, dict):
        # V134: a minta Executive Summary-ben is legyen konkrét ellenfél-játékos értékelés referenciaértékekkel.
        if not tactical_ctx.get("opponent_player_evaluation"):
            tactical_ctx["opp_player_tables"] = {
                "creators": pd.DataFrame([{"player":"Zsolt János Magyar","position":"AM","key_passes":4.0,"progressive_passes":7.0,"shots":2.0,"xg":0.20}]),
                "wide_players": pd.DataFrame([{"player":"Bence Szabó","position":"RM","crosses":8.0,"key_passes":3.0,"lost_balls":4.0}]),
                "finishers": pd.DataFrame([{"player":"András Simon","position":"CF","shots":4.0,"xg":0.60,"goals":1.0}]),
                "weak_links": pd.DataFrame([{"player":"Márk Helembai","position":"CB","lost_balls":7.0,"passes":42.0}]),
            }
            tactical_ctx["opponent_player_evaluation"] = _fpi_build_player_evaluation_v132(tactical_ctx["opp_player_tables"], side="opp", max_rows=8)
        else:
            # A korábbi statikus minta sorokat kiegészítjük referencia-jellegű bizonyítékkal, ha még nincs benne.
            for rr in tactical_ctx.get("opponent_player_evaluation", []):
                if "ref:" not in str(rr.get("Bizonyíték", "")):
                    role = str(rr.get("Szerep", "")).lower()
                    if "befejező" in role:
                        rr["Bizonyíték"] = "lövés: 4.0 (ref: 1–3; sok lövés); xG: 0.6 (ref: 0.20–0.50; magas helyzetminőség)"
                    elif "szélső" in role or "beadó" in role:
                        rr["Bizonyíték"] = "beadás: 8.0 (ref: 3–6; magas beadási volumen); kulcspassz: 3.0 (ref: 1–3; referenciatartományban)"
                    elif "kreatív" in role:
                        rr["Bizonyíték"] = "kulcspassz: 4.0 (ref: 1–3; kiemelt kreatív akció); progresszív passz: 7.0 (ref: 4–8; referenciatartományban)"
    label = "MINTA RIPORT / KTE U19 – GPS + Tactical" if include_tactical else "MINTA RIPORT / KTE U19 – GPS only"
    return build_fpi_product_pdf_bytes(demo_df, latest, "Magas presszing", report_type=report_type, demo_label=label, tactical_context=tactical_ctx)


# =========================================================
# V132 - Own team report + opponent player intelligence helpers
# =========================================================
FPI_FORMATION_OPTIONS_V132 = ["4-2-3-1", "4-3-3", "3-5-2", "3-4-3", "4-4-2", "5-3-2", "5-4-1", "Egyéb"]
FPI_BLOCK_OPTIONS_V132 = ["Magas blokk", "Középső blokk", "Mély blokk", "Vegyes / meccsfüggő"]
FPI_ATTACK_ROUTE_OPTIONS_V132 = ["Szélek", "Centrum / félterületek", "Átmenetek", "Direkt játék", "Vegyes"]


def _fpi_num_v133(v, default: Optional[float] = None) -> Optional[float]:
    try:
        if v is None or v == "":
            return default
        return float(v)
    except Exception:
        return default



def _fpi_ref_eval_value_v134(metric_key: str, value: object, level: str = "player") -> Tuple[str, str]:
    """Rövid referencia/értékelés taktikai PDF-ekhez. Nem abszolút norma, hanem coach-friendly összevetési sáv."""
    try:
        v = float(str(value).replace("%", "").replace(",", "."))
    except Exception:
        return ("nincs referencia", "nem értékelhető")
    pct_keys = {"possession_pct", "pressing_success_pct", "passes_accurate_pct"}
    if metric_key in pct_keys and v <= 1:
        v *= 100
    refs = {
        # csapat
        "possession_pct": (45, 55, "45–55%", "alacsony labdabirtoklás", "domináns labdabirtoklás"),
        "shots": (8, 14, "8–14", "alacsony lövésvolumen", "magas lövésvolumen"),
        "xg": (1.0, 1.8, "1.0–1.8", "kevés / alacsony minőségű helyzet", "erős helyzetminőség"),
        "entries_box": (12, 20, "12–20", "kevés boxjelenlét", "erős boxjelenlét"),
        "final_third_entries": (35, 55, "35–55", "kevés támadóharmad-belépés", "sok támadóharmad-belépés"),
        "key_passes": (5, 10, "5–10", "kevés kulcspassz", "magas kreatív volumen"),
        "corners": (3, 6, "3–6", "kevés pontrúgásnyomás", "pontrúgásveszély magas"),
        "ppda": (8, 12, "8–12", "nagyon agresszív presszing", "passzívabb védekezési aktivitás"),
        "pressing_success_pct": (50, 65, "50–65%", "gyengébb presszinghatékonyság", "erős presszinghatékonyság"),
        "passes_accurate_pct": (78, 86, "78–86%", "labdabiztonsági kockázat", "stabil passzminőség"),
        "crosses": (10, 18, "10–18", "kevés szélső/beadási volumen", "magas szélső/beadási volumen"),
        "recoveries": (45, 65, "45–65", "kevés labdaszerzés", "magas labdaszerzési aktivitás"),
        "lost_balls": (25, 40, "25–40", "kevés labdavesztés", "sok labdavesztés / támadható"),
        "counterattacks": (3, 7, "3–7", "kevés átmeneti veszély", "magas kontraveszély"),
        # játékos
        "player_shots": (1, 3, "1–3", "kevés lövés", "sok lövés"),
        "player_xg": (0.20, 0.50, "0.20–0.50", "alacsony helyzetminőség", "magas helyzetminőség"),
        "player_key_passes": (1, 3, "1–3", "kevés kreatív akció", "kiemelt kreatív akció"),
        "player_progressive_passes": (4, 8, "4–8", "kevés progresszió", "erős progresszió"),
        "player_crosses": (3, 6, "3–6", "kevés beadás", "magas beadási volumen"),
        "player_passes": (35, 60, "35–60", "kevés labdás részvétel", "labdakihozatali hub"),
        "player_recoveries": (4, 8, "4–8", "kevés labdaszerzés", "aktív labdaszerző"),
        "player_interceptions": (2, 5, "2–5", "kevés közbelépés", "sok közbelépés"),
        "player_defensive_challenges": (5, 10, "5–10", "kevés párharc", "magas párharcintenzitás"),
        "player_lost_balls": (3, 6, "3–6", "labdabiztos", "labdavesztési kockázat"),
        "player_minutes_played": (60, 90, "60–90 perc", "rövid játékidő", "teljes/majdnem teljes terhelés"),
    }
    if metric_key not in refs:
        return ("nincs fix sáv", "kontextus alapján értelmezendő")
    lo, hi, label, low_txt, high_txt = refs[metric_key]
    # PPDA és lost_balls típusú mutatóknál az alacsonyabb érték nem feltétlen gyenge.
    if metric_key == "ppda":
        if v < lo:
            return (label, low_txt)
        if v > hi:
            return (label, high_txt)
        return (label, "aktív / kontrollált presszingzóna")
    if metric_key in ["lost_balls", "player_lost_balls"]:
        if v < lo:
            return (label, low_txt)
        if v > hi:
            return (label, high_txt)
        return (label, "elfogadható labdavesztési tartomány")
    if v < lo:
        return (label, low_txt)
    if v > hi:
        return (label, high_txt)
    return (label, "referenciatartományban")


def _fpi_player_metric_summary_v133(r: object) -> str:
    """Játékosszintű mutatóösszefoglaló referenciaértékkel és értékeléssel."""
    if not hasattr(r, "get"):
        return "játékosstatisztika alapján"
    parts: List[str] = []
    def add(label: str, key: str, ref_key: Optional[str] = None, suffix: str = ""):
        v = r.get(key, None)
        if v in [None, "", 0, 0.0]:
            return
        try:
            fv = float(v)
            val_txt = f"{fv:.1f}{suffix}"
        except Exception:
            val_txt = str(v)
        ref_label, eval_txt = _fpi_ref_eval_value_v134(ref_key or f"player_{key}", v, "player")
        if ref_label and ref_label != "nincs referencia":
            parts.append(f"{label}: {val_txt} (ref: {ref_label}; {eval_txt})")
        else:
            parts.append(f"{label}: {val_txt}")
    add("lövés", "shots", "player_shots")
    add("xG", "xg", "player_xg")
    add("gól", "goals", None)
    add("kulcspassz", "key_passes", "player_key_passes")
    add("progresszív passz", "progressive_passes", "player_progressive_passes")
    add("beadás", "crosses", "player_crosses")
    add("passz", "passes", "player_passes")
    add("labdaszerzés", "recoveries", "player_recoveries")
    add("közbelépés", "interceptions", "player_interceptions")
    add("védekező párharc", "defensive_challenges", "player_defensive_challenges")
    add("labdavesztés", "lost_balls", "player_lost_balls")
    add("játékperc", "minutes_played", "player_minutes_played")
    return "; ".join(parts[:5]) if parts else "játékosstatisztika alapján"


def _fpi_player_interpretation_v133(r: object, side: str, role: str) -> str:
    """Értékelés: kategóriázott, edzői nyelvű játékosszintű olvasat.
    V135: az Értékelés oszlopban egyértelműen jelzi, mi kiemelkedő/jó/átlagos/fejlesztendő/figyelendő.
    """
    if not hasattr(r, "get"):
        return role

    metrics = {
        "shots": _fpi_num_v133(r.get("shots"), 0) or 0,
        "xg": _fpi_num_v133(r.get("xg"), 0) or 0,
        "goals": _fpi_num_v133(r.get("goals"), 0) or 0,
        "key_passes": _fpi_num_v133(r.get("key_passes"), 0) or 0,
        "progressive_passes": _fpi_num_v133(r.get("progressive_passes"), 0) or 0,
        "crosses": _fpi_num_v133(r.get("crosses"), 0) or 0,
        "passes": _fpi_num_v133(r.get("passes"), 0) or 0,
        "recoveries": _fpi_num_v133(r.get("recoveries"), 0) or 0,
        "interceptions": _fpi_num_v133(r.get("interceptions"), 0) or 0,
        "defensive_challenges": _fpi_num_v133(r.get("defensive_challenges"), 0) or 0,
        "lost_balls": _fpi_num_v133(r.get("lost_balls"), 0) or 0,
        "minutes_played": _fpi_num_v133(r.get("minutes_played"), 0) or 0,
    }

    rows: List[str] = []

    positive_high = {
        "sok lövés", "magas helyzetminőség", "kiemelt kreatív akció", "erős progresszió",
        "magas beadási volumen", "labdakihozatali hub", "aktív labdaszerző", "sok közbelépés",
        "magas párharcintenzitás", "teljes/majdnem teljes terhelés", "stabil passzminőség",
        "labdabiztos",
    }
    negative_low = {
        "kevés lövés", "alacsony helyzetminőség", "kevés kreatív akció", "kevés progresszió",
        "kevés beadás", "kevés labdás részvétel", "kevés labdaszerzés", "kevés közbelépés",
        "kevés párharc", "rövid játékidő",
    }
    warning = {"labdavesztési kockázat"}

    def add_metric(key: str, label: str, ref_key: str):
        v = metrics.get(key, 0)
        if not v:
            return
        _ref, ev = _fpi_ref_eval_value_v134(ref_key, v, "player")
        ev_l = str(ev).lower()
        if ev_l in positive_high:
            cat = "Kiemelkedő"
        elif ev_l == "referenciatartományban" or ev_l == "elfogadható labdavesztési tartomány":
            cat = "Átlagos / stabil"
        elif ev_l in warning:
            cat = "Figyelendő"
        elif ev_l in negative_low:
            cat = "Fejlesztendő"
        else:
            cat = "Jó" if ev_l and ev_l not in ["kontextus alapján értelmezendő", "nem értékelhető"] else "Átlagos / stabil"
        rows.append(f"{cat}: {label} – {ev}")

    add_metric("shots", "lövésaktivitás", "player_shots")
    add_metric("xg", "helyzetminőség", "player_xg")
    add_metric("key_passes", "kreativitás", "player_key_passes")
    add_metric("progressive_passes", "progresszió", "player_progressive_passes")
    add_metric("crosses", "beadási volumen", "player_crosses")
    add_metric("passes", "labdás részvétel", "player_passes")
    add_metric("recoveries", "labdaszerzés", "player_recoveries")
    add_metric("interceptions", "közbelépés", "player_interceptions")
    add_metric("defensive_challenges", "párharcintenzitás", "player_defensive_challenges")
    add_metric("lost_balls", "labdabiztonság", "player_lost_balls")

    # Befejezési hatékonyság külön olvasat: sok helyzet/xG + nincs gól = figyelendő, nem feltétlen rossz játék.
    if metrics["shots"] >= 3 or metrics["xg"] >= 0.35:
        if metrics["goals"] < max(1, metrics["xg"] * 0.8):
            rows.append("Figyelendő: befejezési hatékonyság – a helyzetminőséghez képest kevés gól")
        else:
            rows.append("Kiemelkedő: befejezés – kapura veszélyes és hatékony")

    # Ha csak átlagos/stabil jelzések vannak, legyen rövid, egyértelmű összegzés.
    if not rows:
        rows.append("Átlagos / stabil: a szerephez tartozó fő mutatók nem jeleznek kiugró eltérést")

    # Dedup és sorrend: előbb a fontosabb kategóriák.
    priority = {"Kiemelkedő": 0, "Jó": 1, "Figyelendő": 2, "Fejlesztendő": 3, "Átlagos / stabil": 4}
    dedup = []
    seen = set()
    for item in rows:
        if item not in seen:
            seen.add(item)
            dedup.append(item)
    dedup.sort(key=lambda x: priority.get(x.split(":", 1)[0], 9))
    return "; ".join(dedup[:5])

def _fpi_player_action_v133(side: str, role: str, interp: str) -> str:
    txt = (role + " " + interp).lower()
    if side == "opp":
        if "progressz" in txt or "labdafelhozó" in txt:
            return "Ne fordulhasson szabadon előre; presszing trigger és oldalra terelés."
        if "kreatív" in txt or "kulcspassz" in txt:
            return "Passzsáv zárása, testhelyzet kontroll, belső irány lezárása."
        if "beadás" in txt or "oldali" in txt:
            return "Oldali 1v1 kontroll, beadásblokkolás, hosszú oldali zárás."
        if "befejező" in txt or "xg" in txt or "helyzet" in txt:
            return "Boxon belüli felvétel korlátozása, második labdák kontrollja."
        if "labdaveszt" in txt:
            return "Célzott presszingpont lehet; labdaátvételnél agresszív nyomás."
        return "Meccstervben külön kontrollpontként kezelni."
    else:
        if "progressz" in txt:
            return "Build-upban tudatosan keresni; első/középső fázisban progressziós opció."
        if "kreatív" in txt or "kulcspassz" in txt:
            return "Támadóharmadban több kapcsolódás, félterületi labdafelvevő szerep."
        if "beadás" in txt or "oldali" in txt:
            return "Szélesség és beadási helyzetek tudatos kialakítása rajta keresztül."
        if "befejez" in txt or "helyzet" in txt:
            return "Helyzetminőség és döntés kontroll; befejező edzésblokk indokolt lehet."
        if "labdaveszt" in txt:
            return "Labdabiztonság / döntésgyorsaság fejlesztési fókusz."
        if "labdaszerzés" in txt:
            return "Védekező átmenetben és második labdáknál támaszkodni rá."
        return "Szerepét a heti tervben célzottan érdemes használni."


def _fpi_df_to_player_eval_rows_v132(table: object, role: str, value_col: str, meaning: str, action: str, max_rows: int = 3, side: str = "opp") -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    if not isinstance(table, pd.DataFrame) or table.empty or "player" not in table.columns:
        return rows
    for _, r in table.head(max_rows).iterrows():
        player = str(r.get("player", "")).strip()
        if not player:
            continue
        pos = str(r.get("position", "")).strip()
        evidence = _fpi_player_metric_summary_v133(r)
        interpretation = _fpi_player_interpretation_v133(r, side=side, role=role)
        action_txt = _fpi_player_action_v133(side, role, interpretation)
        rows.append({"Játékos": player, "Poszt": pos, "Szerep": role, "Bizonyíték": evidence, "Értelmezés": interpretation, "Javaslat": action_txt})
    return rows

def _fpi_build_player_evaluation_v132(player_tables: object, side: str = "opp", max_rows: int = 9) -> List[Dict[str, str]]:
    """Játékosszintű taktikai értékelés.
    V133: nem csak szerepcímkét ad, hanem a rendelkezésre álló mutatókból rövid teljesítmény-értékelést is épít.
    """
    if not isinstance(player_tables, dict):
        return []
    if side == "opp":
        specs = [
            ("creators", "Kreatív kulcsjátékos", "key_passes", "", ""),
            ("progressors", "Labdafelhozó / progresszor", "progressive_passes", "", ""),
            ("build_up", "Labdakihozatal szervező", "passes", "", ""),
            ("finishers", "Befejező / kapura veszélyes játékos", "xg", "", ""),
            ("wide_players", "Szélső / beadó veszély", "crosses", "", ""),
            ("duel_players", "Párharcerős játékos", "defensive_challenges", "", ""),
            ("defenders", "Labdaszerző / védekező láncszem", "interceptions", "", ""),
            ("weak_links", "Támadható láncszem", "lost_balls", "", ""),
        ]
    else:
        specs = [
            ("creators", "Saját kreatív kapcsolódási pont", "key_passes", "", ""),
            ("progressors", "Saját progresszor", "progressive_passes", "", ""),
            ("build_up", "Saját labdakihozatal szervező", "passes", "", ""),
            ("finishers", "Saját befejező / helyzetjátékos", "xg", "", ""),
            ("wide_players", "Saját szélső / beadó opció", "crosses", "", ""),
            ("duel_players", "Saját párharcerős játékos", "defensive_challenges", "", ""),
            ("defenders", "Saját labdaszerző", "interceptions", "", ""),
            ("weak_links", "Fejlesztendő labdabiztonság", "lost_balls", "", ""),
        ]
    rows: List[Dict[str, str]] = []
    seen_players = set()
    for key, role, col, meaning, action in specs:
        for r in _fpi_df_to_player_eval_rows_v132(player_tables.get(key), role, col, meaning, action, max_rows=3, side=side):
            # Egy játékos több szerepben is megjelenhet, de a PDF-ben rövid maradjon: első két legerősebb szerep elég.
            player = r.get("Játékos", "")
            sig_count = sum(1 for x in rows if x.get("Játékos") == player)
            if sig_count >= 2:
                continue
            rows.append(r)
            if len(rows) >= max_rows:
                return rows
    return rows[:max_rows]

def _fpi_player_eval_to_findings_v132(rows: List[Dict[str, str]], side_label: str = "Ellenfél") -> List[Dict[str, str]]:
    out = []
    for r in rows[:6]:
        player = r.get("Játékos", "")
        role = r.get("Szerep", "")
        out.append({
            "Téma": f"{side_label}: {role}",
            "Bizonyíték": f"{player} ({r.get('Poszt','')}): {r.get('Bizonyíték','')}",
            "Edzői következtetés": r.get("Javaslat", ""),
            "Prioritás": "Magas" if side_label == "Ellenfél" and any(k in role.lower() for k in ["kulcs", "progresszor", "labdafelhozó"]) else "Közepes",
        })
    return out



def _fpi_team_metric_reading_v133(label: str, value: str, metric_key: str, own_ctx: Optional[Dict[str, str]] = None) -> str:
    own_ctx = own_ctx or {}
    route = str(own_ctx.get("attack_route", "")).lower()
    block = str(own_ctx.get("defensive_block", "")).lower()
    playmodel = str(own_ctx.get("playmodel", "")).lower()
    try:
        v = float(str(value).replace("%", "").replace(",", "."))
    except Exception:
        v = None
    ref_label, eval_txt = _fpi_ref_eval_value_v134(metric_key, v if v is not None else value, "team")
    prefix = f"Értékelés: {eval_txt}. " if eval_txt else ""
    if metric_key == "possession_pct":
        if v is not None and v >= 55:
            return prefix + "A csapat labdával kontrollképes; dominancia vagy pozíciós támadás esetén ez erősség, de rest defense biztosítás kell."
        if v is not None and v <= 45:
            return prefix + "Kevesebb labdás kontroll; átmeneti vagy direktabb játéktervvel lehet hatékonyabb."
        return prefix + "Kiegyensúlyozott labdabirtoklási kép; a meccstervet inkább a veszélyforrások döntik el."
    if metric_key == "shots":
        return prefix + "A támadó aktivitást jelzi; xG-vel együtt dönthető el, hogy mennyiség vagy valódi helyzetminőség áll mögötte."
    if metric_key == "xg":
        return prefix + "A helyzetek minőségét mutatja; magas értéknél a támadó automatizmusok működnek, alacsonynál a boxba jutás/helyzetminőség fejlesztendő."
    if metric_key == "entries_box":
        return prefix + "A büntetőterületi jelenlétet jelzi; szélső és félterületi támadásoknál kulcsfontosságú."
    if metric_key == "final_third_entries":
        return prefix + "A támadóharmadba jutás gyakorisága; ha magas, a progresszió működik, ha alacsony, az első/középső építési fázis lehet szűk keresztmetszet."
    if metric_key == "key_passes":
        return prefix + "Kreatív áttörések száma; magas értéknél több minőségi utolsó passz és döntési pont azonosítható."
    if metric_key == "corners":
        return prefix + "Pontrúgásnyomás; magas értéknél támadó pontrúgás-rutinokból is lehet versenyelőny."
    if metric_key == "ppda":
        return prefix + "A védekezési aktivitás/presszing egyik proxyja; alacsonyabb érték agresszívebb nyomást jelez."
    if metric_key == "pressing_success_pct":
        return prefix + "Letámadás hatékonysága; magas értéknél presszingből is építhető meccsterv, alacsonynál kontrolláltabb trigger kell."
    if metric_key == "passes_accurate_pct":
        return prefix + "Labdabiztonság és technikai stabilitás; presszing alatt különösen fontos."
    if metric_key == "crosses":
        return prefix + "Szélső játék / beadási volumen; oldali fölény vagy beadásokra építő támadási út jele."
    if metric_key == "recoveries":
        return prefix + "Labdaszerzési aktivitás; visszatámadásnál és középső blokkban fontos."
    if metric_key == "lost_balls":
        return prefix + "Labdabiztonsági kockázat; magas értéknél presszing alatt támadható vagy döntésgyorsasági fejlesztési pont."
    if metric_key == "counterattacks":
        return prefix + "Átmeneti veszély / gyors támadások; rest defense és visszarendeződés szempontból kulcsmutató."
    return prefix + "A saját játékmodell, felállás és heti terhelhetőség alapján értelmezendő."


def _fpi_team_metric_rows_v132(metrics: object) -> List[Tuple[str, str, str, str, str]]:
    """Csapatszintű KPI-k: minden releváns rendelkezésre álló mező + referencia/értékelés."""
    if not isinstance(metrics, dict) or not metrics:
        return []
    label_map = {
        "possession_pct": "Labdabirtoklás", "shots": "Lövések", "xg": "xG", "entries_box": "Box belépések",
        "final_third_entries": "Támadóharmad belépések", "key_passes": "Kulcspasszok", "corners": "Szögletek",
        "ppda": "PPDA", "pressing_success_pct": "Presszing sikeresség", "passes_accurate_pct": "Passzpontosság",
        "crosses": "Beadások", "recoveries": "Labdaszerzések", "lost_balls": "Labdavesztések", "counterattacks": "Kontrák",
    }
    own_ctx = _fpi_own_context_from_session_v132() if "_fpi_own_context_from_session_v132" in globals() else {}
    rows = []
    for k, lab in label_map.items():
        v = metrics.get(k)
        if v in [None, "", 0, 0.0]:
            continue
        try:
            fv = float(v)
            if k in ["possession_pct", "pressing_success_pct", "passes_accurate_pct"] and fv <= 1:
                fv *= 100
            suffix = "%" if k in ["possession_pct", "pressing_success_pct", "passes_accurate_pct"] else ""
            val_txt = f"{fv:.1f}{suffix}"
        except Exception:
            val_txt = str(v)
        ref_label, eval_txt = _fpi_ref_eval_value_v134(k, v, "team")
        ref_txt = f"{ref_label}; {eval_txt}" if ref_label else eval_txt
        rows.append((lab, val_txt, ref_txt, _fpi_team_metric_reading_v133(lab, val_txt, k, own_ctx), k))
    return rows[:14]

def _fpi_own_context_from_session_v132() -> Dict[str, str]:
    return {
        "formation": str(st.session_state.get("clean_own_formation_v132", "")),
        "defensive_block": str(st.session_state.get("clean_own_block_v132", "")),
        "attack_route": str(st.session_state.get("clean_own_attack_route_v132", "")),
        "playmodel": str(st.session_state.get("clean_playmodel_profile_v112", st.session_state.get("clean_playstyle_v112", ""))),
    }


def build_fpi_own_team_profile_pdf_bytes(
    data: pd.DataFrame,
    selected_week: Optional[str] = None,
    playstyle: str = "Kiegyensúlyozott",
    tactical_context: Optional[Dict[str, object]] = None,
    demo_label: str = "",
) -> Optional[bytes]:
    """Külön PDF: saját csapat csapat- és játékosszintű profil. Nem meccsriport, hanem csapatdiagnózis."""
    if SimpleDocTemplate is None:
        return None
    try:
        ctx = _fpi_report_context(data, selected_week, playstyle)
    except Exception:
        return None
    if ctx.get("error"):
        return None
    tactical_context = tactical_context if tactical_context is not None else st.session_state.get("tactical_pro_context", {})
    own_ctx = _fpi_own_context_from_session_v132()
    readiness = int(ctx.get("readiness_score", 70) or 70)
    risk_df = ctx.get("risk_df") if isinstance(ctx.get("risk_df"), pd.DataFrame) else pd.DataFrame()
    high_risk, med_risk = _fpi_count_risk_levels_v126(risk_df) if "_fpi_count_risk_levels_v126" in globals() else (0, 0)
    week = ctx.get("selected_week", selected_week or "")
    own_metrics = (tactical_context or {}).get("own_team_metrics", {}) if isinstance(tactical_context, dict) else {}
    own_players = (tactical_context or {}).get("own_player_tables", {}) if isinstance(tactical_context, dict) else {}
    own_eval = (tactical_context or {}).get("own_player_evaluation", []) if isinstance(tactical_context, dict) else []
    if not own_eval:
        own_eval = _fpi_build_player_evaluation_v132(own_players, side="own", max_rows=10)
    strategy = ((tactical_context or {}).get("strategy_framework") or {}) if isinstance(tactical_context, dict) else {}

    font_name, font_bold = _register_pdf_font()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=0.9*cm, leftMargin=0.9*cm, topMargin=0.7*cm, bottomMargin=0.7*cm)
    styles = getSampleStyleSheet()
    title = ParagraphStyle("FPIOwnTitle", parent=styles["Title"], fontName=font_bold, fontSize=20, leading=23, textColor=colors.HexColor("#0F172A"))
    sub = ParagraphStyle("FPIOwnSub", parent=styles["Normal"], fontName=font_name, fontSize=8.5, leading=10.5, textColor=colors.HexColor("#334155"))
    body = ParagraphStyle("FPIOwnBody", parent=styles["Normal"], fontName=font_name, fontSize=7.6, leading=9.3, textColor=colors.HexColor("#111827"))
    small = ParagraphStyle("FPIOwnSmall", parent=styles["Normal"], fontName=font_name, fontSize=6.5, leading=7.8, textColor=colors.HexColor("#111827"))
    head = ParagraphStyle("FPIOwnHead", parent=styles["Normal"], fontName=font_bold, fontSize=7.0, leading=8.5, alignment=1, textColor=colors.white)

    def clean(v):
        return pdf_safe_text(v).replace("\r", "").strip()
    def P(v, style=body):
        return Paragraph(html.escape(clean(v)).replace("\n", "<br/>") or "—", style)
    def section(txt, color="#DBEAFE"):
        t = Table([[P(txt, ParagraphStyle("FPIOwnH", parent=body, fontName=font_bold, fontSize=11.0, leading=13.0, textColor=colors.HexColor("#0F172A")))]], colWidths=[27.7*cm])
        t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),colors.HexColor(color)),("BOX",(0,0),(-1,-1),0.4,colors.HexColor("#93C5FD")),("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
        return t
    def table(rows, widths, header_bg="#0F172A", row_bgs=None):
        if row_bgs is None:
            row_bgs=[colors.white, colors.HexColor("#F8FAFC")]
        t=Table(rows, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor(header_bg)),("GRID",(0,0),(-1,-1),0.25,colors.HexColor("#CBD5E1")),("VALIGN",(0,0),(-1,-1),"TOP"),("ROWBACKGROUNDS",(0,1),(-1,-1),row_bgs),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3)]))
        return t

    story=[]
    label_prefix = f"{demo_label} | " if demo_label else ""
    story.append(P("Football Performance Intelligence – saját csapat profil", title))
    story.append(P(f"{label_prefix}Hét: {format_week_label(str(week))} | Játékmodell: {playstyle} | Generálva: {datetime.now().strftime('%Y-%m-%d %H:%M')}", sub))
    story.append(Spacer(1,0.18*cm))
    rows=[[P("Terület", head), P("Érték / értelmezés", head)]]
    rows += [
        [P("Readiness", small), P(f"{readiness}/100 – {score_to_label(readiness)}", small)],
        [P("Játékoskockázat", small), P(f"{high_risk} magas / {med_risk} közepes", small)],
        [P("Alapfelállás", small), P(own_ctx.get("formation") or "nincs megadva", small)],
        [P("Védekezési blokk", small), P(own_ctx.get("defensive_block") or "nincs megadva", small)],
        [P("Fő támadási út", small), P(own_ctx.get("attack_route") or "nincs megadva", small)],
        [P("Tactical Framework profil", small), P(str(strategy.get("primary", "nincs számított profil")), small)],
    ]
    story.append(table(rows,[5.5*cm,22.2*cm],header_bg="#0F766E",row_bgs=[colors.HexColor("#F0FDFA"),colors.white]))
    story.append(Spacer(1,0.18*cm))

    metric_rows=_fpi_team_metric_rows_v132(own_metrics)
    story.append(section("Csapatszintű taktikai / játékprofil", "#E0F2FE"))
    if metric_rows:
        tr=[[P("Mutató",head),P("Érték",head),P("Referencia / értékelés",head),P("Edzői olvasat",head)]]
        for lab,val,ref_txt,reading,_key in metric_rows:
            tr.append([P(lab,small),P(val,small),P(ref_txt,small),P(reading,small)])
        story.append(table(tr,[4.4*cm,3.0*cm,5.8*cm,14.5*cm],header_bg="#1E3A8A",row_bgs=[colors.HexColor("#EFF6FF"),colors.white]))
    else:
        story.append(P("Nincs saját csapatstatisztika Excel feltöltve. A riport ilyenkor GPS + megadott játékmodell alapján ad saját csapat profilt.", body))
    story.append(Spacer(1,0.18*cm))

    story.append(section("Játékosszintű saját értékelés", "#DCFCE7"))
    if own_eval:
        pr=[[P("Játékos",head),P("Szerep",head),P("Konkrét mutatók",head),P("Értékelés",head),P("Használati javaslat",head)]]
        for r in own_eval[:10]:
            pr.append([P(r.get("Játékos",""),small),P(r.get("Szerep",""),small),P(r.get("Bizonyíték",""),small),P(r.get("Értelmezés",""),small),P(r.get("Javaslat",""),small)])
        story.append(table(pr,[3.8*cm,4.8*cm,6.0*cm,6.2*cm,6.9*cm],header_bg="#166534",row_bgs=[colors.HexColor("#ECFDF5"),colors.white]))
    else:
        story.append(P("Nincs saját játékosstatisztika Excel. Játékosszintű szerepértékeléshez saját játékos Excel szükséges.", body))
    story.append(Spacer(1,0.18*cm))

    story.append(section("Edzői összegzés", "#FEF3C7"))
    bullets=[
        "Ez a riport nem ellenfélre készülő meccsterv, hanem a saját csapat aktuális profilja.",
        "A saját játékmodell, alapfelállás, blokkmagasság és fő támadási út segít a GPS- és taktikai inputok értelmezésében.",
        "A játékosszintű blokk azt mutatja, kikre építhető a labdakihozatal, progresszió, kreativitás vagy védekező stabilitás.",
    ]
    story.append(P("<br/>".join([f"• {b}" for b in bullets]), body))
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def build_fpi_own_team_profile_sample_pdf_bytes() -> Optional[bytes]:
    demo_raw = build_demo_performance_data()
    demo_df, _, missing = standardize_dataframe(demo_raw)
    if missing:
        return None
    demo_df = add_position_group(demo_df)
    latest = _fpi_latest_week(demo_df)
    tactical_ctx = _build_demo_tactical_context()
    # Demo player tables, hogy a minta PDF-ben is legyen játékosszintű saját csapat blokk.
    tactical_ctx["own_player_tables"] = {
        "creators": pd.DataFrame([{"player":"Kovács M.","position":"AM","key_passes":4.0,"xg":0.22,"shots":2},{"player":"Szabó B.","position":"W","key_passes":3.0,"crosses":7}]),
        "progressors": pd.DataFrame([{"player":"Nagy D.","position":"CM","progressive_passes":9.0,"passes":58},{"player":"Tóth Á.","position":"DM","progressive_passes":7.0,"recoveries":8}]),
        "build_up": pd.DataFrame([{"player":"Varga L.","position":"CB","passes":62.0,"progressive_passes":6},{"player":"Tóth Á.","position":"DM","passes":55.0,"lost_balls":3}]),
        "finishers": pd.DataFrame([{"player":"Farkas Z.","position":"F","shots":5,"xg":0.75,"goals":0}]),
        "wide_players": pd.DataFrame([{"player":"Szabó B.","position":"W","crosses":7,"key_passes":3}]),
        "defenders": pd.DataFrame([{"player":"Farkas Z.","position":"CB","interceptions":6.0,"recoveries":9}]),
        "duel_players": pd.DataFrame([{"player":"Balogh P.","position":"FB","defensive_challenges":11.0,"recoveries":6}]),
    }
    tactical_ctx["own_player_evaluation"] = _fpi_build_player_evaluation_v132(tactical_ctx["own_player_tables"], side="own", max_rows=10)
    old_vals = {"clean_own_formation_v132": st.session_state.get("clean_own_formation_v132"), "clean_own_block_v132": st.session_state.get("clean_own_block_v132"), "clean_own_attack_route_v132": st.session_state.get("clean_own_attack_route_v132")}
    st.session_state["clean_own_formation_v132"] = st.session_state.get("clean_own_formation_v132", "4-2-3-1") or "4-2-3-1"
    st.session_state["clean_own_block_v132"] = st.session_state.get("clean_own_block_v132", "Középső blokk") or "Középső blokk"
    st.session_state["clean_own_attack_route_v132"] = st.session_state.get("clean_own_attack_route_v132", "Átmenetek") or "Átmenetek"
    pdf = build_fpi_own_team_profile_pdf_bytes(demo_df, latest, "Magas presszing", tactical_context=tactical_ctx, demo_label="MINTA RIPORT / KTE U19 – saját csapat profil")
    for k, v in old_vals.items():
        if v is None:
            st.session_state.pop(k, None)
        else:
            st.session_state[k] = v
    return pdf


def _fpi_enrich_tactical_context_v132(executive_ctx: Dict[str, object], own_player_tables: object, opp_player_tables: object) -> Dict[str, object]:
    """Executive context kiegészítése saját/ellenfél játékosszintű értékeléssel."""
    ctx = dict(executive_ctx or {})
    ctx["own_player_tables"] = own_player_tables if isinstance(own_player_tables, dict) else {}
    ctx["opp_player_tables"] = opp_player_tables if isinstance(opp_player_tables, dict) else {}
    existing_opp_eval = list(ctx.get("opponent_player_evaluation") or [])
    existing_own_eval = list(ctx.get("own_player_evaluation") or [])
    opp_eval = existing_opp_eval + _fpi_build_player_evaluation_v132(ctx["opp_player_tables"], side="opp", max_rows=9)
    own_eval = existing_own_eval + _fpi_build_player_evaluation_v132(ctx["own_player_tables"], side="own", max_rows=9)
    # dedupe player-role
    def _dedupe_eval(rows):
        seen=set(); out=[]
        for rr in rows:
            sig=(rr.get("Játékos",""), rr.get("Szerep",""))
            if sig in seen:
                continue
            seen.add(sig); out.append(rr)
        return out
    opp_eval = _dedupe_eval(opp_eval)[:9]
    own_eval = _dedupe_eval(own_eval)[:9]
    ctx["opponent_player_evaluation"] = opp_eval
    ctx["own_player_evaluation"] = own_eval
    findings = list(ctx.get("tactical_findings") or [])
    findings = _fpi_player_eval_to_findings_v132(opp_eval, "Ellenfél") + findings + _fpi_player_eval_to_findings_v132(own_eval[:3], "Saját")
    # duplikátumok kivétele, hogy a PDF ne legyen túl hosszú
    seen = set(); compact = []
    for f in findings:
        sig = (str(f.get("Téma","")), str(f.get("Bizonyíték",""))[:80])
        if sig in seen:
            continue
        seen.add(sig); compact.append(f)
    ctx["tactical_findings"] = compact[:14]
    # A kockázati/fókusz lista elejére bekerül az ellenfél játékosprofil is, így az Executive Summary-ban is látszik.
    risks = list(ctx.get("risks") or [])
    for r in opp_eval[:4]:
        txt = f"{r.get('Játékos')}: {r.get('Szerep')} – {r.get('Javaslat')}"
        if txt not in risks:
            risks.insert(0, txt)
    ctx["risks"] = risks[:10]
    return ctx

# V132 override: a taktikai üzenetek elsőként az ellenfél játékosszintű értékelésből induljanak, ha van ilyen.
def _fpi_top_tactical_messages_v82(tactical_context: Optional[Dict[str, object]], limit: int = 3, gps_context: Optional[Dict[str, object]] = None, readiness: Optional[int] = None, priorities: Optional[List[dict]] = None, week: Optional[str] = None) -> List[str]:
    if not _fpi_has_tactical_signal_v95(tactical_context):
        if gps_context is not None:
            return _fpi_gps_only_conclusions_v95(gps_context, priorities or [], int(readiness or 70), str(week or gps_context.get("selected_week", "")), limit=max(limit, 5))
        return ["GPS-only mód: nincs taktikai input, ezért a javaslat az erőnléti/GPS adatokra épül."]
    out: List[str] = []
    for r in (tactical_context or {}).get("opponent_player_evaluation", [])[:3]:
        out.append(f"{r.get('Játékos')}: {r.get('Szerep')} – {r.get('Javaslat')}")
    findings = (tactical_context or {}).get("tactical_findings") or []
    pdf_first = [f for f in findings if str(f.get("Téma", "")).lower().startswith("pdf") or "PDF" in str(f.get("Forrás", ""))]
    other = [f for f in findings if f not in pdf_first]
    for f in (pdf_first + other):
        if len(out) >= limit:
            break
        theme = str(f.get("Téma", "")).strip()
        decision = str(f.get("Edzői következtetés", "")).strip()
        if theme and decision:
            out.append(f"{theme}: {decision}")
    if not out:
        risks = (tactical_context or {}).get("risks", []) or []
        out = [str(x) for x in risks[:limit]]
    return [_fpi_clean_sentence_v82(x, 180) for x in out[:limit]] or ["Nincs erős taktikai jelzés."]


# =========================================================
# V10.0 - Landing page + import/export workflow
# =========================================================

def _fpi_set_page_v100(page: str) -> None:
    st.session_state["fpi_active_page_v100"] = page
    try:
        st.rerun()
    except Exception:
        pass




def _fpi_home_back_button_v123(key: str = "home_back_v123") -> None:
    """Egységes vissza gomb a főoldalra."""
    if st.button("← Főoldal", use_container_width=True, key=key):
        _fpi_set_page_v100("landing")

def _fpi_mapper_contrast_css_v109() -> None:
    """Mapper / Import / Executive Export kontraszt javítás."""
    st.markdown(
        """
        <style>
        /* Streamlit success/info/warning szövegek legyenek olvashatók sötét témában is */
        div[data-testid="stAlert"] p,
        div[data-testid="stAlert"] div,
        div[data-testid="stAlert"] span {
            color: #0f172a !important;
        }

        /* Expander fejlécek – ne fekete háttéren fekete szöveg legyen */
        details[data-testid="stExpander"] summary,
        details[data-testid="stExpander"] summary p,
        details[data-testid="stExpander"] summary span {
            color: #0f172a !important;
            font-weight: 800 !important;
        }
        details[data-testid="stExpander"] {
            background: #ffffff !important;
            border: 1px solid #d1d5db !important;
            border-radius: 14px !important;
        }

        /* Gombok: sötét gombon fehér szöveg */
        .stButton > button,
        div[data-testid="stDownloadButton"] > button {
            color: #ffffff !important;
            background: #111827 !important;
            border: 1px solid #334155 !important;
            font-weight: 800 !important;
        }
        .stButton > button:hover,
        div[data-testid="stDownloadButton"] > button:hover {
            color: #ffffff !important;
            background: #1f2937 !important;
            border-color: #475569 !important;
        }

        /* Selectbox / input szövegek */
        div[data-baseweb="select"] span,
        div[data-baseweb="select"] div,
        input {
            color: #111827 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _fpi_landing_css_v100() -> None:
    st.markdown(
        """
        <style>
        .fpi-landing-hero {
            border-radius: 30px;
            padding: 34px 38px;
            background: linear-gradient(135deg, #07111f 0%, #0f2a44 46%, #123d66 100%);
            border: 1px solid rgba(255,255,255,.16);
            box-shadow: 0 24px 70px rgba(2, 6, 23, .25);
            color: white;
            margin-bottom: 22px;
        }
        .fpi-landing-kicker {
            display: inline-block;
            padding: 6px 12px;
            border-radius: 999px;
            background: rgba(14,165,233,.18);
            border: 1px solid rgba(125,211,252,.35);
            color: #bae6fd;
            font-weight: 800;
            letter-spacing: .03em;
            font-size: .86rem;
            margin-bottom: 14px;
        }
        .fpi-landing-title {
            font-size: 3.0rem;
            line-height: 1.04;
            font-weight: 900;
            margin: 0 0 14px 0;
        }
        .fpi-landing-sub {
            font-size: 1.08rem;
            line-height: 1.55;
            color: #dbeafe;
            max-width: 980px;
        }
        .fpi-decision-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: 14px;
            margin-top: 24px;
        }
        .fpi-decision-card {
            background: rgba(255,255,255,.08);
            border: 1px solid rgba(255,255,255,.15);
            border-radius: 18px;
            padding: 16px;
            min-height: 132px;
        }
        .fpi-decision-card b {
            display:block;
            color:#fff;
            font-size:1.0rem;
            margin-bottom:8px;
        }
        .fpi-decision-card span {
            color:#cbd5e1;
            font-size:.92rem;
            line-height:1.42;
        }
        .fpi-mode-card {
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-radius: 22px;
            padding: 22px;
            box-shadow: 0 12px 30px rgba(15,23,42,.08);
            min-height: 230px;
        }
        .fpi-mode-card h3 {margin:0 0 8px 0; color:#0f172a;}
        .fpi-mode-card p {color:#475569; line-height:1.45;}
        .fpi-flow-step {
            border-radius: 18px;
            background: #f8fafc;
            border: 1px solid #e2e8f0;
            padding: 16px 18px;
            min-height: 118px;
            height: 118px;
            display: flex;
            flex-direction: column;
            justify-content: flex-start;
            box-sizing: border-box;
        }
        .fpi-flow-step b {
            color:#0f172a;
            display:block;
            min-height: 26px;
            line-height: 1.2;
        }
        .fpi-flow-step div {
            color:#64748b;
            font-size:.92rem;
            margin-top:7px;
            line-height: 1.38;
            min-height: 44px;
        }
        .fpi-section-banner {
            border-radius: 20px;
            padding: 15px 18px;
            margin: 18px 0 12px 0;
            border: 1px solid rgba(255,255,255,.18);
            box-shadow: 0 12px 28px rgba(15,23,42,.10);
        }
        .fpi-section-banner h3 { margin: 0; color: #ffffff; font-size: 1.22rem; font-weight: 950; }
        .fpi-section-banner p { margin: 5px 0 0 0; color: rgba(255,255,255,.88); line-height: 1.35; }
        .fpi-section-gps { background: linear-gradient(135deg,#0f766e,#14b8a6); }
        .fpi-section-tactical { background: linear-gradient(135deg,#1d4ed8,#60a5fa); }
        .fpi-section-settings { background: linear-gradient(135deg,#7c3aed,#a78bfa); }
        .fpi-section-export { background: linear-gradient(135deg,#15803d,#22c55e); }
        .fpi-settings-panel {
            border-radius: 18px;
            padding: 14px 16px;
            background: #f8fafc;
            border: 1px solid #dbeafe;
            margin: 8px 0 14px 0;
        }
        .fpi-settings-panel, .fpi-settings-panel * { color: #0f172a !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_landing_login_panel_v103() -> None:
    """Főoldali belépő / Demo-Pro állapot panel."""
    mode_label = "PRO" if is_pro_mode() else "DEMO"
    mode_color = "#16A34A" if is_pro_mode() else "#2563EB"
    lic = st.session_state.get("license_status", {}) or {}
    club = lic.get("club_name") or lic.get("email") or ""

    st.markdown(
        f"""
        <div style="border-radius:22px;padding:18px 20px;background:#ffffff;border:1px solid #e5e7eb;
                    box-shadow:0 12px 28px rgba(15,23,42,.08);margin-bottom:18px;">
            <div style="display:flex;align-items:center;justify-content:space-between;gap:16px;">
                <div>
                    <div style="font-size:.82rem;font-weight:900;color:#64748b;letter-spacing:.06em;">HOZZÁFÉRÉS</div>
                    <div style="font-size:1.55rem;font-weight:900;color:#0f172a;margin-top:2px;">
                        Aktuális mód: <span style="color:{mode_color};">{mode_label}</span>
                    </div>
                    <div style="color:#64748b;font-size:.95rem;margin-top:4px;">
                        {"Pro hozzáférés aktív" if is_pro_mode() else f"Demo limit: max {DEMO_PLAYER_LIMIT} játékos · max {DEMO_WEEK_LIMIT} hét · max {DEMO_ROW_LIMIT} sor"}
                        {(" · " + html.escape(str(club))) if club else ""}
        </div>
        """,
        unsafe_allow_html=True,
    )

    if is_pro_mode():
        c1, c2 = st.columns([1, 4])
        with c1:
            if st.button("Kijelentkezés", use_container_width=True, key="landing_logout_license_v103"):
                st.session_state.pop("license_status", None)
                st.rerun()
        return

    with st.expander("🔐 Pro belépés / aktiválás", expanded=False):
        email = st.text_input("E-mail", value=st.session_state.get("user_email", ""), placeholder="nev@klub.hu", key="landing_license_email_v103")
        if email:
            st.session_state["user_email"] = email
        license_key = st.text_input("Aktiváló kód", type="password", help="A klubhoz kapott aktiváló kód.", key="landing_license_key_v103")
        if st.button("Pro aktiválása", use_container_width=True, key="landing_activate_license_v103"):
            result = validate_license_supabase(email, license_key)
            if result.get("ok"):
                st.session_state["license_status"] = result
                st.success("Pro hozzáférés aktiválva.")
                st.rerun()
            else:
                st.warning(result.get("message", "Sikertelen aktiválás."))



def render_fpi_landing_page_v100() -> None:
    _fpi_landing_css_v100()
    st.markdown(
        """
        <style>
        .fpi-v137-hero{border-radius:34px;padding:34px 38px;margin:8px 0 18px 0;background:radial-gradient(circle at 10% 8%,rgba(20,184,166,.20),transparent 30%),radial-gradient(circle at 90% 4%,rgba(37,99,235,.18),transparent 30%),linear-gradient(135deg,#ffffff 0%,#e0f2fe 54%,#ecfdf5 100%);border:1px solid #bfdbfe;box-shadow:0 26px 70px rgba(15,23,42,.16);color:#0f172a;}
        .fpi-v137-kicker{display:inline-block;padding:7px 13px;border-radius:999px;background:#ffffff;border:1px solid #93c5fd;color:#0f766e;font-weight:950;letter-spacing:.07em;font-size:.82rem;margin-bottom:12px;}
        .fpi-v137-title{font-size:3.25rem;line-height:.98;font-weight:980;letter-spacing:-.06em;color:#0f172a;margin:0 0 12px 0;max-width:980px;}
        .fpi-v137-sub{font-size:1.12rem;line-height:1.46;color:#334155;max-width:900px;margin-bottom:14px;}
        .fpi-v137-flow{display:flex;flex-wrap:wrap;gap:8px;margin-top:14px;}
        .fpi-v137-flow span{display:inline-block;padding:8px 12px;border-radius:999px;background:#ffffff;border:1px solid #dbeafe;color:#0f172a;font-weight:850;font-size:.88rem;}
        .fpi-v137-login{border-radius:24px;padding:18px 20px;background:#ffffff;border:1px solid #dbeafe;box-shadow:0 14px 34px rgba(15,23,42,.10);margin-bottom:18px;}
        .fpi-v137-login b{font-size:1.15rem;color:#0f172a;}.fpi-v137-login span{color:#475569;}
        .fpi-v137-kpi-grid{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:12px;margin:16px 0 6px 0;}
        .fpi-v137-kpi{border-radius:20px;padding:16px 17px;background:#ffffff;border:1px solid #dbeafe;box-shadow:0 12px 28px rgba(15,23,42,.08);}
        .fpi-v137-kpi strong{display:block;font-size:1.55rem;color:#0f172a;line-height:1;margin-bottom:5px;}.fpi-v137-kpi span{color:#475569;font-weight:750;}
        .fpi-v137-card{border-radius:24px;padding:21px 22px;background:#ffffff;border:1px solid #dbeafe;box-shadow:0 14px 34px rgba(15,23,42,.09);min-height:170px;}
        .fpi-v137-card h3{margin:0 0 8px 0;color:#0f172a;font-weight:950;}.fpi-v137-card p{color:#475569;line-height:1.42;margin:0;}
        .fpi-v137-section-title{font-size:1.28rem;font-weight:950;color:#0f172a;margin:22px 0 10px 0;}
        .stButton > button{border-radius:18px !important;min-height:54px !important;font-weight:950 !important;}
        .stDownloadButton > button{border-radius:18px !important;min-height:48px !important;font-weight:950 !important;}
        </style>
        <div class="fpi-v137-hero">
            <div class="fpi-v137-kicker">FOOTBALL PERFORMANCE INTELLIGENCE</div>
            <div class="fpi-v137-title">Vezetői riport 30 másodperc alatt.</div>
            <div class="fpi-v137-sub">GPS exportból – és opcionálisan taktikai PDF/Excel anyagokból – azonnal kapsz heti állapotképet, játékoskockázatot, referencia-összevetést, ellenfél-specifikus fókuszokat és mikrociklus-javaslatot.</div>
            <div class="fpi-v137-flow"><span>1. GPS / ZIP feltöltés</span><span>2. Heti kontextus</span><span>3. Tactical Pro+ input</span><span>4. Executive PDF</span></div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Belépés / Pro aktiválás – egyszerűen, a fő CTA-k fölött.
    render_landing_login_panel_v103()

    cta1, cta2, cta3 = st.columns([2.2, 1.2, 1.1])
    with cta1:
        if st.button("🚀 Vezetői riport készítése", use_container_width=True, type="primary", key="landing_go_clean_v137"):
            _fpi_set_page_v100("clean")
    with cta2:
        if st.button("📊 Haladó elemző app", use_container_width=True, key="landing_go_app_v137"):
            st.session_state["fpi_app_hub_seen_v137"] = False
            _fpi_set_page_v100("app")
    with cta3:
        if st.button("📚 Metodika", use_container_width=True, key="landing_method_v138"):
            _fpi_set_page_v100("method")

    st.markdown('<div class="fpi-v137-section-title">Mit kapsz a riportban?</div>', unsafe_allow_html=True)
    left_col, right_col = st.columns([1.05, 1.75])
    with left_col:
        st.markdown('<div class="fpi-v137-card" style="min-height:96px;margin-bottom:12px;"><h3>Heti állapotkép</h3><p>Readiness, terhelési kép, trendek és heti kockázati összefoglaló.</p></div>', unsafe_allow_html=True)
        st.markdown('<div class="fpi-v137-card" style="min-height:96px;"><h3>Játékos-kockázat</h3><p>Magas/közepes jelzések, túl- vagy alulterhelési mintázatok és figyelendő játékosok.</p></div>', unsafe_allow_html=True)
    with right_col:
        st.markdown("""
        <div class="fpi-v137-card" style="min-height:216px;">
            <h3>Tactical Pro+ Intelligence Engine</h3>
            <p><b>GPS + játékmodell + ellenfélanyag + heti periodizáció egy rendszerben.</b></p>
            <ul style="margin:10px 0 0 18px;color:#475569;line-height:1.46;padding-left:0;">
                <li><b>Ellenfél-specifikus edzésfókusz:</b> mire készüljön a stáb a héten.</li>
                <li><b>Taktikai prioritások:</b> presszing, rest defense, átmenetek, szélső játék.</li>
                <li><b>Játékosszintű veszélyek:</b> kulcsemberek, progresszorok, befejezők, gyenge láncszemek.</li>
                <li><b>Mikrociklus-terv:</b> GPS-állapothoz és meccstervhez igazított napi fókusz.</li>
            </ul>
            <p style="margin-top:11px;"><b>Nem csak adatokat mutat: edzői döntéseket strukturál.</b></p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="fpi-v137-section-title">Minta riportok</div>', unsafe_allow_html=True)
    m1, m2, m3, m4, m5 = st.columns(5)
    try:
        sample_exec = build_fpi_sample_pdf_bytes("executive")
        sample_full = build_fpi_sample_pdf_bytes("full")
        sample_gps = build_fpi_gps_only_sample_pdf_bytes()
        sample_own = build_fpi_own_team_profile_sample_pdf_bytes() if "build_fpi_own_team_profile_sample_pdf_bytes" in globals() else None
        sample_method = build_fpi_methodology_pdf_bytes_v143() if "build_fpi_methodology_pdf_bytes_v143" in globals() else None
    except Exception:
        sample_exec = sample_full = sample_gps = sample_own = sample_method = None
    with m1:
        if sample_exec: st.download_button("⬇️ Executive", sample_exec, "fpi_minta_executive_summary.pdf", "application/pdf", use_container_width=True, key="sample_exec_v143")
    with m2:
        if sample_gps: st.download_button("⬇️ GPS-only", sample_gps, "fpi_minta_gps_only_report.pdf", "application/pdf", use_container_width=True, key="sample_gps_v143")
    with m3:
        if sample_own: st.download_button("⬇️ Saját csapat", sample_own, "fpi_minta_sajat_csapat_profil.pdf", "application/pdf", use_container_width=True, key="sample_own_v143")
    with m4:
        if sample_full: st.download_button("⬇️ Full report", sample_full, "fpi_minta_full_report.pdf", "application/pdf", use_container_width=True, key="sample_full_v143")
    with m5:
        if sample_method: st.download_button("⬇️ Metodika", sample_method, "fpi_metodika.pdf", "application/pdf", use_container_width=True, key="sample_method_v143")



# =========================================================
# V139 - Session based user defaults fallback
# =========================================================
def _fpi_load_user_defaults_v113() -> Dict[str, object]:
    """Alapbeállítások betöltése jelenleg csak session_state-ből.
    Supabase/perzisztens tárolás később kerül bevezetésre.
    """
    try:
        data = st.session_state.get("fpi_user_defaults_v113", {})
        if isinstance(data, dict):
            return data.copy()
    except Exception:
        pass
    return {}


def _fpi_save_user_defaults_v113(data: Dict[str, object]) -> Tuple[bool, str]:
    """Alapbeállítások mentése ideiglenesen session_state-be.
    Ez nem tartós szerveroldali mentés; éles verzióban Supabase-be költöztethető.
    """
    try:
        st.session_state["fpi_user_defaults_v113"] = dict(data or {})
        return True, "Alapbeállítások elmentve erre a munkamenetre."
    except Exception as exc:
        return False, f"Nem sikerült menteni az alapbeállításokat: {exc}"



# =========================================================
# V140 - Clean workspace section header helper
# =========================================================
def _fpi_section_header_v113(title: str, subtitle: str = "", icon: str = "") -> None:
    """Egységes, olvasható szekciófejléc az import/export munkafolyamatban.
    A v137/v139 merge során a hívás bent maradt, a definíció kimaradt; ez javítja a NameError-t.
    """
    icon_map = {
        "gps": "📡",
        "settings": "⚙️",
        "tactical": "🧠",
        "export": "📄",
        "default": "▣",
    }
    badge = icon_map.get(str(icon or "").lower(), icon_map["default"])
    safe_title = html.escape(str(title or ""))
    safe_sub = html.escape(str(subtitle or ""))
    st.markdown(
        f"""
        <div style="border-radius:22px;padding:16px 19px;margin:18px 0 12px 0;
                    background:linear-gradient(135deg,#ffffff,#f0fdfa);
                    border:1px solid #bfdbfe;box-shadow:0 10px 28px rgba(15,23,42,.08);">
            <div style="display:flex;gap:10px;align-items:center;">
                <div style="font-size:1.45rem;line-height:1;">{badge}</div>
                <div>
                    <div style="font-size:1.22rem;font-weight:950;color:#0f172a;letter-spacing:-.02em;">{safe_title}</div>
                    <div style="font-size:.93rem;color:#475569;line-height:1.35;margin-top:3px;">{safe_sub}</div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# =========================================================
# V141 - Clean workspace hard NameError fallbacks
# =========================================================
def _fpi_idx_v113(options, value, default: int = 0) -> int:
    """Biztonságos selectbox-index segédfüggvény.
    A vezetői riport oldalon több választó használja; merge során kimaradt, ezért NameError-t okozott.
    """
    try:
        opts = list(options or [])
        return opts.index(value) if value in opts else int(default)
    except Exception:
        return int(default or 0)


def _fpi_simple_pdf_blocks_v141(text: str) -> Dict[str, list]:
    txt = (text or "").lower()
    kw = {
        "pressing": ["pressing", "presszing", "letámadás", "letamadas", "ppda"],
        "build_up": ["build-up", "build up", "labdakihozatal", "építkezés", "epitkezes"],
        "direct_play": ["direct", "direkt", "hosszú labda", "hosszu labda"],
        "transition_attack": ["transition", "átmenet", "atmenet", "kontra", "counter"],
        "transition_defense": ["rest defense", "visszarendeződés", "visszarendezodes", "átmeneti védekezés"],
        "wide_play": ["wide", "szélső", "szelso", "beadás", "beadas", "flank", "wing"],
        "central_play": ["half-space", "félterület", "felterulet", "centrum", "central"],
        "set_pieces": ["set piece", "pontrúgás", "pontrugas", "szöglet", "szoglet", "free kick"],
        "chance_creation": ["xg", "helyzet", "shot", "lövés", "loves", "box entry"],
        "key_players": ["key player", "kulcsjátékos", "kulcsjatekos", "veszélyes játékos", "creator", "playmaker"],
        "weakness_risk": ["weakness", "gyengeség", "gyengeseg", "sebezhető", "risk", "kockázat"],
    }
    blocks = {}
    for k, words in kw.items():
        hits = [w for w in words if w in txt]
        if hits:
            blocks[k] = hits[:6]
    return blocks


def _fpi_tactical_pdf_insights(text: str) -> Dict[str, object]:
    """Korai, könnyű PDF-insight fallback a clean import/export oldalhoz.
    A részletes legacy függvény a fájl későbbi részében van, de a page routing előtt még nem létezik.
    """
    text = text or ""
    blocks = _fpi_simple_pdf_blocks_v141(text)
    topics = [{"Téma": k, "Találatok": ", ".join(v)} for k, v in blocks.items()]
    return {
        "blocks": blocks,
        "topics": topics,
        "sportsbase_findings": topics[:8],
        "sportsbase_lines": [],
        "raw_text_len": len(text),
    }


def _fpi_safe_tactical_pdf_insights_v105(text: str, uploaded: bool = False, pages: Optional[list] = None) -> Dict[str, object]:
    data = _fpi_tactical_pdf_insights(text or "")
    data["pdf_uploaded"] = bool(uploaded or text)
    data["pdf_pages"] = len(pages or [])
    data["raw_text_len"] = len(text or "")
    return data


def _fpi_safe_merge_tactical_pdf_insights_v104(own: Optional[Dict[str, object]], opp: Optional[Dict[str, object]]) -> Dict[str, object]:
    own = own or {}; opp = opp or {}
    blocks = {}
    for src in (own.get("blocks") or {}, opp.get("blocks") or {}):
        if isinstance(src, dict):
            for k, v in src.items():
                blocks.setdefault(k, [])
                if isinstance(v, list):
                    blocks[k].extend(v)
                elif v:
                    blocks[k].append(v)
    topics = []
    for src in (own.get("topics") or [], opp.get("topics") or [], own.get("sportsbase_findings") or [], opp.get("sportsbase_findings") or []):
        if isinstance(src, list):
            topics.extend(src)
    return {
        "blocks": blocks,
        "topics": topics[:20],
        "sportsbase_findings": topics[:12],
        "sportsbase_lines": (own.get("sportsbase_lines") or []) + (opp.get("sportsbase_lines") or []),
        "raw_text_len": int(own.get("raw_text_len", 0) or 0) + int(opp.get("raw_text_len", 0) or 0),
        "pdf_uploaded": bool(own.get("pdf_uploaded") or opp.get("pdf_uploaded")),
        "pdf_pages": int(own.get("pdf_pages", 0) or 0) + int(opp.get("pdf_pages", 0) or 0),
    }


def _fpi_safe_build_adaptive_plan_v104(gps_context: Dict[str, object], tactical_ctx: Dict[str, object]) -> Dict[str, object]:
    try:
        return _fpi_build_adaptive_match_training_plan(gps_context or {}, tactical_ctx or {})
    except Exception as exc:
        return {
            "analysis_level": "GPS / alap taktikai fallback",
            "plan_a": "KIE - stabil, kontrollált heti terv",
            "risks": ["A taktikai terv fallback módban készült.", str(exc)[:120]],
            "md_plan": [("MD-4", "Volumen + saját játékmodell", "Stabil heti terhelés."), ("MD-3", "Intenzitás", "Kontrollált HSR/sprint inger."), ("MD-2", "Taktikai nap", "Ellenfél-specifikus fókusz."), ("MD-1", "Aktiváció", "Frissítés és pontrúgások.")],
            "player_focus": [], "tactical_findings": [], "team_comparison": {}, "strategy_framework": {}, "plan_b": ""
        }


def _fpi_safe_tactical_parse_team_excel_v107(df2, mapping=None) -> Dict[str, float]:
    """Korai fallback: ha van taktikai Excel, legalább a numerikus oszlopátlagokat adja vissza."""
    try:
        if not isinstance(df2, pd.DataFrame) or df2.empty:
            return {}
        out = {}
        for c in df2.columns:
            vals = pd.to_numeric(df2[c], errors="coerce").dropna()
            if len(vals):
                key = str(c).strip().lower().replace(" ", "_")[:40]
                out[key] = float(vals.mean())
        return out
    except Exception:
        return {}


def _fpi_safe_tactical_parse_player_excel_v107(df2, mapping=None) -> Dict[str, pd.DataFrame]:
    try:
        if isinstance(df2, pd.DataFrame) and not df2.empty:
            return {"raw": df2.copy()}
    except Exception:
        pass
    return {}


# =========================================================
# V142 - Clean workspace multi tactical upload helpers
# =========================================================
def _fpi_clean_read_table_file_v142(uploaded_file) -> Tuple[pd.DataFrame, str]:
    """Excel/CSV feltöltés biztonságos beolvasása a vezetői riport oldalhoz.
    Nem igényel később definiált mapper függvényeket, ezért a clean workspace korai hívásakor is működik.
    """
    try:
        if uploaded_file is None:
            return pd.DataFrame(), ""
        name = str(getattr(uploaded_file, "name", "file"))
        b = uploaded_file.getvalue()
        low = name.lower()
        if low.endswith(".csv"):
            try:
                return pd.read_csv(io.BytesIO(b)), name
            except Exception:
                return pd.read_csv(io.BytesIO(b), sep=";"), name
        xls = pd.ExcelFile(io.BytesIO(b))
        best_df, best_sheet, best_score = pd.DataFrame(), "", -1
        for sh in xls.sheet_names:
            try:
                df = pd.read_excel(io.BytesIO(b), sheet_name=sh)
                df = df.dropna(how="all")
                score = int(len(df)) + int(sum(pd.to_numeric(df[c], errors="coerce").notna().sum() for c in df.columns))
                sh_norm = str(sh).lower()
                if any(x in sh_norm for x in ["data", "stat", "player", "team", "main"]):
                    score += 100
                if score > best_score:
                    best_df, best_sheet, best_score = df, sh, score
            except Exception:
                continue
        return best_df, f"{name} / {best_sheet}" if best_sheet else name
    except Exception as e:
        return pd.DataFrame(), f"hiba: {e}"


def _fpi_clean_merge_team_excels_v142(files: List[object]) -> Tuple[Dict[str, float], List[Dict[str, object]]]:
    """Több csapat Excel/CSV összeolvasása. A numerikus oszlopok átlagát adja vissza.
    A riportlogika később ezekből tud csapatszintű profilt építeni.
    """
    frames, diag = [], []
    for f in files or []:
        df, src = _fpi_clean_read_table_file_v142(f)
        ok = isinstance(df, pd.DataFrame) and not df.empty
        diag.append({"Fájl": str(getattr(f, "name", src)), "Forrás/lap": src, "Státusz": "OK" if ok else "nem olvasható", "Sor": int(len(df)) if ok else 0, "Oszlop": int(len(df.columns)) if ok else 0})
        if ok:
            df = df.copy()
            df["_source_file"] = str(getattr(f, "name", src))
            frames.append(df)
    if not frames:
        return {}, diag
    merged = pd.concat(frames, ignore_index=True, sort=False)
    metrics = _fpi_safe_tactical_parse_team_excel_v107(merged, None) or {}
    return metrics, diag


def _fpi_clean_merge_player_excels_v142(files: List[object]) -> Tuple[Dict[str, pd.DataFrame], List[Dict[str, object]]]:
    """Több játékos Excel/CSV összeolvasása egy közös raw táblába."""
    frames, diag = [], []
    for f in files or []:
        df, src = _fpi_clean_read_table_file_v142(f)
        ok = isinstance(df, pd.DataFrame) and not df.empty
        diag.append({"Fájl": str(getattr(f, "name", src)), "Forrás/lap": src, "Státusz": "OK" if ok else "nem olvasható", "Sor": int(len(df)) if ok else 0, "Oszlop": int(len(df.columns)) if ok else 0})
        if ok:
            df = df.copy()
            df["_source_file"] = str(getattr(f, "name", src))
            frames.append(df)
    if not frames:
        return {}, diag
    merged = pd.concat(frames, ignore_index=True, sort=False)
    return {"raw": merged}, diag


def _fpi_clean_tactical_diagnostics_box_v142(title: str, pdf_state: Dict[str, object], team_diag: List[Dict[str, object]], player_diag: List[Dict[str, object]]) -> None:
    try:
        pdf_files = pdf_state.get("files", []) or []
        pdf_line = f"PDF: {len(pdf_files)} fájl, {int(pdf_state.get('page_count', 0) or 0)} oldal, {int(pdf_state.get('chars', 0) or 0)} karakter"
        team_ok = sum(1 for x in team_diag if x.get("Státusz") == "OK")
        player_ok = sum(1 for x in player_diag if x.get("Státusz") == "OK")
        st.caption(f"{title} – {pdf_line} | Csapat Excel: {team_ok}/{len(team_diag)} OK | Játékos Excel: {player_ok}/{len(player_diag)} OK")
    except Exception:
        pass


def _fpi_clean_tactical_import_v102(gps_context: Dict[str, object]) -> Dict[str, object]:
    """V142: letisztult, többfájlos taktikai import a vezetői riport oldalhoz.
    Kezeli: saját/ellenfél több PDF, saját/ellenfél csapat Excel(ek), saját/ellenfél játékos Excel(ek).
    Ha nincs taktikai anyag, GPS-only contextet ad vissza hiba nélkül.
    """
    _fpi_section_header_v113(
        "3. Taktikai input",
        "Opcionális, de prémium mód: több saját és ellenfél PDF, csapat Excel és játékos Excel is feltölthető. Az FPI ezeket egy közös taktikai tudásbázissá fűzi össze.",
        "tactical",
    )
    with st.expander("🧠 Taktikai anyagok – saját csapat és ellenfél", expanded=False):
        st.markdown(
            """
            <div class="tactical-readable-box">
            <b>Mit lehet feltölteni?</b><br/>
            • Saját és ellenfél taktikai PDF-ekből akár több fájlt is.<br/>
            • Saját és ellenfél csapatszintű Excel/CSV anyagot.<br/>
            • Saját és ellenfél játékosszintű Excel/CSV anyagot.<br/>
            <span style="color:#475569!important;">A több fájlból az app egy közös saját / ellenfél profilt épít.</span>
            </div>
            """,
            unsafe_allow_html=True,
        )
        left, right = st.columns(2)
        with left:
            st.markdown("#### Saját csapat")
            _, own_state = _fpi_pdf_uploader_v92("Saját taktikai PDF-ek", "own", "clean_own_tactical_pdfs_v142")
            own_team_files = st.file_uploader("Saját csapat Excel / CSV", type=["xlsx", "xls", "csv"], accept_multiple_files=True, key="clean_own_team_excels_v142")
            own_player_files = st.file_uploader("Saját játékos Excel / CSV", type=["xlsx", "xls", "csv"], accept_multiple_files=True, key="clean_own_player_excels_v142")
        with right:
            st.markdown("#### Ellenfél")
            _, opp_state = _fpi_pdf_uploader_v92("Ellenfél taktikai PDF-ek", "opp", "clean_opp_tactical_pdfs_v142")
            opp_team_files = st.file_uploader("Ellenfél csapat Excel / CSV", type=["xlsx", "xls", "csv"], accept_multiple_files=True, key="clean_opp_team_excels_v142")
            opp_player_files = st.file_uploader("Ellenfél játékos Excel / CSV", type=["xlsx", "xls", "csv"], accept_multiple_files=True, key="clean_opp_player_excels_v142")

        own_team_metrics, own_team_diag = _fpi_clean_merge_team_excels_v142(own_team_files or [])
        opp_team_metrics, opp_team_diag = _fpi_clean_merge_team_excels_v142(opp_team_files or [])
        own_player_tables, own_player_diag = _fpi_clean_merge_player_excels_v142(own_player_files or [])
        opp_player_tables, opp_player_diag = _fpi_clean_merge_player_excels_v142(opp_player_files or [])

        _fpi_clean_tactical_diagnostics_box_v142("Saját", own_state, own_team_diag, own_player_diag)
        _fpi_clean_tactical_diagnostics_box_v142("Ellenfél", opp_state, opp_team_diag, opp_player_diag)

        with st.expander("📋 Feltöltött taktikai fájlok ellenőrzése", expanded=False):
            diag_rows = []
            for label, rows in [("Saját csapat Excel", own_team_diag), ("Saját játékos Excel", own_player_diag), ("Ellenfél csapat Excel", opp_team_diag), ("Ellenfél játékos Excel", opp_player_diag)]:
                for r in rows:
                    rr = dict(r); rr["Típus"] = label; diag_rows.append(rr)
            if diag_rows:
                st.dataframe(pd.DataFrame(diag_rows), use_container_width=True)
            else:
                st.caption("Nincs Excel/CSV feltöltés. PDF vagy GPS-only módban is készíthető riport.")

    # A PDF context builder továbbra is a közös own/opp PDF state-ből dolgozik.
    try:
        ctx = _fpi_build_pdf_only_context_from_session_v87(gps_context or {})
    except Exception:
        ctx = None

    has_excel = bool(own_team_metrics or opp_team_metrics or own_player_tables or opp_player_tables)
    has_pdf = bool((own_state or {}).get("has_files") or (opp_state or {}).get("has_files") or (ctx and _fpi_has_tactical_signal_v95(ctx)))

    if isinstance(ctx, dict):
        # Frissítsük a PDF-ből épített contextet az aktuális Excel inputokkal.
        ctx.setdefault("own", {})["team_metrics"] = own_team_metrics
        ctx.setdefault("own", {})["player_tables"] = own_player_tables
        ctx.setdefault("opponent", {})["team_metrics"] = opp_team_metrics
        ctx.setdefault("opponent", {})["player_tables"] = opp_player_tables
        ctx["own_team_metrics"] = own_team_metrics
        ctx["opp_team_metrics"] = opp_team_metrics
        ctx["own_player_tables"] = own_player_tables
        ctx["opp_player_tables"] = opp_player_tables
        ctx["has_own_team_excel"] = bool(own_team_metrics)
        ctx["has_opp_team_excel"] = bool(opp_team_metrics)
        ctx["has_own_player_excel"] = bool(own_player_tables)
        ctx["has_opp_player_excel"] = bool(opp_player_tables)
        ctx["tactical_upload_diag"] = {
            "own_team": own_team_diag, "own_player": own_player_diag,
            "opp_team": opp_team_diag, "opp_player": opp_player_diag,
        }
        if has_excel:
            ctx["analysis_level_label"] = "Full Intelligence – GPS + PDF + többfájlos taktikai Excel"
            st.session_state["tactical_pro_context"] = ctx
        if _fpi_has_tactical_signal_v95(ctx) or has_excel:
            return ctx

    if has_excel:
        # Excel-only taktikai context: PDF nélkül is működjön.
        own_pdf_insights = {"blocks": {}, "topics": [], "sportsbase_findings": [], "sportsbase_lines": [], "raw_text_len": 0, "pdf_uploaded": False, "pdf_pages": 0}
        opp_pdf_insights = {"blocks": {}, "topics": [], "sportsbase_findings": [], "sportsbase_lines": [], "raw_text_len": 0, "pdf_uploaded": False, "pdf_pages": 0}
        tactical_ctx_for_plan = {
            "analysis_level_label": "Full Intelligence – GPS + többfájlos taktikai Excel",
            "pdf_insights": {"blocks": {}, "topics": [], "sportsbase_findings": [], "sportsbase_lines": [], "raw_text_len": 0},
            "team_metrics": opp_team_metrics,
            "player_tables": opp_player_tables,
            "own": {"pdf_insights": own_pdf_insights, "team_metrics": own_team_metrics, "player_tables": own_player_tables},
            "opponent": {"pdf_insights": opp_pdf_insights, "team_metrics": opp_team_metrics, "player_tables": opp_player_tables},
        }
        plan = _fpi_safe_build_adaptive_plan_v104(gps_context or {}, tactical_ctx_for_plan)
        ctx = _build_tactical_executive_context(gps_context or {}, tactical_ctx_for_plan, plan)
        ctx["analysis_level_label"] = "Full Intelligence – GPS + többfájlos taktikai Excel"
        ctx["own_team_metrics"] = own_team_metrics
        ctx["opp_team_metrics"] = opp_team_metrics
        ctx["own_player_tables"] = own_player_tables
        ctx["opp_player_tables"] = opp_player_tables
        ctx["has_own_team_excel"] = bool(own_team_metrics)
        ctx["has_opp_team_excel"] = bool(opp_team_metrics)
        ctx["has_own_player_excel"] = bool(own_player_tables)
        ctx["has_opp_player_excel"] = bool(opp_player_tables)
        ctx["tactical_upload_diag"] = {"own_team": own_team_diag, "own_player": own_player_diag, "opp_team": opp_team_diag, "opp_player": opp_player_diag}
        st.session_state["tactical_pro_context"] = ctx
        return ctx

    # Biztonságos GPS-only fallback
    return {
        "analysis_level_label": "GPS-only",
        "pdf_insights": {"blocks": {}, "topics": [], "sportsbase_findings": [], "sportsbase_lines": [], "raw_text_len": 0},
        "team_metrics": {},
        "player_tables": {},
        "own": {"pdf_insights": {}, "team_metrics": {}, "player_tables": {}},
        "opponent": {"pdf_insights": {}, "team_metrics": {}, "player_tables": {}},
        "has_own_pdf": False,
        "has_opp_pdf": False,
    }

def render_fpi_clean_workspace_v101() -> None:
    """V137: fókuszált import/export munkafolyamat.
    A fő felhasználói út: GPS/ZIP import -> minimális heti kontextus -> opcionális taktika -> PDF export.
    A mapper, alapbeállítások, kapuslogika, játékosszűrés és napokra bontott mikrociklus alapból csukva vannak.
    """
    _fpi_landing_css_v100()
    _fpi_mapper_contrast_css_v109()
    # V139: v118 patch opcionális; ha nincs definiálva, a v119 patch váltja ki.
    if "_fpi_apply_v118_final_light_controls_patch" in globals():
        _fpi_apply_v118_final_light_controls_patch()
    _fpi_apply_v119_all_light_readable_patch()
    user_defaults_clean = _fpi_load_user_defaults_v113()

    st.markdown(
        """
        <style>
        .fpi-work-hero{border-radius:28px;padding:26px 30px;margin:8px 0 18px 0;background:linear-gradient(135deg,#ffffff,#e0f2fe 58%,#ecfdf5);border:1px solid #bfdbfe;box-shadow:0 18px 48px rgba(15,23,42,.13);color:#0f172a;}
        .fpi-work-kicker{display:inline-block;padding:6px 12px;border-radius:999px;background:#f0fdfa;border:1px solid #99f6e4;color:#0f766e;font-size:.78rem;font-weight:950;letter-spacing:.06em;margin-bottom:10px;}
        .fpi-work-title{font-size:2.35rem;line-height:1.04;font-weight:980;letter-spacing:-.045em;color:#0f172a;margin:0 0 8px 0;}
        .fpi-work-sub{color:#475569;font-size:1.02rem;line-height:1.45;max-width:980px;}
        .fpi-workflow-box{border-radius:20px;padding:14px 18px;background:linear-gradient(135deg,#0f766e,#2563eb);color:#ffffff;box-shadow:0 14px 34px rgba(15,23,42,.16);font-size:1.03rem;font-weight:950;letter-spacing:-.01em;text-align:center;margin-top:2px;}
        .fpi-workflow-box span{opacity:.95;font-weight:850;}
        .fpi-step-card{border-radius:22px;padding:15px 17px;background:#ffffff;border:1px solid #dbeafe;box-shadow:0 10px 28px rgba(15,23,42,.08);margin-bottom:12px;}
        .fpi-step-card b{color:#0f172a;font-size:1.06rem;}.fpi-step-card span{color:#475569;}
        .fpi-quick-panel{border-radius:22px;padding:18px 20px;background:#ffffff;border:1px solid #dbeafe;box-shadow:0 14px 34px rgba(15,23,42,.09);margin:8px 0 14px 0;}
        .fpi-quick-panel h3{margin:0 0 10px 0;color:#0f172a;font-weight:950;}.fpi-quick-panel p{color:#475569;margin:0 0 12px 0;}
        div[data-testid="stExpander"]{border-radius:18px !important;border:1px solid #dbeafe !important;background:#ffffff !important;}
        div[data-testid="stExpander"] *{color:#0f172a !important;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    nav1, nav2, nav3, nav4 = st.columns([1.05, 1.25, 1.15, 3.7])
    with nav1:
        if st.button("← Főoldal", use_container_width=True, key="clean_back_landing_v137"):
            _fpi_set_page_v100("landing")
    with nav2:
        if st.button("📊 Haladó app", use_container_width=True, key="clean_go_full_app_v137"):
            st.session_state["fpi_app_hub_seen_v137"] = False
            _fpi_set_page_v100("app")
    with nav3:
        if st.button("📚 Metodika", use_container_width=True, key="clean_go_method_v138"):
            _fpi_set_page_v100("method")
    with nav4:
        st.markdown('<div class="fpi-workflow-box"><span>Fő munkafolyamat:</span> import → heti kontextus → taktikai input → Executive PDF</div>', unsafe_allow_html=True)

    st.markdown(
        """
        <div class="fpi-work-hero">
            <div class="fpi-work-kicker">VEZETŐI RIPORT KÉSZÍTÉSE</div>
            <div class="fpi-work-title">Feltöltésből döntéstámogató PDF.</div>
            <div class="fpi-work-sub">A napi használathoz elég a GPS fájl vagy ZIP, a heti edzésszám, a meccskontextus és – ha van – a taktikai anyag. A többi beállítás lenyitható, de nem kötelező.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # 1. GPS import
    _fpi_section_header_v113(
        "1. GPS import",
        "Brainsports, PlayerTek, Polar Team Pro, Catapult és egyéb Excel/CSV exportok. Egy fájl, több fájl vagy ZIP is használható.",
        "gps",
    )

    provider_choice_v143 = st.selectbox(
        "GPS rendszer",
        FPI_SUPPORTED_GPS_PROVIDERS_V143,
        index=0,
        key="clean_provider_v143",
        help="Automatikus módban az FPI a fejléc és a fájlszerkezet alapján azonosítja a rendszert.",
    )

    with st.expander("📥 GPS fájlok feltöltése", expanded=True):
        st.caption(
            "A fájlok tartalmazhatnak egyetlen edzést/meccset vagy teljes időszakot. "
            "Több Excel/CSV egyszerre is kijelölhető, illetve ZIP-ben több tíz vagy száz fájl is feltölthető."
        )
        gps_a, gps_b, gps_c = st.columns(3)
        with gps_a:
            training_uploads_v143 = st.file_uploader(
                "Edzésfájlok",
                type=["xlsx", "xls", "xlsm", "csv", "zip"],
                accept_multiple_files=True,
                key="clean_training_uploads_v143",
                help="Az itt feltöltött fájlok sorai automatikusan Edzés típusúak lesznek.",
            )
        with gps_b:
            match_uploads_v143 = st.file_uploader(
                "Meccsfájlok",
                type=["xlsx", "xls", "xlsm", "csv", "zip"],
                accept_multiple_files=True,
                key="clean_match_uploads_v143",
                help="Az itt feltöltött fájlok sorai automatikusan Meccs típusúak lesznek.",
            )
        with gps_c:
            mixed_uploads_v143 = st.file_uploader(
                "Vegyes / teljes időszak",
                type=["xlsx", "xls", "xlsm", "csv", "zip"],
                accept_multiple_files=True,
                key="clean_mixed_uploads_v143",
                help="Olyan export, amelyben edzések és meccsek együtt szerepelnek, vagy a fájl maga tartalmazza a típust.",
            )

        demo_col, template_col = st.columns([1, 1.4])
        with demo_col:
            no_real_upload_v143 = not (training_uploads_v143 or match_uploads_v143 or mixed_uploads_v143)
            use_demo_clean = st.toggle(
                "Mintaadat használata",
                value=no_real_upload_v143,
                key="clean_use_demo_v143",
            )
        with template_col:
            template_bytes_clean = create_sample_input_template_bytes()
            if template_bytes_clean is not None:
                st.download_button(
                    "⬇️ Általános Excel sablon",
                    data=template_bytes_clean,
                    file_name="performance_input_sablon.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key="clean_template_download_v143",
                )

    raw_df_clean = None
    selected_sheet_clean = "Mintaadat"
    has_any_upload_v143 = bool(training_uploads_v143 or match_uploads_v143 or mixed_uploads_v143)

    if not has_any_upload_v143 and not use_demo_clean:
        st.info("Tölts fel legalább egy GPS fájlt, vagy kapcsold be a mintaadatot.")
        st.stop()

    if use_demo_clean and not has_any_upload_v143:
        raw_df_clean = build_demo_performance_data()
    else:
        try:
            gps_raw_sheets_v143, gps_report_v143, clean_sig_v143 = _fpi_read_many_gps_files_v143(
                training_uploads_v143,
                match_uploads_v143,
                mixed_uploads_v143,
                provider_choice_v143,
            )
            if st.session_state.get("clean_active_upload_signature_v105") != clean_sig_v143:
                st.session_state.pop("clean_mapped_df_override_v105", None)
                st.session_state.pop("clean_manual_mapping_v105", None)
                st.session_state["clean_active_upload_signature_v105"] = clean_sig_v143

            with st.expander("📦 Importellenőrzés", expanded=False):
                if gps_report_v143 is not None and not gps_report_v143.empty:
                    st.dataframe(gps_report_v143, use_container_width=True, hide_index=True)
                    ok_count = int((gps_report_v143["Státusz"] == "OK").sum()) if "Státusz" in gps_report_v143.columns else 0
                    mapper_count = int((gps_report_v143["Státusz"] == "MAPPER").sum()) if "Státusz" in gps_report_v143.columns else 0
                    error_count = int((gps_report_v143["Státusz"] == "HIBA").sum()) if "Státusz" in gps_report_v143.columns else 0
                    st.caption(f"Automatikus: {ok_count} | Smart Mapper: {mapper_count} | Hiba: {error_count}")

            if not gps_raw_sheets_v143:
                st.error("Nem sikerült értelmezhető GPS adatot beolvasni.")
                st.stop()

            sheets_clean = prepare_uploaded_sheets(gps_raw_sheets_v143)
            sheet_names_clean = list(sheets_clean.keys())
            preferred_names = ["Összes releváns adatlap", "Data", "Összegzés"]
            default_sheet_idx_v143 = 0
            for preferred in preferred_names:
                if preferred in sheet_names_clean:
                    default_sheet_idx_v143 = sheet_names_clean.index(preferred)
                    break
            selected_sheet_clean = st.selectbox(
                "Feldolgozott adatforrás",
                sheet_names_clean,
                index=default_sheet_idx_v143,
                key="clean_sheet_select_v143",
            )
            raw_df_clean = sheets_clean[selected_sheet_clean]
        except Exception as exc:
            st.error(f"Nem sikerült beolvasni a GPS adatokat: {exc}")
            st.stop()

    # Standardizálás + mapper alapból csukva
    if "clean_mapped_df_override_v105" in st.session_state and isinstance(st.session_state["clean_mapped_df_override_v105"], pd.DataFrame):
        df_clean = st.session_state["clean_mapped_df_override_v105"].copy()
        mapping_clean = st.session_state.get("clean_manual_mapping_v105", {})
        missing_clean = []
    else:
        df_clean, mapping_clean, missing_clean = standardize_dataframe(raw_df_clean)

    with st.expander("🧭 GPS Smart Mapper / oszlopfelismerés", expanded=False):
        if missing_clean:
            st.warning("Nem sikerült automatikusan felismerni minden kötelező oszlopot. Állítsd be itt kézzel.")
            st.write("Hiányzó mezők:", ", ".join(missing_clean))
        render_mapping_score(mapping_clean)
        st.dataframe(enhanced_mapping_quality_df(raw_df_clean, mapping_clean), use_container_width=True, hide_index=True)
        st.markdown("#### Gyors kézi mapping")
        cols_clean_map = [""] + sorted([str(c) for c in raw_df_clean.columns], key=lambda x: x.lower())
        editable_fields_clean = ["player_name", "session_type", "start_time", "duration", "match_minutes", "total_distance", "distance_per_min", "max_speed", "sprints", "speed_zone_4", "speed_zone_5", "training_load", "acc_high", "dec_high", "high_efforts"]
        manual_clean = dict(mapping_clean or {})
        grid_clean = st.columns(3)
        for i, field in enumerate(editable_fields_clean):
            default = manual_clean.get(field) or ""
            with grid_clean[i % 3]:
                val = st.selectbox(mapper_label(field), cols_clean_map, index=cols_clean_map.index(default) if default in cols_clean_map else 0, key=f"clean_gps_map_{field}_v137")
            manual_clean[field] = val or None
        if st.button("✅ Gyors mapping alkalmazása", use_container_width=True, key="clean_apply_mapping_v137"):
            try:
                mapped_clean = apply_manual_mapping(raw_df_clean, manual_clean)
                mapped_clean = normalize_combined_fields(mapped_clean, manual_clean)
                mapped_clean = derive_missing_columns(mapped_clean)
                st.session_state["clean_mapped_df_override_v105"] = mapped_clean
                st.session_state["clean_manual_mapping_v105"] = manual_clean
                st.success("Mapping alkalmazva.")
                st.rerun()
            except Exception as e:
                st.error(f"Mapping alkalmazási hiba: {e}")
        if st.button("Haladó mapping a teljes appban", key="clean_advanced_mapping_full_v137"):
            _fpi_set_page_v100("app")

    if missing_clean:
        st.stop()

    df_clean = add_position_group(df_clean)
    with st.expander("🧤 Választók: kapusok / posztlogika", expanded=False):
        df_clean = render_keeper_controls_and_apply(df_clean)
    if df_clean.empty or "week" not in df_clean.columns:
        st.warning("Nincs elemzésre alkalmas hétadat.")
        st.stop()

    weeks_clean = sorted(df_clean["week"].dropna().astype(str).unique().tolist(), key=_fpi_week_sort_key_v99)
    players_clean = sorted(df_clean["player_name"].dropna().astype(str).unique().tolist()) if "player_name" in df_clean.columns else []

    # 2. Heti kontextus és saját csapat heti profil
    _fpi_section_header_v113("2. Heti kontextus", "Csak a legfontosabb kérdések: ellenfél, meccsnap, referencia, heti edzésszám és a saját csapat heti alapprofilja.", "settings")
    st.markdown('<div class="fpi-quick-panel"><h3>Gyors beállítások</h3><p>Az alapfelállás, védekezési blokk és fő támadási út az adott hétre / ellenfélre is értelmezhető, ezért marad nyitva.</p></div>', unsafe_allow_html=True)

    q1, q2, q3, q4 = st.columns(4)
    with q1:
        opponent_clean = st.text_input("Ellenfél", value=st.session_state.get("fpi_match_opponent_v94", user_defaults_clean.get("opponent", "")), key="clean_opponent_v137")
    with q2:
        match_date_clean = st.date_input("Meccsnap", value=st.session_state.get("fpi_match_date_v94", pd.Timestamp.today().date()), key="clean_match_date_v137")
    with q3:
        reference_age_clean = st.selectbox("Korosztály", FPI_REFERENCE_AGE_OPTIONS_V112, index=_fpi_idx_v113(FPI_REFERENCE_AGE_OPTIONS_V112, user_defaults_clean.get("reference_age", "Felnőtt"), 0), key="clean_ref_age_v137")
    with q4:
        reference_level_clean = st.selectbox("Szint", FPI_REFERENCE_LEVEL_OPTIONS_V112, index=_fpi_idx_v113(FPI_REFERENCE_LEVEL_OPTIONS_V112, user_defaults_clean.get("reference_level", "NB II"), 1), key="clean_ref_level_v137")

    q5, q6, q7, q8 = st.columns(4)
    with q5:
        playmodel_profile_clean = st.selectbox("Játékmodell", FPI_PLAYMODEL_OPTIONS_V112, index=_fpi_idx_v113(FPI_PLAYMODEL_OPTIONS_V112, user_defaults_clean.get("playmodel_profile", "Kiegyensúlyozott"), 4), key="clean_playmodel_profile_v137")
    with q6:
        n_train_clean = st.number_input("Heti edzések száma", min_value=0, max_value=7, value=int(user_defaults_clean.get("training_days", 4)), step=1, key="clean_n_train_v137")
    with q7:
        week_type_clean = st.selectbox("Heti cél", FPI_COACH_WEEK_OPTIONS_V112, index=_fpi_idx_v113(FPI_COACH_WEEK_OPTIONS_V112, user_defaults_clean.get("coach_week_type", "Fenntartó hét"), 1), key="clean_week_type_v137")
    with q8:
        selected_playstyle_clean = st.selectbox("Riport játékmodell", FPI_PLAYMODEL_OPTIONS_V112, index=_fpi_idx_v113(FPI_PLAYMODEL_OPTIONS_V112, user_defaults_clean.get("playmodel_profile", "Kiegyensúlyozott"), 4), key="clean_playstyle_v137")

    oc1, oc2, oc3 = st.columns(3)
    with oc1:
        own_formation_clean = st.selectbox("Saját alapfelállás erre a hétre", FPI_FORMATION_OPTIONS_V132, index=_fpi_idx_v113(FPI_FORMATION_OPTIONS_V132, user_defaults_clean.get("own_formation", "4-2-3-1"), 0), key="clean_own_formation_v137")
    with oc2:
        own_block_clean = st.selectbox("Védekezési blokk erre a hétre", FPI_BLOCK_OPTIONS_V132, index=_fpi_idx_v113(FPI_BLOCK_OPTIONS_V132, user_defaults_clean.get("own_block", "Középső blokk"), 1), key="clean_own_block_v137")
    with oc3:
        own_attack_route_clean = st.selectbox("Fő támadási út erre a hétre", FPI_ATTACK_ROUTE_OPTIONS_V132, index=_fpi_idx_v113(FPI_ATTACK_ROUTE_OPTIONS_V132, user_defaults_clean.get("own_attack_route", "Vegyes"), 4), key="clean_own_attack_route_v137")

    match_week_clean = _fpi_iso_week_from_date_v94(match_date_clean)
    default_idx_clean = weeks_clean.index(match_week_clean) if match_week_clean in weeks_clean else (len(weeks_clean)-1 if weeks_clean else 0)
    w1, w2 = st.columns([1, 2.5])
    with w1:
        selected_week_clean = st.selectbox("Elemzett hét", weeks_clean, index=default_idx_clean, format_func=week_label_short, key="clean_week_v137")
    with w2:
        st.caption(f"Aktív referencia: {reference_age_clean} / {reference_level_clean} / játékosonkénti poszt / {playmodel_profile_clean}")

    selected_players_clean = players_clean
    cycle_days_clean = 7
    n_rest_clean = max(0, int(cycle_days_clean) - int(n_train_clean) - 1)
    md_match_day_clean = "MD"
    session_plan_clean = []

    with st.expander("⚙️ Alapbeállítások mentése és haladó heti tervezés", expanded=False):
        _storage_note = st.session_state.get("fpi_defaults_storage_v117", "még nincs mentett alapbeállítás")
        st.caption(f"Mentés helye: {_storage_note}. A mentés gombnyomásra történik, nem automatikusan.")
        if st.session_state.get("fpi_defaults_storage_warning_v117"):
            st.caption(st.session_state.get("fpi_defaults_storage_warning_v117"))
        if st.button("💾 Alapbeállítás mentése ehhez a belépéshez", use_container_width=True, key="clean_save_defaults_v137"):
            ok_save, msg_save = _fpi_save_user_defaults_v113({
                "opponent": opponent_clean,
                "reference_age": reference_age_clean,
                "reference_level": reference_level_clean,
                "coach_week_type": week_type_clean,
                "playmodel_profile": playmodel_profile_clean,
                "cycle_days": int(cycle_days_clean),
                "training_days": int(n_train_clean),
                "rest_days": int(n_rest_clean),
                "md_day": md_match_day_clean,
                "own_formation": own_formation_clean,
                "own_block": own_block_clean,
                "own_attack_route": own_attack_route_clean,
            })
            st.success("Alapbeállítás mentve ehhez a belépési azonosítóhoz.") if ok_save else st.warning(f"Nem sikerült menteni: {msg_save}")

        adv1, adv2, adv3 = st.columns(3)
        with adv1:
            cycle_days_clean = st.number_input("Hány napos a ciklus?", min_value=3, max_value=10, value=int(user_defaults_clean.get("cycle_days", 7)), step=1, key="clean_cycle_days_v137")
        with adv2:
            n_rest_clean = st.number_input("Hány pihenőnap?", min_value=0, max_value=6, value=max(0, int(user_defaults_clean.get("rest_days", n_rest_clean))), step=1, key="clean_n_rest_v137")
        with adv3:
            md_day_options_clean = [f"MD-{i}" for i in range(int(cycle_days_clean)-1, 0, -1)] + ["MD"]
            md_match_day_clean = st.selectbox("Melyik nap az MD?", md_day_options_clean, index=_fpi_idx_v113(md_day_options_clean, user_defaults_clean.get("md_day", "MD"), len(md_day_options_clean)-1), key="clean_md_match_day_v137")

        selected_players_clean = st.multiselect("Játékosok szűrése", players_clean, default=players_clean, key="clean_players_v137")
        total_slots_clean = max(1, min(int(cycle_days_clean), int(n_train_clean) + int(n_rest_clean) + 1))
        md_day_options_clean = [f"MD-{i}" for i in range(int(cycle_days_clean)-1, 0, -1)] + ["MD"]
        st.markdown("##### Napokra bontott mikrociklus – opcionális")
        for i in range(total_slots_clean):
            d1, d2, d3 = st.columns([1, 1.2, 3])
            with d1:
                md_default = min(i, len(md_day_options_clean)-1)
                md_v = st.selectbox(f"Nap {i+1}", md_day_options_clean, index=md_default, key=f"clean_md_day_v137_{i}")
            with d2:
                type_options = ["Edzés", "Pihenő", "Regeneráció", "Aktiváció", "Meccs"]
                default_type = "Meccs" if md_v == "MD" else ("Pihenő" if i >= int(n_train_clean) else "Edzés")
                typ_v = st.selectbox(f"Típus {i+1}", type_options, index=type_options.index(default_type), key=f"clean_md_type_v137_{i}")
            with d3:
                note_v = st.text_input(f"Edzői megjegyzés {i+1}", value="", key=f"clean_md_note_v137_{i}")
            session_plan_clean.append({"md": md_v, "type": typ_v, "note": note_v})

    st.session_state["fpi_match_opponent_v94"] = opponent_clean
    st.session_state["fpi_match_date_v94"] = match_date_clean
    st.session_state["fpi_match_context_v94"] = {"opponent": opponent_clean.strip() if isinstance(opponent_clean, str) else "", "match_date": match_date_clean, "match_week": match_week_clean, "today": pd.Timestamp.today().date(), "today_week": _fpi_iso_week_from_date_v94(pd.Timestamp.today().date())}
    st.session_state["clean_own_formation_v132"] = own_formation_clean
    st.session_state["clean_own_block_v132"] = own_block_clean
    st.session_state["clean_own_attack_route_v132"] = own_attack_route_clean

    ref_profile_clean = f"{reference_age_clean} / {reference_level_clean} / játékosonkénti poszt / {playmodel_profile_clean}"
    st.session_state["fpi_coach_context_v97"] = {
        "reference_profile": ref_profile_clean,
        "reference_age": reference_age_clean,
        "reference_level": reference_level_clean,
        "playmodel_profile": playmodel_profile_clean,
        "selected_playstyle": selected_playstyle_clean,
        "coach_week_type": week_type_clean,
        "cycle_days": int(cycle_days_clean),
        "training_days": int(n_train_clean),
        "rest_days": int(n_rest_clean),
        "md_day": md_match_day_clean,
        "session_plan": session_plan_clean,
        "own_formation": own_formation_clean,
        "own_block": own_block_clean,
        "own_attack_route": own_attack_route_clean,
    }

    with st.expander("📅 Hét / meccsnap ellenőrzés", expanded=False):
        week_context_clean = _fpi_week_context_df_v94(df_clean, match_date_clean)
        warnings_clean = _fpi_match_week_warning_v94(df_clean, selected_week_clean, match_date_clean)
        st.markdown(f"**Mai hét:** {_fpi_iso_week_from_date_v94(pd.Timestamp.today().date())} | **Meccshét:** {match_week_clean} | **Kiválasztott hét:** {selected_week_clean}")
        if not week_context_clean.empty:
            st.dataframe(week_context_clean, use_container_width=True, hide_index=True)
        for w in warnings_clean:
            st.warning(w)

    analysis_clean = df_clean[df_clean["player_name"].astype(str).isin(selected_players_clean)].copy() if selected_players_clean else df_clean.copy()

    try:
        ctx_clean = _fpi_report_context(analysis_clean, selected_week_clean, selected_playstyle_clean)
        readiness_clean = int(ctx_clean.get("readiness_score", 70) or 70)
        tactical_gps_context_clean = {"has_gps": True, "selected_week": selected_week_clean, "readiness_score": readiness_clean, "playstyle": selected_playstyle_clean, "priorities": ctx_clean.get("priorities", []), "periodization_type": ctx_clean.get("periodization_type", "n.a."), "df": analysis_clean}
    except Exception:
        tactical_gps_context_clean = {"has_gps": True, "selected_week": selected_week_clean, "readiness_score": 70, "playstyle": selected_playstyle_clean, "priorities": [], "periodization_type": "n.a.", "df": analysis_clean}

    # 3. Opcionális taktikai input – önálló szekció, csak ha van anyag.
    clean_tactical_context = _fpi_clean_tactical_import_v102(tactical_gps_context_clean)
    st.session_state["fpi_clean_tactical_context_v115"] = clean_tactical_context

    # 4. Exportok – nagy, elsődleges gombok.
    _fpi_section_header_v113("4. Export", "A fő termék az Executive Summary. A többi riport kiegészítő / haladó használatra.", "export")
    safe_week_clean = _safe_filename_week(selected_week_clean)
    ex1, ex2 = st.columns([1.55, 1])
    with ex1:
        exec_pdf_clean = build_fpi_product_pdf_bytes(analysis_clean, selected_week_clean, selected_playstyle_clean, report_type="executive", tactical_context=clean_tactical_context)
        if exec_pdf_clean is not None:
            st.download_button("⬇️ VEZETŐI EXECUTIVE SUMMARY PDF", data=exec_pdf_clean, file_name=f"fpi_executive_summary_{safe_week_clean}.pdf", mime="application/pdf", use_container_width=True, key="clean_export_exec_v137")
    with ex2:
        own_team_pdf_clean = build_fpi_own_team_profile_pdf_bytes(analysis_clean, selected_week_clean, selected_playstyle_clean, tactical_context=clean_tactical_context)
        if own_team_pdf_clean is not None:
            st.download_button("⬇️ Saját csapat profil PDF", data=own_team_pdf_clean, file_name=f"fpi_sajat_csapat_profil_{safe_week_clean}.pdf", mime="application/pdf", use_container_width=True, key="clean_export_own_team_v137")

    more1, more2 = st.columns(2)
    with more1:
        gps_pdf_clean = build_fpi_gps_only_pdf_bytes(analysis_clean, selected_week_clean, selected_playstyle_clean)
        if gps_pdf_clean is not None:
            st.download_button("⬇️ GPS-only PDF", data=gps_pdf_clean, file_name=f"fpi_gps_only_report_{safe_week_clean}.pdf", mime="application/pdf", use_container_width=True, key="clean_export_gps_v137")
    with more2:
        full_pdf_clean = build_fpi_product_pdf_bytes(analysis_clean, selected_week_clean, selected_playstyle_clean, report_type="full", tactical_context=clean_tactical_context)
        if full_pdf_clean is not None:
            st.download_button("⬇️ Full Report PDF", data=full_pdf_clean, file_name=f"fpi_full_report_{safe_week_clean}.pdf", mime="application/pdf", use_container_width=True, key="clean_export_full_v137")

    method_pdf_clean_v143 = build_fpi_methodology_pdf_bytes_v143()
    if method_pdf_clean_v143 is not None:
        st.download_button(
            "⬇️ Metodika PDF",
            data=method_pdf_clean_v143,
            file_name="fpi_metodika.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="clean_export_method_v143",
        )


# =========================================================
# V130 - Early Tactical Framework for Import/Executive Export
# This block is intentionally BEFORE the page routing. The clean import/export
# page stops the Streamlit script early, so the tactical engine must exist here.
# =========================================================
FPI_TACTICAL_DIMENSIONS_V130 = [
    ("pressing", "Letámadás"),
    ("build_up", "Labdakihozatal"),
    ("transitions", "Átmenetek"),
    ("attacking_play", "Támadó játék"),
    ("set_pieces", "Pontrúgások"),
    ("possession", "Labdabirtoklás"),
    ("shot_profile", "Lövésprofil"),
]

FPI_STRATEGY_PALETTE_V130 = [
    {"code": "KON", "name": "Kontra mély blokkból", "style": "Direkt", "block": "Mély", "description": "Mélyebb védekezésből gyors, direkt támadásindítás."},
    {"code": "GAT", "name": "Gyors átmenet", "style": "Direkt", "block": "Közép", "description": "Labdaszerzés után gyors előrejáték, kevés passzból veszély."},
    {"code": "BAT", "name": "Középső blokk + átmenet", "style": "Vegyes", "block": "Közép", "description": "Középső zónás védekezés, majd gyors átmeneti támadás."},
    {"code": "KIE", "name": "Kiegyensúlyozott", "style": "Kiegyensúlyozott", "block": "Közép", "description": "Stabil, kockázatkerülőbb alapjáték, kontrollált intenzitással."},
    {"code": "PRS", "name": "Presszing + átmenet", "style": "Direkt/Presszing", "block": "Közép-magas", "description": "Aktív letámadás, labdaszerzés után gyors támadásvezetés."},
    {"code": "MLT", "name": "Magas letámadás", "style": "Agresszív", "block": "Magas", "description": "Magas blokkból agresszív nyomás és korai labdaszerzés."},
    {"code": "DOM", "name": "Dominancia", "style": "Kontroll", "block": "Magas", "description": "Labdabirtoklásra és területi fölényre épülő meccskontroll."},
    {"code": "POZ", "name": "Pozíciós támadás", "style": "Kontroll", "block": "Közép-magas", "description": "Türelmes építkezés, félterületek és szélesség használata."},
    {"code": "LAB", "name": "Labdatartás mélyebben", "style": "Kontroll", "block": "Alacsony-közép", "description": "Biztonságosabb labdatartás mélyebb szerkezetből."},
]

def _fpi_strategy_palette_rows_any_v130():
    pal = globals().get("FPI_STRATEGY_PALETTE_V129") or FPI_STRATEGY_PALETTE_V130
    return [(x["code"], x["name"], x["style"], x["block"], x["description"]) for x in pal]

# Minimal topic dictionary available to the clean page before the full Tactical Pro module is defined.
if "TACTICAL_TOPIC_TAGS_FPI" not in globals():
    TACTICAL_TOPIC_TAGS_FPI = {
        "formation": {"label": "Formáció / alapfelállás", "keywords": ["formation", "shape", "system", "4-4-2", "4-3-3", "4-2-3-1", "3-5-2", "formáció", "felállás", "játékrendszer"]},
        "build_up": {"label": "Labdakihozatal / támadásépítés", "keywords": ["build-up", "build up", "goal kick", "progression", "progressive pass", "third man", "labdakihozatal", "építkezés", "progresszív", "harmadik ember"]},
        "direct_play": {"label": "Direkt játék / hosszú labda", "keywords": ["direct", "long ball", "second ball", "vertical", "direkt", "hosszú labda", "második labda", "vertikális"]},
        "pressing": {"label": "Letámadás / presszing", "keywords": ["press", "pressing", "high press", "ppda", "pressure", "trigger", "letámadás", "presszing", "nyomás", "magas labdaszerzés"]},
        "defensive_block": {"label": "Védekezési blokk / blokkmagasság", "keywords": ["low block", "mid block", "middle block", "high block", "defensive block", "compact", "mély blokk", "középső blokk", "magas blokk", "kompakt", "blokkmagasság"]},
        "transition_attack": {"label": "Támadó átmenet / kontrák", "keywords": ["transition", "counter", "counterattack", "fast attack", "after regain", "átmenet", "kontra", "gyors támadás", "labdaszerzés után"]},
        "transition_defense": {"label": "Védekező átmenet / rest defense", "keywords": ["defensive transition", "rest defense", "counter prevention", "after loss", "védekező átmenet", "rest defense", "biztosítás", "labdavesztés után"]},
        "chance_creation": {"label": "Helyzetkialakítás / támadóharmad", "keywords": ["chance", "key pass", "box entry", "final third", "xg", "shot", "helyzet", "kulcspassz", "támadóharmad", "várható gól", "lövés"]},
        "wide_play": {"label": "Szélső játék / oldali dominancia", "keywords": ["wide", "wing", "flank", "cross", "overlap", "szélső", "szél", "oldal", "beadás", "átfedés"]},
        "central_play": {"label": "Középső játék / félterületek", "keywords": ["central", "half-space", "between the lines", "zone 14", "középső", "félterület", "vonalak között", "14-es zóna"]},
        "set_pieces": {"label": "Pontrúgások", "keywords": ["set piece", "corner", "free kick", "throw-in", "pontrúgás", "szöglet", "szabadrúgás", "bedobás"]},
        "key_players": {"label": "Kulcsjátékosok", "keywords": ["key player", "creator", "playmaker", "progressor", "dribbler", "kulcsjátékos", "kreatív", "irányító", "progresszor"]},
        "weakness_risk": {"label": "Gyengeségek / kockázatok", "keywords": ["weakness", "risk", "vulnerable", "space behind", "gap", "gyengeség", "kockázat", "sebezhető", "mögötti terület"]},
        "recommendation": {"label": "Javaslat / meccsterv", "keywords": ["recommend", "game plan", "match plan", "plan a", "plan b", "focus", "javaslat", "meccsterv", "terv a", "terv b", "fókusz"]},
    }

def _fpi_tactical_context_has_pdf_v130(tactical_context: Optional[Dict[str, object]]) -> bool:
    if not tactical_context:
        return False
    if tactical_context.get("has_own_pdf") or tactical_context.get("has_opp_pdf"):
        return True
    for k in ["own_pdf_chars", "opp_pdf_chars", "pdf_direct_findings_count", "pdf_direct_lines_count"]:
        try:
            if int(tactical_context.get(k, 0) or 0) > 0:
                return True
        except Exception:
            pass
    pdfi = tactical_context.get("pdf_insights") or {}
    if isinstance(pdfi, dict):
        try:
            if int(pdfi.get("raw_text_len", 0) or 0) > 0:
                return True
        except Exception:
            pass
    return False

def _fpi_has_tactical_signal_v95(tactical_context: Optional[Dict[str, object]]) -> bool:
    if not tactical_context:
        return False
    if _fpi_tactical_context_has_pdf_v130(tactical_context):
        return True
    if tactical_context.get("strategy_framework"):
        return True
    keys = ["tactical_findings", "pdf_provider_lines", "pdf_provider_findings", "own_topics", "opp_topics", "team_comparison", "risks"]
    for k in keys:
        v = tactical_context.get(k)
        if isinstance(v, list) and len(v) > 0:
            return True
        if isinstance(v, dict) and len(v) > 0:
            return True
    if tactical_context.get("has_own_team_excel") or tactical_context.get("has_opp_team_excel") or tactical_context.get("has_own_player_excel") or tactical_context.get("has_opp_player_excel"):
        return True
    return False

def _fpi_topic_present_v130(blocks: Dict[str, object], pdf_topics: List[object], *keys: str) -> bool:
    for k in keys:
        if blocks.get(k):
            return True
    wanted = {str(k).replace("_", " ").lower() for k in keys} | {str(k).lower() for k in keys}
    for item in pdf_topics or []:
        if isinstance(item, dict):
            joined = " ".join(str(v) for v in item.values()).lower()
        else:
            joined = str(item).lower()
        if any(w and w in joined for w in wanted):
            return True
    return False

def _fpi_metric_v130(metrics: Dict[str, object], key: str, default=0):
    try:
        if "_fpi_metric_value_v79" in globals():
            return _fpi_metric_value_v79(metrics, key) or default
    except Exception:
        pass
    try:
        return float((metrics or {}).get(key, default) or default)
    except Exception:
        return default

def _fpi_norm_metric_v130(metric: str, value):
    try:
        if "_fpi_normalized_tactical_metric_v79" in globals():
            return _fpi_normalized_tactical_metric_v79(metric, value)
    except Exception:
        pass
    try:
        return float(value)
    except Exception:
        return None

def _fpi_clamp_v130(x: float, lo: float = 0.0, hi: float = 10.0) -> float:
    try:
        return float(max(lo, min(hi, x)))
    except Exception:
        return lo

def _fpi_tactical_dimension_scores_v130(blocks, pdf_topics, own_team_metrics, opp_team_metrics, pdf_raw_len: int = 0):
    own_pos = _fpi_norm_metric_v130("possession_pct", _fpi_metric_v130(own_team_metrics, "possession_pct", 50)) or 50
    opp_pos = _fpi_norm_metric_v130("possession_pct", _fpi_metric_v130(opp_team_metrics, "possession_pct", 50)) or 50
    opp_ppda = _fpi_norm_metric_v130("ppda", _fpi_metric_v130(opp_team_metrics, "ppda", None))
    own_ppda = _fpi_norm_metric_v130("ppda", _fpi_metric_v130(own_team_metrics, "ppda", None))
    own_shots = _fpi_metric_v130(own_team_metrics, "shots", 0) or 0
    opp_shots = _fpi_metric_v130(opp_team_metrics, "shots", 0) or 0
    own_xg = _fpi_norm_metric_v130("xg", _fpi_metric_v130(own_team_metrics, "xg", 0)) or 0
    opp_xg = _fpi_norm_metric_v130("xg", _fpi_metric_v130(opp_team_metrics, "xg", 0)) or 0
    own_corners = _fpi_metric_v130(own_team_metrics, "corners", 0) or 0
    opp_corners = _fpi_metric_v130(opp_team_metrics, "corners", 0) or 0
    own_counters = _fpi_metric_v130(own_team_metrics, "counterattacks", 0) or 0
    opp_counters = _fpi_metric_v130(opp_team_metrics, "counterattacks", 0) or 0
    own_crosses = _fpi_metric_v130(own_team_metrics, "crosses", 0) or 0
    opp_crosses = _fpi_metric_v130(opp_team_metrics, "crosses", 0) or 0
    confidence_boost = 0.5 if pdf_raw_len > 1000 else 0
    pressing = 4.5 + confidence_boost + (2.0 if _fpi_topic_present_v130(blocks, pdf_topics, "pressing") else 0) + (1.7 if opp_ppda and opp_ppda < 9 else 0) + (0.8 if own_ppda and own_ppda < 9 else 0)
    build_up = 4.8 + confidence_boost + (1.7 if _fpi_topic_present_v130(blocks, pdf_topics, "build_up", "central_play") else 0) + (1.1 if max(own_pos, opp_pos) >= 55 else 0) - (0.8 if _fpi_topic_present_v130(blocks, pdf_topics, "direct_play") else 0)
    transitions = 4.5 + confidence_boost + (2.0 if _fpi_topic_present_v130(blocks, pdf_topics, "transition_attack", "transition_defense") else 0) + (1.6 if own_counters + opp_counters >= 6 else 0)
    attacking = 4.8 + confidence_boost + (1.6 if _fpi_topic_present_v130(blocks, pdf_topics, "chance_creation", "wide_play", "central_play") else 0) + (1.2 if max(own_shots, opp_shots) >= 12 or max(own_xg, opp_xg) >= 1.4 else 0) + (0.6 if own_crosses + opp_crosses >= 22 else 0)
    setpieces = 4.2 + (2.2 if _fpi_topic_present_v130(blocks, pdf_topics, "set_pieces") else 0) + (1.3 if own_corners + opp_corners >= 8 else 0)
    possession = 4.8 + (max(own_pos, opp_pos) - 50) / 7.0 + (1.0 if _fpi_topic_present_v130(blocks, pdf_topics, "build_up") else 0) - (1.0 if _fpi_topic_present_v130(blocks, pdf_topics, "direct_play") else 0)
    shot_profile = 4.5 + (1.2 if _fpi_topic_present_v130(blocks, pdf_topics, "chance_creation") else 0) + (1.0 if max(own_shots, opp_shots) >= 10 else 0) + (1.0 if max(own_xg, opp_xg) >= 1.2 else 0)
    return {"pressing": round(_fpi_clamp_v130(pressing),1), "build_up": round(_fpi_clamp_v130(build_up),1), "transitions": round(_fpi_clamp_v130(transitions),1), "attacking_play": round(_fpi_clamp_v130(attacking),1), "set_pieces": round(_fpi_clamp_v130(setpieces),1), "possession": round(_fpi_clamp_v130(possession),1), "shot_profile": round(_fpi_clamp_v130(shot_profile),1)}

def _fpi_select_tactical_strategy_v130(readiness, blocks, pdf_topics, own_team_metrics, opp_team_metrics, pdf_raw_len: int = 0):
    dims = _fpi_tactical_dimension_scores_v130(blocks, pdf_topics, own_team_metrics, opp_team_metrics, pdf_raw_len)
    own_pos = _fpi_norm_metric_v130("possession_pct", _fpi_metric_v130(own_team_metrics, "possession_pct", 50)) or 50
    opp_pos = _fpi_norm_metric_v130("possession_pct", _fpi_metric_v130(opp_team_metrics, "possession_pct", 50)) or 50
    opp_ppda = _fpi_norm_metric_v130("ppda", _fpi_metric_v130(opp_team_metrics, "ppda", None))
    opp_counters = _fpi_metric_v130(opp_team_metrics, "counterattacks", 0) or 0
    opp_corners = _fpi_metric_v130(opp_team_metrics, "corners", 0) or 0
    scores = {x["code"]: 0.0 for x in FPI_STRATEGY_PALETTE_V130}
    scores["KIE"] += 4.0
    scores["BAT"] += dims["transitions"] * 0.8 + dims["attacking_play"] * 0.25
    scores["GAT"] += dims["transitions"] * 0.9 + (2.0 if _fpi_topic_present_v130(blocks, pdf_topics, "direct_play") else 0)
    scores["KON"] += dims["transitions"] * 0.5 + (2.0 if readiness < 55 else 0)
    scores["PRS"] += dims["pressing"] * 0.95 + dims["transitions"] * 0.35
    scores["MLT"] += dims["pressing"] * 1.15 + (1.5 if opp_ppda and opp_ppda < 8 else 0)
    scores["DOM"] += dims["possession"] * 0.9 + dims["build_up"] * 0.55 + (1.2 if max(own_pos, opp_pos) >= 56 else 0)
    scores["POZ"] += dims["build_up"] * 0.85 + dims["attacking_play"] * 0.55 + (1.0 if _fpi_topic_present_v130(blocks, pdf_topics, "central_play", "wide_play") else 0)
    scores["LAB"] += dims["possession"] * 0.5 + (1.5 if readiness < 60 else 0)
    if opp_counters > 0 or _fpi_topic_present_v130(blocks, pdf_topics, "transition_attack"):
        scores["BAT"] += 2.2; scores["GAT"] += 1.4
    if _fpi_topic_present_v130(blocks, pdf_topics, "set_pieces") or opp_corners > 0:
        scores["KIE"] += 0.6
    if readiness < 55:
        scores["BAT"] += 1.0; scores["KON"] += 1.0; scores["MLT"] -= 2.0; scores["PRS"] -= 1.0
    ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    p, ps = ranked[0]; s, ss = ranked[1]
    rows = {x["code"]: x for x in FPI_STRATEGY_PALETTE_V130}
    total = max(ps + ss, 0.1); pp = int(round(ps / total * 100)); sp = 100 - pp
    detail = []
    if p in ["BAT", "GAT", "KON"]:
        detail.append("átmenetek és rest defense biztosítás")
    if p in ["PRS", "MLT"]:
        detail.append("presszingtrigger és mögöttes biztosítás")
    if p in ["DOM", "POZ", "LAB"]:
        detail.append("labdakihozatal és türelmes progresszió")
    if _fpi_topic_present_v130(blocks, pdf_topics, "wide_play"):
        detail.append("oldali védekezési/támadási kontroll")
    if _fpi_topic_present_v130(blocks, pdf_topics, "set_pieces") or opp_corners > 0:
        detail.append("pontrúgások és második labdák")
    if not detail:
        detail.append("stabil szerkezet és kontrollált kockázat")
    return {"primary_code": p, "primary_name": rows[p]["name"], "secondary_code": s, "secondary_name": rows[s]["name"], "primary_pct": pp, "secondary_pct": sp, "recommendation": f"{p} - {rows[p]['name']}, " + "; ".join(detail[:2]), "plan_b": f"{s} - {rows[s]['name']}", "dimensions": dims, "scores": {k: round(v,2) for k,v in scores.items()}, "detail_focus": detail[:4]}

def _fpi_build_adaptive_match_training_plan(gps_context: Dict[str, object], tactical: Dict[str, object]) -> Dict[str, object]:
    readiness = int((gps_context or {}).get("readiness_score", 70) or 70)
    priorities = (gps_context or {}).get("priorities", []) or []
    pdfi = (tactical or {}).get("pdf_insights") or {}
    blocks = pdfi.get("blocks", {}) if isinstance(pdfi, dict) else {}
    pdf_topics = list((pdfi.get("sportsbase_findings", []) if isinstance(pdfi, dict) else []) or []) + list((pdfi.get("topics", []) if isinstance(pdfi, dict) else []) or [])
    pdf_raw_len = int((pdfi.get("raw_text_len", 0) if isinstance(pdfi, dict) else 0) or 0)
    opp_team_metrics = (tactical or {}).get("team_metrics") or {}
    opp_player_tables = (tactical or {}).get("player_tables") or {}
    own_team_metrics = (((tactical or {}).get("own") or {}).get("team_metrics") or {})
    own_player_tables = (((tactical or {}).get("own") or {}).get("player_tables") or {})
    tactical_findings = []
    try:
        tactical_findings = _fpi_build_excel_driven_tactical_findings_v79(own_team_metrics, opp_team_metrics, own_player_tables, opp_player_tables, pdf_topics)
    except Exception:
        tactical_findings = []
    has_pdf_signal = bool(pdf_raw_len > 0 or any(blocks.values()) or pdf_topics)
    gps_only_mode = not tactical_findings and not has_pdf_signal and not own_team_metrics and not opp_team_metrics
    if gps_only_mode:
        try:
            msgs = _fpi_gps_only_conclusions_v95(gps_context, priorities, readiness, str((gps_context or {}).get("selected_week", "")), limit=6)
        except Exception:
            msgs = ["GPS-only konklúzió: taktikai input nélkül terhelési fókuszú mikrociklus."]
        for msg in msgs:
            tactical_findings.append({"Téma": "GPS-only konklúzió", "Bizonyíték": "GPS / readiness / edzés-meccs arányok", "Edzői következtetés": msg, "Prioritás": "Közepes"})
    elif has_pdf_signal and not tactical_findings:
        tactical_findings.append({"Téma": "Taktikai PDF input", "Bizonyíték": f"PDF szöveg feldolgozva ({pdf_raw_len} karakter)", "Edzői következtetés": "A Tactical Framework aktív: a 7 dimenzió és 9 stratégiai profil alapján készül alapirány és mikrociklus-fókusz.", "Prioritás": "Közepes"})
    risks = []
    for f in tactical_findings:
        if f.get("Prioritás") in ["Magas", "Közepes"]:
            risks.append(f"{f.get('Téma')}: {f.get('Edzői következtetés')}")
    if blocks.get("transition_attack"):
        risks.append("PDF alapján: ellenfél-kontrák / gyors átmenetek kezelése")
    if blocks.get("set_pieces"):
        risks.append("PDF alapján: pontrúgás-védekezés és második labdák")
    if blocks.get("wide_play"):
        risks.append("PDF alapján: szélső játék, beadások, oldali túlterhelések")
    if blocks.get("pressing"):
        risks.append("PDF alapján: presszing kijátszása és első passzsor döntései")
    if not risks:
        risks.append("GPS-alapú terhelési és readiness kockázatok")
    strategy_framework = _fpi_select_tactical_strategy_v130(readiness, blocks, pdf_topics, own_team_metrics, opp_team_metrics, pdf_raw_len=pdf_raw_len) if not gps_only_mode else {}
    plan_a = strategy_framework.get("recommendation") if strategy_framework else "GPS-only - erőnléti fókuszú mikrociklus"
    if strategy_framework.get("plan_b"):
        risks.append(f"B terv / alternatív profil: {strategy_framework.get('plan_b')} ({strategy_framework.get('secondary_pct')}%).")
    md_plan = [("MD+1/MD-5", "Regeneráció / alacsony intenzitás", "Előző terhelés visszarendezése."), ("MD-4", "Volumen + saját játékmodell", "Stabil csapatvolumen és saját labdakihozatal / védekezési alapok."), ("MD-3", "HSR / sprint exponálás + átmenetek", "Meccsintenzitás előkészítése, kontrollált mennyiséggel."), ("MD-2", "Ellenfél-specifikus taktikai nap", "; ".join(risks[:2]) if risks else "Meccsterv."), ("MD-1", "Aktiváció + pontrúgások", "Frissítés, gyors döntések, fix helyzetek.")]
    if gps_only_mode:
        try:
            md_plan = [("MD+1/MD-5", "Regeneráció / monitoring", "Előző terhelés visszarendezése.")] + _fpi_gps_only_md_plan_v95(gps_context, readiness, priorities, str((gps_context or {}).get("selected_week", "")))
        except Exception:
            pass
    else:
        if readiness < 55:
            md_plan[2] = ("MD-3", "Rövid specifikus exponálás", "Csak célzott HSR/sprint inger, alacsony volumen.")
        if any("presszing" in str(r).lower() or "ppda" in str(r).lower() for r in risks):
            md_plan[1] = ("MD-4", "Presszingkijátszás + labdakihozatal", "Ellenfél presszingprofil alapján első passzsor és harmadik ember.")
        if any("kontra" in str(r).lower() or "átmenet" in str(r).lower() for r in risks):
            md_plan[2] = ("MD-3", "Átmeneti játék + HSR/sprint", "Kontrák és gyors átmenetek miatt futóintenzitás + döntésgyorsaság.")
            md_plan[3] = ("MD-2", "Rest defense + kontrák elleni biztosítás", "Ellenfél átmeneti veszélyei miatt.")
        if _fpi_topic_present_v130(blocks, pdf_topics, "set_pieces"):
            md_plan[-1] = ("MD-1", "Aktiváció + pontrúgás fókusz", "Szöglet/pontrúgás profil alapján.")
    player_focus = []
    for f in tactical_findings:
        if "játékos" in str(f.get("Téma", "")).lower() or "progresszor" in str(f.get("Téma", "")).lower():
            player_focus.append(f"{f.get('Téma')}: {f.get('Bizonyíték')} -> {f.get('Edzői következtetés')}")
    return {"analysis_level": (tactical or {}).get("analysis_level_label", "GPS Only"), "plan_a": plan_a, "risks": list(dict.fromkeys(risks))[:6], "md_plan": md_plan, "player_focus": player_focus[:6], "tactical_findings": tactical_findings[:10], "team_comparison": (_fpi_tactical_compare_team_metrics_v79(own_team_metrics, opp_team_metrics) if "_fpi_tactical_compare_team_metrics_v79" in globals() else {}), "strategy_framework": strategy_framework, "plan_b": strategy_framework.get("plan_b", "") if isinstance(strategy_framework, dict) else ""}

def _build_tactical_executive_context(gps_context: Dict[str, object], tactical_ctx: Dict[str, object], plan: Dict[str, object]) -> Dict[str, object]:
    own = (tactical_ctx or {}).get("own", {}) or {}
    opp = (tactical_ctx or {}).get("opponent", {}) or {}
    pdfi = (tactical_ctx or {}).get("pdf_insights", {}) or {}
    own_i = (own.get("pdf_insights") or {})
    opp_i = (opp.get("pdf_insights") or {})
    return {"version": globals().get("TACTICAL_PRO_VERSION", "TACTICAL_EARLY_V130"), "analysis_level": (tactical_ctx or {}).get("analysis_level_label", "GPS + taktikai input"), "has_own_pdf": bool(own_i.get("pdf_uploaded") or own_i.get("raw_text_len", 0)), "has_opp_pdf": bool(opp_i.get("pdf_uploaded") or opp_i.get("raw_text_len", 0)), "own_pdf_pages": int(own_i.get("pdf_pages", 0) or 0), "opp_pdf_pages": int(opp_i.get("pdf_pages", 0) or 0), "own_pdf_chars": int(own_i.get("raw_text_len", 0) or 0), "opp_pdf_chars": int(opp_i.get("raw_text_len", 0) or 0), "has_own_team_excel": bool(own.get("team_metrics")), "has_opp_team_excel": bool(opp.get("team_metrics")), "has_own_player_excel": bool(own.get("player_tables")), "has_opp_player_excel": bool(opp.get("player_tables")), "own_topics": (own_i.get("topics") or [])[:8], "opp_topics": (opp_i.get("topics") or [])[:8], "own_team_metrics": own.get("team_metrics", {}), "opp_team_metrics": opp.get("team_metrics", {}), "plan_a": (plan or {}).get("plan_a", "KIE - Kiegyensúlyozott"), "plan_b": (plan or {}).get("plan_b", ""), "strategy_framework": (plan or {}).get("strategy_framework", {}), "risks": (plan or {}).get("risks", []), "md_plan": (plan or {}).get("md_plan", []), "player_focus": (plan or {}).get("player_focus", []), "tactical_findings": (plan or {}).get("tactical_findings", []), "team_comparison": (plan or {}).get("team_comparison", {}), "own_player_tables": (tactical_ctx or {}).get("own", {}).get("player_tables", {}), "opp_player_tables": (tactical_ctx or {}).get("opponent", {}).get("player_tables", {}), "opponent_player_evaluation": _fpi_build_player_evaluation_v132((tactical_ctx or {}).get("opponent", {}).get("player_tables", {}), side="opp", max_rows=9) if "_fpi_build_player_evaluation_v132" in globals() else [], "own_player_evaluation": _fpi_build_player_evaluation_v132((tactical_ctx or {}).get("own", {}).get("player_tables", {}), side="own", max_rows=9) if "_fpi_build_player_evaluation_v132" in globals() else [], "pdf_provider_lines": (pdfi.get("sportsbase_lines") or []), "pdf_provider_findings": (pdfi.get("sportsbase_findings") or []), "pdf_direct_findings_count": len(pdfi.get("sportsbase_findings") or []), "pdf_direct_lines_count": len(pdfi.get("sportsbase_lines") or [])}


# -----------------------------------------------------------------------------
# UI
# -----------------------------------------------------------------------------

def render_fpi_app_hub_v137() -> None:
    """Haladó app előszobája: a teljes dashboard helyett először navigációs hub."""
    _fpi_landing_css_v100()
    _fpi_apply_v119_all_light_readable_patch()
    st.markdown(
        """
        <style>
        .fpi-hub-hero{border-radius:30px;padding:30px 34px;margin:8px 0 20px 0;background:linear-gradient(135deg,#ffffff,#e0f2fe 55%,#ecfdf5);border:1px solid #bfdbfe;box-shadow:0 22px 58px rgba(15,23,42,.14);}
        .fpi-hub-kicker{display:inline-block;padding:7px 12px;border-radius:999px;background:#f0fdfa;border:1px solid #99f6e4;color:#0f766e;font-weight:950;font-size:.8rem;letter-spacing:.06em;margin-bottom:10px;}
        .fpi-hub-title{font-size:2.55rem;line-height:1;font-weight:980;letter-spacing:-.05em;color:#0f172a;margin-bottom:8px;}
        .fpi-hub-sub{color:#475569;font-size:1.03rem;line-height:1.45;max-width:950px;}
        .fpi-hub-card{border-radius:22px;padding:18px 20px;background:#ffffff;border:1px solid #dbeafe;box-shadow:0 14px 34px rgba(15,23,42,.09);min-height:150px;margin-bottom:12px;}
        .fpi-hub-card h3{margin:0 0 6px 0;color:#0f172a;font-weight:950;}.fpi-hub-card p{margin:0;color:#475569;line-height:1.4;}
        </style>
        <div class="fpi-hub-hero">
            <div class="fpi-hub-kicker">HALADÓ ELEMZŐ APP</div>
            <div class="fpi-hub-title">FPI Hub</div>
            <div class="fpi-hub-sub">A legtöbb felhasználónak elég a Vezetői riport készítése oldal. Ez a felület a részletesebb dashboardokhoz, metodikához és diagnosztikához van.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    nav1, nav2, nav3 = st.columns([1.4, 1.6, 1.6])
    with nav1:
        if st.button("🚀 Vezetői riport készítése", use_container_width=True, type="primary", key="hub_go_clean_v137"):
            _fpi_set_page_v100("clean")
    with nav2:
        if st.button("📊 Teljes dashboard megnyitása", use_container_width=True, key="hub_open_full_v137"):
            st.session_state["fpi_app_hub_seen_v137"] = True
            st.rerun()
    with nav3:
        if st.button("← Főoldal", use_container_width=True, key="hub_back_landing_v137"):
            _fpi_set_page_v100("landing")

    st.markdown("### Gyors belépési pontok")
    l1, l2, l3, l4 = st.columns(4)
    with l1:
        if st.button("📈 Dashboard", use_container_width=True, key="hub_dashboard_link_v138"):
            st.session_state["fpi_app_hub_seen_v137"] = True
            st.rerun()
    with l2:
        if st.button("🧠 Intelligence", use_container_width=True, key="hub_intelligence_link_v138"):
            st.session_state["fpi_app_hub_seen_v137"] = True
            st.rerun()
    with l3:
        if st.button("📚 Metodika Center", use_container_width=True, key="hub_method_link_v138"):
            _fpi_set_page_v100("method")
    with l4:
        if st.button("⚽ Tactical Pro+", use_container_width=True, key="hub_tactical_link_v138"):
            st.session_state["fpi_app_hub_seen_v137"] = True
            st.rerun()

    a,b,c = st.columns(3)
    with a:
        st.markdown('<div class="fpi-hub-card"><h3>Heti összefoglaló</h3><p>Readiness, risk, fő üzenetek és heti döntéstámogatás.</p></div>', unsafe_allow_html=True)
        st.markdown('<div class="fpi-hub-card"><h3>Benchmark</h3><p>Korosztály, szint, poszt és játékmodell alapján értelmezett referencia.</p></div>', unsafe_allow_html=True)
    with b:
        st.markdown('<div class="fpi-hub-card"><h3>Player Risk</h3><p>Játékosszintű figyelmeztetések, túl- és alulterhelési mintázatok.</p></div>', unsafe_allow_html=True)
        st.markdown('<div class="fpi-hub-card"><h3>Tactical Pro+</h3><p>Opcionális PDF/Excel inputból taktikai kontextus, ellenfélveszélyek és meccsfókusz.</p></div>', unsafe_allow_html=True)
    with c:
        st.markdown('<div class="fpi-hub-card"><h3>Mikrociklus</h3><p>Heti edzésszám, readiness és játékmodell alapján javasolt heti terv.</p></div>', unsafe_allow_html=True)
        st.markdown('<div class="fpi-hub-card"><h3>Metodika</h3><p>Readiness, risk, benchmark és Tactical Framework magyarázata.</p></div>', unsafe_allow_html=True)




def render_fpi_methodology_center_v138() -> None:
    """Rövid, vezetői szemléletű metodikai központ első használathoz."""
    _fpi_landing_css_v100()
    _fpi_apply_v119_all_light_readable_patch()
    st.markdown(
        """
        <style>
        .fpi-method-hero{border-radius:30px;padding:30px 34px;margin:8px 0 20px 0;background:linear-gradient(135deg,#ffffff,#e0f2fe 55%,#ecfdf5);border:1px solid #bfdbfe;box-shadow:0 22px 58px rgba(15,23,42,.14);}
        .fpi-method-kicker{display:inline-block;padding:7px 12px;border-radius:999px;background:#f0fdfa;border:1px solid #99f6e4;color:#0f766e;font-weight:950;font-size:.8rem;letter-spacing:.06em;margin-bottom:10px;}
        .fpi-method-title{font-size:2.45rem;line-height:1;font-weight:980;letter-spacing:-.05em;color:#0f172a;margin-bottom:8px;}
        .fpi-method-sub{color:#475569;font-size:1.03rem;line-height:1.45;max-width:980px;}
        .fpi-method-pill{display:inline-block;margin:5px 6px 0 0;padding:7px 11px;border-radius:999px;background:#eff6ff;border:1px solid #bfdbfe;color:#1e3a8a;font-weight:850;font-size:.86rem;}
        div[data-testid="stExpander"]{border-radius:18px !important;border:1px solid #dbeafe !important;background:#ffffff !important;margin-bottom:8px !important;}
        div[data-testid="stExpander"] *{color:#0f172a !important;}
        </style>
        <div class="fpi-method-hero">
            <div class="fpi-method-kicker">FPI METHODOLOGY CENTER</div>
            <div class="fpi-method-title">Mit számol az FPI, és hogyan kell értelmezni?</div>
            <div class="fpi-method-sub">Rövid, szakmai áttekintő azokhoz a kérdésekhez, amelyeket egy vezetőedző, erőnléti edző vagy sportigazgató első ránézésre feltenne.</div>
            <div><span class="fpi-method-pill">GPS terhelés</span><span class="fpi-method-pill">Readiness</span><span class="fpi-method-pill">Risk</span><span class="fpi-method-pill">Benchmark</span><span class="fpi-method-pill">Tactical Pro+</span><span class="fpi-method-pill">Mikrociklus</span></div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    n1,n2,n3 = st.columns([1.2,1.6,4])
    with n1:
        if st.button("← Főoldal", use_container_width=True, key="method_back_landing_v138"):
            _fpi_set_page_v100("landing")
    with n2:
        if st.button("🚀 Riport készítése", use_container_width=True, type="primary", key="method_go_clean_v138"):
            _fpi_set_page_v100("clean")
    with n3:
        st.caption("Az FPI döntéstámogató rendszer: az eredmények edzői, orvosi és teljesítménydiagnosztikai kontextussal együtt értelmezendők.")

    with st.expander("1. Mi az FPI lényege?", expanded=True):
        st.markdown("""
        A **Football Performance Intelligence** a GPS-, benchmark- és opcionálisan taktikai inputokat vezetői döntéstámogatássá alakítja.
        Nem csak grafikonokat ad, hanem rövid, edzői nyelvű konklúziókat: aktuális állapot, kockázatok, referencia-eltérések és következő heti fókusz.
        """)
    with st.expander("2. Mit jelent a Readiness Score?"):
        st.markdown("""
        A Readiness Score a játékos vagy csapat aktuális terhelési állapotának becslése. Figyelembe veszi többek között:
        - elmúlt 3–7 nap terhelését,
        - 4 hetes trendeket,
        - Load / össztáv / HSR / sprint / High Efforts értékeket,
        - meccs- és edzésarányokat,
        - pulzus/HRV adatot, ha rendelkezésre áll.

        **Fontos:** az alacsonyabb readiness nem automatikusan túlterhelést jelent. Okozhatja túlterhelés, alulterhelés vagy kedvezőtlen terhelési mintázat is.
        """)
    with st.expander("3. Mit jelent a Player Risk?"):
        st.markdown("""
        A Player Risk nem sérülés-előrejelzés, hanem figyelmeztető besorolás. A rendszer a túl- vagy alulterhelésre, hirtelen terhelésváltozásra és hiányzó sebességi expozícióra figyel.

        **Alacsony:** stabil terhelési profil.  
        **Közepes:** egy-két figyelmeztető jel, például emelkedő Load vagy gyenge sprint-expozíció.  
        **Magas:** több kedvezőtlen tényező együtt, például terhelésugrás + rossz trend + kiugró HSR/sprint vagy neuromuszkuláris terhelés.
        """)
    with st.expander("4. Milyen benchmarkokat használ?"):
        st.markdown("""
        A benchmark-rendszer nem egyetlen általános átlaghoz hasonlít. A referencia a következő tényezők alapján módosul:
        - korosztály,
        - bajnoki szint,
        - poszt,
        - kapus / mezőnyjátékos logika,
        - játékmodell.

        A benchmarkok nemzetközi szakirodalmi irányok, gyakorlati teljesítménydiagnosztikai tapasztalatok és saját adatbázis logika alapján kerülnek kialakításra.
        """)
    with st.expander("5. Mit tud a Tactical Pro+ modul?"):
        st.markdown("""
        A Tactical Pro+ a GPS-állapotot összekapcsolja a saját játékmodellel, ellenfélanyaggal és heti periodizációval.
        A cél nem hosszú taktikai tanulmány, hanem használható edzői fókusz:
        - ellenfél kulcsveszélyei,
        - játékosszintű ellenfélprofilok,
        - saját csapat heti alapprofilja,
        - Tactical Framework / stratégiai irány,
        - edzésfókuszok MD-napokra bontva.
        """)
    with st.expander("6. Hogyan készül a mikrociklus-javaslat?"):
        st.markdown("""
        A mikrociklus motor az alábbiakat kombinálja:
        - heti edzésszám,
        - meccsnap / kiválasztott hét,
        - readiness és risk állapot,
        - HSR / sprint / High Efforts trend,
        - játékmodell,
        - opcionális ellenfél-specifikus taktikai fókusz.

        Így a terv nem fix 4 edzéses sablon, hanem a megadott heti kontextushoz igazodik.
        """)
    with st.expander("7. Hogyan kell szakmailag használni?"):
        st.markdown("""
        **Vezetőedző:** gyors heti fő üzenetek, taktikai és terhelési fókusz.  
        **Erőnléti edző:** risk, readiness, HSR/sprint/Load trendek és expozíció.  
        **Sportigazgató:** 30 másodperces állapotkép, kockázatok és heti prioritások.  
        **Utánpótlás:** korosztályos és poszt-specifikus összevetés, alul- és túlterhelés korai jelzése.
        """)
    with st.expander("8. Milyen állítást nem tesz az FPI?"):
        st.markdown("""
        Az FPI nem diagnózis, nem orvosi döntés és nem garantált sérülés-előrejelzés.
        A számok döntéstámogató jelzések, amelyeket edzői megfigyeléssel, orvosi információval, wellness/RPE adattal és szakmai kontextussal együtt érdemes értelmezni.
        """)


# =========================================================
# V143 - Universal GPS import engine
# Brainsports / PlayerTek / Polar / Catapult + generic mapper
# =========================================================
FPI_SUPPORTED_GPS_PROVIDERS_V143 = [
    "Automatikus felismerés",
    "Brainsports",
    "PlayerTek",
    "Polar Team Pro",
    "Catapult",
    "Egyéb / Smart Mapper",
]


def _fpi_bytes_from_upload_v143(uploaded_file) -> bytes:
    if uploaded_file is None:
        return b""
    try:
        return uploaded_file.getvalue()
    except Exception:
        try:
            return uploaded_file.read()
        except Exception:
            return b""


def _fpi_read_csv_bytes_v143(data: bytes, name: str = "") -> pd.DataFrame:
    last_err = None
    for enc in ["utf-8-sig", "utf-8", "cp1250", "latin2", "iso-8859-2", "cp1252"]:
        try:
            return pd.read_csv(io.BytesIO(data), sep=None, engine="python", encoding=enc, header=None)
        except Exception as exc:
            last_err = exc
    raise ValueError(f"CSV beolvasási hiba ({name}): {last_err}")


def _fpi_excel_serial_to_datetime_v143(value):
    try:
        if pd.isna(value):
            return pd.NaT
    except Exception:
        pass
    if isinstance(value, (pd.Timestamp, datetime)):
        return pd.to_datetime(value, errors="coerce")
    try:
        num = float(value)
        if 20000 <= num <= 80000:
            return pd.Timestamp("1899-12-30") + pd.to_timedelta(num, unit="D")
    except Exception:
        pass
    return pd.to_datetime(value, errors="coerce", dayfirst=True)


def _fpi_clean_headers_v143(values) -> List[str]:
    result = []
    used = {}
    for idx, value in enumerate(list(values)):
        text = clean_col_name(value)
        if not text or str(text).lower() in {"nan", "none"}:
            text = f"Unnamed_{idx+1}"
        base = text
        used[base] = used.get(base, 0) + 1
        if used[base] > 1:
            text = f"{base}_{used[base]}"
        result.append(text)
    return result


def _fpi_find_header_row_v143(raw: pd.DataFrame, required_terms: List[str], max_scan: int = 40) -> Optional[int]:
    if raw is None or raw.empty:
        return None
    required = [str(x).lower() for x in required_terms]
    best_idx, best_score = None, -1
    for idx in range(min(max_scan, len(raw))):
        vals = [str(x).strip().lower() for x in raw.iloc[idx].tolist() if str(x).strip().lower() not in {"", "nan", "none"}]
        joined = " | ".join(vals)
        score = sum(1 for term in required if term in joined)
        if score > best_score:
            best_idx, best_score = idx, score
    return best_idx if best_score >= max(1, min(2, len(required))) else None


def _fpi_table_from_raw_v143(raw: pd.DataFrame, header_idx: int, stop_markers: Optional[List[str]] = None) -> pd.DataFrame:
    if raw is None or raw.empty or header_idx is None or header_idx >= len(raw):
        return pd.DataFrame()
    headers = _fpi_clean_headers_v143(raw.iloc[header_idx].tolist())
    data = raw.iloc[header_idx + 1 :].copy()
    data.columns = headers
    data = data.dropna(how="all")
    if stop_markers:
        stop_terms = [str(x).lower() for x in stop_markers]
        keep_rows = []
        for _, row in data.iterrows():
            row_text = " | ".join(str(x).strip().lower() for x in row.tolist() if str(x).strip().lower() not in {"", "nan", "none"})
            if any(marker in row_text for marker in stop_terms):
                break
            keep_rows.append(row)
        data = pd.DataFrame(keep_rows, columns=headers) if keep_rows else pd.DataFrame(columns=headers)
    return data.reset_index(drop=True)


def _fpi_detect_provider_v143(sheets: Dict[str, pd.DataFrame], file_name: str = "") -> str:
    name = str(file_name).lower()
    sheet_names = " ".join(str(x).lower() for x in sheets.keys())
    sample_text = ""
    for frame in list(sheets.values())[:4]:
        try:
            sample_text += " " + " ".join(str(x).lower() for x in frame.head(8).fillna("").astype(str).values.ravel().tolist())
        except Exception:
            pass

    if "playertek" in name or "playerteck" in name or all(x in sample_text for x in ["player name", "player load", "distance (km)"]):
        return "PlayerTek"
    if "brainsport" in name or "targets summary" in sample_text or "heart exertion" in sample_text or "main table" in sheet_names:
        return "Brainsports"
    if "polar" in name or "edzési terhelési pontérték" in sample_text or "kardióterhelés" in sample_text:
        return "Polar Team Pro"
    if "catapult" in name or all(x in sample_text for x in ["athlete_name", "activity_name", "total_player_load"]):
        return "Catapult"
    return "Egyéb / Smart Mapper"


def _fpi_session_type_from_hint_v143(series: pd.Series, forced_type: Optional[str], file_name: str = "") -> pd.Series:
    if forced_type in {"Edzés", "Meccs"}:
        return pd.Series([forced_type] * len(series), index=series.index)
    hint = str(file_name).lower()
    if any(x in hint for x in ["match", "game", "meccs"]):
        default = "Meccs"
    elif any(x in hint for x in ["training", "train", "edzes", "edzés"]):
        default = "Edzés"
    else:
        default = "Edzés"
    if series is None:
        return pd.Series(dtype="object")
    out = series.astype(str).apply(normalize_session_type)
    out = out.where(out.isin(["Edzés", "Meccs"]), default)
    return out


def _fpi_prepare_brainsports_v143(
    sheets: Dict[str, pd.DataFrame],
    forced_type: Optional[str],
    file_name: str,
) -> pd.DataFrame:
    candidates = []
    for sheet_name, raw in sheets.items():
        low_name = str(sheet_name).lower()
        if low_name.startswith("ts ") or low_name.startswith("ti "):
            continue
        header_idx = _fpi_find_header_row_v143(raw, ["name", "total distance", "duration"])
        if header_idx is None:
            continue
        table = _fpi_table_from_raw_v143(
            raw,
            header_idx,
            stop_markers=["targets summary", "targeted values", "targets intensity"],
        )
        if table.empty:
            continue
        cols_norm = {clean_col_name(c).lower(): c for c in table.columns}
        if not any("total distance" in k for k in cols_norm):
            continue
        candidates.append((sheet_name, table))

    if not candidates:
        return pd.DataFrame()

    # Prefer the aggregate "Main table..." sheet. Player-specific tabs duplicate rows.
    aggregate = [item for item in candidates if str(item[0]).lower().startswith("main table")]
    selected = aggregate[:1] if aggregate else candidates[:1]
    table = selected[0][1].copy()

    def col_like(*needles):
        for c in table.columns:
            lc = clean_col_name(c).lower()
            if all(n.lower() in lc for n in needles):
                return c
        return None

    out = pd.DataFrame()
    c_name = col_like("name")
    c_split = col_like("split")
    c_start = col_like("start")
    c_duration = col_like("duration")
    c_dist = col_like("total distance")
    c_sprint_dist = col_like("total sprints distance")
    c_sprints = col_like("sprints count")
    c_hsr = col_like("distance(4+5)")
    c_top = col_like("top speed")
    c_he = col_like("high efforts")
    c_hmld = col_like("hmld")
    c_acc = col_like("total accelerations")
    c_dec = col_like("total decelerations")
    c_hr = col_like("heart exertion")

    if c_name is None or c_start is None:
        return pd.DataFrame()

    out["Player Name"] = table[c_name]
    out["Session Type"] = _fpi_session_type_from_hint_v143(
        table[c_split] if c_split else pd.Series([""] * len(table)),
        forced_type,
        file_name,
    )
    out["Session Name"] = table[c_split] if c_split else Path(file_name).stem
    out["Start Time"] = table[c_start].apply(_fpi_excel_serial_to_datetime_v143)
    if c_duration:
        out["Duration"] = table[c_duration]
    if c_dist:
        out["Total Distance"] = table[c_dist]
    if c_sprint_dist:
        out["Sprint Distance"] = table[c_sprint_dist]
    if c_sprints:
        out["Sprints"] = table[c_sprints]
    if c_hsr:
        # Brainsports Distance(4+5) already combines HSR + sprint.
        out["HSR Distance"] = table[c_hsr]
    if c_top:
        top_speed = pd.to_numeric(table[c_top], errors="coerce")
        out["Top Speed"] = top_speed.where(top_speed > 15, top_speed * 3.6)
    if c_he:
        out["High Efforts"] = table[c_he]
    if c_hmld:
        out["HMLD"] = table[c_hmld]
        out["Player Load"] = table[c_hmld]
    if c_acc:
        out["Total Accelerations"] = table[c_acc]
    if c_dec:
        out["Total Decelerations"] = table[c_dec]
    if c_hr:
        out["Heart Exertion"] = table[c_hr]
    return out.dropna(how="all").reset_index(drop=True)


def _fpi_prepare_playertek_v143(
    raw: pd.DataFrame,
    forced_type: Optional[str],
    file_name: str,
) -> pd.DataFrame:
    header_idx = _fpi_find_header_row_v143(raw, ["player name", "distance", "player load"])
    if header_idx is None:
        return pd.DataFrame()
    table = _fpi_table_from_raw_v143(raw, header_idx)
    if table.empty:
        return table

    def pick(*names):
        norm = {clean_col_name(c).lower(): c for c in table.columns}
        for name in names:
            n = clean_col_name(name).lower()
            if n in norm:
                return norm[n]
        for c in table.columns:
            lc = clean_col_name(c).lower()
            if any(clean_col_name(n).lower() in lc for n in names):
                return c
        return None

    out = pd.DataFrame()
    c_name = pick("Player Name")
    c_date = pick("Date")
    c_title = pick("Session Title")
    c_tags = pick("Tags")
    c_duration = pick("Duration")
    c_dist = pick("Distance (km)")
    c_sprints = pick("Sprints")
    c_sprint_dist = pick("Sprint Distance (m)")
    c_acc = pick("Accelerations")
    c_dec = pick("Decelerations")
    c_load = pick("Player Load")
    c_speed = pick("Top Speed (km/h)")
    c_dpm = pick("Distance Per Min (m/min)")
    if c_name is None or c_date is None:
        return pd.DataFrame()

    out["Player Name"] = table[c_name]
    hint_series = table[c_tags] if c_tags else pd.Series([""] * len(table))
    out["Session Type"] = _fpi_session_type_from_hint_v143(hint_series, forced_type, file_name)
    out["Session Name"] = table[c_title] if c_title else Path(file_name).stem
    out["Start Time"] = table[c_date].apply(_fpi_excel_serial_to_datetime_v143)
    if c_duration:
        # PlayerTek duration is seconds in the supplied export.
        out["Duration"] = pd.to_numeric(table[c_duration], errors="coerce") / 60.0
    if c_dist:
        out["Total Distance"] = pd.to_numeric(table[c_dist], errors="coerce") * 1000.0
    if c_sprints:
        out["Sprints"] = table[c_sprints]
    if c_sprint_dist:
        out["Sprint Distance"] = table[c_sprint_dist]
    if c_acc:
        out["Total Accelerations"] = table[c_acc]
    if c_dec:
        out["Total Decelerations"] = table[c_dec]
    if c_load:
        out["Player Load"] = table[c_load]
    if c_speed:
        out["Top Speed"] = table[c_speed]
    if c_dpm:
        out["Distance Per Min"] = table[c_dpm]
    return out.dropna(how="all").reset_index(drop=True)


def _fpi_prepare_polar_v143(
    sheets: Dict[str, pd.DataFrame],
    forced_type: Optional[str],
    file_name: str,
) -> pd.DataFrame:
    chosen = None
    for sheet_name, raw in sheets.items():
        header_idx = _fpi_find_header_row_v143(raw, ["játékos neve", "teljes táv", "kezdési idő"])
        if header_idx is not None:
            chosen = _fpi_table_from_raw_v143(raw, header_idx)
            if not chosen.empty:
                break
    if chosen is None or chosen.empty:
        return pd.DataFrame()
    # Existing aliases already understand Polar Hungarian headings.
    # Only force Training/Match when the user used the separated uploader.
    type_col = None
    for c in chosen.columns:
        if clean_col_name(c).lower() in {"típus", "tipus", "type"}:
            type_col = c
            break
    if forced_type in {"Edzés", "Meccs"}:
        if type_col:
            chosen[type_col] = forced_type
        else:
            chosen["Típus"] = forced_type
    return chosen.reset_index(drop=True)


def _fpi_prepare_catapult_v143(
    sheets: Dict[str, pd.DataFrame],
    forced_type: Optional[str],
    file_name: str,
) -> pd.DataFrame:
    raw_table = None
    for _, raw in sheets.items():
        header_idx = _fpi_find_header_row_v143(raw, ["athlete_name", "activity_name", "total_distance"])
        if header_idx is not None:
            raw_table = _fpi_table_from_raw_v143(raw, header_idx)
            if not raw_table.empty:
                break
    if raw_table is None or raw_table.empty:
        return pd.DataFrame()

    table = raw_table.copy()
    table.columns = [clean_col_name(c) for c in table.columns]
    lookup = {c.lower(): c for c in table.columns}

    def c(name):
        return lookup.get(clean_col_name(name).lower())

    athlete = c("athlete_name")
    activity = c("activity_name")
    date_col = c("date")
    if athlete is None or activity is None or date_col is None:
        return pd.DataFrame()

    numeric_names = [
        "total_duration", "total_distance", "velocity2_band3_total_distance",
        "velocity2_band4_total_distance", "velocity2_band3_total_effort_count",
        "velocity2_band4_total_effort_count", "max_vel", "hsr_efforts",
        "gen2_acceleration_band6plus_total_effort_count",
        "gen2_acceleration_band3plus_total_effort_count",
        "total_player_load", "mean_heart_rate", "max_heart_rate",
    ]
    for name in numeric_names:
        col = c(name)
        if col:
            table[col] = pd.to_numeric(table[col], errors="coerce")

    group_cols = [athlete, activity, date_col]
    pos_col = c("position_name")
    day_col = c("day_name")
    if pos_col:
        group_cols.append(pos_col)
    if day_col:
        group_cols.append(day_col)

    agg = {}
    for name in [
        "total_duration", "total_distance", "velocity2_band3_total_distance",
        "velocity2_band4_total_distance", "velocity2_band3_total_effort_count",
        "velocity2_band4_total_effort_count", "hsr_efforts",
        "gen2_acceleration_band6plus_total_effort_count",
        "gen2_acceleration_band3plus_total_effort_count",
        "total_player_load",
    ]:
        col = c(name)
        if col:
            agg[col] = "sum"
    for name in ["max_vel", "max_heart_rate"]:
        col = c(name)
        if col:
            agg[col] = "max"
    mean_hr = c("mean_heart_rate")
    if mean_hr:
        agg[mean_hr] = "mean"

    grouped = table.groupby(group_cols, dropna=False).agg(agg).reset_index()
    out = pd.DataFrame()
    out["Player Name"] = grouped[athlete]
    out["Session Name"] = grouped[activity]
    if forced_type in {"Edzés", "Meccs"}:
        out["Session Type"] = forced_type
    else:
        hint = grouped[day_col] if day_col and day_col in grouped.columns else grouped[activity]
        out["Session Type"] = hint.astype(str).apply(normalize_session_type)
        inferred = grouped[activity].astype(str).str.lower().str.contains("match|game|meccs", regex=True)
        out.loc[inferred, "Session Type"] = "Meccs"
        out.loc[~inferred & ~out["Session Type"].isin(["Edzés", "Meccs"]), "Session Type"] = "Edzés"
    out["Start Time"] = grouped[date_col].apply(_fpi_excel_serial_to_datetime_v143)
    if pos_col and pos_col in grouped:
        out["Position"] = grouped[pos_col]
    col = c("total_duration")
    if col and col in grouped:
        out["Duration"] = grouped[col] / 60.0
    col = c("total_distance")
    if col and col in grouped:
        out["Total Distance"] = grouped[col]
    v3 = c("velocity2_band3_total_distance")
    v4 = c("velocity2_band4_total_distance")
    if v3 and v3 in grouped:
        out["HSR Distance"] = grouped[v3]
    if v4 and v4 in grouped:
        out["Sprint Distance"] = grouped[v4]
    e3 = c("velocity2_band3_total_effort_count")
    e4 = c("velocity2_band4_total_effort_count")
    if e4 and e4 in grouped:
        out["Sprints"] = grouped[e4]
    elif e3 and e3 in grouped:
        out["Sprints"] = grouped[e3]
    max_vel = c("max_vel")
    if max_vel and max_vel in grouped:
        out["Top Speed"] = grouped[max_vel] * 3.6
    he = c("hsr_efforts")
    if he and he in grouped:
        out["High Efforts"] = grouped[he]
    acc = c("gen2_acceleration_band3plus_total_effort_count")
    if acc and acc in grouped:
        out["Total Accelerations"] = grouped[acc]
    load = c("total_player_load")
    if load and load in grouped:
        out["Player Load"] = grouped[load]
    hr_avg = c("mean_heart_rate")
    hr_max = c("max_heart_rate")
    if hr_avg and hr_avg in grouped:
        out["Average HR"] = grouped[hr_avg]
    if hr_max and hr_max in grouped:
        out["Max HR"] = grouped[hr_max]
    return out.reset_index(drop=True)


def _fpi_read_single_gps_file_v143(
    uploaded_file,
    forced_type: Optional[str] = None,
    provider_override: str = "Automatikus felismerés",
) -> Tuple[Dict[str, pd.DataFrame], List[Dict[str, object]]]:
    name = str(getattr(uploaded_file, "name", "gps_file"))
    data = _fpi_bytes_from_upload_v143(uploaded_file)
    ext = Path(name).suffix.lower()
    report = []
    if not data:
        return {}, [{"Fájl": name, "Státusz": "HIBA", "Rendszer": "", "Sorok": 0, "Megjegyzés": "Üres fájl"}]

    if ext == ".zip":
        sheets = {}
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for member in zf.namelist():
                    if member.endswith("/") or member.startswith("__MACOSX/"):
                        continue
                    m_ext = Path(member).suffix.lower()
                    if m_ext not in {".xlsx", ".xls", ".xlsm", ".csv"}:
                        continue
                    member_bytes = zf.read(member)
                    pseudo = type("FPIUpload", (), {
                        "name": Path(member).name,
                        "getvalue": lambda self, b=member_bytes: b,
                    })()
                    inner_sheets, inner_report = _fpi_read_single_gps_file_v143(
                        pseudo,
                        forced_type=forced_type,
                        provider_override=provider_override,
                    )
                    for key, frame in inner_sheets.items():
                        sheets[f"{Path(member).stem}__{key}"] = frame
                    for row in inner_report:
                        row["Fájl"] = member
                        report.append(row)
            return sheets, report
        except Exception as exc:
            return {}, [{"Fájl": name, "Státusz": "HIBA", "Rendszer": "", "Sorok": 0, "Megjegyzés": str(exc)}]

    try:
        if ext == ".csv":
            raw_sheets = {"CSV": _fpi_read_csv_bytes_v143(data, name)}
        elif ext in {".xlsx", ".xlsm", ".xls"}:
            raw_sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, header=None)
        else:
            return {}, [{"Fájl": name, "Státusz": "KIHAGYVA", "Rendszer": "", "Sorok": 0, "Megjegyzés": "Nem támogatott kiterjesztés"}]

        provider = provider_override if provider_override != "Automatikus felismerés" else _fpi_detect_provider_v143(raw_sheets, name)
        prepared = pd.DataFrame()
        if provider == "Brainsports":
            prepared = _fpi_prepare_brainsports_v143(raw_sheets, forced_type, name)
        elif provider == "PlayerTek":
            raw = next(iter(raw_sheets.values()), pd.DataFrame())
            prepared = _fpi_prepare_playertek_v143(raw, forced_type, name)
        elif provider == "Polar Team Pro":
            prepared = _fpi_prepare_polar_v143(raw_sheets, forced_type, name)
        elif provider == "Catapult":
            prepared = _fpi_prepare_catapult_v143(raw_sheets, forced_type, name)

        if not prepared.empty:
            key = f"{Path(name).stem}__{provider}"
            return {key: prepared}, [{
                "Fájl": name,
                "Státusz": "OK",
                "Rendszer": provider,
                "Sorok": len(prepared),
                "Megjegyzés": "Automatikusan standardizált",
            }]

        # Generic fallback: keep all sheets and let the existing header detector + Smart Mapper work.
        generic = {}
        for sheet_name, raw in raw_sheets.items():
            if raw is None or raw.empty:
                continue
            generic[f"{Path(name).stem}__{sheet_name}"] = raw
        return generic, [{
            "Fájl": name,
            "Státusz": "MAPPER",
            "Rendszer": provider,
            "Sorok": sum(len(x) for x in generic.values()),
            "Megjegyzés": "Általános Smart Mapper szükséges",
        }]
    except Exception as exc:
        return {}, [{"Fájl": name, "Státusz": "HIBA", "Rendszer": "", "Sorok": 0, "Megjegyzés": str(exc)}]


def _fpi_read_many_gps_files_v143(
    training_files: Optional[List[object]],
    match_files: Optional[List[object]],
    mixed_files: Optional[List[object]],
    provider_override: str = "Automatikus felismerés",
) -> Tuple[Dict[str, pd.DataFrame], pd.DataFrame, str]:
    all_sheets: Dict[str, pd.DataFrame] = {}
    rows: List[Dict[str, object]] = []
    signature_parts = []

    groups = [
        ("Edzés", training_files or []),
        ("Meccs", match_files or []),
        (None, mixed_files or []),
    ]
    for forced_type, files in groups:
        for upload in files:
            name = str(getattr(upload, "name", "gps_file"))
            data = _fpi_bytes_from_upload_v143(upload)
            signature_parts.append(f"{forced_type}:{name}:{len(data)}:{hashlib.md5(data[:200000]).hexdigest() if data else ''}")
            sheets, report = _fpi_read_single_gps_file_v143(
                upload,
                forced_type=forced_type,
                provider_override=provider_override,
            )
            for key, frame in sheets.items():
                unique = key
                counter = 2
                while unique in all_sheets:
                    unique = f"{key}_{counter}"
                    counter += 1
                all_sheets[unique] = frame
            for row in report:
                row["Import típus"] = forced_type or "Vegyes / fájlból"
                rows.append(row)

    return all_sheets, pd.DataFrame(rows), hashlib.md5("|".join(signature_parts).encode("utf-8")).hexdigest()


def _fpi_team_level_tactical_messages_v143(
    tactical_context: Dict[str, object],
    tactical_msgs: List[str],
    limit: int = 4,
) -> List[str]:
    opponent_players = []
    for row in (tactical_context or {}).get("opponent_player_evaluation", []) or []:
        name = str(row.get("Játékos", "")).strip().lower()
        if name:
            opponent_players.append(name)

    cleaned = []
    for msg in tactical_msgs or []:
        text = _fpi_clean_sentence_v82(str(msg), 180)
        low = text.lower()
        if any(name in low for name in opponent_players):
            continue
        if text and text not in cleaned:
            cleaned.append(text)

    if not cleaned:
        plan = str((tactical_context or {}).get("plan_a", "KIE – kiegyensúlyozott meccsterv"))
        cleaned = [
            f"Csapatszintű alapirány: {plan}.",
            "A blokk-, presszing- és átmeneti viselkedést egységes csapatelvekkel kell végrehajtani.",
            "A labdavesztés utáni biztosítás és a második labdák kontrollja legyen közös csapatszintű prioritás.",
        ]
    return cleaned[:limit]



# =========================================================
# V144 - Magyar szakmai nyelvi és kontextusmotor
# =========================================================
def _fpi_hu_plain_text_v144(value: object) -> str:
    text = str(value or "")
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</?b>", "", text)
    text = text.replace("&nbsp;", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    replacements = {
        "pressing": "letámadás",
        "presszing": "letámadás",
        "transition": "átmenet",
        "rest defense": "labdavesztés elleni biztosítás",
        "build-up": "labdakihozatal",
        "high block": "magas védekezés",
        "mid block": "középső védekezési blokk",
        "low block": "mély védekezési blokk",
        "wide play": "szélső játék",
        "half-space": "félterület",
        "second ball": "második labda",
        "counterpress": "visszatámadás",
        "counter-press": "visszatámadás",
    }
    for source, target in replacements.items():
        text = re.sub(rf"(?i)\b{re.escape(source)}\b", target, text)
    return text.strip()


def _fpi_context_seed_v144(
    tactical_context: Optional[Dict[str, object]],
    week: Optional[str],
    readiness: Optional[int],
) -> int:
    ctx = tactical_context or {}
    opponent = str(
        ctx.get("opponent_name")
        or st.session_state.get("clean_opponent_v94", "")
        or st.session_state.get("fpi_match_context_v94", {}).get("opponent", "")
    )
    source = "|".join([
        str(week or ""),
        opponent,
        str(ctx.get("plan_a", "")),
        str(ctx.get("plan_b", "")),
        str(ctx.get("own_topics", "")),
        str(ctx.get("opp_topics", "")),
        str(readiness or ""),
    ])
    return int(hashlib.md5(source.encode("utf-8")).hexdigest()[:8], 16)


def _fpi_rotate_v144(items: List[str], seed: int) -> List[str]:
    clean_items = [x for x in items if str(x).strip()]
    if not clean_items:
        return []
    shift = seed % len(clean_items)
    return clean_items[shift:] + clean_items[:shift]


def _fpi_tactical_features_v144(tactical_context: Optional[Dict[str, object]]) -> Dict[str, object]:
    ctx = tactical_context or {}
    all_text = " ".join([
        str(ctx.get("plan_a", "")),
        str(ctx.get("plan_b", "")),
        " ".join(map(str, ctx.get("risks", []) or [])),
        " ".join(map(str, ctx.get("tactical_findings", []) or [])),
        " ".join(map(str, ctx.get("own_topics", []) or [])),
        " ".join(map(str, ctx.get("opp_topics", []) or [])),
        " ".join(map(str, ctx.get("pdf_provider_findings", []) or [])),
    ]).lower()

    def has(*terms):
        return any(term.lower() in all_text for term in terms)

    return {
        "high_press": has("magas letámadás", "high press", "mlt", "presszing", "letámadás"),
        "mid_block": has("középső blokk", "mid block", "bat"),
        "deep_block": has("mély blokk", "low block", "kon"),
        "transition": has("átmenet", "transition", "kontra", "gat", "bat"),
        "possession": has("dominancia", "labdabirtoklás", "pozíciós", "dom", "poz", "lab"),
        "wide": has("szélső", "beadás", "oldali", "wide", "wing"),
        "central": has("középső", "félterület", "vonalak között", "central", "half-space"),
        "set_piece": has("pontrúgás", "szöglet", "szabadrúgás", "set piece"),
        "build_up": has("labdakihozatal", "építkezés", "build-up", "build up"),
        "second_ball": has("második labda", "second ball", "direkt játék"),
        "plan": _fpi_hu_plain_text_v144(ctx.get("plan_a", "KIE – kiegyensúlyozott meccsterv")),
    }


def _fpi_team_level_tactical_messages_v144(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
    limit: int = 3,
) -> List[str]:
    ctx = tactical_context or {}
    features = _fpi_tactical_features_v144(ctx)
    seed = _fpi_context_seed_v144(ctx, week, readiness)
    opponent = str(
        ctx.get("opponent_name")
        or st.session_state.get("clean_opponent_v94", "")
        or st.session_state.get("fpi_match_context_v94", {}).get("opponent", "")
        or "az ellenfél"
    )

    candidates: List[Tuple[str, str]] = []
    plan = features["plan"]
    candidates.append(("alapirany", f"A mérkőzés alapiránya: {plan}."))

    if features["high_press"]:
        candidates.extend([
            ("letamadas", f"{opponent} labdakihozatalát már az első passzoknál érdemes irányítani; a csatár és a szélsők egyszerre zárják a belső passzsávokat."),
            ("letamadas", "A letámadás csak akkor legyen agresszív, ha a középpályás sor mögötte együtt tolódik; különben túl nagy terület nyílhat a vonalak között."),
            ("letamadas", "A visszapassz és a pontatlan oldalváltás legyen közös letámadási jel, hogy a csapat egyszerre induljon meg."),
        ])
    if features["mid_block"]:
        candidates.extend([
            ("blokk", f"{opponent} ellen a középső védekezési blokk adhat stabil kiindulópontot; a cél a belső folyosók lezárása és a játék oldalra terelése."),
            ("blokk", "A két csapatrész közötti távolság maradjon szűk, különösen akkor, amikor az ellenfél a vonalak közé keres passzt."),
            ("blokk", "A szélső oldali kilépést a mögötte lévő védő és a védekező középpályás biztosítsa, hogy ne nyíljon meg a félterület."),
        ])
    if features["transition"]:
        candidates.extend([
            ("atmenet", f"{opponent} labdavesztése után az első előre irányuló passz lehet a fő támadási lehetőség; a szélsők azonnal támadják a védők mögötti területet."),
            ("atmenet", "Saját labdavesztés után az első feladat a középső terület lezárása, csak ezután következzen a labdás játékos megtámadása."),
            ("atmenet", "A labdaszerzés utáni első két döntés legyen egyszerű: biztosító passz vagy azonnali mélységi játék, a helyzettől függően."),
        ])
    if features["possession"]:
        candidates.extend([
            ("labdabirtoklas", f"{opponent} ellen a labdabirtoklás önmagában nem cél; a labdajáratás a védelmi vonalak mozgatását és a félterületek megnyitását szolgálja."),
            ("labdabirtoklas", "A türelmes labdakihozatal mellett legyen előre meghatározott gyorsítási pont, amikor a szélső vagy a támadó középpályás vonalak között labdát kap."),
            ("labdabirtoklas", "A támadás közben legalább három játékos maradjon a labda mögött, hogy a labdavesztés ne vezessen azonnali létszámhátrányos kontrához."),
        ])
    if features["wide"]:
        candidates.extend([
            ("szelso", f"{opponent} szélső játékánál a beadás előtti helyzetet kell megakadályozni; a szélső védő ne maradjon egyedül az egy az egy elleni szituációkban."),
            ("szelso", "Saját támadásban a szélességet az egyik oldalon tartsuk meg, a túloldali szélső pedig időben érkezzen a hosszú oldalra."),
            ("szelso", "Az oldali létszámfölény után ne automatikusan beadás következzen: a visszagurítás és a félterületi passz is legyen előkészített megoldás."),
        ])
    if features["central"]:
        candidates.extend([
            ("kozepso", f"{opponent} vonalak közötti játékát a védekező középpályás és a belső védők közös felelősségeként kell kezelni."),
            ("kozepso", "A középső terület túlterhelésekor a labdás játékos nyomás alá helyezése mellett a mögötte lévő passzkapcsolatot is zárni kell."),
            ("kozepso", "Saját labdabirtoklásnál a támadó középpályás ne ugyanabban a magasságban kérje a labdát, mint a csatár; legyen két külön passzvonal."),
        ])
    if features["set_piece"]:
        candidates.extend([
            ("pontrugas", f"{opponent} pontrúgásainál a második labda környékét külön biztosítani kell; a felszabadítás után ne essen szét a csapat."),
            ("pontrugas", "A támadó pontrúgásoknál legyen kijelölt biztosító játékos a lepattanókra és az ellenfél gyors ellentámadásának megállítására."),
        ])
    if features["build_up"]:
        candidates.extend([
            ("labdakihozatal", f"{opponent} letámadását az első vonal mögötti szabad játékos megtalálásával lehet megbontani; a kapus is legyen aktív passzopció."),
            ("labdakihozatal", "A labdakihozatalnál ne minden játékos kérje egyszerre lábra a labdát: legalább egy mélységi mozgás húzza hátra az ellenfél védelmét."),
        ])
    if features["second_ball"]:
        candidates.extend([
            ("masodik_labda", f"{opponent} direkt játékánál a fejpárbaj önmagában nem elég; a második labda helyének előzetes biztosítása döntő lehet."),
            ("masodik_labda", "A hosszú labdák után a középpályás sor gyorsabban zárjon a párharc köré, hogy a lepattanó ne maradjon az ellenfélnél."),
        ])

    if readiness < 60:
        candidates.extend([
            ("allapot", "A jelenlegi fizikai állapot mellett a taktikai intenzitást rövidebb, pontosan adagolt blokkokban érdemes gyakorolni."),
            ("allapot", "A meccsterv ne épüljön folyamatos, egész pályás letámadásra; inkább előre kijelölt helyzetekben emeljük az intenzitást."),
        ])
    elif readiness >= 75:
        candidates.extend([
            ("allapot", "A csapat aktuális állapota lehetővé teszi az intenzívebb letámadási és átmeneti szakaszokat, de ezek időzítése maradjon tudatos."),
            ("allapot", "A megfelelő frissességre építve bátrabban vállalható magasabb védekezési vonal és gyorsabb labdaszerzés utáni játék."),
        ])
    else:
        candidates.append(("allapot", "A taktikai terv végrehajtható, de a nagy intenzitású szakaszokat célszerű előre kijelölt meccshelyzetekhez kötni."))

    # Existing direct findings are retained when they are natural and team-level.
    for raw in (ctx.get("tactical_findings", []) or []) + (ctx.get("risks", []) or []):
        txt = _fpi_hu_plain_text_v144(_fpi_clean_sentence_v82(raw, 180))
        if txt and not re.search(r"\b[A-ZÁÉÍÓÖŐÚÜŰ][a-záéíóöőúüű]+\s+[A-Z]\.", txt):
            candidates.append(("forras", txt))

    rotated = _fpi_rotate_v144([f"{topic}|||{text}" for topic, text in candidates], seed)
    selected, used_topics, seen = [], set(), set()
    for packed in rotated:
        topic, text = packed.split("|||", 1)
        norm = re.sub(r"\W+", " ", text.lower()).strip()
        if topic in used_topics or norm in seen:
            continue
        selected.append(text)
        used_topics.add(topic)
        seen.add(norm)
        if len(selected) >= limit:
            break
    return selected


def _fpi_fitness_topic_v144(text: str) -> str:
    low = _fpi_hu_plain_text_v144(text).lower()
    groups = [
        ("sebesseg", ["sprint", "hsr", "nagy sebess", "sebesség"]),
        ("volumen", ["össztáv", "volumen", "distance", "táv"]),
        ("terheles", ["load", "terhelési pont", "összterhelés"]),
        ("intenzitas", ["high effort", "intenzitás", "nagy intenzit"]),
        ("readiness", ["readiness", "készenlét", "frissesség"]),
        ("trend", ["4 het", "trend", "előző hetek"]),
        ("regeneracio", ["regener", "pihen", "fárad"]),
        ("pulzus", ["pulzus", "hrv", "heart"]),
        ("gyorsulas", ["gyorsul", "lassul", "neuromuszk"]),
        ("kockazat", ["kockázat", "risk", "figyelendő játékos"]),
    ]
    for topic, terms in groups:
        if any(term in low for term in terms):
            return topic
    return "egyeb"


def _fpi_diverse_fitness_messages_v144(
    ctx: Dict[str, object],
    priorities: List[dict],
    readiness: int,
    week: Optional[str],
    limit: int = 3,
) -> List[str]:
    base = _fpi_top_fitness_messages_v82(ctx, priorities, readiness, 8)
    candidates: List[Tuple[str, str]] = []

    for msg in base:
        clean_msg = _fpi_hu_plain_text_v144(_fpi_clean_sentence_v82(msg, 180))
        if clean_msg:
            candidates.append((_fpi_fitness_topic_v144(clean_msg), clean_msg))

    priority_text = " ".join([
        str(p.get("Teendő", p.get("Cím", p.get("title", ""))))
        for p in priorities or [] if isinstance(p, dict)
    ]).lower()

    if readiness < 55:
        candidates.extend([
            ("readiness", "A csökkent készenléti érték miatt ezen a héten a frissesség visszaépítése élvezzen elsőbbséget."),
            ("regeneracio", "A terhelés után több regenerációs időt és rövidebb, jól elkülönített intenzív blokkokat érdemes tervezni."),
        ])
    elif readiness < 70:
        candidates.append(("readiness", "A jelenlegi készenléti állapot megfelelő, de az intenzív napok után célzott visszatöltésre van szükség."))
    else:
        candidates.append(("readiness", "A csapat frissessége megfelelő alapot ad a tervezett heti fő terhelési ingerhez."))

    if "sprint" in priority_text or "hsr" in priority_text or "nagy sebess" in priority_text:
        candidates.append(("sebesseg", "A sebességi terhelés elmarad a kívánt szinttől; egy kontrollált HSR- és sprintinger indokolt, alacsony ismétlésszámmal."))
    else:
        candidates.append(("sebesseg", "A nagy sebességű futások mennyisége megfelelőnek tűnik; a következő héten inkább a minőség és a teljes regeneráció legyen a cél."))

    if "load" in priority_text or "terhel" in priority_text:
        candidates.append(("terheles", "A heti terhelési pontok eloszlását érdemes egyenletesebbé tenni, hogy ne egyetlen napra kerüljön a fő terhelés."))
    else:
        candidates.append(("terheles", "A heti összterhelés stabil alapot ad; a fő terhelési nap és a meccs közötti távolságot tartsuk meg."))

    candidates.extend([
        ("trend", "A négyhetes trendet a következő héten is érdemes megtartani; hirtelen terhelésugrás helyett fokozatos emelés javasolt."),
        ("volumen", "Az össztáv önmagában nem döntő: a heti volumen mellé megfelelő sebességi és intenzitási inger is szükséges."),
        ("intenzitas", "A nagy intenzitású akciókat rövid, meccsszerű blokkokban érdemes adagolni, hogy a minőség ne romoljon."),
        ("gyorsulas", "A gyorsítások és lassítások száma külön neuromuszkuláris terhelést jelent; ezt a pályaméret és a játékforma megválasztásánál vegyük figyelembe."),
        ("regeneracio", "Az utolsó erősebb edzés után legyen elegendő idő a frissülésre, különösen a magas kockázatú játékosoknál."),
    ])

    seed = _fpi_context_seed_v144({}, week, readiness)
    rotated = _fpi_rotate_v144([f"{topic}|||{text}" for topic, text in candidates], seed)
    selected, used_topics, seen = [], set(), set()
    for packed in rotated:
        topic, text = packed.split("|||", 1)
        norm = re.sub(r"\W+", " ", text.lower()).strip()
        if topic in used_topics or norm in seen:
            continue
        selected.append(text)
        used_topics.add(topic)
        seen.add(norm)
        if len(selected) >= limit:
            break
    return selected



# =========================================================
# V145 - Meccsspecifikus edzői tudásmotor és ismétléskontroll
# =========================================================
def _fpi_extract_coach_text_v145(value: object, max_len: int = 260) -> str:
    """Nyers dict/list repr helyett csak az edző számára használható következtetés."""
    if value is None:
        return ""

    if isinstance(value, dict):
        for key in [
            "Edzői következtetés", "Edzöi következtetés",
            "Javaslat", "Értelmezés", "Következtetés",
            "Megállapítás", "text", "message",
        ]:
            candidate = value.get(key)
            if candidate:
                return _fpi_clean_sentence_v82(_fpi_hu_plain_text_v144(candidate), max_len)
        theme = value.get("Téma") or value.get("Tema") or ""
        evidence = value.get("Bizonyíték") or value.get("Bizonyitek") or ""
        if theme and evidence:
            return _fpi_clean_sentence_v82(
                f"{_fpi_hu_plain_text_v144(theme)}: {_fpi_hu_plain_text_v144(evidence)}",
                max_len,
            )
        return ""

    if isinstance(value, (list, tuple, set)):
        parts = [_fpi_extract_coach_text_v145(item, max_len) for item in value]
        return _fpi_clean_sentence_v82(" ".join(x for x in parts if x), max_len)

    text = _fpi_hu_plain_text_v144(value)
    # Python dictionary/list representation kiszűrése.
    if re.search(r"['\"]Téma['\"]\s*:", text) or re.search(r"['\"]Bizonyíték['\"]\s*:", text):
        try:
            import ast
            parsed = ast.literal_eval(str(value))
            return _fpi_extract_coach_text_v145(parsed, max_len)
        except Exception:
            text = re.sub(r"\{.*?['\"]Edz[őö]i következtetés['\"]\s*:\s*['\"]([^'\"]+)['\"].*?\}", r"\1", text)
            text = re.sub(r"\{.*?\}", "", text)
    return _fpi_clean_sentence_v82(text, max_len)


def _fpi_message_signature_v145(text: object) -> set:
    clean = _fpi_extract_coach_text_v145(text, 500).lower()
    clean = unicodedata.normalize("NFKD", clean)
    clean = "".join(ch for ch in clean if not unicodedata.combining(ch))
    stop = {
        "a", "az", "egy", "es", "hogy", "is", "kell", "legyen", "lehet", "utan",
        "ellenfel", "csapat", "merkozes", "erdemes", "fontos", "jatek", "sajat",
    }
    return {
        token for token in re.findall(r"[a-z0-9]+", clean)
        if len(token) >= 4 and token not in stop
    }


def _fpi_is_near_duplicate_v145(text: object, accepted: List[str], threshold: float = 0.55) -> bool:
    sig = _fpi_message_signature_v145(text)
    if not sig:
        return True
    for existing in accepted:
        other = _fpi_message_signature_v145(existing)
        if not other:
            continue
        similarity = len(sig & other) / max(1, min(len(sig), len(other)))
        if similarity >= threshold:
            return True
    return False


def _fpi_unique_messages_v145(
    messages: List[object],
    limit: int,
    already_used: Optional[List[str]] = None,
) -> List[str]:
    accepted = list(already_used or [])
    result: List[str] = []
    for raw in messages or []:
        text = _fpi_extract_coach_text_v145(raw, 320)
        if not text or _fpi_is_near_duplicate_v145(text, accepted, 0.55):
            continue
        result.append(text)
        accepted.append(text)
        if len(result) >= limit:
            break
    return result


def _fpi_opponent_name_v145(tactical_context: Optional[Dict[str, object]]) -> str:
    ctx = tactical_context or {}
    return str(
        ctx.get("opponent_name")
        or st.session_state.get("clean_opponent_v94", "")
        or st.session_state.get("fpi_match_context_v94", {}).get("opponent", "")
        or "az ellenfél"
    ).strip()


def _fpi_match_plan_messages_v145(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
    limit: int = 4,
) -> List[str]:
    """Rövid, de ténylegesen használható és az adott ellenfélhez kötött meccsterv."""
    ctx = tactical_context or {}
    features = _fpi_tactical_features_v144(ctx)
    opponent = _fpi_opponent_name_v145(ctx)
    plan = _fpi_hu_plain_text_v144(ctx.get("plan_a", "KIE – kiegyensúlyozott meccsterv"))
    seed = _fpi_context_seed_v144(ctx, week, readiness)

    candidates: List[Tuple[str, str]] = [
        (
            "alapirany",
            f"Alapirány: {plan}. A csapat a mérkőzés elején ebből a szerkezetből induljon, "
            f"és csak akkor váltson magasabb kockázatú játékra, ha {opponent} első építési vonala bizonytalanná válik."
        ),
    ]

    if features["high_press"]:
        candidates.extend([
            (
                "letamadas",
                f"Letámadás: {opponent} kapushoz vagy belső védőhöz visszajátszott labdája legyen közös indítójel. "
                "A csatár terelje oldalra a játékot, a szélső zárja a szélső védőt, a belső középpályás pedig vegye el a visszapasszt."
            ),
            (
                "letamadas",
                "A magasabb letámadást ne folyamatosan alkalmazzuk. A rossz labdaátvétel, a háttal álló középpályás és az oldalvonal mellé szorult labdás legyen a három fő kiváltó helyzet."
            ),
        ])

    if features["mid_block"]:
        candidates.extend([
            (
                "blokk",
                f"Védekezési blokk: {opponent} ellen a középső blokkban a belső passzsávok lezárása az elsődleges. "
                "A labdát tereljük oldalra, majd a szélső, a szélső védő és a közeli középpályás együtt szűkítse le a játékteret."
            ),
            (
                "blokk",
                "A csapatrészek közötti távolság ne nyúljon meg. Ha az első sor nem tud nyomást gyakorolni, az egész csapat lépjen vissza öt-tíz métert, ne csak a védősor."
            ),
        ])

    if features["transition"]:
        candidates.extend([
            (
                "tamado_atmenet",
                f"Labdaszerzés után: {opponent} rendezetlen védelme ellen az első pillantás a mélységi területre irányuljon. "
                "Ha nincs tiszta előrejáték, egy biztosító passzal tartsuk meg a labdát, ne erőltessük a második kockázatos átadást."
            ),
            (
                "vedekezo_atmenet",
                "Labdavesztés után: a legközelebbi játékos lassítsa a labdást, a többiek először a középső passzsávot és a csatár felé vezető utat zárják. "
                "A cél nem minden helyzetben az azonnali labdaszerzés, hanem az ellenfél első előrepasszának megakadályozása."
            ),
        ])

    if features["possession"]:
        candidates.extend([
            (
                "labdakihozatal",
                f"Labdabirtoklásban: {opponent} blokkja ellen a labdajáratás akkor értékes, ha oldalváltás vagy félterületi belépés követi. "
                "Az egyik belső középpályás maradjon a labda mögött, hogy a támadás közben is megmaradjon a biztosítás."
            ),
            (
                "tamadas",
                "A támadásokat ne az első szabad beadással fejezzük be. Először keressük a visszagurítást vagy a tizenhatos előtti második hullámot, különösen akkor, ha a büntetőterület zsúfolt."
            ),
        ])

    if features["wide"]:
        candidates.append((
            "szelso",
            f"Szélső védekezés: {opponent} beadásai előtt a labdás játékost kell lassítani. "
            "A közeli szélső segítsen a kilépő védőnek, a túloldali szélső pedig időben zárja a hosszú oldalt."
        ))

    if features["central"]:
        candidates.append((
            "kozepso",
            f"Középső terület: {opponent} vonalak közötti játékosát ne egyetlen védő kövesse. "
            "A védekező középpályás adjon nyomást, a belső védő pedig biztosítsa a mögötte nyíló területet."
        ))

    if features["set_piece"]:
        candidates.append((
            "pontrugas",
            "Pontrúgások: a közvetlen párharcok mellett előre jelöljük ki a lepattanó és a második labda felelőseit. "
            "Felszabadítás után a csapat együtt lépjen ki, hogy ne maradjon újabb beadási lehetőség."
        ))

    if readiness < 60:
        candidates.append((
            "allapot",
            "A csapat aktuális terhelési állapota miatt a meccsterv ne épüljön kilencven perces folyamatos nyomásra. "
            "Jelöljünk ki rövidebb, előre meghatározott intenzív szakaszokat, közöttük stabilabb blokkban védekezve."
        ))
    elif readiness >= 75:
        candidates.append((
            "allapot",
            "A megfelelő frissesség lehetővé teszi az intenzívebb kezdést és a labdavesztés utáni gyorsabb visszatámadást. "
            "Az első húsz percben érdemes nagyobb nyomást helyezni az ellenfél labdakihozatalára."
        ))
    else:
        candidates.append((
            "allapot",
            "A fizikai állapot megfelelő a terv végrehajtásához, de a nagy intenzitású szakaszokat konkrét játékhelyzetekhez kössük, ne folyamatosan alkalmazzuk."
        ))

    rotated = _fpi_rotate_v144([f"{topic}|||{text}" for topic, text in candidates], seed)
    selected, used_topics = [], set()
    for packed in rotated:
        topic, text = packed.split("|||", 1)
        if topic in used_topics:
            continue
        selected.append(text)
        used_topics.add(topic)
        if len(selected) >= limit:
            break
    return selected


def _fpi_contextual_md_plan_rows_v145(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
) -> List[Tuple[str, str, str]]:
    """Minden MD-naphoz konkrét, ellenfélhez és heti állapothoz igazított taktikai cél."""
    base_rows = _fpi_md_plan_rows_v82(
        tactical_context,
        gps_context=gps_context,
        readiness=readiness,
        priorities=priorities,
        week=week,
    )
    if not _fpi_has_tactical_signal_v95(tactical_context):
        return base_rows

    ctx = tactical_context or {}
    features = _fpi_tactical_features_v144(ctx)
    opponent = _fpi_opponent_name_v145(ctx)
    plan = _fpi_hu_plain_text_v144(ctx.get("plan_a", "kiegyensúlyozott meccsterv"))

    result = []
    for day, fitness_goal, original_tactical in base_rows:
        day_text = str(day)
        original = _fpi_extract_coach_text_v145(original_tactical, 140)

        if "MD-4" in day_text:
            if features["build_up"] or features["high_press"]:
                tactical = (
                    f"{opponent} első nyomásának kijátszása és a saját letámadási jelek gyakorlása. "
                    "Nagyobb létszámú játékban rögzítsük, ki tereli oldalra a labdát, és ki zárja a belső visszapasszt."
                )
            else:
                tactical = (
                    f"A {plan} teljes csapatszerkezetének gyakorlása nagyobb területen. "
                    "A csapatrészek közötti távolság és a labda mögötti biztosítás legyen a fő ellenőrzési pont."
                )
        elif "MD-3" in day_text:
            if features["transition"]:
                tactical = (
                    f"{opponent} elleni támadó és védekező átmenetek meccssebességen. "
                    "Labdaszerzés után mélységi döntés, labdavesztés után középső passzsávzárás és visszarendeződés."
                )
            elif features["wide"]:
                tactical = (
                    f"Oldali létszámfölények és beadás előtti védekezés {opponent} szélső veszélyére készülve. "
                    "A gyakorlat végén mindig értékeljük a hosszú oldal és a lepattanó biztosítását."
                )
            else:
                tactical = (
                    f"Az ellenfél-specifikus fő veszélyek gyakorlása magas intenzitással. "
                    f"Az eredeti fókusz: {original or plan}."
                )
        elif "MD-2" in day_text:
            if features["mid_block"] or features["deep_block"]:
                tactical = (
                    f"Középső blokk, oldalra terelés és a vonalak közötti terület lezárása {opponent} várható támadási irányai ellen. "
                    "Rövidebb blokkokban rögzítsük a kilépés és a biztosítás sorrendjét."
                )
            else:
                tactical = (
                    f"Meccstervi részletek: labdakihozatal, első letámadási jel és labdavesztés utáni biztosítás. "
                    "A terhelés csökkenjen, a taktikai pontosság maradjon magas."
                )
        elif "MD-1" in day_text:
            tactical = (
                f"Rövid, alacsony volumenű meccsaktiváció: kezdő helyezkedések, első öt perc forgatókönyve, pontrúgások és a {plan} legfontosabb két szabályának ismétlése."
            )
        else:
            tactical = original or f"{opponent} elleni meccsterv pontosítása."

        result.append((day_text, fitness_goal, _fpi_clean_sentence_v82(tactical, 320)))
    return result


def _fpi_fitness_messages_v145(
    gps_context: Dict[str, object],
    priorities: List[dict],
    readiness: int,
    week: Optional[str],
    limit: int = 3,
) -> List[str]:
    """Három különböző erőnléti téma, kissé részletesebb edzői következtetéssel."""
    short = _fpi_diverse_fitness_messages_v144(gps_context, priorities, readiness, week, 8)
    expanded = []
    for message in short:
        topic = _fpi_fitness_topic_v144(message)
        text = _fpi_extract_coach_text_v145(message, 260)
        suffix = {
            "sebesseg": " A sebességi ingert kevés ismétléssel, teljes pihenőkkel és jó technikai minőség mellett adjuk.",
            "volumen": " A napi terhelést ne csak össztávval, hanem intenzitási eloszlással együtt értékeljük.",
            "terheles": " A fő terhelési nap után legyen egyértelmű visszalépés, hogy a fáradtság ne tolódjon a meccs közelébe.",
            "intenzitas": " A gyakorlatok legyenek rövidek és meccsszerűek; a minőség romlásakor a blokkot zárjuk le.",
            "readiness": " A játékosonkénti eltérések miatt az egyéni módosítás fontosabb lehet, mint a teljes csapat azonos terhelése.",
            "trend": " A következő hét célja a folytonosság megtartása, nem egyetlen kiugró edzés létrehozása.",
            "regeneracio": " A regenerációt a következő napi állapot és az egyéni kockázati jelzések alapján ellenőrizzük.",
            "gyorsulas": " Kis területű játékoknál külön figyeljünk a sok fékezésre és irányváltásra.",
            "pulzus": " A pulzusadatot a külső terheléssel együtt értelmezzük, önmagában ne döntsön.",
            "kockazat": " A magasabb kockázatú játékosoknál csökkentsük az ismétlésszámot, de a szükséges sebességi minőséget tartsuk meg.",
        }.get(topic, "")
        expanded.append((topic, _fpi_clean_sentence_v82(text + suffix, 330)))

    selected, used_topics, accepted = [], set(), []
    for topic, text in expanded:
        if topic in used_topics or _fpi_is_near_duplicate_v145(text, accepted, 0.48):
            continue
        selected.append(text)
        used_topics.add(topic)
        accepted.append(text)
        if len(selected) >= limit:
            break
    return selected


def _fpi_build_executive_blocks_v145(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
) -> Dict[str, List[str]]:
    """Blokkonként egyedi információk; ugyanaz a motívum ne jelenjen meg több helyen."""
    ctx = tactical_context or {}
    used: List[str] = []

    tactical_candidates = _fpi_team_level_tactical_messages_v144(
        ctx, gps_context, readiness, priorities, week, 8
    )
    findings = _fpi_unique_messages_v145(tactical_candidates, 3, used)
    used.extend(findings)

    plan_candidates = _fpi_match_plan_messages_v145(
        ctx, gps_context, readiness, priorities, week, 7
    )
    plan = _fpi_unique_messages_v145(plan_candidates, 4, used)
    used.extend(plan)

    team_candidates = _fpi_team_level_tactical_messages_v144(
        ctx, gps_context, readiness, priorities, week, 12
    ) + _fpi_match_plan_messages_v145(
        ctx, gps_context, readiness, priorities, week, 10
    )
    team_messages = _fpi_unique_messages_v145(team_candidates, 4, used)

    return {
        "findings": findings,
        "plan": plan,
        "team_messages": team_messages,
    }



# =========================================================
# V146 - Coach Intelligence Engine
# Egységes insight-, prioritás-, ismétlés- és narratívamotor
# =========================================================
@dataclass
class FPIInsightV146:
    topic: str
    title: str
    finding: str
    recommendation: str
    source: str = ""
    priority: int = 50
    block: str = ""
    day: str = ""


def _fpi_normalize_any_insight_v146(value: object) -> Optional[FPIInsightV146]:
    """A PDF-generátorba nyers dict/list soha ne kerülhessen."""
    if value is None:
        return None
    if isinstance(value, FPIInsightV146):
        return value
    if isinstance(value, dict):
        topic = str(value.get("Téma") or value.get("Tema") or value.get("topic") or "Taktikai megállapítás")
        finding = str(
            value.get("Edzői következtetés")
            or value.get("Edzöi következtetés")
            or value.get("Következtetés")
            or value.get("Értelmezés")
            or value.get("Megállapítás")
            or value.get("finding")
            or ""
        )
        recommendation = str(
            value.get("Javaslat")
            or value.get("Ajánlás")
            or value.get("Teendő")
            or value.get("recommendation")
            or ""
        )
        source = str(value.get("Bizonyíték") or value.get("Forrás") or value.get("source") or "")
        if not finding and source:
            finding = source
        if not finding and not recommendation:
            return None
        return FPIInsightV146(
            topic=_fpi_hu_plain_text_v144(topic),
            title=_fpi_hu_plain_text_v144(topic),
            finding=_fpi_extract_coach_text_v145(finding, 420),
            recommendation=_fpi_extract_coach_text_v145(recommendation, 420),
            source=_fpi_extract_coach_text_v145(source, 260),
            priority=int(value.get("priority", 50) or 50),
            block=str(value.get("block", "")),
            day=str(value.get("day", "")),
        )
    text = _fpi_extract_coach_text_v145(value, 420)
    if not text:
        return None
    return FPIInsightV146(
        topic=_fpi_fitness_topic_v144(text),
        title="Edzői megállapítás",
        finding=text,
        recommendation="",
    )


def _fpi_strip_raw_repr_v146(text: object) -> str:
    """Utolsó védelmi vonal: sem dict-, sem listarepr nem jelenhet meg."""
    value = str(text or "")
    value = re.sub(r"(?is)\{\s*['\"]Téma['\"].*?\}", "", value)
    value = re.sub(r"(?is)\{\s*['\"]Tema['\"].*?\}", "", value)
    value = re.sub(r"(?is)\{\s*['\"]Bizonyíték['\"].*?\}", "", value)
    value = re.sub(r"(?is)\[\s*\{\s*['\"]Téma['\"].*?\}\s*\]", "", value)
    value = value.replace("{", "").replace("}", "")
    return _fpi_hu_plain_text_v144(value).strip()


def _fpi_render_insight_text_v146(insight: FPIInsightV146, include_source: bool = False) -> str:
    parts = []
    if insight.title:
        parts.append(f"{insight.title}:")
    if insight.finding:
        parts.append(_fpi_strip_raw_repr_v146(insight.finding))
    if insight.recommendation:
        parts.append(_fpi_strip_raw_repr_v146(insight.recommendation))
    if include_source and insight.source:
        parts.append(f"Alapja: {_fpi_strip_raw_repr_v146(insight.source)}")
    return " ".join(x for x in parts if x).strip()


def _fpi_contextual_gps_only_insights_v146(
    gps_context: Dict[str, object],
    priorities: List[dict],
    readiness: int,
    week: Optional[str],
    limit: int = 6,
) -> List[FPIInsightV146]:
    """GPS-only esetben is több száz kombinációs út, témánként egyedi hozzáadott értékkel."""
    ctx = gps_context or {}
    seed = _fpi_context_seed_v144({}, week, readiness)
    priority_text = " ".join(
        str(p.get("Teendő", p.get("Cím", p.get("title", ""))))
        for p in (priorities or []) if isinstance(p, dict)
    ).lower()

    pool: List[FPIInsightV146] = []

    if readiness < 50:
        pool.extend([
            FPIInsightV146("readiness", "Készenlét", "A csapat frissessége jelenleg alacsony.", "A hét első felében csökkentsük a teljes volument, az intenzív ingereket pedig rövid, teljes pihenőkkel tagolt blokkokban adjuk.", priority=95),
            FPIInsightV146("regeneracio", "Regeneráció", "A terhelési állapot alapján a visszatöltés fontosabb a további volumenhalmozásnál.", "A következő erősebb nap előtt ellenőrizzük az egyéni reakciókat, és a magas kockázatú játékosoknál csökkentsük az ismétlésszámot.", priority=92),
        ])
    elif readiness < 65:
        pool.extend([
            FPIInsightV146("readiness", "Készenlét", "A csapat terhelhető, de a frissesség nem optimális.", "A fő terhelési nap maradhat, de az azt követő napon legyen egyértelmű visszalépés volumenben és gyorsításokban.", priority=84),
            FPIInsightV146("regeneracio", "Frissülés", "A hét szerkezete megfelelő lehet, ha az intenzív és regeneráló napok világosan elkülönülnek.", "Ne tegyünk két hasonló neuromuszkuláris terhelésű napot egymás mellé.", priority=78),
        ])
    else:
        pool.extend([
            FPIInsightV146("readiness", "Készenlét", "A csapat aktuális frissessége megfelelő alapot ad a tervezett heti fő ingerhez.", "A minőséget részesítsük előnyben: a fő terhelési napon legyen magas a végrehajtási intenzitás, de ne nőjön feleslegesen az ismétlésszám.", priority=76),
            FPIInsightV146("periodizacio", "Heti felépítés", "A jelenlegi állapot mellett jól elkülöníthető a fő terhelési és az élező nap.", "A legnagyobb terhelés ne kerüljön túl közel a mérkőzéshez, az utolsó 48 órában már a frissesség visszaépítése legyen a cél.", priority=74),
        ])

    if any(x in priority_text for x in ["sprint", "hsr", "sebess"]):
        pool.extend([
            FPIInsightV146("sebesseg", "Sebességi inger", "A nagy sebességű futások és sprintek mennyisége elmarad a kívánt heti zónától.", "Egy célzott sebességi blokk indokolt: kevés ismétlés, hosszú pihenő, teljes technikai kontroll.", priority=93),
            FPIInsightV146("sebesseg", "Sprintminőség", "A hiány nem pusztán volumenprobléma, hanem a maximális sebességhez közeli expozíció hiányát is jelezheti.", "A sprintblokkban a játékosok érjék el a saját csúcssebességük magas hányadát, ne csak több közepes tempójú futást végezzenek.", priority=90),
        ])
    else:
        pool.extend([
            FPIInsightV146("sebesseg", "Sebességi egyensúly", "A heti sebességi terhelés megfelelőnek tűnik.", "A következő héten a sprintmennyiség növelése helyett a jó minőségű, teljesen kipihent végrehajtást tartsuk meg.", priority=69),
        ])

    if any(x in priority_text for x in ["load", "terhel", "volumen", "táv"]):
        pool.extend([
            FPIInsightV146("terheles", "Terheléseloszlás", "A heti terhelés eloszlása egyenetlen vagy túlzottan egy napra koncentrálódik.", "Osszuk szét a volument úgy, hogy a fő nap után ne maradjon tartós fáradtság a meccs közelében.", priority=88),
            FPIInsightV146("volumen", "Heti volumen", "Az össztáv önmagában megfelelő lehet, de a terhelés minőségét az intenzitási összetétel dönti el.", "A következő héten külön ellenőrizzük a HSR-, sprint- és gyorsítási részt, ne csak az összmétert.", priority=82),
        ])
    else:
        pool.extend([
            FPIInsightV146("volumen", "Heti volumen", "A heti összterhelés stabil alapot ad.", "A következő lépés nem feltétlenül több munka, hanem pontosabb napok közötti elosztás és jobb egyéni differenciálás.", priority=67),
        ])

    pool.extend([
        FPIInsightV146("trend", "Négyhetes trend", "A heti adatokat a közelmúlt terhelési mintájához kell viszonyítani.", "Kerüljük a hirtelen ugrást; a következő héten legfeljebb egy fő terhelési tényezőt emeljünk számottevően.", priority=72),
        FPIInsightV146("gyorsulas", "Neuromuszkuláris terhelés", "A gyorsítások, lassítások és irányváltások olyan terhelést adnak, amelyet az össztáv nem mutat meg.", "Kis területű játékoknál külön ellenőrizzük a fékezések számát, és ne ugyanazon a napon halmozzuk a sprintet és a sok irányváltást.", priority=71),
        FPIInsightV146("egyeni", "Egyéni eltérések", "A csapatátlag elfedheti a szélsőséges játékosprofilokat.", "A magas kockázatú vagy alulexponált játékosok kapjanak egyéni módosítást, ne csak a teljes csapatra készüljön közös terhelési döntés.", priority=83),
        FPIInsightV146("meccsterheles", "Meccsterhelés", "Az edzés–meccs arány akkor hasznos, ha külön vizsgáljuk a teljes volument és a nagy intenzitású részt.", "A következő mikrociklusban azt a területet pótoljuk, amely az edzésen aránytalanul elmaradt a meccsigénytől.", priority=80),
    ])

    rotated = _fpi_rotate_v144(pool, seed)
    selected: List[FPIInsightV146] = []
    used_topics = set()
    for item in sorted(rotated, key=lambda x: x.priority, reverse=True):
        if item.topic in used_topics:
            continue
        selected.append(item)
        used_topics.add(item.topic)
        if len(selected) >= limit:
            break
    return selected


def _fpi_specific_match_plan_v146(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
    limit: int = 6,
) -> List[FPIInsightV146]:
    """Hol védekezzünk, hol támadjunk, kit keressünk, kire vigyázzunk."""
    ctx = tactical_context or {}
    opponent = _fpi_opponent_name_v145(ctx)
    features = _fpi_tactical_features_v144(ctx)
    seed = _fpi_context_seed_v144(ctx, week, readiness)

    opp_players = ctx.get("opponent_player_evaluation", []) or []
    own_players = ctx.get("own_player_evaluation", []) or []

    player_threats = []
    for row in opp_players:
        if not isinstance(row, dict):
            continue
        name = str(row.get("Játékos", "")).strip()
        role = _fpi_extract_coach_text_v145(row.get("Szerep", ""), 80)
        eval_text = _fpi_extract_coach_text_v145(row.get("Értelmezés", ""), 170)
        suggestion = _fpi_extract_coach_text_v145(row.get("Javaslat", ""), 180)
        if name:
            player_threats.append((name, role, eval_text, suggestion))

    own_targets = []
    for row in own_players:
        if not isinstance(row, dict):
            continue
        name = str(row.get("Játékos", "")).strip()
        role = _fpi_extract_coach_text_v145(row.get("Szerep", ""), 80)
        eval_text = _fpi_extract_coach_text_v145(row.get("Értelmezés", ""), 160)
        if name:
            own_targets.append((name, role, eval_text))

    pool: List[FPIInsightV146] = []

    if features["mid_block"]:
        pool.append(FPIInsightV146(
            "vedekezesi_zona", "Hol védekezzünk?",
            f"{opponent} ellen a középső blokk tűnik a legjobb kiindulási pontnak.",
            "A csapat a középső harmadban zárja a belső passzsávokat, majd az oldalvonal közelében támadja meg agresszíven a labdást.",
            priority=95,
        ))
    elif features["high_press"]:
        pool.append(FPIInsightV146(
            "vedekezesi_zona", "Hol védekezzünk?",
            f"{opponent} első építési vonalát már a saját térfelén érdemes nyomás alá helyezni.",
            "A magas letámadás csak visszapassz, rossz átvétel vagy oldalvonal mellé szorult labdás esetén induljon, hogy ne nyíljon meg a középső terület.",
            priority=95,
        ))
    else:
        pool.append(FPIInsightV146(
            "vedekezesi_zona", "Hol védekezzünk?",
            f"{opponent} ellen a stabil középső zóna védelme élvezzen elsőbbséget.",
            "Ne húzzuk szét a blokkot felesleges egyéni kilépésekkel; a labdát tereljük az előre kijelölt oldalra.",
            priority=88,
        ))

    if features["wide"]:
        pool.append(FPIInsightV146(
            "tamadasi_zona", "Hol támadjunk?",
            f"{opponent} oldali védekezését létszámfölénnyel és gyors visszaforgatással érdemes bontani.",
            "Az első oldali kombináció után keressük a félterületi visszagurítást vagy a túloldali szélsőt, ne az első beadást erőltessük.",
            priority=93,
        ))
    elif features["central"] or features["possession"]:
        pool.append(FPIInsightV146(
            "tamadasi_zona", "Hol támadjunk?",
            f"{opponent} középpályás és védősora közötti terület lehet a fő célzóna.",
            "A támadó középpályás vagy visszalépő csatár kapjon labdát a vonalak között, majd egyérintős kapcsolattal indítsuk a mélységi futást.",
            priority=93,
        ))
    else:
        pool.append(FPIInsightV146(
            "tamadasi_zona", "Hol támadjunk?",
            f"{opponent} rendezetlen oldalát azonnali oldalváltással vagy gyors átmenettel érdemes keresni.",
            "Labdaszerzés után az első előrepassz a védők mögötti szélső területre menjen, ha a labdás nincs nyomás alatt.",
            priority=88,
        ))

    if player_threats:
        name, role, eval_text, suggestion = player_threats[0]
        pool.append(FPIInsightV146(
            "veszelyes_jatekos", "Kire vigyázzunk?",
            f"{name} ({role}) az ellenfél egyik legfontosabb veszélyforrása. {eval_text}",
            suggestion or "Ne engedjük szabadon labdát átvenni a veszélyes zónában; a közeli biztosító játékos készüljön a második akcióra is.",
            priority=99,
        ))
    if len(player_threats) > 1:
        name, role, eval_text, suggestion = player_threats[1]
        pool.append(FPIInsightV146(
            "masodik_veszely", "Második veszélyforrás",
            f"{name} ({role}) más típusú veszélyt jelent. {eval_text}",
            suggestion or "Az ő oldalán vagy zónájában előre rögzítsük, ki lép ki és ki biztosít mögötte.",
            priority=92,
        ))

    if own_targets:
        name, role, eval_text = own_targets[0]
        pool.append(FPIInsightV146(
            "sajat_kulcsember", "Kit keressünk?",
            f"{name} ({role}) lehet a saját játék egyik fő kapcsolópontja. {eval_text}",
            "A labdakihozatal és a támadó átmenet első vagy második passzában tudatosan keressük, ne csak véletlenszerűen jusson labdához.",
            priority=94,
        ))
    else:
        pool.append(FPIInsightV146(
            "sajat_kulcsember", "Kit keressünk?",
            "A saját progresszív vagy kreatív játékost a vonalak között, illetve az első ellenfél-sor mögött kell labdához juttatni.",
            "Az első építési fázisban legyen előre rögzített passzút hozzá, hogy ne csak oldalirányú labdajáratás történjen.",
            priority=80,
        ))

    if features["transition"]:
        pool.append(FPIInsightV146(
            "elso_tamado_gondolat", "Első támadó gondolat",
            f"{opponent} labdavesztése után azonnal a védők mögötti területet vagy a gyenge oldalt keressük.",
            "Ha az első előrepassz nem tiszta, egy biztosító passzal tartsuk meg a labdát, majd gyors oldalváltással folytassuk.",
            priority=91,
        ))

    if features["set_piece"]:
        pool.append(FPIInsightV146(
            "pontrugas", "Pontrúgásfókusz",
            f"{opponent} pontrúgásainál a második labda és a hosszú oldal külön figyelmet igényel.",
            "Legyen kijelölt játékos a lepattanóra és egy biztosító a gyors ellentámadás megállítására.",
            priority=84,
        ))

    if readiness < 60:
        pool.append(FPIInsightV146(
            "allapothoz_igazitva", "Állapothoz igazítva",
            "A jelenlegi fizikai állapot nem indokol egész mérkőzésen folyamatos magas nyomást.",
            "Az intenzív letámadást rövid, előre kijelölt szakaszokhoz kössük, a köztes időszakokban stabilabb blokkban védekezzünk.",
            priority=90,
        ))

    rotated = _fpi_rotate_v144(pool, seed)
    selected, used_topics = [], set()
    for item in sorted(rotated, key=lambda x: x.priority, reverse=True):
        if item.topic in used_topics:
            continue
        selected.append(item)
        used_topics.add(item.topic)
        if len(selected) >= limit:
            break
    return selected


def _fpi_tactical_findings_v146(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
    limit: int = 6,
) -> List[FPIInsightV146]:
    ctx = tactical_context or {}
    features = _fpi_tactical_features_v144(ctx)
    opponent = _fpi_opponent_name_v145(ctx)
    seed = _fpi_context_seed_v144(ctx, week, readiness)
    pool: List[FPIInsightV146] = []

    raw_sources = []
    for key in ["risks", "tactical_findings", "pdf_provider_findings", "own_topics", "opp_topics"]:
        raw_sources.extend(ctx.get(key, []) or [])
    for raw in raw_sources:
        insight = _fpi_normalize_any_insight_v146(raw)
        if insight and insight.finding:
            insight.title = insight.title or "Taktikai megállapítás"
            insight.priority = max(insight.priority, 70)
            pool.append(insight)

    if features["transition"]:
        pool.append(FPIInsightV146(
            "atmenet", "Átmeneti veszély",
            f"{opponent} gyorsan próbál előre játszani labdaszerzés után.",
            "A labda mögötti biztosítás és az első előrepassz zárása fontosabb, mint az azonnali, szervezetlen visszatámadás.",
            priority=94,
        ))
    if features["wide"]:
        pool.append(FPIInsightV146(
            "szelso", "Szélső veszély",
            f"{opponent} oldali akciói és beadásai kiemelt veszélyt jelenthetnek.",
            "A beadási helyzet előtt kell nyomást gyakorolni; a hosszú oldal és a tizenhatos előtti lepattanó külön felelőst kapjon.",
            priority=92,
        ))
    if features["central"]:
        pool.append(FPIInsightV146(
            "kozepso", "Vonalak közötti játék",
            f"{opponent} a középpályás és védősor közötti területet is keresi.",
            "A védekező középpályás adjon nyomást, a belső védő biztosítson mögötte, hogy egyetlen kilépéssel ne nyíljon meg a kapu felé vezető út.",
            priority=90,
        ))
    if features["build_up"]:
        pool.append(FPIInsightV146(
            "labdakihozatal", "Labdakihozatal",
            f"{opponent} első nyomásának kijátszásához szükség lehet a kapus és a harmadik ember bevonására.",
            "Ne minden játékos kérje lábra a labdát; legyen előre rögzített mélységi mozgás, amely hátrahúzza az ellenfél védelmét.",
            priority=86,
        ))
    if features["second_ball"]:
        pool.append(FPIInsightV146(
            "masodik_labda", "Második labdák",
            f"{opponent} direkt játékánál a párharc utáni terület fontosabb lehet, mint maga az első fejpárbaj.",
            "A középpályás sor gyorsan zárjon a labda várható helyére, hogy a lepattanó ne maradjon az ellenfélnél.",
            priority=88,
        ))
    if features["set_piece"]:
        pool.append(FPIInsightV146(
            "pontrugas", "Pontrúgások",
            f"{opponent} pontrúgásai után a második szituáció is veszélyes lehet.",
            "A felszabadítás után a csapat együtt lépjen ki, és a lepattanóra kijelölt játékos ne hagyja szabadon az újabb beadást.",
            priority=84,
        ))

    rotated = _fpi_rotate_v144(pool, seed)
    selected, used_topics, accepted_texts = [], set(), []
    for item in sorted(rotated, key=lambda x: x.priority, reverse=True):
        if item.topic in used_topics:
            continue
        rendered = _fpi_render_insight_text_v146(item)
        if _fpi_is_near_duplicate_v145(rendered, accepted_texts, 0.70):
            continue
        selected.append(item)
        used_topics.add(item.topic)
        accepted_texts.append(rendered)
        if len(selected) >= limit:
            break
    return selected


def _fpi_contextual_md_plan_rows_v146(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
) -> List[Tuple[str, str, str]]:
    rows = _fpi_contextual_md_plan_rows_v145(
        tactical_context, gps_context, readiness, priorities, week
    )
    result = []
    for day, fitness_goal, tactical_goal in rows:
        tactical_text = _fpi_strip_raw_repr_v146(tactical_goal)
        low = tactical_text.lower()
        if "labdakihozatal" in low or "épít" in low:
            title = "Játékfelépítés"
        elif "átmenet" in low or "labdaszerzés" in low or "labdavesztés" in low:
            title = "Átmenetek"
        elif "blokk" in low or "védekez" in low:
            title = "Védekezés"
        elif "szélső" in low or "beadás" in low or "oldali" in low:
            title = "Szélső játék"
        elif "pontrúgás" in low or "szöglet" in low:
            title = "Pontrúgások"
        elif "letámadás" in low:
            title = "Letámadás"
        elif "aktiváció" in low or "aktiv" in low:
            title = "Aktiváció"
        else:
            title = "Meccsterv"
        result.append((day, fitness_goal, f"{title}: {tactical_text}"))
    return result


def _fpi_coach_blocks_v146(
    tactical_context: Optional[Dict[str, object]],
    gps_context: Optional[Dict[str, object]],
    readiness: int,
    priorities: Optional[List[dict]],
    week: Optional[str],
) -> Dict[str, List[FPIInsightV146]]:
    tactical = _fpi_tactical_findings_v146(
        tactical_context, gps_context, readiness, priorities, week, 7
    )
    plan = _fpi_specific_match_plan_v146(
        tactical_context, gps_context, readiness, priorities, week, 7
    )

    # Lágy deduplikáció: azonos gondolat más mélységben megmaradhat,
    # az 1:1 vagy nagyon közeli ismétlés viszont ne.
    used_texts = [_fpi_render_insight_text_v146(x) for x in tactical]
    clean_plan = []
    for item in plan:
        rendered = _fpi_render_insight_text_v146(item)
        if _fpi_is_near_duplicate_v145(rendered, used_texts, 0.82):
            continue
        clean_plan.append(item)
        used_texts.append(rendered)

    # Csapatszintű edzői üzenetek a taktikai és meccstervi elemek mélyebb,
    # de nem szó szerinti alkalmazási szintjéből.
    team_messages = []
    for item in tactical + clean_plan:
        title = item.title
        finding = item.finding
        recommendation = item.recommendation
        if recommendation:
            team_messages.append(FPIInsightV146(
                topic=f"team_{item.topic}",
                title=title,
                finding=finding,
                recommendation=f"Pályán használható feladat: {recommendation}",
                priority=item.priority,
            ))
        if len(team_messages) >= 5:
            break

    return {
        "tactical": tactical,
        "plan": clean_plan[:6],
        "team": team_messages[:4],
    }



# =========================================================
# V147 - Executive oldalszerkezet + mélyebb erőnléti mikrociklus
# =========================================================
def _fpi_weekly_fitness_rows_v147(
    md_rows: List[Tuple[str, str, str]],
    fitness_insights: List[FPIInsightV146],
    readiness: int,
) -> List[Tuple[str, str, str, str]]:
    """Részletesebb, naphoz kötött erőnléti fókusz és edzői megjegyzés."""
    insight_texts = {
        item.topic: _fpi_render_insight_text_v146(item)
        for item in (fitness_insights or [])
    }
    result = []
    for day, fitness_goal, tactical_goal in md_rows:
        day_text = str(day)
        base_goal = _fpi_strip_raw_repr_v146(fitness_goal)
        tactical_text = _fpi_strip_raw_repr_v146(tactical_goal)

        if "MD-5" in day_text or "MD-6" in day_text:
            extra = (
                "Volumenépítés: közepes intenzitás, nagyobb össztáv, kontrollált gyorsítás- és lassításszám."
            )
            note = (
                "Ez legyen a hét fő volumenadó napja. A sprintterhelést még ne maximalizáljuk; "
                "a cél a terhelési alap felépítése túlzott neuromuszkuláris fáradás nélkül."
            )
        elif "MD-4" in day_text:
            extra = (
                "Fő terhelési nap: HSR, sprint és nagy intenzitású akciók meccsszerű környezetben."
            )
            note = (
                "A sebességi ingert teljes pihenőkkel, kevés jó minőségű ismétléssel adjuk. "
                "A magasabb kockázatú játékosoknál az ismétlésszám csökkenthető, de a szükséges sebesség maradjon meg."
            )
        elif "MD-3" in day_text:
            extra = (
                "Intenzitásfenntartás: rövidebb blokkok, ismételt nagy intenzitású akciók, mérsékelt összvolumen."
            )
            note = (
                "A terhelés legyen meccsszerű, de ne érje el a hét fő terhelési napjának volumenét. "
                "A gyorsítások és fékezések számát külön figyeljük."
            )
        elif "MD-2" in day_text:
            extra = (
                "Terheléscsökkentés: alacsonyabb volumen, rövid gyorsasági érintés, teljes frissülés támogatása."
            )
            note = (
                "Ne halmozzunk új fáradtságot. A sebességi inger csak idegrendszeri aktiváció legyen, "
                "nem kondicionáló blokk."
            )
        elif "MD-1" in day_text:
            extra = (
                "Aktiváció: nagyon alacsony volumen, rövid reakció- és gyorsasági feladatok, hosszú pihenőkkel."
            )
            note = (
                "A játékosok frissessége fontosabb, mint bármely terhelési cél. "
                "A blokkot a minőség első romlásakor zárjuk le."
            )
        elif "Pihen" in base_goal or "regener" in base_goal.lower():
            extra = "Regeneráció: mobilitás, egyéni visszatöltés, szükség esetén kompenzáció."
            note = (
                "A nem játszó vagy alulterhelt játékosok külön kompenzációs tervet kaphatnak; "
                "a magas terhelést kapók valódi regenerációt kapjanak."
            )
        else:
            extra = "Terhelés a heti ciklus aktuális helyéhez és a játékosok egyéni állapotához igazítva."
            note = (
                "Az egyéni eltéréseket a csapatátlag ne fedje el; a kockázati jelzések alapján módosítsuk a volument."
            )

        if readiness < 60:
            note += " A csökkent készenlét miatt a tervezett ismétlésszámból szükség esetén 10–20% elhagyható."
        elif readiness >= 80 and ("MD-4" in day_text or "MD-3" in day_text):
            note += " A jó frissesség lehetővé teszi a magas végrehajtási minőséget, de a felesleges volumen továbbra sem cél."

        combined_goal = f"{base_goal} {extra}".strip()
        result.append((day_text, combined_goal, tactical_text, note))
    return result


def _fpi_fitness_snapshot_rows_v147(
    fitness_insights: List[FPIInsightV146],
    limit: int = 7,
) -> List[List[str]]:
    rows = [["Terület", "Aktuális értelmezés", "Edzői teendő"]]
    for insight in (fitness_insights or [])[:limit]:
        rows.append([
            insight.title or insight.topic,
            _fpi_strip_raw_repr_v146(insight.finding),
            _fpi_strip_raw_repr_v146(insight.recommendation),
        ])
    return rows


# =========================================================
# V143 - Methodology content + PDF export
# =========================================================
FPI_METHODOLOGY_SECTIONS_V143 = [
    (
        "1. Mire épül az FPI metodikája?",
        [
            "A Football Performance Intelligence nemzetközi sporttudományi szakirodalomra, saját és partneri GPS-adatokra, valamint gyakorlati teljesítményelemzési logikára épül.",
            "A rendszer döntéstámogató eszköz: nem diagnózis, nem orvosi döntés és nem garantált sérülés-előrejelzés.",
            "Az eredmények edzői megfigyeléssel, wellness/RPE, orvosi információval és a heti szakmai kontextussal együtt értelmezendők.",
        ],
    ),
    (
        "2. Mit jelent a Readiness?",
        [
            "A Readiness a csapat vagy játékos aktuális terhelési állapotának 0–100 közötti becslése.",
            "Figyelembe vett fő területek: rövid távú terhelés, 3–7 napos állapot, 4 hetes trend, össztáv, Load, HSR, sprinttáv, sprintek, High Efforts és – ha elérhető – pulzus/HRV.",
            "80–100: magas készenlét; 60–79: elfogadható, figyelendő; 40–59: csökkent; 0–39: alacsony.",
            "Az alacsonyabb érték túlterhelést, alulterhelést vagy kedvezőtlen terhelési mintázatot is jelezhet.",
        ],
    ),
    (
        "3. Mit jelent a játékoskockázat?",
        [
            "A Player Risk Score nem sérülésjóslat, hanem korai figyelmeztető besorolás.",
            "A rendszer a hirtelen terhelésváltozást, a négyhetes trendet, a túl- és alulterhelést, a sprint- és nagy sebességű futások mennyiségét, a gyorsításokból és lassításokból adódó terhelést, valamint a nagy intenzitású akciókat vizsgálja.",
            "Alacsony: stabil profil. Közepes: egy vagy több figyelmeztető jel. Magas: több kedvezőtlen tényező egyidejű jelenléte.",
        ],
    ),
    (
        "4. Hogyan működnek a benchmarkok?",
        [
            "A referencia nem egyetlen fix érték: korosztály, bajnoki szint, játékosposzt és játékmodell alapján változik.",
            "A fő referencia-területek: össztáv, terhelési pont, nagy sebességű futás, sprinttáv, sprintek száma és nagy intenzitású akciók.",
            "Csapatszinten a rendszer a játékosállomány posztösszetételéből súlyozott referencia-profilt képez. A kapusok a sebességi benchmarkban kisebb, de nem nulla súlyt kapnak.",
            "A referenciák nemzetközi szakirodalmi tartományokból és saját/partneri adatokból finomított döntéstámogató zónák.",
        ],
    ),
    (
        "5. Mit jelent a 4 hetes trend?",
        [
            "A legutóbbi hét értékeit nem önmagukban, hanem a játékos vagy csapat előző heteinek mintázatához viszonyítjuk.",
            "A trend segít felismerni a hirtelen terhelésnövekedést, a tartós alulexpozíciót, az ingadozó terhelést és a sebességi inger hiányát.",
            "A rendszer a trendet jelzésként használja, nem mechanikus döntési szabályként.",
        ],
    ),
    (
        "6. Hogyan készül a mikrociklus-javaslat?",
        [
            "A motor a heti edzésszámot, a meccsnapot, a készenléti és kockázati állapotot, a nagy sebességű futások és sprintek trendjét, a nagy intenzitású akciókat, a játékmodellt és a rendelkezésre álló taktikai anyagokat kapcsolja össze.",
            "A terv 3–6 edzéses ciklushoz is igazítható, nem fix négynapos sablon.",
            "GPS-only módban a fizikai expozíció és regeneráció vezeti a tervet. Taktikai input esetén az ellenfél-specifikus fókuszok is beépülnek.",
        ],
    ),
    (
        "7. Tactical Framework: 7 dimenzió és 9 stratégia",
        [
            "A 7 taktikai dimenzió: letámadás, labdakihozatal, átmenetek, támadó játék, pontrúgások, labdabirtoklás és lövésprofil.",
            "A két fő értelmezési tengely: játékstílus (direkt–vegyes–kontroll–agresszív) és blokkmagasság (mély–közép–magas).",
            "A 9 stratégiai profil: KON – kontra mély blokkból; GAT – gyors átmenet; BAT – középső blokk + átmenet; KIE – kiegyensúlyozott; PRS – presszing + átmenet; MLT – magas letámadás; DOM – dominancia; POZ – pozíciós támadás; LAB – labdatartás mélyebb szerkezetből.",
            "A stratégia nem önmagában kész meccsterv: a rendszer csapatszintű taktikai megállapításokkal, ellenfél-játékos fókuszokkal és konkrét meccstervi teendőkkel differenciál.",
        ],
    ),
    (
        "8. Hogyan készülnek az edzői üzenetek?",
        [
            "Az erőnléti üzenetek témák szerint készülnek: készenlét, sebességi terhelés, heti volumen, terhelési pont, négyhetes trend, regeneráció, gyorsítások és lassítások, valamint játékoskockázat.",
            "Egy riporton belül ugyanaz a motívum csak egyszer jelenhet meg. Így például az alacsony sprintterhelés nem ismétlődhet három külön sorban eltérő megfogalmazással.",
            "A taktikai üzenetek a saját játékmodellből, az ellenfél erősségeiből és gyengeségeiből, a javasolt stratégiai profilból, az ellenfél játékosértékeléséből és a csapat aktuális fizikai állapotából állnak össze.",
            "A rendszer az adott hét, ellenfél, saját játékmodell, stratégiai profil, játékosértékelés és terhelési állapot alapján állítja össze az üzeneteket. A taktikai célok naponta is változnak: más feladat készül a fő terhelési napra, az átmeneti napra, a meccstervi napra és az aktivációra.",
            "Az üzenetek tudásbázisból és szakmai szabályokból épülnek fel. A rendszer témánként választ, majd hasonlóságvizsgálattal kiszűri, hogy ugyanaz a motívum több blokkban vagy három egymást követő sorban megismétlődjön.",
            "Az Executive Summary logikusan tagolt: az első oldal a vezetői döntést és a heti ciklustervet, a második oldal a fő edzői üzeneteket, a harmadik oldal a játékosszintű fókuszokat tartalmazza. A végén külön erőnléti helyzetkép emeli be a GPS-only motor részletesebb következtetéseit.",
            "A heti ciklusterv minden naphoz erőnléti fókuszt, taktikai fókuszt és rövid edzői megjegyzést rendel. Így a terv nemcsak azt mondja meg, mi legyen a cél, hanem azt is, hogyan és milyen terhelési logikával valósítsuk meg.",
            "A felső vezetői blokk, a meccsterv, a csapatszintű taktikai üzenet, a napi taktikai cél és az ellenfél játékosfókusz eltérő információs mélységet kap. Ugyanaz a téma megjelenhet több helyen, ha más döntési szintet szolgál, de az egy az egyben ismétlődő mondatokat a rendszer kiszűri.",
            "A meccsterv konkrét kérdésekre válaszol: hol védekezzünk, hol támadjunk, kit keressünk, kire vigyázzunk, mi legyen az első támadó gondolat, és mely játékhelyzeteket kell edzésen gyakorolni.",
            "GPS-only módban ugyanaz a tudás- és kombinációs motor dolgozik, csak taktikai input nélkül. A rendszer a készenlét, kockázat, össztáv, terhelési pont, nagy sebességű futások, sprintek, nagy intenzitású akciók, gyorsítások, lassítások, négyhetes trend, edzés–meccs arány és heti periodizáció alapján állít össze többféle erőnléti üzenetet.",
            "Az erőnléti és taktikai üzenetek témánként, prioritás szerint és kontextusfüggően készülnek. Fél év riportjait visszanézve nem ugyanazok a mondatok ismétlődnek, hanem az adott hét és mérkőzés saját problémái, döntési pontjai és pályán használható feladatai kerülnek előtérbe.",
            "A szövegmotor nem szabadon kitalált állításokat készít: kizárólag az elérhető GPS- és taktikai inputokból, valamint előre meghatározott szakmai szabályokból dolgozik.",
        ],
    ),
    (
        "9. Hogyan kell szakmailag használni?",
        [
            "Vezetőedző: gyors heti fő üzenetek, taktikai és terhelési fókusz.",
            "Erőnléti edző: readiness, risk, Load, HSR, sprint és High Efforts trendek.",
            "Sportigazgató: 30 másodperces állapotkép, kockázatok és prioritások.",
            "Utánpótlás: korosztályos és poszt-specifikus összevetés, alul- és túlterhelés korai jelzése.",
        ],
    ),
]


def build_fpi_methodology_pdf_bytes_v143() -> Optional[bytes]:
    if SimpleDocTemplate is None:
        return None
    output = io.BytesIO()
    try:
        regular_font, bold_font = _register_pdf_fonts() if "_register_pdf_fonts" in globals() else ("Helvetica", "Helvetica-Bold")
    except Exception:
        regular_font, bold_font = "Helvetica", "Helvetica-Bold"

    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "FPI_METHOD_TITLE_V143",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=10,
    )
    subtitle = ParagraphStyle(
        "FPI_METHOD_SUB_V143",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#475569"),
        spaceAfter=10,
    )
    heading = ParagraphStyle(
        "FPI_METHOD_HEAD_V143",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#0F766E"),
        spaceBefore=8,
        spaceAfter=5,
    )
    body = ParagraphStyle(
        "FPI_METHOD_BODY_V143",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=9.2,
        leading=12.5,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=4,
    )

    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )
    story = [
        Paragraph("Football Performance Intelligence – Metodika", title),
        Paragraph(
            "A Metodika Center kibontott, exportálható változata. Rövid, szakmai és transzparens áttekintés arról, hogy az FPI milyen adatokat és döntési logikát használ.",
            subtitle,
        ),
        Spacer(1, 0.15 * cm),
    ]
    for section_title, paragraphs in FPI_METHODOLOGY_SECTIONS_V143:
        story.append(Paragraph(pdf_safe_text(section_title), heading))
        for paragraph in paragraphs:
            story.append(Paragraph("• " + pdf_safe_text(paragraph), body))
        story.append(Spacer(1, 0.10 * cm))
    story.append(Spacer(1, 0.18 * cm))
    story.append(Paragraph(
        "Szakmai megjegyzés: az FPI döntéstámogató rendszer. A kimeneteket edzői, orvosi és teljesítménydiagnosztikai információkkal együtt kell értelmezni.",
        subtitle,
    ))
    doc.build(story)
    output.seek(0)
    return output.read()


def render_fpi_methodology_center_v143() -> None:
    _fpi_landing_css_v100()
    c1, c2, c3 = st.columns([1, 1, 4])
    with c1:
        if st.button("← Főoldal", use_container_width=True, key="method_v143_back"):
            _fpi_set_page_v100("landing")
    with c2:
        if st.button("⚡ Riport", use_container_width=True, key="method_v143_clean"):
            _fpi_set_page_v100("clean")
    with c3:
        st.caption("Rövid, átlátható szakmai magyarázat – a képletek részletezése nélkül.")

    st.markdown(
        """
        <div style="border-radius:26px;padding:24px 28px;background:linear-gradient(135deg,#ffffff,#e0f2fe,#ecfdf5);border:1px solid #bfdbfe;box-shadow:0 16px 44px rgba(15,23,42,.12);margin:8px 0 18px 0;">
            <div style="font-size:.8rem;font-weight:950;color:#0f766e;letter-spacing:.07em;">FPI METHODOLOGY CENTER</div>
            <div style="font-size:2.2rem;font-weight:980;color:#0f172a;letter-spacing:-.04em;margin-top:5px;">Mit számol a rendszer, és hogyan értelmezzük?</div>
            <div style="color:#475569;margin-top:7px;">A legfontosabb szakmai kérdések rövid, kibontott magyarázata.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    method_pdf = build_fpi_methodology_pdf_bytes_v143()
    if method_pdf:
        st.download_button(
            "⬇️ Teljes metodika PDF",
            data=method_pdf,
            file_name="fpi_metodika.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="method_pdf_download_v143",
        )
    for title_text, paragraphs in FPI_METHODOLOGY_SECTIONS_V143:
        with st.expander(title_text, expanded=False):
            for paragraph in paragraphs:
                st.markdown(f"- {paragraph}")


# Override the older methodology renderer with the shared V143 content.
render_fpi_methodology_center_v138 = render_fpi_methodology_center_v143


# Default: első oldal / landing page. A teljes import-export app csak gomb után indul.
if "fpi_active_page_v100" not in st.session_state:
    st.session_state["fpi_active_page_v100"] = "landing"

active_page_v101 = st.session_state.get("fpi_active_page_v100", "landing")
if active_page_v101 == "landing":
    render_fpi_landing_page_v100()
    st.stop()
if active_page_v101 == "clean":
    render_fpi_clean_workspace_v101()
    st.stop()
if active_page_v101 == "method":
    render_fpi_methodology_center_v138()
    st.stop()

if active_page_v101 == "app" and not st.session_state.get("fpi_app_hub_seen_v137", False):
    render_fpi_app_hub_v137()
    st.stop()

render_fpi_hero()

top_back_col, top_clean_col, top_method_col, top_title_col = st.columns([1, 2.2, 1.4, 3.0])
with top_back_col:
    if st.button("← Főoldal", use_container_width=True, key="back_to_landing_v100"):
        _fpi_set_page_v100("landing")
with top_clean_col:
    if st.button("⚡ Input + vezetői export oldal", use_container_width=True, key="full_to_clean_v101", type="primary"):
        _fpi_set_page_v100("clean")
with top_method_col:
    if st.button("📚 Metodika", use_container_width=True, key="full_to_method_v138"):
        _fpi_set_page_v100("method")
with top_title_col:
    st.caption("Haladó elemző felület / részletes dashboardok / diagnosztika")

# V12.5: a haladó appban nincs külön bal oldali belépési/import panel.
# A fő munkafolyamat továbbra is az Import / Executive Export oldal.
st.markdown(
    """
    <style>
    [data-testid="stSidebar"]{display:none !important;}
    .block-container{padding-left:2rem !important;padding-right:2rem !important;}
    </style>
    <div class="fpi-top-nav-card">
        <b>Haladó elemző felület</b><br>
        <span>A fő munkafolyamat az <b>Input + vezetői export oldal</b>. Ott tölthető fel a GPS, a saját/ellenfél taktikai PDF/Excel, és ott készül az Executive Summary. Ez a rész a haladó elemzésekhez van.</span>
    </div>
    """,
    unsafe_allow_html=True,
)
use_demo_data = False
uploaded = None
if "clean_mapped_df_override_v105" not in st.session_state:
    use_demo_data = st.toggle("Minta riport mintaadatokkal", value=True, key="full_demo_data_v115")

# V115: a teljes app lehetőleg a Import / Executive Export oldalon már feldolgozott GPS adatot használja.
# Így a Dashboard / Tactical Pro+ / többi fül nem duplikál importot és mappert.
if "clean_mapped_df_override_v105" in st.session_state and isinstance(st.session_state["clean_mapped_df_override_v105"], pd.DataFrame) and not st.session_state["clean_mapped_df_override_v105"].empty:
    raw_df = st.session_state["clean_mapped_df_override_v105"].copy()
    selected_sheet = "Import oldal"
    uploaded = None
    use_demo_data = False
    st.session_state["mapped_df_override"] = raw_df.copy()
    mapping = st.session_state.get("clean_manual_mapping_v105", st.session_state.get("manual_mapping", {}))
    missing_core = []
elif uploaded is None and not use_demo_data:
    st.info("Tölts fel GPS/terhelési Excel fájlt az Import / export oldalon, vagy kapcsold be a minta riportot.")
    st.stop()
elif use_demo_data and uploaded is None:
    raw_df = build_demo_performance_data()
    selected_sheet = "Mintaadatok"
else:
    sheets = read_excel_all(uploaded)
    sheets = prepare_uploaded_sheets(sheets)
    sheet_names = list(sheets.keys())
    with st.expander("📄 Munkalap kiválasztása", expanded=False):
        selected_sheet = st.selectbox("Melyik munkalapot használjuk?", sheet_names, index=0)
    if st.session_state.get("mapper_selected_sheet") != selected_sheet:
        st.session_state.pop("mapped_df_override", None)
        st.session_state.pop("manual_mapping", None)
        st.session_state["mapper_selected_sheet"] = selected_sheet
    raw_df = sheets[selected_sheet]

# Ha a felhasználó a mapperrel már alkalmazott kézi mappinget, azt használjuk.
if "mapped_df_override" in st.session_state and isinstance(st.session_state["mapped_df_override"], pd.DataFrame) and not st.session_state["mapped_df_override"].empty:
    df = st.session_state["mapped_df_override"].copy()
    mapping = st.session_state.get("manual_mapping", {})
    missing_core = []
else:
    df, mapping, missing_core = standardize_dataframe(raw_df)

st.session_state['last_raw_df'] = raw_df

if missing_core:
    st.error(f"Hiányzó alapmezők: {', '.join(missing_core)}")
    st.write("Oszlopmapping:", mapping)
    st.info("Nyisd le lent a Smart Excel Mappert. Most már nem állítjuk meg az appot, hogy kézzel javítható legyen a mapping.")
    # Minimális üres standard df, hogy a mapper és a nyers adatnézet elérhető maradjon.
    df = pd.DataFrame(columns=["player_name", "session_type", "start_time", "session_date", "week"])

df = add_position_group(df)
df = render_keeper_controls_and_apply(df)
df, demo_limit_info = apply_demo_limits(df)

if df.empty or "week" not in df.columns or df["week"].dropna().empty:
    st.warning("A fájl még nem értelmezhető elemzésre. Javítsd a kötelező mezőket a Smart Excel Mapperben.")
    render_emergency_mapper(raw_df, mapping, missing_core if "missing_core" in globals() else [])
    st.stop()

render_demo_limit_notice(demo_limit_info if 'demo_limit_info' in globals() else {})

weeks = sorted(df["week"].dropna().unique().tolist()) if "week" in df.columns else []
players = sorted(df["player_name"].dropna().unique().tolist()) if "player_name" in df.columns else []
session_types = sorted(df["session_type"].dropna().unique().tolist()) if "session_type" in df.columns else []

if not weeks or not players:
    render_emergency_mapper(raw_df, mapping if "mapping" in globals() else {}, missing_core if "missing_core" in globals() else [])
    st.stop()

with st.expander("🧭 Hétfelismerés / diagnosztika", expanded=False):
    if st.session_state.get("week_rescue_applied"):
        with st.expander("Hétfelismerés diagnosztika V6.2", expanded=False):
            st.json(st.session_state.get("week_rescue_applied"))

with st.expander("⚙️ Meccskontextus és edzői beállítások", expanded=False):
    st.header("Meccskontextus")
    today_v94 = pd.Timestamp.today().date()
    opponent_v94 = st.text_input("Ellenfél neve", value=st.session_state.get("fpi_match_opponent_v94", ""))
    match_date_v94 = st.date_input("Meccsnap", value=st.session_state.get("fpi_match_date_v94", today_v94))
    match_week_v94 = _fpi_iso_week_from_date_v94(match_date_v94)
    st.session_state["fpi_match_opponent_v94"] = opponent_v94
    st.session_state["fpi_match_date_v94"] = match_date_v94
    st.session_state["fpi_match_context_v94"] = {
        "opponent": opponent_v94.strip() if isinstance(opponent_v94, str) else "",
        "match_date": match_date_v94,
        "match_week": match_week_v94,
        "today": today_v94,
        "today_week": _fpi_iso_week_from_date_v94(today_v94),
    }
    st.caption(f"Mai hét: {_fpi_iso_week_from_date_v94(today_v94)} | Meccshét: {match_week_v94}")

    st.header("Edzői kontextus / referencia V2")
    reference_age_v97 = st.selectbox("Korosztály", FPI_REFERENCE_AGE_OPTIONS_V112, index=0, key="app_ref_age_v112")
    reference_level_v97 = st.selectbox("Szint", FPI_REFERENCE_LEVEL_OPTIONS_V112, index=1, key="app_ref_level_v112")
    coach_week_type_v97 = st.selectbox("Mi a hét célja?", FPI_COACH_WEEK_OPTIONS_V112, index=1, key="app_week_type_v112")
    playmodel_profile_v97 = st.selectbox("Játékmodell profil", FPI_PLAYMODEL_OPTIONS_V112, index=4, key="app_playmodel_profile_v112")
    ref_profile_v97 = f"{reference_age_v97} / {reference_level_v97} / játékosonkénti poszt / {playmodel_profile_v97}"
    st.caption(f"Aktív referencia: {ref_profile_v97}. Nincs globális referencia poszt; a poszt játékosonként kerül értelmezésre.")

    cycle_days_v97 = st.number_input("Hány napos a ciklus?", min_value=3, max_value=10, value=7, step=1, key="app_cycle_days_v112")
    n_train_v97 = st.number_input("Hány edzés lesz?", min_value=0, max_value=6, value=4, step=1, key="app_n_train_v112")
    n_rest_v97 = st.number_input("Hány pihenőnap?", min_value=0, max_value=5, value=1, step=1, key="app_n_rest_v112")
    md_day_options_v97 = [f"MD-{i}" for i in range(int(cycle_days_v97)-1, 0, -1)] + ["MD"]
    md_match_day_v97 = st.selectbox("Melyik nap az MD?", md_day_options_v97, index=len(md_day_options_v97)-1, key="app_md_match_day_v112")
    session_plan_v97 = []
    st.caption("Add meg, melyik napon van edzés/pihenő/meccs. Így a mikrociklus nem fix 4 edzésre épül.")
    total_slots_v97 = max(1, min(int(cycle_days_v97), int(n_train_v97) + int(n_rest_v97) + (1 if md_match_day_v97 == "MD" else 0)))
    for i in range(total_slots_v97):
        c1, c2 = st.columns([1, 2])
        with c1:
            md_v = st.selectbox(f"Nap {i+1}", md_day_options_v97, index=min(i, len(md_day_options_v97)-1), key=f"md_day_v112_{i}")
        with c2:
            type_options = ["Edzés", "Pihenő", "Regeneráció", "Aktiváció", "Meccs"]
            default_type = "Meccs" if md_v == "MD" else ("Pihenő" if i >= int(n_train_v97) else "Edzés")
            typ_v = st.selectbox(f"Típus {i+1}", type_options, index=type_options.index(default_type), key=f"md_type_v112_{i}")
        note_v = st.text_input(f"Edzői megjegyzés {i+1}", value="", key=f"md_note_v112_{i}")
        session_plan_v97.append({"md": md_v, "type": typ_v, "note": note_v})
    st.session_state["fpi_coach_context_v97"] = {
        "reference_profile": ref_profile_v97,
        "reference_age": reference_age_v97,
        "reference_level": reference_level_v97,
        "playmodel_profile": playmodel_profile_v97,
        "coach_week_type": coach_week_type_v97,
        "cycle_days": int(cycle_days_v97),
        "training_days": int(n_train_v97),
        "rest_days": int(n_rest_v97),
        "md_day": md_match_day_v97,
        "session_plan": session_plan_v97,
    }

    st.header("Szűrők")
    default_week_idx_v94 = weeks.index(match_week_v94) if match_week_v94 in weeks else (len(weeks) - 1 if weeks else 0)
    selected_week = st.selectbox("Hét", weeks, index=default_week_idx_v94, format_func=week_label_short)
    selected_playstyle = st.selectbox("Játékmodell", FPI_PLAYMODEL_OPTIONS_V112, index=4)
    st.caption(PLAYSTYLE_OPTIONS.get(selected_playstyle, "Játékmodell-alapú referencia és következtetés."))
    selected_types = st.multiselect("Típus", session_types, default=session_types)
    selected_players = st.multiselect("Játékosok", players, default=players)

    st.header("Performance memória")
    if is_pro_mode():
        use_memory = st.checkbox("Korábbi mentett adatok bevonása", value=False)
        save_current_to_memory = st.button("Aktuális feltöltés mentése memoryba", use_container_width=True)
    else:
        use_memory = False
        save_current_to_memory = False
        st.caption("🔒 Performance memória Pro funkció.")

filtered = df[
    (df["week"] == selected_week)
    & (df["session_type"].isin(selected_types))
    & (df["player_name"].isin(selected_players))
]


# -----------------------------------------------------------------------------
# V120 - Empty state + compact table renderer
# -----------------------------------------------------------------------------
def fpi_empty_state(title: str, body: str = "", icon: str = "ℹ️") -> None:
    """Szép, kis helyigényű üres állapot a nagy üres fehér panelek helyett."""
    st.markdown(
        f"""
        <div class="fpi-empty-state-v120">
            <div class="fpi-empty-icon-v120">{html.escape(str(icon))}</div>
            <div>
                <div class="fpi-empty-title-v120">{html.escape(str(title))}</div>
                <div class="fpi-empty-body-v120">{html.escape(str(body))}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def fpi_compact_table(df_in: pd.DataFrame, max_rows: int = 12) -> None:
    """Streamlit dataframe helyett HTML table: nem hagy nagy üres komponenst és mindig olvasható."""
    if df_in is None or not isinstance(df_in, pd.DataFrame) or df_in.empty:
        fpi_empty_state("Nincs megjeleníthető adat", "A táblázat csak akkor jelenik meg, ha van értelmezhető sor.", "📊")
        return
    show_df = df_in.head(max_rows).copy()
    # túl hosszú cellák rövidítése, hogy ne törje szét a UI-t
    for c in show_df.columns:
        show_df[c] = show_df[c].apply(lambda x: "" if pd.isna(x) else str(x)[:120])
    st.markdown(
        '<div class="fpi-table-wrap-v120">' + show_df.to_html(index=False, escape=True, classes="fpi-mini-table-v120") + '</div>',
        unsafe_allow_html=True,
    )

st.markdown(
    """
    <style>
    .fpi-empty-state-v120 {
        display:flex;
        align-items:center;
        gap:14px;
        padding:18px 20px;
        margin:10px 0 14px 0;
        border-radius:18px;
        background:linear-gradient(135deg,#f8fafc,#eef6ff);
        border:1px solid #cbd5e1;
        box-shadow:0 8px 22px rgba(15,23,42,.06);
        color:#0f172a !important;
        min-height:72px;
    }
    .fpi-empty-state-v120 * { color:#0f172a !important; }
    .fpi-empty-icon-v120 {
        width:42px;height:42px;border-radius:999px;
        display:flex;align-items:center;justify-content:center;
        background:#dbeafe;border:1px solid #bfdbfe;
        font-size:1.25rem;flex:0 0 auto;
    }
    .fpi-empty-title-v120 { font-weight:950;font-size:1.02rem;margin-bottom:3px; }
    .fpi-empty-body-v120 { color:#475569 !important;font-size:.92rem;line-height:1.35; }
    .fpi-table-wrap-v120 {
        background:#ffffff;
        border:1px solid #cbd5e1;
        border-radius:18px;
        overflow:hidden;
        box-shadow:0 8px 22px rgba(15,23,42,.06);
        margin:10px 0 14px 0;
    }
    table.fpi-mini-table-v120 {
        width:100%;border-collapse:collapse;background:#ffffff;color:#0f172a;font-size:.92rem;
    }
    table.fpi-mini-table-v120 th {
        background:#eaf3ff;color:#0f172a;font-weight:950;text-align:left;padding:10px 12px;border-bottom:1px solid #cbd5e1;
    }
    table.fpi-mini-table-v120 td {
        color:#0f172a;padding:9px 12px;border-bottom:1px solid #e2e8f0;vertical-align:top;
    }
    table.fpi-mini-table-v120 tr:nth-child(even) td { background:#f8fafc; }
    /* Üres streamlit dataframe/plotly konténerek ne uralják a képet */
    [data-testid="stDataFrame"]:empty,
    .element-container:has([data-testid="stDataFrame"]:empty) { display:none !important; }
    </style>
    """,
    unsafe_allow_html=True,
)

# V9.4 - Hét és meccsnap kontextus ellenőrzése
match_ctx_v94 = _fpi_selected_match_context_v94()
week_context_df_v94 = _fpi_week_context_df_v94(df, match_ctx_v94.get("match_date"))
week_warnings_v94 = _fpi_match_week_warning_v94(df, selected_week, match_ctx_v94.get("match_date"))

with st.expander("📅 Hét / meccsnap / feltöltött fájlok ellenőrzése", expanded=True):
    st.markdown(
        f"**Mai nap:** {match_ctx_v94.get('today')} | **Mai hét:** {match_ctx_v94.get('today_week')} | "
        f"**Ellenfél:** {match_ctx_v94.get('opponent') or 'n.a.'} | **Meccsnap:** {match_ctx_v94.get('match_date')} | "
        f"**Meccshét:** {match_ctx_v94.get('match_week')} | **Kiválasztott hét:** {selected_week}"
    )
    if week_context_df_v94 is not None and not week_context_df_v94.empty:
        fpi_compact_table(week_context_df_v94, max_rows=8)
    else:
        fpi_empty_state("Nincs feltöltött fájlösszefoglaló", "A hétellenőrzéshez előbb tölts fel GPS vagy taktikai fájlt.", "📁")
    if week_warnings_v94:
        for w in week_warnings_v94:
            st.warning(w)
    else:
        st.success("A kiválasztott hét, a mai nap és a meccsnap alapján nincs nyilvánvaló hétkeveredés.")

if save_current_to_memory:
    ok, msg = save_to_memory(df)
    if ok:
        st.sidebar.success(msg)
    else:
        st.sidebar.error(msg)

analysis_df_full = merge_with_memory(df, use_memory)
analysis_base_df = analysis_df_full[analysis_df_full["player_name"].isin(selected_players)]
base_insights = team_insights(analysis_base_df, selected_week)
micro_insights = microcycle_insights(analysis_base_df, selected_week)
style_insights = playstyle_insights(analysis_base_df, selected_week, selected_playstyle)
pattern_insights = build_pattern_insights(analysis_base_df, selected_week)
readiness_score, readiness_components, readiness_reasons = calculate_readiness_score(analysis_base_df, selected_week, selected_playstyle)
weekly_fingerprints = build_weekly_fingerprints(analysis_base_df)
current_fp = weekly_fingerprints[weekly_fingerprints["week"] == selected_week]
periodization_type = current_fp["periodizacios_tipus"].iloc[0] if not current_fp.empty and "periodizacios_tipus" in current_fp.columns else "Nincs elég adat"
all_insights = sorted(base_insights + micro_insights + style_insights + pattern_insights, key=lambda x: SEVERITY_RANK.get(x.severity, 9))[:16]
coaching_priorities = build_adaptive_recommendations(all_insights, readiness_score, periodization_type, pattern_insights, selected_playstyle)
weekly_summary_text = build_weekly_summary(all_insights, selected_week, selected_playstyle)
all_insights = humanize_insights(all_insights)
coaching_priorities = humanize_priority_list(coaching_priorities)
weekly_summary_text = coach_friendly_phrase(weekly_summary_text)
weekly_summary_text += (
    f"\n\nMeccskészültség: {readiness_score}/100 ({score_to_label(readiness_score)})"
    f"\nPeriodizációs besorolás: {periodization_type}"
)
player_risk_df = calculate_player_risk(analysis_base_df, selected_week)
high_risk_count, medium_risk_count = _fpi_count_risk_levels_v126(player_risk_df)
week_status_info = week_completeness_summary(analysis_base_df, selected_week)
past_week_review_df, past_week_review_text = build_past_week_review(
    analysis_base_df,
    selected_week,
    selected_playstyle,
)
current_remaining_plan_df, current_remaining_text = build_current_remaining_days_plan(
    analysis_base_df,
    selected_week,
    selected_playstyle,
    readiness_score,
    periodization_type,
    player_risk_df,
)
next_week_plan_df, next_week_plan_text = build_next_week_plan_v5(
    analysis_base_df,
    selected_week,
    selected_playstyle,
    readiness_score,
    periodization_type,
    player_risk_df,
    past_week_review_df,
    current_remaining_plan_df,
)
next_microcycle_plan_df = next_week_plan_df
player_next_actions_df = build_player_next_actions(player_risk_df, analysis_base_df, selected_week)
forward_summary_text = (
    "MÚLT HÉT -> ERRE A HÉTRE\n"
    + str(past_week_review_text)
    + "\n\nAKTUÁLIS HÉT -> HÁTRALÉVŐ NAPOK\n"
    + str(current_remaining_text)
    + "\n\nJÖVŐ HÉT\n"
    + str(next_week_plan_text)
)




def _fpi_benchmark_browser_df_v124(age: str, level: str, position: str, playmodel: str, metric_filter: str = "Összes mutató") -> pd.DataFrame:
    """Benchmark böngésző táblázat a Metodika oldalhoz.
    V12.5: mindig ad vissza értelmezhető sort; a szűrők nem üríthetik ki véletlenül a táblát.
    """
    metric_names = {
        "total_distance": "Össztáv",
        "training_load": "Load / terhelési pont",
        "hsr_distance": "HSR / nagysebességű futás",
        "sprint_distance": "Sprint táv",
        "sprints": "Sprint darabszám",
        "high_efforts": "High Efforts",
    }
    metric_alias = {v: k for k, v in metric_names.items()}
    # Régebbi címkék / esetleges elgépelések kezelése
    metric_alias.update({
        "Load": "training_load",
        "Terhelési pont": "training_load",
        "HSR": "hsr_distance",
        "Sprint": "sprint_distance",
        "High efforts": "high_efforts",
    })
    prof = _fpi_build_reference_profile_v112(age, level, position, playmodel)
    ranges = prof.get("ranges", {}) or {}
    wanted_metric = None if metric_filter in [None, "", "Összes mutató"] else metric_alias.get(str(metric_filter), None)
    rows = []
    for metric, label in metric_names.items():
        if wanted_metric and metric != wanted_metric:
            continue
        rng = ranges.get(metric) or FPI_REFERENCE_BASE_RANGES_V112.get(metric)
        if not rng:
            continue
        weekly_label, avg_label, low, high, avg_low, avg_high = rng
        rows.append({
            "Mutató": label,
            "Korosztály": age,
            "Bajnoki szint": level,
            "Poszt": position,
            "Játékmodell": playmodel,
            "Heti célzóna": weekly_label,
            "Edzésátlag célzóna": avg_label,
            "Heti alsó %": low,
            "Heti felső %": high,
            "Edzésátlag alsó %": avg_low,
            "Edzésátlag felső %": avg_high,
        })
    if not rows:
        # Biztonsági fallback: ha a szűrőcímke nem található, mutassuk az összes mutatót.
        for metric, label in metric_names.items():
            rng = ranges.get(metric) or FPI_REFERENCE_BASE_RANGES_V112.get(metric)
            if not rng:
                continue
            weekly_label, avg_label, low, high, avg_low, avg_high = rng
            rows.append({
                "Mutató": label, "Korosztály": age, "Bajnoki szint": level, "Poszt": position, "Játékmodell": playmodel,
                "Heti célzóna": weekly_label, "Edzésátlag célzóna": avg_label,
                "Heti alsó %": low, "Heti felső %": high, "Edzésátlag alsó %": avg_low, "Edzésátlag felső %": avg_high,
            })
    return pd.DataFrame(rows)


def _fpi_render_benchmark_browser_table_v125(df: pd.DataFrame) -> None:
    """Olvasható benchmark táblázat dataframe mellett/helyett."""
    if df is None or df.empty:
        st.warning("Ehhez a szűréshez nincs megjeleníthető benchmark. Válassz másik mutatót vagy profilt.")
        return
    st.dataframe(df, use_container_width=True, hide_index=True)
    # Extra, biztosan látható kompakt HTML táblázat, ha a Streamlit táblázat stílusa furcsán viselkedne.
    rows_html = []
    for _, r in df.iterrows():
        rows_html.append(
            f"<tr><td>{html.escape(str(r.get('Mutató','')))}</td>"
            f"<td>{html.escape(str(r.get('Heti célzóna','')))}</td>"
            f"<td>{html.escape(str(r.get('Edzésátlag célzóna','')))}</td>"
            f"<td>{html.escape(str(r.get('Korosztály','')))} / {html.escape(str(r.get('Bajnoki szint','')))} / {html.escape(str(r.get('Poszt','')))} / {html.escape(str(r.get('Játékmodell','')))}</td></tr>"
        )
    st.markdown(
        """
        <style>
        .fpi-benchmark-table{width:100%;border-collapse:separate;border-spacing:0;margin-top:12px;background:#ffffff;border:1px solid #dbeafe;border-radius:16px;overflow:hidden;}
        .fpi-benchmark-table th{background:#e0f2fe;color:#0f172a;padding:10px 12px;text-align:left;font-weight:950;border-bottom:1px solid #bfdbfe;}
        .fpi-benchmark-table td{background:#ffffff;color:#0f172a;padding:10px 12px;border-bottom:1px solid #e5e7eb;}
        .fpi-benchmark-table tr:last-child td{border-bottom:none;}
        </style>
        <table class="fpi-benchmark-table">
            <thead><tr><th>Mutató</th><th>Heti célzóna</th><th>Edzésátlag célzóna</th><th>Aktív profil</th></tr></thead>
            <tbody>
        """ + "".join(rows_html) + """
            </tbody>
        </table>
        """,
        unsafe_allow_html=True,
    )


def render_methodology_tab() -> None:
    """Football Performance Intelligence metodikai oldal – transzparens, de nem túlzottan képletszintű kommunikáció."""
    st.markdown("## 📚 Football Performance Intelligence metodika")
    st.markdown(
        """
        <div class="fpi-summary-card">
            <h3>Mit csinál a Football Performance Intelligence?</h3>
            <p>
            A <b>Football Performance Intelligence</b> döntéstámogató rendszer. A GPS-adatokból és opcionálisan
            feltöltött taktikai PDF/Excel anyagokból olyan vezetői információkat készít, amelyek segítik
            az edzői, erőnléti és teljesítménydiagnosztikai döntéseket.
            </p>
            <p>
            A metodika <b>nemzetközi sporttudományi szakirodalomra, gyakorlati teljesítménydiagnosztikai tapasztalatokra,
            saját adatokra és saját fejlesztésű elemzési logikára</b> épül. A rendszer nem orvosi diagnózist és nem
            automatikus döntést ad, hanem adatvezérelt szakmai támogatást.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="fpi-action-card">
            <b>Fontos értelmezés:</b><br>
            A Football Performance Intelligence által számított pontszámok és kockázati szintek becslések,
            amelyek több terhelési és teljesítménymutató együttes értelmezésén alapulnak. Az eredményeket
            az edzői megfigyelésekkel, orvosi információkkal, wellness/RPE adatokkal és klubspecifikus
            kontextussal együtt érdemes értelmezni.
        </div>
        """,
        unsafe_allow_html=True,
    )

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(
            """
            <div class="fpi-kpi-panel">
                <div class="label">Adatforrás</div>
                <div class="value">GPS+</div>
                <div class="note">GPS export az alap; taktikai PDF/Excel opcionálisan beépíthető.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with c2:
        st.markdown(
            """
            <div class="fpi-kpi-panel">
                <div class="label">Értelmezési logika</div>
                <div class="value">4 hét</div>
                <div class="note">A friss heti adatokat saját előzményhez és referenciazónákhoz viszonyítja.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with c3:
        st.markdown(
            """
            <div class="fpi-kpi-panel">
                <div class="label">Kimenet</div>
                <div class="value">Döntés</div>
                <div class="note">Readiness, risk, benchmark és mikrociklus-javaslat.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("### 1. Adatimport és Smart Mapper")
    st.write(
        "A rendszer a feltöltött GPS exportból automatikusan próbálja felismerni a fontos oszlopokat. "
        "Magyar és angol mezőneveket is kezel, például: játékosnév, dátum, edzés/meccs típus, össztáv, "
        "Load, HSR, sprint, gyorsulás, lassítás, High Efforts, játékperc és poszt."
    )
    st.info(
        "A Smart Mapper célja, hogy különböző GPS-rendszerek exportjaiból is egységes Football Performance Intelligence adatszerkezet készüljön."
    )

    st.markdown("### 2. Dátum-, hét- és meccskörnyezet felismerése")
    st.write(
        "A Football Performance Intelligence ISO hét alapján csoportosítja az adatokat, és figyeli, hogy a feltöltött fájl "
        "mely heteket tartalmazza. A meccsnap, meccshét és kiválasztott hét összevetése segít kiszűrni, ha véletlenül "
        "másik hét vagy több hét adatai keverednek a riportba."
    )
    st.write(
        "A rendszer kezeli a normál dátumokat, dátum+idő mezőket, valamint több exportban előforduló szöveges dátumformákat is."
    )

    st.markdown("### 3. Readiness Score – mit vesz figyelembe?")
    st.write(
        "A Readiness Score a csapat vagy játékos aktuális terhelési állapotának becslése. Nem egyetlen mérőszám, hanem "
        "több rövid és középtávú terhelési jel együttes értelmezése."
    )
    st.markdown(
        """
        <div class="fpi-clean-card">
        <b>A readiness értelmezésében szerepet kaphat:</b><br>
        • heti Load és annak változása az előző hetekhez képest<br>
        • elmúlt 3–7 nap terhelési képe<br>
        • 4 hetes terhelési trendek<br>
        • HSR / nagysebességű futás<br>
        • sprinttáv és sprintdarabszám<br>
        • High Efforts, gyorsulások és lassítások<br>
        • edzés–meccs arány és meccs előtti frissítés/taper logika<br>
        • játékmodellhez való illeszkedés<br>
        • pulzus vagy HRV, ha a fájlban rendelkezésre áll
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.caption("A readiness nem orvosi állapotjelző, hanem terhelési mintázatokból képzett döntéstámogató jelzés.")

    st.markdown("### 4. Player Risk Score – mit jelent a kockázati besorolás?")
    st.write(
        "A Player Risk Score a túlterhelés, alulterhelés vagy hirtelen terhelésváltozás korai felismerését támogatja. "
        "A rendszer elsősorban a játékos saját előző heteivel hasonlítja össze az aktuális hetet, ezért nem csak abszolút számokat néz."
    )
    st.markdown(
        """
        <div class="fpi-clean-card">
        <b>A kockázati értékelés fő bemenetei:</b><br>
        • heti Load változás saját 4 hetes átlaghoz képest<br>
        • sprinttáv és maximális sebesség expozíció<br>
        • High Efforts, gyorsulás és lassítás jellegű neuromuszkuláris terhelés<br>
        • játékperc / exposure változás<br>
        • max sebesség trendje saját korábbi csúcshoz képest<br>
        • kapus vagy mezőnyjátékos szerep szerinti eltérő súlyozás<br>
        • alacsony terhelés utáni hirtelen terhelésnövekedés lehetősége
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.warning(
        "A kockázati szint nem sérülés-előrejelzés. A cél az, hogy a szakmai stáb időben észrevegye a szokatlan terhelési mintákat."
    )

    st.markdown("### 5. Benchmark Engine – hogyan választ referenciát?")
    st.write(
        "A benchmarkrendszer nem egyetlen általános átlaghoz hasonlít. A referenciaértékek a kiválasztott profil alapján "
        "aktualizálódnak: korosztály, bajnoki szint, poszt és játékmodell szerint. Ha van posztadat, a rendszer játékosonként "
        "súlyoz; ha nincs, mezőnyjátékos fallbacket használ."
    )
    st.markdown(
        """
        <div class="fpi-clean-card">
        <b>A referenciazónák fő dimenziói:</b><br>
        • korosztály: Felnőtt, U21, U19, U17, U16, U15, U14, U13<br>
        • szint: NB I, NB II, NB III, Akadémia, Regionális, Megye I, Egyéb<br>
        • poszt: kapus, hátvéd, középpályás, szélső, csatár és ezek bontásai<br>
        • játékmodell: dominancia, magas presszing, átmeneti játék, direkt játék, kiegyensúlyozott<br>
        • mutató: össztáv, Load, HSR, sprinttáv, sprintdarabszám, High Efforts
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander("🔎 Benchmark böngésző – szűrés mutató, korosztály, bajnoki szint és poszt szerint", expanded=False):
        st.caption("A célzónák azt mutatják, hogy a heti edzésösszeg és az egy edzésre jutó átlag hány százaléka az adott profil meccsreferenciájának.")
        bc1, bc2, bc3, bc4, bc5 = st.columns([1, 1, 1, 1, 1.2])
        metric_options_v124 = [
            "Összes mutató",
            "Össztáv",
            "Load / terhelési pont",
            "HSR / nagysebességű futás",
            "Sprint táv",
            "Sprint darabszám",
            "High Efforts",
        ]
        with bc1:
            bench_metric_v124 = st.selectbox("Mutató", metric_options_v124, index=0, key="method_benchmark_metric_v128")
        with bc2:
            bench_age_v124 = st.selectbox("Korosztály", FPI_REFERENCE_AGE_OPTIONS_V112, index=0, key="method_benchmark_age_v128")
        with bc3:
            bench_level_v124 = st.selectbox("Bajnoki szint", FPI_REFERENCE_LEVEL_OPTIONS_V112, index=1, key="method_benchmark_level_v128")
        with bc4:
            bench_position_v124 = st.selectbox("Poszt", FPI_REFERENCE_POSITION_OPTIONS_V112, index=4, key="method_benchmark_position_v128")
        with bc5:
            bench_playmodel_v124 = st.selectbox("Játékmodell", FPI_PLAYMODEL_OPTIONS_V112, index=4, key="method_benchmark_playmodel_v128")

        benchmark_df_v124 = _fpi_benchmark_browser_df_v124(
            bench_age_v124,
            bench_level_v124,
            bench_position_v124,
            bench_playmodel_v124,
            bench_metric_v124,
        )
        _fpi_render_benchmark_browser_table_v125(benchmark_df_v124)
        st.info("Példa: ha a HSR heti célzóna 150–250%, akkor a teljes heti nagysebességű futás célja az adott profil meccsreferenciájának kb. 1,5–2,5-szerese.")

    st.markdown("### 6. Mikrociklus Motor – hogyan készül a javaslat?")
    st.write(
        "A mikrociklus javaslat a kiválasztott hét állapotából, a heti célból, a meccsnapból, a pihenőnapokból, "
        "az edzésszámból, a readinessből, a risk jelzésekből és a benchmarkeltérésekből épül."
    )
    st.markdown(
        """
        <div class="fpi-clean-card">
        <b>A mikrociklus tervezésében figyelembe vett fő elemek:</b><br>
        • aktuális heti volumen és Load<br>
        • HSR és sprintinger hiánya vagy többlete<br>
        • High Efforts és neuromuszkuláris terhelés<br>
        • játékosonkénti kockázati jelzések<br>
        • heti típus: regenerációs, stabilizáló, terhelésfokozó, fejlesztő, formaidőzítő, mérkőzésre felkészítő<br>
        • meccsnap és MD-struktúra<br>
        • játékmodell és opcionális taktikai prioritások
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### 7. Kapuskezelés és posztsúlyozás")
    st.write(
        "A kapusokat a rendszer nem ugyanazzal a sebesség- és sprintlogikával értékeli, mint a mezőnyjátékosokat. "
        "Kapusoknál a teljes Load, High Efforts, játékperc/exposure és neuromuszkuláris jellegű változások nagyobb hangsúlyt kapnak, "
        "míg a HSR/sprint elvárás csökkentett súllyal szerepel."
    )

    st.markdown("### 8. Taktikai integráció")
    st.write(
        "Ha taktikai PDF vagy Excel is rendelkezésre áll, a Football Performance Intelligence nem csak GPS-only módon működik. "
        "A taktikai inputból ellenfélprofil, játékmodell, kiemelt veszélyek, taktikai prioritások és meccsterv-jellegű szempontok "
        "kerülhetnek be az összegzésbe és a mikrociklus javaslatba."
    )


    st.markdown("### 9. Tactical Framework - 7 dimenzióból 9 stratégiai profil")
    st.write(
        "A Tactical Pro+ nem csak kulcsszavakat keres. A taktikai inputokat egy egységes Football Performance Intelligence "
        "keretrendszerbe rendezi: 7 taktikai dimenzióból képez csapatprofilt, majd ezt 9 stratégiai profilhoz viszonyítja. "
        "A cél nem kész meccsterv helyettesítése, hanem az edzői döntés gyorsítása és strukturálása."
    )
    st.markdown(
        """
        <div class="fpi-clean-card">
        <b>A 7 taktikai dimenzió:</b><br>
        1. Letámadás<br>
        2. Labdakihozatal<br>
        3. Átmenetek<br>
        4. Támadó játék<br>
        5. Pontrúgások<br>
        6. Labdabirtoklás<br>
        7. Lövésprofil
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.caption("A dimenziók 0-10 skálán értelmezhetők. Nem abszolút scouting-minősítés, hanem döntéstámogató profil.")
    st.markdown(
        """
        <div class="fpi-clean-card">
        <b>A két fő taktikai tengely:</b><br>
        • Játékstílus: Direkt -> Vegyes -> Kiegyensúlyozott -> Kontroll -> Agresszív<br>
        • Blokkmagasság: Mély -> Alacsony-közép -> Közép -> Közép-magas -> Magas
        </div>
        """,
        unsafe_allow_html=True,
    )
    try:
        st.dataframe(
            pd.DataFrame(_fpi_global_strategy_palette_rows_v129(), columns=["Kód", "Stratégia", "Játékstílus", "Blokkmagasság", "Jelentés"]),
            use_container_width=True,
            hide_index=True,
        )
    except Exception:
        pass
    st.info(
        "A rendszer nem feltétlenül csak egy címkét használ. A gyakorlatban elsődleges profil, alternatív profil és meccsspecifikus fókuszok is készülhetnek. Például: BAT elsődleges, POZ/KIE alternatív megoldással, rest defense és átmeneti fókuszokkal."
    )

    st.markdown("### 10. Mit nem állít a rendszer?")
    st.error(
        "A Football Performance Intelligence nem állítja, hogy egy pontszám önmagában megmondja a sérülést, a teljesítményt vagy a mérkőzés kimenetelét. "
        "A rendszer célja az, hogy a szakmai stáb gyorsabban lássa a fontos eltéréseket, trendeket és döntési pontokat."
    )

    st.markdown("### 11. Technikai státusz")
    st.json({
        "FPI_VERSION": FPI_IMPORT_ENGINE_VERSION,
        "Smart Mapper": "aktív",
        "Week Rescue Engine": "aktív",
        "Keeper Logic": "aktív",
        "Minutes Normalization": "aktív",
        "Readiness Engine": "trend + terhelés + sebességi expozíció + játékmodell",
        "Risk Engine": "saját 4 hetes előzmény + aktuális heti eltérés",
        "Benchmark Engine": "korosztály + szint + poszt + játékmodell",
        "Microcycle Engine": "readiness + benchmark + risk + MD-struktúra",
        "Tactical Framework": "7 dimenzió -> 9 stratégiai profil -> meccsspecifikus fókusz",
        "Tactical Pro+": "opcionális inputként aktív",
    })


# =========================================================
# TACTICAL PRO+ MERGE MODULE V7.0
# PDF: automatikus HU/EN témakinyerés
# Excel: Smart Tactical Mapper csapat- és játékosstatisztikákra
# Adaptive Intelligence: ha van taktikai adat, figyelembe veszi; ha nincs, GPS-only módban fut
# =========================================================

TACTICAL_PRO_VERSION = "TACTICAL_PRO_MERGED_V070_2026_06_17"

TACTICAL_TOPIC_TAGS_FPI = {
    "formation": {"label": "Formáció / alapfelállás", "keywords": ["formation", "shape", "system", "line-up", "lineup", "starting xi", "formáció", "felállás", "játékrendszer", "4-4-2", "4-2-3-1", "4-3-3", "3-5-2", "3-4-3", "5-3-2", "5-4-1"]},
    "build_up": {"label": "Labdakihozatal / támadásépítés", "keywords": ["build-up", "build up", "first phase", "second phase", "goal kick", "short goal kick", "progression", "progressive pass", "third man", "pivot", "half-space", "switch of play", "labdakihozatal", "támadásépítés", "építkezés", "kirúgás", "progresszív passz", "harmadik ember", "félterület", "oldalváltás"]},
    "direct_play": {"label": "Direkt játék / hosszú labda", "keywords": ["direct play", "long ball", "long pass", "second ball", "aerial duel", "target man", "vertical", "direct attack", "direkt játék", "hosszú labda", "második labda", "felívelés", "fejpárbaj", "céljátékos", "vertikális"]},
    "pressing": {"label": "Letámadás / presszing", "keywords": ["press", "pressing", "high press", "mid press", "low press", "counterpress", "ppda", "pressure", "pressing trigger", "trap", "high recovery", "letámadás", "presszing", "magas letámadás", "visszatámadás", "nyomás", "trigger", "csapda", "magas labdaszerzés"]},
    "defensive_block": {"label": "Védekezési blokk / blokkmagasság", "keywords": ["low block", "mid block", "middle block", "high block", "defensive block", "compact", "defensive line", "line height", "deep defending", "mély blokk", "középső blokk", "magas blokk", "védekezési blokk", "kompakt", "védelmi vonal", "blokkmagasság", "mély védekezés"]},
    "transition_attack": {"label": "Támadó átmenet / kontrák", "keywords": ["transition", "attacking transition", "counterattack", "counter attack", "counter-attacks", "fast attack", "quick attack", "after regain", "after winning", "átmenet", "támadó átmenet", "kontra", "kontratámadás", "gyors támadás", "labdaszerzés után", "labdanyerés után"]},
    "transition_defense": {"label": "Védekező átmenet / rest defense", "keywords": ["defensive transition", "after losing", "after loss", "rest defense", "counter prevention", "cover behind", "védekező átmenet", "labdavesztés után", "kontrák elleni védekezés", "átmeneti védekezés", "biztosítás", "visszarendeződés"]},
    "chance_creation": {"label": "Helyzetkialakítás / támadóharmad", "keywords": ["chance creation", "key pass", "shot assist", "box entry", "penalty area", "final third", "through ball", "cutback", "cross", "xg", "expected goals", "helyzetkialakítás", "kulcspassz", "box entry", "tizenhatos", "támadóharmad", "mélységi passz", "visszagurítás", "beadás", "várható gól"]},
    "wide_play": {"label": "Szélső játék / oldali dominancia", "keywords": ["wide play", "wing", "flank", "left side", "right side", "overlap", "underlap", "fullback", "crossing", "side dominance", "szélső játék", "szél", "bal oldal", "jobb oldal", "átfedés", "aláfutás", "beadás", "oldali dominancia"]},
    "central_play": {"label": "Középső játék / félterületek", "keywords": ["central", "middle", "half-space", "between the lines", "pocket", "zone 14", "inside channel", "central overload", "középen", "középső", "félterület", "vonalak között", "zseb", "14-es zóna", "belső csatorna"]},
    "set_pieces": {"label": "Pontrúgások", "keywords": ["set piece", "set pieces", "corner", "corner kick", "free kick", "throw-in", "attacking corner", "defensive corner", "near post", "far post", "aerial", "header", "pontrúgás", "pontrúgások", "szöglet", "szabadrúgás", "bedobás", "rövid oldal", "hosszú oldal", "fejpárbaj", "fejes"]},
    "key_players": {"label": "Kulcsjátékosok", "keywords": ["key player", "danger man", "main threat", "top scorer", "creator", "playmaker", "progressor", "target man", "dribbler", "1v1", "finisher", "kulcsjátékos", "veszélyes játékos", "fő veszély", "gólkirály", "kreatív játékos", "irányító", "progresszor", "céljátékos", "cselező", "egy az egy", "befejező"]},
    "weakness_risk": {"label": "Gyengeségek / kockázatok", "keywords": ["weakness", "weaknesses", "risk", "risks", "vulnerable", "vulnerability", "exposed", "space behind", "gap", "mistake", "error", "turnover", "lost balls", "danger", "threat", "gyengeség", "gyengeségek", "kockázat", "sebezhető", "sebezhetőség", "nyitott terület", "mögötti terület", "rés", "hiba", "labdavesztés", "veszély"]},
    "strength": {"label": "Erősségek", "keywords": ["strength", "strengths", "strong", "advantage", "edge", "dominant", "effective", "efficient", "erősség", "erősségek", "erős", "előny", "domináns", "hatékony", "kiemelkedő"]},
    "recommendation": {"label": "Javaslat / meccsterv", "keywords": ["recommendation", "recommend", "should", "game plan", "match plan", "plan a", "plan b", "solution", "exploit", "avoid", "focus", "priority", "target", "javaslat", "ajánlás", "meccsterv", "mérkőzésterv", "terv a", "terv b", "megoldás", "kihasználni", "elkerülni", "fókusz", "prioritás"]},
}

TACTICAL_TEAM_ALIASES_FPI = {
    "possession_pct": ["ball possession", "possession", "possession %", "labdabirtoklás", "labdabirtoklás %", "birtoklás"],
    "shots": ["shots", "total shots", "attempts", "lövés", "lövések", "összes lövés"],
    "xg": ["xg", "expected goals", "várható gól", "várható gólok"],
    "entries_box": ["box entries", "entries into box", "penalty box entries", "tizenhatosba belépés", "büntetőterület", "16-osba belépés"],
    "final_third_entries": ["final third entries", "entries to final third", "támadóharmad belépések", "utolsó harmad"],
    "key_passes": ["key passes", "shot assists", "chances created", "kulcspassz", "helyzetkialakítás", "lövést előkészítő"],
    "corners": ["corners", "corner kicks", "szögletek", "szöglet"],
    "ppda": ["ppda", "passes allowed per defensive action", "passz engedett védekező akciónként"],
    "pressing_success_pct": ["pressing success", "successful pressing", "team pressing successful", "letámadás sikeresség", "presszing sikeresség"],
    "passes_accurate_pct": ["pass accuracy", "passes accurate", "passing accuracy", "passzpontosság", "átadáspontosság"],
    "crosses": ["crosses", "successful crosses", "beadások", "beadás"],
    "recoveries": ["recoveries", "ball recoveries", "regains", "labdaszerzések", "visszaszerzések"],
    "lost_balls": ["lost balls", "losses", "turnovers", "labdavesztések", "elvesztett labdák"],
    "counterattacks": ["counterattacks", "counter attacks", "fast attacks", "kontrák", "kontratámadások", "gyors támadások"],
}

TACTICAL_PLAYER_ALIASES_FPI = {
    "player": ["player", "player name", "name", "játékos", "játékos neve", "név"],
    "position": ["position", "pos", "role", "poszt", "pozíció"],
    "minutes_played": ["minutes played", "minutes", "mins", "played", "játékperc", "percek", "játszott perc"],
    "passes": ["passes", "pass", "passzok", "passz"],
    "progressive_passes": ["progressive passes", "progressive pass", "progresszív passz", "előrehaladó passz"],
    "key_passes": ["key passes", "shot assists", "kulcspassz", "helyzetkialakítás"],
    "shots": ["shots", "attempts", "lövések", "lövés"],
    "xg": ["xg", "expected goals", "várható gól"],
    "assists": ["assists", "assist", "gólpassz"],
    "recoveries": ["recoveries", "ball recoveries", "labdaszerzések"],
    "interceptions": ["interceptions", "interception", "közbelépés", "labdaszerzés"],
    "defensive_challenges": ["defensive challenges", "defensive duels", "védekező párharc", "párharc"],
    "crosses": ["crosses", "beadások", "beadás"],
}


# V7.6 – kibővített HU/EN taktikai szótár.
# Cél: ne csak SportsBase/egyféle riport kulcsszavaira reagáljon, hanem Provision/Wyscout/Hudl jellegű
# taktikai PDF-ekben előforduló kifejezésekre is.
TACTICAL_TOPIC_TAGS_FPI.update({
    "formation": {
        "label": "Formáció / alapfelállás",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("formation", {}).get("keywords", []) + [
            "formation", "shape", "system", "structure", "line-up", "lineup", "starting eleven", "starting xi",
            "base formation", "out of possession shape", "in possession shape",
            "formáció", "felállás", "alapfelállás", "játékrendszer", "szerkezet", "kezdőcsapat",
            "labdával", "labda nélkül", "védekezésben", "támadásban",
            "4-4-2", "4-2-3-1", "4-3-3", "3-5-2", "3-4-3", "5-3-2", "5-4-1", "4-1-4-1", "4-3-1-2",
        ])),
    },
    "build_up": {
        "label": "Labdakihozatal / támadásépítés",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("build_up", {}).get("keywords", []) + [
            "build-up", "build up", "buildout", "first phase", "second phase", "goal kick", "goal kicks",
            "short goal kick", "deep build", "progression", "progressive pass", "progressive passes",
            "progressive carry", "third man", "centre back", "center back", "fullback", "pivot", "six",
            "number 6", "half-space", "switch of play", "circulation", "positional attack", "possession phase",
            "labdakihozatal", "támadásépítés", "építkezés", "első fázis", "második fázis", "kirúgás",
            "rövid kirúgás", "progresszió", "progresszív passz", "progresszív labdavezetés", "harmadik ember",
            "belső védő", "szélső védő", "hatossal", "6-os", "félterület", "oldalváltás", "labdajáratás",
            "pozíciós támadás", "labdabirtoklási fázis",
        ])),
    },
    "direct_play": {
        "label": "Direkt játék / hosszú labda",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("direct_play", {}).get("keywords", []) + [
            "direct play", "long ball", "long pass", "second ball", "aerial duel", "target man", "flick-on",
            "vertical", "early forward", "long distribution", "direct attack", "route one", "directness",
            "direkt játék", "hosszú labda", "hosszú passz", "második labda", "felívelés", "fejpárbaj",
            "céljátékos", "lecsorgó", "vertikális", "korai előrejáték", "direkt támadás",
        ])),
    },
    "pressing": {
        "label": "Letámadás / presszing",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("pressing", {}).get("keywords", []) + [
            "press", "pressing", "high press", "mid press", "low press", "counterpress", "counter-press",
            "ppda", "pressure", "pressing trigger", "trap", "press trap", "forced turnover", "high recovery",
            "challenge intensity", "defensive actions", "intensity", "aggressive press", "recoveries high",
            "letámadás", "presszing", "magas letámadás", "középső presszing", "visszatámadás", "nyomás",
            "presszing trigger", "trigger", "csapda", "pressingcsapda", "kikényszerített labdavesztés",
            "magas labdaszerzés", "védekező akció", "intenzitás", "agresszív letámadás",
        ])),
    },
    "defensive_block": {
        "label": "Védekezési blokk / blokkmagasság",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("defensive_block", {}).get("keywords", []) + [
            "low block", "mid block", "middle block", "high block", "defensive block", "compact", "compactness",
            "defensive line", "back line", "line height", "block height", "deep defending", "drop", "retreat",
            "defensive shape", "out of possession", "defending third",
            "mély blokk", "középső blokk", "magas blokk", "védekezési blokk", "kompakt", "kompaktság",
            "védelmi vonal", "védősor", "blokkmagasság", "mély védekezés", "visszazár", "visszarendeződés",
            "labda nélkül", "védekező szerkezet",
        ])),
    },
    "transition_attack": {
        "label": "Támadó átmenet / kontrák",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("transition_attack", {}).get("keywords", []) + [
            "transition", "attacking transition", "offensive transition", "counterattack", "counter attack", "counter-attacks",
            "fast attack", "quick attack", "break", "breakaway", "after regain", "after winning", "regain and go",
            "direct attacks", "fast attacks", "counter attacks", "attacks after regain",
            "átmenet", "támadó átmenet", "kontra", "kontratámadás", "gyors támadás", "gyors átmenet",
            "labdaszerzés után", "labdanyerés után", "visszaszerzés után", "indítás", "megindulás",
        ])),
    },
    "transition_defense": {
        "label": "Védekező átmenet / rest defense",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("transition_defense", {}).get("keywords", []) + [
            "defensive transition", "after losing", "after loss", "rest defense", "counter attack prevention",
            "counter prevention", "defend transition", "negative transition", "cover behind", "protection behind",
            "rest-defence", "loss of possession", "turnover",
            "védekező átmenet", "labdavesztés után", "rest defense", "kontrák elleni védekezés",
            "átmeneti védekezés", "negatív átmenet", "biztosítás", "mögöttes biztosítás", "visszarendeződés",
            "labdavesztés", "elvesztett labda",
        ])),
    },
    "chance_creation": {
        "label": "Helyzetkialakítás / támadóharmad",
        "keywords": [
            "chance creation", "key pass", "key passes", "shot assist", "shot assists", "box entry", "box entries",
            "penalty area", "final third", "final-third", "final third entries", "entries to final third",
            "entries into box", "through ball", "cutback", "cross", "low cross", "deep cross", "half-space cross",
            "xg", "expected goals", "shots", "shots on target", "big chances", "goal attempts",
            "helyzetkialakítás", "kulcspassz", "lövést előkészítő", "box entry", "tizenhatos", "büntetőterület",
            "támadóharmad", "mélységi passz", "visszagurítás", "beadás", "lapos beadás", "xg", "várható gól",
            "lövések", "kaput eltaláló lövés", "nagy helyzet",
        ],
    },
    "wide_play": {
        "label": "Szélső játék / oldali dominancia",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("wide_play", {}).get("keywords", []) + [
            "wide play", "wing", "flank", "left side", "right side", "overlap", "underlap", "fullback", "wingback",
            "crossing", "side dominance", "touchline", "wide overload", "flank overload", "attacks left", "attacks right",
            "szélső játék", "szél", "oldal", "bal oldal", "jobb oldal", "átfedés", "aláfutás", "szélső védő",
            "wingback", "beadás", "oldali dominancia", "oldalvonal", "oldali túlterhelés", "bal oldali", "jobb oldali",
        ])),
    },
    "central_play": {
        "label": "Középső játék / félterületek",
        "keywords": [
            "central", "middle", "half-space", "half space", "between the lines", "pocket", "zone 14", "inside channel",
            "central overload", "interior", "attacking midfielder", "number 10", "third man run",
            "középen", "középső", "félterület", "félterületek", "vonalak között", "zseb", "14-es zóna",
            "belső csatorna", "középső túlterhelés", "belső középpályás", "10-es",
        ],
    },
    "set_pieces": {
        "label": "Pontrúgások",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("set_pieces", {}).get("keywords", []) + [
            "set piece", "set pieces", "corner", "corners", "corner kick", "free kick", "throw-in", "throw in",
            "penalty", "attacking corner", "defensive corner", "near post", "far post", "second ball", "aerial", "header",
            "pontrúgás", "pontrúgások", "szöglet", "szögletek", "szabadrúgás", "bedobás", "büntető",
            "támadó szöglet", "védekező szöglet", "rövid oldal", "hosszú oldal", "második labda", "fejpárbaj", "fejes",
        ])),
    },
    "key_players": {
        "label": "Kulcsjátékosok",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("key_players", {}).get("keywords", []) + [
            "key player", "danger man", "main threat", "top scorer", "creator", "playmaker", "progressor", "target man",
            "dribbler", "1v1", "one-v-one", "finisher", "captain", "most dangerous", "player to watch",
            "kulcsjátékos", "veszélyes játékos", "fő veszély", "gólkirály", "kreatív játékos", "irányító",
            "progresszor", "céljátékos", "cselező", "egy az egy", "befejező", "csapatkapitány", "legveszélyesebb",
        ])),
    },
    "weakness_risk": {
        "label": "Gyengeségek / kockázatok",
        "keywords": list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI.get("weakness_risk", {}).get("keywords", []) + [
            "weakness", "weaknesses", "risk", "risks", "vulnerable", "vulnerability", "exposed", "space behind",
            "gap", "mistake", "error", "turnover", "lost balls", "losses", "danger", "threat", "problem",
            "conceded", "conceding", "dangerous area",
            "gyengeség", "gyengeségek", "kockázat", "sebezhető", "sebezhetőség", "nyitott terület",
            "mögötti terület", "rés", "hiba", "labdavesztés", "elvesztett labda", "veszély", "fenyegetés", "probléma",
            "kapott gól", "veszélyes zóna",
        ])),
    },
    "strength": {
        "label": "Erősségek",
        "keywords": [
            "strength", "strengths", "strong", "advantage", "edge", "dominant", "effective", "efficient",
            "best", "successful", "high value", "threat",
            "erősség", "erősségek", "erős", "előny", "domináns", "hatékony", "kiemelkedő", "sikeres", "veszélyes",
        ],
    },
    "goalkeeper": {
        "label": "Kapus szerepe",
        "keywords": [
            "goalkeeper", "keeper", "gk", "sweeper keeper", "distribution", "long kick", "short pass from gk",
            "goalkeeper involvement", "keeper distribution",
            "kapus", "hálóőr", "kapusjáték", "kapus kirúgás", "kapus passz", "hosszú kirúgás", "rövid kirúgás",
        ],
    },
    "match_dynamics": {
        "label": "Meccsdinamika / fázisok",
        "keywords": [
            "first half", "second half", "opening phase", "late phase", "last 15", "tempo", "rhythm", "momentum",
            "game state", "when leading", "when trailing", "after goal", "minutes", "phase",
            "első félidő", "második félidő", "kezdő fázis", "végjáték", "utolsó 15", "tempó", "ritmus",
            "momentum", "meccsállapot", "vezetésnél", "hátrányban", "gól után", "percek", "fázis",
        ],
    },
    "recommendation": {
        "label": "Javaslat / meccsterv",
        "keywords": [
            "recommendation", "recommend", "should", "we should", "game plan", "match plan", "plan a", "plan b",
            "solution", "exploit", "avoid", "focus", "priority", "target", "press here", "attack here",
            "javaslat", "ajánlás", "meccsterv", "mérkőzésterv", "terv a", "terv b", "megoldás",
            "kihasználni", "elkerülni", "fókusz", "prioritás", "célpont", "itt presszing", "itt támadni",
        ],
    },
})


# V7.7 – provider-specifikus és fallback taktikai kulcsszavak
# Ezek a SportsBase / Provision / Wyscout / Hudl jellegű PDF-ekben gyakori, de nem mindig
# klasszikus "coach sentence" formában megjelenő címkéket is témákhoz kötik.
TACTICAL_TOPIC_TAGS_FPI.setdefault("team_stats", {"label": "Csapatstatisztikák / fő KPI-k", "keywords": []})
TACTICAL_TOPIC_TAGS_FPI.setdefault("event_actions", {"label": "Események / akciótípusok", "keywords": []})
TACTICAL_TOPIC_TAGS_FPI.setdefault("player_stats", {"label": "Játékosstatisztikák", "keywords": []})

TACTICAL_TOPIC_TAGS_FPI["team_stats"]["keywords"] = list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI["team_stats"].get("keywords", []) + [
    "main statistics", "team statistics", "match statistics", "summary statistics", "statistics",
    "ball possession", "possession", "passes accurate", "pass accuracy", "shots", "shots on target",
    "xg", "expected goals", "corners", "fouls", "offsides", "duels", "aerial duels", "recoveries",
    "csapatstatisztika", "fő statisztikák", "statisztika", "labdabirtoklás", "passzpontosság",
    "lövések", "kaput eltaláló lövések", "várható gól", "szögletek", "párharcok", "labdaszerzések",
]))
TACTICAL_TOPIC_TAGS_FPI["event_actions"]["keywords"] = list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI["event_actions"].get("keywords", []) + [
    "episode search", "episodes", "passes", "ball recoveries", "lost balls", "challenges", "carry",
    "carry >10m", "counterattacks", "counter attacks", "positional attacks", "ball possession",
    "set pieces", "sequences", "attacks", "progressions", "final third entries", "box entries",
    "action zone", "destination zone", "heat map", "pass map", "shot map", "touch map",
    "epizód", "epizódok", "passzok", "labdaszerzések", "labdavesztések", "párharcok",
    "labdavezetés", "kontrák", "pozíciós támadások", "pontrúgások", "akciózóna", "célzóna",
    "hőtérkép", "passztérkép", "lövéstérkép", "érintések térképe",
]))
TACTICAL_TOPIC_TAGS_FPI["player_stats"]["keywords"] = list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI["player_stats"].get("keywords", []) + [
    "player statistics", "player stats", "players", "minutes played", "goals", "assists", "key passes",
    "progressive passes", "defensive duels", "interceptions", "recoveries", "xg", "xa",
    "játékosstatisztika", "játékosok", "játékperc", "gólok", "gólpasszok", "kulcspasszok",
    "progresszív passzok", "védekező párharcok", "közbelépések", "labdaszerzések",
]))

# Csatoljuk a provider-címkéket konkrét taktikai témákhoz is, hogy a meccsterv motor használja őket.
for _topic, _extra in {
    "transition_attack": ["counterattacks", "counter attacks", "fast attacks", "attacks after regain", "kontrák", "kontratámadások", "gyors támadások"],
    "transition_defense": ["lost balls", "losses", "turnovers", "labdavesztések", "elvesztett labdák"],
    "build_up": ["passes", "pass map", "progressions", "ball possession", "sequences", "passzok", "passztérkép", "progressziók", "labdabirtoklás"],
    "pressing": ["ball recoveries", "recoveries", "high recoveries", "challenges", "defensive actions", "labdaszerzések", "visszaszerzések", "párharcok", "védekező akciók"],
    "chance_creation": ["shots", "shot map", "box entries", "final third entries", "key passes", "xg", "lövéstérkép", "lövések", "box entries", "támadóharmad", "kulcspasszok", "várható gól"],
    "wide_play": ["crosses", "attacks left", "attacks right", "left side", "right side", "beadások", "bal oldal", "jobb oldal"],
    "set_pieces": ["set pieces", "corners", "corner kicks", "free kicks", "throw ins", "pontrúgások", "szögletek", "szabadrúgások", "bedobások"],
    "key_players": ["player statistics", "players", "goals", "assists", "key passes", "játékosok", "gólok", "gólpasszok", "kulcspasszok"],
}.items():
    if _topic in TACTICAL_TOPIC_TAGS_FPI:
        TACTICAL_TOPIC_TAGS_FPI[_topic]["keywords"] = list(dict.fromkeys(TACTICAL_TOPIC_TAGS_FPI[_topic].get("keywords", []) + _extra))

def _fpi_tactical_norm(x: object) -> str:
    import unicodedata as _unicodedata
    s = _unicodedata.normalize("NFKD", str(x or "").lower().replace("\\u00ad", " "))
    s = "".join(ch for ch in s if not _unicodedata.combining(ch))
    return re.sub(r"\\s+", " ", s).strip()
def _fpi_tactical_score(col: object, aliases: List[str], field: str = "") -> int:
    src = _fpi_tactical_norm(col)
    if not src:
        return 0
    best = 0
    for alias in aliases + [field]:
        a = _fpi_tactical_norm(alias)
        if not a:
            continue
        if src == a:
            best = max(best, 100)
        elif src.replace(" ", "") == a.replace(" ", ""):
            best = max(best, 96)
        elif a in src or src in a:
            best = max(best, 86 if len(a) >= 4 else 70)
        else:
            stoks, atoks = set(src.split()), set(a.split())
            if stoks and atoks:
                best = max(best, int(len(stoks & atoks) / max(1, len(atoks)) * 78))
    return int(best)

def _fpi_tactical_suggest_mapping(df: pd.DataFrame, aliases: Dict[str, List[str]]) -> Dict[str, Optional[str]]:
    mapping, used = {}, set()
    for field, als in aliases.items():
        scored = sorted([(_fpi_tactical_score(c, als, field), c) for c in df.columns if c not in used], reverse=True, key=lambda x: x[0])
        if scored and scored[0][0] >= 58:
            mapping[field] = scored[0][1]
            used.add(scored[0][1])
        else:
            mapping[field] = None
    return mapping

def _fpi_tactical_detect_header(raw: pd.DataFrame, aliases: Dict[str, List[str]]) -> int:
    all_aliases = [a for als in aliases.values() for a in als]
    best_i, best = 0, -999
    for i in range(min(40, len(raw))):
        cells = [str(v).strip() for v in raw.iloc[i].tolist() if str(v).strip().lower() not in ["", "nan", "none"]]
        if not cells:
            continue
        joined = _fpi_tactical_norm(" | ".join(cells))
        score = sum(4 for a in all_aliases if _fpi_tactical_norm(a) and _fpi_tactical_norm(a) in joined)
        numeric_like = sum(1 for c in cells if re.fullmatch(r"-?\d+(?:[.,]\d+)?%?", c))
        text_like = len(cells) - numeric_like
        if text_like >= 3:
            score += 8
        if numeric_like > text_like:
            score -= 8
        if score > best:
            best_i, best = i, score
    return best_i

def _fpi_tactical_read_best_excel(file_bytes: bytes, aliases: Dict[str, List[str]]) -> Tuple[pd.DataFrame, str, List[dict]]:
    debug = []
    try:
        xls = pd.ExcelFile(io.BytesIO(file_bytes))
    except Exception as e:
        st.warning(f"Taktikai Excel nem olvasható: {e}")
        return pd.DataFrame(), "", debug
    best_df, best_sheet, best_score = pd.DataFrame(), "", -999
    for sheet in xls.sheet_names:
        try:
            raw = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet, header=None)
        except Exception:
            continue
        if raw.empty:
            continue
        h = _fpi_tactical_detect_header(raw, aliases)
        cols = [str(x).strip() if str(x).strip().lower() not in ["", "nan", "none"] else f"col_{j+1}" for j, x in enumerate(raw.iloc[h].tolist())]
        df2 = raw.iloc[h + 1:].copy()
        df2.columns = cols
        df2 = df2.dropna(how="all")
        score = 0
        for field, als in aliases.items():
            if max([_fpi_tactical_score(c, als, field) for c in df2.columns] or [0]) >= 58:
                score += 10
        if "main statistics" in _fpi_tactical_norm(sheet):
            score += 20
        if score > best_score:
            best_df, best_sheet, best_score = df2, sheet, score
        debug.append({"sheet": sheet, "header_row": h, "score": score, "columns": cols[:25]})
    return best_df, best_sheet, debug

def _fpi_tactical_mapper_ui(uploaded_file, aliases: Dict[str, List[str]], state_prefix: str, title: str) -> Tuple[pd.DataFrame, Dict[str, Optional[str]]]:
    if uploaded_file is None:
        return pd.DataFrame(), {}
    file_bytes = uploaded_file.getvalue()
    df2, sheet_name, debug = _fpi_tactical_read_best_excel(file_bytes, aliases)
    if df2.empty:
        return df2, {}
    signature = hashlib.md5(file_bytes).hexdigest()[:12]
    map_key = f"{state_prefix}_tactical_mapping"
    sig_key = f"{state_prefix}_tactical_sig"
    if st.session_state.get(sig_key) != signature:
        st.session_state[sig_key] = signature
        st.session_state[map_key] = _fpi_tactical_suggest_mapping(df2, aliases)
    mapping = dict(st.session_state.get(map_key, _fpi_tactical_suggest_mapping(df2, aliases)))
    cols = [""] + sorted([str(c) for c in df2.columns], key=lambda x: x.lower())
    with st.expander(f"🧭 {title} – Smart Tactical Mapper", expanded=True):
        st.caption(f"Felismert munkalap: {sheet_name or 'n.a.'}")
        grid = st.columns(3)
        for i, field in enumerate(aliases.keys()):
            default = mapping.get(field) or ""
            with grid[i % 3]:
                mapping[field] = st.selectbox(field, cols, index=cols.index(default) if default in cols else 0, key=f"{state_prefix}_{field}_{signature}") or None
        st.session_state[map_key] = mapping
        q = [{"Mező": k, "Oszlop": v or "", "Bizonyosság": _fpi_tactical_score(v, aliases[k], k) if v else 0} for k, v in mapping.items()]
        st.dataframe(pd.DataFrame(q), use_container_width=True)
        if st.checkbox("Előnézet", key=f"{state_prefix}_preview_{signature}"):
            st.dataframe(df2.head(15), use_container_width=True)
        if st.checkbox("Munkalap diagnosztika", key=f"{state_prefix}_debug_{signature}"):
            st.json(debug)
    return df2, mapping

def _fpi_tactical_numeric(s: pd.Series) -> pd.Series:
    return pd.to_numeric(s.astype(str).str.replace(",", ".", regex=False).str.replace("%", "", regex=False), errors="coerce")

def _fpi_tactical_parse_team_excel(df2: pd.DataFrame, mapping: Dict[str, Optional[str]]) -> Dict[str, float]:
    metrics = {}
    if df2 is None or df2.empty:
        return metrics
    total_row = None
    try:
        total_mask = df2.iloc[:, 0].astype(str).str.strip().str.lower().eq("total")
        if total_mask.any():
            total_row = df2[total_mask].iloc[0]
    except Exception:
        pass
    avg_fields = {"possession_pct", "pressing_success_pct", "passes_accurate_pct", "ppda", "xg"}
    for field, col in mapping.items():
        if not col or col not in df2.columns:
            continue
        if total_row is not None:
            val = coerce_cell_value(total_row.get(col)) if "coerce_cell_value" in globals() else total_row.get(col)
            metrics[field] = float(val) if isinstance(val, (int, float, np.number)) else float(pd.to_numeric(str(val).replace(",", ".").replace("%", ""), errors="coerce") or 0)
        else:
            nums = _fpi_tactical_numeric(df2[col]).dropna()
            if nums.empty:
                continue
            metrics[field] = float(nums.mean() if field in avg_fields else nums.sum())
    return metrics

def _fpi_tactical_parse_player_excel(df2: pd.DataFrame, mapping: Dict[str, Optional[str]]) -> Dict[str, pd.DataFrame]:
    empty = {"creators": pd.DataFrame(), "progressors": pd.DataFrame(), "build_up": pd.DataFrame(), "defenders": pd.DataFrame(), "duel_players": pd.DataFrame()}
    if df2 is None or df2.empty:
        return empty
    out = pd.DataFrame()
    for field, col in mapping.items():
        if col and col in df2.columns:
            out[field] = df2[col]
    if "player" not in out.columns:
        return empty
    if "minutes_played" not in out.columns:
        out["minutes_played"] = 999
    out["player"] = out["player"].astype(str).str.strip()
    out["minutes_played"] = _fpi_tactical_numeric(out["minutes_played"]).fillna(0)
    out = out[out["minutes_played"] >= 1].copy()
    for c in ["passes", "progressive_passes", "key_passes", "interceptions", "defensive_challenges", "shots", "xg", "assists", "recoveries", "crosses"]:
        out[c] = _fpi_tactical_numeric(out[c]).fillna(0) if c in out.columns else 0
    if "position" not in out.columns:
        out["position"] = ""
    return {
        "creators": out.sort_values("key_passes", ascending=False)[["player", "position", "key_passes"]].head(5).reset_index(drop=True),
        "progressors": out.sort_values("progressive_passes", ascending=False)[["player", "position", "progressive_passes"]].head(5).reset_index(drop=True),
        "build_up": out.sort_values("passes", ascending=False)[["player", "position", "passes"]].head(5).reset_index(drop=True),
        "defenders": out.sort_values("interceptions", ascending=False)[["player", "position", "interceptions"]].head(5).reset_index(drop=True),
        "duel_players": out.sort_values("defensive_challenges", ascending=False)[["player", "position", "defensive_challenges"]].head(5).reset_index(drop=True),
    }

def _fpi_tactical_page_text_pdfplumber_v78(page) -> str:
    """pdfplumber oldalszöveg – text + layout + words + tables."""
    chunks = []
    for kwargs in [
        {"x_tolerance": 1, "y_tolerance": 3},
        {"x_tolerance": 2, "y_tolerance": 4},
        {"layout": True, "x_tolerance": 1, "y_tolerance": 3},
    ]:
        try:
            txt = page.extract_text(**kwargs) or ""
            if txt.strip():
                chunks.append(txt)
        except Exception:
            pass

    try:
        words = page.extract_words(use_text_flow=True, keep_blank_chars=False) or []
        word_text = " ".join(str(w.get("text", "")).strip() for w in words if str(w.get("text", "")).strip())
        if word_text.strip():
            chunks.append(word_text)
    except Exception:
        try:
            words = page.extract_words() or []
            word_text = " ".join(str(w.get("text", "")).strip() for w in words if str(w.get("text", "")).strip())
            if word_text.strip():
                chunks.append(word_text)
        except Exception:
            pass

    try:
        tables = page.extract_tables() or []
        for tbl in tables:
            for row in tbl:
                vals = [str(c).strip() for c in (row or []) if str(c).strip() and str(c).strip().lower() not in ["none", "nan"]]
                if vals:
                    chunks.append(" | ".join(vals))
    except Exception:
        pass

    return _fpi_clean_pdf_text_v78("\n".join(chunks))


def _fpi_clean_pdf_text_v78(txt: str) -> str:
    txt = str(txt or "")
    txt = txt.replace("\ufb01", "fi").replace("\ufb02", "fl")
    txt = txt.replace("‐", "-").replace("‑", "-").replace("–", "-").replace("—", "-")
    txt = txt.replace("\x00", " ")
    txt = re.sub(r"[ \t]+", " ", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return txt.strip()


def _fpi_tactical_extract_with_pdfplumber_v78(file_bytes: bytes, fname: str, max_pages: int) -> Tuple[str, List[dict]]:
    pages, texts = [], []
    if pdfplumber is None:
        return "", pages
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages[:max_pages]):
                txt = _fpi_tactical_page_text_pdfplumber_v78(page)
                pages.append({
                    "file": fname, "page": i + 1, "reader": "pdfplumber",
                    "chars": len(txt or ""), "has_text": bool((txt or "").strip()),
                    "text": txt[:3000],
                })
                if txt.strip():
                    texts.append(f"[{fname} / oldal {i+1} / pdfplumber]\n{txt}")
    except Exception as e:
        pages.append({"file": fname, "page": None, "reader": "pdfplumber", "chars": 0, "has_text": False, "error": str(e), "text": ""})
    return "\n\n".join(texts), pages


def _fpi_tactical_extract_with_pymupdf_v78(file_bytes: bytes, fname: str, max_pages: int) -> Tuple[str, List[dict]]:
    pages, texts = [], []
    if not PYMUPDF_AVAILABLE or fitz is None:
        return "", pages
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for i in range(min(len(doc), max_pages)):
            page = doc[i]
            chunks = []
            for mode in ["text", "blocks", "words"]:
                try:
                    if mode == "text":
                        t = page.get_text("text") or ""
                    elif mode == "blocks":
                        blocks = page.get_text("blocks") or []
                        t = "\n".join(str(b[4]).strip() for b in blocks if len(b) >= 5 and str(b[4]).strip())
                    else:
                        words = page.get_text("words") or []
                        t = " ".join(str(w[4]).strip() for w in words if len(w) >= 5 and str(w[4]).strip())
                    if t.strip():
                        chunks.append(t)
                except Exception:
                    pass
            txt = _fpi_clean_pdf_text_v78("\n".join(chunks))
            pages.append({
                "file": fname, "page": i + 1, "reader": "pymupdf",
                "chars": len(txt or ""), "has_text": bool((txt or "").strip()),
                "text": txt[:3000],
            })
            if txt.strip():
                texts.append(f"[{fname} / oldal {i+1} / pymupdf]\n{txt}")
        doc.close()
    except Exception as e:
        pages.append({"file": fname, "page": None, "reader": "pymupdf", "chars": 0, "has_text": False, "error": str(e), "text": ""})
    return "\n\n".join(texts), pages


def _fpi_tactical_extract_with_pypdf_v78(file_bytes: bytes, fname: str, max_pages: int) -> Tuple[str, List[dict]]:
    pages, texts = [], []
    if not PYPDF_AVAILABLE or PdfReader is None:
        return "", pages
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        for i, page in enumerate(reader.pages[:max_pages]):
            try:
                txt = _fpi_clean_pdf_text_v78(page.extract_text() or "")
            except Exception:
                txt = ""
            pages.append({
                "file": fname, "page": i + 1, "reader": "pypdf",
                "chars": len(txt or ""), "has_text": bool((txt or "").strip()),
                "text": txt[:3000],
            })
            if txt.strip():
                texts.append(f"[{fname} / oldal {i+1} / pypdf]\n{txt}")
    except Exception as e:
        pages.append({"file": fname, "page": None, "reader": "pypdf", "chars": 0, "has_text": False, "error": str(e), "text": ""})
    return "\n\n".join(texts), pages


def _fpi_tactical_extract_pdf_text(files: List[object], max_pages: int = 120) -> Tuple[str, List[dict]]:
    """V7.8 több PDF-olvasós taktikai motor.

    Sorrend:
    1. pdfplumber text/layout/words/tables
    2. PyMuPDF text/blocks/words
    3. pypdf / PyPDF2 fallback

    A végén mindig a legtöbb karaktert adó readert választja, de a diagnosztikában mindhárom látszik.
    """
    all_pages, chosen_texts = [], []
    for f in files or []:
        if f is None:
            continue
        fname = getattr(f, "name", "pdf")
        try:
            file_bytes = f.getvalue()
        except Exception:
            continue

        attempts = []
        for reader_name, fn in [
            ("pdfplumber", _fpi_tactical_extract_with_pdfplumber_v78),
            ("pymupdf", _fpi_tactical_extract_with_pymupdf_v78),
            ("pypdf", _fpi_tactical_extract_with_pypdf_v78),
        ]:
            txt, pages = fn(file_bytes, fname, max_pages)
            attempts.append((reader_name, txt, pages, len(txt or "")))
            all_pages.extend(pages)

        best_reader, best_text, best_pages, best_chars = max(attempts, key=lambda x: x[3]) if attempts else ("none", "", [], 0)
        all_pages.append({
            "file": fname,
            "page": "BEST",
            "reader": best_reader,
            "chars": best_chars,
            "has_text": best_chars > 0,
            "text": (best_text or "")[:3000],
        })
        if best_text.strip():
            chosen_texts.append(best_text)

    full_text = _fpi_clean_pdf_text_v78("\n\n".join(chosen_texts))
    return full_text, all_pages

def _fpi_tactical_split_units_v76(text: str) -> List[str]:
    raw = str(text or "").replace("\r", "\n")
    lines = []
    for block in raw.splitlines():
        block = re.sub(r"\s+", " ", block).strip()
        if not block:
            continue
        # Ha egy teljes oldal egy sorban jön ki, mondatokra / KPI szeletekre vágjuk.
        if len(block) > 280:
            parts = re.split(r"(?<=[.!?])\s+| \| | • |; ", block)
            lines.extend([p.strip() for p in parts if len(p.strip()) >= 4])
        else:
            lines.append(block)
    return lines

def _fpi_tactical_context_lines(text: str, topic: str, limit: int = 10) -> List[str]:
    cfg = TACTICAL_TOPIC_TAGS_FPI.get(topic, {})
    kws = [_fpi_tactical_norm(k) for k in cfg.get("keywords", [])]
    units = _fpi_tactical_split_units_v76(text)
    out = []
    for i, line in enumerate(units):
        ln = _fpi_tactical_norm(line)
        if any(k and k in ln for k in kws):
            for j in range(max(0, i - 1), min(len(units), i + 2)):
                candidate = units[j].strip()
                if len(candidate) >= 4:
                    out.append(candidate[:900])
        if len(dict.fromkeys(out)) >= limit:
            break
    return list(dict.fromkeys(out))[:limit]


# =========================================================
# V8.4 - LEGACY TACTICAL PDF ENGINE 1:1 CORE
# A korábbi külön Tactical app PDF -> insight motorjának prefixelve visszaemelt változata.
# Nem UI-t emelünk át, hanem a témadetektáló / briefing blokk logikát.
# =========================================================

FPI_LEGACY_TACTICAL_TOPIC_TAGS = {
    "formation": {
        "label_hu": "Formáció / alapfelállás",
        "keywords": [
            "formation", "shape", "system", "line-up", "lineup", "starting eleven", "starting xi", "structure",
            "formáció", "felállás", "alapfelállás", "játékrendszer", "szerkezet", "kezdőcsapat", "kezdő tizenegy",
            "4-4-2", "4-2-3-1", "4-3-3", "3-5-2", "3-4-3", "5-3-2", "5-4-1",
        ],
    },
    "build_up": {
        "label_hu": "Labdakihozatal / támadásépítés",
        "keywords": [
            "build-up", "build up", "buildout", "first phase", "second phase", "goal kick", "short goal kick",
            "deep build", "progression", "progressive pass", "progressive passes", "progressive carry", "third man",
            "centre back", "center back", "fullback", "pivot", "six", "number 6", "half-space", "switch of play",
            "labdakihozatal", "támadásépítés", "építkezés", "első fázis", "második fázis", "kirúgás",
            "rövid kirúgás", "progresszió", "progresszív passz", "progresszív labdavezetés", "harmadik ember",
            "belső védő", "szélső védő", "hatossal", "6-os", "félterület", "oldalváltás",
        ],
    },
    "direct_play": {
        "label_hu": "Direkt játék / hosszú labda",
        "keywords": [
            "direct play", "long ball", "long pass", "second ball", "aerial duel", "target man", "flick-on",
            "vertical", "early forward", "long distribution", "direct attack", "route one",
            "direkt játék", "hosszú labda", "hosszú passz", "második labda", "felívelés", "fejpárbaj",
            "céljátékos", "lecsorgó", "vertikális", "korai előrejáték", "direkt támadás",
        ],
    },
    "pressing": {
        "label_hu": "Letámadás / presszing",
        "keywords": [
            "press", "pressing", "high press", "mid press", "low press", "counterpress", "counter-press",
            "ppda", "pressure", "pressing trigger", "trap", "press trap", "forced turnover", "high recovery",
            "challenge intensity", "defensive actions", "intensity", "aggressive press",
            "letámadás", "presszing", "magas letámadás", "középső presszing", "visszatámadás", "nyomás",
            "presszing trigger", "trigger", "csapda", "pressingcsapda", "kikényszerített labdavesztés",
            "magas labdaszerzés", "védekező akció", "intenzitás", "agresszív letámadás",
        ],
    },
    "defensive_block": {
        "label_hu": "Védekezési blokk / blokkmagasság",
        "keywords": [
            "low block", "mid block", "middle block", "high block", "defensive block", "compact", "compactness",
            "defensive line", "back line", "line height", "block height", "deep defending", "drop", "retreat",
            "mély blokk", "középső blokk", "magas blokk", "védekezési blokk", "kompakt", "kompaktság",
            "védelmi vonal", "védősor", "blokkmagasság", "mély védekezés", "visszazár", "visszarendeződés",
        ],
    },
    "transition_attack": {
        "label_hu": "Támadó átmenet / kontrák",
        "keywords": [
            "transition", "attacking transition", "offensive transition", "counterattack", "counter attack", "counter-attacks",
            "fast attack", "quick attack", "break", "breakaway", "after regain", "after winning", "regain and go",
            "átmenet", "támadó átmenet", "kontra", "kontratámadás", "gyors támadás", "gyors átmenet",
            "labdaszerzés után", "labdanyerés után", "visszaszerzés után", "indítás", "megindulás",
        ],
    },
    "transition_defense": {
        "label_hu": "Védekező átmenet / rest defense",
        "keywords": [
            "defensive transition", "after losing", "after loss", "rest defense", "counter attack prevention",
            "counter prevention", "defend transition", "negative transition", "cover behind", "protection behind",
            "védekező átmenet", "labdavesztés után", "rest defense", "kontrák elleni védekezés",
            "átmeneti védekezés", "negatív átmenet", "biztosítás", "mögöttes biztosítás", "visszarendeződés",
        ],
    },
    "chance_creation": {
        "label_hu": "Helyzetkialakítás / támadóharmad",
        "keywords": [
            "chance creation", "key pass", "shot assist", "box entry", "penalty area", "final third", "final-third",
            "through ball", "cutback", "cross", "low cross", "deep cross", "half-space cross", "xg", "expected goals",
            "helyzetkialakítás", "kulcspassz", "lövést előkészítő", "box entry", "tizenhatos", "büntetőterület",
            "támadóharmad", "mélységi passz", "visszagurítás", "beadás", "lapos beadás", "xg", "várható gól",
        ],
    },
    "wide_play": {
        "label_hu": "Szélső játék / oldali dominancia",
        "keywords": [
            "wide play", "wing", "flank", "left side", "right side", "overlap", "underlap", "fullback", "wingback",
            "crossing", "side dominance", "touchline", "wide overload", "flank overload",
            "szélső játék", "szél", "oldal", "bal oldal", "jobb oldal", "átfedés", "aláfutás", "szélső védő",
            "wingback", "beadás", "oldali dominancia", "oldalvonal", "oldali túlterhelés",
        ],
    },
    "central_play": {
        "label_hu": "Középső játék / félterületek",
        "keywords": [
            "central", "middle", "half-space", "half space", "between the lines", "pocket", "zone 14", "inside channel",
            "central overload", "interior", "attacking midfielder", "number 10",
            "középen", "középső", "félterület", "félterületek", "vonalak között", "zseb", "14-es zóna",
            "belső csatorna", "középső túlterhelés", "belső középpályás", "10-es",
        ],
    },
    "set_pieces": {
        "label_hu": "Pontrúgások",
        "keywords": [
            "set piece", "set pieces", "corner", "corner kick", "free kick", "throw-in", "throw in", "penalty",
            "attacking corner", "defensive corner", "near post", "far post", "second ball", "aerial", "header",
            "pontrúgás", "pontrúgások", "szöglet", "szabadrúgás", "bedobás", "büntető", "támadó szöglet",
            "védekező szöglet", "rövid oldal", "hosszú oldal", "második labda", "fejpárbaj", "fejes",
        ],
    },
    "key_players": {
        "label_hu": "Kulcsjátékosok",
        "keywords": [
            "key player", "danger man", "main threat", "top scorer", "creator", "playmaker", "progressor", "target man",
            "dribbler", "1v1", "one-v-one", "finisher", "captain", "most dangerous",
            "kulcsjátékos", "veszélyes játékos", "fő veszély", "gólkirály", "kreatív játékos", "irányító",
            "progresszor", "céljátékos", "cselező", "egy az egy", "befejező", "csapatkapitány", "legveszélyesebb",
        ],
    },
    "weakness_risk": {
        "label_hu": "Gyengeségek / kockázatok",
        "keywords": [
            "weakness", "weaknesses", "risk", "risks", "vulnerable", "vulnerability", "exposed", "space behind",
            "gap", "mistake", "error", "turnover", "lost balls", "losses", "danger", "threat", "problem",
            "gyengeség", "gyengeségek", "kockázat", "sebezhető", "sebezhetőség", "nyitott terület", "mögötti terület",
            "rés", "hiba", "labdavesztés", "elvesztett labda", "veszély", "fenyegetés", "probléma",
        ],
    },
    "strength": {
        "label_hu": "Erősségek",
        "keywords": [
            "strength", "strengths", "strong", "advantage", "edge", "dominant", "effective", "efficient",
            "erősség", "erősségek", "erős", "előny", "domináns", "hatékony", "kiemelkedő",
        ],
    },
    "goalkeeper": {
        "label_hu": "Kapus szerepe",
        "keywords": [
            "goalkeeper", "keeper", "gk", "sweeper keeper", "distribution", "long kick", "short pass from gk",
            "kapus", "hálóőr", "kapusjáték", "kapus kirúgás", "kapus passz", "hosszú kirúgás", "rövid kirúgás",
        ],
    },
    "match_dynamics": {
        "label_hu": "Meccsdinamika / fázisok",
        "keywords": [
            "first half", "second half", "opening phase", "late phase", "last 15", "tempo", "rhythm", "momentum",
            "game state", "when leading", "when trailing", "after goal", "minutes", "phase",
            "első félidő", "második félidő", "kezdő fázis", "végjáték", "utolsó 15", "tempó", "ritmus",
            "momentum", "meccsállapot", "vezetésnél", "hátrányban", "gól után", "percek", "fázis",
        ],
    },
    "recommendation": {
        "label_hu": "Javaslat / meccsterv",
        "keywords": [
            "recommendation", "recommend", "should", "we should", "game plan", "match plan", "plan a", "plan b",
            "solution", "exploit", "avoid", "focus", "priority", "target", "press here", "attack here",
            "javaslat", "ajánlás", "meccsterv", "mérkőzésterv", "terv a", "terv b", "megoldás",
            "kihasználni", "elkerülni", "fókusz", "prioritás", "célpont", "itt presszing", "itt támadni",
        ],
    },
}
FPI_LEGACY_TACTICAL_TOPIC_ORDER = [
    "formation", "build_up", "direct_play", "pressing", "defensive_block", "transition_attack",
    "transition_defense", "chance_creation", "wide_play", "central_play", "set_pieces", "key_players",
    "weakness_risk", "strength", "goalkeeper", "match_dynamics", "recommendation",
]

def _fpi_legacy_norm_for_tagging(text: object) -> str:
    s = unicodedata.normalize("NFKD", str(text or "").lower())
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", s).strip()

def _fpi_legacy_unique_keep_order(items: List[str]) -> List[str]:
    out = []
    seen = set()
    for x in items or []:
        key = str(x).strip()
        if key and key not in seen:
            out.append(key)
            seen.add(key)
    return out

def _fpi_legacy_tactical_keyword_hits(text: str, keywords: List[str]) -> int:
    norm = _fpi_legacy_norm_for_tagging(text)
    hits = 0
    for kw in keywords:
        k = _fpi_legacy_norm_for_tagging(kw)
        if k and k in norm:
            hits += 1
    return hits

def _fpi_legacy_extract_context_lines_by_topic(text: str, topic_key: str, limit: int = 8, context_radius: int = 1) -> List[str]:
    cfg = FPI_LEGACY_TACTICAL_TOPIC_TAGS.get(topic_key, {})
    keywords = cfg.get("keywords", [])
    lines = [x.strip() for x in str(text or "").splitlines() if x.strip()]
    selected = []
    for i, line in enumerate(lines):
        if _fpi_legacy_tactical_keyword_hits(line, keywords) > 0:
            start = max(0, i - context_radius)
            end = min(len(lines), i + context_radius + 1)
            for j in range(start, end):
                if len(lines[j]) >= 4:
                    selected.append(lines[j])
        if len(_fpi_legacy_unique_keep_order(selected)) >= limit:
            break
    return _fpi_legacy_unique_keep_order(selected)[:limit]

def _fpi_legacy_detect_tactical_topics(text: str) -> Tuple[Dict[str, dict], List[str]]:
    rows = {}
    for key in FPI_LEGACY_TACTICAL_TOPIC_ORDER:
        cfg = FPI_LEGACY_TACTICAL_TOPIC_TAGS[key]
        lines = _fpi_legacy_extract_context_lines_by_topic(text, key, limit=10, context_radius=0)
        hit_count = sum(_fpi_legacy_tactical_keyword_hits(line, cfg["keywords"]) for line in lines)
        rows[key] = {
            "label_hu": cfg["label_hu"],
            "hit_count": hit_count,
            "lines": lines[:8],
            "confidence": min(100, hit_count * 18 + len(lines) * 5),
        }
    detected = [k for k, v in rows.items() if v["hit_count"] > 0 or v["lines"]]
    detected = sorted(detected, key=lambda k: (rows[k]["confidence"], rows[k]["hit_count"]), reverse=True)
    return rows, detected

def _fpi_legacy_build_universal_briefing_blocks(text: str) -> Dict[str, List[str]]:
    topic_rows, detected = _fpi_legacy_detect_tactical_topics(text)
    return {
        "opponent_identity": _fpi_legacy_extract_context_lines_by_topic(text, "formation", limit=4, context_radius=1),
        "build_up": _fpi_legacy_extract_context_lines_by_topic(text, "build_up", limit=8, context_radius=1),
        "pressing": _fpi_legacy_extract_context_lines_by_topic(text, "pressing", limit=8, context_radius=1),
        "defensive_block": _fpi_legacy_extract_context_lines_by_topic(text, "defensive_block", limit=8, context_radius=1),
        "transition_attack": _fpi_legacy_extract_context_lines_by_topic(text, "transition_attack", limit=8, context_radius=1),
        "transition_defense": _fpi_legacy_extract_context_lines_by_topic(text, "transition_defense", limit=8, context_radius=1),
        "set_pieces": _fpi_legacy_extract_context_lines_by_topic(text, "set_pieces", limit=8, context_radius=1),
        "key_players": _fpi_legacy_extract_context_lines_by_topic(text, "key_players", limit=8, context_radius=1),
        "risks": _fpi_legacy_extract_context_lines_by_topic(text, "weakness_risk", limit=8, context_radius=1),
        "recommendations": _fpi_legacy_extract_context_lines_by_topic(text, "recommendation", limit=8, context_radius=1),
        "detected_topics": [topic_rows[k]["label_hu"] for k in detected[:10]],
    }

def _fpi_legacy_extract_lines_with_keywords(text: str, keywords: List[str], limit: int = 6) -> List[str]:
    out = []
    lines = [x.strip() for x in str(text or "").splitlines() if x.strip()]
    keyword_norms = [_fpi_legacy_norm_for_tagging(k) for k in keywords]
    for line in lines:
        line_norm = _fpi_legacy_norm_for_tagging(line)
        if any(k and k in line_norm for k in keyword_norms):
            out.append(line)
        if len(out) >= limit:
            break
    return _fpi_legacy_unique_keep_order(out)

def _fpi_legacy_infer_formation(text: str) -> Optional[str]:
    m = re.search(r"\b([3-5]-[1-5]-[1-5](?:-[1-3])?)\b", text or "")
    if m:
        return m.group(1)
    return None

def _fpi_legacy_extract_player_names_from_pdf(text: str, limit: int = 6) -> List[str]:
    names = re.findall(r"\b[A-ZÁÉÍÓÖŐÚÜŰ][a-záéíóöőúüű\-]+(?:\s+[A-ZÁÉÍÓÖŐÚÜŰ][a-záéíóöőúüű\-]+)+\b", text or "")
    return _fpi_legacy_unique_keep_order(names)[:limit]

def _fpi_legacy_build_pdf_insights(text: str) -> Dict[str, object]:
    """A külön Tactical app Universal Tactical PDF Reader motorja, prefixelve."""
    formation = _fpi_legacy_infer_formation(text)
    topic_rows, detected_topics = _fpi_legacy_detect_tactical_topics(text)
    universal_blocks = _fpi_legacy_build_universal_briefing_blocks(text)

    dna_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("build_up", [])[:3]
        + universal_blocks.get("pressing", [])[:3]
        + universal_blocks.get("defensive_block", [])[:3]
        + _fpi_legacy_extract_lines_with_keywords(
            text,
            FPI_LEGACY_TACTICAL_TOPIC_TAGS["build_up"]["keywords"]
            + FPI_LEGACY_TACTICAL_TOPIC_TAGS["pressing"]["keywords"]
            + FPI_LEGACY_TACTICAL_TOPIC_TAGS["transition_attack"]["keywords"]
            + FPI_LEGACY_TACTICAL_TOPIC_TAGS["direct_play"]["keywords"],
            limit=10,
        )
    )[:10]

    risk_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("risks", [])
        + _fpi_legacy_extract_lines_with_keywords(
            text,
            FPI_LEGACY_TACTICAL_TOPIC_TAGS["weakness_risk"]["keywords"]
            + FPI_LEGACY_TACTICAL_TOPIC_TAGS["transition_defense"]["keywords"],
            limit=10,
        )
    )[:10]

    set_piece_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("set_pieces", [])
        + _fpi_legacy_extract_lines_with_keywords(text, FPI_LEGACY_TACTICAL_TOPIC_TAGS["set_pieces"]["keywords"], limit=8)
    )[:8]

    dynamics_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("transition_attack", [])[:3]
        + universal_blocks.get("transition_defense", [])[:3]
        + _fpi_legacy_extract_lines_with_keywords(text, FPI_LEGACY_TACTICAL_TOPIC_TAGS["match_dynamics"]["keywords"], limit=8)
    )[:8]

    pressing_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("pressing", [])
        + _fpi_legacy_extract_lines_with_keywords(text, FPI_LEGACY_TACTICAL_TOPIC_TAGS["pressing"]["keywords"], limit=8)
    )[:8]

    build_up_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("build_up", [])
        + universal_blocks.get("direct_play", [])
        + _fpi_legacy_extract_lines_with_keywords(
            text,
            FPI_LEGACY_TACTICAL_TOPIC_TAGS["build_up"]["keywords"] + FPI_LEGACY_TACTICAL_TOPIC_TAGS["direct_play"]["keywords"],
            limit=8,
        )
    )[:8]

    player_threat_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("key_players", [])
        + universal_blocks.get("chance_creation", [])
        + _fpi_legacy_extract_lines_with_keywords(
            text,
            FPI_LEGACY_TACTICAL_TOPIC_TAGS["key_players"]["keywords"]
            + FPI_LEGACY_TACTICAL_TOPIC_TAGS["chance_creation"]["keywords"],
            limit=10,
        )
    )[:10]

    detected_names = _fpi_legacy_extract_player_names_from_pdf(text, limit=12)

    recommendation_lines = _fpi_legacy_unique_keep_order(
        universal_blocks.get("recommendations", [])
        + _fpi_legacy_extract_lines_with_keywords(text, FPI_LEGACY_TACTICAL_TOPIC_TAGS["recommendation"]["keywords"], limit=10)
    )[:10]

    topic_summary_rows = []
    for key in detected_topics[:12]:
        row = topic_rows[key]
        topic_summary_rows.append({
            "Téma": row["label_hu"],
            "Kulcs": key,
            "Találat": row["hit_count"],
            "Bizonyosság": row["confidence"],
            "Minta": " | ".join(row["lines"][:2]),
        })

    return {
        "formation": formation or "n.a.",
        "dna_lines": dna_lines,
        "risk_lines": risk_lines,
        "set_piece_lines": set_piece_lines,
        "dynamics_lines": dynamics_lines,
        "pressing_lines": pressing_lines,
        "build_up_lines": build_up_lines,
        "player_threat_lines": player_threat_lines,
        "recommendation_lines": recommendation_lines,
        "detected_names": detected_names,
        "universal_blocks": universal_blocks,
        "detected_topics": universal_blocks.get("detected_topics", []),
        "topic_debug": topic_summary_rows,
        "reader_version": "Legacy Tactical PDF Reader 1:1 core imported into FPI v8.4",
    }


def _fpi_tactical_top_terms_v76(text: str, limit: int = 35) -> List[dict]:
    stop = set("""
    the and for with from that this into their your are was were have has had not but
    egy meg hogy van vagy mint ahol után előtt illetve saját ellenfél csapat játékos
    és az ez ha is de már csak nem total average percent percentage page oldal
    """.split())
    words = re.findall(r"[A-Za-zÁÉÍÓÖŐÚÜŰáéíóöőúüű][A-Za-zÁÉÍÓÖŐÚÜŰáéíóöőúüű\\-]{3,}", str(text or ""))
    counts = {}
    for w in words:
        n = _fpi_tactical_norm(w)
        if n in stop or len(n) < 4:
            continue
        counts[n] = counts.get(n, 0) + 1
    return [{"szó": k, "db": v} for k, v in sorted(counts.items(), key=lambda kv: kv[1], reverse=True)[:limit]]

def _fpi_full_text_keyword_hits_v77(text: str, keywords: List[str]) -> int:
    norm = _fpi_tactical_norm(text)
    hits = 0
    for kw in keywords:
        k = _fpi_tactical_norm(kw)
        if k and k in norm:
            hits += 1
    return hits

def _fpi_find_keyword_windows_v77(text: str, keywords: List[str], limit: int = 6) -> List[str]:
    """Ha a soralapú context nem talál semmit, a teljes szövegből vág ablakot a kulcsszó köré."""
    raw = re.sub(r"\\s+", " ", str(text or "")).strip()
    norm_raw = _fpi_tactical_norm(raw)
    out = []
    for kw in keywords:
        k = _fpi_tactical_norm(kw)
        if not k:
            continue
        pos = norm_raw.find(k)
        if pos >= 0:
            # norm és raw index nem tökéletesen egyezik, de közelítő ablaknak elég
            start = max(0, pos - 180)
            end = min(len(raw), pos + 280)
            snippet = raw[start:end].strip()
            if snippet:
                out.append(snippet)
        if len(out) >= limit:
            break
    return list(dict.fromkeys(out))[:limit]

def _fpi_fallback_topics_from_terms_v77(text: str, top_terms: List[dict]) -> List[dict]:
    """Ha nincs erős kulcsszó-találat, gyakori szavakból ad óvatos, de használható fallback témákat."""
    term_text = " ".join(str(x.get("szó", "")) for x in (top_terms or []))
    combined = f"{text[:5000]} {term_text}"
    fallback_rules = [
        ("team_stats", "Csapatstatisztikák / fő KPI-k", ["statistics", "possession", "shots", "passes", "xg", "labdabirtoklás", "lövések", "passz"]),
        ("event_actions", "Események / akciótípusok", ["episodes", "passes", "recoveries", "lost", "actions", "counter", "passzok", "labdaszerzések", "labdavesztések", "akció"]),
        ("build_up", "Labdakihozatal / támadásépítés", ["passes", "possession", "progression", "build", "passz", "labdabirtoklás", "progresszió"]),
        ("chance_creation", "Helyzetkialakítás / támadóharmad", ["shots", "xg", "box", "final", "key", "cross", "lövések", "tizenhatos", "kulcspassz", "beadás"]),
        ("transition_attack", "Támadó átmenet / kontrák", ["counter", "fast", "transition", "regain", "kontra", "átmenet", "gyors"]),
        ("set_pieces", "Pontrúgások", ["corner", "corners", "set", "free", "szöglet", "pontrúgás", "szabadrúgás"]),
        ("key_players", "Kulcsjátékosok", ["player", "players", "goals", "assists", "játékos", "gól", "gólpassz"]),
    ]
    rows = []
    for key, label, kws in fallback_rules:
        h = _fpi_full_text_keyword_hits_v77(combined, kws)
        if h > 0:
            rows.append({
                "Téma": label,
                "Kulcs": key,
                "Találat": h,
                "Bizonyosság": min(55, 25 + h * 7),
                "Minta": "Fallback: a PDF szövegében gyakori KPI/akció kifejezések alapján.",
                "Fallback": True,
            })
    return sorted(rows, key=lambda x: x["Bizonyosság"], reverse=True)


# =========================================================
# V8.5 - PROVIDER-AWARE PDF EXTRACTION
# SportsBase kompatibilis, de általános fallbackekkel is működik.
# A cél: ne csak témaszavakat találjon, hanem konkrét taktikai adatpontokat is.
# =========================================================

def _fpi_float_v85(x: object, default: float = 0.0) -> float:
    try:
        s = str(x).replace(",", ".").replace("%", "").strip()
        if not s or s in ["—", "-", "None", "nan"]:
            return default
        return float(s)
    except Exception:
        return default

def _fpi_int_v85(x: object, default: int = 0) -> int:
    try:
        s = str(x).replace(",", ".").strip()
        if not s or s in ["—", "-", "None", "nan"]:
            return default
        return int(float(s))
    except Exception:
        return default

def _fpi_ratio_pair_v85(text: str, label_regex: str) -> Optional[Tuple[int, int, int, int]]:
    """Két csapat értékeit olvassa ilyen mintából: Label 9 / 2 9 / 4."""
    try:
        m = re.search(label_regex + r"\s+(\d+)\s*/\s*(\d+)\s+(\d+)\s*/\s*(\d+)", text, flags=re.I)
        if m:
            return tuple(_fpi_int_v85(g) for g in m.groups())
    except Exception:
        pass
    return None

def _fpi_extract_match_teams_v85(text: str) -> Dict[str, str]:
    # SportsBase első oldalak: "Ajka 1:2 Csakvari TK"
    m = re.search(r"\n([A-Za-zÁÉÍÓÖŐÚÜŰáéíóöőúüű .'-]+)\s+\d+\s*:\s*\d+\s+([A-Za-zÁÉÍÓÖŐÚÜŰáéíóöőúüű .'-]+)\n", "\n" + str(text or "") + "\n")
    if m:
        return _fpi_mark_own_opponent_teams_v95({"team_a": m.group(1).strip(), "team_b": m.group(2).strip()})
    # fallback: Match report line
    m = re.search(r"MATCH REPORT\s+(.+?)\s+(.+?)\s+Hungary", str(text or ""), flags=re.I | re.S)
    if m:
        return _fpi_mark_own_opponent_teams_v95({"team_a": re.sub(r"\s+", " ", m.group(1)).strip(), "team_b": re.sub(r"\s+", " ", m.group(2)).strip()})
    return _fpi_mark_own_opponent_teams_v95({"team_a": "KTE / Kecskemét", "team_b": "Ellenfél"})

def _fpi_extract_formations_v85(text: str) -> Dict[str, str]:
    forms = re.findall(r"Match start\s+([3-5][–\-\u2013][1-5][–\-\u2013][1-5](?:[–\-\u2013][1-3])?)", str(text or ""), flags=re.I)
    forms = [f.replace("–", "-").replace("—", "-") for f in forms]
    out = {}
    if len(forms) >= 1:
        out["team_a_start"] = forms[0]
    if len(forms) >= 2:
        out["team_b_start"] = forms[1]
    end_forms = re.findall(r"Match end\s+([3-5][–\-\u2013][1-5][–\-\u2013][1-5](?:[–\-\u2013][1-3])?)", str(text or ""), flags=re.I)
    end_forms = [f.replace("–", "-").replace("—", "-") for f in end_forms]
    if len(end_forms) >= 1:
        out["team_a_end"] = end_forms[0]
    if len(end_forms) >= 2:
        out["team_b_end"] = end_forms[1]
    return out

def _fpi_extract_sportsbase_team_stats_v85(text: str) -> Dict[str, object]:
    """SportsBase Team statistics / Match info jellegű PDF-ekből konkrét KPI-k."""
    raw = str(text or "")
    flat = re.sub(r"\s+", " ", raw)
    teams = _fpi_extract_match_teams_v85(raw)
    formations = _fpi_extract_formations_v85(raw)
    metrics = {"teams": teams, "formations": formations, "provider": "SportsBase-like"}

    # Match info indexes
    m = re.search(r"([\d.]+)\s*/\s*([\d.]+)\s+xG\s*/\s*per shot\s+([\d.]+)\s*/\s*([\d.]+)", flat, flags=re.I)
    if m:
        metrics["xg_a"] = _fpi_float_v85(m.group(1))
        metrics["xg_per_shot_a"] = _fpi_float_v85(m.group(2))
        metrics["xg_b"] = _fpi_float_v85(m.group(3))
        metrics["xg_per_shot_b"] = _fpi_float_v85(m.group(4))

    m = re.search(r"([\d.]+)\s+PPDA\s+([\d.]+)", flat, flags=re.I)
    if m:
        metrics["ppda_a"] = _fpi_float_v85(m.group(1))
        metrics["ppda_b"] = _fpi_float_v85(m.group(2))

    m = re.search(r"([\d.]+)\s*/\s*([\d.]+)\s+Speed of passes\s*/\s*accurate\s+([\d.]+)\s*/\s*([\d.]+)", flat, flags=re.I)
    if m:
        metrics["pass_speed_a"] = _fpi_float_v85(m.group(1))
        metrics["accurate_pass_speed_a"] = _fpi_float_v85(m.group(2))
        metrics["pass_speed_b"] = _fpi_float_v85(m.group(3))
        metrics["accurate_pass_speed_b"] = _fpi_float_v85(m.group(4))

    # Team statistics
    pair = _fpi_ratio_pair_v85(flat, r"Shots\s*/\s*on target")
    if pair:
        metrics["shots_a"], metrics["shots_on_target_a"], metrics["shots_b"], metrics["shots_on_target_b"] = pair

    m = re.search(r"\bxG\s+([\d.]+)\s+([\d.]+)", flat, flags=re.I)
    if m:
        # team statistics page can overwrite match info with same xG; okay
        metrics["xg_a"] = _fpi_float_v85(m.group(1))
        metrics["xg_b"] = _fpi_float_v85(m.group(2))

    m = re.search(r"Ball possession\s+(\d+)%\s+(\d+)%", flat, flags=re.I)
    if m:
        metrics["possession_a"] = _fpi_float_v85(m.group(1))
        metrics["possession_b"] = _fpi_float_v85(m.group(2))

    pair = _fpi_ratio_pair_v85(flat, r"Key passes")
    if pair:
        metrics["key_passes_a"], metrics["key_passes_acc_a"], metrics["key_passes_b"], metrics["key_passes_acc_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Passes\s*/\s*accurate")
    if pair:
        metrics["passes_a"], metrics["passes_acc_a"], metrics["passes_b"], metrics["passes_acc_b"] = pair
        metrics["pass_accuracy_a"] = round(metrics["passes_acc_a"] / max(metrics["passes_a"], 1) * 100, 1)
        metrics["pass_accuracy_b"] = round(metrics["passes_acc_b"] / max(metrics["passes_b"], 1) * 100, 1)

    pair = _fpi_ratio_pair_v85(flat, r"Passes into the penalty box")
    if pair:
        metrics["box_passes_a"], metrics["box_passes_acc_a"], metrics["box_passes_b"], metrics["box_passes_acc_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Crosses")
    if pair:
        metrics["crosses_a"], metrics["crosses_acc_a"], metrics["crosses_b"], metrics["crosses_acc_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Passes into the final third of the pitch")
    if pair:
        metrics["final_third_passes_a"], metrics["final_third_passes_acc_a"], metrics["final_third_passes_b"], metrics["final_third_passes_acc_b"] = pair

    # Attacks block
    pair = _fpi_ratio_pair_v85(flat, r"Attacks\s*/\s*with shots")
    if pair:
        metrics["attacks_a"], metrics["attacks_with_shots_a"], metrics["attacks_b"], metrics["attacks_with_shots_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Positional attacks\s*/\s*with shots")
    if pair:
        metrics["pos_attacks_a"], metrics["pos_attacks_shots_a"], metrics["pos_attacks_b"], metrics["pos_attacks_shots_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Counter-attacks\s*/\s*with shots")
    if pair:
        metrics["counter_attacks_a"], metrics["counter_attacks_shots_a"], metrics["counter_attacks_b"], metrics["counter_attacks_shots_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Set-piece attacks\s*/\s*with shots")
    if pair:
        metrics["set_piece_attacks_a"], metrics["set_piece_attacks_shots_a"], metrics["set_piece_attacks_b"], metrics["set_piece_attacks_shots_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Corners\s*/\s*with shots")
    if pair:
        metrics["corner_attacks_a"], metrics["corner_attacks_shots_a"], metrics["corner_attacks_b"], metrics["corner_attacks_shots_b"] = pair

    # Lost balls / recoveries
    pair = _fpi_ratio_pair_v85(flat, r"Lost balls\s*/\s*in own half")
    if pair:
        metrics["lost_balls_a"], metrics["lost_own_half_a"], metrics["lost_balls_b"], metrics["lost_own_half_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Ball recoveries\s*/\s*in opp\. half")
    if pair:
        metrics["recoveries_a"], metrics["recoveries_opp_half_a"], metrics["recoveries_b"], metrics["recoveries_opp_half_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Pressing\s*/\s*successful")
    if pair:
        metrics["pressing_a"], metrics["pressing_success_a"], metrics["pressing_b"], metrics["pressing_success_b"] = pair
        metrics["pressing_success_pct_a"] = round(metrics["pressing_success_a"] / max(metrics["pressing_a"], 1) * 100, 1)
        metrics["pressing_success_pct_b"] = round(metrics["pressing_success_b"] / max(metrics["pressing_b"], 1) * 100, 1)

    pair = _fpi_ratio_pair_v85(flat, r"High pressing\s*/\s*successful")
    if pair:
        metrics["high_press_a"], metrics["high_press_success_a"], metrics["high_press_b"], metrics["high_press_success_b"] = pair

    pair = _fpi_ratio_pair_v85(flat, r"Low pressing\s*/\s*successful")
    if pair:
        metrics["low_press_a"], metrics["low_press_success_a"], metrics["low_press_b"], metrics["low_press_success_b"] = pair

    return metrics

def _fpi_sportsbase_metric_lines_v85(metrics: Dict[str, object]) -> List[str]:
    if not metrics or len(metrics.keys()) <= 3:
        return []
    ta = (metrics.get("teams") or {}).get("team_a", "A csapat")
    tb = (metrics.get("teams") or {}).get("team_b", "B csapat")
    lines = []
    if metrics.get("formations"):
        fa = metrics["formations"].get("team_a_start")
        fb = metrics["formations"].get("team_b_start")
        if fa or fb:
            lines.append(f"Formáció: {ta} {fa or 'n.a.'}, {tb} {fb or 'n.a.'}.")
    if metrics.get("xg_a") or metrics.get("xg_b"):
        better = tb if metrics.get("xg_b", 0) > metrics.get("xg_a", 0) else ta
        lines.append(f"xG: {ta} {metrics.get('xg_a', 0):.2f}, {tb} {metrics.get('xg_b', 0):.2f}; helyzetminőségben {better} állt jobban.")
    if metrics.get("shots_a") or metrics.get("shots_b"):
        lines.append(f"Lövések: {ta} {metrics.get('shots_a', 0)}/{metrics.get('shots_on_target_a', 0)} kaput találó, {tb} {metrics.get('shots_b', 0)}/{metrics.get('shots_on_target_b', 0)} kaput találó.")
    if metrics.get("possession_a") or metrics.get("possession_b"):
        poss_team = ta if metrics.get("possession_a", 0) > metrics.get("possession_b", 0) else tb
        lines.append(f"Labdabirtoklás: {ta} {metrics.get('possession_a', 0):.0f}%, {tb} {metrics.get('possession_b', 0):.0f}%; kontrollban {poss_team} volt fölényben.")
    if metrics.get("ppda_a") or metrics.get("ppda_b"):
        press_team = ta if metrics.get("ppda_a", 99) < metrics.get("ppda_b", 99) else tb
        lines.append(f"PPDA: {ta} {metrics.get('ppda_a', 0):.2f}, {tb} {metrics.get('ppda_b', 0):.2f}; presszingintenzitásban {press_team} aktívabb.")
    if metrics.get("counter_attacks_a") or metrics.get("counter_attacks_b"):
        lines.append(f"Kontrák: {ta} {metrics.get('counter_attacks_a', 0)}/{metrics.get('counter_attacks_shots_a', 0)} lövéssel, {tb} {metrics.get('counter_attacks_b', 0)}/{metrics.get('counter_attacks_shots_b', 0)} lövéssel.")
    if metrics.get("set_piece_attacks_a") or metrics.get("set_piece_attacks_b"):
        lines.append(f"Pontrúgásos támadások: {ta} {metrics.get('set_piece_attacks_a', 0)}/{metrics.get('set_piece_attacks_shots_a', 0)} lövéssel, {tb} {metrics.get('set_piece_attacks_b', 0)}/{metrics.get('set_piece_attacks_shots_b', 0)} lövéssel.")
    if metrics.get("key_passes_a") or metrics.get("key_passes_b"):
        key_team = ta if metrics.get("key_passes_a", 0) > metrics.get("key_passes_b", 0) else tb
        lines.append(f"Kulcspasszok: {ta} {metrics.get('key_passes_a', 0)}, {tb} {metrics.get('key_passes_b', 0)}; kreatív előny: {key_team}.")
    if metrics.get("crosses_a") or metrics.get("crosses_b"):
        wide_team = ta if metrics.get("crosses_a", 0) > metrics.get("crosses_b", 0) else tb
        lines.append(f"Beadások: {ta} {metrics.get('crosses_a', 0)}, {tb} {metrics.get('crosses_b', 0)}; szélső játékban {wide_team} aktívabb.")
    if metrics.get("lost_balls_a") or metrics.get("lost_balls_b"):
        risk_team = ta if metrics.get("lost_own_half_a", 0) > metrics.get("lost_own_half_b", 0) else tb
        lines.append(f"Labdavesztés saját térfélen: {ta} {metrics.get('lost_own_half_a', 0)}, {tb} {metrics.get('lost_own_half_b', 0)}; átmeneti kockázat: {risk_team}.")
    if metrics.get("pressing_a") or metrics.get("pressing_b"):
        lines.append(f"Presszing: {ta} {metrics.get('pressing_a', 0)}/{metrics.get('pressing_success_a', 0)} ({metrics.get('pressing_success_pct_a', 0):.0f}%), {tb} {metrics.get('pressing_b', 0)}/{metrics.get('pressing_success_b', 0)} ({metrics.get('pressing_success_pct_b', 0):.0f}%).")
    return lines


def _fpi_find_first_pair_float_v86(flat: str, label: str) -> Tuple[float, float]:
    m = re.search(label + r"\s+([\d.,]+)\s+([\d.,]+)", flat, flags=re.I)
    if not m:
        return 0.0, 0.0
    return _fpi_float_v85(m.group(1)), _fpi_float_v85(m.group(2))

def _fpi_find_first_pair_percent_v86(flat: str, label: str) -> Tuple[float, float]:
    m = re.search(label + r"\s+(\d+)\s*%\s+(\d+)\s*%", flat, flags=re.I)
    if not m:
        return 0.0, 0.0
    return _fpi_float_v85(m.group(1)), _fpi_float_v85(m.group(2))

def _fpi_find_ratio_pair_v86(flat: str, label: str) -> Tuple[int, int, int, int]:
    m = re.search(label + r"\s+(\d+)\s*/\s*(\d+)\s+(\d+)\s*/\s*(\d+)", flat, flags=re.I)
    if not m:
        return 0, 0, 0, 0
    return tuple(_fpi_int_v85(x) for x in m.groups())

def _fpi_direct_pdf_extract_v86(text: str, role: str = "PDF") -> Dict[str, object]:
    """Brutálisan direkt, report-provider-független fallback.
    Nem a bonyolult topic engine-re vár, hanem a teljes PDF szövegben keres konkrét futball KPI mintákat.
    SportsBase-en nagyon erős, más riportokon pedig általános KPI fallbackként működik.
    """
    raw = str(text or "")
    flat = re.sub(r"\s+", " ", raw)
    teams = _fpi_extract_match_teams_v85(raw)
    forms = _fpi_extract_formations_v85(raw)
    ta = teams.get("team_a", "A csapat")
    tb = teams.get("team_b", "B csapat")

    data = {"role": role, "teams": teams, "formations": forms, "lines": [], "findings": []}
    def add_line(s):
        if s and s not in data["lines"]:
            data["lines"].append(s)
    def add_finding(title, evidence, decision, priority="Közepes"):
        data["findings"].append({"Téma": title, "Bizonyíték": evidence, "Edzői következtetés": decision, "Prioritás": priority, "Forrás": role})

    # Formáció
    fa, fb = forms.get("team_a_start"), forms.get("team_b_start")
    if fa or fb:
        add_line(f"Formáció: {ta} {fa or 'n.a.'}, {tb} {fb or 'n.a.'}.")
        add_finding(f"{role}: formációs alaphelyzet", f"{ta}: {fa or 'n.a.'}; {tb}: {fb or 'n.a.'}.", "A meccstervet a formációs párosítás és a szélesség/félterület viszonyok alapján kell pontosítani.", "Közepes")

    # xG: először team statistics xG, fallback match info xG/per shot
    xa, xb = _fpi_find_first_pair_float_v86(flat, r"\bxG\b")
    if not (xa or xb):
        m = re.search(r"([\d.]+)\s*/\s*[\d.]+\s+xG\s*/\s*per shot\s+([\d.]+)\s*/\s*[\d.]+", flat, flags=re.I)
        if m:
            xa, xb = _fpi_float_v85(m.group(1)), _fpi_float_v85(m.group(2))
    if xa or xb:
        add_line(f"xG: {ta} {xa:.2f}, {tb} {xb:.2f}.")
        if xb > xa * 1.12:
            add_finding(f"{role}: ellenfél jobb helyzetminőség", f"xG: {ta} {xa:.2f}, {tb} {xb:.2f}.", "Boxvédekezés, belső zónák és átmenetek kontrollja kiemelt.", "Magas")
        elif xa > xb * 1.12:
            add_finding(f"{role}: saját oldal jobb helyzetminőség", f"xG: {ta} {xa:.2f}, {tb} {xb:.2f}.", "Támadóbb alapirány vállalható, de a rest defense biztosítás maradjon fókusz.", "Közepes")

    # Labdabirtoklás
    pa, pb = _fpi_find_first_pair_percent_v86(flat, r"Ball possession")
    if pa or pb:
        add_line(f"Labdabirtoklás: {ta} {pa:.0f}%, {tb} {pb:.0f}%.")
        if abs(pa-pb) >= 10:
            better = ta if pa > pb else tb
            add_finding(f"{role}: labdabirtoklási fölény", f"{better} birtokolt többet ({max(pa,pb):.0f}%).", "A labdabirtoklás önmagában nem elég: össze kell vetni xG-vel és kontraveszéllyel.", "Közepes")

    # Lövések
    sa, sota, sb, sotb = _fpi_find_ratio_pair_v86(flat, r"Shots\s*/\s*on target")
    if sa or sb:
        add_line(f"Lövések/kaput találó: {ta} {sa}/{sota}, {tb} {sb}/{sotb}.")
        if sotb > sota:
            add_finding(f"{role}: ellenfél pontosabb lövésprofil", f"Kaput találó lövések: {ta} {sota}, {tb} {sotb}.", "Lövőzónák zárása és box előtti nyomás szükséges.", "Közepes")

    # Kulcspassz
    ka, kaa, kb, kba = _fpi_find_ratio_pair_v86(flat, r"Key passes")
    if ka or kb:
        add_line(f"Kulcspasszok: {ta} {ka}, {tb} {kb}.")
        if kb > ka * 1.5:
            add_finding(f"{role}: ellenfél kreatív fölény", f"Kulcspasszok: {ta} {ka}, {tb} {kb}.", "Félterületek, 10-es zóna és visszagurítások kontrollja kiemelt.", "Magas")
        elif ka > kb * 1.5:
            add_finding(f"{role}: saját kreatív fölény", f"Kulcspasszok: {ta} {ka}, {tb} {kb}.", "A támadóharmadba jutás fenntartható, a befejezési minőségre kell fókuszálni.", "Közepes")

    # Kontra
    ca, csa, cb, csb = _fpi_find_ratio_pair_v86(flat, r"Counter-attacks\s*/\s*with shots")
    if ca or cb:
        add_line(f"Kontrák/lövéssel: {ta} {ca}/{csa}, {tb} {cb}/{csb}.")
        if cb >= ca and csb > csa:
            add_finding(f"{role}: ellenfél kontra-veszély", f"Kontrák/lövés: {ta} {ca}/{csa}, {tb} {cb}/{csb}.", "MD-3 átmeneti játék + HSR, MD-2 rest defense és labdavesztés utáni első reakció.", "Magas")
        elif ca > cb:
            add_finding(f"{role}: saját kontraaktivitás", f"Kontrák: {ta} {ca}, {tb} {cb}.", "Gyors átmeneti játék támadható út, mögöttes biztosítással.", "Közepes")

    # Pontrúgás
    spa, spsa, spb, spsb = _fpi_find_ratio_pair_v86(flat, r"Set-piece attacks\s*/\s*with shots")
    if spa or spb:
        add_line(f"Pontrúgásos támadások/lövéssel: {ta} {spa}/{spsa}, {tb} {spb}/{spsb}.")
        if spb and spsb >= spsa:
            add_finding(f"{role}: pontrúgás-veszély", f"Pontrúgásos támadások/lövések: {ta} {spa}/{spsa}, {tb} {spb}/{spsb}.", "MD-1 pontrúgás-védekezés, első kontakt és második labdák kontrollja.", "Magas")

    # PPDA
    ppdaa, ppdab = _fpi_find_first_pair_float_v86(flat, r"\bPPDA\b")
    if ppdaa or ppdab:
        add_line(f"PPDA: {ta} {ppdaa:.2f}, {tb} {ppdab:.2f}.")
        if ppdaa and ppdab:
            active = ta if ppdaa < ppdab else tb
            add_finding(f"{role}: presszingprofil", f"PPDA: {ta} {ppdaa:.2f}, {tb} {ppdab:.2f}.", f"Presszingintenzitásban {active} aktívabb; ehhez kell igazítani a labdakihozatalt.", "Közepes")

    # Pressing
    pra, prsa, prb, prsb = _fpi_find_ratio_pair_v86(flat, r"Pressing\s*/\s*successful")
    if pra or prb:
        add_line(f"Pressing/sikeres: {ta} {pra}/{prsa}, {tb} {prb}/{prsb}.")
        if pra and prb:
            pcta = prsa / max(pra, 1) * 100
            pctb = prsb / max(prb, 1) * 100
            if abs(pcta-pctb) >= 15:
                better = ta if pcta > pctb else tb
                add_finding(f"{role}: pressing hatékonysági különbség", f"Pressing sikeresség: {ta} {pcta:.0f}%, {tb} {pctb:.0f}%.", f"{better} hatékonyabban presszingelt; triggerzónák és kijátszás felülvizsgálandó.", "Közepes")

    # Lost balls own half
    lba, loha, lbb, lohb = _fpi_find_ratio_pair_v86(flat, r"Lost balls\s*/\s*in own half")
    if lba or lbb:
        add_line(f"Labdavesztés/saját térfélen: {ta} {lba}/{loha}, {tb} {lbb}/{lohb}.")
        if lohb > loha:
            add_finding(f"{role}: ellenfél saját térfeles labdavesztés", f"Saját térfeles labdavesztés: {ta} {loha}, {tb} {lohb}.", "Magasabb presszinggel vagy középső csapdával támadható lehet.", "Közepes")
        elif loha > lohb:
            add_finding(f"{role}: saját oldal labdavesztési kockázat", f"Saját térfeles labdavesztés: {ta} {loha}, {tb} {lohb}.", "Labdakihozatalban egyszerűbb döntések és biztosítás szükséges.", "Közepes")

    # Ha xG és possession ellentmond
    if xa or xb:
        if pa > pb + 10 and xb > xa:
            add_finding(f"{role}: kontroll vs veszély ellentmondás", f"{ta} többet birtokolt ({pa:.0f}%), de {tb} xG-je magasabb ({xb:.2f}).", "A labdabirtoklási fölény mellett az ellenfél hatékony átmeneteit/helyzetminőségét kell levédeni.", "Magas")
        if pb > pa + 10 and xa > xb:
            add_finding(f"{role}: ellenfél kontroll, saját hatékonyabb helyzetek", f"{tb} többet birtokolt ({pb:.0f}%), de {ta} xG-je magasabb ({xa:.2f}).", "Középső blokk + gyors átmenet működő irány lehet.", "Magas")

    data["lines"] = data["lines"][:14]
    data["findings"] = data["findings"][:12]
    return data


def _fpi_sportsbase_findings_v85(metrics: Dict[str, object], pdf_role: str = "PDF") -> List[dict]:
    findings = []
    if not metrics or len(metrics.keys()) <= 3:
        return findings
    ta = (metrics.get("teams") or {}).get("team_a", "A csapat")
    tb = (metrics.get("teams") or {}).get("team_b", "B csapat")

    def add(title, evidence, decision, priority="Közepes"):
        findings.append({"Téma": title, "Bizonyíték": evidence, "Edzői következtetés": decision, "Prioritás": priority, "Forrás": pdf_role})

    # xG / shot quality
    if metrics.get("xg_a") or metrics.get("xg_b"):
        xa, xb = metrics.get("xg_a", 0), metrics.get("xg_b", 0)
        if xb > xa * 1.15:
            add(f"{pdf_role}: {tb} jobb helyzetminőséget hozott", f"xG: {ta} {xa:.2f}, {tb} {xb:.2f}.", "Boxvédekezés, belső zónák és átmeneti védekezés külön fókusz.", "Magas")
        elif xa > xb * 1.15:
            add(f"{pdf_role}: {ta} jobb helyzetminőséget hozott", f"xG: {ta} {xa:.2f}, {tb} {xb:.2f}.", "Támadó alapirány vállalhatóbb, de minőségi befejezés és rest defense kell.", "Közepes")

    # Possession vs efficiency
    if metrics.get("possession_a") and metrics.get("xg_a") is not None and metrics.get("xg_b") is not None:
        if metrics.get("possession_a", 0) > metrics.get("possession_b", 0) + 10 and metrics.get("xg_b", 0) > metrics.get("xg_a", 0):
            add(f"{pdf_role}: labdabirtoklás vs hatékonyság ellentmondás", f"{ta} birtokolt többet ({metrics.get('possession_a'):.0f}%), de {tb} xG-je magasabb ({metrics.get('xg_b'):.2f}).", "Nem elég kontrollálni a labdát: az ellenfél átmeneti/helyzetminőségi veszélyét kell zárni.", "Magas")

    # Counters
    if metrics.get("counter_attacks_a") or metrics.get("counter_attacks_b"):
        ca, cb = metrics.get("counter_attacks_a", 0), metrics.get("counter_attacks_b", 0)
        csa, csb = metrics.get("counter_attacks_shots_a", 0), metrics.get("counter_attacks_shots_b", 0)
        if cb >= ca and csb > csa:
            add(f"{pdf_role}: ellenfél kontrái veszélyesek", f"Kontrák/lövés: {ta} {ca}/{csa}, {tb} {cb}/{csb}.", "MD-3 átmeneti játék + HSR, MD-2 rest defense és labdavesztés utáni első reakció.", "Magas")
        elif ca > cb:
            add(f"{pdf_role}: saját/első csapat kontraaktivitás", f"Kontrák: {ta} {ca}, {tb} {cb}.", "A gyors átmeneti játék támadható út, de mögöttes biztosítás kell.", "Közepes")

    # Set pieces
    if metrics.get("set_piece_attacks_a") or metrics.get("set_piece_attacks_b"):
        spa, spb = metrics.get("set_piece_attacks_a", 0), metrics.get("set_piece_attacks_b", 0)
        ssa, ssb = metrics.get("set_piece_attacks_shots_a", 0), metrics.get("set_piece_attacks_shots_b", 0)
        if spb > 0 and ssb >= ssa:
            add(f"{pdf_role}: pontrúgás-veszély", f"Pontrúgásos támadások/lövések: {ta} {spa}/{ssa}, {tb} {spb}/{ssb}.", "MD-1 pontrúgás-védekezés, első kontakt és második labdák kontrollja.", "Magas")

    # Key passes / creative player zones
    if metrics.get("key_passes_a") or metrics.get("key_passes_b"):
        ka, kb = metrics.get("key_passes_a", 0), metrics.get("key_passes_b", 0)
        if kb > ka * 1.5:
            add(f"{pdf_role}: ellenfél kreatív előny", f"Kulcspasszok: {ta} {ka}, {tb} {kb}.", "Kulcspassz-sávok, félterületek és 10-es körüli védekezés kiemelt.", "Magas")
        elif ka > kb * 1.5:
            add(f"{pdf_role}: saját kreatív előny", f"Kulcspasszok: {ta} {ka}, {tb} {kb}.", "A támadóharmadba jutás fenntartható, minőségi befejezésekre kell fókuszálni.", "Közepes")

    # Pressing
    if metrics.get("ppda_a") or metrics.get("ppda_b"):
        pa, pb = metrics.get("ppda_a", 0), metrics.get("ppda_b", 0)
        if pa and pb:
            active = ta if pa < pb else tb
            add(f"{pdf_role}: presszingprofil", f"PPDA: {ta} {pa:.2f}, {tb} {pb:.2f}.", f"Presszingintenzitásban {active} aktívabb; ehhez kell igazítani a labdakihozatalt és a trigger-zónákat.", "Közepes")

    return findings[:10]


def _fpi_tactical_pdf_insights(text: str) -> Dict[str, object]:
    """V8.4: a régi Tactical app build_pdf_insights motorja dolgozik.
    A visszatérő struktúrát FPI-kompatibilissé alakítjuk, hogy a meglévő Match Plan / PDF / Excel logika is használja.
    """
    raw_text = str(text or "")
    legacy = _fpi_legacy_build_pdf_insights(raw_text)
    sportsbase_metrics = _fpi_extract_sportsbase_team_stats_v85(raw_text)
    sportsbase_lines = _fpi_sportsbase_metric_lines_v85(sportsbase_metrics)
    sportsbase_findings = _fpi_sportsbase_findings_v85(sportsbase_metrics, pdf_role="PDF")
    direct_pdf = _fpi_direct_pdf_extract_v86(raw_text, role="PDF")
    # Ha a SportsBase parser nem fogott eleget, a direkt parser sorait és következtetéseit is használjuk.
    sportsbase_lines = list(dict.fromkeys((sportsbase_lines or []) + (direct_pdf.get("lines") or [])))
    sportsbase_findings = (sportsbase_findings or []) + (direct_pdf.get("findings") or [])

    topic_rows = legacy.get("topic_debug", []) or []
    blocks = {
        "formation": legacy.get("universal_blocks", {}).get("opponent_identity", []),
        "build_up": legacy.get("build_up_lines", []) or legacy.get("universal_blocks", {}).get("build_up", []),
        "direct_play": legacy.get("universal_blocks", {}).get("direct_play", []),
        "pressing": legacy.get("pressing_lines", []) or legacy.get("universal_blocks", {}).get("pressing", []),
        "defensive_block": legacy.get("universal_blocks", {}).get("defensive_block", []),
        "transition_attack": legacy.get("dynamics_lines", []) or legacy.get("universal_blocks", {}).get("transition_attack", []),
        "transition_defense": legacy.get("risk_lines", []) or legacy.get("universal_blocks", {}).get("transition_defense", []),
        "chance_creation": legacy.get("player_threat_lines", []) or legacy.get("universal_blocks", {}).get("chance_creation", []),
        "wide_play": legacy.get("universal_blocks", {}).get("wide_play", []),
        "central_play": legacy.get("universal_blocks", {}).get("central_play", []),
        "set_pieces": legacy.get("set_piece_lines", []),
        "key_players": legacy.get("player_threat_lines", []) or legacy.get("universal_blocks", {}).get("key_players", []),
        "weakness_risk": legacy.get("risk_lines", []),
        "strength": legacy.get("universal_blocks", {}).get("strength", []),
        "goalkeeper": legacy.get("universal_blocks", {}).get("goalkeeper", []),
        "match_dynamics": legacy.get("dynamics_lines", []),
        "recommendation": legacy.get("recommendation_lines", []),
        "sportsbase_metrics": sportsbase_lines,
        "provider_metrics": sportsbase_lines,
    }

    # Provider-aware topics: ha konkrét SportsBase KPI-k vannak, ezek is témának számítanak.
    for sb_line in sportsbase_lines[:8]:
        topic_rows.append({
            "Téma": "SportsBase KPI / taktikai adat",
            "Kulcs": "sportsbase_metrics",
            "Találat": 1,
            "Bizonyosság": 75,
            "Minta": sb_line,
        })

    # Legacy detected topics -> FPI topics. Ha nincs topic_debug, detected_topicsből fallback.
    topics = []
    for r in topic_rows:
        if r.get("Találat", 0) or r.get("Minta"):
            topics.append({
                "Téma": r.get("Téma", ""),
                "Kulcs": r.get("Kulcs", ""),
                "Találat": r.get("Találat", 0),
                "Bizonyosság": r.get("Bizonyosság", 0),
                "Minta": r.get("Minta", ""),
                "Forrás": "Legacy Tactical PDF Engine",
            })
    if not topics:
        for t in legacy.get("detected_topics", [])[:10]:
            topics.append({
                "Téma": t,
                "Kulcs": "",
                "Találat": 1,
                "Bizonyosság": 40,
                "Minta": "Legacy Tactical PDF Engine által felismert téma.",
                "Forrás": "Legacy Tactical PDF Engine",
            })

    preview = re.sub(r"\s+", " ", raw_text).strip()[:2200]
    return {
        "formation": legacy.get("formation", "n.a."),
        "blocks": blocks,
        "topics": topics[:20],
        "raw_text_len": len(raw_text),
        "preview": preview,
        "top_terms": _fpi_tactical_top_terms_v76(raw_text),
        "reader_version": legacy.get("reader_version", "Legacy Tactical PDF Engine"),
        "legacy": legacy,
        "legacy_detected_topics": legacy.get("detected_topics", []),
        "legacy_dna_lines": legacy.get("dna_lines", []),
        "legacy_risk_lines": legacy.get("risk_lines", []),
        "legacy_set_piece_lines": legacy.get("set_piece_lines", []),
        "legacy_dynamics_lines": legacy.get("dynamics_lines", []),
        "legacy_pressing_lines": legacy.get("pressing_lines", []),
        "legacy_build_up_lines": legacy.get("build_up_lines", []),
        "legacy_player_threat_lines": legacy.get("player_threat_lines", []),
        "legacy_recommendation_lines": legacy.get("recommendation_lines", []),
        "legacy_detected_names": legacy.get("detected_names", []),
        "sportsbase_metrics": sportsbase_metrics,
        "sportsbase_lines": sportsbase_lines,
        "sportsbase_findings": sportsbase_findings,
        "direct_pdf": direct_pdf,
    }

def _fpi_analysis_level(has_gps: bool, has_pdf: bool, has_team_excel: bool, has_player_excel: bool) -> Tuple[int, str]:
    if has_gps and has_pdf and has_team_excel and has_player_excel:
        return 4, "Full Intelligence – GPS + taktikai PDF + csapat Excel + játékos Excel"
    if has_gps and has_pdf and has_team_excel:
        return 3, "GPS + Tactical Team Intelligence"
    if has_gps and has_pdf:
        return 2, "GPS + Tactical PDF Intelligence"
    if has_gps:
        return 1, "GPS Only – Performance Intelligence"
    return 0, "Nincs elegendő adat"


def _fpi_metric_value_v79(metrics: Dict[str, float], key: str, default: float = 0.0) -> float:
    try:
        v = metrics.get(key, default) if isinstance(metrics, dict) else default
        if v in [None, "", "-", "nan"]:
            return default
        return float(v)
    except Exception:
        return default

def _fpi_normalized_tactical_metric_v79(key: str, value: float) -> float:
    """PDF/riportbarát normalizálás.
    Néhány taktikai export összesíti a százalékokat/PPDA-t több meccsre, ezért a nyers érték
    irreálisan nagy lehet. Itt csak riport-értelmezéshez normalizálunk, nem írjuk felül az alapadatot.
    """
    try:
        v = float(value)
    except Exception:
        return 0.0
    if key in ["possession_pct", "pressing_success_pct", "passes_accurate_pct"]:
        if 0 < v <= 1:
            return v * 100
        if v > 1000:
            return v / 100
        if v > 100:
            return v / 10
        return v
    if key == "ppda":
        if v > 1000:
            return v / 1000
        if v > 100:
            return v / 10
        return v
    if key == "xg":
        if v > 100:
            return v / 10
        return v
    return v

def _fpi_tactical_metric_label_v79(key: str) -> str:
    return {
        "possession_pct": "Labdabirtoklás",
        "shots": "Lövések",
        "xg": "xG",
        "entries_box": "Box entries",
        "key_passes": "Kulcspasszok",
        "corners": "Szögletek",
        "ppda": "PPDA",
        "pressing_success_pct": "Pressing sikeresség",
        "passes_accurate_pct": "Passzpontosság",
        "counterattacks": "Kontrák",
        "recoveries": "Labdaszerzések",
        "lost_balls": "Labdavesztések",
        "crosses": "Beadások",
    }.get(key, key)

def _fpi_tactical_compare_team_metrics_v79(own_metrics: Dict[str, float], opp_metrics: Dict[str, float]) -> List[dict]:
    rows = []
    metric_keys = ["possession_pct", "shots", "xg", "entries_box", "key_passes", "corners", "ppda", "pressing_success_pct", "counterattacks", "recoveries", "lost_balls", "crosses"]
    for k in metric_keys:
        own_raw = _fpi_metric_value_v79(own_metrics, k, 0.0)
        opp_raw = _fpi_metric_value_v79(opp_metrics, k, 0.0)
        if own_raw == 0 and opp_raw == 0:
            continue
        own = _fpi_normalized_tactical_metric_v79(k, own_raw)
        opp = _fpi_normalized_tactical_metric_v79(k, opp_raw)
        diff = own - opp
        if k == "ppda":
            better = "Saját intenzívebb presszing" if own < opp else "Ellenfél intenzívebb presszing" if opp < own else "Hasonló"
        else:
            better = "Saját előny" if diff > 0 else "Ellenfél előny" if diff < 0 else "Hasonló"
        rows.append({
            "metric": k,
            "Mutató": _fpi_tactical_metric_label_v79(k),
            "Saját": own,
            "Ellenfél": opp,
            "Eltérés": diff,
            "Értelmezés": better,
        })
    return rows

def _fpi_player_table_top_v79(player_tables: Dict[str, pd.DataFrame], key: str, value_col: str) -> Tuple[str, float]:
    try:
        df = player_tables.get(key)
        if not isinstance(df, pd.DataFrame) or df.empty or "player" not in df.columns:
            return "", 0.0
        row = df.iloc[0]
        return str(row.get("player", "")), float(row.get(value_col, 0) or 0)
    except Exception:
        return "", 0.0

def _fpi_build_excel_driven_tactical_findings_v79(
    own_team: Dict[str, float],
    opp_team: Dict[str, float],
    own_players: Dict[str, pd.DataFrame],
    opp_players: Dict[str, pd.DataFrame],
    pdf_topics: Optional[List[dict]] = None,
) -> List[dict]:
    findings = []
    comparisons = _fpi_tactical_compare_team_metrics_v79(own_team, opp_team)
    by_key = {r["metric"]: r for r in comparisons}

    def add(title, evidence, decision, priority="Közepes"):
        findings.append({"Téma": title, "Bizonyíték": evidence, "Edzői következtetés": decision, "Prioritás": priority})

    # V8.5: Provider-aware PDF findings elsőbbség, ha konkrét KPI-t nyertünk ki a PDF-ből.
    pdf_topics = pdf_topics or []
    for row in pdf_topics:
        if isinstance(row, dict) and row.get("Edzői következtetés"):
            tema = str(row.get("Téma", "PDF taktikai adat"))
            if not tema.startswith("PDF"):
                tema = "PDF: " + tema
            add(tema, str(row.get("Bizonyíték", "")), str(row.get("Edzői következtetés", "")), str(row.get("Prioritás", "Közepes")))

    # V8.4: PDF insightok elsőként kerülnek be, ha a régi Tactical engine konkrét témát/sort talált.
    for row in pdf_topics[:8]:
        tema = str(row.get("Téma", "")).strip()
        sample = str(row.get("Minta", "")).strip()
        if not tema:
            continue
        tnorm = _fpi_tactical_norm(tema)
        if "presszing" in tnorm or "letamadas" in tnorm or "press" in tnorm:
            add("PDF: presszing / letámadás", sample or tema, "Presszingkijátszás, első passzsor és harmadik emberes megoldások külön edzésfókuszt kapjanak.", "Magas")
        elif "atmenet" in tnorm or "kontra" in tnorm or "transition" in tnorm:
            add("PDF: átmenetek / kontrák", sample or tema, "MD-3 átmeneti játék + HSR/sprint inger, MD-2 rest defense biztosítás.", "Magas")
        elif "pontrugas" in tnorm or "szoglet" in tnorm or "set" in tnorm:
            add("PDF: pontrúgások", sample or tema, "MD-1 pontrúgás-védekezés, első kontakt és második labdák kontrollja.", "Magas")
        elif "szel" in tnorm or "beadas" in tnorm or "wide" in tnorm:
            add("PDF: szélső játék", sample or tema, "Oldali 1v1 védekezés, beadásblokkolás és hosszú oldali zárás.", "Közepes")
        elif "kulcsj" in tnorm or "player" in tnorm:
            add("PDF: kulcsjátékosok", sample or tema, "A PDF-ben jelzett játékosok passzsávjait és döntési testhelyzetét kontrollálni kell.", "Közepes")
        else:
            add(f"PDF: {tema}", sample or "PDF-ből felismert taktikai téma.", "Videóval ellenőrizendő, majd a heti taktikai blokkba beépíthető.", "Közepes")

    # Team Excel alapján
    if "shots" in by_key:
        r = by_key["shots"]
        if r["Ellenfél"] > r["Saját"] * 1.10:
            add("Ellenfél lövésvolumen előnye", f"Ellenfél lövések: {r['Ellenfél']:.1f}, saját: {r['Saját']:.1f}.", "A lövőzónák zárása, box előtti nyomás és második labdák kontrollja kiemelt.", "Magas")
        elif r["Saját"] > r["Ellenfél"] * 1.10:
            add("Saját lövésvolumen előny", f"Saját lövések: {r['Saját']:.1f}, ellenfél: {r['Ellenfél']:.1f}.", "A támadóharmadba jutás fenntartható, a minőségi befejezésekre kell fókuszálni.", "Közepes")

    if "xg" in by_key:
        r = by_key["xg"]
        if r["Ellenfél"] > r["Saját"] * 1.10:
            add("Ellenfél magasabb xG-profil", f"Ellenfél xG: {r['Ellenfél']:.1f}, saját: {r['Saját']:.1f}.", "Nem csak lövésszám, hanem helyzetminőség ellen is védekezni kell: boxvédekezés és belső zónák.", "Magas")
        elif r["Saját"] > r["Ellenfél"] * 1.10:
            add("Saját xG-előny", f"Saját xG: {r['Saját']:.1f}, ellenfél: {r['Ellenfél']:.1f}.", "A meccsterv támadó oldalon vállalhatóbb lehet, ha a readiness ezt elbírja.", "Közepes")

    if "possession_pct" in by_key:
        r = by_key["possession_pct"]
        if r["Saját"] - r["Ellenfél"] > 5:
            add("Saját labdabirtoklási előny", f"Saját labdabirtoklás: {r['Saját']:.1f}%, ellenfél: {r['Ellenfél']:.1f}%.", "Türelmesebb labdabirtoklás / POZ-KIE irány támogatható, de rest defense biztosítás kell.", "Közepes")
        elif r["Ellenfél"] - r["Saját"] > 5:
            add("Ellenfél labdabirtoklási fölény várható", f"Ellenfél labdabirtoklás: {r['Ellenfél']:.1f}%, saját: {r['Saját']:.1f}%.", "Középső blokk + átmenet, labdaszerzés utáni első passz és kontraindítás fókusz.", "Magas")

    if "ppda" in by_key:
        r = by_key["ppda"]
        if r["Ellenfél"] < r["Saját"] * 0.85:
            add("Ellenfél aktívabb presszingprofil", f"Ellenfél PPDA: {r['Ellenfél']:.1f}, saját: {r['Saját']:.1f}.", "Presszingkijátszás, kapus/CB első döntés, harmadik emberes megoldások MD-2 fókuszban.", "Magas")
        elif r["Saját"] < r["Ellenfél"] * 0.85:
            add("Saját presszingelőny", f"Saját PPDA: {r['Saját']:.1f}, ellenfél: {r['Ellenfél']:.1f}.", "Magasabb labdaszerzésre épülő PRS opció reális, ha a GPS readiness megfelelő.", "Közepes")

    if "corners" in by_key:
        r = by_key["corners"]
        if r["Ellenfél"] > r["Saját"] * 1.10:
            add("Ellenfél pontrúgás-terhelés", f"Ellenfél szögletek: {r['Ellenfél']:.1f}, saját: {r['Saját']:.1f}.", "MD-1 pontrúgás-védekezés, zónák és második labdák kötelező blokk.", "Magas")

    if "recoveries" in by_key:
        r = by_key["recoveries"]
        if r["Ellenfél"] > r["Saját"] * 1.08:
            add("Ellenfél labdaszerzési aktivitás", f"Ellenfél labdaszerzések: {r['Ellenfél']:.1f}, saját: {r['Saját']:.1f}.", "Labdabiztonság, első érintés és visszatámadás elleni biztosítás fontos.", "Közepes")
        elif r["Saját"] > r["Ellenfél"] * 1.08:
            add("Saját labdaszerzési előny", f"Saját labdaszerzések: {r['Saját']:.1f}, ellenfél: {r['Ellenfél']:.1f}.", "Aktívabb presszing vagy középső blokkban labdaszerzésre építő terv működhet.", "Közepes")

    # Player Excel alapján
    own_creator, own_creator_val = _fpi_player_table_top_v79(own_players, "creators", "key_passes")
    opp_creator, opp_creator_val = _fpi_player_table_top_v79(opp_players, "creators", "key_passes")
    if opp_creator:
        add("Ellenfél kreatív kulcsjátékos", f"{opp_creator}: {opp_creator_val:.1f} kulcspassz.", "A kreatív játékos passzsávjait és testhelyzetét kontrollálni kell.", "Magas")
    if own_creator:
        add("Saját kreatív kapcsolódási pont", f"{own_creator}: {own_creator_val:.1f} kulcspassz.", "A saját támadást érdemes rajta keresztül strukturálni, különösen átmenetben vagy félterületben.", "Közepes")

    own_prog, own_prog_val = _fpi_player_table_top_v79(own_players, "progressors", "progressive_passes")
    opp_prog, opp_prog_val = _fpi_player_table_top_v79(opp_players, "progressors", "progressive_passes")
    if opp_prog:
        add("Ellenfél progresszor", f"{opp_prog}: {opp_prog_val:.1f} progresszív passz.", "Nyomástrigger: ha ő kapja szabadon, zárni kell a belső passzsávot.", "Magas")
    if own_prog:
        add("Saját progresszor", f"{own_prog}: {own_prog_val:.1f} progresszív passz.", "Build-upban ő lehet a gyors előrehaladás egyik kulcsa.", "Közepes")

    opp_duel, opp_duel_val = _fpi_player_table_top_v79(opp_players, "duel_players", "defensive_challenges")
    if opp_duel:
        add("Ellenfél párharcerős játékos", f"{opp_duel}: {opp_duel_val:.1f} védekező párharc.", "Kerülni kell az izolált 1v1-et vele szemben; kombinációval vagy oldalváltással bontani.", "Közepes")

    # PDF téma fallback, ha van
    for row in (pdf_topics or [])[:3]:
        tema = row.get("Téma")
        if tema and not any(tema in f["Téma"] for f in findings):
            add(f"PDF téma: {tema}", row.get("Minta", "PDF-ből felismert taktikai jelzés."), "A PDF-jelzés ellenőrzendő videóval, majd edzéscélként beépíthető.", "Közepes")

    if not findings:
        add("Nincs elég taktikai adatból képzett következtetés", "A taktikai Excel/PDF mapping nem adott értelmezhető különbséget.", "Ellenőrizd a Tactical Mapper mezőit, különösen: xG, lövések, PPDA, labdabirtoklás, kulcspasszok, játékperc.", "Alacsony")

    return findings[:10]



# =========================================================
# V129 - Tactical Framework: 7 dimenzió -> 9 stratégiai profil
# =========================================================
FPI_TACTICAL_DIMENSIONS_V129 = [
    ("pressing", "Letámadás"),
    ("build_up", "Labdakihozatal"),
    ("transitions", "Átmenetek"),
    ("attacking_play", "Támadó játék"),
    ("set_pieces", "Pontrúgások"),
    ("possession", "Labdabirtoklás"),
    ("shot_profile", "Lövésprofil"),
]

FPI_STRATEGY_PALETTE_V129 = [
    {"code": "KON", "name": "Kontra mély blokkból", "style": "Direkt", "block": "Mély", "description": "Mélyebb védekezésből gyors, direkt támadásindítás."},
    {"code": "GAT", "name": "Gyors átmenet", "style": "Direkt", "block": "Közép", "description": "Labdaszerzés után gyors előrejáték, kevés passzból veszély."},
    {"code": "BAT", "name": "Középső blokk + átmenet", "style": "Vegyes", "block": "Közép", "description": "Középső zónás védekezés, majd gyors átmeneti támadás."},
    {"code": "KIE", "name": "Kiegyensúlyozott", "style": "Kiegyensúlyozott", "block": "Közép", "description": "Stabil, kockázatkerülőbb alapjáték, kontrollált intenzitással."},
    {"code": "PRS", "name": "Presszing + átmenet", "style": "Direkt/Presszing", "block": "Közép-magas", "description": "Aktív letámadás, labdaszerzés után gyors támadásvezetés."},
    {"code": "MLT", "name": "Magas letámadás", "style": "Agresszív", "block": "Magas", "description": "Magas blokkból agresszív nyomás és korai labdaszerzés."},
    {"code": "DOM", "name": "Dominancia", "style": "Kontroll", "block": "Magas", "description": "Labdabirtoklásra és területi fölényre épülő meccskontroll."},
    {"code": "POZ", "name": "Pozíciós támadás", "style": "Kontroll", "block": "Közép-magas", "description": "Türelmes építkezés, félterületek és szélesség használata."},
    {"code": "LAB", "name": "Labdatartás mélyebben", "style": "Kontroll", "block": "Alacsony-közép", "description": "Biztonságosabb labdatartás mélyebb szerkezetből."},
]


def _fpi_global_strategy_palette_rows_v129() -> List[Tuple[str, str, str, str, str]]:
    """Globális stratégiai paletta a Metodika és Tactical Pro+ nézethez."""
    return [(x["code"], x["name"], x["style"], x["block"], x["description"]) for x in FPI_STRATEGY_PALETTE_V129]


def _fpi_clamp_v129(x: float, lo: float = 0.0, hi: float = 10.0) -> float:
    try:
        return float(max(lo, min(hi, x)))
    except Exception:
        return lo


def _fpi_topic_present_v129(blocks: Dict[str, object], pdf_topics: List[object], *keys: str) -> bool:
    keys_l = {str(k).lower() for k in keys}
    for k in keys:
        if blocks.get(k):
            return True
    for item in pdf_topics or []:
        if isinstance(item, dict):
            joined = " ".join(str(item.get(c, "")) for c in ["Téma", "Tema", "label", "topic", "Szöveg", "Megjegyzés"]).lower()
        else:
            joined = str(item).lower()
        if any(k.replace("_", " ") in joined or k in joined for k in keys_l):
            return True
    return False


def _fpi_tactical_dimension_scores_v129(
    blocks: Dict[str, object],
    pdf_topics: List[object],
    own_team_metrics: Dict[str, object],
    opp_team_metrics: Dict[str, object],
) -> Dict[str, float]:
    """7 dimenziós taktikai profil 0-10 skálán.
    A cél nem egzakt scouting, hanem strukturált, összehasonlítható döntési jel.
    """
    own_pos = _fpi_normalized_tactical_metric_v79("possession_pct", _fpi_metric_value_v79(own_team_metrics, "possession_pct")) or 50
    opp_pos = _fpi_normalized_tactical_metric_v79("possession_pct", _fpi_metric_value_v79(opp_team_metrics, "possession_pct")) or 50
    own_ppda = _fpi_normalized_tactical_metric_v79("ppda", _fpi_metric_value_v79(own_team_metrics, "ppda")) or None
    opp_ppda = _fpi_normalized_tactical_metric_v79("ppda", _fpi_metric_value_v79(opp_team_metrics, "ppda")) or None
    own_shots = _fpi_metric_value_v79(own_team_metrics, "shots") or 0
    opp_shots = _fpi_metric_value_v79(opp_team_metrics, "shots") or 0
    own_xg = _fpi_normalized_tactical_metric_v79("xg", _fpi_metric_value_v79(own_team_metrics, "xg")) or 0
    opp_xg = _fpi_normalized_tactical_metric_v79("xg", _fpi_metric_value_v79(opp_team_metrics, "xg")) or 0
    own_corners = _fpi_metric_value_v79(own_team_metrics, "corners") or 0
    opp_corners = _fpi_metric_value_v79(opp_team_metrics, "corners") or 0
    own_counters = _fpi_metric_value_v79(own_team_metrics, "counterattacks") or 0
    opp_counters = _fpi_metric_value_v79(opp_team_metrics, "counterattacks") or 0
    own_crosses = _fpi_metric_value_v79(own_team_metrics, "crosses") or 0
    opp_crosses = _fpi_metric_value_v79(opp_team_metrics, "crosses") or 0

    pressing = 4.5
    if _fpi_topic_present_v129(blocks, pdf_topics, "pressing"):
        pressing += 2.0
    if opp_ppda and opp_ppda < 9:
        pressing += 1.7
    if own_ppda and own_ppda < 9:
        pressing += 0.8

    build_up = 4.8
    if _fpi_topic_present_v129(blocks, pdf_topics, "build_up", "possession", "central_play"):
        build_up += 1.7
    if own_pos >= 55 or opp_pos >= 55:
        build_up += 1.1
    if _fpi_topic_present_v129(blocks, pdf_topics, "direct_play"):
        build_up -= 0.8

    transitions = 4.5
    if _fpi_topic_present_v129(blocks, pdf_topics, "transition_attack", "transition_defense"):
        transitions += 2.0
    if own_counters + opp_counters >= 6:
        transitions += 1.6

    attacking_play = 4.8
    if _fpi_topic_present_v129(blocks, pdf_topics, "chance_creation", "wide_play", "central_play"):
        attacking_play += 1.6
    if max(own_shots, opp_shots) >= 12 or max(own_xg, opp_xg) >= 1.4:
        attacking_play += 1.2
    if own_crosses + opp_crosses >= 22:
        attacking_play += 0.6

    set_pieces = 4.2
    if _fpi_topic_present_v129(blocks, pdf_topics, "set_pieces"):
        set_pieces += 2.2
    if own_corners + opp_corners >= 8:
        set_pieces += 1.3

    possession = 4.8 + (max(own_pos, opp_pos) - 50) / 7.0
    if _fpi_topic_present_v129(blocks, pdf_topics, "possession", "build_up"):
        possession += 1.0
    if _fpi_topic_present_v129(blocks, pdf_topics, "direct_play"):
        possession -= 1.0

    shot_profile = 4.5
    if _fpi_topic_present_v129(blocks, pdf_topics, "chance_creation"):
        shot_profile += 1.2
    if max(own_shots, opp_shots) >= 10:
        shot_profile += 1.0
    if max(own_xg, opp_xg) >= 1.2:
        shot_profile += 1.0

    return {
        "pressing": round(_fpi_clamp_v129(pressing), 1),
        "build_up": round(_fpi_clamp_v129(build_up), 1),
        "transitions": round(_fpi_clamp_v129(transitions), 1),
        "attacking_play": round(_fpi_clamp_v129(attacking_play), 1),
        "set_pieces": round(_fpi_clamp_v129(set_pieces), 1),
        "possession": round(_fpi_clamp_v129(possession), 1),
        "shot_profile": round(_fpi_clamp_v129(shot_profile), 1),
    }


def _fpi_select_tactical_strategy_v129(
    readiness: int,
    blocks: Dict[str, object],
    pdf_topics: List[object],
    own_team_metrics: Dict[str, object],
    opp_team_metrics: Dict[str, object],
) -> Dict[str, object]:
    """9 stratégiai profil + meccsspecifikus differenciálás.
    Nem csak egy kódot ad: elsődleges profil, másodlagos profil, arány, dimenziók és fókuszok is készülnek.
    """
    dims = _fpi_tactical_dimension_scores_v129(blocks, pdf_topics, own_team_metrics, opp_team_metrics)
    own_pos = _fpi_normalized_tactical_metric_v79("possession_pct", _fpi_metric_value_v79(own_team_metrics, "possession_pct")) or 50
    opp_pos = _fpi_normalized_tactical_metric_v79("possession_pct", _fpi_metric_value_v79(opp_team_metrics, "possession_pct")) or 50
    opp_ppda = _fpi_normalized_tactical_metric_v79("ppda", _fpi_metric_value_v79(opp_team_metrics, "ppda")) or None
    opp_counters = _fpi_metric_value_v79(opp_team_metrics, "counterattacks") or 0
    opp_corners = _fpi_metric_value_v79(opp_team_metrics, "corners") or 0
    own_xg = _fpi_normalized_tactical_metric_v79("xg", _fpi_metric_value_v79(own_team_metrics, "xg")) or 0
    opp_xg = _fpi_normalized_tactical_metric_v79("xg", _fpi_metric_value_v79(opp_team_metrics, "xg")) or 0

    # Pontozás a 9 palettára. Egy meccs több profilhoz is közel lehet, ezért top2-t adunk vissza.
    scores = {x["code"]: 0.0 for x in FPI_STRATEGY_PALETTE_V129}
    scores["KIE"] += 4.0
    scores["BAT"] += dims["transitions"] * 0.8 + dims["attacking_play"] * 0.25
    scores["GAT"] += dims["transitions"] * 0.9 + (2.0 if _fpi_topic_present_v129(blocks, pdf_topics, "direct_play") else 0.0)
    scores["KON"] += dims["transitions"] * 0.5 + (2.0 if readiness < 55 else 0.0)
    scores["PRS"] += dims["pressing"] * 0.95 + dims["transitions"] * 0.35
    scores["MLT"] += dims["pressing"] * 1.15 + (1.5 if opp_ppda and opp_ppda < 8 else 0.0)
    scores["DOM"] += dims["possession"] * 0.9 + dims["build_up"] * 0.55 + (1.2 if max(own_pos, opp_pos) >= 56 else 0.0)
    scores["POZ"] += dims["build_up"] * 0.85 + dims["attacking_play"] * 0.55 + (1.0 if _fpi_topic_present_v129(blocks, pdf_topics, "central_play", "wide_play") else 0.0)
    scores["LAB"] += dims["possession"] * 0.5 + (1.5 if readiness < 60 else 0.0)

    if opp_counters > 0 or _fpi_topic_present_v129(blocks, pdf_topics, "transition_attack"):
        scores["BAT"] += 2.2
        scores["GAT"] += 1.4
    if _fpi_topic_present_v129(blocks, pdf_topics, "set_pieces") or opp_corners > 0:
        scores["KIE"] += 0.6
    if own_xg and opp_xg and own_xg > opp_xg * 1.10 and readiness >= 65:
        scores["POZ"] += 1.2
        scores["DOM"] += 1.0
    if readiness < 55:
        scores["BAT"] += 1.0
        scores["KON"] += 1.0
        scores["MLT"] -= 2.0
        scores["PRS"] -= 1.0

    ranked = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    primary_code, primary_score = ranked[0]
    secondary_code, secondary_score = ranked[1]
    code_to_row = {x["code"]: x for x in FPI_STRATEGY_PALETTE_V129}
    total = max(primary_score + secondary_score, 0.1)
    primary_pct = int(round(primary_score / total * 100))
    secondary_pct = 100 - primary_pct

    detail_parts = []
    if primary_code in ["BAT", "GAT", "KON"]:
        if opp_counters > 0 or _fpi_topic_present_v129(blocks, pdf_topics, "transition_attack"):
            detail_parts.append("átmenetek és rest defense biztosítás")
        if _fpi_topic_present_v129(blocks, pdf_topics, "wide_play"):
            detail_parts.append("oldali védekezési kontroll")
    if primary_code in ["PRS", "MLT"]:
        detail_parts.append("presszingtrigger és mögöttes biztosítás")
    if primary_code in ["DOM", "POZ", "LAB"]:
        detail_parts.append("labdakihozatal és türelmes progresszió")
    if _fpi_topic_present_v129(blocks, pdf_topics, "set_pieces") or opp_corners > 0:
        detail_parts.append("pontrúgások és második labdák")
    if not detail_parts:
        detail_parts.append("stabil szerkezet és kontrollált kockázat")

    primary = code_to_row[primary_code]
    secondary = code_to_row[secondary_code]
    recommendation = f"{primary_code} – {primary['name']}, " + "; ".join(detail_parts[:2])
    b_plan = f"{secondary_code} – {secondary['name']}"
    return {
        "primary_code": primary_code,
        "primary_name": primary["name"],
        "secondary_code": secondary_code,
        "secondary_name": secondary["name"],
        "primary_pct": primary_pct,
        "secondary_pct": secondary_pct,
        "recommendation": recommendation,
        "plan_b": b_plan,
        "dimensions": dims,
        "scores": {k: round(v, 2) for k, v in scores.items()},
        "detail_focus": detail_parts[:4],
    }


def _fpi_strategy_framework_to_rows_v129(framework: Dict[str, object]) -> List[Dict[str, object]]:
    dims = framework.get("dimensions", {}) if isinstance(framework, dict) else {}
    out = []
    for key, label in FPI_TACTICAL_DIMENSIONS_V129:
        out.append({"Dimenzió": label, "Érték (0-10)": dims.get(key, "n.a.")})
    return out

def _fpi_build_adaptive_match_training_plan(gps_context: Dict[str, object], tactical: Dict[str, object]) -> Dict[str, object]:
    readiness = int(gps_context.get("readiness_score", 70) or 70)
    priorities = gps_context.get("priorities", []) or []
    pdfi = tactical.get("pdf_insights") or {}
    opp_team_metrics = tactical.get("team_metrics") or {}
    opp_player_tables = tactical.get("player_tables") or {}
    own_team_metrics = ((tactical.get("own") or {}).get("team_metrics") or {})
    own_player_tables = ((tactical.get("own") or {}).get("player_tables") or {})
    blocks = pdfi.get("blocks", {}) if isinstance(pdfi, dict) else {}
    pdf_topics = pdfi.get("topics", []) if isinstance(pdfi, dict) else []
    if isinstance(pdfi, dict):
        pdf_topics = list(pdfi.get("sportsbase_findings", []) or []) + list(pdf_topics or [])

    tactical_findings = _fpi_build_excel_driven_tactical_findings_v79(
        own_team_metrics,
        opp_team_metrics,
        own_player_tables,
        opp_player_tables,
        pdf_topics,
    )

    gps_only_mode = not tactical_findings and not blocks and not own_team_metrics and not opp_team_metrics and not pdf_topics
    if gps_only_mode:
        for msg in _fpi_gps_only_conclusions_v95(gps_context, priorities, readiness, str(gps_context.get("selected_week", "")), limit=6):
            tactical_findings.append({
                "Téma": "GPS-only konklúzió",
                "Bizonyíték": "GPS / readiness / edzés-meccs arányok",
                "Edzői következtetés": msg,
                "Prioritás": "Közepes",
            })

    risks = []
    for f in tactical_findings:
        if f.get("Prioritás") in ["Magas", "Közepes"]:
            risks.append(f"{f.get('Téma')}: {f.get('Edzői következtetés')}")
    if blocks.get("transition_attack"):
        risks.append("PDF alapján: ellenfél-kontrák / gyors átmenetek kezelése")
    if blocks.get("set_pieces"):
        risks.append("PDF alapján: pontrúgás-védekezés és második labdák")
    if blocks.get("wide_play"):
        risks.append("PDF alapján: szélső játék, beadások, oldali túlterhelések")
    if blocks.get("pressing"):
        risks.append("PDF alapján: presszing kijátszása és első passzsor döntései")
    if not risks:
        risks.append("GPS-alapú terhelési és readiness kockázatok")

    # Plan A: Excel + PDF + GPS együtt - V129 taktikai keretrendszerrel
    opp_pos = _fpi_normalized_tactical_metric_v79("possession_pct", _fpi_metric_value_v79(opp_team_metrics, "possession_pct"))
    own_pos = _fpi_normalized_tactical_metric_v79("possession_pct", _fpi_metric_value_v79(own_team_metrics, "possession_pct"))
    opp_ppda = _fpi_normalized_tactical_metric_v79("ppda", _fpi_metric_value_v79(opp_team_metrics, "ppda"))
    own_ppda = _fpi_normalized_tactical_metric_v79("ppda", _fpi_metric_value_v79(own_team_metrics, "ppda"))
    opp_counters = _fpi_metric_value_v79(opp_team_metrics, "counterattacks")
    opp_corners = _fpi_metric_value_v79(opp_team_metrics, "corners")
    opp_xg = _fpi_normalized_tactical_metric_v79("xg", _fpi_metric_value_v79(opp_team_metrics, "xg"))
    own_xg = _fpi_normalized_tactical_metric_v79("xg", _fpi_metric_value_v79(own_team_metrics, "xg"))

    strategy_framework = _fpi_select_tactical_strategy_v129(
        readiness, blocks, pdf_topics, own_team_metrics, opp_team_metrics
    )
    plan_a = strategy_framework.get("recommendation", "KIE – Kiegyensúlyozott")
    if strategy_framework.get("plan_b"):
        risks.append(f"B terv / alternatív profil: {strategy_framework.get('plan_b')} ({strategy_framework.get('secondary_pct')}%).")

    md_plan = [
        ("MD+1/MD-5", "Regeneráció / alacsony intenzitás", "Előző terhelés visszarendezése."),
        ("MD-4", "Volumen + saját játékmodell", "Stabil csapatvolumen és saját labdakihozatal / védekezési alapok."),
        ("MD-3", "HSR / sprint exponálás + átmenetek", "Meccsintenzitás előkészítése, de kontrollált mennyiséggel."),
        ("MD-2", "Ellenfél-specifikus taktikai nap", "; ".join(risks[:2]) if risks else "Meccsterv."),
        ("MD-1", "Aktiváció + pontrúgások", "Frissítés, gyors döntések, fix helyzetek."),
    ]
    if gps_only_mode:
        md_plan = [("MD+1/MD-5", "Regeneráció / monitoring", "Előző terhelés visszarendezése.")] + _fpi_gps_only_md_plan_v95(gps_context, readiness, priorities, str(gps_context.get("selected_week", "")))
        plan_a = "GPS-only – erőnléti fókuszú mikrociklus"
    if (not gps_only_mode) and readiness < 55:
        md_plan[2] = ("MD-3", "Rövid specifikus exponálás", "Csak célzott HSR/sprint inger, alacsony volumen.")
    if (not gps_only_mode) and any("presszing" in r.lower() or "ppda" in r.lower() for r in risks):
        md_plan[1] = ("MD-4", "Presszingkijátszás + labdakihozatal", "Ellenfél presszingprofil / PPDA alapján első passzsor és harmadik ember.")
    if (not gps_only_mode) and any("kontra" in r.lower() or "átmenet" in r.lower() for r in risks):
        md_plan[2] = ("MD-3", "Átmeneti játék + HSR/sprint", "Kontrák és gyors átmenetek miatt futóintenzitás + döntésgyorsaság.")
        md_plan[3] = ("MD-2", "Rest defense + kontrák elleni biztosítás", "Ellenfél átmeneti veszélyei miatt.")
    if (not gps_only_mode) and (opp_corners > 0 or blocks.get("set_pieces")):
        md_plan[-1] = ("MD-1", "Aktiváció + pontrúgás fókusz", "Szöglet/pontrúgás profil alapján.")

    player_focus = []
    for f in tactical_findings:
        if "játékos" in f.get("Téma", "").lower() or "progresszor" in f.get("Téma", "").lower():
            player_focus.append(f"{f.get('Téma')}: {f.get('Bizonyíték')} -> {f.get('Edzői következtetés')}")
    if not player_focus and isinstance(opp_player_tables, dict):
        for key, label in [("creators", "ellenfél kreatív játékos"), ("progressors", "ellenfél progresszor"), ("duel_players", "ellenfél párharcerős játékos")]:
            dfp = opp_player_tables.get(key)
            if isinstance(dfp, pd.DataFrame) and not dfp.empty and "player" in dfp.columns:
                player_focus.append(f"{dfp.iloc[0]['player']} – {label}")
    if not player_focus and priorities:
        for p in priorities[:3]:
            if isinstance(p, dict):
                player_focus.append(p.get("Teendő", p.get("Cím", "Játékosszintű monitoring")))

    return {
        "analysis_level": tactical.get("analysis_level_label", "GPS Only"),
        "plan_a": plan_a,
        "risks": list(dict.fromkeys(risks))[:6],
        "md_plan": md_plan,
        "player_focus": player_focus[:6],
        "tactical_findings": tactical_findings[:10],
        "team_comparison": _fpi_tactical_compare_team_metrics_v79(own_team_metrics, opp_team_metrics),
        "strategy_framework": strategy_framework if 'strategy_framework' in locals() else {},
        "plan_b": (strategy_framework.get("plan_b") if 'strategy_framework' in locals() and isinstance(strategy_framework, dict) else ""),
    }


def _merge_tactical_pdf_insights(own_insights: Dict[str, object], opp_insights: Dict[str, object]) -> Dict[str, object]:
    """Saját + ellenfél taktikai PDF insightok összefűzése úgy, hogy a régi logika se vesszen el."""
    merged_blocks = {}
    for key in TACTICAL_TOPIC_TAGS_FPI.keys():
        own_lines = ((own_insights or {}).get("blocks") or {}).get(key, []) or []
        opp_lines = ((opp_insights or {}).get("blocks") or {}).get(key, []) or []
        merged_blocks[key] = [f"Saját: {x}" for x in own_lines[:4]] + [f"Ellenfél: {x}" for x in opp_lines[:4]]

    topics = []
    for source_label, src in [("Saját", own_insights or {}), ("Ellenfél", opp_insights or {})]:
        for row in src.get("topics", []) or []:
            r = dict(row)
            r["Forrás"] = source_label
            topics.append(r)
    topics = sorted(topics, key=lambda x: x.get("Bizonyosság", 0), reverse=True)

    formation = (opp_insights or {}).get("formation") or (own_insights or {}).get("formation") or "n.a."
    sportsbase_findings = []
    for role, src in [("Saját PDF", own_insights or {}), ("Ellenfél PDF", opp_insights or {})]:
        src_findings = list(src.get("sportsbase_findings", []) or [])
        src_findings += list(((src.get("direct_pdf") or {}).get("findings") or []))
        for f in src_findings:
            ff = dict(f)
            ff["Forrás"] = role
            sportsbase_findings.append(ff)
    sportsbase_lines = [f"Saját: {x}" for x in ((own_insights or {}).get("sportsbase_lines", []) or [])[:8]] + [f"Ellenfél: {x}" for x in ((opp_insights or {}).get("sportsbase_lines", []) or [])[:8]]

    return {
        "formation": formation,
        "blocks": merged_blocks,
        "topics": topics[:18],
        "raw_text_len": int((own_insights or {}).get("raw_text_len", 0) or 0) + int((opp_insights or {}).get("raw_text_len", 0) or 0),
        "sportsbase_findings": sportsbase_findings[:12],
        "sportsbase_lines": sportsbase_lines[:16],
        "own": own_insights or {},
        "opponent": opp_insights or {},
    }

def _tactical_key_numbers_summary(metrics: Dict[str, float]) -> str:
    if not metrics:
        return "Nincs értelmezhető taktikai csapat KPI."
    label_map = {
        "possession_pct": "Labdabirtoklás",
        "shots": "Lövések",
        "xg": "xG",
        "entries_box": "Box entries",
        "key_passes": "Kulcspasszok",
        "corners": "Szögletek",
        "ppda": "PPDA",
        "pressing_success_pct": "Pressing %",
        "passes_accurate_pct": "Passzpontosság %",
        "counterattacks": "Kontrák",
        "recoveries": "Labdaszerzések",
        "lost_balls": "Labdavesztések",
        "crosses": "Beadások",
    }
    parts = []
    for k, lab in label_map.items():
        v = metrics.get(k)
        if v not in [None, 0, 0.0, ""]:
            try:
                fv = _fpi_normalized_tactical_metric_v79(k, float(v))
                if k in ["possession_pct", "pressing_success_pct", "passes_accurate_pct"]:
                    parts.append(f"{lab}: {fv:.1f}%")
                else:
                    parts.append(f"{lab}: {fv:.1f}")
            except Exception:
                parts.append(f"{lab}: {v}")
    return " | ".join(parts[:10]) if parts else "Nincs kiemelkedő taktikai KPI."


def _build_tactical_executive_context(gps_context: Dict[str, object], tactical_ctx: Dict[str, object], plan: Dict[str, object]) -> Dict[str, object]:
    own = tactical_ctx.get("own", {}) if tactical_ctx else {}
    opp = tactical_ctx.get("opponent", {}) if tactical_ctx else {}
    analysis_level = tactical_ctx.get("analysis_level_label", "GPS Only")
    return {
        "version": TACTICAL_PRO_VERSION,
        "analysis_level": analysis_level,
        "has_own_pdf": bool((own.get("pdf_insights") or {}).get("pdf_uploaded") or (own.get("pdf_insights") or {}).get("raw_text_len", 0)),
        "has_opp_pdf": bool((opp.get("pdf_insights") or {}).get("pdf_uploaded") or (opp.get("pdf_insights") or {}).get("raw_text_len", 0)),
        "own_pdf_pages": int((own.get("pdf_insights") or {}).get("pdf_pages", 0) or 0),
        "opp_pdf_pages": int((opp.get("pdf_insights") or {}).get("pdf_pages", 0) or 0),
        "has_own_team_excel": bool(own.get("team_metrics")),
        "has_opp_team_excel": bool(opp.get("team_metrics")),
        "has_own_player_excel": bool(own.get("player_tables")),
        "has_opp_player_excel": bool(opp.get("player_tables")),
        "own_topics": ((own.get("pdf_insights") or {}).get("topics") or [])[:8],
        "opp_topics": ((opp.get("pdf_insights") or {}).get("topics") or [])[:8],
        "own_pdf_chars": int((own.get("pdf_insights") or {}).get("raw_text_len", 0) or 0),
        "opp_pdf_chars": int((opp.get("pdf_insights") or {}).get("raw_text_len", 0) or 0),
        "own_pdf_reader": str((own.get("pdf_insights") or {}).get("reader_version", "")),
        "opp_pdf_reader": str((opp.get("pdf_insights") or {}).get("reader_version", "")),
        "own_team_metrics": own.get("team_metrics", {}),
        "opp_team_metrics": opp.get("team_metrics", {}),
        "plan_a": plan.get("plan_a", "KIE – Kiegyensúlyozott"),
        "risks": plan.get("risks", []),
        "md_plan": plan.get("md_plan", []),
        "player_focus": plan.get("player_focus", []),
        "tactical_findings": plan.get("tactical_findings", []),
        "team_comparison": plan.get("team_comparison", []),
        "own_player_tables": (tactical_ctx or {}).get("own", {}).get("player_tables", {}),
        "opp_player_tables": (tactical_ctx or {}).get("opponent", {}).get("player_tables", {}),
        "opponent_player_evaluation": _fpi_build_player_evaluation_v132((tactical_ctx or {}).get("opponent", {}).get("player_tables", {}), side="opp", max_rows=9) if "_fpi_build_player_evaluation_v132" in globals() else [],
        "own_player_evaluation": _fpi_build_player_evaluation_v132((tactical_ctx or {}).get("own", {}).get("player_tables", {}), side="own", max_rows=9) if "_fpi_build_player_evaluation_v132" in globals() else [],
        "pdf_provider_lines": ((tactical_ctx.get("pdf_insights") or {}).get("sportsbase_lines", []) or []),
        "pdf_provider_findings": ((tactical_ctx.get("pdf_insights") or {}).get("sportsbase_findings", []) or []),
        "pdf_direct_findings_count": len(((tactical_ctx.get("pdf_insights") or {}).get("sportsbase_findings", []) or [])),
        "pdf_direct_lines_count": len(((tactical_ctx.get("pdf_insights") or {}).get("sportsbase_lines", []) or [])),
    }

def render_tactical_pro_module(gps_context: Dict[str, object]) -> None:
    st.markdown("## 🧠 Tactical Pro+ / Adaptive Intelligence")
    st.markdown(
        """
        <div class="tactical-readable-box">
        <b>Hogyan működik?</b><br>
        GPS-alapon önállóan is működik. Ha saját csapatról és/vagy ellenfélről taktikai PDF-et,
        csapat Excelt vagy játékos Excelt töltesz fel, azokat beépíti a meccstervbe és a heti edzésterv-javaslatba.
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander("📥 Taktikai inputok – saját csapat és ellenfél", expanded=True):
        own_col, opp_col = st.columns(2)
        with own_col:
            st.markdown("### Saját csapat")
            own_pdfs, own_pdf_state_v92 = _fpi_pdf_uploader_v92("Saját taktikai PDF-ek", "own", "tactical_pro_own_pdfs_v92")
            own_team_xlsx = st.file_uploader("Saját csapatstatisztika Excel", type=["xlsx", "xls"], key="tactical_pro_own_team_xlsx")
            own_player_xlsx = st.file_uploader("Saját játékosstatisztika Excel", type=["xlsx", "xls"], key="tactical_pro_own_player_xlsx")
        with opp_col:
            st.markdown("### Ellenfél")
            opp_pdfs, opp_pdf_state_v92 = _fpi_pdf_uploader_v92("Ellenfél taktikai PDF-ek", "opp", "tactical_pro_opp_pdfs_v92")
            opp_team_xlsx = st.file_uploader("Ellenfél csapatstatisztika Excel", type=["xlsx", "xls"], key="tactical_pro_opp_team_xlsx")
            opp_player_xlsx = st.file_uploader("Ellenfél játékosstatisztika Excel", type=["xlsx", "xls"], key="tactical_pro_opp_player_xlsx")

        # V8.8: a PDF-eket azonnal byte-ként is eltároljuk, hogy exportkor ne vesszen el
        # a feltöltött fájlobjektum tartalma / ne legyen 0 oldalas PDF context.
        # V9.2: PDF state-et már a _fpi_pdf_uploader_v92 kezeli.
        # Nem írjuk felül / nem töröljük itt, mert ez okozhatta a 0 fájlos állapotot.

        st.caption("Tipp: a feltöltött fájl törléséhez használd a fájlnév melletti kis X-et. Ha nem látszik, ez a verzió javítja a kontrasztot. Teljes resethez frissítsd az oldalt vagy használd az alábbi gombot.")
        if st.button("🧹 Tactical feltöltések / mapping reset", key="tactical_pro_reset_upload_mapping"):
            for k in list(st.session_state.keys()):
                if str(k).startswith(("own_team_tactical", "opp_team_tactical", "own_player_tactical", "opp_player_tactical", "tactical_pro_context", "tactical_pro_own_pdf_bytes_v88", "tactical_pro_opp_pdf_bytes_v88", "tactical_pro_own_pdf_text_store_v89", "tactical_pro_opp_pdf_text_store_v89", "tactical_pro_own_pdf_upload_manager_v92", "tactical_pro_opp_pdf_upload_manager_v92")):
                    st.session_state.pop(k, None)
            st.success("A taktikai mapping/session állapot törölve. A fájlok törléséhez szükség esetén frissítsd az oldalt.")

    has_gps = bool(gps_context.get("has_gps", True))
    has_pdf = bool(opp_pdfs or own_pdfs)
    has_team_excel = opp_team_xlsx is not None or own_team_xlsx is not None
    has_player_excel = opp_player_xlsx is not None or own_player_xlsx is not None
    level, level_label = _fpi_analysis_level(has_gps, has_pdf, has_team_excel, has_player_excel)
    st.markdown(
        f"""
        <div class="tactical-readable-box">
        <b>Elemzési szint:</b> Level {level}<br>
        <b>Mit jelent:</b> {level_label}
        </div>
        """,
        unsafe_allow_html=True,
    )

        # V9.2: minden PDF feldolgozás a stabil upload manager state-ből jön.
    own_pdf_state_v92 = _fpi_get_pdf_upload_state_v92("own")
    opp_pdf_state_v92 = _fpi_get_pdf_upload_state_v92("opp")

    own_pdf_text = own_pdf_state_v92.get("text", "")
    opp_pdf_text = opp_pdf_state_v92.get("text", "")
    own_pdf_pages = own_pdf_state_v92.get("pages", [])
    opp_pdf_pages = opp_pdf_state_v92.get("pages", [])

    own_pdf_uploaded = bool(own_pdf_state_v92.get("has_files"))
    opp_pdf_uploaded = bool(opp_pdf_state_v92.get("has_files"))

    own_pdf_insights = _fpi_safe_tactical_pdf_insights_v105(own_pdf_text, uploaded=own_pdf_uploaded, pages=own_pdf_pages)
    opp_pdf_insights = _fpi_safe_tactical_pdf_insights_v105(opp_pdf_text, uploaded=opp_pdf_uploaded, pages=opp_pdf_pages)
    own_pdf_insights["pdf_uploaded"] = own_pdf_uploaded
    own_pdf_insights["pdf_pages"] = len([p for p in own_pdf_pages if p.get("has_text") or p.get("text")])
    opp_pdf_insights["pdf_uploaded"] = opp_pdf_uploaded
    opp_pdf_insights["pdf_pages"] = len([p for p in opp_pdf_pages if p.get("has_text") or p.get("text")])
    merged_pdf_insights = _merge_tactical_pdf_insights(own_pdf_insights, opp_pdf_insights)

    own_team_metrics, opp_team_metrics = {}, {}
    own_player_tables, opp_player_tables = {}, {}

    if own_team_xlsx is not None:
        own_team_df, own_team_mapping = _fpi_tactical_mapper_ui(own_team_xlsx, TACTICAL_TEAM_ALIASES_FPI, "own_team_tactical", "Saját csapat Excel")
        own_team_metrics = _fpi_safe_tactical_parse_team_excel_v107(own_team_df, own_team_mapping)

    if opp_team_xlsx is not None:
        opp_team_df, opp_team_mapping = _fpi_tactical_mapper_ui(opp_team_xlsx, TACTICAL_TEAM_ALIASES_FPI, "opp_team_tactical", "Ellenfél csapat Excel")
        opp_team_metrics = _fpi_safe_tactical_parse_team_excel_v107(opp_team_df, opp_team_mapping)

    if own_player_xlsx is not None:
        own_player_df, own_player_mapping = _fpi_tactical_mapper_ui(own_player_xlsx, TACTICAL_PLAYER_ALIASES_FPI, "own_player_tactical", "Saját játékos Excel")
        own_player_tables = _fpi_safe_tactical_parse_player_excel_v107(own_player_df, own_player_mapping)

    if opp_player_xlsx is not None:
        opp_player_df, opp_player_mapping = _fpi_tactical_mapper_ui(opp_player_xlsx, TACTICAL_PLAYER_ALIASES_FPI, "opp_player_tactical", "Ellenfél játékos Excel")
        opp_player_tables = _fpi_safe_tactical_parse_player_excel_v107(opp_player_df, opp_player_mapping)

    # A régi plan-motor ellenfél fókuszú volt. Megtartjuk, de kibővített, merge-elt PDF insighttal és ellenfél KPI-okkal etetjük.
    tactical_ctx_for_plan = {
        "analysis_level_label": level_label,
        "pdf_insights": merged_pdf_insights,
        "team_metrics": opp_team_metrics,
        "player_tables": opp_player_tables,
        "own": {"pdf_insights": own_pdf_insights, "team_metrics": own_team_metrics, "player_tables": own_player_tables},
        "opponent": {"pdf_insights": opp_pdf_insights, "team_metrics": opp_team_metrics, "player_tables": opp_player_tables},
    }
    plan = _fpi_build_adaptive_match_training_plan(gps_context, tactical_ctx_for_plan)
    executive_ctx = _build_tactical_executive_context(gps_context, tactical_ctx_for_plan, plan)
    if "_fpi_enrich_tactical_context_v132" in globals():
        executive_ctx = _fpi_enrich_tactical_context_v132(executive_ctx, own_player_tables, opp_player_tables)
    st.session_state["tactical_pro_context"] = executive_ctx

    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Adaptive szint", f"Level {level}")
    k2.metric("Saját PDF oldalak", len([p for p in own_pdf_pages if p.get("has_text")]))
    k3.metric("Ellenfél PDF oldalak", len([p for p in opp_pdf_pages if p.get("has_text")]))
    k4.metric("Taktikai KPI-k", len([v for v in {**own_team_metrics, **opp_team_metrics}.values() if v not in [0, 0.0, None]]))

    st.caption(
        f"PDF reader státusz: saját {len(own_pdf_text or '')} karakter / {len([p for p in own_pdf_pages if p.get('has_text') or p.get('text')])} oldal; "
        f"ellenfél {len(opp_pdf_text or '')} karakter / {len([p for p in opp_pdf_pages if p.get('has_text') or p.get('text')])} oldal."
    )
    if not own_pdf_uploaded and not opp_pdf_uploaded:
        st.warning("Jelenleg a Tactical Pro+ PDF feltöltő üres (0 fájl). A V9.2-ben a PDF feltöltés külön stabil managerrel fut; töltsd fel újra a PDF-et itt, és azonnal látnod kell a PDF OK sort.")

    if own_pdf_uploaded and not own_pdf_text:
        st.warning("Saját PDF feltöltve, de nem sikerült szöveget kinyerni belőle. Valószínűleg képalapú/scannelt PDF, ezért a taktikai témák nem jelennek meg.")
    if opp_pdf_uploaded and not opp_pdf_text:
        st.warning("Ellenfél PDF feltöltve, de nem sikerült szöveget kinyerni belőle. Valószínűleg képalapú/scannelt PDF, ezért a taktikai témák nem jelennek meg.")

    with st.expander("🔎 PDF diagnosztika – mit sikerült kiolvasni?", expanded=False):
        diag_rows = []
        for side, pages, insights in [
            ("Saját", own_pdf_pages, own_pdf_insights),
            ("Ellenfél", opp_pdf_pages, opp_pdf_insights),
        ]:
            real_pages = [p for p in pages if isinstance(p.get("page"), int)]
            best_rows = [p for p in pages if p.get("page") == "BEST"]
            best_reader = best_rows[-1].get("reader", "n.a.") if best_rows else "n.a."
            best_chars = best_rows[-1].get("chars", 0) if best_rows else 0
            diag_rows.append({
                "Oldal": side,
                "Fájl feltöltve": "igen" if (side == "Saját" and own_pdf_uploaded) or (side == "Ellenfél" and opp_pdf_uploaded) else "nem",
                "Valós PDF oldalak": len({(p.get("file"), p.get("page")) for p in real_pages if isinstance(p.get("page"), int)}),
                "Legjobb reader": best_reader,
                "Legjobb karakter": best_chars,
                "V8.8 byte-store fájlok": len(st.session_state.get("tactical_pro_own_pdf_bytes_v88", []) if side == "Saját" else st.session_state.get("tactical_pro_opp_pdf_bytes_v88", [])),
                "Riportba kerülő karakter": int(insights.get("raw_text_len", 0) or 0),
                "Felismert témák": len(insights.get("topics", []) or []),
            })
        st.dataframe(pd.DataFrame(diag_rows), use_container_width=True, hide_index=True)

        with st.expander("Reader-részletek oldalanként", expanded=False):
            all_reader_rows = []
            for side, pages in [("Saját", own_pdf_pages), ("Ellenfél", opp_pdf_pages)]:
                for p in pages:
                    all_reader_rows.append({
                        "Oldal": side,
                        "Fájl": p.get("file"),
                        "PDF oldal": p.get("page"),
                        "Reader": p.get("reader"),
                        "Karakter": p.get("chars"),
                        "Van szöveg": p.get("has_text"),
                        "Hiba": p.get("error", ""),
                    })
            if all_reader_rows:
                st.dataframe(pd.DataFrame(all_reader_rows), use_container_width=True, hide_index=True)

        d1, d2 = st.columns(2)
        with d1:
            st.markdown("**Saját PDF előnézet**")
            st.text_area("Saját PDF – első kinyert szövegrész", own_pdf_insights.get("preview", ""), height=180, key="own_pdf_preview_v76")
            if own_pdf_insights.get("top_terms"):
                st.dataframe(pd.DataFrame(own_pdf_insights.get("top_terms", [])), use_container_width=True, hide_index=True)
        with d2:
            st.markdown("**Ellenfél PDF előnézet**")
            st.text_area("Ellenfél PDF – első kinyert szövegrész", opp_pdf_insights.get("preview", ""), height=180, key="opp_pdf_preview_v76")
            if opp_pdf_insights.get("top_terms"):
                st.dataframe(pd.DataFrame(opp_pdf_insights.get("top_terms", [])), use_container_width=True, hide_index=True)

    st.markdown("### Taktikai stratégiai paletta")
    st.caption("A kódok jelentése, hogy a javasolt Plan A / Plan B ne legyen félreérthető.")
    try:
        st.dataframe(pd.DataFrame(_strategy_palette_pdf_rows(), columns=["Kód", "Stratégia", "Egyszerű jelentés"]), use_container_width=True, hide_index=True)
    except Exception:
        pass


    if executive_ctx.get("strategy_framework"):
        sfw = executive_ctx.get("strategy_framework") or {}
        st.markdown("### 0. Tactical Framework - 7 dimenzió és stratégiai profil")
        cfw1, cfw2, cfw3 = st.columns(3)
        with cfw1:
            st.metric("Elsődleges profil", f"{sfw.get('primary_code', '')} - {sfw.get('primary_name', '')}", f"{sfw.get('primary_pct', '')}%")
        with cfw2:
            st.metric("Alternatív profil", f"{sfw.get('secondary_code', '')} - {sfw.get('secondary_name', '')}", f"{sfw.get('secondary_pct', '')}%")
        with cfw3:
            st.metric("Javasolt alapirány", str(executive_ctx.get("plan_a", ""))[:42])
        try:
            st.dataframe(pd.DataFrame(_fpi_strategy_framework_to_rows_v129(sfw)), use_container_width=True, hide_index=True)
        except Exception:
            pass
        focus = sfw.get("detail_focus") or []
        if focus:
            st.markdown("**Meccsspecifikus differenciálás:** " + "; ".join(str(x) for x in focus))

    st.markdown("### 1. Match Plan AI – javasolt meccsterv")
    st.markdown(f"**Plan A:** {plan['plan_a']}")
    st.markdown("**Fő kockázatok / fókuszok:**")
    for r in plan["risks"]:
        st.markdown(f"- {r}")

    st.markdown("### 2. Taktikai Excel + PDF következtetések")
    if executive_ctx.get("tactical_findings"):
        st.dataframe(pd.DataFrame(executive_ctx.get("tactical_findings")), use_container_width=True, hide_index=True)
    else:
        st.caption("Nincs még értelmezhető taktikai következtetés. Ellenőrizd a Team/Player Excel mappinget.")

    if executive_ctx.get("pdf_provider_lines"):
        with st.expander("✅ PDF-ből konkrétan kinyert SportsBase / provider adatok", expanded=True):
            for line in executive_ctx.get("pdf_provider_lines", [])[:16]:
                st.markdown(f"- {line}")
    elif own_pdf_uploaded or opp_pdf_uploaded:
        st.error("PDF fel van töltve, de ebből a verzióból sem jött ki konkrét provider/KPI adat. Ilyenkor a PDF diagnosztikában az első 2000 karaktert és a reader-részleteket kell nézni.")

    st.markdown("### 3. Saját vs ellenfél gyors összevetés")
    comp_rows = [
        {"Oldal": "Saját csapat", "PDF témák": len(own_pdf_insights.get("topics", []) or []), "Csapat KPI": _tactical_key_numbers_summary(own_team_metrics)},
        {"Oldal": "Ellenfél", "PDF témák": len(opp_pdf_insights.get("topics", []) or []), "Csapat KPI": _tactical_key_numbers_summary(opp_team_metrics)},
    ]
    st.dataframe(pd.DataFrame(comp_rows), use_container_width=True, hide_index=True)

    st.markdown("### 4. Integrált mikrociklus – erőnléti + taktikai cél")
    try:
        md_df = pd.DataFrame(_combined_md_rows(executive_ctx), columns=["Nap", "Erőnléti cél", "Taktikai cél", "Indoklás"])
    except Exception:
        md_df = pd.DataFrame(plan["md_plan"], columns=["Nap", "Taktikai cél", "Indoklás"])
    st.dataframe(md_df, use_container_width=True, hide_index=True)

    st.markdown("### 5. Játékosszintű/taktikai fókusz")
    if plan["player_focus"]:
        for p in plan["player_focus"]:
            st.markdown(f"- {p}")
    else:
        st.caption("Nincs külön játékos Excel vagy kiemelt játékos. GPS-alapú monitoring marad aktív.")

    st.markdown("### 6. PDF-ből felismert taktikai témák")
    tcol1, tcol2 = st.columns(2)
    with tcol1:
        st.markdown("**Saját csapat PDF témák**")
        if own_pdf_insights.get("topics"):
            st.dataframe(pd.DataFrame(own_pdf_insights["topics"]), use_container_width=True)
        else:
            st.caption("Nincs saját taktikai PDF vagy nem volt felismerhető szöveg.")
    with tcol2:
        st.markdown("**Ellenfél PDF témák**")
        if opp_pdf_insights.get("topics"):
            st.dataframe(pd.DataFrame(opp_pdf_insights["topics"]), use_container_width=True)
        else:
            st.caption("Nincs ellenfél taktikai PDF vagy nem volt felismerhető szöveg.")

    with st.expander("PDF szövegkörnyezetek témánként – saját + ellenfél"):
        for key, lines in (merged_pdf_insights.get("blocks") or {}).items():
            if lines:
                st.markdown(f"**{TACTICAL_TOPIC_TAGS_FPI.get(key, {}).get('label', key)}**")
                for line in lines[:8]:
                    st.caption(line)

    export_payload = {
        "version": TACTICAL_PRO_VERSION,
        "analysis_level": level_label,
        "plan_a": plan["plan_a"],
        "risks": plan["risks"],
        "md_plan": [{"Nap": a, "Fókusz": b, "Miért": c} for a, b, c in plan["md_plan"]],
        "player_focus": plan["player_focus"],
        "own_detected_topics": own_pdf_insights.get("topics", []),
        "opp_detected_topics": opp_pdf_insights.get("topics", []),
        "own_team_metrics": own_team_metrics,
        "opp_team_metrics": opp_team_metrics,
    }
    st.download_button("⬇️ Tactical Pro+ JSON export", data=json.dumps(export_payload, ensure_ascii=False, indent=2).encode("utf-8"), file_name="fpi_tactical_pro_plus_export.json", mime="application/json", use_container_width=True)




# =========================================================
# V130 - Final Methodology tab override: visible 9 strategies + benchmark browser
# =========================================================
def _fpi_strategy_framework_to_rows_v130(framework: Dict[str, object]) -> List[Dict[str, object]]:
    dims = framework.get("dimensions", {}) if isinstance(framework, dict) else {}
    return [{"Dimenzió": label, "Érték (0-10)": dims.get(key, "n.a.")} for key, label in (globals().get("FPI_TACTICAL_DIMENSIONS_V129") or FPI_TACTICAL_DIMENSIONS_V130)]

def render_methodology_tab() -> None:
    st.markdown("## 📚 Football Performance Intelligence metodika")
    st.markdown("""
    <div class="fpi-summary-card">
    <h3>Metodikai alapelv</h3>
    <p>A <b>Football Performance Intelligence</b> döntéstámogató rendszer. GPS-adatokból és opcionálisan taktikai PDF/Excel inputokból készít readiness-, risk-, benchmark-, mikrociklus- és Tactical Framework alapú vezetői információt.</p>
    <p>A metodika <b>nemzetközi sporttudományi szakirodalomra, saját adatokra, gyakorlati teljesítménydiagnosztikai tapasztalatokra és saját fejlesztésű elemzési logikára</b> épül. Nem orvosi diagnózis és nem automatikus döntés, hanem szakmai döntéstámogatás.</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### 1. Readiness Score")
    st.write("A játékos/csapat aktuális terhelési állapotának becslése. Figyelembe veszi a heti Load trendet, az elmúlt 3-7 nap terhelését, a 4 hetes mintázatot, a HSR/sprint expozíciót, a High Efforts jellegű intenzív mozgásokat, a taper/frissítés logikát és - ha van - pulzus/HRV jellegű belső terhelési adatokat.")
    st.markdown("**Értelmezés:** 80-100 magas készenlét; 60-79 elfogadható; 40-59 csökkent; 0-39 alacsony. Az alacsonyabb readiness nem csak túlterhelést, hanem alulterhelést vagy kedvezőtlen terhelési mintázatot is jelenthet.")

    st.markdown("### 2. Player Risk Score")
    st.write("A risk score játékosonként számolódik, és elsősorban a saját előző 4 hetes mintázathoz viszonyít. Nem csak az abszolút magas érték lehet rizikó, hanem a szokatlan emelkedés, visszaesés, hiányzó sebességi expozíció vagy több kedvezőtlen jel együttes megjelenése is.")
    st.markdown("**Kategóriák:** alacsony <45; közepes 45-69; magas >=70.")

    st.markdown("### 3. Benchmark Engine")
    st.write("A referencia nem egyetlen fix benchmark. A rendszer korosztály, bajnoki szint, játékosposzt és játékmodell alapján módosított célzónákat használ: össztáv, Load, HSR, sprint táv, sprint darabszám és High Efforts mutatókra.")
    with st.expander("🔎 Benchmark böngésző / szűrő", expanded=True):
        ages = globals().get("FPI_AGE_GROUPS_V112", ["Felnőtt", "U21", "U19", "U17", "U16", "U15", "U14", "U13"])
        levels = globals().get("FPI_COMPETITION_LEVELS_V112", ["NB I", "NB II", "NB III", "Akadémia", "Regionális", "Megye I", "Egyéb"])
        positions = globals().get("FPI_POSITIONS_V112", ["Kapus", "Középhátvéd", "Szélső hátvéd", "Védekező középpályás", "Középpályás", "Támadó középpályás", "Szélső", "Csatár"])
        playmodels = globals().get("FPI_PLAYMODEL_OPTIONS_V112", ["Dominancia", "Magas presszing", "Átmeneti játék", "Direkt játék", "Kiegyensúlyozott"])
        metrics = ["Összes mutató", "Össztáv", "Load / terhelési pont", "HSR / nagysebességű futás", "Sprint táv", "Sprint darabszám", "High Efforts"]
        c1, c2, c3, c4, c5 = st.columns(5)
        with c1: age = st.selectbox("Korosztály", ages, index=0, key="method_age_v130")
        with c2: level = st.selectbox("Bajnoki szint", levels, index=1 if "NB II" in levels else 0, key="method_level_v130")
        with c3: pos = st.selectbox("Poszt", positions, index=min(4, len(positions)-1), key="method_pos_v130")
        with c4: play = st.selectbox("Játékmodell", playmodels, index=playmodels.index("Kiegyensúlyozott") if "Kiegyensúlyozott" in playmodels else 0, key="method_play_v130")
        with c5: metric = st.selectbox("Mutató", metrics, index=0, key="method_metric_v130")
        try:
            dfb = _fpi_benchmark_browser_df_v124(age, level, pos, play, metric)
            _fpi_render_benchmark_browser_table_v125(dfb)
        except Exception as e:
            st.warning(f"A benchmark böngésző nem tudott megjelenni: {e}")

    st.markdown("### 4. Mikrociklus Motor")
    st.write("A mikrociklus motor MD-napokban gondolkodik. Figyelembe veszi a hét típusát, a meccsnapot, az edzésszámot, a pihenőnapokat, a readiness értéket, a risk jelzéseket, a benchmark eltéréseket és a játékmodellt.")

    st.markdown("### 5. Tactical Framework - 7 dimenzióból 9 stratégiai profil")
    st.write("A Tactical Pro+ a taktikai inputokat egységes keretrendszerbe rendezi. Először 7 dimenzió mentén profilt képez, majd ezt 9 stratégiai profilhoz viszonyítja. A rendszer nem feltétlenül csak egy címkét ad: elsődleges profil, alternatív profil és meccsspecifikus fókuszok is készülhetnek.")
    dim_rows = [{"Dimenzió": label, "Mit jelent?": desc} for label, desc in [
        ("Letámadás", "presszing, nyomás, labdaszerzés magassága"),
        ("Labdakihozatal", "építkezés, progresszió, első/második fázis"),
        ("Átmenetek", "támadó és védekező transition, rest defense"),
        ("Támadó játék", "helyzetkialakítás, box entry, szélső/félterületi játék"),
        ("Pontrúgások", "szöglet, szabadrúgás, második labdák"),
        ("Labdabirtoklás", "kontroll, dominancia, labdajáratás"),
        ("Lövésprofil", "lövésszám, xG, helyzetminőség"),
    ]]
    # V131: st.dataframe egyes böngésző/CSS kombinációkban üres fehér dobozként jelent meg.
    # Ezért a metodikai táblázatokat statikus, olvasható HTML táblaként rendereljük.
    def _fpi_methodology_html_table_v131(rows, headers):
        def esc(x):
            return html.escape(str(x if x is not None else ""))
        head = "".join(f"<th>{esc(h)}</th>" for h in headers)
        body = ""
        for r in rows:
            if isinstance(r, dict):
                vals = [r.get(h, "") for h in headers]
            else:
                vals = list(r)
            body += "<tr>" + "".join(f"<td>{esc(v)}</td>" for v in vals) + "</tr>"
        return f"""
        <div class="fpi-method-table-wrap">
          <table class="fpi-method-table">
            <thead><tr>{head}</tr></thead>
            <tbody>{body}</tbody>
          </table>
        </div>
        """

    st.markdown(_fpi_methodology_html_table_v131(dim_rows, ["Dimenzió", "Mit jelent?"]), unsafe_allow_html=True)
    st.markdown("**Két fő tengely:** játékstílus: direkt → vegyes → kiegyensúlyozott → kontroll → agresszív; blokkmagasság: mély → alacsony-közép → közép → közép-magas → magas.")
    strat_rows = _fpi_strategy_palette_rows_any_v130()
    st.markdown(_fpi_methodology_html_table_v131(strat_rows, ["Kód", "Stratégia", "Játékstílus", "Blokkmagasság", "Jelentés"]), unsafe_allow_html=True)
    st.info("A 9 stratégia nem lezárt címke. A gyakorlatban a Football Performance Intelligence elsődleges profilt, alternatív profilt és konkrét meccsfókuszokat ad. Példa: BAT elsődleges, POZ/KIE alternatív profillal, rest defense és átmeneti fókuszokkal.")

    st.markdown("### 6. Értelmezési korlát")
    st.error("A Football Performance Intelligence nem állítja, hogy egy pontszám önmagában megmondja a sérülést, a teljesítményt vagy a mérkőzés kimenetelét. A cél: gyorsabb, strukturáltabb szakmai döntéstámogatás.")
    st.markdown("### 7. Technikai státusz")
    st.json({"Import/Export oldal": "fő munkafolyamat", "Readiness Engine": "trend + terhelés + sebességi expozíció + játékmodell", "Risk Engine": "saját 4 hetes előzmény + aktuális heti eltérés", "Benchmark Engine": "korosztály + szint + poszt + játékmodell", "Tactical Framework": "7 dimenzió -> 9 stratégiai profil -> alternatív profil + meccsfókusz", "Tactical Pro+": "opcionális inputként aktív"})

# Final signal override after all original definitions.
def _fpi_has_tactical_signal_v95(tactical_context: Optional[Dict[str, object]]) -> bool:
    return _fpi_tactical_context_has_pdf_v130(tactical_context) or bool(tactical_context and (tactical_context.get("strategy_framework") or tactical_context.get("has_own_team_excel") or tactical_context.get("has_opp_team_excel") or tactical_context.get("has_own_player_excel") or tactical_context.get("has_opp_player_excel") or tactical_context.get("tactical_findings") or tactical_context.get("own_topics") or tactical_context.get("opp_topics") or tactical_context.get("pdf_provider_lines") or tactical_context.get("pdf_provider_findings") or tactical_context.get("team_comparison")))


# Tabok
tab_exec, tab_intro, tab1, tab_premium, tab_export, tab_methodology, tab_tactical_pro, tab_intel, tab_micro, tab_risk, tab2, tab3, tab4, tab5 = st.tabs([
    "🏠 Dashboard",
    "ℹ️ Rendszer",
    "📌 Áttekintő",
    "🎛️ Cockpit",
    "📄 Export",
    "📚 Metodika",
    "🧠 Tactical Pro+",
    "🧠 Intelligence",
    "📅 Mikrociklus",
    "🚨 Kockázat",
    "🎯 Javaslatok",
    "👤 Játékosok",
    "✅ Adatminőség",
    "🧾 Nyers adatok",
])

with tab_methodology:
    render_methodology_tab()

with tab_tactical_pro:
    tactical_gps_context = {
        "has_gps": True,
        "selected_week": selected_week if 'selected_week' in globals() else None,
        "readiness_score": readiness_score if 'readiness_score' in globals() else 70,
        "playstyle": selected_playstyle if 'selected_playstyle' in globals() else "Kiegyensúlyozott",
        "priorities": coaching_priorities if 'coaching_priorities' in globals() else [],
        "periodization_type": periodization_type if 'periodization_type' in globals() else "n.a.",
    }
    st.subheader("🧠 Tactical Pro+")
    st.caption("V115: a taktikai PDF/Excel import kizárólag az Import / export oldalon van. Itt már csak az ott betöltött kontextus összefoglalója látszik.")
    clean_tctx_v115 = st.session_state.get("fpi_clean_tactical_context_v115")
    if clean_tctx_v115:
        level = clean_tctx_v115.get("analysis_level") or clean_tctx_v115.get("level_label") or "Tactical context aktív"
        topics = clean_tctx_v115.get("topics") or clean_tctx_v115.get("opp_detected_topics") or []
        st.success(f"Aktív Tactical Pro+ kontextus: {level}")
        if topics:
            st.markdown("**Felismerhető taktikai témák:** " + ", ".join([str(x) for x in topics[:10]]))
        with st.expander("Tactical context technikai összefoglaló", expanded=False):
            st.json({k: v for k, v in clean_tctx_v115.items() if k not in ["raw_text", "df", "own_team_df", "opp_team_df", "own_player_df", "opp_player_df"]})
    else:
        st.info("Nincs betöltött taktikai kontextus. Töltsd fel a Tactical Pro+ fájlokat az Import / export oldalon. GPS-only módban ez teljesen rendben van.")




with tab_exec:
    st.markdown(
        """
        <div class="hero-box">
            <div class="hero-title">⭐ Vezetői riport</div>
            <div class="hero-sub">
                Egyetlen oldal vezetőedzőnek / sportigazgatónak: meccskészültség, fő kockázatok,
                top edzői teendők és exportálható vezetői csomag.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    mode_col1, mode_col2, mode_col3 = st.columns([1, 1, 2])
    with mode_col1:
        st.metric("Mód", "PRO" if is_pro_mode() else "DEMO")
    with mode_col2:
        st.metric("Hét", format_week_label(selected_week))
    with mode_col3:
        if not is_pro_mode():
            st.caption("DEMO mód – Pro exportok korlátozva.")

    k1, k2, k3, k4, k5 = st.columns(5)
    with k1:
        st.metric("Meccskészültség", f"{readiness_score}/100", score_to_label(readiness_score))
    with k2:
        st.metric("Magas risk", high_risk_count)
    with k3:
        st.metric("Közepes risk", medium_risk_count)
    with k4:
        st.metric("Insight", len(all_insights))
    with k5:
        mem_weeks = analysis_base_df["week"].nunique() if "week" in analysis_base_df.columns else 0
        st.metric("Elemzett hetek", mem_weeks)

    st.markdown("### Heti vezetői összefoglaló")
    st.info(week_status_info.get("message", ""))
    render_weekly_summary_card(weekly_summary_text)

    with st.expander("🧭 Múlt / aktuális / jövő fókusz", expanded=False):
        st.markdown("**Múlt hét -> erre a hétre**")
        render_weekly_summary_card(past_week_review_text)
        st.markdown("**Aktuális hét -> hátralévő napok**")
        render_weekly_summary_card(current_remaining_text)
        st.markdown("**Jövő hét**")
        render_weekly_summary_card(next_week_plan_text)

    left, right = st.columns([1.05, 1])
    with left:
        st.markdown("### Top 3 edzői teendő")
        render_coaching_priorities(coaching_priorities)

    with right:
        st.markdown("### Játékos risk gyorsnézet")
        if player_risk_df is not None and not player_risk_df.empty:
            show_cols = [c for c in ["Játékos", "Kockázati szint", "Risk score", "Kockázati pontszám", "Fő ok", "Fő okok"] if c in player_risk_df.columns]
            risk_view_v120 = player_risk_df[show_cols].head(8) if show_cols else player_risk_df.head(8)
            fpi_compact_table(risk_view_v120, max_rows=8)
        else:
            fpi_empty_state("Nincs játékos risk adat", "A risk gyorsnézet akkor jelenik meg, ha van elég heti játékos- és terhelésadat.", "🛡️")

    st.markdown("### Letölthető Football Performance Intelligence riportok")
    st.caption("A régi vezetői PDF/Word/Excel/CSV gombok kikerültek. Itt csak a termékriportok maradnak: GPS-only, Executive Summary és Full Report.")
    safe_week_main = _safe_filename_week(selected_week)

    gps_only_live_pdf_main = build_fpi_gps_only_pdf_bytes(analysis_base_df.copy(), selected_week, selected_playstyle)
    if gps_only_live_pdf_main is not None:
        st.download_button(
            "⬇️ GPS-only PDF riport",
            data=gps_only_live_pdf_main,
            file_name=f"fpi_gps_only_report_{safe_week_main}.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="download_gps_only_live_v96_main",
        )


    st.markdown("### Football Performance Intelligence riportok – éles export")
    st.caption("Három kimenet: GPS-only Report, Executive Summary és Full Report.")
    live_report_base = analysis_base_df.copy()
    lr1, lr2, lr3 = st.columns(3)
    with lr1:
        gps_only_pack_pdf = build_fpi_gps_only_pdf_bytes(live_report_base, selected_week, selected_playstyle)
        if gps_only_pack_pdf is not None:
            st.download_button(
                "⬇️ GPS-only PDF Report",
                data=gps_only_pack_pdf,
                file_name=f"fpi_gps_only_report_{safe_week_main}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="download_fpi_gps_only_v96",
            )
    with lr2:
        exec_pack_pdf = build_fpi_product_pdf_bytes(
            live_report_base,
            selected_week,
            selected_playstyle,
            report_type="executive",
            tactical_context=_fpi_context_for_export_v87(tactical_gps_context if "tactical_gps_context" in globals() else {}),
        )
        if exec_pack_pdf is not None:
            st.download_button(
                "⬇️ Executive Summary PDF",
                data=exec_pack_pdf,
                file_name=f"fpi_executive_summary_{safe_week_main}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="download_fpi_exec_v83",
            )
    with lr3:
        full_pack_pdf = build_fpi_product_pdf_bytes(
            live_report_base,
            selected_week,
            selected_playstyle,
            report_type="full",
            tactical_context=_fpi_context_for_export_v87(tactical_gps_context if "tactical_gps_context" in globals() else {}),
        )
        if full_pack_pdf is not None:
            st.download_button(
                "⬇️ Full Report PDF",
                data=full_pack_pdf,
                file_name=f"fpi_full_report_{safe_week_main}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="download_fpi_full_v83",
            )



with tab_intro:
    render_system_intro_page()

with tab1:
    st.subheader("Csapat áttekintő")
    st.caption("A főoldal célja: 30 másodperc alatt látni, mi fontos a héten.")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        metric_card("Játékosok", filtered["player_name"].nunique())
    with col2:
        metric_card("Sessionök száma", len(filtered))
    with col3:
        td = filtered["total_distance"].sum() if "total_distance" in filtered.columns else np.nan
        metric_card("Össztáv", f"{td:,.0f} m" if pd.notna(td) else "—")
    with col4:
        load = filtered["training_load"].sum() if "training_load" in filtered.columns else np.nan
        metric_card("Terhelési pont", f"{load:,.0f}" if pd.notna(load) else "—")

    st.markdown("### Heti vezetői összefoglaló")
    render_weekly_summary_card(weekly_summary_text)
    r1, r2, r3 = st.columns(3)
    with r1:
        st.metric("Meccskészültség", f"{readiness_score}/100", score_to_label(readiness_score))
    with r2:
        st.metric("Periodizáció", periodization_type)
    with r3:
        mem_weeks = analysis_base_df["week"].nunique() if "week" in analysis_base_df.columns else 0
        st.metric("Elemzett hetek", mem_weeks)
    st.markdown("### Top 3 adaptív edzői teendő")
    render_coaching_priorities(coaching_priorities)

    st.markdown("### Múlt hét alapján: javaslat erre a hétre")
    render_weekly_summary_card(past_week_review_text)

    st.markdown("### Aktuális hét: maradék napok")
    render_weekly_summary_card(current_remaining_text)
    if current_remaining_plan_df is not None and not current_remaining_plan_df.empty:
        st.dataframe(current_remaining_plan_df, use_container_width=True, hide_index=True)

    st.markdown("### Jövő hét / következő mikrociklus")
    render_weekly_summary_card(next_week_plan_text)
    if next_week_plan_df is not None and not next_week_plan_df.empty:
        st.dataframe(next_week_plan_df, use_container_width=True, hide_index=True)

    weekly = aggregate_weekly(df[df["session_type"].isin(selected_types)])
    st.markdown("### Heti trendek")
    trend_options = available_metric_options(
        weekly,
        ["training_load", "total_distance", "sprint_distance", "hsr_distance", "distance_per_min", "max_speed", "dec_count"],
    )
    if trend_options:
        chart_metric = st.selectbox("Trendmutató", trend_options, format_func=metric_name, index=0)
        fig = px.line(
            weekly,
            x="week",
            y=chart_metric,
            color="session_type",
            markers=True,
            title=f"Heti trend: {metric_name(chart_metric)}",
        )
        fig.update_layout(xaxis_title="Hét", yaxis_title=metric_name(chart_metric), legend_title="Típus", template="plotly_white")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("Nincs elérhető trendmutató.")

    st.markdown("### Edzés vs meccs profil az aktuális héten")
    profile_cols = available_metric_options(
        filtered,
        ["total_distance", "distance_per_min", "sprint_distance", "hsr_distance", "max_speed", "dec_count"],
    )
    if profile_cols:
        prof = filtered.groupby("session_type", as_index=False)[profile_cols].mean(numeric_only=True)
        prof = prof.rename(columns={c: metric_name(c) for c in profile_cols})
        prof = prof.rename(columns={"session_type": "Típus"})
        st.dataframe(prof, use_container_width=True, hide_index=True)
    else:
        st.info("Nincs elérhető edzés-meccs összehasonlító mutató.")




with tab_premium:
    st.subheader("🎛️ Cockpit")
    st.caption("Cool vezetői nézet: readiness, risk, load, sprint, periodizáció és top teendők egy képernyőn.")
    render_hero(selected_week, selected_playstyle, readiness_score, periodization_type)
    k1, k2, k3, k4 = st.columns(4)
    with k1:
        render_premium_kpi("Meccskészültség", f"{readiness_score}/100", score_to_label(readiness_score), score_to_color(readiness_score))
    with k2:
        render_premium_kpi("Magas kockázatos játékos", str(high_risk_count), "Játékos kockázati motor alapján", "#ef4444" if high_risk_count else "#22c55e")
    with k3:
        mem_weeks = analysis_base_df["week"].nunique() if "week" in analysis_base_df.columns else 0
        render_premium_kpi("Elemzett hetek", str(mem_weeks), "Memory + aktuális feltöltés", "#38bdf8")
    with k4:
        render_premium_kpi("Periodizáció", periodization_type, "Automatikus heti besorolás", "#a78bfa")
    st.markdown("### Meccskészültség komponensek")
    if readiness_components:
        comp_df = pd.DataFrame([{"Komponens": k, "Pont": round(v, 1)} for k, v in readiness_components.items()])
        fig = px.bar(comp_df, x="Komponens", y="Pont", range_y=[0, 100], title="Meccskészültség komponensek")
        fig.update_layout(xaxis_title="", yaxis_title="Pont", template="plotly_white")
        st.plotly_chart(fig, use_container_width=True)
    st.markdown("### Top 3 edzői teendő")
    render_coaching_priorities(coaching_priorities)
    st.markdown("### Top játékos kockázat")
    render_risk_cards(player_risk_df, limit=5)
    if weekly_fingerprints is not None and not weekly_fingerprints.empty:
        st.markdown("### Multi-week performance fingerprint")
        trend_metric_options = [c for c in ["training_load", "sprint_distance", "distance_per_min", "max_speed", "dec_count"] if c in weekly_fingerprints.columns]
        if trend_metric_options:
            tm = st.selectbox("Vezetői trendmutató", trend_metric_options, format_func=metric_name)
            fig = px.line(weekly_fingerprints, x="week", y=tm, markers=True, title=f"Performance memória trend: {metric_name(tm)}")
            fig.update_layout(xaxis_title="Hét", yaxis_title=metric_name(tm), template="plotly_white")
            st.plotly_chart(fig, use_container_width=True)


with tab_export:
    st.subheader("📄 Export")
    st.caption("Egy helyen minden vezetői információ és export: összefoglaló, readiness, teendők, player risk, insightok.")

    st.markdown(
        """
        <div class="export-panel">
            <h3 style="margin-top:0;">Vezetői csomag</h3>
            <p style="color:#e0f2fe;">
                Ezt érdemes megmutatni vagy elküldeni a vezetőedzőnek: rövid összefoglaló,
                top teendők, readiness, kockázati lista és insightok.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    executive_df = build_executive_summary_df(
        selected_week,
        selected_playstyle,
        readiness_score,
        periodization_type,
        weekly_summary_text,
        high_risk_count if "high_risk_count" in globals() else 0,
        medium_risk_count if "medium_risk_count" in globals() else 0,
    )
    insight_export_df_export = build_insight_export_df(all_insights)
    priorities_df = pd.DataFrame(coaching_priorities)
    risk_export_df = player_risk_df if "player_risk_df" in globals() else pd.DataFrame()

    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.metric("Meccskészültség", f"{readiness_score}/100", score_to_label(readiness_score))
    with k2:
        st.metric("Periodizáció", periodization_type)
    with k3:
        st.metric("Magas kockázat", high_risk_count if "high_risk_count" in globals() else 0)
    with k4:
        st.metric("Insightok", len(all_insights))

    st.markdown("### Heti vezetői összefoglaló")
    render_wrapped_table(executive_df)

    st.markdown("### Mit jelentenek a fő fogalmak?")
    render_wrapped_table(build_plain_language_explanation())

    st.markdown("### Top 3 edzői teendő")
    if not priorities_df.empty:
        render_wrapped_table(priorities_df)
    else:
        st.info("Nincs kiemelt teendő.")

    st.markdown("### Játékos risk")
    if risk_export_df is not None and not risk_export_df.empty:
        render_wrapped_table(risk_export_df.head(10))
    else:
        st.info("Nincs player kockázati adat.")

    st.markdown("### Insightok")
    render_wrapped_table(insight_export_df_export)

    st.markdown("### Export")
    st.caption("A régi vezetői PDF/Word/Excel/CSV export gombok kikerültek. A használható termékriportok lent érhetők el.")
    safe_week = _safe_filename_week(selected_week)

    # -------------------------------------------------------------------------
    # V6.1 - Product Report Pack visszaemelve az Export fülre is
    # -------------------------------------------------------------------------
    st.markdown("### Football Performance Intelligence riportok – minta PDF-ek")
    st.caption("Két riport: Executive Summary és Full Report. A minta ugyanazt a struktúrát használja, mint az éles export.")
    if 'build_fpi_sample_pdf_bytes' in globals():
        try:
            sample_exec_pdf_bytes_export = build_fpi_sample_pdf_bytes("executive")
            sample_full_pdf_bytes_export = build_fpi_sample_pdf_bytes("full")
            sample_gps_only_pdf_bytes_export = build_fpi_gps_only_sample_pdf_bytes()
        except Exception:
            sample_exec_pdf_bytes_export = sample_full_pdf_bytes_export = sample_gps_only_pdf_bytes_export = None
        sample_method_pdf_bytes_export = build_fpi_methodology_pdf_bytes_v143() if "build_fpi_methodology_pdf_bytes_v143" in globals() else None
        sm1, sm2, sm3, sm4 = st.columns(4)
        with sm1:
            if sample_exec_pdf_bytes_export is not None:
                st.download_button("⬇️ Minta Executive Summary", data=sample_exec_pdf_bytes_export, file_name="fpi_minta_executive_summary.pdf", mime="application/pdf", use_container_width=True, key="download_sample_exec_v143_export")
        with sm2:
            if sample_full_pdf_bytes_export is not None:
                st.download_button("⬇️ Minta Full Report", data=sample_full_pdf_bytes_export, file_name="fpi_minta_full_report.pdf", mime="application/pdf", use_container_width=True, key="download_sample_full_v143_export")
        with sm3:
            if sample_gps_only_pdf_bytes_export is not None:
                st.download_button("⬇️ Minta GPS-only Report", data=sample_gps_only_pdf_bytes_export, file_name="fpi_minta_gps_only_report.pdf", mime="application/pdf", use_container_width=True, key="download_sample_gps_only_v143_export")
        with sm4:
            if sample_method_pdf_bytes_export is not None:
                st.download_button("⬇️ Metodika PDF", data=sample_method_pdf_bytes_export, file_name="fpi_metodika.pdf", mime="application/pdf", use_container_width=True, key="download_sample_method_v143_export")
        if sample_exec_pdf_bytes_export is None and sample_full_pdf_bytes_export is None and sample_gps_only_pdf_bytes_export is None:
            st.info("A minta PDF exporthoz a reportlab csomag szükséges.")

    st.markdown("### Football Performance Intelligence riportok – éles PDF-ek")
    st.caption("Az aktuális feltöltött adatokból: Executive Summary + Full Report.")
    if 'build_fpi_product_pdf_bytes' in globals():
        live_report_base_export = analysis_base_df.copy() if 'analysis_base_df' in globals() else df.copy()
        er1, er2, er3 = st.columns(3)
        with er1:
            gps_only_pack_pdf_export = build_fpi_gps_only_pdf_bytes(live_report_base_export, selected_week, selected_playstyle)
            if gps_only_pack_pdf_export is not None:
                st.download_button("⬇️ GPS-only PDF Report", data=gps_only_pack_pdf_export, file_name=f"fpi_gps_only_report_{safe_week}.pdf", mime="application/pdf", use_container_width=True, key="download_fpi_gps_only_v96_export")
        with er2:
            exec_pack_pdf_export = build_fpi_product_pdf_bytes(
                live_report_base_export,
                selected_week,
                selected_playstyle,
                report_type="executive",
                tactical_context=_fpi_context_for_export_v87(tactical_gps_context if "tactical_gps_context" in globals() else {}),
            )
            if exec_pack_pdf_export is not None:
                st.download_button("⬇️ Executive Summary PDF", data=exec_pack_pdf_export, file_name=f"fpi_executive_summary_{safe_week}.pdf", mime="application/pdf", use_container_width=True, key="download_fpi_exec_v83_export")
        with er3:
            full_pack_pdf_export = build_fpi_product_pdf_bytes(
                live_report_base_export,
                selected_week,
                selected_playstyle,
                report_type="full",
                tactical_context=_fpi_context_for_export_v87(tactical_gps_context if "tactical_gps_context" in globals() else {}),
            )
            if full_pack_pdf_export is not None:
                st.download_button("⬇️ Full Report PDF", data=full_pack_pdf_export, file_name=f"fpi_full_report_{safe_week}.pdf", mime="application/pdf", use_container_width=True, key="download_fpi_full_v83_export")
    else:
        st.warning("A Product Report Pack függvényei nem érhetők el ebben a fájlban.")


with tab_intel:
    st.subheader("🧠 Intelligence")
    st.caption("V2.5: readiness score, periodizáció, performance memória, multi-week mintázatok és adaptív ajánlások.")

    c1, c2 = st.columns([1, 1])
    with c1:
        render_score_card(
            "Meccskészültség",
            readiness_score,
            score_to_label(readiness_score),
            readiness_reasons,
        )
    with c2:
        st.markdown("### Meccskészültség komponensek")
        if readiness_components:
            comp_df = pd.DataFrame([
                {"Komponens": k, "Pont": round(v, 1)}
                for k, v in readiness_components.items()
            ])
            fig = px.bar(comp_df, x="Komponens", y="Pont", range_y=[0, 100], title="Meccskészültség komponensek")
            fig.update_layout(xaxis_title="", yaxis_title="Pont")
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Nincs elég adat a komponensekhez.")

    st.markdown("### Periodizációs besorolás")
    st.info(f"Az aktuális hét besorolása: **{periodization_type}**")

    st.markdown("### Heti lenyomat / performance memória")
    if weekly_fingerprints.empty:
        st.info("Nincs elég heti adat a fingerprint táblához.")
    else:
        show_cols = [
            "week", "periodizacios_tipus", "players", "training_load", "total_distance",
            "sprint_distance", "distance_per_min", "max_speed", "dec_count"
        ]
        show_cols = [c for c in show_cols if c in weekly_fingerprints.columns]
        fp_show = weekly_fingerprints[show_cols].copy()
        fp_show = fp_show.rename(columns={
            "week": "Hét",
            "periodizacios_tipus": "Periodizáció",
            "players": "Játékosok",
            "training_load": "Terhelési pont",
            "total_distance": "Össztáv",
            "sprint_distance": "Sprinttáv",
            "distance_per_min": "Táv/perc",
            "max_speed": "Max sebesség",
            "dec_count": "Lassítások",
        })
        st.dataframe(fp_show, use_container_width=True, hide_index=True)

        trend_metric_options = [c for c in ["training_load", "sprint_distance", "distance_per_min", "max_speed", "dec_count"] if c in weekly_fingerprints.columns]
        if trend_metric_options:
            tm = st.selectbox("Memória trendmutató", trend_metric_options, format_func=metric_name)
            fig = px.line(weekly_fingerprints, x="week", y=tm, markers=True, title=f"Többhetes trend: {metric_name(tm)}")
            fig.update_layout(xaxis_title="Hét", yaxis_title=metric_name(tm))
            st.plotly_chart(fig, use_container_width=True)

    st.markdown("### Többhetes mintázatok")
    if pattern_insights:
        render_insight_cards(pattern_insights)
    else:
        st.success("Nem látható kiemelt többhetes negatív mintázat az aktuális adatok alapján.")

    st.markdown("### Adaptív ajánlórendszer")
    render_coaching_priorities(coaching_priorities)


with tab_micro:
    st.subheader("📅 Mikrociklus")
    st.caption("A rendszer a meccsnaphoz viszonyítva értelmezi a heti struktúrát: MD-4, MD-3, MD-2, MD-1, MD, MD+1.")

    micro_df = build_microcycle_table(analysis_base_df, selected_week)
    if micro_df.empty:
        st.info("Nincs elérhető mikrociklus adat az aktuális szűrésre.")
    else:
        match_day = detect_match_day(analysis_base_df[analysis_base_df["week"] == selected_week])
        if match_day is not None:
            st.markdown(
                f"<span class='micro-pill'>Meccsnap: {match_day.strftime('%Y-%m-%d')}</span>"
                f"<span class='micro-pill'>Játékmodell: {html.escape(selected_playstyle)}</span>",
                unsafe_allow_html=True,
            )
        else:
            st.warning("Az aktuális héten nincs automatikusan azonosítható meccsnap.")

        show_cols = [
            "session_date_dt", "md_label", "session_type", "load_index", "total_distance",
            "distance_per_min", "sprint_distance", "hsr_distance", "max_speed", "dec_count", "játékosok"
        ]
        show_cols = [c for c in show_cols if c in micro_df.columns]
        show_micro = micro_df[show_cols].copy()
        if "session_date_dt" in show_micro.columns:
            show_micro["session_date_dt"] = show_micro["session_date_dt"].dt.strftime("%Y-%m-%d")
        show_micro = show_micro.rename(columns={
            "session_date_dt": "Dátum",
            "md_label": "MD-nap",
            "session_type": "Típus",
            "load_index": "Load index",
            "total_distance": "Össztáv",
            "distance_per_min": "Táv/perc",
            "sprint_distance": "Sprinttáv",
            "hsr_distance": "Nagy sebességű táv",
            "max_speed": "Max sebesség",
            "dec_count": "Lassítások",
            "játékosok": "Játékosok",
        })
        st.dataframe(show_micro, use_container_width=True, hide_index=True)

        chart_cols = available_metric_options(micro_df, ["load_index", "sprint_distance", "distance_per_min", "max_speed", "dec_count"])
        if chart_cols:
            metric = st.selectbox("Mikrociklus grafikon mutató", chart_cols, format_func=lambda x: "Load index" if x == "load_index" else metric_name(x))
            fig = px.bar(
                micro_df,
                x="md_label",
                y=metric,
                color="session_type",
                hover_data=["session_date_dt"],
                title=f"Mikrociklus profil: {'Load index' if metric == 'load_index' else metric_name(metric)}",
            )
            fig.update_layout(xaxis_title="Meccsnaphoz viszonyított nap", yaxis_title="Load index" if metric == "load_index" else metric_name(metric), legend_title="Típus")
            st.plotly_chart(fig, use_container_width=True)

        st.markdown("### Mikrociklus megállapítások")
        if micro_insights:
            render_insight_cards(micro_insights)
        else:
            st.success("A mikrociklus struktúrában nem látszik kiemelt figyelmeztetés az aktuális szabályok alapján.")



    st.markdown("## 🧭 Múlt hét -> aktuális hét -> jövő hét")
    st.caption("A rendszer külön kezeli: múlt hét teljes elemzése, aktuális hét hátralévő napjai, jövő heti mikrociklus.")

    st.markdown("### 1) Múlt hét teljes elemzése és javaslat erre a hétre")
    render_weekly_summary_card(past_week_review_text)
    if past_week_review_df is not None and not past_week_review_df.empty:
        st.dataframe(past_week_review_df, use_container_width=True, hide_index=True)

    st.markdown("### 2) Aktuális hét eddig feltöltött napjai és hátralévő napok")
    render_weekly_summary_card(current_remaining_text)
    if current_remaining_plan_df is not None and not current_remaining_plan_df.empty:
        st.dataframe(current_remaining_plan_df, use_container_width=True, hide_index=True)

    st.markdown("### 3) Jövő hét / következő mikrociklus az eddigiek alapján")
    render_weekly_summary_card(next_week_plan_text)
    if next_week_plan_df is not None and not next_week_plan_df.empty:
        st.dataframe(next_week_plan_df, use_container_width=True, hide_index=True)

    st.markdown("### 4) Játékosszintű következő teendők")
    if player_next_actions_df is not None and not player_next_actions_df.empty:
        st.dataframe(player_next_actions_df, use_container_width=True, hide_index=True)


with tab_risk:
    st.subheader("🚨 Kockázat")
    st.caption("Játékosszintű, többhetes eltérésalapú risk scoring: load spike, sprintprofil, lassítások, max sebesség.")
    if player_risk_df.empty:
        st.info("Nincs elég adat a játékosszintű risk engine-hez. Legalább több hét játékosszintű adat kell hozzá.")
    else:
        render_risk_cards(player_risk_df, limit=8)
        st.markdown("### Kockázati tábla")
        risk_show_df = player_risk_df.copy()
        if "Kockázati pontszám" in risk_show_df.columns:
            risk_show_df["Értelmezés"] = pd.to_numeric(risk_show_df["Kockázati pontszám"], errors="coerce").apply(
                lambda x: "Magas: azonnali edzői kontroll" if pd.notna(x) and x >= 70 else ("Közepes: figyelendő" if pd.notna(x) and x >= 40 else "Alacsony: rendben")
            )
        risk_cols_v116 = [c for c in ["Játékos", "Szerep", "Típus", "Játékperc", "Kockázati pontszám", "Kockázati szint", "Értelmezés", "Fő okok"] if c in risk_show_df.columns]
        st.dataframe(risk_show_df[risk_cols_v116] if risk_cols_v116 else risk_show_df, use_container_width=True, hide_index=True)
        fig = px.bar(player_risk_df.head(20), x="Játékos", y="Kockázati pontszám", title="Játékos risk score")
        fig.update_layout(xaxis_title="Játékos", yaxis_title="Kockázati pontszám (0–100)", xaxis_tickangle=-45, template="plotly_white")
        st.plotly_chart(fig, use_container_width=True)
        st.download_button("⬇️ Játékos risk export CSV", data=player_risk_df.to_csv(index=False).encode("utf-8-sig"), file_name=f"player_risk_{_safe_filename_week(selected_week)}.csv", mime="text/csv", use_container_width=True,
            key="download_button_unique_10",
        )

with tab2:
    st.subheader("🎯 Javaslatok")
    st.caption("Szabályalapú performance motor: AI nélkül is ad szakmai következtetést és javaslatot.")

    insights = all_insights

    st.markdown("### Heti összefoglaló")
    render_weekly_summary_card(weekly_summary_text)

    st.markdown("### Top 3 adaptív edzői teendő")
    render_coaching_priorities(coaching_priorities)

    st.markdown("### Coach-friendly insight kártyák")
    render_insight_cards(insights)

    st.markdown("### Exportálható insight tábla")
    insight_export_df = build_insight_export_df(insights)
    render_wrapped_table(insight_export_df)

    with st.expander("Táblázat szerkesztés nélküli nézetben"):
        st.data_editor(
            insight_export_df,
            use_container_width=True,
            hide_index=True,
            disabled=True,
            column_config={
                "Mit látunk?": st.column_config.TextColumn("Mit látunk?", width="large"),
                "Miért fontos?": st.column_config.TextColumn("Miért fontos?", width="large"),
                "Javaslat": st.column_config.TextColumn("Javaslat", width="large"),
                "Megállapítás": st.column_config.TextColumn("Megállapítás", width="medium"),
            },
        )

    st.markdown("#### Riport exportálása")
    safe_week = _safe_filename_week(selected_week)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.download_button(
            "⬇️ Excel riport",
            data=insights_to_excel_bytes(insight_export_df, selected_week),
            file_name=f"performance_riport_{safe_week}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        
            key="download_button_unique_11",
        )
    with c2:
        csv_bytes = insight_export_df.to_csv(index=False).encode("utf-8-sig")
        st.download_button(
            "⬇️ CSV",
            data=csv_bytes,
            file_name=f"performance_riport_{safe_week}.csv",
            mime="text/csv",
            use_container_width=True,
        
            key="download_button_unique_12",
        )
    with c3:
        word_bytes = insights_to_word_bytes(insight_export_df, selected_week)
        if word_bytes is not None:
            st.download_button(
                "⬇️ Word riport",
                data=word_bytes,
                file_name=f"performance_riport_{safe_week}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            
            key="download_button_unique_13",
        )
        else:
            st.info("Word exporthoz add hozzá a requirements.txt fájlhoz: python-docx")
    with c4:
        pdf_bytes = insights_to_pdf_bytes(insight_export_df, selected_week)
        if pdf_bytes is not None:
            st.download_button(
                "⬇️ PDF riport",
                data=pdf_bytes,
                file_name=f"performance_riport_{safe_week}.pdf",
                mime="application/pdf",
                use_container_width=True,
            
            key="download_button_unique_14",
        )
        else:
            st.info("PDF exporthoz add hozzá a requirements.txt fájlhoz: reportlab")
    st.markdown("#### Vezetői V3 PDF export")
    premium_pdf = build_premium_pdf_bytes(insight_export_df, selected_week, readiness_score, periodization_type, weekly_summary_text, coaching_priorities, player_risk_df, selected_playstyle)
    if premium_pdf is not None:
        st.download_button("⬇️ Vezetői V3 PDF riport", data=premium_pdf, file_name=f"premium_performance_riport_{safe_week}.pdf", mime="application/pdf", use_container_width=True,
            key="download_button_unique_15",
        )

with tab3:
    st.subheader("👤 Játékosok")
    pw = player_weekly(df)
    pw_current = pw[(pw["week"] == selected_week) & (pw["player_name"].isin(selected_players))]
    if not pw_current.empty:
        rank_options = available_metric_options(
            pw_current,
            ["training_load", "total_distance", "sprint_distance", "hsr_distance", "max_speed", "dec_count", "high_efforts"],
        )
        if rank_options:
            metric = st.selectbox("Játékos rangsor mutató", rank_options, format_func=metric_name, index=0)
            rank = pw_current.groupby("player_name", as_index=False)[metric].sum().sort_values(metric, ascending=False)
            fig = px.bar(rank.head(25), x="player_name", y=metric, title=f"Játékosrangsor: {metric_name(metric)}")
            fig.update_layout(xaxis_title="Játékos", yaxis_title=metric_name(metric), xaxis_tickangle=-45, template="plotly_white")
            st.plotly_chart(fig, use_container_width=True)

        st.markdown("### Játékos heti összesítő tábla")
        show = pw_current.rename(columns={"player_name": "Játékos", "week": "Hét", "session_type": "Típus"})
        show = show.rename(columns={c: metric_name(c) for c in pw_current.columns})
        st.dataframe(show, use_container_width=True, hide_index=True)
    else:
        st.info("Nincs adat az aktuális szűrésre.")

with tab4:
    st.subheader("✅ Adatminőség")
    st.caption("Ez azért fontos, mert a sportadat a valóságban mindig kicsit koszos.")

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        metric_card("Nyers sor", len(raw_df))
    with c2:
        metric_card("Standardizált sor", len(df))
    with c3:
        metric_card("Hetek", df["week"].nunique())
    with c4:
        metric_card("Típusok", ", ".join(session_types))

    st.markdown("### Performance memória állapot")
    mem_df = load_memory_df()
    if mem_df.empty:
        st.info("Nincs még mentett lokális memory adat.")
    else:
        st.success(f"Mentett memory: {len(mem_df)} sor, {mem_df['week'].nunique() if 'week' in mem_df.columns else 'ismeretlen'} hét.")
    st.markdown("### Automatikus oszlopmapping")
    map_df = pd.DataFrame([
        {"standard mező": k, "forrás oszlop": v if v is not None else "NINCS"}
        for k, v in mapping.items()
    ])
    st.dataframe(map_df, use_container_width=True, hide_index=True)

    st.markdown("### Hiányzó értékek a standard mezőkben")
    na = df.isna().mean().sort_values(ascending=False).reset_index()
    na.columns = ["oszlop", "hiányzó arány"]
    st.dataframe(na.head(30), use_container_width=True, hide_index=True)

with tab5:
    st.subheader("Nyers / standardizált adatok")
    st.markdown("### Standardizált adat")
    st.dataframe(df.head(500), use_container_width=True)
    st.markdown("### Nyers adat")
    st.dataframe(raw_df.head(200), use_container_width=True)


st.divider()
st.caption("V5.7 GENERAL EXCEL IMPORT – Data lap + általános fejlécfelismerés + védett dátumkezelés + ISO hetek.")

# -----------------------------------------------------------------------------
# V4.4 Smart Excel Mapper + License UI
# -----------------------------------------------------------------------------
with st.expander("🧩 Smart Excel Mapper + License / oszlopmapping ellenőrzése", expanded=False):
    st.write("Ha más klub más szerkezetű Excel-exportot tölt fel, itt ellenőrizhető és javítható, hogy melyik forrásoszlop melyik standard mezőre menjen.")
    raw_df_for_mapper = locals().get("raw_df", st.session_state.get("last_raw_df", pd.DataFrame()))
    if isinstance(raw_df_for_mapper, pd.DataFrame) and not raw_df_for_mapper.empty:
        raw_df = raw_df_for_mapper
        current_mapping = st.session_state.get("manual_mapping", None)
        if current_mapping is None:
            current_mapping = suggest_mapping(raw_df)
        render_mapping_score(current_mapping)
        st.session_state["manual_mapping"] = current_mapping

        profile_upload = st.file_uploader("Korábban mentett mapping profil betöltése (.json)", type=["json"], key="mapping_profile_upload")
        if profile_upload is not None:
            loaded_mapping = load_mapping_profile(profile_upload)
            if loaded_mapping:
                st.session_state["manual_mapping"] = loaded_mapping
                current_mapping = loaded_mapping
                st.success("Mapping profil betöltve.")

        st.dataframe(enhanced_mapping_quality_df(raw_df, current_mapping), use_container_width=True)

        if st.button("♻️ Mapping override törlése", use_container_width=True, key="clear_mapping_override"):
            st.session_state.pop("mapped_df_override", None)
            st.session_state.pop("manual_mapping", None)
            st.rerun()

        st.markdown("#### Kézi javítás")
        source_options = [""] + list(raw_df.columns)
        cols = st.columns(2)
        editable_mapping = dict(current_mapping)
        important_fields = CORE_REQUIRED + [
            "position", "duration", "total_distance", "distance_per_min", "max_speed",
            "sprints", "speed_zone_4", "speed_zone_5", "training_load", "acc_high", "dec_high"
        ]
        for idx, std_col in enumerate(important_fields):
            with cols[idx % 2]:
                current_val = editable_mapping.get(std_col) or ""
                default_idx = source_options.index(current_val) if current_val in source_options else 0
                editable_mapping[std_col] = st.selectbox(
                    f"{std_col} → forrásoszlop",
                    source_options,
                    index=default_idx,
                    key=f"mapper_select_{std_col}",
                ) or None

        c_apply, c_export = st.columns(2)
        with c_apply:
            if st.button("✅ Mapping alkalmazása erre a fájlra", use_container_width=True, key="apply_manual_mapping"):
                mapped_df, applied_mapping, missing = apply_mapping_to_raw(raw_df, editable_mapping)
                if missing:
                    st.error(f"Hiányzó kötelező mezők: {', '.join(missing)}")
                else:
                    st.session_state["manual_mapping"] = applied_mapping
                    st.session_state["mapped_df_override"] = mapped_df
                    st.success(f"Mapping alkalmazva. Feldolgozott sorok: {len(mapped_df)}")
                    st.rerun()
        with c_export:
            st.download_button(
                "⬇️ Mapping profil mentése",
                data=export_mapping_profile(editable_mapping),
                file_name="gps_mapping_profile.json",
                mime="application/json",
                use_container_width=True,
                key="download_mapping_profile",
            )
    else:
        st.info("Mapping ellenőrzéshez előbb tölts fel egy Excel/CSV fájlt.")
